#!/usr/bin/env python3
"""Build the S1-CL3 AT1 ASSESSOR instrument (.docx) by populating the Kangan template.

Institutional compliance document (NOT a YAT-branded artefact): loads the official Kangan
'Project Assessment - Assessor' template and fills it in, preserving the Kangan structure and
styles exactly (Details, Teacher/Assessor instructions, Marking Guide, Instructions to Student,
Benchmark). Mirrors the approved CL1/CL2 AT1 assessor instruments.

AT1 = Team Setup (individual) — a single deliverable:
  Team Plan  (BSBXTW401 element 1 + the task-allocation slice of element 2)

AT1 is the team / project setup only — how the student will lead the improvement team. The
technical analysis and the improvement design are AT2; the implementation is AT3. This keeps the
evidence unambiguous: AT1 is 401-led, AT2 integrates both units, AT3 is 504-led.

Knowledge evidence: the Team Plan carries a written Knowledge Evidence appendix (T6) for the
BSBXTW401 element-1 knowledge. The worked model answer lives in the YAT-branded exemplar
(build_cl3_team_plan_exemplar); this document carries the task instructions and the marking guide
with bidirectional UoC traceability.

Usage:  python scripts/build_cl3_at1_assessor.py [output.docx]
Default: S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT1/AT1-Team-Setup-Assessor.docx
"""
import sys
from pathlib import Path

from docx import Document  # noqa: E402

TEMPLATE = "templates/Project Assessment - Assessor.docx"


# ---------- cell / table helpers (preserve template styles) ----------

def set_cell(cell, lines):
    """Replace a cell's content with one paragraph per line, keeping the cell's style."""
    if isinstance(lines, str):
        lines = [lines]
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    for r in list(first.runs):
        r._element.getparent().remove(r._element)
    first.add_run(lines[0])
    style = first.style
    for line in lines[1:]:
        np = cell.add_paragraph()
        np.style = style
        np.add_run(line)


def instr_row(table, label):
    """Return the col-1 cell of the instructions row whose col-0 starts with `label`."""
    for row in table.rows:
        if row.cells[0].text.strip().lower().startswith(label.lower()):
            return row.cells[1]
    raise KeyError(label)


def clear_table_rows(table, keep):
    """Remove all rows after the first `keep` rows."""
    for row in table.rows[keep:]:
        row._element.getparent().remove(row._element)


def add_section_row(table, text):
    row = table.add_row()
    set_cell(row.cells[0], text)
    for r in row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell(row.cells[1], "")
    return row


def add_criterion_row(table, text):
    row = table.add_row()
    set_cell(row.cells[0], text)
    set_cell(row.cells[1], "Yes        No")
    return row


# ---------- content ----------

DETAILS = {
    "qualification": "ICT50220 Diploma of Information Technology",
    "units": [
        "BSBXTW401 Lead and facilitate a team",
    ],
    "task_title": "AT1 – Team Setup",
    "task_number": "1 of 3",
}

OVERVIEW = (
    "Students are assessed individually on setting up the improvement team for the Ledgerline "
    "cloud-infrastructure improvement engagement. Working as the team lead of an MP Tech Solutions "
    "(MTS) improvement team, each student produces a Team Plan — how the team will work: its "
    "objectives and responsibilities, per-member performance expectations, accountability "
    "strategies, contingency/risk planning, and the role and task allocation. This is the "
    "team-setup phase; the team analyses and designs the improvement in AT2 and implements it in "
    "AT3. AT1 is completed individually."
)

TASKS = [
    "In this project the student takes on the role of the team lead of an MTS improvement team "
    "engaged by YAT College to improve the cloud infrastructure of its Ledgerline (Accounting) "
    "system, following YAT's India-campus partnership. AT1 is completed individually and has one "
    "deliverable:",
    "• Team Plan — the team lead's plan for how the improvement team will work: common objectives "
    "and responsibilities, per-member performance expectations and behaviours, accountability "
    "strategies, contingency/risk planning, and the allocation of the improvement work across the "
    "team. Produced in the YAT Team Plan template.",
    "The technical analysis and the improvement design are assessed in AT2; AT1 sets up the team "
    "and the project.",
]

RESOURCES = [
    "Teacher/assessor supplied resources / Access to:",
    "• The YAT scenario site / intranet — the Ledgerline Cloud Infrastructure Improvement project "
    "(Engagement Role Brief, Improvement Requirements, ICT Manager Consultation Notes), which "
    "frames the engagement the team plan is for.",
    "• The YAT Team Plan template (intranet Templates section).",
    "• The YAT Document Archive — examples of previous MTS team plans.",
]

CRITERIA_STATEMENT = (
    "To receive a Satisfactory outcome for this assessment task, the student must complete every "
    "criterion in the marking guide below to a satisfactory standard. Where a criterion is not yet "
    "satisfactory, the student is given feedback and a further attempt per the Second attempt "
    "provisions."
)

CONDITIONS = [
    "C1 — The YAT scenario site / intranet is accessible to the student throughout the assessment.",
    "C2 — The student has access to the YAT Team Plan template.",
    "C3 — AT1 is completed individually; each student authors their own Team Plan.",
    "C4 — The improvement team and its membership (which the Team Plan plans for) are confirmed "
    "with the assessor; teams are formed at the assessor's discretion.",
]

# Marking guide: the Team Plan criteria. UoC traceability is carried in the Benchmark.
TEAM_PLAN = [
    "T1 — Team objectives and outcomes: identifies the common objectives of the improvement team, "
    "the responsibilities, and the required outcome(s).",
    "T2 — Performance expectations: uses performance plans to establish expected outcomes, goals "
    "and behaviours for individual team members, in line with the team objective and relevant "
    "policies.",
    "T3 — Accountability strategies: selects appropriate strategies to ensure team members are "
    "accountable for their roles and responsibilities.",
    "T4 — Contingency / risk planning: plans for contingencies that could impact the team.",
    "T5 — Task allocation: allocates the improvement work (the four dimensions — security, "
    "reliability, scalability, cost) to team members with appropriate instruction, considering "
    "contingencies.",
    "T6 — Knowledge Evidence appendix: written contextual responses linking the plan to the "
    "organisational requirements, the legislative requirements relevant to the workplace, and the "
    "typical workplace contingencies relevant to this team.",
    "T7 — Document quality: the Team Plan uses the YAT Team Plan template, plain professional "
    "English, and is complete and internally consistent.",
]

# Benchmark: per-criterion UoC evidenced (bidirectional traceability). KE is allocated to the
# written KE appendix (T6); the body criteria carry the PC / PE / FS they demonstrate.
BENCHMARK = [
    ("Team Plan (BSBXTW401, element 1 + task allocation)", [
        ("T1", "[BSBXTW401 PC 1.1] identify common objectives of the workplace team, responsibilities "
               "and required outcome(s)."),
        ("T2", "[BSBXTW401 PC 1.2] use performance plans to establish expected outcomes, goals and "
               "behaviours for individual team members in accordance with the team objective and "
               "relevant policies."),
        ("T3", "[BSBXTW401 PC 1.3] select appropriate strategies to ensure team members are "
               "accountable for their roles and responsibilities."),
        ("T4", "[BSBXTW401 PC 1.4] plan for contingencies that could impact the team."),
        ("T5", "[BSBXTW401 PC 2.2] allocate tasks to team members and provide appropriate "
               "instruction; [PE 1] assign tasks with appropriate instruction, considering required "
               "contingencies."),
        ("T6", "[BSBXTW401 KE 1] organisational requirements relevant to workplace teams; [KE 2] "
               "legislative requirements relevant to the workplace; [KE 9] typical workplace "
               "contingencies — evidenced as written contextual responses against the student's own "
               "team plan."),
        ("T7", "[BSBXTW401 FS Get the work done] plans, organises and implements work activities in "
               "line with organisational policies and procedures."),
    ]),
]

STUDENT_INTRO = [
    ("The engagement and your role", "Heading 2"),
    ("YAT College is a Registered Training Organisation (RTO) based at 175 Cremorne Street, Cremorne "
     "VIC. Following a new campus partnership in India, YAT wants the cloud infrastructure of its "
     "Ledgerline (Accounting) system confirmed as stable, reliable and fit for purpose, and "
     "compliant with the Indian regulatory requirements that now apply — and improved where it "
     "falls short.", "Assessor text"),
    ("You are the team lead of an MP Tech Solutions (MTS) improvement team, reporting to Pat Lin "
     "(MTS Senior Consultant), who liaises with Sam Walker, the YAT ICT Manager. In AT1 you set the "
     "team up: you plan how your team will work so it can analyse and design the improvement in AT2 "
     "and implement it in AT3. AT1 is your own individual work, and it is about leading the team — "
     "the technical analysis and design come in AT2.", "Assessor text"),
    ("The engagement framing is provided on the intranet (the Ledgerline Cloud Infrastructure "
     "Improvement project: the Engagement Role Brief, Improvement Requirements, and Consultation "
     "Notes). You produce the team plan — it is not provided to you. (The provided Engagement Role "
     "Brief sets the consulting roles and reporting; your Team Plan is your own plan for your team.)",
     "Assessor text"),
]

STUDENT_TASK = [
    ("Your task — the Team Plan", "Heading 2"),
    ("Using the YAT Team Plan template (download from the intranet's Templates section), produce a "
     "plan for how you will lead the improvement team. Your plan must cover:", "Assessor text"),
    ("• The team's common objectives, responsibilities and required outcomes for the engagement.", "Assessor text"),
    ("• Per-member performance expectations — the expected outcomes, goals and behaviours for each "
     "team member, in line with the team objective and YAT's policies.", "Assessor text"),
    ("• Accountability strategies — how you will ensure each member is accountable for their role.", "Assessor text"),
    ("• Contingency / risk planning — the contingencies that could impact the team and how you will "
     "plan for them.", "Assessor text"),
    ("• Task allocation — how you allocate the improvement work across the team. Each member owns "
     "one improvement dimension (security, reliability, scalability or cost); allocate with clear "
     "instruction and account for contingencies.", "Assessor text"),
    ("Knowledge Evidence appendix (required). At the end of your Team Plan, add a Knowledge "
     "Evidence appendix and answer the following in your own words, about your own plan:", "Assessor text"),
    ("• Which organisational requirements (workplace policies, codes of conduct, organisational "
     "reputation and culture) shape how your team must operate, and how does your plan reflect "
     "them?", "Assessor text"),
    ("• Which legislative requirements are relevant to your team's work, and how does your plan "
     "account for them?", "Assessor text"),
    ("• What typical workplace contingencies could impact your team (for example unplanned leave, "
     "re-allocation of work, succession for a key role), and how does your plan handle them?", "Assessor text"),
    ("Resources — on the YAT intranet: the Ledgerline Cloud Infrastructure Improvement project "
     "(Engagement Role Brief, Improvement Requirements, Consultation Notes) and the YAT Team Plan "
     "template.", "Assessor text"),
    ("Submit the populated Team Plan (.docx) with the Knowledge Evidence appendix completed.", "Assessor text"),
]

TIPS = [
    ("Tips for success", "Heading 2"),
    ("AT1 is about leading the team, not the technology. Plan how your team will work — objectives, "
     "expectations, accountability, contingencies and who does what. The technical analysis and "
     "design are AT2; do not start them here.", "Assessor text"),
    ("Your Team Plan is yours; the Engagement Role Brief is provided. The provided Engagement Role "
     "Brief sets the MTS / YAT consulting roles and reporting. Your Team Plan is your own plan for "
     "your improvement team — don't just restate the role brief.", "Assessor text"),
    ("Allocate by dimension. Each member owns one of security, reliability, scalability or cost. "
     "Make the allocation clear, with the instruction each member needs, and allow for "
     "contingencies.", "Assessor text"),
    ("Write to your own plan. The Knowledge Evidence appendix asks about your own team and "
     "engagement — answer in your own words, not from a textbook.", "Assessor text"),
]


def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details ----
    t_details = doc.tables[0]
    set_cell(t_details.rows[1].cells[1], DETAILS["qualification"])
    set_cell(t_details.rows[2].cells[1], DETAILS["units"])
    set_cell(t_details.rows[3].cells[1], DETAILS["task_title"])
    set_cell(t_details.rows[4].cells[1], DETAILS["task_number"])

    # ---- Table 1: Teacher/Assessor instructions ----
    t_instr = doc.tables[1]
    set_cell(instr_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell(instr_row(t_instr, "Task"), TASKS)
    set_cell(instr_row(t_instr, "Resources required"), RESOURCES)
    set_cell(instr_row(t_instr, "Assessment criteria"), CRITERIA_STATEMENT)
    cond_row = t_instr.add_row()
    set_cell(cond_row.cells[0], "Assessment Conditions & Setup Requirements")
    for r in cond_row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell(cond_row.cells[1], CONDITIONS)

    # ---- Table 2: Marking Guide ----
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)  # keep header rows
    add_section_row(t_mark, "Team Plan")
    for c in TEAM_PLAN:
        add_criterion_row(t_mark, c)

    # ---- Instructions to Student ----
    doc.add_paragraph("Instructions to Student", style="Heading 1")
    for text, style in (STUDENT_INTRO + STUDENT_TASK + TIPS):
        doc.add_paragraph(text, style=style)

    # ---- Benchmark (UoC traceability) ----
    doc.add_paragraph("Marking Benchmark — UoC traceability", style="Heading 1")
    doc.add_paragraph(
        "For each criterion, the unit-of-competency items it evidences. Every PC, PE, KE and FS "
        "item allocated to AT1 in the cluster assessment plan is covered below.", style="Assessor text")
    for part_title, rows in BENCHMARK:
        doc.add_paragraph(part_title, style="Heading 2")
        for cid, uoc in rows:
            p = doc.add_paragraph(style="Assessor text")
            run = p.add_run(f"{cid}  —  ")
            run.bold = True
            p.add_run(uoc)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT1/AT1-Team-Setup-Assessor.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
