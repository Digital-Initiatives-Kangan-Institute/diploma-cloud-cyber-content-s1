#!/usr/bin/env python3
"""Build the S1-CL3 AT2 ASSESSOR instrument (.docx) by populating the Kangan template.

Institutional compliance document (NOT a YAT-branded artefact): loads the official Kangan
'Project Assessment - Assessor' template and fills it in, preserving the Kangan structure and
styles exactly. Mirrors the approved CL1/CL2/CL3-AT1 assessor instruments.

AT2 = Team Implementation (group, with individual evidence) — BSBXTW401 elements 1-4. The team
of four (one cloud component each — network / compute / database / storage) plans the
implementation and writes the infrastructure-as-code for the APPROVED improvement design,
integrating it into one deployable template across a series of led working meetings. Each student
plans, coordinates, supports and monitors the team — leading at least one meeting — and is assessed
INDIVIDUALLY on the BSBXTW401 leadership competencies.

The CloudFormation the team writes is the VEHICLE for the teamwork; it is NOT assessed against
ICTCLD504 (that is ICTCLD505, done in CL2). The individual technical deployment is AT3.

Usage:  python scripts/s1_cl3/build_s1_cl3_at2_assessor.py [output.docx]
Default: S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT2/AT2-Team-Implementation-Assessor.docx
"""
import sys
from pathlib import Path

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Assessor.docx")

sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_tables import add_criterion_row, add_section_row, clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402


# ---------- content ----------

DETAILS = {
    "qualification": "ICT50220 Diploma of Information Technology",
    "units": [
        "BSBXTW401 Lead and facilitate a team",
    ],
    "task_title": "AT2 – Team Implementation",
    "task_number": "2 of 3",
}

OVERVIEW = (
    "Students are assessed individually, within a team, on leading and facilitating the "
    "implementation of the approved Ledgerline cloud-infrastructure improvement. Working in a team "
    "of four — one cloud component each (network, compute, database, storage) — the team plans the "
    "implementation and writes the infrastructure-as-code for the approved design, integrating it "
    "into one deployable template across a series of led working meetings. Each student plans, "
    "coordinates, supports and monitors the team, leading at least one meeting, and is assessed "
    "individually on the BSBXTW401 leadership competencies. The CloudFormation the team writes is "
    "the vehicle for the teamwork; the technical deployment is assessed individually in AT3."
)

TASKS = [
    "In this project the students are members of an MP Tech Solutions (MTS) improvement team "
    "implementing the approved Ledgerline cloud-infrastructure improvement for YAT College. The "
    "team is four members, each owning one cloud component (network, compute, database, storage). "
    "AT2 has two parts and individual leadership evidence:",
    "• Part A — Plan: as a team, produce the team plan and the implementation plan — the team's "
    "objectives and accountabilities, per-member performance expectations, contingencies, and the "
    "allocation of one component's infrastructure-as-code write to each member.",
    "• Part B — Lead the work: write the CloudFormation for the approved design (each member their "
    "component), integrate it into one deployable template, and validate it for a team sign-off — "
    "run across a series of working meetings. Each member leads at least one meeting.",
    "• Individually, each member: leads at least one working meeting (assessor-observed), keeps a "
    "reflection on two challenges/conflicts, and produces a performance review of the team's work.",
    "The approved design and the existing-state baseline are provided; the team implements the "
    "agreed design. The deployment and its verification are AT3.",
]

RESOURCES = [
    "Teacher/assessor supplied resources / Access to:",
    "• The approved improvement design — the Accounting System Cloud Architecture Approved "
    "Improvement Design (intranet, Ledgerline Cloud Infrastructure Improvement project) — the design "
    "the team writes the infrastructure-as-code for.",
    "• The baseline lab-pack — the existing-state CloudFormation the team's upgrade targets.",
    "• The YAT Team Plan and Implementation Plan templates (intranet Templates section).",
    "• An AWS Academy lab and infrastructure-as-code tooling (editor, cfn-lint / change-set preview) "
    "to write and validate the templates.",
]

CRITERIA_STATEMENT = (
    "To receive a Satisfactory outcome for this assessment task, the student must complete every "
    "criterion in the marking guide below to a satisfactory standard. The criteria are assessed "
    "individually for each student, on the leadership evidence they personally produce. Where a "
    "criterion is not yet satisfactory, the student is given feedback and a further attempt per the "
    "Second attempt provisions."
)

CONDITIONS = [
    "C1 — Teams of four, one cloud component each (network / compute / database / storage); teams "
    "are formed at the assessor's discretion.",
    "C2 — The approved improvement design and the baseline lab-pack are provided to the team.",
    "C3 — AT2 runs as rotating-chair working meetings; each student chairs at least one, and the "
    "assessor observes each chair and completes the led-meeting observation checklist (below).",
    "C4 — Individual evidence: each student's led-meeting observation, reflection, and performance "
    "review are their own work.",
    "C5 — During the meetings the assessor stimulates at least one conflict or challenge for "
    "students to manage; how the assessor does so is the assessor's conduct and is out of scope of "
    "the student instrument.",
]

# Marking guide — Part A (Plan), Part B (Lead the work, observed), Part C (Monitor & reflect),
# plus the team deliverable + quality. UoC traceability is in the Benchmark.
PART_A = [
    "I1 — Team plan: identifies the team's common objectives, responsibilities and required "
    "outcomes; sets per-member performance expectations, goals and behaviours; selects "
    "accountability strategies; and plans for contingencies that could impact the team.",
    "I2 — Task allocation: allocates one component's infrastructure-as-code write to each member "
    "based on expertise / development potential, with appropriate instruction and allowing for "
    "contingencies.",
]

PART_B = [
    "I3 — Communicate and facilitate: communicates the team's common objectives and "
    "responsibilities, facilitates open and respectful collaboration between members (including "
    "those from diverse backgrounds), and identifies cross-collaboration opportunities.",
    "I4 — Support and resolve: coaches and supports members toward the team goals, facilitates the "
    "team to identify, brainstorm, report and resolve task-related issues, and uses problem-solving "
    "to deal with team, task or individual challenges — including managing a conflict.",
    "I5 — Lead a working meeting (observed): leads / facilitates at least one working meeting; the "
    "assessor records the led-meeting observation checklist for the student as chair. Primary "
    "performance evidence.",
]

PART_C = [
    "I6 — Monitor team performance: measures team-member performance against the agreed work plans, "
    "provides timely and constructive feedback, identifies specific learning and development "
    "opportunities, and implements action plans (the performance review).",
    "I7 — Reflection: reflects on two conflicts or challenges encountered — what happened, how it "
    "was handled, what was learned, and what to do differently next time.",
    "I8 — Knowledge Evidence appendix: written contextual responses on the organisational and "
    "legislative requirements relevant to the team, facilitation and coaching techniques, "
    "communication methods (including cross-cultural), professional behaviours to role-model, and "
    "typical workplace contingencies.",
]

DELIVERABLE = [
    "I9 — Integrated implementation: the team produces the infrastructure-as-code for the approved "
    "design (each member their component), integrated into one deployable template, validated and "
    "signed off as ready to deploy.",
    "I10 — Document quality: the team plan, the reflection and the performance review use plain "
    "professional English and are complete and internally consistent.",
]

# Led-meeting observation checklist — the behaviours the assessor records for each student as chair.
OBSERVATION_CHECKLIST = [
    "Communicated the team's common objectives and the meeting's purpose.",
    "Allocated / confirmed tasks and gave appropriate instruction.",
    "Facilitated open, respectful collaboration between members (including diverse perspectives).",
    "Coached / supported a member toward the team goals.",
    "Facilitated the team to identify and resolve a task-related issue.",
    "Managed a conflict or challenge constructively.",
    "Monitored progress against the work plan and gave constructive feedback.",
]

# Benchmark: per-criterion UoC evidenced (bidirectional traceability). Every PC/PE/KE/FS item
# allocated to AT2 in the cluster assessment plan (BSBXTW401 el 1-4) is covered below.
BENCHMARK = [
    ("Part A — Plan team outcomes and allocate (BSBXTW401 element 1 + 2.2)", [
        ("I1", "[BSBXTW401 PC 1.1] identify common objectives, responsibilities and required "
               "outcome(s); [PC 1.2] use performance plans to establish expected outcomes, goals and "
               "behaviours; [PC 1.3] select accountability strategies; [PC 1.4] plan for "
               "contingencies."),
        ("I2", "[BSBXTW401 PC 2.2] allocate tasks based on staff expertise or development potential "
               "and provide appropriate instructions; [PE 1] assign tasks with appropriate "
               "instruction, considering required contingencies."),
    ]),
    ("Part B — Coordinate and support the team (BSBXTW401 elements 2-3, observed)", [
        ("I3", "[BSBXTW401 PC 2.1] communicate common team objectives and responsibilities; [PC 2.3] "
               "facilitate open and respectful communication and collaboration, considering diverse "
               "backgrounds; [PC 2.4] identify cross-collaboration opportunities."),
        ("I4", "[BSBXTW401 PC 3.1] provide coaching to staff; [PC 3.2] support individuals toward "
               "common goals; [PC 3.3] facilitate the team to identify, brainstorm, report and "
               "resolve task issues; [PC 3.4] use problem solving for team/task/individual "
               "challenges; [PE 2] provide feedback and assistance; [PE 5] manage conflicts and "
               "challenges."),
        ("I5", "[BSBXTW401 PC 2.1] communicate objectives, [PC 2.3] facilitate collaboration, "
               "[PC 3.1] coach and support, [PC 3.3] resolve a task issue, [PE 5] manage a conflict — "
               "observed by the assessor for the student as chair in the led meeting (the primary "
               "performance evidence; recorded on the observation checklist); [BSBXTW401 FS Interact "
               "with others] appropriate communication, relationship-building and conflict management."),
    ]),
    ("Part C — Monitor performance and reflect (BSBXTW401 element 4 + knowledge)", [
        ("I6", "[BSBXTW401 PC 4.1] measure performance against agreed work plans; [PC 4.2] provide "
               "timely, constructive feedback; [PC 4.3] identify learning and development "
               "opportunities; [PC 4.4] implement action plans; [PE 3] collate feedback on individual "
               "and team performance; [PE 4] identify and implement development opportunities."),
        ("I7", "[BSBXTW401 KE 5] strategies for conflict resolution and negotiation; [KE 10] teamwork "
               "challenges (difficulties performing tasks, conflicts, risks, inappropriate "
               "behaviour) — reflected on against the student's own experience."),
        ("I8", "[BSBXTW401 KE 1] organisational requirements; [KE 2] legislative requirements; [KE 3] "
               "facilitation techniques; [KE 4] mentoring and coaching techniques; [KE 6] methods and "
               "styles of communication; [KE 7] cross-cultural communication and communication with "
               "individuals with special needs; [KE 8] professional behaviours to role-model; [KE 9] "
               "typical workplace contingencies."),
    ]),
    ("Team deliverable and quality (Foundation Skills)", [
        ("I9", "[BSBXTW401 FS Get the work done] plans, organises and implements work activities in "
               "line with organisational policies; [FS Navigate the world of work] understands and "
               "explains ethical, legal, regulatory and organisational responsibilities to the team. "
               "(The CloudFormation write is the team-work vehicle; it is not assessed against "
               "ICTCLD504.)"),
        ("I10", "[BSBXTW401 FS Interact with others] appropriate written communication in the team "
                "plan, reflection and performance review."),
    ]),
]

STUDENT_INTRO = [
    ("The engagement and your role", "Heading 2"),
    ("YAT College has approved the improvement design for the cloud infrastructure of its Ledgerline "
     "(Accounting) system. Your MP Tech Solutions (MTS) improvement team now implements it: you "
     "write the infrastructure-as-code for the approved design and integrate it into one deployable "
     "template, ready to be deployed.", "Assessor text"),
    ("Your team is four members, each owning one cloud component — network, compute, database, or "
     "storage. AT2 is teamwork: you plan how the team will work, you write your component, and you "
     "lead and facilitate the team through the implementation. You are assessed individually on how "
     "you lead — your planning, coordination, support, and monitoring of the team — not on the "
     "technology (that is AT3, where each of you deploys the result individually).", "Assessor text"),
    ("The approved design and the existing-state baseline are provided on the intranet. You implement "
     "the agreed design — you do not redesign it.", "Assessor text"),
]

STUDENT_TASK = [
    ("Part A — plan the implementation (team)", "Heading 2"),
    ("As a team, produce a team plan and an implementation plan covering:", "Assessor text"),
    ("• The team's common objectives, responsibilities and required outcomes; per-member performance "
     "expectations, goals and behaviours; accountability strategies; and contingency planning.", "Assessor text"),
    ("• The allocation of the infrastructure-as-code write — one cloud component to each member "
     "(network, compute, database, storage) — with clear instruction and allowance for "
     "contingencies; and the sequencing and dependencies for integrating the components into one "
     "deployable template.", "Assessor text"),
    ("Part B — lead the work (team, individually observed)", "Heading 2"),
    ("Write the CloudFormation for the approved design (each of you your own component) and integrate "
     "it into one deployable template, validated and signed off as ready to deploy. Do this across a "
     "series of working meetings, with the chair rotating so each of you leads at least one.", "Assessor text"),
    ("When you chair a meeting, you are leading and facilitating the team: communicate the "
     "objectives, allocate and instruct, facilitate collaboration, coach and support members, help "
     "resolve task issues, and manage any conflict. Your assessor observes you as chair.", "Assessor text"),
    ("Individually, you also produce:", "Assessor text"),
    ("• A reflection on two conflicts or challenges you encountered — what happened, how you handled "
     "it, what you learned, and what you would do differently.", "Assessor text"),
    ("• A performance review of the team's performance on the work — measuring against the plan, with "
     "feedback, development opportunities and action plans.", "Assessor text"),
    ("• A Knowledge Evidence appendix — written responses, in your own words, on the organisational "
     "and legislative requirements relevant to your team, the facilitation and coaching techniques "
     "you used, communication methods (including across cultures), the professional behaviours you "
     "role-modelled, and the typical contingencies your team had to allow for.", "Assessor text"),
    ("Submit your team's plan and integrated implementation, and your individual reflection, "
     "performance review and Knowledge Evidence appendix.", "Assessor text"),
]

TIPS = [
    ("Tips for success", "Heading 2"),
    ("AT2 is about leading the team, not the technology. You are marked on how you plan, coordinate, "
     "support and monitor the team — the CloudFormation is what the team works on together, but your "
     "deployment is assessed in AT3.", "Assessor text"),
    ("Lead at least one meeting well. Your led meeting is your primary evidence — communicate the "
     "objectives, allocate clearly, facilitate everyone, coach where needed, resolve an issue, and "
     "handle conflict constructively.", "Assessor text"),
    ("Implement the approved design — don't redesign it. The design is agreed; your job is to "
     "realise it as code, together, and integrate the components cleanly.", "Assessor text"),
    ("Write to your own experience. The reflection, performance review and Knowledge Evidence are "
     "about your own team and your own leadership — answer in your own words.", "Assessor text"),
]


def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details ----
    t_details = doc.tables[0]
    set_cell_content(t_details.rows[1].cells[1], DETAILS["qualification"])
    set_cell_content(t_details.rows[2].cells[1], DETAILS["units"])
    set_cell_content(t_details.rows[3].cells[1], DETAILS["task_title"])
    set_cell_content(t_details.rows[4].cells[1], DETAILS["task_number"])

    # ---- Table 1: Teacher/Assessor instructions ----
    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Resources required"), RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA_STATEMENT)
    cond_row = t_instr.add_row()
    set_cell_content(cond_row.cells[0], "Assessment Conditions & Setup Requirements")
    for r in cond_row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell_content(cond_row.cells[1], CONDITIONS)

    # ---- Table 2: Marking Guide ----
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)  # keep header rows
    add_section_row(t_mark, "Part A — Plan")
    for c in PART_A:
        add_criterion_row(t_mark, c)
    add_section_row(t_mark, "Part B — Lead the work (observed)")
    for c in PART_B:
        add_criterion_row(t_mark, c)
    add_section_row(t_mark, "Part C — Monitor and reflect")
    for c in PART_C:
        add_criterion_row(t_mark, c)
    add_section_row(t_mark, "Team deliverable and quality")
    for c in DELIVERABLE:
        add_criterion_row(t_mark, c)

    # ---- Instructions to Student ----
    doc.add_paragraph("Instructions to Student", style="Heading 1")
    for text, style in (STUDENT_INTRO + STUDENT_TASK + TIPS):
        doc.add_paragraph(text, style=style)

    # ---- Led-meeting observation checklist (assessor recording tool) ----
    doc.add_paragraph("Led-meeting observation checklist", style="Heading 1")
    doc.add_paragraph(
        "Complete one checklist per student, for the meeting they chair. Record evidence of each "
        "behaviour and a Satisfactory / Not yet satisfactory judgement.", style="Assessor text")
    for b in OBSERVATION_CHECKLIST:
        doc.add_paragraph("• " + b, style="Assessor text")

    # ---- Benchmark (UoC traceability) ----
    doc.add_paragraph("Marking Benchmark — UoC traceability", style="Heading 1")
    doc.add_paragraph(
        "For each criterion, the unit-of-competency items it evidences. Every PC, PE, KE and FS "
        "item allocated to AT2 in the cluster assessment plan is covered below.", style="Assessor text")
    for part_title, rows in BENCHMARK:
        doc.add_paragraph(part_title, style="Heading 2")
        for cid, uoc in rows:
            p = doc.add_paragraph(style="Assessor text")
            run = p.add_run(f"{cid}  —  ")
            run.bold = True
            p.add_run(uoc)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT2/AT2-Team-Implementation-Assessor.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
