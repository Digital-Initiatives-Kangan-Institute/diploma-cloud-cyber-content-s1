#!/usr/bin/env python3
"""Build the S1-CL3 AT1 ASSESSOR instrument (.docx) by populating the Kangan template.

Institutional compliance document (NOT a YAT-branded artefact): loads the official Kangan
'Project Assessment - Assessor' template and fills it in, preserving the Kangan structure and
styles exactly (Details, Teacher/Assessor instructions, Marking Guide, Instructions to Student,
Benchmark). Mirrors the approved CL1/CL2 AT1 assessor instruments.

AT1 = Design (individual) — ICTCLD504 elements 1 (analyse) + 2 (design), two parts:
  Part A  Solution Design  — analyse the single-AZ Ledgerline baseline and design the whole
          improvement (four components: network / compute / database / storage; across all four
          optimisation concerns: security, reliability, scalability, cost; incl. the India
          residency slice), documented and justified.
  Part B  Design presentation — present the proposed architecture for review and obtain sign-off
          to proceed.

The worked model answer lives in the YAT-branded exemplar (build_s1_cl3_at1_solution_design_exemplar).
This document carries the task instructions and the marking guide with bidirectional UoC traceability.
The CloudFormation write and the deployment are AT2 / AT3; AT1 is the design.

Usage:  python scripts/s1_cl3/build_s1_cl3_at1_assessor.py [output.docx]
Default: S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT1/AT1-Design-Assessor.docx
"""
import sys
from pathlib import Path

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Assessor.docx")


# ---------- cell / table helpers (preserve template styles) ----------

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_tables import add_criterion_row, add_section_row, clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402


# ---------- content ----------

DETAILS = {
    "qualification": "ICT50220 Diploma of Information Technology",
    "units": [
        "ICTCLD504 Improve cloud-based infrastructure",
    ],
    "task_title": "AT1 – Design",
    "task_number": "1 of 3",
}

OVERVIEW = (
    "Students are assessed individually on designing the cloud-infrastructure improvement for the "
    "YAT Ledgerline (Accounting) system. Working as an MP Tech Solutions (MTS) consultant, each "
    "student analyses the current single-Availability-Zone baseline and designs an improvement "
    "across all four optimisation concerns — security, reliability, scalability and cost — including "
    "the infrastructure changes needed to meet the applicable Indian regulatory requirements, then "
    "presents the proposed architecture for review and obtains sign-off to proceed. The deliverables "
    "are a Solution Design (Part A) and a design presentation (Part B). The improvement is to the "
    "cloud infrastructure only; the application and its data are out of scope. AT1 is completed "
    "individually."
)

TASKS = [
    "In this project the student takes on the role of an MTS consultant engaged by YAT College to "
    "improve the cloud infrastructure of its Ledgerline (Accounting) system, following YAT's "
    "India-campus partnership. AT1 is completed individually and has two parts:",
    "• Part A — Solution Design: analyse the current single-AZ baseline and design the improvement "
    "across the four components (network, compute, database, storage) and all four optimisation "
    "concerns (security, reliability, scalability, cost), including the infrastructure changes that "
    "meet the Indian regulatory requirements. Documented and justified in the YAT Solution Design "
    "template, with a Knowledge Evidence appendix.",
    "• Part B — Design presentation: present the proposed architecture to the required personnel for "
    "review, respond to questions, and obtain sign-off to proceed to deployment.",
    "The team-based implementation planning is AT2 and the deployment is AT3; AT1 is the individual "
    "design and its approval.",
]

RESOURCES = [
    "Teacher/assessor supplied resources / Access to:",
    "• The YAT scenario site / intranet — the Ledgerline Cloud Infrastructure Improvement project "
    "(Engagement Role Brief, Improvement Requirements, ICT Manager Consultation Notes, Indian "
    "Regulatory Requirements) and the current-state ICT records (the Accounting System Cloud "
    "Architecture Baseline Design, Infrastructure Specifications, Application Specification, "
    "Operational Costing) — the baseline the student analyses and improves.",
    "• The YAT Solution Design template (intranet Templates section).",
    "• A room and schedule for the Part B design presentation.",
]

CRITERIA_STATEMENT = (
    "To receive a Satisfactory outcome for this assessment task, the student must complete every "
    "criterion in the marking guide below to a satisfactory standard. Where a criterion is not yet "
    "satisfactory, the student is given feedback and a further attempt per the Second attempt "
    "provisions."
)

CONDITIONS = [
    "C1 — The YAT scenario site / intranet (the Ledgerline Cloud Infrastructure Improvement project "
    "and the current-state ICT records) is accessible to the student throughout the assessment.",
    "C2 — The student has access to the YAT Solution Design template.",
    "C3 — AT1 is completed individually; each student authors their own Solution Design and delivers "
    "their own presentation.",
    "C4 — The Part B presentation is scheduled and observed by the assessor, who records the "
    "presentation observation and the sign-off decision.",
]

# Marking guide — Part A (Solution Design) and Part B (presentation). UoC traceability is in the Benchmark.
PART_A = [
    "D1 — Architecture review: reviews and evaluates the current single-AZ Ledgerline infrastructure "
    "and identifies the business impact of its design (single points of failure, the idle-profile "
    "cost, scalability headroom, the compliance position).",
    "D2 — Options and confirmed decisions: identifies architectural options and patterns, assesses "
    "their benefits and differences against the current business needs, and confirms the design "
    "decisions — including recognising that a Multi-AZ database is not available for Ledgerline (the "
    "discovered Multi-AZ limitation) and ruling it out with a cost-versus-benefit justification.",
    "D3 — Goals and metrics: sets measurable security, reliability, scalability and cost goals and "
    "confirms the performance metrics the design is measured against.",
    "D4 — Improvement design: improves compute, storage, database and network across all four "
    "optimisation concerns into one integrated architecture (application-tier Multi-AZ plus database "
    "backup/restore and cross-Region DR for reliability — the database itself cannot be Multi-AZ, a "
    "Ledgerline constraint; elastic-on-demand scalability; in-place security hardening; proportionate "
    "cost optimisation).",
    "D5 — Regulatory compliance: assesses the infrastructure against the Indian Regulatory "
    "Requirements and includes the infrastructure changes that close the gaps (the light India "
    "residency slice — CERT-In logs and books-of-account retrievability).",
    "D6 — Documented and justified: the architecture is documented and justified in the Solution "
    "Design, each improvement carrying a proportionate cost-versus-benefit rationale (sound "
    "engineering, not gold-plating).",
    "D7 — Knowledge Evidence appendix: written contextual responses (object storage; avoiding single "
    "points of failure and testing for them; cloud features for each optimisation concern) about the "
    "student's own design.",
]

PART_B = [
    "ASSESSOR FOCUS — common error to probe. The database tier CANNOT be made Multi-AZ: Ledgerline is "
    "vendor-certified single-instance only (the Multi-AZ database limitation discovered during the cloud "
    "migration — see the current-state ICT records and the Cloud Migration Technical Finding). A candidate "
    "who proposes a Multi-AZ or mirrored database has missed this constraint — probe it directly in "
    "questioning: did they find the discovered-limitation record, and did they weigh the only real route to "
    "database failover (replacing the accounting product — new licence, full data migration, staff retraining "
    "and change management, and the delivery risk) against accepting backup/restore plus cross-Region DR? "
    "Proposing a Multi-AZ database without recognising the constraint and making that cost-versus-benefit case "
    "is inadequate research and analysis — reflect it in D1, D2 and D4. NOTE: application-tier Multi-AZ (an "
    "Auto Scaling group across AZs behind the internal ALB) is correct and expected — do not penalise it; "
    "only the database is constrained.",
    "D8 — Design presentation: presents the proposed architecture for review to the required "
    "personnel using appropriate industry language, and responds to questions to confirm the design.",
    "D9 — Sign-off: obtains sign-off to proceed to deployment from the required personnel.",
]

QUALITY = [
    "D10 — Document quality: the Solution Design uses the YAT Solution Design template, plain "
    "professional English, and is complete and internally consistent.",
]

# Benchmark: per-criterion UoC evidenced (bidirectional traceability). Every PC/PE/KE/FS item
# allocated to AT1 in the cluster assessment plan (ICTCLD504 el 1-2) is covered below.
BENCHMARK = [
    ("Part A — Solution Design (ICTCLD504 element 1 — analyse)", [
        ("D1", "[ICTCLD504 PC 1.1] identify and review the business's cloud architecture design; "
               "[PC 1.2] evaluate the architecture and identify the business impact of design "
               "decisions."),
        ("D2", "[ICTCLD504 PC 1.3] identify design patterns and architectural options; [PC 1.4] "
               "determine and assess the benefits and differences against the current business model "
               "and needs; [PC 1.5] confirm system design decisions according to business needs."),
        ("D3", "[ICTCLD504 PC 1.6] set business goals for security, reliability, high-performance and "
               "cost; [PC 2.1] evaluate and confirm performance metrics; [PE 3] determine performance "
               "metrics and business goals."),
    ]),
    ("Part A — Solution Design (ICTCLD504 element 2 — design & improve)", [
        ("D4", "[ICTCLD504 PC 2.2] select and improve compute, storage, database and network "
               "resources according to business needs; [PC 2.3] review and improve the architecture "
               "to enhance security, reliability, scalability and cost optimisation; [PE 1] assess, "
               "identify and improve cloud architecture according to design decisions."),
        ("D5", "[ICTCLD504 PC 1.2] business impact of the regulatory requirements; [PC 2.3] improve "
               "the architecture to meet them (the India residency slice). [AC 5] specific "
               "requirements and legislative requirements."),
        ("D6", "[ICTCLD504 PC 2.4] document the proposed architecture; [FS Writing] writes technical "
               "data logically; [KE 4] design principles for cloud applications; [KE 5] migrating "
               "principles; [KE 9] features of cloud services to improve the four concerns."),
        ("D7", "[ICTCLD504 KE 1] industry technology standards; [KE 2] industry-standard hardware/"
               "software products; [KE 3] methods and impacts of cloud adoption; [KE 6] use of object "
               "storage for static web sites; [KE 8] testing/debugging incl. avoiding single point "
               "failures — written contextual responses against the student's own design."),
    ]),
    ("Part B — Design presentation (ICTCLD504 element 2 close)", [
        ("D8", "[ICTCLD504 PC 2.4] present the proposed architecture for review to required personnel; "
               "[FS Oral communication] presents proposed solutions using appropriate industry "
               "language and confirms requirements through questioning."),
        ("D9", "[ICTCLD504 PC 2.5] obtain sign off to proceed to deployment with required personnel."),
    ]),
    ("Document quality (Foundation Skills)", [
        ("D10", "[ICTCLD504 FS Reading] interprets complex technical and operational documentation; "
                "[FS Writing] writes and edits technical data in a logical manner. (FS Problem solving "
                "and Self-management are co-evidenced through the analysis and design choices.)"),
    ]),
]

STUDENT_INTRO = [
    ("The engagement and your role", "Heading 2"),
    ("YAT College is a Registered Training Organisation (RTO) based at 175 Cremorne Street, Cremorne "
     "VIC. Following a new campus partnership in India, YAT wants the cloud infrastructure of its "
     "Ledgerline (Accounting) system confirmed as stable, reliable and fit for purpose, and "
     "compliant with the Indian regulatory requirements that now apply — and improved where it falls "
     "short. The improvement is to the cloud infrastructure only; the Ledgerline application and its "
     "data are not yours to change.", "Assessor text"),
    ("You are an MP Tech Solutions (MTS) consultant, reporting to Pat Lin (MTS Senior Consultant), "
     "who liaises with Sam Walker, the YAT ICT Manager. In AT1 you analyse the current cloud "
     "infrastructure and design the improvement, then present it for approval. AT1 is your own "
     "individual work. (The team plans and builds the agreed design in AT2, and each engineer "
     "deploys it in AT3.)", "Assessor text"),
    ("The engagement framing and the current-state records are provided on the intranet (the "
     "Ledgerline Cloud Infrastructure Improvement project, and the Accounting System Baseline Design, "
     "Infrastructure Specifications, Application Specification and Operational Costing). The analysis "
     "and the design are yours to produce.", "Assessor text"),
]

STUDENT_TASK = [
    ("Part A — your Solution Design", "Heading 2"),
    ("Using the YAT Solution Design template (download from the intranet's Templates section), "
     "analyse the current single-AZ Ledgerline infrastructure and design the improvement. Your "
     "Solution Design must:", "Assessor text"),
    ("• Review and evaluate the current architecture and identify the business impact of its design "
     "(its single points of failure, its idle-profile cost, its scalability headroom, and its "
     "position against the Indian regulatory requirements).", "Assessor text"),
    ("• Set measurable security, reliability, scalability and cost goals and the performance metrics "
     "you will measure the design against.", "Assessor text"),
    ("• Design the improvement across the four components — network, compute, database and storage — "
     "improving each across all four concerns into one integrated architecture, and include the "
     "infrastructure changes that meet the Indian regulatory requirements.", "Assessor text"),
    ("• Document and justify each improvement on a cost-versus-benefit basis — proportionate to an "
     "internal, business-hours finance system (sound engineering, not gold-plating).", "Assessor text"),
    ("Knowledge Evidence appendix (required). At the end of your Solution Design, add a Knowledge "
     "Evidence appendix and answer the following in your own words, about your own design:", "Assessor text"),
    ("• How does the accounting system's use of object storage differ from a system that serves "
     "objects publicly (for example a website), and how would you provision that storage if it were "
     "needed here?", "Assessor text"),
    ("• How does your design avoid single points of failure, and how would you test that it does?", "Assessor text"),
    ("• For each of the four optimisation concerns, name a cloud feature your design uses and the "
     "improvement it delivers.", "Assessor text"),
    ("Part B — your design presentation", "Heading 2"),
    ("Present your proposed architecture to the required personnel (your assessor, in the role of "
     "the YAT ICT Manager / MTS Senior Consultant) for review. Walk through the design and its "
     "justification, respond to questions, and obtain sign-off to proceed to deployment.", "Assessor text"),
    ("Submit the populated Solution Design (.docx) with the Knowledge Evidence appendix completed, "
     "and deliver the Part B presentation at the scheduled time.", "Assessor text"),
]

TIPS = [
    ("Tips for success", "Heading 2"),
    ("Improve the infrastructure, not the application. The Ledgerline application and its data are "
     "fixed inputs — you are improving the cloud infrastructure underneath them.", "Assessor text"),
    ("Cover all four concerns on the whole system. Security, reliability, scalability and cost each "
     "need to be addressed — don't leave one out.", "Assessor text"),
    ("Be proportionate. This is an internal, business-hours finance system; justify each improvement "
     "against the business need rather than adding capability by default.", "Assessor text"),
    ("Scalability means the ability to scale on demand, not a forecast that load will grow — design "
     "for elastic headroom, and be ready to explain how you would demonstrate it.", "Assessor text"),
    ("Write to your own design. The Knowledge Evidence appendix and the presentation are about your "
     "own design choices — answer and present in your own words.", "Assessor text"),
]


def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details ----
    t_details = doc.tables[0]
    set_cell_content(t_details.rows[1].cells[1], DETAILS["qualification"])
    set_cell_content(t_details.rows[2].cells[1], DETAILS["units"])
    set_cell_content(t_details.rows[3].cells[1], DETAILS["task_title"])
    set_cell_content(t_details.rows[4].cells[1], DETAILS["task_number"])

    # ---- Table 1: Teacher/Assessor instructions ----
    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Resources required"), RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA_STATEMENT)
    cond_row = t_instr.add_row()
    set_cell_content(cond_row.cells[0], "Assessment Conditions & Setup Requirements")
    for r in cond_row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell_content(cond_row.cells[1], CONDITIONS)

    # ---- Table 2: Marking Guide ----
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)  # keep header rows
    add_section_row(t_mark, "Part A — Solution Design")
    for c in PART_A:
        add_criterion_row(t_mark, c)
    add_section_row(t_mark, "Part B — Design Presentation")
    for c in PART_B:
        add_criterion_row(t_mark, c)
    add_section_row(t_mark, "Document quality")
    for c in QUALITY:
        add_criterion_row(t_mark, c)

    # ---- Instructions to Student ----
    doc.add_paragraph("Instructions to Student", style="Heading 1")
    for text, style in (STUDENT_INTRO + STUDENT_TASK + TIPS):
        doc.add_paragraph(text, style=style)

    # ---- Benchmark (UoC traceability) ----
    doc.add_paragraph("Marking Benchmark — UoC traceability", style="Heading 1")
    doc.add_paragraph(
        "For each criterion, the unit-of-competency items it evidences. Every PC, PE, KE and FS "
        "item allocated to AT1 in the cluster assessment plan is covered below.", style="Assessor text")
    for part_title, rows in BENCHMARK:
        doc.add_paragraph(part_title, style="Heading 2")
        for cid, uoc in rows:
            p = doc.add_paragraph(style="Assessor text")
            run = p.add_run(f"{cid}  —  ")
            run.bold = True
            p.add_run(uoc)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT1/AT1-Design-Assessor.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
