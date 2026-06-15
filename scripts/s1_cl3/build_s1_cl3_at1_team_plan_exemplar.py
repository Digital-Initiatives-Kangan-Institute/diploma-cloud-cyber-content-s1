#!/usr/bin/env python3
"""Build the worked Team Plan EXEMPLAR (.docx) — S1-CL3 AT1 Part A assessor marking reference.

A complete, worked model answer for the YAT / MTS Team Plan that a student authors in AT1
Part A (BSBXTW401 element 1 + the task-allocation slice of element 2). Reuses the YAT branding
helpers (build_bc_template) and the exemplar prose/table helpers (build_bc_exemplar), so it
carries identical YAT brand styling and retains the UoC `Evidences:` tags (this artefact is
AT1-linked and internal to assessors).

The blank student-facing Team Plan template is derived from this exemplar by stripping the
worked content and the UoC tags and replacing them with section guidance (build_team_plan_template).

Scenario: the MTS improvement team engaged to improve YAT's Ledgerline (Accounting) cloud
infrastructure (S1-CL3). The team is four members, one per improvement dimension (security,
reliability, scalability, cost); the team-lead role rotates. AT1 is individual — each student
authors their own Team Plan for their team; this is one worked example.

Usage:  python scripts/build_cl3_team_plan_exemplar.py [output.docx]
Default: S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT1/AT1-exemplar-team-plan.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_bullet_list  # noqa: E402
import build_bc_template as bc   # noqa: E402  (branding palette + styles + header/footer)
import build_bc_exemplar as ex   # noqa: E402  (uoc / para / bullets / etable helpers)

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path):
    doc = Document()
    bc.configure_styles(doc)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    bc.build_header_footer(sec)

    # ---- COVER ----
    bc.wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(bc.ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(bc.GREY)
    bc.paragraph_bottom_rule(doc.add_paragraph(), bc.TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Team Plan")
    sub = doc.add_paragraph().add_run("Ledgerline Cloud Infrastructure Improvement")
    sub.font.size = Pt(15); sub.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA); sub.bold = True
    note = doc.add_paragraph().add_run(
        "Assessor exemplar — internal marking reference (not for distribution to students)")
    note.italic = True; note.font.size = Pt(10); note.font.color.rgb = RGBColor.from_string(bc.GREY)
    doc.add_paragraph()

    cover = [
        ("Prepared for", "Pat Lin, MTS Senior Consultant (engagement supervision)"),
        ("Prepared by", "MTS Improvement Team Lead — engagement-setup phase"),
        ("Engagement", "Ledgerline Cloud Infrastructure Improvement — S1-CL3"),
        ("Team", "MTS improvement team of four — one member per improvement dimension"),
        ("Document version", "v1.0"),
        ("Classification", "Internal — MTS / YAT"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        bc.set_cell_borders(cells[0]); bc.set_cell_borders(cells[1]); bc.shade_cell(cells[0], bc.CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    bc.add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h2 = lambda t: doc.add_paragraph(t, style="Heading 2")

    h1("1. Purpose and team objectives")
    ex.uoc(doc, "[BSBXTW401 PC 1.1]")
    ex.para(doc, "MP Tech Solutions (MTS) has been engaged by YAT College to improve the cloud "
                 "infrastructure of its Ledgerline (Accounting) system, following YAT's India-campus "
                 "partnership. This Team Plan sets out how our improvement team will work to deliver the "
                 "engagement. It is the team lead's plan; the team-lead role rotates, and each member "
                 "leads at least one working session.")
    ex.para(doc, "Common objective. Deliver one integrated improvement to Ledgerline's cloud "
                 "infrastructure — confirmed stable, reliable, fit for purpose, and compliant with the "
                 "Indian regulatory requirements — proposed to YAT in an approved business case (AT2) and "
                 "implemented and signed off (AT3), within YAT's budget and Change Management Procedure.")
    ex.para(doc, "Responsibilities. Each member owns one improvement dimension (security, reliability, "
                 "scalability, cost) through the analysis, the design and the implementation, and "
                 "contributes their dimension to the team's single integrated architecture. The team lead "
                 "coordinates, chairs the working session in their rotation, and is the point of contact "
                 "with the MTS Senior Consultant (Pat Lin) and the YAT ICT Manager (Sam Walker).")
    ex.para(doc, "Required outcomes. Two outcomes gate the engagement: the deploy sign-off on the "
                 "Improvement Business Case at the end of AT2, and the final sign-off on the implemented "
                 "improvements at the end of AT3.")

    h1("2. Team structure and dimension ownership")
    ex.uoc(doc, "[BSBXTW401 PC 1.1] (responsibilities)")
    ex.para(doc, "The team is four members, each owning one improvement dimension so that the team's "
                 "integrated architecture addresses all four. The team-lead role rotates across the "
                 "working sessions.")
    ex.etable(doc, ["Member", "Dimension owned", "Focus of the dimension"],
              [["Maya N. (lead, this phase)", "Security",
                "IAM, encryption, network/security layers, and the controls that meet the Indian "
                "regulatory requirements (log residency, breach-reporting readiness)."],
               ["Theo R.", "Reliability",
                "Fault tolerance, backup and recovery, and availability appropriate to a business-hours "
                "finance system."],
               ["Ina P.", "Scalability",
                "Compute/storage elasticity and performance headroom for period-end load."],
               ["Cooper L.", "Cost",
                "Right-sizing, pricing models (reserved vs on-demand), and ongoing cost efficiency."]],
              widths=[4.2, 3.0, 8.3])

    h1("3. Performance expectations")
    ex.uoc(doc, "[BSBXTW401 PC 1.2]")
    ex.para(doc, "Each member's expected outcomes, goals and behaviours, set in line with the team "
                 "objective and YAT's policies (the Change Management Procedure, the Security and Incident "
                 "Response Policy, and the Privacy / Data Handling Policy).")
    ex.etable(doc, ["Member (dimension)", "Expected outcomes", "Goals", "Expected behaviours"],
              [["Maya N. (Security)",
                "A security/compliance improvement design and as-built record for the owned dimension.",
                "Close the compliance gaps identified in analysis; meet the security goals proportionately.",
                "Designs to the requirements; raises risks early; respects the change discipline."],
               ["Theo R. (Reliability)",
                "A reliability improvement design and as-built record.",
                "Meet the agreed reliability goals without over-engineering a business-hours system.",
                "Justifies each change on the business need; tests against the metrics."],
               ["Ina P. (Scalability)",
                "A scalability improvement design and as-built record.",
                "Provide proportionate headroom for period-end load.",
                "Measures before and after; avoids gold-plating."],
               ["Cooper L. (Cost)",
                "A cost-optimisation design and as-built record.",
                "Keep the improvement cost-justified and the ongoing cost impact explicit.",
                "Brings the cost view to every design decision; flags cost risk."]],
              widths=[3.4, 4.0, 4.0, 4.1])

    h1("4. Accountability strategies")
    ex.uoc(doc, "[BSBXTW401 PC 1.3]")
    ex.para(doc, "Members are held accountable for their roles and responsibilities through:")
    add_bullet_list(doc, [
        "Owned-dimension deliverables — each member is named against their dimension's analysis "
        "contribution, design and as-built record, so the work is individually attributable.",
        "The two sign-off gates — the AT2 deploy sign-off and the AT3 final sign-off are the formal "
        "accountability checkpoints with YAT.",
        "The led-meeting record — when a member chairs, the assessor's observation record captures how "
        "they led; the team's decisions are minuted against owners and due dates.",
        "Peer review at integration — each dimension's design is reviewed by the team before it is "
        "integrated, so gaps and conflicts surface before the business case.",
        "Escalation path — unresolved issues are escalated to the team lead, then to Pat Lin (MTS "
        "Senior Consultant), per the Engagement Role Brief.",
    ])

    h1("5. Task allocation")
    ex.uoc(doc, "[BSBXTW401 PC 2.2] · [BSBXTW401 PE 1]")
    ex.para(doc, "The improvement work is allocated by dimension, with the instruction each member needs "
                 "and an allowance for the contingencies in §6. Each member carries their dimension across "
                 "all three phases.")
    ex.etable(doc, ["Member", "Allocated tasks (AT1 → AT2 → AT3)", "Instruction / notes"],
              [["Maya N. (Security)",
                "Contribute the security & compliance findings to the analysis; design the security "
                "improvement and the controls that meet the Indian Regulatory Requirements; implement, "
                "test and document them.",
                "Work to the Indian Regulatory Requirements as provided; the Compliance area owns the "
                "law. Coordinate with Cost on any licensing impact."],
               ["Theo R. (Reliability)",
                "Reliability findings → reliability design (backups/recovery/availability) → implement "
                "and test against the reliability metrics.",
                "Keep proportionate to a business-hours finance system; align recovery with Finance's "
                "expectations."],
               ["Ina P. (Scalability)",
                "Scalability findings → scalability design (elasticity/headroom) → implement and "
                "load-test.",
                "Measure period-end load first; avoid standing capacity that Cost would flag."],
               ["Cooper L. (Cost)",
                "Cost baseline and findings → cost-optimisation design and the business-case costings → "
                "implement and verify the cost impact.",
                "Owns the business-case cost view; reconcile every dimension's design against budget."]],
              widths=[3.0, 7.5, 5.0])

    h1("6. Contingency planning")
    ex.uoc(doc, "[BSBXTW401 PC 1.4]")
    ex.para(doc, "The contingencies most likely to impact a small four-person team on a time-boxed "
                 "engagement, and the plan for each.")
    ex.etable(doc, ["Contingency", "Likelihood", "Plan"],
              [["Unplanned leave or absence of a member",
                "Medium",
                "Cross-brief at each working session so no dimension is a single point of failure; the "
                "absent member's owned tasks are picked up by the lead or an adjacent-dimension owner "
                "from the documented design notes; deadlines re-planned if needed."],
               ["Re-allocation of work tasks",
                "Medium",
                "A small schedule buffer before each sign-off gate; adjacent dimensions (security/cost, "
                "reliability/scalability) are paired so re-allocation stays within a familiar area."],
               ["Succession for an important team role (the lead)",
                "Low",
                "The rotating-chair model means every member can lead; a short written handover passes "
                "the open actions and decisions to the next chair."],
               ["A dimension's design conflicts with another at integration",
                "Medium",
                "Peer review at integration (§4) surfaces conflicts before the business case; the team "
                "lead arbitrates and minutes the decision."]],
              widths=[5.0, 2.3, 8.2])

    h1("7. Team operating norms")
    ex.uoc(doc, "[BSBXTW401 FS Get the work done]")
    ex.para(doc, "The team works through a series of rotating-chair working meetings, each chaired by a "
                 "different member. Decisions are minuted against an owner and a due date. Production-"
                 "affecting work follows YAT's Change Management Procedure and avoids the monthly "
                 "accounting close and end-of-financial-year periods. Communication with YAT runs through "
                 "the team lead to the YAT ICT Manager, with Pat Lin (MTS Senior Consultant) as the "
                 "supervision and escalation point.")

    # ---- KNOWLEDGE EVIDENCE ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1("Appendix — Knowledge Evidence")
    ex.para(doc, "Underlying knowledge as applied to this team and this engagement.")
    ex.etable(doc, ["Question", "Response"],
              [["Which organisational requirements (workplace policies, codes of conduct, organisational "
                "reputation and culture) shape how the team must operate, and how does the plan reflect "
                "them? [BSBXTW401 KE 1]",
                "YAT's Change Management Procedure (all production change is gated — reflected in the "
                "two sign-off gates and the change discipline in §7), the Security and Incident Response "
                "Policy and the Privacy / Data Handling Policy (the security/compliance dimension designs "
                "to them), and YAT's reputation as an RTO handling student and financial data (the team "
                "errs toward proportionate, well-justified change rather than risk)."],
               ["Which legislative requirements are relevant to the team's work, and how does the plan "
                "account for them? [BSBXTW401 KE 2]",
                "The Indian regulatory requirements arising from the India-campus operation (DPDP Act, "
                "CERT-In log residency, financial-records obligations) — provided by the Compliance area "
                "and carried by the security dimension; and Australian obligations (the Privacy Act / "
                "Australian Privacy Principles) for YAT data generally. The team designs to these "
                "requirements; it does not interpret the law (that is the Compliance area's role)."],
               ["What typical workplace contingencies could impact the team, and how does the plan "
                "handle them? [BSBXTW401 KE 9]",
                "Unplanned leave/absence, re-allocation of work tasks, and succession for an important "
                "role (here, the rotating lead). Each is addressed in §6 — cross-briefing so no dimension "
                "is a single point of failure, a schedule buffer and paired adjacent dimensions for "
                "re-allocation, and the rotating-chair model plus a written handover for succession."]],
              widths=[6.5, 9.0])

    # ---- SIGN-OFF ----
    h1("Sign-off")
    ex.etable(doc, ["Role", "Name", "Date", "Signature"],
              [["Prepared by (team lead)", "Maya N. (MTS improvement team)", "", ""],
               ["Reviewed by", "Pat Lin (MTS Senior Consultant)", "", ""],
               ["Noted by", "Sam Walker (YAT ICT Manager)", "", ""]],
              widths=[5.0, 5.5, 2.5, 3.0])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT1/AT1-exemplar-team-plan.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
