#!/usr/bin/env python3
"""Build the S1-CL3 per-UoC Assessment Mapping docx files (ICTCLD504, BSBXTW401).

Mirrors the CL1/CL2 mapping docs: the institutional "Assessment Mapping Tool" template,
filled with each unit's items (PC/PE/KE/FS/AC rows) and the AT1/AT2/AT3 columns carrying
the criterion code(s) that evidence each item.

CL3 is a THREE-AT cluster (like CL1), and the write-is-the-seam split means each unit lands
in only some ATs:
  * ICTCLD504  -> AT1 (D#, el 1-2 design) + AT3 (E#, el 3-4 deploy).  Not AT2.
  * BSBXTW401  -> AT2 (I#, el 1-4 team leadership).                   Not AT1/AT3.

Two data sources, both already in the repo (no hand-typed UoC text):
  * item text + order  — parsed from units_of_competency/<UNIT>_Complete_R*.md
  * item -> AT/criterion — inverted from the AT1/AT2/AT3 assessor benchmarks
    (build_s1_cl3_at{1,2,3}_assessor.BENCHMARK). Each benchmark is inverted under its own
    AT, so the AT column is known directly (no prefix-splitting needed).

FS/AC closest-fit: the benchmarks tag most FS inline (auto-mapped). Any FS skill with no
benchmark tag, and the resource-access ACs (which map to the instrument Conditions, not to
marking criteria), are filled by judgement via FS_MAP / AC_MAP below — the CL1/CL2 convention
(the assessor instruments are NOT retro-updated to claim FS/AC; noted mapping artefact).

USAGE:
    python scripts/s1_cl3/build_s1_cl3_mapping_docs.py --dump 504
    python scripts/s1_cl3/build_s1_cl3_mapping_docs.py --dump all
    python scripts/s1_cl3/build_s1_cl3_mapping_docs.py --build all
"""
import copy
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_s1_cl3_at1_assessor as a1   # noqa: E402
import build_s1_cl3_at2_assessor as a2   # noqa: E402
import build_s1_cl3_at3_assessor as a3   # noqa: E402

from docx import Document  # noqa: E402

REPO = Path(__file__).resolve().parents[2]  # scripts/s1_cl3/ -> content repo root
UOC = REPO / "S1-CL3-Cloud-Infrastructure-Improvement" / "units_of_competency"
MAPPINGS = REPO / "S1-CL3-Cloud-Infrastructure-Improvement" / "mappings"
TEMPLATE = REPO / "kangan-templates" / "Assessment Mapping Tool.docx"

AT1_TITLE = "Cloud Infrastructure Improvement: Design (individual)"
AT2_TITLE = "Team Implementation: lead the CloudFormation write (group)"
AT3_TITLE = "Implement: deploy, test and finalise (individual)"

AT_BENCHMARKS = [("AT1", a1.BENCHMARK), ("AT2", a2.BENCHMARK), ("AT3", a3.BENCHMARK)]

UNITS = {
    "504": ("ICTCLD504", "Improve cloud infrastructure", "ICTCLD504_Complete_R1.md"),
    "401": ("BSBXTW401", "Lead and facilitate a team", "BSBXTW401_Complete_R2.md"),
}

# FS closest-fit fills — ONLY for FS skills the benchmarks do not already tag.
# Keyed by unit code -> skill name -> {AT1, AT2, AT3}. Populated after --dump review.
FS_MAP = {
    "ICTCLD504": {},
    "BSBXTW401": {},
}

# AC -> instrument Conditions map (resource ACs map to the Conditions section, not to
# marking criteria). One entry per AC in source order; final AC is the assessor-requirement
# tick row. Populated after --dump review.
AC_MAP = {
    # ICTCLD504 (8 ACs in source order). Resource/access ACs are satisfied in the AT3 deploy
    # (C3 = the AWS Academy lab environment with console/CLI access); the managed DB also rides
    # the provided lab-pack (C2). AC5 (specific/legislative requirements) is genuinely ASSESSED in
    # the design, so it auto-maps to D5 (AT1) and is left to that. AC8 = assessor requirement.
    "ICTCLD504": [
        {"AT3": "C3"},               # AC1 cloud vendor service provider
        {"AT3": "C2, C3"},           # AC2 cloud managed database service
        {"AT3": "C3"},               # AC3 cloud management console / SDK / CLI
        {"AT3": "C3"},               # AC4 integrated development environment (IDE)
        {},                          # AC5 specific/industry/legislative requirements -> auto D5 (AT1)
        {"AT3": "C3"},               # AC6 internet and web browser
        {"AT3": "C3"},               # AC7 SSH / RDP client
        {"AT1": "✓", "AT3": "✓"},    # AC8 assessor requirement (the two ICTCLD504 ATs)
    ],
    # BSBXTW401 (2 ACs). The leadership is assessed in AT2; the simulated environment is the
    # team-of-four MTS engagement (C1) within the provided scenario (C2). AC2 = assessor requirement.
    "BSBXTW401": [
        {"AT2": "C1, C2"},           # AC1 a safe working or simulated environment
        {"AT2": "✓"},                # AC2 assessor requirement
    ],
}


# ----------------------------------------------------------------------------
# Source .md parser  (verbatim from build_s1_cl2_mapping_docs)
# ----------------------------------------------------------------------------

def _sections(text):
    out, cur, buf = {}, None, []
    for line in text.splitlines():
        m = re.match(r"^#\s+(.*)", line)
        if m:
            if cur is not None:
                out.setdefault(cur, "\n".join(buf))
            cur, buf = m.group(1).strip(), []
        else:
            buf.append(line)
    if cur is not None:
        out.setdefault(cur, "\n".join(buf))
    return out


def _bullets(body):
    items = []
    for ln in body.splitlines():
        stripped = ln.strip()
        if not stripped.startswith("- "):
            continue
        text = re.sub(r"^-\s+", "", stripped).strip()
        indented = len(ln) - len(ln.lstrip()) > 0
        if indented and items:
            items[-1] = items[-1] + "\n  - " + text
        else:
            items.append(text)
    return items


def parse_source_unit(unit_code, fname):
    text = (UOC / fname).read_text(encoding="utf-8")
    secs = _sections(text)

    pcs = []          # [(pc_num, pc_text, element_text)]
    epc = secs.get("Elements and Performance Criteria", "")
    for ln in epc.splitlines():
        if not ln.strip().startswith("|"):
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if len(cells) < 2:
            continue
        elem, pc_cell = cells[0], cells[1]
        if elem.upper().startswith("ELEMENTS") or elem.startswith("Elements describe"):
            continue
        if not re.match(r"^\d+\.", elem):
            continue
        for piece in pc_cell.split("<br>"):
            piece = piece.strip()
            m = re.match(r"^(\d+\.\d+)\s+(.*)", piece)
            if m:
                pcs.append((m.group(1), m.group(2).strip(), elem))

    fss = []          # [(skill, description)]
    fsec = secs.get("Foundation Skills", "")
    for ln in fsec.splitlines():
        if not ln.strip().startswith("|"):
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if len(cells) < 2:
            continue
        skill, desc = cells[0], cells[1]
        if skill.upper() == "SKILL" or not skill or skill.startswith("---"):
            continue
        fss.append((skill, desc.replace("<br>", " ")))

    pes = _bullets(secs.get("Performance Evidence", ""))
    kes = _bullets(secs.get("Knowledge Evidence", ""))

    acsec = secs.get("Assessment Conditions", "")
    acs = _bullets(acsec)
    for ln in acsec.splitlines():
        if ln.strip().lower().startswith("assessors of this unit must satisfy"):
            acs.append(ln.strip())
    return {"pcs": pcs, "fss": fss, "pes": pes, "kes": kes, "acs": acs}


# ----------------------------------------------------------------------------
# Benchmark inverter: (unit, section, item) -> {AT1:[...], AT2:[...], AT3:[...]}
# ----------------------------------------------------------------------------

def _expand(section, numbering):
    out = []
    for part in numbering.split(","):
        part = part.strip().replace("–", "-").replace("—", "-")
        if not part:
            continue
        if "-" in part and section != "FS":
            a, b = (x.strip() for x in part.split("-", 1))
            if section == "PC" and "." in a and "." in b:
                ea, ia = a.split("."); eb, ib = b.split(".")
                if ea == eb:
                    out += [f"{ea}.{i}" for i in range(int(ia), int(ib) + 1)]
                else:
                    out += [a, b]
            else:
                try:
                    out += [str(i) for i in range(int(a), int(b) + 1)]
                except ValueError:
                    out += [a, b]
        else:
            out.append(part)
    return out


def _parse_tags(s):
    """Yield (unit, section, item) from a benchmark UoC string (inherits unit prefix)."""
    cur_unit = None
    for m in re.finditer(r"\[([^\]]+)\]", s):
        inner = m.group(1).strip()
        um = re.match(r"((?:ICTCLD|BSBXTW)\d{3})\s+(.*)", inner)
        if um:
            cur_unit, rest = um.group(1), um.group(2)
        else:
            rest = inner
        sm = re.match(r"(PC|PE|KE|AC|FS)\s+(.*)", rest)
        if not sm or cur_unit is None:
            continue
        section, numbering = sm.group(1), sm.group(2).strip()
        if section == "FS":
            yield (cur_unit, "FS", numbering)
        else:
            for item in _expand(section, numbering):
                yield (cur_unit, section, item)


def invert_benchmarks():
    out = {}
    for at, bench in AT_BENCHMARKS:
        for _part_title, rows in bench:
            for cid, uoc in rows:
                for tag in _parse_tags(uoc):
                    d = out.setdefault(tag, {"AT1": [], "AT2": [], "AT3": []})
                    if cid not in d[at]:
                        d[at].append(cid)
    return out


def _codes(inv, unit, section, item, at):
    return ", ".join(inv.get((unit, section, item), {}).get(at, []))


# ----------------------------------------------------------------------------
# Dump (verification)
# ----------------------------------------------------------------------------

def dump(unit_key):
    code, title, fname = UNITS[unit_key]
    items = parse_source_unit(code, fname)
    inv = invert_benchmarks()
    print(f"\n================  {code}  {title}  ================")
    print(f"counts: PC={len(items['pcs'])} FS={len(items['fss'])} "
          f"PE={len(items['pes'])} KE={len(items['kes'])} AC={len(items['acs'])}")

    def line(section, item, label):
        cells = [_codes(inv, code, section, item, at) for at in ("AT1", "AT2", "AT3")]
        flag = "" if any(cells) else "   <-- UNMAPPED"
        print(f"  {label:<7} AT1[{cells[0]:<10}] AT2[{cells[1]:<10}] AT3[{cells[2]:<10}]{flag}")

    print("\n-- PCs --")
    for num, txt, _elem in items["pcs"]:
        line("PC", num, num)
    print("\n-- PEs --")
    for i, _txt in enumerate(items["pes"], 1):
        line("PE", str(i), f"PE{i}")
    print("\n-- KEs --")
    for i, _txt in enumerate(items["kes"], 1):
        line("KE", str(i), f"KE{i}")
    print("\n-- FSs (skill name must match the benchmark [FS ...] tag) --")
    for skill, _desc in items["fss"]:
        cells = [_codes(inv, code, "FS", skill, at) for at in ("AT1", "AT2", "AT3")]
        fm = FS_MAP.get(code, {}).get(skill, {})
        merged = [c or fm.get(at, "") for c, at in zip(cells, ("AT1", "AT2", "AT3"))]
        flag = "" if any(merged) else "   <-- UNMAPPED (needs FS_MAP)"
        print(f"  {skill:<28} AT1[{merged[0]:<10}] AT2[{merged[1]:<10}] AT3[{merged[2]:<10}]{flag}")
    print("\n-- ACs (source order; resource ACs need AC_MAP -> Conditions) --")
    for i, txt in enumerate(items["acs"], 1):
        cells = [_codes(inv, code, "AC", str(i), at) for at in ("AT1", "AT2", "AT3")]
        am = AC_MAP.get(code, [])
        amrow = am[i - 1] if i - 1 < len(am) else {}
        merged = [c or amrow.get(at, "") for c, at in zip(cells, ("AT1", "AT2", "AT3"))]
        flag = "" if any(merged) else "   <-- UNMAPPED (needs AC_MAP)"
        print(f"  AC{i:<2} AT1[{merged[0]:<8}] AT2[{merged[1]:<8}] AT3[{merged[2]:<8}]{flag}  {txt[:55]}")


# ----------------------------------------------------------------------------
# docx generation
# ----------------------------------------------------------------------------

T_DETAILS, T_DESC, T_AC, T_PC, T_PE, T_KE, T_FS = 1, 2, 3, 4, 5, 6, 7
# AT1/AT2/AT3 column indices per table (AT4/AT5 columns exist in the template; left blank).
AT_COLS = {"AC": (1, 2, 3), "PC": (2, 3, 4), "PE": (1, 2, 3), "KE": (1, 2, 3), "FS": (3, 4, 5)}
CLEAR_COLS = {"AC": (4, 5), "PC": (5, 6), "PE": (4, 5), "KE": (4, 5), "FS": (6, 7)}


def _put(cell, text):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r.text = ""
    lines = (text or "").split("\n")
    (p0.runs[0] if p0.runs else p0.add_run("")).text = lines[0]
    for ln in lines[1:]:
        cell.add_paragraph().add_run(ln)


def _reset(table, n_header=2):
    rows = list(table.rows)
    template_tr = copy.deepcopy(rows[n_header]._tr)
    for row in rows[n_header:]:
        row._tr.getparent().remove(row._tr)
    return template_tr


def _add(table, template_tr):
    table._tbl.append(copy.deepcopy(template_tr))
    return table.rows[-1]


def _fill_ats(row, section, codes3):
    for col, val in zip(AT_COLS[section], codes3):
        _put(row.cells[col], val)
    for col in CLEAR_COLS[section]:
        _put(row.cells[col], "")


def build_unit(unit_key):
    code, title, fname = UNITS[unit_key]
    items = parse_source_unit(code, fname)
    inv = invert_benchmarks()
    doc = Document(str(TEMPLATE))

    det = doc.tables[T_DETAILS]
    _put(det.rows[1].cells[1], code)
    _put(det.rows[1].cells[3], title)
    _put(det.rows[1].cells[5], "1")
    _put(det.rows[2].cells[1], "ICT50220")
    _put(det.rows[2].cells[3], "Diploma of Information Technology")

    desc = doc.tables[T_DESC]
    for r, t in ((2, AT1_TITLE), (3, AT2_TITLE), (4, AT3_TITLE)):
        _put(desc.rows[r].cells[1], t)
        _put(desc.rows[r].cells[2], "1.0")

    # PCs
    pc_tbl = doc.tables[T_PC]
    tmpl = _reset(pc_tbl)
    last_elem = None
    for num, txt, elem in items["pcs"]:
        row = _add(pc_tbl, tmpl)
        _put(row.cells[0], elem if elem != last_elem else "")
        last_elem = elem
        _put(row.cells[1], f"{num} {txt}")
        _fill_ats(row, "PC", [_codes(inv, code, "PC", num, at) for at in ("AT1", "AT2", "AT3")])

    # PEs / KEs
    for tnum, key, sect in ((T_PE, "pes", "PE"), (T_KE, "kes", "KE")):
        tbl = doc.tables[tnum]
        tmpl = _reset(tbl)
        for i, txt in enumerate(items[key], 1):
            row = _add(tbl, tmpl)
            _put(row.cells[0], txt)
            _fill_ats(row, sect, [_codes(inv, code, sect, str(i), at) for at in ("AT1", "AT2", "AT3")])

    # FSs (auto from benchmark tags; FS_MAP fills gaps)
    fs_tbl = doc.tables[T_FS]
    tmpl = _reset(fs_tbl)
    for skill, fdesc in items["fss"]:
        row = _add(fs_tbl, tmpl)
        _put(row.cells[0], skill)
        _put(row.cells[1], "")
        _put(row.cells[2], fdesc)
        fm = FS_MAP.get(code, {}).get(skill, {})
        merged = [_codes(inv, code, "FS", skill, at) or fm.get(at, "") for at in ("AT1", "AT2", "AT3")]
        _fill_ats(row, "FS", merged)

    # ACs (auto for tagged ACs; AC_MAP fills resource ACs -> Conditions)
    ac_tbl = doc.tables[T_AC]
    tmpl = _reset(ac_tbl)
    am = AC_MAP.get(code, [])
    for i, txt in enumerate(items["acs"], 1):
        row = _add(ac_tbl, tmpl)
        _put(row.cells[0], txt)
        amrow = am[i - 1] if i - 1 < len(am) else {}
        merged = [_codes(inv, code, "AC", str(i), at) or amrow.get(at, "") for at in ("AT1", "AT2", "AT3")]
        _fill_ats(row, "AC", merged)

    MAPPINGS.mkdir(parents=True, exist_ok=True)
    out = MAPPINGS / f"{code}_Assessment_Mapping.docx"
    doc.save(str(out))
    print(f"Wrote {out}  (PC {len(items['pcs'])}, PE {len(items['pes'])}, "
          f"KE {len(items['kes'])}, FS {len(items['fss'])}, AC {len(items['acs'])})")


def main():
    if len(sys.argv) >= 3 and sys.argv[1] in ("--dump", "--build"):
        targets = ["504", "401"] if sys.argv[2] == "all" else [sys.argv[2]]
        for t in targets:
            (dump if sys.argv[1] == "--dump" else build_unit)(t)
    else:
        print("Usage: build_s1_cl3_mapping_docs.py {--dump|--build} {504|401|all}")
        sys.exit(2)


if __name__ == "__main__":
    main()
