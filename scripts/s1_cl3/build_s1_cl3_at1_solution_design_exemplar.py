#!/usr/bin/env python3
"""Build the CL3 AT1 Solution Design — the agreed Ledgerline improvement ("to be") design.

Dual purpose (one source, two outputs via the `exemplar` flag):
  * exemplar=True  -> the AT1 ASSESSOR model answer: carries inline UoC "Evidences:" tags and a
                      Knowledge-Evidence appendix (assessor reference, not for student distribution).
  * exemplar=False -> the clean, in-world APPROVED DESIGN provided to the team at the start of AT2:
                      UoC tags and the KE appendix stripped, no assessment meta-language.

The design is a proportionate improvement of the single-AZ Ledgerline (Accounting System) baseline —
ICTCLD504 el 1 (analyse) + el 2 (design). Four components (network / compute / database / storage),
each improved across all four optimisation concerns (security, reliability, scalability, cost), plus a
light India residency slice (CERT-In logs + Companies-Act books-of-account only). Provisioned as
parameterised CloudFormation. HA = Multi-AZ. Scalability = elastic-capacity-on-demand (demonstrable),
not a forecast of load growth. Cost-versus-benefit justification rides here (no separate business case).

The CloudFormation *write* is the AT2 team work and is NOT assessed against 504 (it is ICTCLD505,
done in CL2); the deploy-and-operate is AT3. This document is the design those build on.

Reuses the CL2 Solution Design exemplar pattern (build_s1_cl2_at1_solution_design_exemplar).

Usage:  python scripts/s1_cl3/build_s1_cl3_at1_solution_design_exemplar.py
        (writes BOTH the exemplar and the stripped approved design)
"""
import sys
from pathlib import Path

sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_body_paragraph, add_bullet_list  # noqa: E402
from helpers.docx_tables import add_data_table  # noqa: E402
from helpers.uoc_tags import add_uoc_evidence_tag  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from helpers.yat_brand import ADDRESS, CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402
from helpers.yat_docx_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path, exemplar=True):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # UoC tag helper — emitted only in the assessor exemplar
    def ev(tag):
        if exemplar:
            add_uoc_evidence_tag(doc, tag)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Solution Design")
    sub = doc.add_paragraph().add_run("Ledgerline Cloud Infrastructure Improvement — Approved Design")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    if exemplar:
        note = doc.add_paragraph().add_run("Assessor exemplar — AT1 marking reference with inline UoC mapping (not for distribution to students)")
        note.italic = True; note.font.size = Pt(10); note.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT Ledgerline Cloud Infrastructure Improvement"),
        ("Engagement reference", "YAT-MTS-2026-006"),
        ("Document version", "v1.0"),
        ("Prepared by", "MTS Improvement Team"),
        ("Date", "[DD/MM/YYYY]"),
        ("Submitted to", "Sam Walker (YAT ICT Manager) · YAT Finance · YAT Compliance"),
        ("Related documents", "Improvement Requirements; Indian Regulatory Requirements; Accounting "
                              "System Cloud Architecture — Baseline Design; Infrastructure Specifications; Operational Costing"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    # 1 Purpose and Scope
    h1("1. Purpose and Scope")
    ev("[ICTCLD504 PC 1.1] identify and review the business's cloud architecture design")
    add_body_paragraph(doc, "This Solution Design sets out the improvements to the cloud infrastructure of the YAT Accounting "
                 "System (Ledgerline). Following YAT's India-campus partnership, the infrastructure must be confirmed "
                 "stable, reliable, fit for purpose, and compliant with the applicable Indian regulatory requirements, "
                 "and improved across security, reliability, scalability and cost. The design improves the existing "
                 "single-Availability-Zone deployment in Sydney; it is the basis for the implementation (the "
                 "infrastructure is provisioned as code) and the as-deployed verification.")
    add_body_paragraph(doc, "In scope: the cloud infrastructure across four components — network, compute, database, and "
                 "storage — and a light India residency slice for regulatory logs and financial records. Out of scope: "
                 "the Ledgerline application and its financial data, which are preserved unchanged (IR-4); the legal "
                 "interpretation of the Indian obligations (owned by YAT Compliance — treated here as fixed inputs); and "
                 "payroll (outsourced, not on this system). Improvements are proportionate to an internal, business-hours "
                 "finance system (IR-2) — sound engineering, not gold-plating — and each is justified on a cost-versus-"
                 "benefit basis (IR-6).")

    # 2 Design Inputs and Requirements
    h1("2. Design Inputs and Requirements")
    h3("2.1 Inputs")
    ev("[ICTCLD504 PC 1.1] review the existing design according to business needs")
    add_body_paragraph(doc, "This design is built from the Improvement Requirements (IR-1 to IR-7), the Indian Regulatory "
                 "Requirements, the Accounting System Cloud Architecture Baseline Design, the Infrastructure "
                 "Specifications, and the Operational Costing. Ledgerline is an internal, staff-only, business-hours "
                 "workload (~60 users, ~15-25 concurrent, peaks of ~45-55 at month-end and EOFY, effectively idle "
                 "overnight and at weekends), reached over the Site-to-Site VPN — not public-facing. The application "
                 "runs on Windows Server with a Microsoft SQL Server database; financial records carry a 7-year "
                 "retention obligation.")
    h3("2.2 Requirements the design must meet")
    add_data_table(doc, ["Ref", "Requirement", "Source"],
              [["IR-1", "Stable and reliable in business-hours operation; at least preserve the current service level", "Improvement Requirements"],
               ["IR-2", "Fit for purpose and proportionate to an internal business-hours finance system (no gold-plating)", "Improvement Requirements"],
               ["IR-3", "Brought into compliance with the applicable Indian regulatory requirements", "Indian Regulatory Requirements"],
               ["IR-4", "Application and financial data preserved — infrastructure only, no app change, no data loss", "Improvement Requirements"],
               ["IR-5", "Operable and maintainable by YAT in-house ICT; monitoring/alerting; favour reproducible IaC", "Improvement Requirements"],
               ["IR-6", "Each improvement justified on cost-versus-benefit; operating-cost impact made explicit", "Improvement Requirements"],
               ["IR-7", "Production change follows Change Management; avoids the monthly close / EOFY Restricted Period", "Improvement Requirements"]],
              widths=[1.2, 11.0, 3.8])

    # 3 Review of the Existing Architecture (504 el 1 — analyse)
    h1("3. Review of the Existing Architecture")
    ev("[ICTCLD504 PC 1.2] evaluate the architecture and identify business impact of design decisions")
    add_body_paragraph(doc, "The baseline is a single-AZ deployment in ap-southeast-2 (Sydney): an EC2 application tier "
                 "(Windows Server, Auto Scaling group min 1 / max 2) in one Availability Zone behind an internal "
                 "Application Load Balancer; an Amazon RDS for SQL Server database (single-AZ); document attachments on "
                 "Amazon S3; gp3 EBS volumes; reached over the Site-to-Site VPN. It is functional but was provisioned "
                 "for the migration, with high-availability hardening explicitly deferred.")
    h3("3.1 Single points of failure")
    ev("[ICTCLD504 PC 1.2] identify business impact")
    add_bullet_list(doc, [
        "The application tier runs in a single Availability Zone — an AZ or instance failure takes the system down.",
        "The database is single-AZ with no standby — an AZ or instance failure means an outage and a restore-from-backup, "
        "which a finance system the business depends on cannot absorb cleanly during business hours.",
        "Network and data subnets exist in one AZ only, so nothing downstream can be made resilient until the network "
        "spans two AZs.",
    ])
    h3("3.2 Current state against the four concerns and compliance")
    add_data_table(doc, ["Concern", "Current state", "Gap"],
              [["Security", "Internal-only (VPN), encrypted at rest, security groups in place; patching/role posture not hardened",
                "Tighten least-privilege, managed patching, secrets and TLS (encryption at rest is already in the baseline)"],
               ["Reliability", "Single-AZ compute and database; 99.5% business-hours target",
                "Remove the single-AZ SPOFs (Multi-AZ) so AZ/instance failure is survived automatically"],
               ["Scalability", "ASG target-tracking exists; subnets/endpoints not sized for elastic growth",
                "Confirm elastic-on-demand headroom across the tiers (not extra capacity for load not expected)"],
               ["Cost", "Instances run 24/7 though the system is idle overnight; storage untiered; SQL licensing not reviewed",
                "Business-hours scheduling, storage lifecycle, right-sizing and a licensing review"],
               ["Compliance", "All logs and records sit in Sydney; no India residency",
                "CERT-In logs (180-day, in India) and Companies-Act books retrievable in India"]],
              widths=[2.6, 7.4, 6.0])
    h3("3.3 Goals and metrics")
    ev("[ICTCLD504 PC 1.6] set business goals for security, reliability, performance and cost · [ICTCLD504 PC 2.1] confirm performance metrics")
    add_data_table(doc, ["Concern", "Goal", "Metric"],
              [["Security", "Least-privilege, secrets-managed, managed-patch posture (baseline already encrypted at rest)", "0 plaintext secrets; managed-patch compliance reported; least-privilege roles/SGs in place"],
               ["Reliability", "Survive an AZ/instance failure automatically within business hours", "RTO < 15 min on AZ failure (auto-failover); 99.9% business-hours availability"],
               ["Scalability", "Elastic on demand for the month-end/EOFY peak", "Demonstrated scale-out under a controlled load test; no manual capacity change"],
               ["Cost", "Cut spend on the idle profile without losing service", "Compute hours reduced outside business hours; storage cost/GB reduced via tiering"],
               ["Compliance", "India log + books residency met", "CERT-In logs retained 180 days in ap-south-1; books-of-account retrievable in India"]],
              widths=[2.4, 6.2, 7.4])
    h3("3.4 Review findings summary")
    add_body_paragraph(doc, "Reviewed against IR-1 to IR-7, the improvement is: remove the single-AZ SPOFs (Multi-AZ across "
                 "compute and database on a two-AZ network), harden security, realise the idle-profile and storage cost "
                 "savings, confirm elastic headroom, and add the light India residency slice for regulatory logs and "
                 "financial records. The application, its SQL Server stack and its data are preserved (IR-4); every change "
                 "is proportionate and cost-justified (IR-2, IR-6).")

    # 4 Architecture Design (504 el 2 — design)
    h1("4. Architecture Design")
    ev("[ICTCLD504 PC 1.3] identify design patterns and options · [ICTCLD504 PC 1.4] assess benefits and differences against current needs · [ICTCLD504 PC 1.5] confirm system design decisions")
    h3("4.1 Assumptions and constraints")
    add_bullet_list(doc, [
        "The baseline is implemented and running; this design improves it in place, preserving the application and data (IR-4).",
        "The system stays internal (Site-to-Site VPN); there is no public ingress.",
        "Provisioning is parameterised CloudFormation across four component stacks (network / compute / database / storage), "
        "each independently authorable and deployable, integrated by cross-stack references.",
        "Production change follows Change Management and avoids the Restricted Period (monthly close, EOFY) (IR-7).",
    ])
    h3("4.2 Accounts and regions")
    ev("[ICTCLD504 PC 2.2] select resources according to business needs")
    add_body_paragraph(doc, "ap-southeast-2 (Sydney) remains the primary region for the whole system and its data — the DPDP "
                 "Act permits the India-operation personal data to remain in Australia. ap-south-1 (Mumbai) is introduced "
                 "for one purpose only: the light residency slice (§4.7) — the CERT-In system/access logs (retained in "
                 "India for 180 days) and a periodic export of the India-operation books-of-account for Companies-Act "
                 "retrievability. The main system and database stay in Sydney.")
    h3("4.3 Design approach — four components")
    add_body_paragraph(doc, "The improvement is expressed as four cloud-infrastructure components, each improved across all "
                 "four optimisation concerns. They are provisioned as separate, parameterised CloudFormation stacks so "
                 "they can be authored independently and integrated into one deployable system.")

    h3("4.4 Component 1 — Network (VPC)")
    ev("[ICTCLD504 PC 2.2] improve network resources · [ICTCLD504 PC 2.3] improve for security, reliability, scalability and cost")
    add_data_table(doc, ["Concern", "Improvement"],
              [["Security", "Tiered private subnets (app / data), NACLs between tiers, least-privilege security groups, VPC Flow Logs enabled (feed the India log slice); no public ingress, Site-to-Site VPN retained"],
               ["Reliability", "Extend the VPC across ap-southeast-2a and 2b — app and data subnets in both AZs, redundant egress — so the tiers above can be made AZ-resilient"],
               ["Scalability", "Subnet CIDRs sized with headroom; interface/gateway VPC endpoints (S3, SSM, Secrets Manager, CloudWatch) so AWS-service traffic scales without a NAT bottleneck"],
               ["Cost", "VPC endpoints cut NAT-gateway and data-transfer charges; NAT sized to the business-hours profile"]],
              widths=[2.4, 13.2])

    h3("4.5 Component 2 — Compute (EC2 application tier)")
    ev("[ICTCLD504 PC 2.2] improve compute resources · [ICTCLD504 PC 2.3] improve for the four concerns")
    add_data_table(doc, ["Concern", "Improvement"],
              [["Security", "Hardened Windows AMI, SSM Patch Manager for managed patching, least-privilege IAM instance role (no long-lived keys), security groups limited to ALB-to-app and app-to-database"],
               ["Reliability", "Auto Scaling group spanning 2a/2b behind the internal ALB; health-checked instances replaced automatically; an AZ loss leaves healthy capacity in the other"],
               ["Scalability", "Target-tracking on CPU (retained, absorbs the month-end peak); right-sized instance family; demonstrable scale-out under a controlled load test"],
               ["Cost", "Scheduled stop outside business hours (the system is idle overnight and at weekends — the largest single saving), right-sized instances, gp3 volumes"]],
              widths=[2.4, 13.2])

    h3("4.6 Component 3 — Database (RDS for SQL Server)")
    ev("[ICTCLD504 PC 2.2] improve database resources · [ICTCLD504 PC 2.3] improve for the four concerns")
    add_data_table(doc, ["Concern", "Improvement"],
              [["Security", "Encryption in transit (TLS); credentials in Secrets Manager with rotation; security group restricting access to the app tier; not publicly accessible (encryption at rest is already enabled in the baseline)"],
               ["Reliability", "Multi-AZ deployment — synchronous standby in the second AZ with automatic failover; automated backups and point-in-time recovery"],
               ["Scalability", "Right-sized instance class; storage autoscaling enabled; a read replica noted as headroom for month-end reporting (not provisioned by default — proportionate)"],
               ["Cost", "Right-sized to measured load; a Savings Plan / Reserved Instance for the steady baseline; SQL Server licensing model reviewed (License Included vs BYOL) and made explicit; backup retention tuned (long-term financial-records retention handled by export to S3, not live RDS)"]],
              widths=[2.4, 13.2])

    h3("4.7 Component 4 — Storage (S3 + EBS) and the India residency slice")
    ev("[ICTCLD504 PC 2.2] improve storage resources · [ICTCLD504 PC 2.3] improve for the four concerns")
    add_data_table(doc, ["Concern", "Improvement"],
              [["Security", "S3 Block Public Access; bucket policies; S3 access logging (encryption at rest is already enabled in the baseline)"],
               ["Reliability", "S3 durability with versioning for document attachments; automated EBS snapshots; lifecycle policies"],
               ["Scalability", "S3 absorbs unbounded document growth (~4 GB/yr) with no capacity planning; EBS gp3 volumes resized live (elastic volumes) as the footprint grows"],
               ["Cost", "S3 lifecycle / Intelligent-Tiering moves older scanned documents to IA/Glacier — a large saving over the 7-year retention"]],
              widths=[2.4, 13.2])
    add_body_paragraph(doc, "Residency slice (owned by this component): an ap-south-1 (Mumbai) S3 bucket holds (a) the CERT-In "
                 "system/access logs — VPC Flow Logs, CloudTrail, ALB/application logs — with a >=180-day lifecycle "
                 "and quick retrievability for the 6-hour reporting obligation, and (b) a periodic, encrypted, "
                 "version-locked export of the India-operation books-of-account for Companies-Act retrievability. The "
                 "main system and primary data remain in Sydney.")

    h3("4.8 Monitoring and operability")
    ev("[ICTCLD504 PC 2.1] confirm performance metrics for the application")
    add_body_paragraph(doc, "CloudWatch metrics, dashboards and alarms per component provide the measures the goals (§3.3) "
                 "are tested against and keep the environment operable by YAT in-house ICT (IR-5); CloudTrail records the "
                 "audit trail (also feeding the India log slice). Alerting covers AZ-failover, scaling events, and "
                 "backup/lifecycle health.")
    h3("4.9 Provisioning — infrastructure as code")
    add_body_paragraph(doc, "The whole design is provisioned as parameterised CloudFormation, partitioned into four component "
                 "stacks (network / compute / database / storage) with exported cross-stack references for the "
                 "dependencies, so each can be authored, deployed and updated independently and the system reproduced "
                 "consistently (IR-5).")
    h3("4.10 Goals and metrics — designed state")
    ev("[ICTCLD504 PC 1.5] confirm design decisions · [ICTCLD504 PC 2.3] improve for the four concerns")
    add_body_paragraph(doc, "Each goal in §3.3 is met by the design above: reliability by Multi-AZ across compute and "
                 "database on a two-AZ network (auto-failover); scalability by the retained ASG target-tracking plus "
                 "elastic storage and endpoint headroom (demonstrable on demand, not over-provisioned); cost by the "
                 "business-hours schedule, storage tiering, right-sizing and the licensing review; security by the "
                 "hardening, encryption and least-privilege posture; compliance by the India residency slice.")
    h3("4.11 Single points of failure removed")
    add_body_paragraph(doc, "The single-AZ compute and single-AZ database SPOFs (§3.1) are removed: the application tier "
                 "runs across two AZs behind the internal ALB, and the database runs Multi-AZ with automatic failover, on "
                 "a network that now spans both AZs.")

    # 5 Implementation Sequencing
    h1("5. Implementation Sequencing")
    ev("[ICTCLD504 PC 2.5] obtain sign off to proceed to deployment")
    add_body_paragraph(doc, "Changes are applied in dependency order, under Change Management and outside the Restricted "
                 "Period (IR-7). The build is provisioned from the component CloudFormation stacks.")
    add_data_table(doc, ["#", "Change", "Expected impact"],
              [["1", "Network stack — extend the VPC to two AZs, add subnets, endpoints, flow logs", "None — additive"],
               ["2", "Storage stack — encryption, versioning, lifecycle, and the ap-south-1 residency bucket", "None — additive"],
               ["3", "Database stack — convert RDS to Multi-AZ; encryption; secrets", "Brief failover blip; scheduled in a maintenance window"],
               ["4", "Compute stack — ASG across two AZs, hardening, business-hours schedule", "Rolling; capacity maintained throughout"]],
              widths=[0.9, 10.6, 4.1])

    # 6 Verification Plan
    h1("6. Verification Plan")
    add_body_paragraph(doc, "How the design is shown to meet the goals (§3.3); the evidence is produced at deployment.")
    add_bullet_list(doc, [
        "Reliability — force an AZ/instance failure (reboot-with-failover on RDS; terminate an app instance) and confirm automatic recovery within the RTO, with no data loss.",
        "Scalability — a controlled load test ramps to the month-end peak; confirm the ASG scales out and back and storage/endpoints absorb growth, with no manual capacity change.",
        "Security — confirm encryption at rest/in transit, least-privilege roles/SGs, patch compliance, and no plaintext secrets.",
        "Cost — confirm the business-hours schedule stops/starts compute and the storage lifecycle transitions older objects; report the operating-cost delta.",
        "Compliance — confirm CERT-In logs are written to and retained in ap-south-1 (>=180 days) and the books-of-account export is retrievable in India.",
    ])

    # 7 Out of Scope
    h1("7. Out of Scope")
    add_bullet_list(doc, [
        "The Ledgerline application and its financial data — preserved unchanged (IR-4).",
        "The legal interpretation of the Indian obligations — owned by YAT Compliance; treated here as fixed inputs.",
        "Payroll — outsourced to an external bureau; not on this system.",
        "Full localisation of the system into India — deliberately avoided; only the regulatory logs and books-of-account export sit in India (the light slice).",
    ])

    # 8 References
    h1("8. References")
    add_bullet_list(doc, [
        "Improvement Requirements (Ledgerline Cloud Infrastructure Improvement)",
        "Indian Regulatory Requirements (DPDP Act 2023; CERT-In Directions 2022; Companies Act 2013 / Income-tax retention)",
        "Accounting System Cloud Architecture — Baseline Design",
        "Accounting System Infrastructure Specifications; Operational Costing",
        "AWS Well-Architected Framework — Reliability, Security, Cost Optimization and Performance Efficiency pillars",
    ])

    # 9 Review and Approval
    h1("9. Review and Approval")
    ev("[ICTCLD504 PC 2.4] document and present the proposed architecture for review")
    add_body_paragraph(doc, "The design was reviewed against IR-1 to IR-7 and the goals (§3.3). It is presented to the "
                 "required personnel for review, and sign-off to proceed to deployment is obtained at that presentation; "
                 "the record below captures that approval.")
    add_data_table(doc, ["Role", "Name", "Decision", "Date"],
              [["Prepared by", "MTS Improvement Team", "Submitted", "[date]"],
               ["Reviewed by", "MTS Senior Consultant", "Approved for submission", "[date]"],
               ["Approved by", "Sam Walker (YAT ICT Manager)", "Approved — cleared to proceed to deployment", "[date]"]],
              widths=[3.6, 5.0, 5.4, 2.0])

    # Appendix A — Knowledge Evidence (exemplar only)
    if exemplar:
        doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
        h1("Appendix A — Knowledge Evidence (assessor reference)")
        add_body_paragraph(doc, "Model responses to the AT1 contextual knowledge questions (ICTCLD504 design knowledge). Each "
                     "asks the candidate to reason about their own design. NOT part of the design distributed to the AT2 team.")

        def ke_qa(tag, question, answer):
            add_uoc_evidence_tag(doc, tag)
            qp = doc.add_paragraph(); qr = qp.add_run(question); qr.bold = True; qr.font.size = Pt(10.5)
            add_body_paragraph(doc, answer)

        ke_qa("[ICTCLD504 KE 6] — use of object storage for static web sites",
              "Q1. Ledgerline stores scanned document attachments in object storage but is not a static web site. "
              "How does an accounting system's use of object storage differ from a system that serves objects such as "
              "images for a public web site, and how would you provision that storage if it were needed here?",
              "Ledgerline uses S3 as a durable, private store for document attachments (invoices, POs, receipts) read "
              "by the application over the VPN — access is authenticated, internal, and low-rate, so the design "
              "emphasises encryption, Block Public Access, versioning and lifecycle tiering for the 7-year retention. A "
              "static web site instead serves objects (images, pages) publicly and at high request rates, so it would "
              "add a CDN (CloudFront) in front of S3 for edge caching and latency, public-read with origin access "
              "control rather than block-public, and cache headers — none of which Ledgerline needs because it is "
              "internal and not object-serving to the public. If Ledgerline ever had to serve objects publicly, the "
              "provisioning would be S3 + CloudFront with an origin access control and tuned caching.")

        ke_qa("[ICTCLD504 KE 8] — testing and debugging techniques, including techniques to avoid single point failures",
              "Q2. Explain how your design avoids single points of failure and how you would test that it does.",
              "The baseline's single points of failure are the single-AZ application tier and single-AZ database. The "
              "design removes them by spanning the network across two AZs, running the application Auto Scaling group "
              "across both AZs behind the internal load balancer, and running the database Multi-AZ with a synchronous "
              "standby and automatic failover. Testing: terminate an application instance and confirm the ASG replaces "
              "it with no outage; reboot the database with failover and confirm the standby takes over within the "
              "recovery objective with no data loss; remove an AZ's capacity and confirm the system stays available "
              "from the other AZ.")

        ke_qa("[ICTCLD504 KE 9] — features of cloud services, including techniques to improve security, reliability, scalability and cost",
              "Q3. For each of the four optimisation concerns, name a cloud feature your design uses and the improvement it delivers.",
              "Security — KMS encryption, Secrets Manager and least-privilege IAM roles remove plaintext secrets and "
              "tighten access. Reliability — Multi-AZ RDS and a cross-AZ Auto Scaling group survive an AZ or instance "
              "failure automatically. Scalability — ASG target-tracking and elastic gp3/S3 storage add capacity on "
              "demand without manual change or over-provisioning. Cost — a business-hours start/stop schedule on the "
              "idle-overnight compute and S3 lifecycle tiering on the growing document store cut spend without reducing "
              "service.")

    # Document control
    h1("Document control")
    add_data_table(doc, ["Field", "Value"],
              [["Document version", "v1.0 — Solution Design (Ledgerline Cloud Infrastructure Improvement)"],
               ["Author", "MTS Improvement Team"],
               ["Engagement", "YAT Ledgerline Cloud Infrastructure Improvement"],
               ["Provisioned as", "Parameterised CloudFormation — network / compute / database / storage component stacks"],
               ["Implemented by", "the as-deployed Deployment Report"]],
              widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    # __file__-anchored (NOT cwd-relative) so output can't land in a phantom tree.
    content_root = Path(__file__).resolve().parents[2]          # diploma-cloud-cyber-content
    website = content_root.parent / "diploma-cloud-cyber-website"  # sibling under the umbrella
    # Exemplar (assessor marking reference) stays in the content repo's AT1 folder.
    exemplar_out = content_root / "S1-CL3-Cloud-Infrastructure-Improvement" / "assessments" / "AT1" / "AT1-exemplar-solution-design.docx"
    # Approved design (in-world, distributed at AT2) lives in the website documents folder; print to PDF there.
    approved_out = website / "public" / "documents" / "YAT-Accounting-Improved-Solution-Design.docx"
    build(str(exemplar_out), exemplar=True)
    build(str(approved_out), exemplar=False)
