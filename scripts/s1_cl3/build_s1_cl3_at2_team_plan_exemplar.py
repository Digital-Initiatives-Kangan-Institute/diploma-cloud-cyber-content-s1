#!/usr/bin/env python3
"""Build the worked Team Plan EXEMPLAR (.docx) — S1-CL3 AT2 Part A assessor marking reference.

A complete, worked model answer for the YAT / MTS Team Plan that a student produces in AT2 Part A
(BSBXTW401 element 1 + the task-allocation slice of element 2). Carries YAT brand styling and the
UoC `Evidences:` tags (assessor-internal).

Current model (write-is-the-seam): the improvement Solution Design is ALREADY approved (AT1, individual)
and provided to the team. In AT2 the team implements it — writing the approved design as one integrated,
deployable CloudFormation template, DIVIDED BY COMPONENT (network / compute / database / storage), one
component per member, integrated and validated for a team sign-off. Each member then deploys and operates
the whole system individually in AT3. There is no business case (the cost-benefit rides in the Solution
Design). The team plan allocates COMPONENTS, not optimisation concerns — each component is improved across
all four concerns (security, reliability, scalability, cost) per the approved design.

The blank student-facing Team Plan template is derived from this exemplar by stripping the worked content
and the UoC tags (build_team_plan_template).

Usage:  python scripts/s1_cl3/build_s1_cl3_at2_team_plan_exemplar.py [output.docx]
Default: S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT2/AT2-exemplar-team-plan.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_body_paragraph, add_bullet_list  # noqa: E402
from helpers.docx_tables import add_data_table  # noqa: E402
from helpers.uoc_tags import add_uoc_evidence_tag  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from helpers.yat_brand import ADDRESS, CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402
from helpers.yat_docx_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path):
    doc = Document()
    configure_styles(doc)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Team Plan")
    sub = doc.add_paragraph().add_run("Ledgerline Cloud Infrastructure Improvement")
    sub.font.size = Pt(15); sub.font.color.rgb = RGBColor.from_string(TERRACOTTA); sub.bold = True
    note = doc.add_paragraph().add_run(
        "Assessor exemplar — internal marking reference (not for distribution to students)")
    note.italic = True; note.font.size = Pt(10); note.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()

    cover = [
        ("Prepared for", "Pat Lin, MTS Senior Consultant (engagement supervision)"),
        ("Prepared by", "MTS Improvement Team Lead — implementation phase"),
        ("Engagement", "Ledgerline Cloud Infrastructure Improvement — S1-CL3"),
        ("Team", "MTS improvement team of four — one cloud component per member"),
        ("Document version", "v1.0"),
        ("Classification", "Internal — MTS / YAT"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")

    h1("1. Purpose and team objectives")
    add_uoc_evidence_tag(doc, "[BSBXTW401 PC 1.1]")
    add_body_paragraph(doc, "MP Tech Solutions (MTS) has been engaged by YAT College to improve the cloud "
                 "infrastructure of its Ledgerline (Accounting) system, following YAT's India-campus "
                 "partnership. The improvement Solution Design has been approved by YAT and is the input to "
                 "this phase. This Team Plan sets out how our improvement team will implement that design — "
                 "writing it as infrastructure-as-code. It is the team lead's plan; the team-lead role "
                 "rotates, and each member leads at least one working session.")
    add_body_paragraph(doc, "Common objective. Implement the approved improvement as **one integrated, "
                 "deployable CloudFormation template** — divided into four cloud components, one written by "
                 "each member, integrated by cross-stack references and validated by the team. The "
                 "infrastructure-as-code is the vehicle for the teamwork; each member then deploys and "
                 "operates the whole improved system individually.")
    add_body_paragraph(doc, "Responsibilities. Each member owns one cloud component (network, compute, "
                 "database, storage), writing its CloudFormation to the approved design and improving it "
                 "across all four optimisation concerns (security, reliability, scalability, cost) as that "
                 "design directs. The team lead coordinates, chairs the working session in their rotation, "
                 "and is the point of contact with the MTS Senior Consultant (Pat Lin) and the YAT ICT "
                 "Manager (Sam Walker).")
    add_body_paragraph(doc, "Required outcomes. The team's deliverable is one integrated, validated "
                 "CloudFormation template that implements the approved design; the **team sign-off on that "
                 "template** is the gate that closes the build. The design approval has already been obtained "
                 "(the Solution Design presentation); the final sign-off on the deployed system is obtained "
                 "individually at deployment.")

    h1("2. Team structure and component ownership")
    add_uoc_evidence_tag(doc, "[BSBXTW401 PC 1.1] (responsibilities)")
    add_body_paragraph(doc, "The team is four members, each owning one cloud component so that the team's "
                 "single integrated template covers the whole system. Components are allocated by expertise; "
                 "the team-lead role rotates across the working sessions.")
    add_data_table(doc, ["Member", "Component owned", "Scope of the component"],
              [["Maya N. (lead, this phase)", "Network (VPC)",
                "VPC across two Availability Zones, tiered subnets, route tables, NACLs and security groups, "
                "VPC endpoints, and flow logs - the network foundation the other components build on and "
                "reference."],
               ["Theo R.", "Compute",
                "EC2 launch template, the Auto Scaling group spread across two AZs, the internal Application "
                "Load Balancer, and the scaling policy."],
               ["Ina P.", "Database",
                "Amazon RDS for SQL Server - single-instance (Ledgerline does not support a Multi-AZ "
                "database), encrypted, with the DB subnet group, backup/point-in-time-restore configuration "
                "and the database security group."],
               ["Cooper L.", "Storage",
                "The S3 buckets (attachments, backups) with encryption, versioning and lifecycle; EBS "
                "configuration; and the India residency slice (the ap-south-1 CERT-In logs + books-of-account "
                "buckets)."]],
              widths=[4.0, 3.0, 8.5])

    h1("3. Performance expectations")
    add_uoc_evidence_tag(doc, "[BSBXTW401 PC 1.2]")
    add_body_paragraph(doc, "Each member's expected outcomes, goals and behaviours, set in line with the team "
                 "objective and YAT's policies (the Change Management Procedure, the Security and Incident "
                 "Response Policy, and the Privacy / Data Handling Policy).")
    add_data_table(doc, ["Member (component)", "Expected outcomes", "Goals", "Expected behaviours"],
              [["Maya N. (Network)",
                "A validated network component that the compute and database components import cleanly.",
                "Provide the two-AZ network foundation the approved design requires, with no public ingress.",
                "Publishes stable cross-stack exports early; raises interface changes before they break others."],
               ["Theo R. (Compute)",
                "A validated compute component (ASG across two AZs behind the internal ALB).",
                "Deliver the application-tier high availability the design specifies, proportionately.",
                "Builds to the network's exports; tests health-check and scaling behaviour."],
               ["Ina P. (Database)",
                "A validated database component - single-instance, encrypted, backed up.",
                "Meet the recovery-based reliability goal (backup/PITR), not Multi-AZ, per the design.",
                "Keeps the database untouched by later updates where the lab requires; documents the constraint."],
               ["Cooper L. (Storage)",
                "Validated storage (buckets + lifecycle) and the India residency slice.",
                "Deliver durable, cost-tiered storage and the IR-3 residency buckets.",
                "Reconciles encryption/lifecycle with the design; keeps the residency slice region-correct."]],
              widths=[3.4, 4.0, 4.0, 4.1])

    h1("4. Accountability strategies")
    add_uoc_evidence_tag(doc, "[BSBXTW401 PC 1.3]")
    add_body_paragraph(doc, "Members are held accountable for their roles and responsibilities through:")
    add_bullet_list(doc, [
        "Owned-component deliverables - each member is named against their component's CloudFormation in the "
        "integrated template, so the work is individually attributable.",
        "The team sign-off gate - the team validates the integrated template before it is handed on for "
        "deployment; that validation is the formal accountability checkpoint.",
        "The led-meeting record - when a member chairs, the assessor's observation record captures how they "
        "led; the team's decisions are minuted against owners and due dates.",
        "Peer review at integration - each component's cross-stack interface (exports it provides, imports it "
        "needs) is reviewed by the team, so reference mismatches surface before integration fails.",
        "Escalation path - unresolved issues are escalated to the team lead, then to Pat Lin (MTS Senior "
        "Consultant), per the Engagement Role Brief.",
    ])

    h1("5. Task allocation")
    add_uoc_evidence_tag(doc, "[BSBXTW401 PC 2.2] · [BSBXTW401 PE 1]")
    add_body_paragraph(doc, "The infrastructure-as-code is allocated one component per member, with the "
                 "instruction each member needs and an allowance for the contingencies in §6. Each member "
                 "writes their component to the approved design and integrates it into the team's single "
                 "deployable template.")
    add_data_table(doc, ["Member", "Component allocated", "Instruction / notes"],
              [["Maya N. (Network)",
                "Write the network CloudFormation (VPC, two-AZ subnets, route tables, NACLs, security groups, "
                "endpoints, flow logs); export the values the other components import.",
                "Land the exports first so compute and database are not blocked; coordinate CIDR/AZ choices "
                "with the whole team."],
               ["Theo R. (Compute)",
                "Write the compute CloudFormation (launch template, ASG across two AZs, internal ALB, scaling "
                "policy), importing the network's subnets and security groups.",
                "Confirm the ALB and ASG span both AZs; align the app security group with the database's."],
               ["Ina P. (Database)",
                "Write the database CloudFormation (RDS for SQL Server single-instance, encrypted, DB subnet "
                "group, backups), importing the network's data subnets and security group.",
                "Keep the instance single-AZ per the Ledgerline constraint; coordinate the SG rule with "
                "Compute."],
               ["Cooper L. (Storage)",
                "Write the storage CloudFormation (S3 buckets with encryption/versioning/lifecycle, EBS) and "
                "the separate ap-south-1 residency stack.",
                "Keep the residency buckets in the Mumbai region; reconcile lifecycle/retention with the "
                "design and the Indian Regulatory Requirements."]],
              widths=[3.0, 7.5, 5.0])

    h1("6. Contingency planning")
    add_uoc_evidence_tag(doc, "[BSBXTW401 PC 1.4]")
    add_body_paragraph(doc, "The contingencies most likely to impact a small four-person team on a time-boxed "
                 "implementation, and the plan for each.")
    add_data_table(doc, ["Contingency", "Likelihood", "Plan"],
              [["Unplanned leave or absence of a member",
                "Medium",
                "Cross-brief each component's cross-stack interface at every working session so no component "
                "is a single point of failure; the absent member's component is picked up by the lead or an "
                "adjacent-component owner from the documented interface; deadlines re-planned if needed."],
               ["Re-allocation of work tasks",
                "Medium",
                "A small schedule buffer before the integration/sign-off; adjacent components (network/"
                "compute, database/storage) are paired so re-allocation stays within a familiar area."],
               ["Succession for an important team role (the lead)",
                "Low",
                "The rotating-chair model means every member can lead; a short written handover passes the "
                "open actions and decisions to the next chair."],
               ["A component's cross-stack interface conflicts with another at integration",
                "Medium",
                "Peer review of the exports/imports (§4) surfaces reference mismatches before integration; "
                "the team lead arbitrates and minutes the decision."]],
              widths=[5.0, 2.3, 8.2])

    h1("7. Team operating norms")
    add_uoc_evidence_tag(doc, "[BSBXTW401 FS Get the work done]")
    add_body_paragraph(doc, "The team works through a series of rotating-chair working meetings, each chaired "
                 "by a different member. Decisions are minuted against an owner and a due date. The eventual "
                 "production deployment follows YAT's Change Management Procedure and avoids the monthly "
                 "accounting close and end-of-financial-year periods. Communication with YAT runs through the "
                 "team lead to the YAT ICT Manager, with Pat Lin (MTS Senior Consultant) as the supervision "
                 "and escalation point.")

    # ---- KNOWLEDGE EVIDENCE ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1("Appendix — Knowledge Evidence")
    add_body_paragraph(doc, "Underlying knowledge as applied to this team and this engagement.")
    add_data_table(doc, ["Question", "Response"],
              [["Which organisational requirements (workplace policies, codes of conduct, organisational "
                "reputation and culture) shape how the team must operate, and how does the plan reflect "
                "them? [BSBXTW401 KE 1]",
                "YAT's Change Management Procedure (all production change is gated - reflected in the change "
                "discipline in §7), the Security and Incident Response Policy and the Privacy / Data Handling "
                "Policy (the network, database and storage components are written to them), and YAT's "
                "reputation as an RTO handling student and financial data (the team errs toward proportionate, "
                "well-justified infrastructure rather than risk)."],
               ["Which legislative requirements are relevant to the team's work, and how does the plan "
                "account for them? [BSBXTW401 KE 2]",
                "The Indian regulatory requirements arising from the India-campus operation (DPDP Act, "
                "CERT-In log residency, financial-records obligations) - provided by the Compliance area and "
                "carried by the storage component's residency slice; and Australian obligations (the Privacy "
                "Act / Australian Privacy Principles) for YAT data generally. The team builds the "
                "infrastructure to these requirements; it does not interpret the law (that is the Compliance "
                "area's role)."],
               ["What typical workplace contingencies could impact the team, and how does the plan "
                "handle them? [BSBXTW401 KE 9]",
                "Unplanned leave/absence, re-allocation of work tasks, and succession for an important role "
                "(here, the rotating lead). Each is addressed in §6 - cross-briefing the component interfaces "
                "so no component is a single point of failure, a schedule buffer and paired adjacent "
                "components for re-allocation, and the rotating-chair model plus a written handover for "
                "succession."]],
              widths=[6.5, 9.0])

    # ---- SIGN-OFF ----
    h1("Sign-off")
    add_data_table(doc, ["Role", "Name", "Date", "Signature"],
              [["Prepared by (team lead)", "Maya N. (MTS improvement team)", "", ""],
               ["Reviewed by", "Pat Lin (MTS Senior Consultant)", "", ""],
               ["Noted by", "Sam Walker (YAT ICT Manager)", "", ""]],
              widths=[5.0, 5.5, 2.5, 3.0])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT2/AT2-exemplar-team-plan.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
