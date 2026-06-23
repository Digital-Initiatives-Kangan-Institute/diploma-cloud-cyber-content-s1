#!/usr/bin/env python3
"""Build the S1-CL3 AT3 ASSESSOR instrument (.docx) by populating the Kangan template.

Institutional compliance document (NOT a YAT-branded artefact): loads the official Kangan
'Project Assessment - Assessor' template and fills it in, preserving the Kangan structure and
styles exactly. Mirrors the approved CL1/CL2/CL3 assessor instruments.

AT3 = Implement (individual) — ICTCLD504 elements 3 (deploy, monitor, test) + 4 (finalise). Each
student deploys the existing-state baseline and applies the approved improvement as a CloudFormation
change-set in their own AWS Academy lab, then monitors, tests and demonstrates the whole improved
system across all four optimisation concerns (security, reliability, scalability, cost), applies
short-term refinements, documents the as-deployed result and a long-term strategy, and obtains final
sign-off. The deliverable is a Deployment Report.

The CloudFormation the team wrote in AT2 is the thing deployed here; the deploy-and-operate is the
individual ICTCLD504 evidence.

Usage:  python scripts/s1_cl3/build_s1_cl3_at3_assessor.py [output.docx]
Default: S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT3/AT3-Implement-Assessor.docx
"""
import sys
from pathlib import Path

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Assessor.docx")

sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_tables import add_criterion_row, add_section_row, clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402


# ---------- content ----------

DETAILS = {
    "qualification": "ICT50220 Diploma of Information Technology",
    "units": [
        "ICTCLD504 Improve cloud-based infrastructure",
    ],
    "task_title": "AT3 – Implement",
    "task_number": "3 of 3",
}

OVERVIEW = (
    "Students are assessed individually on deploying and operating the approved Ledgerline "
    "cloud-infrastructure improvement. Each student deploys the existing-state baseline and applies "
    "the approved improvement as a CloudFormation change-set in their own AWS Academy lab, then "
    "monitors, tests and demonstrates the whole improved system across all four optimisation "
    "concerns — security, reliability, scalability and cost — applies short-term refinements from the "
    "results, documents the as-deployed result and a long-term improvement strategy, and obtains "
    "final sign-off. The deliverable is a Deployment Report. AT3 is completed individually."
)

TASKS = [
    "In this project the student is an MP Tech Solutions (MTS) engineer deploying the approved "
    "Ledgerline cloud-infrastructure improvement for YAT College. AT3 is completed individually, in "
    "the student's own lab, and has one deliverable — a Deployment Report — produced by:",
    "• Deploying the existing-state baseline, then applying the approved improvement as a "
    "CloudFormation change-set to the running stack (deploying the approved architecture).",
    "• Monitoring and measuring the deployed architecture against the performance metrics and "
    "business goals; testing and demonstrating security, reliability, scalability and cost "
    "optimisation on the whole deployed system (including, for reliability, an application-tier "
    "failover demonstration and a database backup/restore and DR demonstration, and a scale-out "
    "demonstration for scalability).",
    "• Applying short-term refinements to the deployed resources according to the test results.",
    "• Documenting the as-deployed architecture and test results (highlighting the changes from the "
    "approved design), describing a long-term improvement strategy, and obtaining final sign-off.",
    "The approved design was produced in AT1 and written as code by the team in AT2; AT3 is each "
    "engineer's own deployment, demonstration and finalisation of the whole system.",
]

RESOURCES = [
    "Teacher/assessor supplied resources / Access to:",
    "• The baseline lab-pack — the existing-state CloudFormation (single-AZ Ledgerline) the student "
    "deploys and then upgrades.",
    "• The approved improvement design and the team's validated implementation (the CloudFormation "
    "change-set applied to the baseline).",
    "• Assessor reference combined template (the lab-pack improved.yaml) — an optional fallback the "
    "assessor may supply if a student's AT2 team implementation is not appropriate to deploy in AT3, so "
    "the individual can deploy and be assessed on their own AT3 work rather than be blocked by the "
    "team's result.",
    "• The YAT Deployment Report template (intranet Templates section).",
    "• The AWS Academy lab environment specified in the lab-pack, with the console / CLI access and "
    "the resources the deployment requires.",
]

CRITERIA_STATEMENT = (
    "To receive a Satisfactory outcome for this assessment task, the student must complete every "
    "criterion in the marking guide below to a satisfactory standard. Where a criterion is not yet "
    "satisfactory, the student is given feedback and a further attempt per the Second attempt "
    "provisions."
)

CONDITIONS = [
    "C1 — AT3 is completed individually; each student deploys and operates the system in their own "
    "AWS Academy lab and produces their own Deployment Report.",
    "C2 — The baseline lab-pack, the approved design, and the team's validated implementation are "
    "provided to the student.",
    "C3 — The student deploys into the AWS Academy lab environment specified in the lab-pack, with "
    "console / CLI access and the resources the deployment requires.",
    "C4 — The final sign-off is recorded by the assessor on review of the deployed system and the "
    "Deployment Report.",
]

# Marking guide — Part A (Deploy & operate, el 3) and Part B (Finalise, el 4). UoC traceability in the Benchmark.
PART_A = [
    "E1 — Deploy: deploys the existing-state baseline and applies the approved improvement as a "
    "CloudFormation change-set to the running stack, using the console / CLI (deploying the approved "
    "architecture).",
    "E2 — Monitor and measure: monitors and measures the deployed architecture against the "
    "performance metrics and business goals (e.g. CloudWatch metrics and alarms).",
    "E3 — Test and demonstrate: tests and demonstrates security, reliability, scalability and cost "
    "optimisation on the deployed resources — for reliability, an application-tier failover "
    "demonstration (terminate an instance; the Auto Scaling group recovers across AZs) and a database "
    "backup/restore and cross-Region DR demonstration (the database is single-instance — no Multi-AZ "
    "failover); and a scalability (scale-out) demonstration on the whole system.",
    "E4 — Refine: applies short-term refinements to the deployed resources according to the test "
    "results (e.g. tunes a scaling policy, an alarm threshold, a security group, or a size).",
]

PART_B = [
    "E5 — As-deployed documentation: documents the as-deployed architecture and the test results, "
    "highlighting the changes and improvements from the approved design.",
    "E6 — Long-term strategy: describes long-term improvement strategies and their benefits as "
    "applied to the deployed resources.",
    "E7 — Final sign-off: obtains final sign-off from the required personnel.",
]

QUALITY = [
    "E8 — Document quality: the Deployment Report uses the YAT Deployment Report template, plain "
    "professional English, is complete and internally consistent, and includes the evidence "
    "(screenshots / logs / test output) of the deployment and the demonstrations.",
]

# Benchmark: per-criterion UoC evidenced. Every PC/PE/KE/FS item allocated to AT3 in the cluster
# assessment plan (ICTCLD504 el 3-4) is covered below.
BENCHMARK = [
    ("Part A — Deploy, monitor and test (ICTCLD504 element 3)", [
        ("E1", "[ICTCLD504 PC 3.1] deploy the approved architecture on the cloud platform; [PE 2] "
               "deploy at least one architecture design; [PE 4] use cloud management consoles, SDKs "
               "or command line tools."),
        ("E2", "[ICTCLD504 PC 3.2] monitor and measure the architecture against performance metrics "
               "and business goals; [KE 10] techniques, methods and industry-standard metrics and "
               "goals used to monitor performance of cloud resources."),
        ("E3", "[ICTCLD504 PC 3.3] test and demonstrate security, reliability, scalability and cost "
               "optimisation of the deployed resources; [PE 2] test and measure the design against "
               "principles, metrics and goals; [KE 7] testing and debugging techniques, including "
               "techniques to avoid single point failures."),
        ("E4", "[ICTCLD504 PC 3.4] apply short-term refinements to the deployed resources according "
               "to test results."),
    ]),
    ("Part B — Finalise improvements (ICTCLD504 element 4)", [
        ("E5", "[ICTCLD504 PC 4.1] document the as-deployed architecture and test results, "
               "highlighting changes and improvements from the approved design; [PE 5] create "
               "documentation of deployment and testing steps; [FS Writing] writes technical data "
               "logically."),
        ("E6", "[ICTCLD504 PC 4.2] describe long-term improvement strategies and their benefits as "
               "applied to the deployed resources."),
        ("E7", "[ICTCLD504 PC 4.3] obtain final sign off from required personnel."),
    ]),
    ("Document quality (Foundation Skills)", [
        ("E8", "[ICTCLD504 PE 5] create documentation of deployment and testing steps; [FS Writing] "
               "writes and edits technical data logically; [FS Problem solving] uses formal and "
               "intuitive processes to identify issues, evaluate strategies and consider "
               "implementation contingencies; [FS Self-management] applies cloud-computing knowledge "
               "to troubleshoot — co-evidenced through the testing, troubleshooting and refinement."),
    ]),
]

STUDENT_INTRO = [
    ("The engagement and your role", "Heading 2"),
    ("YAT College has approved the improvement to the cloud infrastructure of its Ledgerline "
     "(Accounting) system, and your team has written the infrastructure-as-code for it. In AT3 you "
     "deploy it: you bring up the existing-state baseline and apply the approved improvement, then "
     "show that the improved system meets its goals.", "Assessor text"),
    ("You are an MP Tech Solutions (MTS) engineer. AT3 is your own individual work, in your own AWS "
     "Academy lab: you deploy and operate the whole system, demonstrate all four optimisation "
     "concerns on it, refine it from what your tests show, and document and finalise the result.", "Assessor text"),
    ("The baseline, the approved design and your team's validated implementation are provided. The "
     "system's data is not your concern — the database deploys empty; you are improving and proving "
     "the infrastructure.", "Assessor text"),
]

STUDENT_TASK = [
    ("Your task — deploy, demonstrate and finalise", "Heading 2"),
    ("In your own AWS Academy lab, produce a Deployment Report (using the YAT Deployment Report "
     "template) that shows you have:", "Assessor text"),
    ("• Deployed the existing-state baseline, then applied the approved improvement as a "
     "CloudFormation change-set to the running stack (preview the change-set, then apply it).", "Assessor text"),
    ("• Monitored and measured the deployed system against its performance metrics and business "
     "goals.", "Assessor text"),
    ("• Tested and demonstrated all four optimisation concerns on the whole system — security, "
     "reliability (an application-tier failover demonstration plus a database backup/restore and DR "
     "demonstration), scalability (a scale-out demonstration), and cost optimisation.", "Assessor text"),
    ("• Applied short-term refinements to the deployed resources based on what your tests showed.", "Assessor text"),
    ("• Documented the as-deployed architecture and test results (highlighting the changes from the "
     "approved design), described a long-term improvement strategy, and obtained final sign-off.", "Assessor text"),
    ("Submit the Deployment Report with the evidence (screenshots, logs, test output) of the "
     "deployment and the demonstrations.", "Assessor text"),
]

TIPS = [
    ("Tips for success", "Heading 2"),
    ("Deploy the whole system, demonstrate the whole system. You are assessed on the improved "
     "infrastructure end to end, across all four concerns — not just one part of it.", "Assessor text"),
    ("Demonstrate, don't just describe. Reliability and scalability need to be shown — trigger an "
     "application-tier failover (terminate an instance and watch the ASG recover across AZs), run a "
     "database restore / DR drill, and trigger a scale-out, capturing the evidence for each.", "Assessor text"),
    ("Refine from your results. Your tests will show something worth tuning — make the refinement and "
     "record why.", "Assessor text"),
    ("Tear down cleanly. Delete your stack and end the lab when you are done, to free the lab "
     "credits.", "Assessor text"),
]


def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details ----
    t_details = doc.tables[0]
    set_cell_content(t_details.rows[1].cells[1], DETAILS["qualification"])
    set_cell_content(t_details.rows[2].cells[1], DETAILS["units"])
    set_cell_content(t_details.rows[3].cells[1], DETAILS["task_title"])
    set_cell_content(t_details.rows[4].cells[1], DETAILS["task_number"])

    # ---- Table 1: Teacher/Assessor instructions ----
    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Resources required"), RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA_STATEMENT)
    cond_row = t_instr.add_row()
    set_cell_content(cond_row.cells[0], "Assessment Conditions & Setup Requirements")
    for r in cond_row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell_content(cond_row.cells[1], CONDITIONS)

    # ---- Table 2: Marking Guide ----
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)  # keep header rows
    add_section_row(t_mark, "Part A — Deploy and operate")
    for c in PART_A:
        add_criterion_row(t_mark, c)
    add_section_row(t_mark, "Part B — Finalise")
    for c in PART_B:
        add_criterion_row(t_mark, c)
    add_section_row(t_mark, "Document quality")
    for c in QUALITY:
        add_criterion_row(t_mark, c)

    # ---- Instructions to Student ----
    doc.add_paragraph("Instructions to Student", style="Heading 1")
    for text, style in (STUDENT_INTRO + STUDENT_TASK + TIPS):
        doc.add_paragraph(text, style=style)

    # ---- Benchmark (UoC traceability) ----
    doc.add_paragraph("Marking Benchmark — UoC traceability", style="Heading 1")
    doc.add_paragraph(
        "For each criterion, the unit-of-competency items it evidences. Every PC, PE, KE and FS "
        "item allocated to AT3 in the cluster assessment plan is covered below.", style="Assessor text")
    for part_title, rows in BENCHMARK:
        doc.add_paragraph(part_title, style="Heading 2")
        for cid, uoc in rows:
            p = doc.add_paragraph(style="Assessor text")
            run = p.add_run(f"{cid}  —  ")
            run.bold = True
            p.add_run(uoc)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT3/AT3-Implement-Assessor.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
