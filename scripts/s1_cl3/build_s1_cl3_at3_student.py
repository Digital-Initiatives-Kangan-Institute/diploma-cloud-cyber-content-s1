#!/usr/bin/env python3
"""Build the S1-CL3 AT3 STUDENT instrument (.docx) by populating the Kangan template.

The student-facing version of AT3 (Implement), derived from the assessor instrument
(build_s1_cl3_at3_assessor) so the shared content is single-sourced and cannot drift. Loads the
official Kangan 'Project Assessment - Student.docx' template. Student-facing content only: Details,
Student instructions, the assessment criteria (what they're marked against), and the detailed task
instructions for the deployment. The marking benchmark and its UoC traceability are NOT included.

Usage:  python scripts/s1_cl3/build_s1_cl3_at3_student.py [output.docx]
Default: S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT3/AT3-Implement-Student.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_s1_cl3_at3_assessor as a  # noqa: E402  (cluster content)
sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_tables import add_criterion_row, add_section_row, clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Student.docx")


def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details ----
    t_details = doc.tables[0]
    set_cell_content(t_details.rows[2].cells[1], a.DETAILS["qualification"])
    set_cell_content(t_details.rows[3].cells[1], a.DETAILS["units"])
    set_cell_content(t_details.rows[4].cells[1], a.DETAILS["task_title"])
    set_cell_content(t_details.rows[5].cells[1], a.DETAILS["task_number"])

    # ---- Table 1: Student instructions ----
    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), a.OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), a.TASKS)
    set_cell_content(find_instruction_row(t_instr, "Resources required"), a.RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), a.CRITERIA_STATEMENT)

    # ---- Table 2: Assessment criteria (no UoC/benchmark) ----
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)
    add_section_row(t_mark, "Part A — Deploy and operate")
    for c in a.PART_A:
        add_criterion_row(t_mark, c)
    add_section_row(t_mark, "Part B — Finalise")
    for c in a.PART_B:
        add_criterion_row(t_mark, c)
    add_section_row(t_mark, "Document quality")
    for c in a.QUALITY:
        add_criterion_row(t_mark, c)

    # ---- Detailed task instructions (student-facing prose; NO marking benchmark / UoC) ----
    doc.add_paragraph("Instructions to Student", style="Heading 1")
    for text, style in (a.STUDENT_INTRO + a.STUDENT_TASK + a.TIPS):
        doc.add_paragraph(text, style=style)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT3/AT3-Implement-Student.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
