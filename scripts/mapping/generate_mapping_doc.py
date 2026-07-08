#!/usr/bin/env python3
"""General, contract-driven Assessment Mapping document generator — one engine for all clusters.

Replaces the per-cluster build_s1_clN_mapping_docs.py generators, which diverged (CL1 hand-authored +
in-place; CL2 flat-list + prefix split; CL3 per-AT split). The mapping doc is a derived artefact, so
the engine is a pure function of the contract inputs (see docs/mapping-document-standard.md):

  rows        — parsed from the unit's source UoC (units_of_competency/<UNIT>_Complete_*.md)
  AT columns  — the AT assessor benchmarks, inverted PER AT (each criterion lands under the column of
                the benchmark it came from; the criterion-prefix scheme is then just readability)
  FS / AC     — closest-fit judgement (FS_MAP / AC_MAP), the only non-derived layer

A cluster is described by one entry in CLUSTERS below; its assessor BENCHMARKs and FS_MAP / AC_MAP are
sourced from the existing per-cluster modules, so the engine reproduces the committed docx exactly
(prove it with --check before --build). Per-cluster policy is two flags: n_ats, and whether FS/AC take
benchmark codes before the closest-fit map.

USAGE:
  python scripts/mapping/generate_mapping_doc.py --check {cl2|cl3|all}   # table content vs committed
  python scripts/mapping/generate_mapping_doc.py --build {cl2|cl3|all}   # (re)write the docx
"""
import argparse
import copy
import importlib
import re
import sys
import tempfile
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from docx import Document  # noqa: E402

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
AT_LABELS = ("AT1", "AT2", "AT3")


# ============================================================================
# Course registry — the CLUSTERS registry, the qualification metadata, and the unit-code prefixes come
# from the CONTENT REPO's mapping_registry.py (this engine is generic and course-agnostic). Resolve it
# via --registry <dir|file>; otherwise the engine's sibling scripts/ (works when the engine sits inside
# the content repo). REPO (the content repo root) is derived from the registry's own location, so the
# engine works wherever it lives.
# ============================================================================

def _registry_dir(argv):
    for i, a in enumerate(argv):
        if a == "--registry" and i + 1 < len(argv):
            p = Path(argv[i + 1]).resolve()
            return p.parent if p.is_file() else p
    return Path(__file__).resolve().parents[1]  # scripts/ (sibling of mapping/)


sys.path.insert(0, str(_registry_dir(sys.argv)))
import mapping_registry as _reg  # noqa: E402

CLUSTERS = _reg.CLUSTERS
QUALIFICATION = _reg.QUALIFICATION
UNIT_PREFIXES = _reg.UNIT_PREFIXES
REPO = Path(_reg.__file__).resolve().parents[1]  # content repo root (registry at <repo>/scripts/)
_UNIT_RE = r"((?:" + "|".join(UNIT_PREFIXES) + r")\d{3})\s+(.*)"


def load_config(key):
    spec = CLUSTERS[key]
    sys.path.insert(0, str(REPO / "scripts" / spec["pkg"]))
    bmod = importlib.import_module(spec["build_mod"])
    uoc_dir = REPO / spec["cluster_dir"] / "units_of_competency"
    if spec.get("source") == "data":
        return _data_config(key, spec, bmod, uoc_dir)
    assessors = [importlib.import_module(m) for m in spec["assessors"]]
    units = []
    for code, title, *_ in bmod.UNITS.values():
        units.append(dict(code=code, title=title,
                          uoc_md=next(uoc_dir.glob(f"{code}_Complete_*.md"))))
    return dict(
        key=key, n_ats=spec["n_ats"],
        template=Path(bmod.TEMPLATE), mappings_dir=Path(bmod.MAPPINGS),
        at_titles=[getattr(bmod, f"AT{i + 1}_TITLE") for i in range(spec["n_ats"])],
        at_benchmarks=[m.BENCHMARK for m in assessors],
        fs_map=bmod.FS_MAP, ac_map=bmod.AC_MAP,
        fs_source=spec["fs_source"], ac_source=spec["ac_source"],
        units=units,
    )


def _split(codes):
    return [c.strip() for c in str(codes).split(",") if c.strip()]


def _data_config(key, spec, bmod, uoc_dir):
    """Config for a legacy DATA_*-driven cluster: synthesise the per-AT inversion from the
    hand-authored item->{AT:codes} dicts (PC/PE/KE/FS/AC), so the engine renders it like any other."""
    direct_inv, units = {}, []

    def put(code, sec, item, ats):
        for at in AT_LABELS:
            v = ats.get(at, "")
            if v:
                d = direct_inv.setdefault((code, sec, str(item)), {a: [] for a in AT_LABELS})
                d[at] = _split(v)

    for fname, data in bmod.UNIT_DATA.values():
        code = fname.split("_Assessment_Mapping")[0]
        units.append(dict(code=code, title=spec["unit_titles"][code],
                          uoc_md=next(uoc_dir.glob(f"{code}_Complete_*.md"))))
        for num, ats in data.get("pcs", {}).items():
            put(code, "PC", num, ats)
        pe_entries = data.get("pes") or [e[1] for e in data.get("pes_split", [])]
        for i, ats in enumerate(pe_entries, 1):
            put(code, "PE", i, ats)
        for i, ats in enumerate(data.get("kes", []), 1):
            put(code, "KE", i, ats)
        for skill, ats in data.get("fss", {}).items():
            put(code, "FS", skill, ats)
        for i, (_label, ats) in enumerate(data.get("acs", []), 1):
            put(code, "AC", i, ats)

    return dict(
        key=key, n_ats=spec["n_ats"],
        template=Path(getattr(bmod, "TEMPLATE", REPO / "kangan-templates" / "Assessment Mapping Tool.docx")),
        mappings_dir=Path(bmod.MAPPINGS),
        at_titles=spec["at_titles"], at_benchmarks=None, direct_inv=direct_inv,
        fs_map={}, ac_map={}, fs_source=spec["fs_source"], ac_source=spec["ac_source"],
        units=units,
    )


# ============================================================================
# Source UoC parser
# ============================================================================

def _sections(text):
    out, cur, buf = {}, None, []
    for line in text.splitlines():
        m = re.match(r"^#\s+(.*)", line)
        if m:
            if cur is not None:
                out.setdefault(cur, "\n".join(buf))
            cur, buf = m.group(1).strip(), []
        else:
            buf.append(line)
    if cur is not None:
        out.setdefault(cur, "\n".join(buf))
    return out


def _bullets(body):
    raw = body.splitlines()
    top = [ln for ln in raw if re.match(r"^- ", ln)]
    # Parent-bullet special case (e.g. ICTICT517 PE "For one organisation:" + 6 sub-bullets):
    # a single top-level bullet ending ':' means its sub-bullets are the assessable items.
    if len(top) == 1 and top[0].rstrip().endswith(":"):
        return [re.sub(r"^\s*-\s+", "", ln).strip() for ln in raw if re.match(r"^\s+- ", ln)]
    items = []
    for ln in raw:
        s = ln.strip()
        if not s.startswith("- "):
            continue
        text = re.sub(r"^-\s+", "", s).strip()
        if (len(ln) - len(ln.lstrip())) > 0 and items:
            items[-1] = items[-1] + "\n  - " + text
        else:
            items.append(text)
    return items


def parse_source_unit(md_path):
    secs = _sections(Path(md_path).read_text(encoding="utf-8"))
    pcs = []
    for ln in secs.get("Elements and Performance Criteria", "").splitlines():
        if not ln.strip().startswith("|"):
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if len(cells) < 2:
            continue
        elem, pc_cell = cells[0], cells[1]
        if elem.upper().startswith("ELEMENTS") or elem.startswith("Elements describe"):
            continue
        if not re.match(r"^\d+\.", elem):
            continue
        for piece in pc_cell.split("<br>"):
            m = re.match(r"^(\d+\.\d+)\s+(.*)", piece.strip())
            if m:
                pcs.append((m.group(1), m.group(2).strip(), elem))
    fss = []
    for ln in secs.get("Foundation Skills", "").splitlines():
        if not ln.strip().startswith("|"):
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if len(cells) < 2:
            continue
        skill, desc = cells[0], cells[1]
        if skill.upper() == "SKILL" or not skill or skill.startswith("---"):
            continue
        fss.append((skill, desc.replace("<br>", " ")))
    acsec = secs.get("Assessment Conditions", "")
    acs = _bullets(acsec)
    for ln in acsec.splitlines():
        if ln.strip().lower().startswith("assessors of this unit must satisfy"):
            acs.append(ln.strip())
    return {"pcs": pcs, "fss": fss,
            "pes": _bullets(secs.get("Performance Evidence", "")),
            "kes": _bullets(secs.get("Knowledge Evidence", "")),
            "acs": acs}


# ============================================================================
# Benchmark inverter — (unit, section, item) -> {AT1: [...], AT2: [...], AT3: [...]}
# ============================================================================

def _expand(section, numbering):
    out = []
    for part in numbering.split(","):
        part = part.strip().replace("–", "-").replace("—", "-")
        if not part:
            continue
        if "-" in part and section != "FS":
            a, b = (x.strip() for x in part.split("-", 1))
            if section == "PC" and "." in a and "." in b:
                ea, ia = a.split("."); eb, ib = b.split(".")
                if ea == eb:
                    out += [f"{ea}.{i}" for i in range(int(ia), int(ib) + 1)]
                else:
                    out += [a, b]
            else:
                try:
                    out += [str(i) for i in range(int(a), int(b) + 1)]
                except ValueError:
                    out += [a, b]
        else:
            out.append(part)
    return out


def _parse_tags(s):
    cur_unit = None
    for m in re.finditer(r"\[([^\]]+)\]", s):
        inner = m.group(1).strip()
        um = re.match(_UNIT_RE, inner)
        if um:
            cur_unit, rest = um.group(1), um.group(2)
        else:
            rest = inner
        sm = re.match(r"(PC|PE|KE|AC|FS)\s+(.*)", rest)
        if not sm or cur_unit is None:
            continue
        section, numbering = sm.group(1), sm.group(2).strip()
        if section == "FS":
            yield (cur_unit, "FS", numbering)
        else:
            for item in _expand(section, numbering):
                yield (cur_unit, section, item)


def invert(at_benchmarks):
    """Per-AT inversion: each criterion lands under the column of the benchmark it came from."""
    out = {}
    for idx, bench in enumerate(at_benchmarks):
        at = AT_LABELS[idx]
        for _part_title, rows in bench:
            for cid, uoc in rows:
                for tag in _parse_tags(uoc):
                    d = out.setdefault(tag, {"AT1": [], "AT2": [], "AT3": []})
                    if cid not in d[at]:
                        d[at].append(cid)
    return out


def _codes(inv, unit, section, item, at):
    return ", ".join(inv.get((unit, section, item), {}).get(at, []))


# ============================================================================
# docx fill (verbatim-shared mechanics)
# ============================================================================

T_DETAILS, T_DESC, T_AC, T_PC, T_PE, T_KE, T_FS = 1, 2, 3, 4, 5, 6, 7
AT_COLS = {"AC": (1, 2, 3), "PC": (2, 3, 4), "PE": (1, 2, 3), "KE": (1, 2, 3), "FS": (3, 4, 5)}
CLEAR_COLS = {"AC": (4, 5), "PC": (5, 6), "PE": (4, 5), "KE": (4, 5), "FS": (6, 7)}


def _put(cell, text):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r.text = ""
    lines = (text or "").split("\n")
    (p0.runs[0] if p0.runs else p0.add_run("")).text = lines[0]
    for ln in lines[1:]:
        cell.add_paragraph().add_run(ln)


def _reset(table, n_header=2):
    rows = list(table.rows)
    template_tr = copy.deepcopy(rows[n_header]._tr)
    for row in rows[n_header:]:
        row._tr.getparent().remove(row._tr)
    return template_tr


def _add(table, template_tr):
    table._tbl.append(copy.deepcopy(template_tr))
    return table.rows[-1]


def _fill_ats(row, section, codes3):
    for col, val in zip(AT_COLS[section], codes3):
        _put(row.cells[col], val)
    for col in CLEAR_COLS[section]:
        _put(row.cells[col], "")


# ============================================================================
# Resolve per-item AT codes, honouring the FS/AC source policy
# ============================================================================

def _at3(codes3, n_ats):
    """Blank the AT3 column for a 2-AT cluster."""
    return codes3 if n_ats >= 3 else [codes3[0], codes3[1], ""]


def resolve_bench(inv, unit, section, item, n_ats):
    return _at3([_codes(inv, unit, section, item, at) for at in AT_LABELS], n_ats)


def resolve_mapped(inv, unit, section, item, map_entry, source, n_ats):
    bench = [_codes(inv, unit, section, item, at) for at in AT_LABELS]
    if source == "benchmark_then_map":
        merged = [bench[i] or map_entry.get(AT_LABELS[i], "") for i in range(3)]
    else:  # map_only
        merged = [map_entry.get(AT_LABELS[i], "") for i in range(3)]
    return _at3(merged, n_ats)


# ============================================================================
# Build
# ============================================================================

def build_unit(cfg, unit, out_path):
    items = parse_source_unit(unit["uoc_md"])
    inv = cfg.get("direct_inv") or invert(cfg["at_benchmarks"])
    code, n_ats = unit["code"], cfg["n_ats"]
    doc = Document(str(cfg["template"]))

    det = doc.tables[T_DETAILS]
    _put(det.rows[1].cells[1], code)
    _put(det.rows[1].cells[3], unit["title"])
    _put(det.rows[1].cells[5], "1")
    _put(det.rows[2].cells[1], QUALIFICATION["code"])
    _put(det.rows[2].cells[3], QUALIFICATION["title"])

    desc = doc.tables[T_DESC]
    for i, title in enumerate(cfg["at_titles"]):
        _put(desc.rows[2 + i].cells[1], title)
        _put(desc.rows[2 + i].cells[2], "1.0")

    pc_tbl = doc.tables[T_PC]
    tmpl = _reset(pc_tbl)
    last_elem = None
    for num, txt, elem in items["pcs"]:
        row = _add(pc_tbl, tmpl)
        _put(row.cells[0], elem if elem != last_elem else "")
        last_elem = elem
        _put(row.cells[1], f"{num} {txt}")
        _fill_ats(row, "PC", resolve_bench(inv, code, "PC", num, n_ats))

    for tnum, key, sect in ((T_PE, "pes", "PE"), (T_KE, "kes", "KE")):
        tbl = doc.tables[tnum]
        tmpl = _reset(tbl)
        for i, txt in enumerate(items[key], 1):
            row = _add(tbl, tmpl)
            _put(row.cells[0], txt)
            _fill_ats(row, sect, resolve_bench(inv, code, sect, str(i), n_ats))

    fs_tbl = doc.tables[T_FS]
    tmpl = _reset(fs_tbl)
    fs_map = cfg["fs_map"].get(code, {})
    for skill, fdesc in items["fss"]:
        row = _add(fs_tbl, tmpl)
        _put(row.cells[0], skill)
        _put(row.cells[1], "")
        _put(row.cells[2], fdesc)
        _fill_ats(row, "FS", resolve_mapped(inv, code, "FS", skill,
                                            fs_map.get(skill, {}), cfg["fs_source"], n_ats))

    ac_tbl = doc.tables[T_AC]
    tmpl = _reset(ac_tbl)
    ac_map = cfg["ac_map"].get(code, [])
    for i, txt in enumerate(items["acs"], 1):
        row = _add(ac_tbl, tmpl)
        _put(row.cells[0], txt)
        entry = ac_map[i - 1] if i - 1 < len(ac_map) else {}
        _fill_ats(row, "AC", resolve_mapped(inv, code, "AC", str(i), entry, cfg["ac_source"], n_ats))

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# ============================================================================
# Reproduction check — table cell content of generated vs committed
# ============================================================================

def table_content(docx_path):
    with zipfile.ZipFile(docx_path) as z:
        tree = ET.parse(z.open("word/document.xml"))
    body = tree.getroot().find(f"{W}body")
    out = []
    for tbl in body.iter(f"{W}tbl"):
        rows = []
        for tr in tbl.findall(f"{W}tr"):
            cells = []
            for tc in tr.findall(f"{W}tc"):
                paras = ["".join(n.text or "" for n in p.iter() if n.tag == f"{W}t")
                         for p in tc.findall(f"{W}p")]
                cells.append("\n".join(paras))
            rows.append(cells)
        out.append(rows)
    return out


def check_unit(cfg, unit):
    committed = cfg["mappings_dir"] / f"{unit['code']}_Assessment_Mapping.docx"
    if not committed.exists():
        print(f"  {unit['code']}: no committed docx to compare ({committed.name})")
        return False
    with tempfile.TemporaryDirectory() as td:
        gen = build_unit(cfg, unit, Path(td) / "gen.docx")
        a, b = table_content(committed), table_content(gen)
    if a == b:
        print(f"  {unit['code']}: IDENTICAL table content ({len(a)} tables)")
        return True
    diffs = 0
    print(f"  {unit['code']}: DIFFERS")
    for ti, (ta, tb) in enumerate(zip(a, b)):
        for ri, (ra, rb) in enumerate(zip(ta, tb)):
            if ra != rb:
                diffs += 1
                if diffs <= 12:
                    print(f"    table {ti} row {ri}: committed={ra}")
                    print(f"    table {ti} row {ri}: generated={rb}")
        if len(ta) != len(tb):
            print(f"    table {ti}: row count committed={len(ta)} generated={len(tb)}")
    if len(a) != len(b):
        print(f"    table count committed={len(a)} generated={len(b)}")
    return False


def build_cluster(key, only=None):
    """Build a cluster's mapping docs to their committed locations. `only` filters by unit code.
    Used by the thin per-cluster wrappers; returns the list of written paths."""
    cfg = load_config(key)
    written = []
    for unit in cfg["units"]:
        if only and unit["code"] not in only:
            continue
        out = build_unit(cfg, unit, cfg["mappings_dir"] / f"{unit['code']}_Assessment_Mapping.docx")
        written.append(out)
    return written


def main():
    try:
        sys.stdout.reconfigure(encoding="utf-8")  # AC tick (✓) + en-dashes survive a cp1252 console
    except (AttributeError, ValueError):
        pass
    ap = argparse.ArgumentParser(description="General Assessment Mapping document generator.")
    ap.add_argument("--check", metavar="CLUSTER", help="cl2 | cl3 | all")
    ap.add_argument("--build", metavar="CLUSTER", help="cl2 | cl3 | all")
    ap.add_argument("--registry", metavar="PATH",
                    help="course mapping_registry.py (dir or file); default: the engine's sibling scripts/")
    args = ap.parse_args()
    mode = "check" if args.check else "build"
    target = args.check or args.build
    if not target:
        ap.error("give --check or --build {cl2|cl3|all}")
    keys = list(CLUSTERS) if target == "all" else [target]

    ok = True
    for key in keys:
        cfg = load_config(key)
        print(f"\n=== {key} ({mode}) ===")
        for unit in cfg["units"]:
            if mode == "check":
                ok &= check_unit(cfg, unit)
            else:
                out = build_unit(cfg, unit, cfg["mappings_dir"] / f"{unit['code']}_Assessment_Mapping.docx")
                print(f"  wrote {out}")
    if mode == "check":
        print("\nRESULT:", "PASS — engine reproduces committed docx" if ok else "FAIL — see diffs")
        sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()
