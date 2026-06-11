#!/usr/bin/env python3
"""Build the S1-CL3 AT1 STUDENT instrument (.docx) by populating the Kangan template.

The student-facing version of AT1 (Team Setup), derived from the assessor instrument
(build_cl3_at1_assessor) so the shared content is single-sourced and cannot drift. Loads the
official Kangan 'Project Assessment - Student.docx' template and fills it in, keeping the Kangan
structure and styles. Student-facing content only: Details, Student instructions, the assessment
criteria (what they're marked against), and the detailed task instructions for the Team Plan.
The assessor-only material — the marking benchmark and its UoC traceability — is NOT included.

Usage:  python scripts/build_cl3_at1_student.py [output.docx]
Default: S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT1/AT1-Team-Setup-Student.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_cl3_at1_assessor as a  # noqa: E402  (shared helpers + content — single source of truth)

from docx import Document  # noqa: E402

TEMPLATE = "templates/Project Assessment - Student.docx"


def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details (student name / ID / assessor name+email left blank to fill) ----
    t_details = doc.tables[0]
    a.set_cell(t_details.rows[2].cells[1], a.DETAILS["qualification"])
    a.set_cell(t_details.rows[3].cells[1], a.DETAILS["units"])
    a.set_cell(t_details.rows[4].cells[1], a.DETAILS["task_title"])
    a.set_cell(t_details.rows[5].cells[1], a.DETAILS["task_number"])

    # ---- Table 1: Student instructions ----
    t_instr = doc.tables[1]
    a.set_cell(a.instr_row(t_instr, "Assessment overview"), a.OVERVIEW)
    a.set_cell(a.instr_row(t_instr, "Task"), a.TASKS)
    a.set_cell(a.instr_row(t_instr, "Resources required"), a.RESOURCES)
    a.set_cell(a.instr_row(t_instr, "Assessment criteria"), a.CRITERIA_STATEMENT)
    # Time allowed, Location, Results/Second attempt, Important information left as template text.

    # ---- Table 2: Assessment criteria (what the student is marked against — no UoC/benchmark) ----
    t_mark = doc.tables[2]
    a.clear_table_rows(t_mark, 2)  # keep the two header rows
    a.add_section_row(t_mark, "Team Plan")
    for c in a.TEAM_PLAN:
        a.add_criterion_row(t_mark, c)

    # ---- Detailed task instructions (student-facing prose — same as the assessor doc's
    #      'Instructions to Student' section; NO marking benchmark / UoC) ----
    doc.add_paragraph("Instructions to Student", style="Heading 1")
    for text, style in (a.STUDENT_INTRO + a.STUDENT_TASK + a.TIPS):
        doc.add_paragraph(text, style=style)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL3-Cloud-Infrastructure-Improvement/assessments/AT1/AT1-Team-Setup-Student.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
