#!/usr/bin/env python3
"""Build the Website Baseline Solution Design (.docx) — in-world document.

The as-built cloud architecture of YAT's 2023 Website Cloud Migration — YAT's first
cloud project, a deliberately low-risk pilot. A simple single-Availability-Zone Sydney
deployment (single EC2 LAMP instance + single-AZ RDS MySQL + S3 backups), with NO load
balancer, NO autoscaling, NO Multi-AZ and NO disaster recovery — those are stated as
accepted residual risks and explicitly out of scope. The non-HA sections (§4 HA review,
§4.x SPOFs) record the baseline's gaps rather than removing them.

This is the website baseline of record: the state CL3 improves from (team-led improvement)
and the state CL2 DR-planning practises against (as a supplied snapshot). An in-world
supplied artefact (no UoC tags) — the branded Solution Design document.

Output to the website documents folder for printing to PDF, then linked from the
website-cloud-migration project's Cloud Architecture Baseline page.

Usage:  python scripts/build_website_baseline_solution_design.py [output.docx]
Default: ../diploma-cloud-cyber-website/public/documents/YAT-Website-Baseline-Solution-Design.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_bullet_list  # noqa: E402
import build_bc_template as bc   # noqa: E402
import build_s1_cl1_at1_bc_exemplar as ex   # noqa: E402  (etable, para, bullets)

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def na(doc, reason):
    p = doc.add_paragraph()
    r = p.add_run(f"Not applicable — {reason}")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    p.paragraph_format.space_after = Pt(6)
    return p


def diagram_placeholder(doc, caption, source):
    """A bordered, shaded drop-zone for a network diagram — replace with the exported image."""
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    bc.set_cell_borders(cell)
    bc.shade_cell(cell, bc.CREAM)
    cell.width = Cm(16.6)
    t.rows[0].height = Cm(9.5)
    t.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[ NETWORK TOPOLOGY DIAGRAM — PASTE HERE ]")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(source)
    r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = RGBColor.from_string(bc.GREY)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = RGBColor.from_string(bc.GREY)
    return t


def build(path):
    doc = Document()
    bc.configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    bc.build_header_footer(sec)

    # ---- COVER ----
    bc.wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(bc.ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(bc.GREY)
    bc.paragraph_bottom_rule(doc.add_paragraph(), bc.TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Solution Design")
    sub = doc.add_paragraph().add_run("YAT Website Cloud Architecture — Baseline Design")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT Website Cloud Migration (2023) — YAT's first cloud project"),
        ("Document type", "Technical design (Solution Design)"),
        ("Version", "v1.0 — As-built baseline"),
        ("Authored by", "MTS Consulting, in consultation with YAT ICT"),
        ("Approved by", "Pat Lin (MTS Senior Consultant) · Sam Walker (YAT ICT Manager)"),
        ("Status", "Closed engagement — retained as the website baseline of record"),
        ("Classification", "Internal — YAT ICT, and MTS personnel on signed MSA"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        bc.set_cell_borders(cells[0]); bc.set_cell_borders(cells[1]); bc.shade_cell(cells[0], bc.CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    bc.add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Purpose and Scope")
    ex.para(doc, "This document records the baseline AWS architecture MTS implemented for the YAT Website "
                 "Cloud Migration — YAT's first move to the cloud. The website (marketing site, course "
                 "catalogue, and online enquiry / application intake) and its CMS were re-hosted from a single "
                 "on-premises web server onto a simple, cost-minimal cloud deployment in Sydney. The engagement "
                 "was a deliberately low-risk pilot: the goal was a working, supportable cloud deployment at "
                 "minimum cost and complexity — not a highly available one.")
    h3("In scope of this design")
    add_bullet_list(doc, [
        "The cloud hosting foundation for the YAT public website and its CMS.",
        "Compute, networking, identity, storage, database and a basic monitoring baseline sufficient to run the website as a single-instance LAMP workload in AWS.",
        "Single-region, single-Availability-Zone deployment in ap-southeast-2 (Sydney).",
        "Migration of the existing website content, CMS database, and uploaded media; production cutover after the enrolment-enquiry peak.",
    ])
    h3("Out of scope — deferred (accepted pilot limitations)")
    add_bullet_list(doc, [
        "High availability of any kind: no load balancer, no autoscaling, no Multi-AZ database, no cross-AZ resilience.",
        "Disaster recovery: no cross-region copies, no DR runbook, no tested recovery objectives.",
        "Content delivery / edge caching (CDN) and media offload to object storage.",
        "Web application firewall and advanced security hardening beyond the baseline.",
    ])
    h3("Out of MTS scope entirely — YAT responsibility")
    add_bullet_list(doc, [
        "Website content authoring, information architecture, and visual redesign (YAT Marketing — the site was migrated as-is).",
        "Ongoing operation and support of the website after handover (YAT ICT).",
        "Domain-registrar account administration beyond the cutover DNS change (YAT ICT).",
    ])

    h1("2. Design Inputs and Requirements")
    h3("2.1 Inputs")
    add_bullet_list(doc, [
        "Website Migration Requirements — the requirements the migration had to meet.",
        "Engagement Role Brief — engagement scope; re-host (no redesign) and content preservation.",
        "ICT Manager Consultation Notes — the requirements consultation (drivers, availability stance, integrations, cutover constraints).",
        "ICT Environment Overview — the on-premises current state being migrated from.",
    ])
    h3("2.2 Requirements the design must meet")
    ex.etable(doc, ["Requirement", "Target / note"],
              [["Region / data residency", "ap-southeast-2 (Sydney); enquiry/application PII within Australia (Privacy Act, APP 8)"],
               ["Application stack", "Preserved — existing PHP/MySQL CMS on a LAMP stack, migrated as-is"],
               ["Availability", "No formal SLA for the pilot; single-AZ acceptable (a known, accepted limitation)"],
               ["Recovery", "Routine database backups for basic protection; no committed RTO/RPO"],
               ["Load profile", "Light most of the year; predictable Jan–Feb enrolment-enquiry peak"],
               ["Integrations", "Enquiry/application form → email to Admissions + forward to student-admin intake; HTTPS"],
               ["Cost", "Smallest sensible footprint; no year-round redundant capacity"]],
              widths=[5.0, 11.0])

    h1("3. Review of the Existing (On-Premises) Architecture")
    ex.para(doc, "Unlike a greenfield build, this engagement migrated a running system. The source was a single "
                 "on-premises web server at the Cremorne campus hosting the CMS (PHP/MySQL on a LAMP stack), its "
                 "MySQL database, and uploaded media on local disk, behind the campus internet connection. It was "
                 "a single point of failure with no redundancy, strained during the enrolment-enquiry peak, and "
                 "approaching end of warranty. The cloud baseline below re-hosts this same stack as-is; it "
                 "improves the hardware and elasticity position but deliberately reproduces the single-instance "
                 "shape (now in AWS) rather than re-architecting for availability.")

    h1("4. Architecture Design")
    h3("4.1 Assumptions and constraints")
    ex.etable(doc, ["#", "Assumption / constraint", "Source"],
              [["A1", "Region must be ap-southeast-2 (Sydney) for data residency", "Migration Requirements; Privacy Policy"],
               ["A2", "CMS application stack preserved: PHP/MySQL on a LAMP stack, migrated as-is", "Engagement Role Brief"],
               ["A3", "Light load most of the year, with a Jan–Feb enrolment-enquiry peak", "Consultation Notes"],
               ["A4", "Pilot is cost-minimal: smallest sensible single-instance footprint", "Consultation Notes"],
               ["A5", "Enquiry/application PII must remain within Australia", "Privacy Act 1988, APP 8"],
               ["A6", "Baseline is NOT high-availability and has NO disaster recovery", "Scoping decision (accepted residual risk)"]],
              widths=[1.0, 9.0, 6.0])
    h3("4.2 AWS account and region")
    ex.para(doc, "An AWS account scoped to the engagement, in region ap-southeast-2 (Sydney); no deployment "
                 "outside Australian regions. A single Availability Zone (ap-southeast-2a) hosts the entire "
                 "baseline — there is no second AZ.")
    h3("4.3 Identity and Access Management (IAM)")
    ex.etable(doc, ["Group / role", "Purpose", "Indicative permissions"],
              [["YAT-ICT-Admins", "YAT ICT day-to-day ops post-handover", "Console access to EC2/RDS/S3/CloudWatch; no IAM changes"],
               ["MTS-Consultants", "MTS during build + cutover", "Admin during build; removed at handover"],
               ["Website-Instance role", "EC2 instance profile for the website", "S3 read/write (backups); CloudWatch logs/metrics"]],
              widths=[3.5, 5.0, 7.5])
    add_bullet_list(doc, [
        "MFA required for all human users (User Access Policy; ACSC Essential Eight).",
        "No long-lived access keys for humans; EC2 access to AWS via the instance profile.",
        "Configuration decision left to the implementer: the administrative SSH access path to the instance.",
    ])
    h3("4.4 Network topology")
    ex.para(doc, "VPC 10.0.0.0/16, with DNS hostnames/resolution and VPC flow logs enabled. Single-AZ subnets "
                 "in ap-southeast-2a:")
    ex.etable(doc, ["Subnet", "CIDR", "Tier", "Internet-facing?"],
              [["public-web-a", "10.0.1.0/24", "Website EC2 (CMS)", "Yes"],
               ["private-data-a", "10.0.21.0/24", "Database (RDS)", "No"]],
              widths=[4.0, 3.5, 5.5, 3.0])
    add_bullet_list(doc, [
        "Internet Gateway for the public subnet; the EC2 instance has an Elastic IP and serves traffic directly (no load balancer in the baseline).",
        "Route tables: public-web → IGW; private-data → no internet route.",
        "There is no second AZ and no -b subnets — the single-AZ shape is the defining property of the baseline.",
    ])
    diagram_placeholder(doc,
                        "Figure 4.4 — Website baseline network topology (single-AZ, single instance).",
                        "Source diagram: website-baseline-single-az.drawio")
    h3("4.5 Compute (single EC2 instance)")
    add_bullet_list(doc, [
        "A single EC2 instance (e.g. t3.small/medium — final type an implementer decision) running the LAMP stack and the CMS, in public-web-a with an Elastic IP.",
        "No Auto Scaling Group and no load balancer: the instance is the website. This is the principal single point of failure, accepted for the pilot.",
        "EBS: gp3 root volume holding the OS, the CMS application, and uploaded media (the CMS stores uploads on local disk by default).",
    ])
    h3("4.6 TLS / HTTPS")
    add_bullet_list(doc, [
        "HTTPS terminates on the EC2 instance itself (there is no load balancer or CDN to terminate TLS at the edge).",
        "The certificate mechanism (e.g. an automatically-renewed certificate on the instance vs an imported certificate) is an implementer decision.",
    ])
    h3("4.7 Database (RDS)")
    add_bullet_list(doc, [
        "Amazon RDS for MySQL (preserves the existing CMS schema and data); engine version confirmed against the CMS at build time.",
        "A small general-purpose instance class (e.g. db.t3.small — an implementer decision); gp3 storage sized to footprint + growth.",
        "Multi-AZ DISABLED (single-AZ baseline); storage encryption enabled (KMS).",
        "Placed in private-data-a; not publicly accessible; 7-day automated backups; maintenance in an overnight window.",
    ])
    h3("4.8 Storage")
    ex.etable(doc, ["Resource", "Type", "Purpose"],
              [["EC2 root volume", "EBS gp3", "OS, CMS application, and uploaded media"],
               ["Backups bucket", "S3", "Nightly database dumps and media snapshots"]],
              widths=[4.0, 4.5, 7.5])
    add_bullet_list(doc, [
        "Backups bucket: block all public access; SSE-S3 encryption; versioning enabled.",
        "Media is served from the instance, not from object storage or a CDN — an accepted pilot simplification (and a candidate first improvement later).",
    ])
    h3("4.9 Security")
    ex.etable(doc, ["Security group", "Inbound", "Outbound"],
              [["sg-web", "HTTP:80 + HTTPS:443 from 0.0.0.0/0; SSH:22 from an admin source (implementer decision)", "MySQL:3306 to sg-db; HTTPS to AWS APIs / updates"],
               ["sg-db", "MySQL:3306 from sg-web only", "none"]],
              widths=[3.0, 8.0, 5.0])
    add_bullet_list(doc, [
        "Encryption in transit: HTTPS to visitors; TLS from EC2 to RDS.",
        "Encryption at rest: EBS, RDS, and S3 all enabled.",
        "No web application firewall in the baseline (a deferred improvement).",
        "Operates under the AWS Shared Responsibility Model — AWS secures the cloud; YAT/MTS secure the OS, application, IAM, data and access in the cloud.",
    ])
    h3("4.10 Monitoring (baseline)")
    ex.para(doc, "Standard CloudWatch metrics for the EC2 instance and the RDS database, with a small set of "
                 "baseline alarms:")
    ex.etable(doc, ["Alarm", "Threshold"],
              [["EC2 instance status check failed", "Any failed check"],
               ["EC2 CPU high", "≥ 80% over 10 min"],
               ["RDS CPU high", "≥ 80% over 10 min"],
               ["RDS free storage low", "< 15%"]],
              widths=[8.0, 8.0])
    add_bullet_list(doc, ["Logging: web/CMS and RDS logs → CloudWatch Logs; VPC flow logs enabled. No cross-AZ or replica-lag alarms (there is no second AZ or replica)."])
    h3("4.11 Naming and tagging conventions")
    ex.para(doc, "Naming pattern yat-web-<resource-type>-<env> (e.g. yat-web-ec2-prod). Mandatory tags:")
    ex.etable(doc, ["Tag", "Value"],
              [["Project", "YAT-Website-Migration"], ["Environment", "Production"], ["Owner", "YAT-ICT"],
               ["ManagedBy", "MTS during build → YAT-ICT post-handover"],
               ["CostCentre", "YAT-Website"], ["DataClassification", "Internal; PII (enquiry/application data)"]],
              widths=[5.0, 11.0])
    h3("4.12 Backup")
    ex.etable(doc, ["Resource", "Mechanism", "Retention"],
              [["RDS database", "Automated daily backups + transaction logs", "7 days"],
               ["EC2 EBS volume", "Daily AMI/EBS snapshot (Data Lifecycle Manager)", "7 days"],
               ["Database / media", "Nightly dump + media snapshot to the S3 backups bucket", "Versioned"]],
              widths=[4.5, 7.5, 4.0])
    add_bullet_list(doc, ["Cross-Region backup copies are out of scope for the baseline — there is no disaster-recovery tier."])
    h3("4.13 Recovery objectives — baseline state")
    ex.para(doc, "The baseline's recovery posture is backup-based and unmeasured: recovery means rebuilding the "
                 "single instance and restoring from the most recent backup. No RTO or RPO was committed for the "
                 "pilot. There is no second-AZ or second-region fallback. This is the principal gap a future "
                 "engagement (and any disaster-recovery plan) would address.")
    h3("4.14 Single points of failure (known and accepted)")
    ex.para(doc, "The baseline is single-AZ and single-instance by design. The following are known single points "
                 "of failure, accepted as residual risk for the pilot and recorded for a future engagement:")
    add_bullet_list(doc, [
        "The single EC2 instance — its failure takes the whole website offline.",
        "The single Availability Zone — an AZ-level event takes everything down.",
        "The single RDS instance — no standby; failure or maintenance interrupts the site.",
        "Uploaded media on the instance's EBS volume — not redundant or offloaded.",
    ])
    h3("4.15 Configuration decisions left to the implementer")
    ex.etable(doc, ["#", "Decision", "Why left open"],
              [["C1", "EC2 instance type", "Size against the website's light load + peak"],
               ["C2", "RDS instance class + storage size", "Size from the CMS data footprint + growth"],
               ["C3", "TLS certificate mechanism on the instance", "Implementer's operational preference"],
               ["C4", "Administrative SSH access path", "Adapt to YAT ICT's admin arrangements"],
               ["C5", "CMS / MySQL engine versions", "Confirm against the migrated CMS"]],
              widths=[1.0, 7.0, 8.0])

    h1("5. Implementation Sequencing")
    ex.para(doc, "Because this migrated a live website, the build was sequenced around a cutover timed to avoid "
                 "the enrolment-enquiry peak:")
    add_bullet_list(doc, [
        "Provision the foundation: account, VPC, single-AZ subnets, security groups, IAM.",
        "Stand up the EC2 instance and the RDS database; install the LAMP stack and the CMS.",
        "Migrate content: export the on-premises CMS database and import to RDS; copy uploaded media and assets to the instance.",
        "Verify on the cloud deployment pre-cutover: content complete, course catalogue correct, enquiry/application form submits and notifies.",
        "Cut over after the enrolment-enquiry peak: switch DNS to the Elastic IP; ramp down the on-premises server.",
        "Stabilise, then hand over documentation and the operations runbook to YAT ICT.",
    ])

    h1("6. Verification Plan")
    ex.para(doc, "The baseline was verified functionally: homepage and key pages load over HTTPS; the course "
                 "catalogue renders; the enquiry/application form submits, emails Admissions, and forwards to the "
                 "student-admin intake; migrated content and submitted records reconcile against the source.")
    na(doc, "failure-simulation and availability verification do not apply to this baseline — there is no "
            "redundancy to fail over to. Resilience testing is a matter for a future high-availability and "
            "disaster-recovery engagement.")

    h1("7. Out of Scope")
    ex.para(doc, "Stated explicitly — these are the deliberate gaps in the baseline, and the natural inputs to "
                 "any future improvement or disaster-recovery work:")
    add_bullet_list(doc, [
        "Load balancer (ALB) and an Auto Scaling Group across multiple AZs.",
        "Multi-AZ RDS; cross-AZ subnets; replica/standby database.",
        "Content delivery network (CDN) and media offload to object storage.",
        "Web application firewall and advanced security hardening.",
        "Disaster recovery: cross-Region copies, a DR runbook, and tested recovery objectives.",
        "And, out of MTS scope entirely: website content/redesign and ongoing support (YAT).",
    ])

    h1("8. References")
    add_bullet_list(doc, [
        "Website Migration Requirements; Engagement Role Brief; ICT Manager Consultation Notes.",
        "ICT Environment Overview (on-premises current state).",
        "Privacy Policy; User Access Policy; Industry Standards Reference (AWS Well-Architected, ACSC Essential Eight).",
    ])

    h1("Document control")
    ex.etable(doc, ["Field", "Value"],
              [["Document version", "v1.0 — As-built baseline"],
               ["Authored by", "MTS Consulting, in consultation with YAT ICT"],
               ["Approved by", "Pat Lin (MTS Senior Consultant) · Sam Walker (YAT ICT Manager)"],
               ["Engagement", "YAT Website Cloud Migration (2023) — closed 12 May 2023"],
               ["Status", "Retained as the website baseline of record (single-AZ, no HA, no DR)"]],
              widths=[5.0, 11.0])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website/public/documents/YAT-Website-Baseline-Solution-Design.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
