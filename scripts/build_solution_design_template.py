#!/usr/bin/env python3
"""Build the YAT / MTS Solution Design template (.docx) from the brand pack.

The third document type alongside Business Case and Deployment Report — the cloud
solution/architecture design that a Deployment Report then implements. Generalised from
the AT3 HA Design outline: the architecture-design core is generic; the change-specific
sections (review of an existing baseline, implementation sequencing/rollback, simulation
plan) carry the "Not applicable — reason" convention so a greenfield design can omit them.
Assessment scaffolding (UoC notes, student ID) stripped.

Usage:  python scripts/build_solution_design_template.py [output.docx]
Default: ../diploma-cloud-cyber-website/public/templates/YAT-Solution-Design-Template.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_bc_template as bc                      # noqa: E402
import build_deployment_report_template as dr       # noqa: E402  (applic, callout)

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path):
    doc = Document()
    bc.configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    bc.build_header_footer(sec)

    # ---- COVER ----
    bc.wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(bc.ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(bc.GREY)
    bc.paragraph_bottom_rule(doc.add_paragraph(), bc.TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Solution Design")
    sub = doc.add_paragraph().add_run("[ Solution / initiative name ]")
    sub.font.size = Pt(15); sub.italic = True; sub.font.color.rgb = RGBColor.from_string(bc.GREY)
    doc.add_paragraph()
    cover = [
        ("Engagement", "[ Engagement name ]"),
        ("Engagement reference", "[ Reference ]"),
        ("Document version", "[ e.g. v1.0 ]"),
        ("Prepared by", "[ Consultant name, role ]"),
        ("Date", "[ DD/MM/YYYY ]"),
        ("Submitted to", "[ Acceptance authority / sponsor ]"),
        ("Related documents", "[ requirements + any baseline design this supersedes ]"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in cover:
        cells = ct.add_row().cells
        bc.set_cell_borders(cells[0]); bc.set_cell_borders(cells[1]); bc.shade_cell(cells[0], bc.CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        vr = cells[1].paragraphs[0].add_run(v); vr.font.size = Pt(10); vr.italic = True
        vr.font.color.rgb = RGBColor.from_string(bc.GREY)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS + convention ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    doc.add_paragraph("How to use this template", style="Heading 1")
    dr.callout(doc, [
        ("One design for any cloud solution.", "Use this for a greenfield architecture or for hardening / "
         "changing an existing one."),
        ("Complete every section.", "Where a section does not apply, mark it “Not applicable — [reason]” "
         "rather than deleting it."),
        ("Sections flagged with an applicability note", "(terracotta) are the ones a greenfield design "
         "usually marks Not applicable — reviewing an existing baseline, implementation sequencing on a "
         "live system, and the simulation plan."),
        ("This design is implemented by a Deployment Report.", "Business Case (why) → Solution Design "
         "(what / how) → Deployment Report (what was built)."),
    ])
    doc.add_paragraph("Contents", style="Heading 1")
    bc.add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Purpose and Scope")
    bc.guidance(doc, "What this design covers and what it doesn't (≤ ½ page): the solution being designed, "
                     "the objective it serves, and what is in / out of scope.")
    bc.response(doc)

    h1("2. Design Inputs and Requirements")
    h3("2.1 Inputs")
    bc.guidance(doc, "The source documents this design is built from — the approved business case, the "
                     "requirements, the workload/application specification, and (if hardening an existing "
                     "system) the baseline design and the report of what was actually built.")
    bc.response(doc)
    h3("2.2 Requirements the design must meet")
    bc.guidance(doc, "State the targets the design is held to — availability, recovery, performance, "
                     "security, data residency. Add rows as needed.")
    bc.table(doc, ["Requirement / metric", "Target"],
             [["Availability", "[ e.g. 99.9% ]"],
              ["RPO (acceptable data loss)", "[ e.g. ≤ 1 hour ]"],
              ["RTO (time to recover)", "[ e.g. ≤ 4 hours ]"],
              ["Support response (severity-1)", "[ e.g. ≤ 1 hour ]"],
              ["Data residency", "[ e.g. Australia ]"],
              ["[ add others ]", "[ … ]"]],
             widths=[9.0, 7.0])

    h1("3. Review of the Existing Architecture")
    dr.applic(doc, "designs that harden or change an existing system (a greenfield design has no baseline to review)")
    h3("3.1 Architecture review")
    bc.guidance(doc, "Review the existing architecture against the requirements, layer by layer (compute, "
                     "load balancing, database, network, monitoring) — where each currently meets or fails "
                     "the targets.")
    bc.response(doc)
    h3("3.2 Single points of failure")
    bc.table(doc, ["Component", "Failure mode", "Consequence"],
             [["[ … ]", "[ … ]", "[ … ]"], ["[ … ]", "[ … ]", "[ … ]"]],
             widths=[5.0, 5.0, 6.0])
    h3("3.3 Recovery objectives — current state")
    bc.table(doc, ["Component", "Current RPO", "Current RTO", "Meets target?"],
             [["[ … ]", "[ … ]", "[ … ]", "[ Yes/No ]"], ["[ … ]", "[ … ]", "[ … ]", "[ Yes/No ]"]],
             widths=[5.5, 3.5, 3.5, 3.5])
    h3("3.4 Components requiring vertical scaling")
    bc.table(doc, ["Component", "Vertical scale required for", "Availability impact during scale"],
             [["[ … ]", "[ … ]", "[ … ]"]],
             widths=[5.0, 5.0, 6.0])
    h3("3.5 Review findings summary")
    bc.guidance(doc, "Summarise the gap between the current architecture and the requirements: what's met, "
                     "what isn't, and which components drive the gap (≤ 250 words).")
    bc.response(doc)

    h1("4. Architecture Design")
    bc.guidance(doc, "The design proper. For a greenfield design this is the full architecture; for a change, "
                     "describe the design relative to the baseline.")
    h3("4.1 Assumptions and constraints")
    bc.response(doc)
    h3("4.2 AWS account and region")
    bc.guidance(doc, "Region(s) and account; note any data-residency constraint and any new region (e.g. a "
                     "cross-Region backup destination).")
    bc.response(doc)
    for n, title, hint in [
        ("4.3", "Identity and access management (IAM)", "groups/roles/users, MFA, instance profiles — or, for a change, any additions"),
        ("4.4", "Network topology", "VPC, subnets, AZ distribution, gateways, route tables (use the subnet table below)"),
        ("4.5", "Compute (EC2 + Auto Scaling)", "instance type, ASG capacity + AZ spread, scaling policy, health checks"),
        ("4.6", "Load balancing (ALB)", "the load balancer, target group, listener, AZ coverage"),
        ("4.7", "Database (RDS)", "engine + version, instance class, storage, encryption, backups, Multi-AZ (if designed)"),
        ("4.8", "Storage (EBS + S3)", "volumes, buckets, encryption, public-access settings, cross-Region copy (if designed)"),
        ("4.9", "Security", "the tiered security-group model, encryption in transit + at rest, any HA-related adjustments"),
        ("4.10", "Monitoring", "the alarms and the service-level/availability tracking (use the alarm table below)"),
        ("4.11", "Naming and tagging conventions", "the tagging scheme (Project, Environment, Owner, CostCentre, DataClassification, AZ)"),
        ("4.12", "Backup", "the backup baseline and any cross-Region copy / retention design"),
    ]:
        h3(f"{n} {title}")
        bc.guidance(doc, f"Cover: {hint}.")
        if n == "4.4":
            bc.table(doc, ["Subnet", "CIDR", "AZ", "Purpose"],
                     [["[ … ]", "[ … ]", "[ … ]", "[ … ]"], ["[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
                     widths=[4.0, 3.5, 4.0, 4.5])
        elif n == "4.10":
            bc.table(doc, ["Alarm", "Metric", "Threshold", "Triggers"],
                     [["[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
                     widths=[4.0, 4.0, 3.5, 4.5])
        else:
            bc.response(doc)
    h3("4.13 Recovery objectives — designed state")
    bc.table(doc, ["Component", "Designed RPO", "Designed RTO", "Notes"],
             [["[ … ]", "[ … ]", "[ … ]", "[ … ]"], ["Overall service", "[ … ]", "[ … ]", "[ meets target ]"]],
             widths=[5.0, 3.5, 3.5, 4.0])
    h3("4.14 Components requiring vertical scaling — designed state")
    bc.table(doc, ["Component", "Vertical scale required for", "Availability impact (designed)"],
             [["[ … ]", "[ … ]", "[ … ]"]],
             widths=[5.0, 5.0, 6.0])
    h3("4.15 Single points of failure removed")
    dr.applic(doc, "designs that hardened an existing system (cross-reference §3.2)")
    bc.table(doc, ["SPOF (from §3.2)", "Mitigation in this design"],
             [["[ … ]", "[ … ]"], ["[ … ]", "[ … ]"]],
             widths=[7.0, 9.0])

    h1("5. Implementation Sequencing")
    dr.applic(doc, "deployments into a running system, where order and rollback matter")
    bc.guidance(doc, "The order the changes are applied, with per-change duration, expected impact, a "
                     "verification step, and a rollback if it fails. State the total window and buffer.")
    bc.table(doc, ["#", "Change", "Duration", "Expected impact", "Verification", "Rollback"],
             [["1", "[ … ]", "[ … ]", "[ none ]", "[ … ]", "[ … ]"],
              ["2", "[ … ]", "[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[0.8, 4.0, 2.0, 3.0, 3.0, 3.0])

    h1("6. Simulation and Verification Plan")
    dr.applic(doc, "resilience / high-availability designs (the plan the Deployment Report then executes)")
    h3("6.1 Failure simulation plan")
    bc.table(doc, ["#", "Simulation", "Method", "Expected outcome", "Verification"],
             [["F1", "[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[0.8, 3.0, 3.5, 4.5, 4.0])
    h3("6.2 Resize simulation plan")
    bc.table(doc, ["#", "Simulation", "Method", "Expected outcome", "Verification"],
             [["R1", "[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[0.8, 3.0, 3.5, 4.5, 4.0])
    h3("6.3 Verification criteria")
    bc.guidance(doc, "The success criteria — what evidence will demonstrate the design works as intended.")
    bc.response(doc)

    h1("7. Out of Scope")
    bc.guidance(doc, "What this design deliberately does not address, and why (e.g. cross-Region active-active "
                     "DR, application-layer HA, items owned by another team).")
    bc.response(doc)

    h1("8. References")
    bc.guidance(doc, "The source documents and standards informing this design (requirements, baseline "
                     "design, application specification, reference architectures, industry standards), with "
                     "external sources cited with access dates.")
    bc.response(doc)

    h1("9. Review and Approval")
    bc.guidance(doc, "The completed design is submitted to the accepting authority / superior for review. "
                     "The reviewer records their feedback below; the author records how each point was "
                     "addressed; the reviewer then signs off on the design before it is implemented.")
    h3("9.1 Reviewer feedback and author response")
    bc.table(doc, ["#", "Reviewer feedback / comment", "Author response", "Resulting change"],
             [["1", "[ … ]", "[ … ]", "[ … ]"],
              ["2", "[ … ]", "[ … ]", "[ … ]"],
              ["3", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[0.8, 6.0, 5.2, 4.0])
    h3("9.2 Sign-off")
    bc.guidance(doc, "The accepting authority records their decision. “Approved with comments” means the "
                     "design proceeds subject to the changes recorded in §9.1.")
    bc.table(doc, ["Role", "Name", "Decision", "Date", "Signature"],
             [["Prepared by (author)", "[ name, role ]", "—", "[ DD/MM/YYYY ]", ""],
              ["Reviewed and approved by (accepting authority)", "[ name, role ]",
               "[ Approved / Approved with comments / Rejected ]", "[ DD/MM/YYYY ]", ""]],
             widths=[4.6, 3.2, 3.6, 2.4, 2.2])

    h1("Document control")
    bc.table(doc, ["Field", "Value"],
             [["Document version", "[ v1.0 ]"],
              ["Author", "[ Name, role ]"],
              ["Engagement", "[ Engagement name ]"],
              ["Supersedes", "[ any baseline design this replaces ]"],
              ["Implemented by", "[ the Deployment Report that builds this design ]"],
              ["Approval status", "[ Pending / Approved / Rejected with comments ]"]],
             widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website/public/templates/YAT-Solution-Design-Template.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
