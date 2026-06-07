#!/usr/bin/env python3
"""Build the S1-CL2 AT2 ASSESSOR instrument (.docx) by populating the Kangan template.

Institutional compliance document (NOT YAT-branded): loads the official Kangan
'Project Assessment - Assessor' template and fills it in, mirroring the CL1 / AT1 shape.
AT2 = DR Implementation — a single written Deployment Report implementing the AT1-approved
design: the audit-log microservice provisioned as infrastructure-as-code, with monitoring,
built and tested in the AWS Academy lab.

The worked model lives in the YAT-branded exemplar (build_at2_dr_implementation_exemplar);
this document carries the task instructions and the marking guide with bidirectional UoC
traceability. Shared cell/table helpers are reused from build_at1_assessor.

Usage:  python scripts/build_at2_assessor.py [output.docx]
Default: S1-CL2-Cloud-Disaster-Recovery/assessments/AT2/AT2-DR-Implementation-Assessor.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_at1_assessor as a1  # noqa: E402  (shared helpers: set_cell, instr_row, etc.)

from docx import Document  # noqa: E402

TEMPLATE = "templates/Project Assessment - Assessor.docx"

DETAILS = {
    "qualification": "ICT50220 Diploma of Information Technology",
    "units": [
        "ICTCLD503 Implement web-scale cloud infrastructure",
        "ICTCLD505 Implement cloud infrastructure with code",
    ],
    "task_title": "AT2 – DR Implementation",
    "task_number": "2 of 2",
}

OVERVIEW = (
    "Students are assessed on implementing the YAT LMS Global Expansion design approved at AT1. Working as an "
    "MP Tech Solutions (MTS) consultant in the AWS Academy lab, the student provisions the audit-log "
    "microservice as infrastructure-as-code, configures monitoring, tests and troubleshoots the build, and "
    "documents it in a Deployment Report submitted to YAT ICT management for sign-off. This is the "
    "implementation phase of the engagement; the design and plan come from AT1."
)

TASKS = [
    "In this project the student implements, in the AWS Academy lab, the microservice and infrastructure "
    "designed and approved in AT1, and documents it in a single Deployment Report. The build has four "
    "streams:",
    "• Operate the predefined infrastructure-as-code templates — deploy, configure, update and delete "
    "cloud resources from the supplied templates, and troubleshoot template errors.",
    "• Author an infrastructure-as-code template — write, parameterise, update and redeploy a template "
    "that provisions the microservice's resources.",
    "• Build the microservice — deploy and configure the serverless audit-log microservice (API "
    "Gateway → queue → function → datastore), and confirm it functions.",
    "• Monitor — configure metrics and alarms, then document, test, hand over and obtain sign-off.",
]

RESOURCES = [
    "Teacher/assessor supplied resources / Access to:",
    "• The YAT scenario site / intranet — the LMS Global Expansion project, and the AT1 Solution Design "
    "and Disaster Recovery Plan (the approved design this build implements).",
    "• The YAT Deployment Report template (intranet Templates section).",
    "• AWS Academy Cloud Foundations [104469] + AWS Academy Cloud Architecting [172221] — the lab "
    "environment for the build.",
    "• Assessor-supplied lab artefacts — the predefined infrastructure-as-code templates the student "
    "operates, the pre-prepared microservice code elements, and the webhook payload contract.",
]

CRITERIA_STATEMENT = (
    "To receive a Satisfactory outcome for this assessment task, the student must complete every criterion in "
    "the marking guide below to a satisfactory standard. Where a criterion is not yet satisfactory, the student "
    "is given feedback and a further attempt per the Second attempt provisions."
)

CONDITIONS = [
    "C1 — The YAT scenario site / intranet (including the AT1 Solution Design and DR Plan) is accessible to "
    "the student throughout the assessment.",
    "C2 — AWS Academy lab access (Cloud Foundations / Cloud Architecting) is available for the student to "
    "build, deploy and test.",
    "C3 — The assessor supplies the predefined infrastructure-as-code templates, the microservice code "
    "elements, and the webhook payload contract.",
    "C4 — Region note: the design places the audit log in the India region (ap-south-1), which the AWS "
    "Academy lab does not offer; us-west-2 stands in for it and region is a template parameter.",
    "C5 — The assessor role-plays YAT ICT management (Sam Walker, ICT Manager) for the build sign-off.",
]

# Single-part AT2: one Deployment Report. Criteria grouped under one section.
CRITERIA = [
    "D1 — Scope and context: the report states what was deployed, the lab / region setup (incl. the region "
    "simulation), and what is in and out of scope.",
    "D2 — Operate predefined IaC templates: deploys, confirms, updates and deletes cloud resources from the "
    "predefined templates.",
    "D3 — Author an IaC template: authors a template provisioning a related set of resources, parameterised "
    "for reuse, then updates/redeploys and removes it.",
    "D4 — Microservice build: deploys and configures the serverless microservice (API Gateway → queue → "
    "function → datastore) and confirms it functions.",
    "D5 — Troubleshooting: diagnoses and fixes template and microservice errors, with a troubleshooting log.",
    "D6 — Monitoring: configures metrics and alarms (e.g. queue depth, function errors) with thresholds and "
    "notification.",
    "D7 — Tooling: uses the cloud console, CLI and/or SDK for the build.",
    "D8 — Configuration decisions: states and justifies the implementation decisions (IaC service, datastore, "
    "decoupling, compute, region parameterisation).",
    "D9 — Testing and validation: documents the tests run — IaC deploy/update/delete, own-template "
    "parameterise/redeploy, microservice functional, durability, and alarm tests.",
    "D10 — IaC user documentation: produces user documentation for operating the stack (deploy/update/delete "
    "with parameters).",
    "D11 — Feedback and sign-off: confirms/seeks/responds to feedback and obtains final sign-off from the "
    "required personnel.",
    "D12 — Knowledge Evidence: written contextual responses linking the build to the underlying IaC and "
    "microservice concepts (and the shared cloud foundations).",
    "D13 — Document quality: the Deployment Report uses the YAT template, plain professional English, and is "
    "complete with cross-referenced evidence appendices.",
]

BENCHMARK = [
    ("Deployment Report — DR Implementation (ICTCLD503 build + ICTCLD505)", [
        ("D1", "[ICTCLD503 AC 5] / [ICTCLD505 AC 4] scenario data sources; scope of the implemented build."),
        ("D2", "[ICTCLD505 PC 1.1–1.4] benefits/automation/issues/service selection; [PC 2.1–2.6] template "
               "syntax, review, deploy/update/delete, confirm, remove, troubleshoot; [PE 1] deploy/update/remove "
               "using templates; [KE 3, 4, 5, 6, 10]."),
        ("D3", "[ICTCLD505 PC 3.1–3.7] author/deploy/update/parameterise/remove/troubleshoot own template; "
               "[PE 2] create, run and update at least one own template; [KE 8, 9]."),
        ("D4", "[ICTCLD503 PC 3.1–3.3] review design/code, deploy and configure, test and confirm functioning; "
               "[PE 3] deploy a microservice using serverless technologies."),
        ("D5", "[ICTCLD503 PC 3.4] troubleshoot and fix errors; [ICTCLD505 PC 2.6, 3.7] test/troubleshoot "
               "template errors; [ICTCLD505 KE 7] / [ICTCLD503 KE 5] testing and debugging techniques."),
        ("D6", "[ICTCLD503 PC 4.1] set up metrics and trigger scaling alarms; [ICTCLD501 KE 6] techniques and "
               "methods to monitor and create alerts."),
        ("D7", "[ICTCLD503 PE 4] / [ICTCLD505 PE 3] use cloud management consoles, SDKs or command line tools."),
        ("D8", "[ICTCLD505 KE 10, 11] methods/metrics to provision and manage templates; configuration choices "
               "justified."),
        ("D9", "[ICTCLD503 PC 3.3] test and confirm the application functions; [ICTCLD505 PC 2.6, 3.7] test and "
               "troubleshoot."),
        ("D10", "[ICTCLD505 PC 4.1] create user documentation including IaC templates; [PE 4] create user "
                "documentation; [ICTCLD505 FS Writing]."),
        ("D11", "[ICTCLD503 PC 4.2] confirm, seek and respond to feedback; [PC 4.3] obtain final sign-off; "
                "[ICTCLD505 PC 4.2] obtain final sign-off; [ICTCLD505 FS Oral communication]."),
        ("D12", "[ICTCLD505 KE 3–9] IaC concepts; [ICTCLD503 KE 5] testing/debugging; [ICTCLD503 KE 1, 2] / "
                "[ICTCLD505 KE 1, 2] industry standards and standard hardware/software/storage."),
        ("D13", "[ICTCLD505 FS Writing] prepares user documentation in a logical manner using required syntax "
                "and language; [ICTCLD503 FS Writing]."),
    ]),
]

STUDENT_INTRO = [
    ("The engagement and your role", "Heading 2"),
    ("This is the implementation phase of the YAT LMS Global Expansion engagement. The design and disaster "
     "recovery plan were produced and approved in AT1; in AT2 you build the part of that design that this "
     "cluster implements — the audit-log microservice and its infrastructure-as-code — in the AWS Academy lab, "
     "and document it in a Deployment Report.", "Assessor text"),
    ("You remain an MP Tech Solutions (MTS) consultant reporting to Pat Lin (MTS Senior Consultant), with "
     "Sam Walker (YAT ICT Manager) as the accepting authority for the build.", "Assessor text"),
]

STUDENT_TASK = [
    ("The Deployment Report", "Heading 2"),
    ("Using the YAT Deployment Report template (download from the intranet's Templates section), document the "
     "build you perform in the lab. Your report must cover four streams:", "Assessor text"),
    ("• Operate the predefined infrastructure-as-code templates — deploy, configure, update and delete "
     "resources from the supplied templates, and troubleshoot template errors.", "Assessor text"),
    ("• Author your own infrastructure-as-code template — write a parameterised template that provisions "
     "the microservice's resources, then update/redeploy and remove it.", "Assessor text"),
    ("• Build the microservice — deploy and configure the serverless audit-log microservice from the "
     "supplied code elements and the webhook contract, and confirm it functions.", "Assessor text"),
    ("• Monitor, test, document and hand over — configure metrics and alarms, run and record your tests, "
     "produce the IaC user documentation, and obtain sign-off.", "Assessor text"),
    ("Region note: the design places the audit log in the India region (ap-south-1). The AWS Academy lab does "
     "not offer ap-south-1, so use us-west-2 as the India stand-in and make region a template parameter — the "
     "mechanics are real, only the geography is simulated.", "Assessor text"),
    ("Complete the Knowledge Evidence responses (contextual short answers about your own build) and attach "
     "your evidence — screenshots, the templates, and the test/troubleshooting evidence — in the appendices.",
     "Assessor text"),
    ("AT2 is submitted to the LMS as the populated Deployment Report (.docx) with the templates and evidence "
     "appendices completed.", "Assessor text"),
]

TIPS = [
    ("Tips for success", "Heading 2"),
    ("Build what AT1 approved. Your report implements the Solution Design from AT1 — keep the microservice "
     "and the IaC consistent with that design.", "Assessor text"),
    ("Parameterise, don't hard-code. The clearest way to evidence code reuse (and the region simulation) is "
     "to expose region and environment as template parameters and redeploy with different values.", "Assessor text"),
    ("Capture evidence as you go. Screenshot each stack state, the datastore item, the test request/response "
     "and the alarms — with the region visible — rather than reconstructing them afterwards.", "Assessor text"),
    ("Keep a troubleshooting log. Recording the errors you hit and how you fixed them is itself assessed "
     "(testing and debugging).", "Assessor text"),
    ("Sign-off is part of the assessment. You cannot achieve Satisfactory without obtaining the manager's "
     "sign-off on the build.", "Assessor text"),
]


def build(path):
    doc = Document(TEMPLATE)

    # ---- Details ----
    t = doc.tables[0]
    a1.set_cell(t.rows[1].cells[1], DETAILS["qualification"])
    a1.set_cell(t.rows[2].cells[1], DETAILS["units"])
    a1.set_cell(t.rows[3].cells[1], DETAILS["task_title"])
    a1.set_cell(t.rows[4].cells[1], DETAILS["task_number"])

    # ---- Instructions ----
    ti = doc.tables[1]
    a1.set_cell(a1.instr_row(ti, "Assessment overview"), OVERVIEW)
    a1.set_cell(a1.instr_row(ti, "Task"), TASKS)
    a1.set_cell(a1.instr_row(ti, "Resources required"), RESOURCES)
    a1.set_cell(a1.instr_row(ti, "Assessment criteria"), CRITERIA_STATEMENT)
    cond = ti.add_row()
    a1.set_cell(cond.cells[0], "Assessment Conditions & Setup Requirements")
    for r in cond.cells[0].paragraphs[0].runs:
        r.bold = True
    a1.set_cell(cond.cells[1], CONDITIONS)

    # ---- Marking guide ----
    tm = doc.tables[2]
    a1.clear_table_rows(tm, 2)
    a1.add_section_row(tm, "Deployment Report")
    for c in CRITERIA:
        a1.add_criterion_row(tm, c)

    # ---- Instructions to Student ----
    doc.add_paragraph("Instructions to Student", style="Heading 1")
    for text, style in (STUDENT_INTRO + STUDENT_TASK + TIPS):
        doc.add_paragraph(text, style=style)

    # ---- Benchmark ----
    doc.add_paragraph("Marking Benchmark — UoC traceability", style="Heading 1")
    doc.add_paragraph(
        "For each criterion, the unit-of-competency items it evidences. Every PC, PE, KE and FS item allocated "
        "to AT2 in the cluster assessment plan is covered below.", style="Assessor text")
    for part_title, rows in BENCHMARK:
        doc.add_paragraph(part_title, style="Heading 2")
        for cid, uoc in rows:
            p = doc.add_paragraph(style="Assessor text")
            run = p.add_run(f"{cid}  —  "); run.bold = True
            p.add_run(uoc)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL2-Cloud-Disaster-Recovery/assessments/AT2/AT2-DR-Implementation-Assessor.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
