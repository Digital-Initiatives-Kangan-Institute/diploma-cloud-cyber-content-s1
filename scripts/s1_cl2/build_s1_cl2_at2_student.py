#!/usr/bin/env python3
"""Build the S1-CL2 AT2 STUDENT instrument (.docx) by populating the Kangan template.

Student-facing version of AT2, derived from the assessor instrument (build_s1_cl2_at2_assessor) so
the shared content is single-sourced. Loads the official Kangan 'Project Assessment - Student'
template; includes Details, Student instructions, the assessment criteria (what they're marked
against), the detailed task instructions, and the provided artefacts as Appendices A and B
(the data-store template + the microservice code/contract students operate and build with).
The assessor-only material — the marking benchmark, its UoC traceability and assessor notes — is
NOT included.

Usage:  python scripts/build_at2_student.py [output.docx]
Default: S1-CL2-Cloud-Disaster-Recovery/assessments/AT2/AT2-Microservice-IaC-Implementation-Student.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_s1_cl2_at2_assessor as a2   # noqa: E402  shared AT2 content (single source of truth)
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_tables import add_criterion_row, add_section_row, clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402
from helpers.docx_code import add_code_block  # noqa: E402

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Student.docx")


def build(path):
    doc = Document(TEMPLATE)

    # ---- Details (student name / ID / assessor name+email left blank) ----
    t = doc.tables[0]
    set_cell_content(t.rows[2].cells[1], a2.DETAILS["qualification"])
    set_cell_content(t.rows[3].cells[1], a2.DETAILS["units"])
    set_cell_content(t.rows[4].cells[1], a2.DETAILS["task_title"])
    set_cell_content(t.rows[5].cells[1], a2.DETAILS["task_number"])

    # ---- Student instructions ----
    ti = doc.tables[1]
    set_cell_content(find_instruction_row(ti, "Assessment overview"), a2.OVERVIEW)
    set_cell_content(find_instruction_row(ti, "Task"), a2.TASKS)
    set_cell_content(find_instruction_row(ti, "Resources required"), a2.RESOURCES)
    set_cell_content(find_instruction_row(ti, "Assessment criteria"), a2.CRITERIA_STATEMENT)

    # ---- Assessment criteria (no UoC / benchmark) ----
    tm = doc.tables[2]
    clear_table_rows(tm, 2)
    add_section_row(tm, "Deployment Report")
    for c in a2.CRITERIA:
        add_criterion_row(tm, c)

    # ---- Detailed task instructions (no marking benchmark) ----
    doc.add_paragraph("Instructions to Student", style="Heading 1")
    for text, style in (a2.STUDENT_INTRO + a2.STUDENT_TASK + a2.TIPS):
        doc.add_paragraph(text, style=style)

    # ---- Appendix A: provided data-store template (students operate this) ----
    for text, style in a2.APPENDIX_A:
        doc.add_paragraph(text, style=style)
    add_code_block(doc, a2.DATASTORE_YAML)
    for text, style in a2.APPENDIX_A_AFTER:
        doc.add_paragraph(text, style=style)
    add_code_block(doc, a2.DATASTORE_OPS)

    # ---- Appendix B: provided microservice code + contract (students build with this) ----
    for text, style in a2.APPENDIX_B_INTRO:
        doc.add_paragraph(text, style=style)
    add_code_block(doc, a2.HANDLER_LINES)
    for text, style in a2.APPENDIX_B_MID:
        doc.add_paragraph(text, style=style)
    add_code_block(doc, a2.CONTRACT_LINES)
    for text, style in a2.APPENDIX_B_USE:
        doc.add_paragraph(text, style=style)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL2-Cloud-Disaster-Recovery/assessments/AT2/AT2-Microservice-IaC-Implementation-Student.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
