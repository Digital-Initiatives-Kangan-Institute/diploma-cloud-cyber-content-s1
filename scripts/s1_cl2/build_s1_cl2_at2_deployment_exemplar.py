#!/usr/bin/env python3
"""Build the S1-CL2 AT2 Deployment Report EXEMPLAR (.docx) - assessor marking reference.

A worked model answer for the YAT Website Global Expansion Cloud Microservice & IaC Implementation
(S1-CL2 AT2), implementing the AT1-approved design: the audit-log microservice. The student
OPERATES a provided data-store template (DynamoDB) and AUTHORS their own template for the
microservice (API Gateway + SQS + Lambda) that writes to it, with monitoring, built and tested
in the AWS Academy lab. Covers ICTCLD503 element 3 (microservice build), ICTCLD505 (IaC -
operate the provided template + author own), the monitoring strand, and the build documentation
+ sign-off.

Reuses the Deployment Report shell (as CL1 AT2) with a serverless + IaC build narrative (S4)
and testing (S6). Evidence is described, not captured: "[SCREENSHOT - should show ...]". Retains
UoC `Evidences:` tags and the Knowledge Evidence appendix responses (assessor layers an org template
omits). The lab deploys to us-east-1 (per the region-substitution standard); region is the deploy
target, not a template parameter.

Reuses the docx brand helpers (build_bc_template) and prose/table helpers (build_s1_cl1_at1_bc_exemplar).

Usage:  python scripts/build_at2_microservice_iac_exemplar.py [output.docx]
Default: S1-CL2-Cloud-Disaster-Recovery/assessments/AT2/AT2-exemplar-deployment-report.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_body_paragraph, add_bullet_list  # noqa: E402
from helpers.docx_tables import add_data_table  # noqa: E402
from helpers.uoc_tags import add_uoc_evidence_tag  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from helpers.yat_brand import ADDRESS, CHARCOAL, CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402
from helpers.yat_docx_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def na(doc, reason):
    p = doc.add_paragraph()
    r = p.add_run(f"Not applicable - {reason}")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    p.paragraph_format.space_after = Pt(6)
    return p


def ev(doc, kind, desc):
    """Described-evidence placeholder, e.g. [SCREENSHOT - should show ...]."""
    p = doc.add_paragraph()
    tag = p.add_run(f"[{kind}] ")
    tag.bold = True
    tag.font.size = Pt(9.5)
    tag.font.color.rgb = RGBColor.from_string(TEAL)
    d = p.add_run("- " + desc)
    d.font.size = Pt(9.5)
    d.italic = True
    d.font.color.rgb = RGBColor.from_string(GREY)
    p.paragraph_format.space_after = Pt(6)
    return p


def code(doc, lines):
    """A short monospaced code/snippet block."""
    for ln in lines:
        p = doc.add_paragraph()
        r = p.add_run(ln)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(CHARCOAL)
        p.paragraph_format.space_after = Pt(0)


def build(path):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Deployment Report")
    sub = doc.add_paragraph().add_run("YAT Website Global Expansion - Cloud Microservice & IaC Implementation")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    note = doc.add_paragraph().add_run("Assessor exemplar - internal marking reference (not for distribution to students)")
    note.italic = True; note.font.size = Pt(10); note.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT Website Global Expansion & Disaster Recovery - Implementation"),
        ("Engagement reference", "YAT-MTS-2026-003"),
        ("Report version", "v1.0"),
        ("Prepared by", "MTS Consultant"),
        ("Date submitted", "[DD/MM/YYYY]"),
        ("Submitted to", "Sam Walker (YAT ICT Manager) - Pat Lin (MTS Senior Consultant)"),
        ("Related documents", "Website Global Expansion Solution Design (the approved design); the Disaster "
                              "Recovery Plan"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    # 1 Executive Summary
    h1("1. Executive Summary")
    add_uoc_evidence_tag(doc, "[ICTCLD503 PE 3] - [ICTCLD505 PE 1, 2] - overview of the implemented build")
    add_body_paragraph(doc, "This Deployment Report documents the Phase 2 implementation of the YAT Website Global Expansion: "
                 "the audit-log microservice, built and tested in the AWS Academy lab. It implements the "
                 "audit-log microservice from the Solution Design approved at AT1.")
    add_body_paragraph(doc, "The build was delivered in three streams: operating the provided infrastructure-as-code "
                 "template for the audit-log data store (a DynamoDB table); authoring a parameterised template "
                 "that provisions the microservice itself (API Gateway, queue and writer function) wired to that "
                 "store; and deploying and testing the serverless microservice end to end. Monitoring was "
                 "configured, the build was tested and troubleshooted, documented for operational handover, and "
                 "signed off.")

    # 2 Engagement Context
    h1("2. Engagement Context")
    add_uoc_evidence_tag(doc, "[ICTCLD503 AC 5] - [ICTCLD505 AC 4] (scenario data sources used)")
    add_body_paragraph(doc, "Following approval of the design at AT1, MTS implemented the audit-log microservice for the "
                 "India-cohort access logs. The implementation was built in the AWS Academy lab environment "
                 "using the services, the provided data-store template, the supplied microservice code, and the "
                 "webhook contract supplied for the engagement.")
    add_body_paragraph(doc, "Lab / region note: the AT1 design places the audit log in the India region (ap-south-1) for "
                 "data residency. In the AWS Academy Learner Lab the build deploys to us-east-1 "
                 "[scenario: ap-south-1 (India) | deploy: us-east-1]; the region is the deploy target (passed to "
                 "the deploy command), so the same templates target us-east-1 in the lab and ap-south-1 in "
                 "production - the mechanics are real, only the geography is configurable.")

    # 3 Scope of Deployment
    h1("3. Scope of Deployment")
    add_body_paragraph(doc, "In scope: the audit-log microservice and the infrastructure-as-code that provisions it - "
                 "operating the provided data-store template, authoring the microservice template, the "
                 "serverless build (API Gateway, SQS, Lambda) wired to the provided table, the monitoring, the "
                 "infrastructure-as-code user documentation, and the build sign-off.")
    add_body_paragraph(doc, "Out of scope: the website web-tier scaling changes (CloudFront / caching / WAF - designed at "
                 "AT1, a separate workstream, not built here); the disaster-recovery plan and any recovery drill "
                 "(the DR Plan, AT1); and production cutover. The microservice integrates with the website only "
                 "through the webhook contract, so no change to the live website was required.")

    # 4 Build / Change Narrative
    h1("4. Build and Implementation Narrative")
    h3("4.1 Lab environment and access")
    add_uoc_evidence_tag(doc, "[ICTCLD505 AC 1, 5, 7, 8] (vendor platform, IDE, CLI/SSH, console)")
    add_body_paragraph(doc, "Work was performed in the AWS Academy lab via the console, the AWS CLI, and an IDE. The "
                 "writer function ran under a least-privilege execution role (write to the one DynamoDB table "
                 "and read from the queue only); no long-lived credentials were created. The lab product is the "
                 "AWS Academy Learner Lab and the deployment region is us-east-1.")
    ev(doc, "SCREENSHOT", "should show the AWS Academy lab session and the function's execution-role policy.")

    h3("4.2 Operating the provided data-store template")
    add_uoc_evidence_tag(doc, "[ICTCLD505 PC 1.1-1.4] - [ICTCLD505 PC 2.1-2.6] - [ICTCLD505 PE 1, 3]")
    add_body_paragraph(doc, "Benefits and service selection (PC 1.1-1.4): infrastructure-as-code was chosen over manual "
                 "console provisioning because it makes the stack repeatable, reviewable and parameterised - "
                 "exactly what a multi-environment, multi-region design needs. AWS CloudFormation was selected "
                 "as the IaC service: it is native to the platform, needs no extra tooling, and the lab supports "
                 "it.")
    add_body_paragraph(doc, "Template syntax and review (PC 2.1, 2.2): the provided data-store template (datastore.yaml) "
                 "was read to confirm the resource it creates and its outputs - a single DynamoDB audit table, "
                 "with the table name exported for the microservice stack to consume - before deploying. The "
                 "CloudFormation template language (Resources, Parameters, Outputs, intrinsic functions such as "
                 "!Ref and !Sub) is summarised in the Knowledge Evidence appendix (Q2).")
    add_body_paragraph(doc, "Deploy, update, delete (PC 2.3-2.5, PE 1): the provided template was used to deploy the table, "
                 "confirm it in the console, update the stack (changing the EnvName parameter and observing the "
                 "change set), and then delete a stack to confirm clean teardown.")
    add_data_table(doc, ["Action", "Command / method", "Result"],
              [["Deploy", "aws cloudformation deploy --template-file datastore.yaml --stack-name yat-audit-store",
                "CREATE_COMPLETE; the table visible in the console"],
               ["Update", "change EnvName; aws cloudformation deploy (change set reviewed)",
                "UPDATE_COMPLETE; only the intended resource changed"],
               ["Delete", "aws cloudformation delete-stack --stack-name yat-audit-store",
                "DELETE_COMPLETE; the table removed"]],
              widths=[2.6, 8.4, 5.0])
    add_body_paragraph(doc, "Confirm and troubleshoot (PC 2.4, 2.6): the provided template did not deploy as "
                 "supplied - its KeySchema referenced the partition key event_id, which the AttributeDefinitions "
                 "did not declare; the error was read from the stack events, the attribute definition corrected, "
                 "and the stack redeployed and confirmed with describe-stacks - see the troubleshooting log in S6.6.")
    ev(doc, "SCREENSHOT", "should show the data-store stack at CREATE_COMPLETE with the DynamoDB table.")

    h3("4.3 Authoring the microservice infrastructure-as-code template")
    add_uoc_evidence_tag(doc, "[ICTCLD505 PC 3.1-3.7] - [ICTCLD505 PE 2]")
    add_body_paragraph(doc, "A template was authored from scratch to provision the microservice's own resources (PC 3.2): "
                 "the SQS queue, the Lambda writer function, and the HTTP API - wired to the audit table created "
                 "by the provided data-store stack (its name supplied as a parameter, or imported from the "
                 "stack export). It was deployed, then updated to add the API stage (PC 3.3), and parameterised "
                 "so the same template deploys to a named environment without editing the body (PC 3.5) - the "
                 "key code-reuse outcome. The deployment region is the deploy target, not a template "
                 "parameter.")
    code(doc, [
        "Parameters:",
        "  EnvName:        { Type: String, Default: dev }",
        "  AuditTableName: { Type: String }     # the table from the provided data-store stack",
        "Resources:",
        "  EventQueue:  { Type: AWS::SQS::Queue }",
        "  WriterFn:    { Type: AWS::Lambda::Function, Properties: { Runtime: python3.12,",
        "                 Environment: { Variables: { AUDIT_TABLE: !Ref AuditTableName } }, ... } }",
        "  HttpApi:     { Type: AWS::ApiGatewayV2::Api, Properties: { ProtocolType: HTTP } }",
    ])
    add_body_paragraph(doc, "The template follows industry practice: parameters not hard-coding, a least-privilege "
                 "role for the writer, tagging (Project/Environment/Owner/DataClassification), and Outputs that "
                 "export the API endpoint. It was redeployed with a different parameter set to confirm reuse, "
                 "then the stack was removed cleanly (PC 3.6); template errors were tested and fixed (PC 3.7).")
    ev(doc, "FILE", "the authored template is attached in Appendix C (the microservice template).")

    h3("4.4 The audit-log microservice")
    add_uoc_evidence_tag(doc, "[ICTCLD503 PC 3.1-3.4] - [ICTCLD503 PE 3, 4]")
    add_body_paragraph(doc, "Reviewing the design and code (PC 3.1): the supplied microservice code (the writer Lambda) and "
                 "the webhook payload contract were reviewed against the AT1 design before deployment. The data "
                 "flow is: an event webhook -> API Gateway -> SQS -> Lambda -> the provided DynamoDB table.")
    add_body_paragraph(doc, "Deploy and configure (PC 3.2): the serverless services were deployed and configured via the "
                 "authored template - the API Gateway POST /events route, the queue as the Lambda event source, "
                 "and the Lambda writing append-only items to the provided table (its name supplied to the "
                 "function as AUDIT_TABLE). The webhook contract is the single integration point:")
    code(doc, [
        "POST /events",
        "{ event_id, occurred_at, user_ref, cohort, event_type, source_ip }",
    ])
    add_body_paragraph(doc, "Test and confirm functioning (PC 3.3): sample events were posted and confirmed persisted "
                 "exactly once in the provided DynamoDB table (see S6.3). Troubleshoot (PC 3.4): an IAM "
                 "permission error and a payload-validation bug were diagnosed from the Lambda logs and fixed "
                 "(S6.6). Tooling (PE 4): the build used the console, the AWS CLI and the SDK.")
    ev(doc, "SCREENSHOT", "should show a DynamoDB item written by the Lambda after a test POST to /events.")

    h3("4.5 Monitoring and alarms")
    add_uoc_evidence_tag(doc, "[ICTCLD503 PC 4.1]")
    add_body_paragraph(doc, "Metrics and alarms (PC 4.1): CloudWatch alarms were configured on the queue depth "
                 "(ApproximateNumberOfMessagesVisible - the scaling signal for the SQS-driven function: a "
                 "sustained backlog means the writer is not keeping up), the Lambda error rate and throttles, "
                 "and the DynamoDB write throttles, each notifying an SNS topic. The queue-depth alarm is the "
                 "scaling-relevant metric, and the same operational signal supports the DR plan's detection step "
                 "for the audit path.")
    ev(doc, "SCREENSHOT", "should show the CloudWatch alarms list with the queue-depth alarm in OK state.")

    # 5 Configuration Decisions
    h1("5. Configuration Decisions")
    add_data_table(doc, ["Decision", "Choice", "Why"],
              [["IaC service", "AWS CloudFormation", "Native to the platform; no extra tooling; supported in the lab"],
               ["Datastore", "DynamoDB, on-demand (provided template)", "Append-only audit log; scales to zero and to spikes; supplied as the data-store template the student operates"],
               ["Decoupling", "SQS between API Gateway and Lambda", "Buffers spikes so events queue rather than drop"],
               ["Compute", "Lambda (python3.12)", "Serverless; scales with event volume; no servers to run"],
               ["Region", "Deploy target - us-east-1 in the lab; ap-south-1 in production", "Same template deploys to us-east-1 or production by changing the deploy target, not the body"]],
              widths=[3.4, 5.0, 7.6])
    add_body_paragraph(doc, "Industry standards: the build uses standard managed cloud services and "
                 "the JSON/HTTPS and IAM standards common to the platform; storage is the managed NoSQL "
                 "key-value service (DynamoDB) appropriate to an append-only log.")

    # 6 Testing and Validation
    h1("6. Testing and Validation")
    add_uoc_evidence_tag(doc, "[ICTCLD503 PC 3.3, 3.4] - [ICTCLD505 PC 2.6, 3.7]")
    h3("6.1 Provided-template operate verification")
    add_body_paragraph(doc, "Confirmed the provided data-store template deploys, updates (via a reviewed change set) and "
                 "deletes cleanly, with the console and describe-stacks as evidence (PC 2.4).")
    ev(doc, "SCREENSHOT", "stack events showing CREATE_COMPLETE -> UPDATE_COMPLETE -> DELETE_COMPLETE.")
    h3("6.2 Own-template parameterise and redeploy")
    add_body_paragraph(doc, "Deployed the authored microservice template, then redeployed it with a second parameter set "
                 "(different EnvName) to confirm configuration reuse without editing the template (PC 3.4, 3.5).")
    h3("6.3 Microservice functional test")
    add_body_paragraph(doc, "Posted sample events to /events and confirmed each was written exactly once to the provided "
                 "DynamoDB table, with the expected attributes, and that an invalid event was rejected (PC 3.3).")
    ev(doc, "SCREENSHOT", "the test request, the 200 response, and the resulting DynamoDB item.")
    h3("6.4 Durability test")
    add_body_paragraph(doc, "Forced a downstream stall (throttled the writer) and confirmed events queued in SQS and then "
                 "drained without loss once the writer recovered - the decoupling works as designed.")
    h3("6.5 Monitoring / alarm test")
    add_body_paragraph(doc, "Drove the queue depth above the threshold and confirmed the alarm transitioned to ALARM and "
                 "notified the SNS topic, then returned to OK as the backlog cleared (PC 4.1).")
    h3("6.6 Troubleshooting log")
    add_data_table(doc, ["Symptom", "Cause", "Fix"],
              [["Provided data-store template would not deploy as supplied", "KeySchema referenced event_id, but AttributeDefinitions declared a different attribute (id)",
                "Corrected the attribute definition to event_id; redeployed (PC 2.6)"],
               ["Lambda AccessDenied on PutItem", "Execution role lacked dynamodb:PutItem on the table",
                "Scoped a PutItem statement to the provided table ARN"],
               ["Events accepted but not stored", "Payload-validation bug rejected valid cohort values",
                "Fixed the schema check; re-tested (S6.3)"]],
              widths=[5.2, 5.4, 5.4])

    # 7 Operational Handover
    h1("7. Operational Handover")
    h3("7.1 Infrastructure-as-code user documentation")
    add_uoc_evidence_tag(doc, "[ICTCLD505 PC 4.1] - [ICTCLD505 PE 4] - [ICTCLD505 FS Writing]")
    add_body_paragraph(doc, "User documentation was produced so YAT ICT can operate the stack: prerequisites; how to deploy, "
                 "update and delete both stacks (the provided data store, then the microservice) with parameters; "
                 "the parameters and their meaning; the Outputs (the API endpoint); and how to roll back. It is "
                 "included as Appendix C with the templates.")
    h3("7.2 Access and runbook references")
    add_body_paragraph(doc, "Access is via the YAT AWS account roles; the operational runbook references the alarms (S4.5) "
                 "and the redeploy procedure (S7.1).")
    h3("7.3 Known limitations and what's next")
    add_bullet_list(doc, [
        "Built in the AWS Academy Learner Lab; deployed to us-east-1, with ap-south-1 the production target, "
        "set by the deploy target.",
        "The website web-tier scaling (CloudFront / caching / WAF) is designed at AT1 but not built here - "
        "a separate workstream.",
    ])
    h3("7.4 Documentation filing")
    add_body_paragraph(doc, "The report, the templates and the evidence are filed in YAT ICT shared documentation under the "
                 "Records Management Policy.")
    h3("7.5 Feedback record")
    add_uoc_evidence_tag(doc, "[ICTCLD503 PC 4.2] confirm, seek and respond to feedback with required personnel")
    add_data_table(doc, ["Feedback received", "From", "Response", "Action"],
              [["Confirm the audit data stays in the India-region store", "Sam Walker (YAT ICT Manager)",
                "Confirmed - the table is provisioned in the deployment region (us-east-1 in the lab; the production target is the India region, ap-south-1); region is the deploy target",
                "Noted in S2 and S4.2"],
               ["Confirm events can't be lost under load", "Pat Lin (MTS Senior Consultant)",
                "The SQS buffer queues spikes; demonstrated in S6.4", "Recorded in S6.4"]],
              widths=[5.0, 3.0, 4.6, 2.8])
    h3("7.6 Sign-off")
    add_uoc_evidence_tag(doc, "[ICTCLD503 PC 4.3] - [ICTCLD505 PC 4.2] - [ICTCLD505 FS Oral communication]")
    add_data_table(doc, ["Role", "Name", "Date", "Acceptance"],
              [["Prepared by", "MTS Consultant", "[date]", "Submitted"],
               ["Reviewed by", "Pat Lin (MTS Senior Consultant)", "[date]", "Approved for submission"],
               ["Approved by", "Sam Walker (YAT ICT Manager)", "[date]",
                "Build accepted - microservice and IaC implemented, tested and documented; cleared for production deployment"]],
              widths=[4.0, 5.0, 2.5, 4.0])

    # Appendix A - Knowledge Evidence (contextual KE responses; the mandatory written KE location)
    h1("Appendix A - Knowledge Evidence")
    add_uoc_evidence_tag(doc, "[ICTCLD505 KE 3-11] - [ICTCLD503 KE 1, 2, 5] - [ICTCLD505 KE 1, 2] (written contextual KE)")
    h3("Q1 - Why infrastructure as code, not manual provisioning?")
    add_body_paragraph(doc, "[KE 3] In my build, IaC made the stack repeatable and reviewable, let me tear down and "
                 "rebuild identically, and parameterised the environment so the same template deploys to "
                 "different environments and regions (the region is set at deploy time). Manual console "
                 "provisioning would not be reproducible or auditable.")
    h3("Q2 - IaC services, template syntax, and tooling")
    add_body_paragraph(doc, "[KE 4, 5, 6] I used AWS CloudFormation (other options include Terraform and the AWS CDK). A "
                 "CloudFormation template declares Parameters, Resources and Outputs in YAML, with intrinsic "
                 "functions such as !Ref and !Sub for values and dependencies. The tooling to execute it is the "
                 "AWS CLI (cloudformation deploy / delete-stack) and the console.")
    h3("Q3 - Parameterisation for reuse")
    add_body_paragraph(doc, "[KE 8] My microservice template exposes EnvName and AuditTableName as parameters, so one "
                 "template deploys dev and prod and connects to the provided data store without editing the "
                 "body; the deployment region is supplied to the deploy command, so the same template targets "
                 "us-east-1 in the lab or the production India region (ap-south-1).")
    h3("Q4 - Testing and debugging")
    add_body_paragraph(doc, "[KE 7 / 503 KE 5] I debugged template errors from the CloudFormation stack events (the "
                 "provided template's KeySchema/AttributeDefinitions mismatch as supplied) and microservice "
                 "errors from the Lambda CloudWatch logs (an IAM AccessDenied and a payload-validation bug) - "
                 "see the S6.6 log.")
    h3("Q5 - Industry IaC practices and metrics")
    add_body_paragraph(doc, "[KE 9, 10, 11] Industry practice I applied: parameters over hard-coding, least-privilege "
                 "roles, consistent tagging, Outputs for integration values, and change sets to preview updates. "
                 "Stack status and drift detection are the standard signals for managing templates.")
    h3("Q6 - Industry standards and standard products")
    add_body_paragraph(doc, "[503 KE 1, 2 / 505 KE 1, 2] The build uses standard cloud-industry technologies - HTTPS/JSON "
                 "for the API, IAM for access, and the managed serverless and NoSQL services (Lambda, DynamoDB, "
                 "SQS, API Gateway). DynamoDB is the standard managed key-value store, appropriate for an "
                 "append-only audit log.")

    # Appendices (build / IaC / test evidence + reflections)
    h1("Appendix B - Build evidence (screenshots)")
    add_body_paragraph(doc, "Collected console screenshots referenced through S4-S6 (lab session, the provided data-store "
                 "stack at CREATE_COMPLETE, the DynamoDB item, the alarms, the test request/response).")
    h1("Appendix C - Infrastructure-as-code templates and user documentation")
    add_body_paragraph(doc, "The authored microservice template, the provided data-store template operated, and the "
                 "infrastructure-as-code user documentation (deploy/update/delete both stacks with parameters).")
    h1("Appendix D - Test and troubleshooting evidence")
    add_body_paragraph(doc, "The functional, durability and alarm test evidence, and the troubleshooting log (S6.6).")
    h1("Appendix E - Reflections")
    add_body_paragraph(doc, "Short reflections on what the student would do differently and what the build confirmed about "
                 "the AT1 design.")

    # Document control
    h1("Document control")
    add_data_table(doc, ["Field", "Value"],
              [["Document version", "v1.0 - Deployment Report (Cloud Microservice & IaC Implementation)"],
               ["Author", "MTS Consultant"],
               ["Engagement", "YAT Website Global Expansion & Disaster Recovery"],
               ["Implements", "Website Global Expansion Solution Design (AT1)"],
               ["Companion document", "Disaster Recovery Plan"]],
              widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL2-Cloud-Disaster-Recovery/assessments/AT2/AT2-exemplar-deployment-report.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
