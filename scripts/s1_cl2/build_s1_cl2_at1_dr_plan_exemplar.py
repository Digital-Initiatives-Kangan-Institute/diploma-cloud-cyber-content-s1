#!/usr/bin/env python3
"""Build the AT1 Disaster Recovery Plan EXEMPLAR (.docx) — assessor marking reference.

A complete, worked model answer for the YAT Website Global Expansion DR Plan (S1-CL2 AT1
Part B, ICTCLD501), on the YAT brand. Structured on ICTCLD501's element sequence for the
DR-plan *document* — prepare, impact analysis, develop solutions, finalise plan — with a
per-section UoC mapping (ex.uoc "Evidences:" line) so it doubles as the marking reference
and the basis the student template is stripped back from.

Scope discipline:
  * This is PURE disaster recovery — "what we do if the website goes down". The web-scale
    architecture and the audit-log microservice are the companion Solution Design (Part A),
    NOT here; data residency is an input, respected simply by recovering into a second
    *Australian* region so no new cross-border movement of Australian-held data occurs.
  * Elements 1-4 only. ICTCLD501 element 5 (verbal walkthrough, feedback, lodgement,
    sign-off) is the AT1 **Part C** presentation deliverable, not a section of this document.

Figures are indicative and internally consistent (RTO <= 4 h / RPO <= 1 h; primary
ap-southeast-2 Sydney, recovery ap-southeast-4 Melbourne).

Reuses the docx brand helpers (build_bc_template) + the exemplar helpers (build_s1_cl1_at1_bc_exemplar).

Usage:  python scripts/build_at1_dr_plan_exemplar.py [output.docx]
Default: S1-CL2-Cloud-Disaster-Recovery/assessments/AT1/AT1-exemplar-dr-plan.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_bc_template as bc  # noqa: E402  (shared branding helpers + palette)
import build_s1_cl1_at1_bc_exemplar as ex  # noqa: E402  (uoc, para, bullets, etable)

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path):
    doc = Document()
    bc.configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    bc.build_header_footer(sec)

    # ---- COVER ----
    bc.wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(bc.ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(bc.GREY)
    bc.paragraph_bottom_rule(doc.add_paragraph(), bc.TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Disaster Recovery Plan")
    sub = doc.add_paragraph().add_run("YAT Website Global Expansion — Disaster Recovery Plan")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    note = doc.add_paragraph().add_run("Assessor exemplar — internal marking reference (not for distribution to students)")
    note.italic = True; note.font.size = Pt(10); note.font.color.rgb = RGBColor.from_string(bc.GREY)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT Website Global Expansion & Disaster Recovery"),
        ("Engagement reference", "YAT-MTS-2026-003"),
        ("Plan version", "v1.0"),
        ("Prepared by", "MTS Consultant"),
        ("Date submitted", "[DD/MM/YYYY]"),
        ("Submitted to", "Sam Walker (YAT ICT Manager) · Dana Mercer (Marketing & Admissions Manager) · Pat Lin (MTS Senior Consultant)"),
        ("Related documents", "Website Global Expansion Solution Design (the system this plan protects); "
                              "the Functional & Non-Functional Requirements (RTO / RPO)"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        bc.set_cell_borders(cells[0]); bc.set_cell_borders(cells[1]); bc.shade_cell(cells[0], bc.CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    bc.add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")

    # 1 Executive Summary
    h1("1. Executive Summary")
    ex.uoc(doc, "[ICTCLD501 PE 1] — overview of the developed and evaluated plan")
    ex.para(doc, "This Disaster Recovery Plan protects the YAT public website after its expansion to serve the "
                 "new India-campus audience. As the global enrolment front-door, the website is now business-"
                 "critical — and it must survive the loss of an entire cloud region, a risk the existing "
                 "in-region high-availability design does not address. The plan analyses three major risk "
                 "events (loss of the primary region, destructive data loss, and a cyber-security "
                 "compromise), sets the recovery objectives at RTO ≤ 4 hours and RPO ≤ 1 hour, and "
                 "recommends a backup-and-restore strategy that recovers the website into a second Australian "
                 "region (Melbourne, ap-southeast-4).")
    ex.para(doc, "Recovery is driven from infrastructure-as-code plus automated and cross-region backups, "
                 "which together rebuild and restore the platform inside the four-hour window. Recovery is "
                 "kept within Australia, so restoring the service introduces no new cross-border movement of "
                 "the website's Australian-held data. The plan was validated and signed off at the "
                 "design-approval presentation (§7).")

    # 2 Engagement Context & Scope
    h1("2. Engagement Context and Scope")
    ex.uoc(doc, "[ICTCLD501 PC 1.1] identify disaster recovery plan requirements according to business needs")
    ex.para(doc, "The website was migrated to AWS and hardened for in-region high availability (Multi-AZ "
                 "database, a cross-AZ Auto Scaling group behind a load balancer, media served from object "
                 "storage) under prior engagements. This engagement extends it for the India campus, and this "
                 "plan adds the region-level disaster recovery that the partnership makes necessary. The "
                 "business requirement, confirmed with the YAT ICT Manager, is that a disaster affecting the "
                 "primary region must lose no more than one hour of data (RPO ≤ 1 h) and the website must be "
                 "back in service within four hours (RTO ≤ 4 h).")
    ex.para(doc, "In scope: the website's web / application tier and its database, serving both the Australian "
                 "and the India prospective-student audiences. Out of scope: in-region component or "
                 "Availability-Zone failure (already handled by the existing HA design), application-layer "
                 "recovery of the website CMS software (a YAT responsibility), and the separate India "
                 "audit-log store (an independently-durable service designed in the companion Solution "
                 "Design). The plan designs to the recovery objectives; it does not set them.")

    # 3 Current Recovery Position
    h1("3. Current Recovery Position")
    ex.uoc(doc, "[ICTCLD501 PC 1.2] existing organisational recovery plans · [ICTCLD501 PC 1.3] vendor disaster recovery plan and SLAs")
    ex.para(doc, "Existing arrangements (PC 1.2): the only previously documented disaster recovery plan is the "
                 "deprecated on-premises DR plan, which the move to the cloud superseded and which no longer "
                 "applies. The current website platform provides in-region resilience — a Multi-AZ database "
                 "with automatic failover, an Auto Scaling group and load balancer across two Availability "
                 "Zones, media held in Amazon S3, and automated database backups (per the Baseline Design, "
                 "governed by the YAT Backup and Retention Policy). These protect against the loss of a "
                 "component or an Availability Zone, but not against the loss of the whole region; YAT has no "
                 "current region-level recovery plan, which is the gap this plan closes.")
    ex.para(doc, "Vendor provisions and SLAs (PC 1.3): AWS operates each region as multiple independent "
                 "Availability Zones and offers the native cross-region protections this plan relies on — "
                 "automated RDS backups with point-in-time restore, cross-region snapshot copy, S3 "
                 "eleven-nines durability with cross-region replication for media, and CloudFront for global "
                 "delivery. The engagement is covered by AWS Business Support (severity-1 response within one "
                 "hour), and recovery operates under the AWS shared-responsibility model: AWS for the regional "
                 "infrastructure, YAT for the data, configuration, and recovery actions.")

    # 4 Impact Analysis
    h1("4. Impact Analysis")
    ex.uoc(doc, "[ICTCLD501 PC 2.1–2.5] · [ICTCLD501 PE 2]")
    ex.para(doc, "Method: risks were identified in a workshop with YAT ICT and a review of the expanded "
                 "architecture, then rated on likelihood and impact to give a severity, with the recovery "
                 "objectives derived from the business's tolerance for downtime and data loss.")
    ex.para(doc, "Recovery objectives (PC 2.1): RTO ≤ 4 hours and RPO ≤ 1 hour, carried from the "
                 "Functional & Non-Functional Requirements. RPO ≤ 1 h means recovery may lose at most the last "
                 "hour of data; RTO ≤ 4 h means service is restored within four hours of an incident being "
                 "declared. These figures drove the choice of backup cadence and recovery strategy in §5.")
    ex.para(doc, "Data managed (PC 2.3): the website database (approximately 15–25 GB) holds the CMS content, "
                 "the course catalogue, and the enquiry / application records — which include prospective-"
                 "student personal information; uploaded media (brochures, images, course PDFs) is held "
                 "separately in Amazon S3. This data is subject to the Australian Privacy Act; keeping recovery "
                 "within Australia avoids any new cross-border handling of it. (The India-cohort access logs are "
                 "held in India by the separate audit-log service and are outside this plan.)")
    ex.para(doc, "Risk environment: the principal risks to a public-facing, mission-critical website now "
                 "relied on by a partner institution for enrolment are assessed below.")
    ex.etable(doc, ["#", "Risk event", "Likelihood", "Impact", "Severity", "Drives"],
              [["RE1", "Loss of the primary region (region-wide AWS outage)", "Low", "Critical", "High",
                "Cross-region recovery capability"],
               ["RE2", "Destructive data loss (corruption or accidental deletion)", "Medium", "High", "High",
                "Point-in-time backups (RPO)"],
               ["RE3", "Cyber-security compromise (ransomware / defacement / credential abuse)", "Medium", "Critical", "High",
                "Isolated, immutable, cross-region backups"]],
              widths=[0.9, 5.4, 2.0, 1.8, 1.8, 4.1])
    ex.para(doc, "Severity of impact and disruption (PC 2.4): RE1 and RE3 are the most severe — both can take "
                 "the public website fully offline (cutting off the enrolment front-door), and RE3 also "
                 "threatens data integrity and the institution's public reputation with the partner. RE2 is "
                 "contained to the affected data but can still breach the RPO if backups are inadequate.")
    ex.para(doc, "Plan exclusions (PC 2.2): single-component and single-AZ failures are excluded — they are "
                 "handled by the existing in-region HA design, not this plan. Application-layer faults within "
                 "the website CMS software are excluded as a YAT responsibility. The plan assumes recovery into "
                 "at least one healthy AWS region remains possible; a simultaneous loss of both the primary and "
                 "recovery regions is outside its scope.")
    ex.para(doc, "Outcomes (PC 2.5): this impact analysis is recorded in YAT ICT shared documentation per the "
                 "Backup and Retention and Records Management policies, and is the basis for the recovery "
                 "strategy in §5.")

    # 5 Recovery Strategy
    h1("5. Recovery Strategy")
    ex.uoc(doc, "[ICTCLD501 PC 3.1–3.4]")
    ex.para(doc, "Range of solutions evaluated (PC 3.1): the four standard cloud DR strategies were "
                 "assessed against the RTO/RPO targets and cost.")
    ex.etable(doc, ["Option", "RTO", "RPO", "Relative cost", "Fit"],
              [["Backup & Restore (recommended)", "Hours", "Minutes–1 h", "Low",
                "Meets ≤ 4 h / ≤ 1 h at lowest cost and complexity"],
               ["Pilot Light", "~1 h", "Minutes", "Medium", "Faster than needed; ongoing standby cost"],
               ["Warm Standby", "Minutes", "Seconds", "High", "Exceeds the requirement; not justified"],
               ["Multi-region Active-Active", "Near-zero", "Near-zero", "Very high",
                "Far exceeds the requirement; disproportionate"]],
              widths=[4.4, 2.0, 2.4, 2.4, 4.3])
    ex.para(doc, "Recommended approach: Backup & Restore into Melbourne (ap-southeast-4). Automated RDS "
                 "backups give point-in-time restore (RPO in minutes); a cross-region snapshot copy to the "
                 "recovery region runs at least hourly (within the 1-hour RPO); S3 cross-region replication "
                 "keeps the media available; and the infrastructure-as-code templates rebuild the stack in the "
                 "recovery region by changing the target-region parameter. Recovering into a second Australian "
                 "region keeps the website's Australian-held data onshore. This is the simplest strategy that "
                 "meets the objectives — consistent with the requirement to avoid unnecessary cost and "
                 "complexity.")
    ex.para(doc, "Vendor protections and risk prioritisation (PC 3.2): the strategy leans on AWS-native "
                 "protections — durable snapshots, cross-region copy, and a separate backup vault with "
                 "restricted access for ransomware resistance. Risks are prioritised RE1 (region loss) first — "
                 "addressed by cross-region recovery — then RE2 (data loss) via point-in-time restore, then "
                 "RE3 (compromise) via isolated, immutable backups held in a different region.")
    ex.para(doc, "Insurance (PC 3.3): YAT holds a cyber and business-interruption insurance policy. It was "
                 "reviewed as a complementary control, not a substitute for recovery: the policy transfers "
                 "residual financial risk (e.g. breach-response and liability costs) that the technical plan "
                 "cannot eliminate, while this plan reduces the likelihood and duration of a claimable event. "
                 "The policy's notification and evidence requirements are noted in the recovery runbook.")
    ex.para(doc, "Other components (PC 3.4): the strategy also includes a recovery runbook, a declaration and "
                 "escalation path, a stakeholder communication plan (including the India partner and YAT "
                 "Marketing), CDN / DNS failover to the recovered endpoint, and the order of recovery across "
                 "the tiers. The approach aligns with ISO/IEC 27001 and 27002 (information security) and "
                 "ISO/IEC 27031 (ICT readiness for business continuity).")

    # 6 The Disaster Recovery Plan
    h1("6. The Disaster Recovery Plan")
    ex.uoc(doc, "[ICTCLD501 PC 4.1–4.3] · [ICTCLD501 PE 3]")
    ex.para(doc, "Alignment to risk (PC 4.1): the recovery actions below are aligned to the prioritised risks "
                 "from §4 — the cross-region rebuild answers RE1, the point-in-time restore answers RE2, and "
                 "restoring from the isolated backup vault answers RE3.")
    ex.para(doc, "Detection and declaration: a region-level event is detected by CloudWatch alarms, the AWS "
                 "Health Dashboard, and the availability monitoring on the website endpoints; breaching the "
                 "detection threshold pages the on-call YAT ICT officer, who declares a disaster and starts the "
                 "runbook.")
    ex.para(doc, "Recovery steps, timelines and service providers (PC 4.2):")
    ex.etable(doc, ["Step", "Action", "Owner", "Target (cumulative)"],
              [["1", "Detect and declare the disaster; notify stakeholders and the India partner", "YAT ICT on-call", "0:20"],
               ["2", "Provision the recovery stack in ap-southeast-4 from the IaC templates (region parameter = ap-southeast-4)", "YAT ICT", "1:30"],
               ["3", "Restore the database from the latest cross-region snapshot / point-in-time copy; confirm the media is served from the replicated S3 bucket", "YAT ICT", "2:45"],
               ["4", "Verify the integrity of the restored data and reconcile to the RPO checkpoint", "YAT ICT", "3:00"],
               ["5", "Validate connectivity and core website functions; smoke-test the catalogue and the enquiry / application intake for both audiences", "YAT ICT", "3:30"],
               ["6", "Cut CDN / DNS over to the recovered endpoint; confirm visitors are served", "YAT ICT", "3:50"],
               ["7", "Post-recovery confirmation, insurer / stakeholder notifications as required, incident log", "YAT ICT", "4:00"]],
              widths=[1.0, 8.3, 3.0, 3.2])
    ex.para(doc, "How the plan meets the objectives (PE 3): RPO ≤ 1 h is met because point-in-time backups and "
                 "the at-least-hourly cross-region copy cap data loss below one hour; RTO ≤ 4 h is met because "
                 "the templated rebuild and the database restore complete inside the four-hour budget, as the "
                 "step timeline shows. Service providers in the plan are AWS (regional infrastructure and "
                 "recovery services), YAT ICT (declares and executes recovery), and MTS (designed the plan and "
                 "handed it over).")
    ex.para(doc, "Documentation (PC 4.3): this plan is the documented disaster recovery plan; it is versioned, "
                 "lodged in YAT ICT shared documentation under the Records Management Policy, and registered "
                 "against the Change Management Procedure, with a review on each major architecture change and "
                 "at least annually. It is presented for approval at the design-approval presentation before "
                 "implementation begins (ICTCLD501 element 5 — the AT1 Part C presentation).")

    # 7 Plan Validation and Approval (element 5 — conducted at the AT1 Part C presentation; recorded here)
    h1("7. Plan Validation and Approval")
    ex.uoc(doc, "[ICTCLD501 PC 5.1–5.4] — conducted and observed at the AT1 Part C presentation; this section records the outcome")
    ex.para(doc, "Walkthrough (PC 5.1): the plan was presented to the YAT ICT Manager and stakeholders at the "
                 "design-approval presentation, including a tabletop run of the RE1 region-loss scenario against "
                 "the §6 steps, which confirmed the actions and that the timeline stays within the four-hour RTO.")
    ex.para(doc, "Feedback sought and responded to (PC 5.2):")
    ex.etable(doc, ["Feedback received", "From", "Response", "Resulting action"],
              [["Confirm recovering into a second region keeps Australian-held data onshore",
                "Sam Walker (YAT ICT Manager)",
                "Confirmed — the recovery region is Melbourne (ap-southeast-4), so the website's Australian-held data stays in Australia throughout",
                "Recorded in §5"],
               ["Check the recovery cost is proportionate",
                "Robin Ng (CFO)",
                "Backup-and-restore was chosen specifically as the lowest-cost option meeting the objectives (§5)",
                "Cost rationale recorded in §5"]],
              widths=[5.2, 3.0, 4.3, 3.0])
    ex.para(doc, "Lodgement (PC 5.3): the approved plan is lodged in YAT ICT shared documentation under the "
                 "Records Management Policy and registered against the Change Management Procedure, with a "
                 "review on each major architecture change and at least annually.")
    ex.para(doc, "Sign-off (PC 5.4):")
    ex.etable(doc, ["Role", "Name", "Date", "Acceptance"],
              [["Prepared by", "MTS Consultant", "[date]", "Submitted"],
               ["Reviewed by", "Pat Lin (MTS Senior Consultant)", "[date]", "Approved for submission"],
               ["Approved by", "Sam Walker (YAT ICT Manager)", "[date]",
                "DR Plan accepted — objectives, strategy and recovery steps confirmed; cleared to proceed to implementation"]],
              widths=[4.0, 5.0, 2.5, 4.0])

    # Appendix A — Knowledge Evidence (the Part B "B15" deliverable: written contextual responses)
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1("Appendix A — Knowledge Evidence")
    ex.para(doc, "Written contextual responses for ICTCLD501 — the mandatory knowledge-evidence location for "
                 "Part B. Each question asks the candidate to reason about their own plan; the answers below "
                 "are the model responses (assessor reference).")

    def ke_qa(tag, question, answer):
        ex.uoc(doc, tag)
        qp = doc.add_paragraph(); qr = qp.add_run(question); qr.bold = True; qr.font.size = Pt(10.5)
        ex.para(doc, answer)

    ke_qa("[ICTCLD501 KE 1] — risk environments in cloud/ICT environments",
          "Q1. Your impact analysis settled on three major risk events. What is distinctive about the risk "
          "environment of a public-facing, cloud-hosted website, and why do region loss, destructive data "
          "loss and cyber-compromise capture it?",
          "The website runs on shared cloud infrastructure, so its risk environment includes provider-level "
          "failures — a whole AWS region becoming unavailable — that an on-premises system would not face, "
          "alongside the universal risks to data integrity and security. Being public and internet-facing "
          "raises both the likelihood and the impact of a cyber-compromise. The three events span the failure "
          "domains that matter above the Availability-Zone level the existing HA design already covers: "
          "infrastructure (RE1), data (RE2) and security (RE3). Single-component and single-AZ failures are "
          "excluded because the in-region HA design handles them.")

    ke_qa("[ICTCLD501 KE 2] — data analysis methodologies to determine the risk environment",
          "Q2. How did you determine the risk environment — what method did you use to identify and rate the "
          "risks, and why is it appropriate here?",
          "I identified candidate risks in a structured workshop with YAT ICT together with a review of the "
          "expanded architecture, then rated each on likelihood and impact to derive a severity. This "
          "qualitative likelihood-by-impact method suits a plan with a small number of high-consequence events "
          "and little historical incident data: it surfaces the events that actually drive the recovery "
          "objectives without implying a false precision a quantitative model could not support.")

    ke_qa("[ICTCLD501 KE 3] — disaster recovery techniques applicable to cloud environments",
          "Q3. You evaluated four cloud DR strategies and recommended Backup & Restore. Explain the recovery "
          "techniques available in the cloud and why Backup & Restore fits this website rather than Pilot "
          "Light, Warm Standby or Active-Active.",
          "The standard cloud DR techniques trade recovery speed against cost: Backup & Restore rebuilds from "
          "backups (RTO in hours); Pilot Light keeps core data replicated and scales the rest on failover; "
          "Warm Standby runs a smaller always-on copy; and Multi-region Active-Active runs a full duplicate. "
          "The website's targets are RTO ≤ 4 h and RPO ≤ 1 h. Backup & Restore meets both at the lowest cost "
          "and complexity, especially with IaC-driven rebuild and cross-region backups; the faster options "
          "exceed the requirement and carry ongoing standby cost the brief explicitly asks to avoid.")

    ke_qa("[ICTCLD501 KE 4] — ISO/IEC 27001, 27002 and 27031 standards",
          "Q4. Which information-security and business-continuity standards informed your plan, and how does "
          "the plan reflect them?",
          "ISO/IEC 27001 and 27002 (information-security management and controls) inform the security of the "
          "backups — isolated, immutable, access-restricted vaults for ransomware resistance — and the "
          "handling of the personal information the database holds. ISO/IEC 27031 (ICT readiness for business "
          "continuity) frames the plan's shape: detection and declaration, defined recovery objectives, a "
          "documented and tested recovery procedure, and clear roles — which is why the plan carries a "
          "runbook, an escalation path and a validated step timeline rather than just a backup configuration.")

    ke_qa("[ICTCLD501 KE 5] — RTO and RPO standards and techniques",
          "Q5. Explain RTO and RPO in your own words, the values you set for the website, and the techniques "
          "you used to meet each.",
          "RPO is the maximum acceptable amount of data loss — how far back recovery may have to go — and RTO "
          "is the maximum acceptable time to restore service. For the website I set RPO ≤ 1 hour and RTO ≤ 4 "
          "hours, carried from the requirements. RPO is met by point-in-time database backups plus an "
          "at-least-hourly cross-region snapshot copy, which together cap data loss below an hour. RTO is met "
          "because the IaC templates rebuild the stack in the recovery region and the database restore "
          "completes inside the four-hour budget, as the §6 step timeline demonstrates.")

    ke_qa("[ICTCLD501 KE 6] — techniques and methods to monitor and create alerts in cloud environments",
          "Q6. How would a region-level disaster be detected to trigger this plan, and what monitoring and "
          "alerting would you put in place so the on-call team is notified in time to meet the RTO?",
          "Detection relies on layered cloud monitoring: CloudWatch alarms on the website's health and key "
          "metrics, the AWS Health Dashboard for region-level service events, and external (synthetic) "
          "availability monitoring of the public endpoints from multiple locations. When a threshold is "
          "breached — for example the endpoints failing from several locations while AWS reports a regional "
          "impairment — the alarm raises an alert that pages the on-call YAT ICT officer, who declares the "
          "disaster and starts the runbook. Fast automated detection is what makes the four-hour RTO "
          "achievable, because the recovery clock only starts once the disaster is declared.")

    # Document control
    h1("Document control")
    ex.etable(doc, ["Field", "Value"],
              [["Document version", "v1.0 — Disaster Recovery Plan"],
               ["Author", "MTS Consultant"],
               ["Engagement", "YAT Website Global Expansion & Disaster Recovery"],
               ["Companion document", "Website Global Expansion Solution Design"],
               ["Successor document", "Website Global Expansion Deployment Report (Phase 2 implementation)"]],
              widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL2-Cloud-Disaster-Recovery/assessments/AT1/AT1-exemplar-dr-plan.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
