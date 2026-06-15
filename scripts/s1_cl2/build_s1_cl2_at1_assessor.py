#!/usr/bin/env python3
"""Build the S1-CL2 AT1 ASSESSOR instrument (.docx) by populating the Kangan template.

This is an institutional compliance document, NOT a YAT-branded artefact: it loads the
official Kangan 'Project Assessment - Assessor' template and fills it in, preserving the
Kangan structure and styles exactly (Details, Teacher/Assessor instructions, Marking Guide,
Instructions to Student, Benchmark). It mirrors the shape of the approved CL1 AT1 assessor
instrument.

AT1 = Cloud Expansion: Design & DR Plan — three parts:
  Part A  Solution Design  (ICTCLD503 design: web-scale + microservice)
  Part B  DR Plan          (ICTCLD501)
  Part C  Presentation     (ICTCLD501 element 5 — the approval gate; verbal KE Q&A)

Vehicle (per scenario-flow): CL2 ASSESSES on the YAT website. The website is a public,
unauthenticated site (anonymous visitors), which is why the web-scale design carries edge
delivery, WAF / bot / DDoS protection and SEO — the deliberate contrast with the LMS practice
vehicle (an authenticated cohort).

Knowledge evidence: each report carries a written Knowledge Evidence appendix (the mandatory
location — A12 for the Solution Design's 503 KE; B15 for the DR Plan's 501 KE, including the
detection/alerting KE 6), with the same KE re-covered verbally at the Part C Q&A (C5). The
worked model answers live in the YAT-branded exemplars (build_at1_*_exemplar); this document
carries the task instructions and the marking guide with bidirectional UoC traceability.

Usage:  python scripts/build_at1_assessor.py [output.docx]
Default: S1-CL2-Cloud-Disaster-Recovery/assessments/AT1/AT1-Design-DR-Plan-Assessor.docx
"""
import sys
from pathlib import Path

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

TEMPLATE = "templates/Project Assessment - Assessor.docx"


# ---------- cell / table helpers (preserve template styles) ----------

def set_cell(cell, lines):
    """Replace a cell's content with one paragraph per line, keeping the cell's style."""
    if isinstance(lines, str):
        lines = [lines]
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    for r in list(first.runs):
        r._element.getparent().remove(r._element)
    first.add_run(lines[0])
    style = first.style
    for line in lines[1:]:
        np = cell.add_paragraph()
        np.style = style
        np.add_run(line)


def instr_row(table, label):
    """Return the col-1 cell of the instructions row whose col-0 starts with `label`."""
    for row in table.rows:
        if row.cells[0].text.strip().lower().startswith(label.lower()):
            return row.cells[1]
    raise KeyError(label)


def clear_table_rows(table, keep):
    """Remove all rows after the first `keep` rows."""
    for row in table.rows[keep:]:
        row._element.getparent().remove(row._element)


def add_section_row(table, text):
    row = table.add_row()
    set_cell(row.cells[0], text)
    # bold the section label
    for r in row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell(row.cells[1], "")
    return row


def add_criterion_row(table, text):
    row = table.add_row()
    set_cell(row.cells[0], text)
    set_cell(row.cells[1], "Yes        No")
    return row


# ---------- content ----------

DETAILS = {
    "qualification": "ICT50220 Diploma of Information Technology",
    "units": [
        "ICTCLD501 Develop cloud disaster recovery plans",
        "ICTCLD503 Implement web-scale cloud infrastructure",
    ],
    "task_title": "AT1 – Cloud Expansion: Design & DR Plan",
    "task_number": "1 of 2",
}

OVERVIEW = (
    "Students are assessed on the cloud design and disaster-recovery planning for the YAT website global "
    "expansion. Working as an MP Tech Solutions (MTS) consultant, the student produces a Solution Design "
    "(Part A) and a Disaster Recovery Plan (Part B) for the expanded public website, then presents both to "
    "YAT ICT management for approval (Part C). This is the design-and-planning phase of the engagement; the "
    "approved design and plan are implemented in AT2."
)

TASKS = [
    "In this project the student takes on the role of an MTS consultant engaged by YAT College to design "
    "the changes that prepare the public website for its India-campus expansion and to plan its disaster "
    "recovery. The project has three parts:",
    "• Part A — Solution Design: a written design covering the web-scale architecture (serving a "
    "global, anonymous public audience) and an audit-log microservice. Produced in the YAT Solution "
    "Design template.",
    "• Part B — Disaster Recovery Plan: a written DR plan for the designed website — risk and "
    "impact analysis, recovery objectives (RTO/RPO), recovery strategy, the recovery runbook, and "
    "validation. Produced in the YAT Disaster Recovery Plan template.",
    "• Part C — Presentation: the student presents the Solution Design and DR Plan to YAT ICT "
    "management (role-played by the assessor) for approval, responds to questions, and obtains sign-off.",
]

RESOURCES = [
    "Teacher/assessor supplied resources / Access to:",
    "• The YAT scenario site / intranet — the Website Global Expansion project (engagement brief, "
    "requirements, ICT manager consultation notes, data residency & sovereignty requirements), the ICT "
    "documents (website specification, HA-hardened solution design baseline), and the deprecated on-prem DR plan.",
    "• The YAT Solution Design template and YAT Disaster Recovery Plan template (intranet Templates section).",
    "• AWS Academy Learner Lab — the authorised lab environment (referenced for the design; the build "
    "itself is AT2).",
    "• The YAT Document Archive — examples of previous MTS solution designs and DR plans.",
]

CRITERIA_STATEMENT = (
    "To receive a Satisfactory outcome for this assessment task, the student must complete every criterion in "
    "the marking guide below to a satisfactory standard. All three parts (A, B and C) must be satisfactory. "
    "Where a criterion is not yet satisfactory, the student is given feedback and a further attempt per the "
    "Second attempt provisions."
)

CONDITIONS = [
    "C1 — The YAT scenario site / intranet is accessible to the student throughout the assessment.",
    "C2 — The student has access to the YAT Solution Design and Disaster Recovery Plan templates.",
    "C3 — AWS Academy Learner Lab access is available for the student to reference cloud services in the design.",
    "C4 — For Part C, the assessor role-plays YAT ICT management (Sam Walker, ICT Manager) and conducts "
    "the approval meeting, capturing feedback and sign-off.",
    "C5 — The presentation may be delivered in person on campus or via video conference (confirm with the "
    "assessor).",
]

# Marking guide: (criterion text) per part. UoC traceability is carried in the Benchmark.
PART_A = [
    "A1 — Purpose and scope: states the two design changes (scale for a global, anonymous public audience; "
    "add the audit-log microservice) and what is in and out of scope.",
    "A2 — Design inputs and requirements: identifies the source documents and the requirements the design "
    "must meet, including data residency as an input constraint.",
    "A3 — Review of the existing architecture: reviews the current website baseline against the requirements "
    "and identifies the gaps the design closes.",
    "A4 — Web-scale design — scaling by layer: designs architecture changes using cloud services that "
    "scale network, compute and storage as utilisation increases.",
    "A5 — Web-scale design — global delivery: designs the changes needed to serve a global, anonymous "
    "public audience (edge delivery and discoverability).",
    "A6 — Web-scale design — availability and security: confirms availability and security are "
    "maintained through the changes, including protection of the public attack surface (web exploits, bots, "
    "DDoS), and reviews the design against the requirements.",
    "A7 — Web-scale component choices: justifies the component choices across SQL/NoSQL, monolith/"
    "microservice, virtual/container/serverless compute, and CDN/in-memory store.",
    "A8 — Microservice — services and data transactions: identifies the microservice(s) and the data "
    "transaction(s) each handles, per business needs.",
    "A9 — Microservice — supporting cloud services: determines the cloud services that support the "
    "microservice architecture.",
    "A10 — Microservice — architecture: designs the microservice architecture (cohesion/coupling, "
    "persistent storage, API/messaging/queuing), with the interface/integration contract.",
    "A11 — Architecture documented and justified: the web-scale and microservice design changes are "
    "documented and justified.",
    "A12 — Knowledge Evidence appendix: written contextual responses linking the student's design choices "
    "to the underlying web-scale concepts (components, infrastructure, web-scaling principles).",
    "A13 — Document quality: the Solution Design uses the YAT template, plain professional English, and is "
    "complete and internally consistent.",
]

PART_B = [
    "B1 — Plan requirements and context: identifies the DR plan requirements according to business needs.",
    "B2 — Current recovery position: determines the existing organisational recovery arrangements and the "
    "vendor (cloud provider) DR provisions and SLAs.",
    "B3 — Recovery objectives: determines RTO and RPO according to business needs and explains them.",
    "B4 — Data managed: estimates the amount and security level of data managed.",
    "B5 — Risk assessment: assesses at least three major risk events with likelihood, impact and severity, "
    "and evaluates the severity of impact and disruption.",
    "B6 — Plan exclusions: assesses potential risks and plan exclusions according to business requirements.",
    "B7 — Impact-analysis outcomes: documents the outcomes of the impact analysis per organisational "
    "policies and procedures.",
    "B8 — Recovery strategy: develops a range of DR solutions and recommends one that meets the RTO/RPO, "
    "with rationale.",
    "B9 — Vendor protections and prioritisation: determines vendor protections and prioritises the risks.",
    "B10 — Insurance: assesses external insurance protection and its suitability.",
    "B11 — Other recovery components: identifies the other DR plan components, aligned to relevant "
    "continuity standards.",
    "B12 — The recovery plan: aligns recovery to the prioritised risks and outlines the steps, timelines, "
    "service providers and detection/alerting.",
    "B13 — Meeting the objectives: shows how the plan reaches the RTO and RPO targets (covering at least "
    "three major risk events).",
    "B14 — Plan documented: the DR plan is documented according to business needs and requirements.",
    "B15 — Knowledge Evidence appendix: written contextual responses linking the plan to DR concepts "
    "(risk environment, data-analysis method, DR techniques, ISO standards, RTO/RPO, and the monitoring/"
    "alerting used for disaster detection).",
]

PART_C = [
    "C1 — Verbal walkthrough: conducts a verbal walkthrough of the Solution Design and DR Plan with the "
    "required personnel.",
    "C2 — Feedback: seeks and responds to feedback during the presentation.",
    "C3 — Lodgement: lodges the DR plan according to organisational and legislative protocol.",
    "C4 — Sign-off: obtains final sign-off from the required personnel.",
    "C5 — Knowledge Evidence Q&A: answers the assessor's contextual questions on the student's own design "
    "and plan choices (the verbal knowledge-evidence location, re-covering the A12 / B15 appendices).",
    "C6 — Oral communication: uses listening and questioning techniques and industry language appropriate "
    "to the audience.",
]

# Benchmark: per-criterion UoC evidenced (bidirectional traceability). KE is allocated to the
# written KE appendices (A12 / B15); the body criteria carry the PC / PE / FS they demonstrate.
BENCHMARK = [
    ("Part A — Solution Design (ICTCLD503 design)", [
        ("A1", "[ICTCLD503 PC 1.1] determine and confirm cloud web-scaling needs."),
        ("A2", "[ICTCLD503 PC 1.1] confirm needs against business needs; data residency carried as an input."),
        ("A3", "[ICTCLD503 PC 1.2] review architecture for the web application according to business needs."),
        ("A4", "[ICTCLD503 PC 1.3] identify cloud services to scale; [PC 1.4] design changes scale network, "
               "compute and storage; [PE 1] design at least one scaling architecture."),
        ("A5", "[ICTCLD503 PC 1.5] determine architecture changes to scale for a global user base."),
        ("A6", "[ICTCLD503 PC 1.6] check availability and security maintained and review the design as required; "
               "[PE 5] apply web-scaling principles and technologies."),
        ("A7", "[ICTCLD503 PC 1.7] supported by the documented justification of the web-scale component choices."),
        ("A8", "[ICTCLD503 PC 2.1] identify microservices and data transactions; [PE 2] design at least one "
               "microservice architecture."),
        ("A9", "[ICTCLD503 PC 2.2] determine cloud services to support the microservice architecture."),
        ("A10", "[ICTCLD503 PC 2.3] design microservice architecture using cloud services."),
        ("A11", "[ICTCLD503 PC 1.7] document and justify architecture changes; [PC 2.4] document and justify "
                "microservice design; [FS Writing]."),
        ("A12", "[ICTCLD503 KE 3] web-scale components (SQL/NoSQL, monolith/microservice, compute models, "
                "CDN/in-memory); [KE 4] web-scale infrastructure (cohesive/coupled, managed datastore, "
                "API/messaging/queuing); [KE 6] web-scaling principles and technologies — evidenced as written "
                "contextual responses against the student's own design."),
        ("A13", "[ICTCLD503 FS Writing] complex documentation in the required format and language."),
    ]),
    ("Part B — DR Plan (ICTCLD501, elements 1–4)", [
        ("B1", "[ICTCLD501 PC 1.1] identify DR plan requirements according to business needs."),
        ("B2", "[ICTCLD501 PC 1.2] determine existing organisational recovery plans; [PC 1.3] identify vendor "
               "DR plan and SLAs."),
        ("B3", "[ICTCLD501 PC 2.1] determine RTO and RPO according to business needs."),
        ("B4", "[ICTCLD501 PC 2.3] estimate amount and security level of data managed."),
        ("B5", "[ICTCLD501 PC 2.4] evaluate severity of impact and disruption; [PE 2] determine likelihood and "
               "impact of risk events."),
        ("B6", "[ICTCLD501 PC 2.2] assess potential risks and plan exclusions according to business requirements."),
        ("B7", "[ICTCLD501 PC 2.5] document outcomes of impact analysis per organisational policies and procedures."),
        ("B8", "[ICTCLD501 PC 3.1] develop a range of DR solutions."),
        ("B9", "[ICTCLD501 PC 3.2] determine vendor protections and prioritise risks."),
        ("B10", "[ICTCLD501 PC 3.3] assess external insurance protection levels and suitability."),
        ("B11", "[ICTCLD501 PC 3.4] identify other DR solution components."),
        ("B12", "[ICTCLD501 PC 4.1] align DR risk potential to business requirements; [PC 4.2] outline steps, "
                "timelines, key features and service providers."),
        ("B13", "[ICTCLD501 PE 1] develop and evaluate a DR plan covering at least three major risk events; "
                "[PE 3] document how the plan reaches RTO/RPO targets."),
        ("B14", "[ICTCLD501 PC 4.3] document the DR plan according to business needs and requirements; [FS Writing]."),
        ("B15", "[ICTCLD501 KE 1] risk environments; [KE 2] data-analysis methodologies; [KE 3] DR techniques for "
                "cloud; [KE 4] ISO 27001/27002/27031 standards; [KE 5] RTO/RPO standards and techniques; [KE 6] "
                "techniques/methods to monitor and create alerts (disaster detection) — evidenced as written "
                "contextual responses against the student's own plan."),
    ]),
    ("Part C — Presentation (ICTCLD501 element 5)", [
        ("C1", "[ICTCLD501 PC 5.1] conduct verbal walkthrough of the DR plan with required personnel."),
        ("C2", "[ICTCLD501 PC 5.2] seek and respond to feedback as required."),
        ("C3", "[ICTCLD501 PC 5.3] lodge the DR plan according to organisation and legislative protocol."),
        ("C4", "[ICTCLD501 PC 5.4] obtain final sign-off from required personnel."),
        ("C5", "Verbal re-coverage of [ICTCLD501 KE 1–6] and [ICTCLD503 KE 3, 4, 6] — the assessor probes the "
               "student's own design and plan choices (the written KE is the primary evidence in the A12 / B15 "
               "appendices; this is the verbal confirmation)."),
        ("C6", "[ICTCLD501 FS Oral communication] listening/questioning and industry language for the audience."),
    ]),
]

STUDENT_INTRO = [
    ("The engagement and your role", "Heading 2"),
    ("YAT College is a Registered Training Organisation (RTO) based at 175 Cremorne Street, Cremorne VIC. "
     "Following a new partnership with a training institution in India, YAT is expanding its public website — "
     "its global enrolment front-door — to serve prospective students in both countries.", "Assessor text"),
    ("You are an MP Tech Solutions (MTS) consultant reporting to Pat Lin (MTS Senior Consultant). Pat liaises "
     "with Sam Walker, the YAT ICT Manager (Dana Mercer, Marketing & Admissions Manager, is the website "
     "business owner). In AT1 you design the changes the expansion needs and plan the website's disaster "
     "recovery, then present both to YAT ICT management for approval. The approved design and plan are "
     "implemented in AT2.", "Assessor text"),
]

STUDENT_PART_A = [
    ("Part A — Solution Design", "Heading 2"),
    ("Using the YAT Solution Design template (download from the intranet's Templates section), produce a "
     "solution design for the website global expansion covering two strands:", "Assessor text"),
    ("• Web-scale design — the architecture changes that let the website serve a global, anonymous public "
     "audience while scaling network, compute and storage, maintaining availability and security, and "
     "protecting the public attack surface.", "Assessor text"),
    ("• Microservice design — a serverless audit-log microservice that records India-cohort access "
     "events, including the interface/webhook contract.", "Assessor text"),
    ("Treat the data residency requirements as an input that shapes your design — do not write a separate "
     "compliance plan.", "Assessor text"),
    ("Knowledge Evidence appendix (required). At the end of your Solution Design, add a Knowledge Evidence "
     "appendix and answer the following questions in your own words, about your own design (the assessor "
     "will also ask you about these in the Part C Q&A):", "Assessor text"),
    ("• Your design makes four component choices — SQL vs NoSQL, monolith vs microservice, the compute model, "
     "and CDN vs in-memory caching. Explain the function and trade-off of each, and why you chose as you did "
     "for this website.", "Assessor text"),
    ("• Your audit-log microservice is described as highly cohesive and loosely coupled, using a managed "
     "datastore and API / messaging / queuing services. Explain what those properties mean and how your "
     "design exhibits them.", "Assessor text"),
    ("• Which web-scaling principles and technologies does your design apply, and how does each address the "
     "website's global, anonymous, read-heavy load?", "Assessor text"),
    ("Resources for Part A — on the YAT intranet: the Website Global Expansion project (engagement brief, "
     "requirements, ICT manager consultation notes, data residency & sovereignty requirements), the website "
     "specification and the HA-hardened solution design baseline, and the YAT Solution Design template.", "Assessor text"),
    ("Part A is submitted as the populated Solution Design (.docx) with the Knowledge Evidence appendix "
     "completed.", "Assessor text"),
]

STUDENT_PART_B = [
    ("Part B — Disaster Recovery Plan", "Heading 2"),
    ("Using the YAT Disaster Recovery Plan template (download from the intranet's Templates section), produce a "
     "disaster recovery plan for the website designed in Part A. Disaster recovery answers ‘what do we do if "
     "the website goes down’ — keep it focused on recovery, not on the design or the regulatory "
     "requirements.", "Assessor text"),
    ("Your plan must cover: the plan requirements and current recovery position; an impact analysis with your "
     "recovery objectives (RTO/RPO), the data managed, and at least three major risk events; a recovery "
     "strategy with options evaluated and one recommended; and the recovery runbook with steps, timelines and "
     "service providers.", "Assessor text"),
    ("Knowledge Evidence appendix (required). At the end of your DR Plan, add a Knowledge Evidence appendix "
     "and answer the following questions in your own words, about your own plan:", "Assessor text"),
    ("• What is distinctive about the risk environment of a public-facing, cloud-hosted website, and why do "
     "your three major risk events capture it?", "Assessor text"),
    ("• What method did you use to identify and rate the risks, and why is it appropriate here?", "Assessor text"),
    ("• Explain the cloud disaster-recovery techniques available, and why your recommended strategy fits this "
     "website's objectives better than the alternatives.", "Assessor text"),
    ("• Which information-security and business-continuity standards informed your plan, and how does the plan "
     "reflect them?", "Assessor text"),
    ("• Explain RTO and RPO in your own words, the values you set, and the techniques you used to meet each.", "Assessor text"),
    ("• How would a region-level disaster be detected to trigger this plan, and what monitoring and alerting "
     "would you put in place so the on-call team is notified in time to meet the RTO?", "Assessor text"),
    ("A deprecated on-premises DR plan is on the intranet — it was superseded by the move to the cloud and "
     "needs replacing. Use it for context only; your plan is a new cloud DR plan.", "Assessor text"),
    ("Part B is submitted as the populated Disaster Recovery Plan (.docx) with the Knowledge Evidence appendix "
     "completed.", "Assessor text"),
]

STUDENT_PART_C = [
    ("Part C — Presentation for approval", "Heading 2"),
    ("Present your Solution Design and Disaster Recovery Plan to YAT ICT management (role-played by your "
     "assessor as Sam Walker, YAT ICT Manager) and seek approval to proceed to implementation.", "Assessor text"),
    ("The presentation event:", "Assessor text"),
    ("• Duration: 10–15 minutes presenting + 5 minutes of questions, feedback and sign-off.", "Assessor text"),
    ("• Format: in person on campus or via video conference (confirm with your assessor).", "Assessor text"),
    ("• Submit the completed Solution Design and DR Plan at least 48 hours in advance.", "Assessor text"),
    ("During the event you will: walk the manager through your design and DR plan; answer questions during "
     "Q&A (the assessor probes both your design and your underlying knowledge); capture feedback; lodge the "
     "plan; and obtain sign-off.", "Assessor text"),
    ("Part C is submitted as the presentation deck (.pptx) plus the completed feedback and sign-off records.", "Assessor text"),
]

TIPS = [
    ("Tips for success", "Heading 2"),
    ("Keep the three concerns separate. The design (web-scale + microservice) goes in the Solution Design; "
     "recovery goes in the DR Plan; data residency is an input both respect — it is not its own document.", "Assessor text"),
    ("Remember it is a public site. The website serves anonymous visitors from the open internet, so the "
     "web-scale design has to address public exposure (caching at the edge, WAF, bots, DDoS) and "
     "discoverability — not just latency.", "Assessor text"),
    ("Design before you plan recovery. You cannot plan recovery for a system you have not designed — do "
     "Part A first, then Part B for that design.", "Assessor text"),
    ("Make it the simplest thing that meets the requirements. Justify your choices against the requirements; "
     "avoid over-engineering (no unnecessary replication or standby).", "Assessor text"),
    ("Prepare for the Q&A. The assessor will probe both your design and your underlying knowledge — be "
     "ready to explain why you chose each component.", "Assessor text"),
    ("Sign-off is part of the assessment. You cannot achieve Satisfactory on Part C without obtaining the "
     "manager's sign-off at the presentation.", "Assessor text"),
]


def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details ----
    t_details = doc.tables[0]
    set_cell(t_details.rows[1].cells[1], DETAILS["qualification"])
    set_cell(t_details.rows[2].cells[1], DETAILS["units"])
    set_cell(t_details.rows[3].cells[1], DETAILS["task_title"])
    set_cell(t_details.rows[4].cells[1], DETAILS["task_number"])

    # ---- Table 1: Teacher/Assessor instructions ----
    t_instr = doc.tables[1]
    set_cell(instr_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell(instr_row(t_instr, "Task"), TASKS)
    set_cell(instr_row(t_instr, "Resources required"), RESOURCES)
    set_cell(instr_row(t_instr, "Assessment criteria"), CRITERIA_STATEMENT)
    # add a Conditions row at the end (as CL1 did), matching the table style
    cond_row = t_instr.add_row()
    set_cell(cond_row.cells[0], "Assessment Conditions & Setup Requirements")
    for r in cond_row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell(cond_row.cells[1], CONDITIONS)

    # ---- Table 2: Marking Guide ----
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)  # keep 'Assessment criteria' + 'Criteria | Satisfactory?' header rows
    add_section_row(t_mark, "Part A — Solution Design")
    for c in PART_A:
        add_criterion_row(t_mark, c)
    add_section_row(t_mark, "Part B — Disaster Recovery Plan")
    for c in PART_B:
        add_criterion_row(t_mark, c)
    add_section_row(t_mark, "Part C — Presentation")
    for c in PART_C:
        add_criterion_row(t_mark, c)

    # ---- Instructions to Student ----
    doc.add_paragraph("Instructions to Student", style="Heading 1")
    for text, style in (STUDENT_INTRO + STUDENT_PART_A + STUDENT_PART_B + STUDENT_PART_C + TIPS):
        doc.add_paragraph(text, style=style)

    # ---- Benchmark (UoC traceability) ----
    doc.add_paragraph("Marking Benchmark — UoC traceability", style="Heading 1")
    doc.add_paragraph(
        "For each criterion, the unit-of-competency items it evidences. Every PC, PE, KE and FS item allocated "
        "to AT1 in the cluster assessment plan is covered below.", style="Assessor text")
    for part_title, rows in BENCHMARK:
        doc.add_paragraph(part_title, style="Heading 2")
        for cid, uoc in rows:
            p = doc.add_paragraph(style="Assessor text")
            run = p.add_run(f"{cid}  —  ")
            run.bold = True
            p.add_run(uoc)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL2-Cloud-Disaster-Recovery/assessments/AT1/AT1-Design-DR-Plan-Assessor.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
