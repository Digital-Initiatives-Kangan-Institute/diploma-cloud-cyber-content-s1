#!/usr/bin/env python3
"""Build the AT1 Solution Design EXEMPLAR (.docx) — assessor marking reference.

A complete, worked model answer for the YAT Website Global Expansion Solution Design (S1-CL2 AT1
Part A, ICTCLD503 *design*). It mirrors the YAT Solution Design template's section structure
(build_solution_design_template.py) so the exemplar and the student template stay in parity, and
carries a per-section UoC mapping (ex.uoc "Evidences:" line) so it doubles as the marking
reference. Covers the two 503 design strands: the web-scale architecture (§5) and the audit-log
microservice (§6). Disaster recovery is deliberately NOT here — it is the companion DR Plan
(Part B). Data residency is treated as a fixed input, not a deliverable.

Knowledge Evidence (503 KE 3/4/6) is evidenced in the Knowledge Evidence appendix (the "A12"
deliverable), not inline in the body. Element-level approval/sign-off is the AT1 **Part C**
presentation, not a section of this document — §11 keeps only the internal design review (PC 1.6).

The website is YAT's PUBLIC, UNAUTHENTICATED shopfront — anonymous visitors from the open
internet and search engines — which is the deliberate contrast with the LMS practice vehicle
(an authenticated cohort) and is why this design carries CDN/edge, WAF + bot/DDoS protection,
and SEO that an internal authenticated system would not need.

Figures are indicative and internally consistent with the DR Plan exemplar (primary
ap-southeast-2 Sydney; India slice ap-south-1 Mumbai; availability >= 99.9%).

Reuses the docx brand helpers (build_bc_template) + the exemplar helpers (build_s1_cl1_at1_bc_exemplar).

Usage:  python scripts/build_at1_solution_design_exemplar.py [output.docx]
Default: S1-CL2-Cloud-Disaster-Recovery/assessments/AT1/AT1-exemplar-solution-design.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_body_paragraph, add_bullet_list  # noqa: E402
from helpers.docx_tables import add_data_table  # noqa: E402
from helpers.uoc_tags import add_uoc_evidence_tag  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from brand import ADDRESS, CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Solution Design")
    sub = doc.add_paragraph().add_run("YAT Website Global Expansion — Solution Design")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    note = doc.add_paragraph().add_run("Assessor exemplar — internal marking reference (not for distribution to students)")
    note.italic = True; note.font.size = Pt(10); note.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT Website Global Expansion & Disaster Recovery"),
        ("Engagement reference", "YAT-MTS-2026-003"),
        ("Document version", "v1.0"),
        ("Prepared by", "MTS Consultant"),
        ("Date", "[DD/MM/YYYY]"),
        ("Submitted to", "Sam Walker (YAT ICT Manager) · Dana Mercer (Marketing & Admissions Manager) · Pat Lin (MTS Senior Consultant)"),
        ("Related documents", "Functional & Non-Functional Requirements; Data Residency & Sovereignty "
                              "Requirements; Website Cloud Architecture — HA-Hardened Design (baseline); companion Disaster Recovery Plan"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    # 1 Purpose and Scope
    h1("1. Purpose and Scope")
    add_uoc_evidence_tag(doc, "[ICTCLD503 PC 1.1] determine and confirm cloud web-scaling needs")
    add_body_paragraph(doc, "This Solution Design sets out the architecture changes that prepare the YAT public website to "
                 "serve the new India-campus market alongside its existing Australian audience. The website is "
                 "YAT's public, unauthenticated shopfront — its users are anonymous visitors arriving from the "
                 "open internet and from search engines. Two changes are required: (1) scale the website to "
                 "serve a global, anonymous public audience while maintaining availability and security, and "
                 "(2) add a microservice that records India-cohort access events to a store held in India, "
                 "meeting the data-residency requirements. This design is the basis for the Phase 2 "
                 "implementation (the Deployment Report) and for the companion Disaster Recovery Plan.")
    add_body_paragraph(doc, "In scope: the web-tier scaling design (edge delivery, caching, public-traffic protection) and "
                 "the audit-log microservice design. Out of scope: disaster recovery (the companion DR Plan), "
                 "the build and its evidence (the Phase 2 Deployment Report), and the legal determination of "
                 "the residency obligations (owned by YAT compliance — this design treats them as fixed "
                 "inputs).")

    # 2 Design Inputs and Requirements
    h1("2. Design Inputs and Requirements")
    h3("2.1 Inputs")
    add_uoc_evidence_tag(doc, "[ICTCLD503 PC 1.1] confirm web-scaling needs according to business needs")
    add_body_paragraph(doc, "This design is built from the Functional & Non-Functional Requirements, the Data Residency & "
                 "Sovereignty Requirements, the Website Specification, and the existing HA-Hardened Solution "
                 "Design (the baseline). The partnership opens the website to a large new market of anonymous "
                 "prospective students far from the Sydney region, so the website must serve that audience with "
                 "acceptable latency, absorb a higher and spikier public read load, stay discoverable in search, "
                 "and preserve its availability and security under greater public exposure.")
    h3("2.2 Requirements the design must meet")
    add_data_table(doc, ["Ref", "Requirement", "Source"],
              [["R1", "Serve a global, anonymous public audience (AU + India) with acceptable latency", "Functional & Non-Functional Requirements"],
               ["R2", "Scale network, compute and storage as utilisation increases", "Functional & Non-Functional Requirements"],
               ["R3", "Maintain availability (>= 99.9%) and security through the changes", "Requirements / Baseline Design"],
               ["R4", "Protect the public attack surface (web exploits, bots, DDoS) and stay search-discoverable", "Functional & Non-Functional Requirements"],
               ["R5", "India-cohort access logs retained in India (>= 180 days)", "Data Residency Requirements (CERT-In)"],
               ["R6", "Main website data may remain in Australia (light residency slice)", "Data Residency Requirements (DPDP)"]],
              widths=[1.4, 8.6, 6.0])

    # 3 Review of the Existing Architecture
    h1("3. Review of the Existing Architecture")
    add_uoc_evidence_tag(doc, "[ICTCLD503 PC 1.2] review architecture for web application according to business needs")
    add_body_paragraph(doc, "The current architecture (HA-Hardened Solution Design) is a single-region, multi-AZ "
                 "deployment in ap-southeast-2 (Sydney): an Application Load Balancer across two Availability "
                 "Zones, an Auto Scaling group of EC2 instances running the LAMP / CMS stack, a Multi-AZ RDS "
                 "database, and Amazon S3 for media. It is highly available within the region but was sized for "
                 "an Australian audience served entirely from Sydney.")
    add_body_paragraph(doc, "Reviewed layer by layer against R1–R4, four gaps drive this design:")
    add_bullet_list(doc, [
        "Edge / global delivery — anonymous India visitors are served entirely from Sydney; there is no CDN, "
        "so static and cacheable content travels the full distance on every request, hurting latency (R1).",
        "Read-path scale — the Auto Scaling group scales compute, but there is no caching layer to absorb a "
        "larger, spikier public read load off the database (R2).",
        "Public exposure — as an unauthenticated public site reaching a wider audience, it needs protection at "
        "the edge against web exploits, scraping bots and denial-of-service, and it must stay search-"
        "discoverable; the baseline has neither a CDN edge nor edge security in front of it (R4).",
        "Logging residency — all access logging currently lands in Sydney; there is no mechanism to keep "
        "India-cohort access logs in India (R5).",
    ])
    add_body_paragraph(doc, "The compute and load-balancing tiers scale adequately and are retained largely unchanged; the "
                 "gaps are at the edge, the data read path, public-traffic protection, and logging residency — "
                 "addressed in §5 and §6.")

    # 4 Architecture Design
    h1("4. Architecture Design")
    add_uoc_evidence_tag(doc, "[ICTCLD503 PC 1.4] design architecture changes using cloud services (overview) · [ICTCLD503 PC 1.5] scale for a global user base")
    h3("4.1 Assumptions and constraints")
    add_bullet_list(doc, [
        "The baseline is fully implemented: the website runs highly available in ap-southeast-2 (multi-AZ ALB / ASG / RDS, media on S3).",
        "The change is incremental — this design modifies the baseline rather than rebuilding it.",
        "Data residency is a fixed input: India-cohort access logs must be in India; the main website database "
        "may remain in Australia (the light residency slice).",
    ])
    h3("4.2 AWS accounts and regions")
    add_body_paragraph(doc, "ap-southeast-2 (Sydney) remains the primary application and data region. ap-south-1 (Mumbai) "
                 "is introduced for one purpose only — to hold the India-cohort audit log (the residency slice, "
                 "§6). Global, anonymous visitors are reached through edge delivery (§5.3), not by replicating "
                 "the application into other regions. The main website database stays in Sydney; only the access "
                 "logs are residency-bound.")
    h3("4.3 Design relative to the baseline")
    add_body_paragraph(doc, "The baseline's IAM, VPC / network, compute (ALB + Auto Scaling group), database (Multi-AZ "
                 "RDS) and S3 media storage are retained. This design adds: a global edge / caching layer with "
                 "edge security in front of the website (§5), a read / session cache (§5), and the audit-log "
                 "microservice in India (§6). Infrastructure layers not described here are unchanged from the "
                 "baseline.")

    # 5 Web-scale Design
    h1("5. Web-scale Design")
    add_uoc_evidence_tag(doc, "[ICTCLD503 PC 1.3, 1.4, 1.5, 1.6, 1.7] · [ICTCLD503 PE 1, 5]")
    h3("5.1 Web-scaling needs")
    add_body_paragraph(doc, "The drivers (PC 1.1, 1.3): a large new audience of anonymous public visitors, geographically "
                 "distant from Sydney, with a read-heavy and spiky demand profile (marketing campaigns, the "
                 "enrolment-enquiry period). The design must scale network, compute and storage elastically "
                 "(R2) and serve the global, anonymous base with acceptable latency (R1) while preserving "
                 "availability and security under greater public exposure (R3, R4).")
    h3("5.2 Scaling by layer")
    add_body_paragraph(doc, "Services selected to scale the website (PC 1.3) and how each layer scales as utilisation "
                 "increases (PC 1.4):")
    add_data_table(doc, ["Layer", "Service", "How it scales"],
              [["Edge / network", "CloudFront CDN + Route 53 latency routing",
                "Serves cached / static content from edge locations worldwide; India visitors hit the nearest edge, not Sydney"],
               ["Compute", "ALB + EC2 Auto Scaling (retained)",
                "Target-tracking policies add / remove instances on demand; retuned for the higher, spikier public load"],
               ["Caching / read path", "ElastiCache (Redis) — new",
                "Absorbs session and hot-read load off the database as anonymous concurrency grows"],
               ["Storage", "S3 (retained), served via CloudFront",
                "Effectively unlimited media storage; edge-cached for global delivery"],
               ["Database", "RDS Multi-AZ (retained, Sydney)",
                "Vertical scaling, with a read replica as future headroom (not required at launch)"]],
              widths=[3.4, 5.0, 7.6])
    h3("5.3 Global delivery and discoverability")
    add_body_paragraph(doc, "A global, anonymous audience is served (PC 1.5) by pushing content to the edge with "
                 "CloudFront and routing visitors to the lowest-latency entry point with Route 53. The website "
                 "is read-heavy and highly cacheable — marketing pages, the course catalogue, media — so the "
                 "dominant share of traffic is served from the nearest edge and India visitors get acceptable "
                 "latency without replicating the application or its data into India. Because the catalogue and "
                 "marketing pages are the primary acquisition channel for the new market, the edge configuration "
                 "preserves clean, crawlable URLs and fast first-byte times so the site stays discoverable by "
                 "search engines (R4).")
    h3("5.4 Web-scale component choices")
    add_body_paragraph(doc, "Each web-scale component is chosen where it fits:")
    add_data_table(doc, ["Choice", "Decision", "Why"],
              [["SQL vs NoSQL", "SQL (RDS) for website content / enquiry data; NoSQL (DynamoDB) for the audit log",
                "Relational integrity for the CMS and enquiry / application records; append-only high-write key-value for logs (§6)"],
               ["Monolith vs microservice", "Keep the website (CMS) a monolith; split the audit function into a microservice",
                "The CMS performs well and isn't the bottleneck; the new, independently-scaling concern is isolated (§6)"],
               ["Compute model", "Container / VM (EC2 + ASG) for the site; serverless (Lambda) for the microservice",
                "Steady site load suits always-on instances; bursty event load suits scale-to-zero serverless"],
               ["CDN + in-memory store", "CloudFront (CDN) + ElastiCache (in-memory)",
                "Edge delivery for global, anonymous latency; in-memory caching for read scale"]],
              widths=[3.6, 5.8, 6.6])
    h3("5.5 Availability and security maintained")
    add_body_paragraph(doc, "Availability is preserved (PC 1.6) — the multi-AZ ALB / ASG / RDS design is unchanged, and "
                 "CloudFront and ElastiCache are themselves resilient managed services. Security is preserved "
                 "and strengthened for the wider public exposure (R4): TLS end-to-end; AWS WAF at the "
                 "CloudFront edge for the OWASP-class web exploits; AWS Shield and WAF rate-based and bot rules "
                 "to blunt scraping, credential-stuffing of the public application form, and denial-of-service; "
                 "and the origin locked to accept only CloudFront traffic. There is no change to the IAM or "
                 "data-at-rest posture. The design was reviewed against R1–R4 and meets them.")
    add_body_paragraph(doc, "This is one architecture that scales networking, compute and storage for the multi-tier "
                 "website (PE 1) using established web-scaling principles — edge offload, horizontal compute "
                 "scaling, read caching (PE 5) — and the rationale for each change is recorded above, "
                 "satisfying the requirement to document and justify the architecture changes (PC 1.7).")

    # 6 Microservice and Serverless Design
    h1("6. Microservice and Serverless Design")
    add_uoc_evidence_tag(doc, "[ICTCLD503 PC 2.1, 2.2, 2.3, 2.4] · [ICTCLD503 PE 2]")
    h3("6.1 Microservices and data transactions")
    add_body_paragraph(doc, "One microservice is introduced (PC 2.1) — the Access-Log Service. Its single responsibility "
                 "is to record India-cohort access events durably and in-region. The data transaction: on each "
                 "relevant website event (page view, enquiry / application submission) for an India-cohort "
                 "visitor, the website emits an event and the service persists one append-only audit record. It "
                 "never reads back into the website and holds no other state — highly cohesive and loosely "
                 "coupled.")
    h3("6.2 Supporting cloud services")
    add_body_paragraph(doc, "The service is serverless and event-driven, hosted in ap-south-1 (India) (PC 2.2):")
    add_data_table(doc, ["Concern", "Service", "Why"],
              [["Ingress / API", "API Gateway (regional, ap-south-1)", "Receives the website event webhook over HTTPS"],
               ["Messaging / queuing", "SQS queue", "Buffers events so a spike or downstream slowness never blocks or loses one"],
               ["Compute", "AWS Lambda", "Serverless; scales to zero and to spikes, with no servers to manage"],
               ["Persistent storage", "DynamoDB (ap-south-1)", "Append-only NoSQL audit log, kept in India (R5 / CERT-In >= 180 days)"]],
              widths=[3.6, 5.0, 7.4])
    h3("6.3 Microservice architecture")
    add_body_paragraph(doc, "Flow (PC 2.3): website event webhook → API Gateway (ap-south-1) → SQS → Lambda → DynamoDB "
                 "(ap-south-1). The website is the event producer and knows only the webhook contract; "
                 "everything behind it can change independently. The design is highly cohesive and loosely "
                 "coupled (the service does one thing; the website integrates only through the contract); it "
                 "uses a managed database / storage service for persistent data (DynamoDB); and it connects its "
                 "parts with API, messaging and queuing services (API Gateway + SQS). Because the store is in "
                 "ap-south-1, the access logs are resident in India by construction.")
    h3("6.4 Interface / integration contract")
    add_body_paragraph(doc, "The contract the website calls is the single coupling point (PC 2.4): POST /events with a "
                 "JSON body { event_id, occurred_at, user_ref, cohort, event_type, source_ip }, authenticated "
                 "with a signing key. Defining it here lets the website and the service evolve independently. "
                 "The contract is deliberately generic — any event producer that meets it can use the same "
                 "service unchanged, only the payload differing.")
    add_body_paragraph(doc, "This is one microservice architecture for a simple application (PE 2), designed on serverless "
                 "cloud services, with each choice justified above — satisfying the requirement to document and "
                 "justify the microservice design (PC 2.4).")

    # 7 Implementation Sequencing
    h1("7. Implementation Sequencing")
    add_body_paragraph(doc, "The changes are applied in an order that holds availability throughout; the build and its "
                 "evidence are the Phase 2 Deployment Report.")
    add_data_table(doc, ["#", "Change", "Expected impact"],
              [["1", "Stand up the audit-log microservice in ap-south-1 (API Gateway → SQS → Lambda → DynamoDB) and point the website webhook at it", "None — additive"],
               ["2", "Introduce ElastiCache and move the session / read path onto it", "None — cache warms transparently"],
               ["3", "Put CloudFront + WAF / Shield in front of the ALB and switch Route 53 to the distribution", "None — origin unchanged; cutover at DNS"],
               ["4", "Retune the Auto Scaling group scaling policies for the public, spikier load profile", "None"]],
              widths=[0.9, 10.6, 4.5])

    # 8 Simulation and Verification Plan
    h1("8. Simulation and Verification Plan")
    add_body_paragraph(doc, "How the design will be shown to work; the evidence is produced in Phase 2.")
    h3("8.1 Web-scale verification")
    add_bullet_list(doc, [
        "A load test ramps anonymous concurrent visitors to the projected global peak; the Auto Scaling group "
        "scales out and back, edge offload and cache hit-rate are measured, and p95 latency for an "
        "India-region client stays within target (R1, R2).",
        "Confirm multi-AZ failover still works, WAF blocks the OWASP baseline, the rate-based / bot rules "
        "throttle abusive traffic, and the origin rejects non-CloudFront traffic (R3, R4).",
    ])
    h3("8.2 Microservice verification")
    add_bullet_list(doc, [
        "Fire sample webhook events; confirm each is persisted exactly once in the ap-south-1 DynamoDB table.",
        "Force a downstream stall; confirm events queue in SQS and drain without loss.",
        "Confirm the audit store location is ap-south-1 (R5).",
    ])

    # 9 Out of Scope
    h1("9. Out of Scope")
    add_bullet_list(doc, [
        "Disaster recovery — covered by the companion DR Plan (Part B).",
        "The implementation and its evidence — the Phase 2 Deployment Report (AT2).",
        "The legal determination of the residency obligations — owned by YAT compliance; treated here as "
        "fixed inputs.",
        "Replicating the main website database into India — deliberately avoided; the residency slice is the "
        "access logs only, and replication would add cost and complexity for no required benefit.",
    ])

    # 10 References
    h1("10. References")
    add_bullet_list(doc, [
        "Functional & Non-Functional Requirements (YAT Website Global Expansion)",
        "Data Residency & Sovereignty Requirements (DPDP Act 2023; CERT-In Directions 2022)",
        "Website Specification",
        "Website Cloud Architecture — HA-Hardened Design (the baseline)",
        "AWS Well-Architected Framework — Performance Efficiency and Reliability pillars",
    ])

    # 11 Review and Approval
    h1("11. Review and Approval")
    add_uoc_evidence_tag(doc, "[ICTCLD503 PC 1.6] review design as required")
    add_body_paragraph(doc, "Before submission the design was reviewed internally against R1–R4 and the requirements, and "
                 "by the MTS Senior Consultant. Stakeholder approval of the design is obtained at the AT1 "
                 "design-walkthrough presentation, together with the companion DR Plan (Part C); the sign-off "
                 "below records that approval.")
    h3("11.1 Review feedback and author response")
    add_data_table(doc, ["#", "Review feedback", "Author response", "Resulting change"],
              [["1", "Confirm anonymous India visitors are served without replicating data into India",
                "Edge delivery (CloudFront) serves the global base; only the access logs sit in India",
                "Clarified in §4.2 and §5.3"],
               ["2", "Confirm the public exposure is addressed, not just latency",
                "WAF + Shield + rate-based / bot rules added at the edge; SEO discoverability preserved",
                "Recorded in §5.3 and §5.5"],
               ["3", "Confirm the microservice cannot lose events under load",
                "An SQS buffer decouples ingress from the writer, so spikes queue rather than drop",
                "Recorded in §6.2 / §6.3"]],
              widths=[0.8, 5.4, 5.4, 4.0])

    h3("11.2 Sign-off")
    add_data_table(doc, ["Role", "Name", "Decision", "Date"],
              [["Prepared by", "MTS Consultant", "Submitted", "[date]"],
               ["Reviewed by", "Pat Lin (MTS Senior Consultant)", "Approved for submission", "[date]"],
               ["Approved by", "Sam Walker (YAT ICT Manager)", "Approved at design walkthrough — cleared to implement", "[date]"]],
              widths=[3.6, 5.0, 5.4, 2.0])

    # Appendix A — Knowledge Evidence (the Part A "A12" deliverable: written contextual responses)
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1("Appendix A — Knowledge Evidence")
    add_body_paragraph(doc, "Written contextual responses for ICTCLD503 (design knowledge) — the mandatory "
                 "knowledge-evidence location for Part A. Each question asks the candidate to reason about "
                 "their own design; the answers below are the model responses (assessor reference).")

    def ke_qa(tag, question, answer):
        add_uoc_evidence_tag(doc, tag)
        qp = doc.add_paragraph(); qr = qp.add_run(question); qr.bold = True; qr.font.size = Pt(10.5)
        add_body_paragraph(doc, answer)

    ke_qa("[ICTCLD503 KE 3] — functions, benefits and differences of web-scale cloud components "
          "(SQL/NoSQL, monolith/microservice, compute models, CDN/in-memory stores)",
          "Q1. Your design makes four component choices — SQL vs NoSQL, monolith vs microservice, the compute "
          "model, and CDN vs in-memory caching. Explain the function and trade-off of each and why you chose "
          "as you did for this website.",
          "SQL (RDS) gives relational integrity and querying for the CMS content and the enquiry / application "
          "records; NoSQL (DynamoDB) gives cheap, high-write, append-only storage that suits the audit log — "
          "so each data shape goes to the store that fits it. A monolith keeps related code and data together "
          "and is simplest when one component performs well, which the CMS does; a microservice isolates a "
          "single, independently-scaling concern, which is why only the new audit function is split out. For "
          "compute, VMs/containers (EC2 + ASG) suit steady, always-on site load, while serverless (Lambda) "
          "suits the bursty, scale-to-zero event load of the microservice. A CDN caches and serves content "
          "from edge locations close to anonymous global visitors (latency); an in-memory store (ElastiCache) "
          "caches hot reads and sessions close to the application (read scale) — they solve different "
          "bottlenecks, so the design uses both.")

    ke_qa("[ICTCLD503 KE 4] — definitions, functions, features and uses of web-scale cloud infrastructure "
          "(cohesive/loosely-coupled systems; managed database/storage; API, messaging and queuing services)",
          "Q2. Your audit-log service is described as highly cohesive and loosely coupled, using a managed "
          "datastore and API / messaging / queuing services. Explain what those properties mean and how your "
          "design exhibits them.",
          "High cohesion means a component does one well-defined job: the service only records access events "
          "and holds no other state. Loose coupling means components depend on each other only through a "
          "stable interface, not internals: the website and the service share just the webhook contract, so "
          "either can change behind it. A managed database / storage service (DynamoDB) provides durable "
          "persistent storage without running database servers. API, messaging and queuing services connect "
          "the parts: API Gateway is the HTTPS ingress, and an SQS queue buffers events between ingress and "
          "the Lambda writer so a spike or a downstream stall queues rather than drops or blocks. Together "
          "these give a small, resilient, independently-scalable service.")

    ke_qa("[ICTCLD503 KE 6] — web-scaling principles and technologies",
          "Q3. Which web-scaling principles and technologies does your design apply, and how does each address "
          "the website's global, anonymous, read-heavy load?",
          "The design applies three core web-scaling principles. Edge offload: push cacheable content to a CDN "
          "(CloudFront) so most requests are served close to the visitor and never reach the origin — the "
          "biggest win for a read-heavy public site with a distant audience. Horizontal scaling: add and "
          "remove stateless compute on demand (the EC2 Auto Scaling group on target-tracking policies) rather "
          "than relying on a bigger single server, so capacity tracks the spiky public load. Caching the read "
          "path: an in-memory store (ElastiCache) absorbs hot reads and session lookups off the database as "
          "anonymous concurrency grows. Keeping the application stateless behind the load balancer is what "
          "lets the compute and edge tiers scale independently without sticky-session constraints.")

    # Document control
    h1("Document control")
    add_data_table(doc, ["Field", "Value"],
              [["Document version", "v1.0 — Solution Design"],
               ["Author", "MTS Consultant"],
               ["Engagement", "YAT Website Global Expansion & Disaster Recovery"],
               ["Companion document", "Disaster Recovery Plan"],
               ["Implemented by", "Website Global Expansion Deployment Report (Phase 2)"]],
              widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL2-Cloud-Disaster-Recovery/assessments/AT1/AT1-exemplar-solution-design.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
