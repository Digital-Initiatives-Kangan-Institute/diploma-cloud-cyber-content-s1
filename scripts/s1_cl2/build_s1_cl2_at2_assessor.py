#!/usr/bin/env python3
"""Build the S1-CL2 AT2 ASSESSOR instrument (.docx) by populating the Kangan template.

Institutional compliance document (NOT YAT-branded): loads the official Kangan
'Project Assessment - Assessor' template and fills it in, mirroring the CL1 / AT1 shape.
AT2 = Cloud Microservice & IaC Implementation - a single written Deployment Report that
implements the audit-log microservice approved at AT1: the student OPERATES a provided
data-store template, AUTHORS their own infrastructure-as-code template for the microservice
(using the supplied code), deploys and monitors it, and documents it for sign-off.

The worked model lives in the YAT-branded exemplar (build_at2_microservice_iac_exemplar);
this document carries the task instructions, the marking guide with bidirectional UoC
traceability + assessor notes, and the provided artefacts as Appendices A and B. Shared
cell/table helpers are reused from build_s1_cl2_at1_assessor.

Usage:  python scripts/build_at2_assessor.py [output.docx]
Default: S1-CL2-Cloud-Disaster-Recovery/assessments/AT2/AT2-Microservice-IaC-Implementation-Assessor.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_s1_cl2_at1_assessor as a1  # noqa: E402  (shared helpers: set_cell, instr_row, etc.)

from docx import Document  # noqa: E402
from docx.shared import Pt  # noqa: E402

TEMPLATE = "templates/Project Assessment - Assessor.docx"

DETAILS = {
    "qualification": "ICT50220 Diploma of Information Technology",
    "units": [
        "ICTCLD503 Implement web-scale cloud infrastructure",
        "ICTCLD505 Implement cloud infrastructure with code",
    ],
    "task_title": "AT2 - Cloud Microservice & IaC Implementation",
    "task_number": "2 of 2",
}

OVERVIEW = (
    "Students are assessed on implementing the audit-log microservice from the YAT Website Global Expansion design "
    "approved at AT1. Working as an MP Tech Solutions (MTS) consultant in the AWS Academy lab, the student "
    "OPERATES a provided infrastructure-as-code template for the audit-log data store, AUTHORS and deploys their "
    "own infrastructure-as-code template for the microservice (using the supplied code) that writes to that "
    "store, configures monitoring, tests and troubleshoots the build, and documents it in a Deployment Report "
    "submitted to YAT ICT management for sign-off. This is the implementation phase of the engagement; the "
    "design and plan come from AT1."
)

TASKS = [
    "In this project the student implements, in the AWS Academy lab, the audit-log microservice designed and "
    "approved in AT1, and documents it in a single Deployment Report. The build has four streams:",
    "• Operate the provided data-store template - deploy, confirm, update and delete the supplied "
    "DynamoDB audit-store template, and troubleshoot it if it does not deploy as supplied (Appendix A).",
    "• Author your own infrastructure-as-code template - write, parameterise, update and redeploy a "
    "template that provisions the microservice (API, queue and function), wired to the provided data store.",
    "• Build the microservice - deploy and configure the serverless microservice from the supplied code "
    "and webhook contract (Appendix B), and confirm it writes a record to the data store.",
    "• Monitor, document and hand over - configure a metric and alarm, test, produce the infrastructure-"
    "as-code user documentation, and obtain sign-off.",
]

RESOURCES = [
    "Teacher/assessor supplied resources / Access to:",
    "• The YAT scenario site / intranet - the Website Global Expansion project, and the AT1 Solution Design "
    "and Disaster Recovery Plan (the approved design this build implements).",
    "• The YAT Deployment Report template (intranet Templates section).",
    "• The AWS Academy Learner Lab for the build (the deployment region is confirmed and advised before the "
    "assessment runs - see Conditions).",
    "• The provided assessment artefacts (Appendices A and B): the data-store infrastructure-as-code "
    "template, the microservice code, and the webhook payload contract.",
]

CRITERIA_STATEMENT = (
    "To receive a Satisfactory outcome for this assessment task, the student must complete every criterion in "
    "the marking guide below to a satisfactory standard. Where a criterion is not yet satisfactory, the student "
    "is given feedback and a further attempt per the Second attempt provisions."
)

CONDITIONS = [
    "C1 - The YAT scenario site / intranet (including the AT1 Solution Design and DR Plan) is accessible to "
    "the student throughout the assessment.",
    "C2 - AWS Academy Learner Lab access is available for the student to build, deploy and test. The "
    "deployment region is confirmed and advised to students before the assessment runs.",
    "C3 - The assessor supplies the provided artefacts in Appendices A and B: the data-store template, the "
    "microservice code, and the webhook payload contract.",
    "C4 - Design context: the AT1 design places the audit store in the India region (ap-south-1) for data "
    "residency. The deployment region in the lab is the lab's (advised before the build); the region is the "
    "deploy target (the --region flag on the deploy command), not a template parameter, so the same template "
    "targets the lab region or production. The mechanics are assessed, not the geography.",
    "C5 - The assessor role-plays YAT ICT management (Sam Walker, ICT Manager) for the build sign-off.",
]

# Single-part AT2: one Deployment Report. Criteria grouped under one section.
CRITERIA = [
    "D1 - Scope and context: the report states what was deployed, the lab / region setup, and what is in and "
    "out of scope (the microservice + its IaC; the web-tier scaling and DR are AT1, not built here).",
    "D2 - Operate the provided data-store template: reviews it (resources and dependencies), deploys, confirms, "
    "updates and deletes the DynamoDB store from the provided template (Appendix A).",
    "D3 - Author an IaC template: authors a parameterised template provisioning the microservice's resources "
    "(API, queue, function) wired to the provided store, then updates/redeploys and removes it.",
    "D4 - Microservice build: deploys and configures the serverless microservice from the provided code "
    "(Appendix B), and confirms it writes a record to the data store.",
    "D5 - Troubleshooting: diagnoses and fixes the fault in the provided data-store template (if it does not "
    "deploy as supplied) and any microservice errors, with a troubleshooting log.",
    "D6 - Monitoring: configures a metric and alarm (e.g. queue depth, function errors, or a scaling alarm) "
    "with threshold and notification.",
    "D7 - Tooling: uses the cloud console, CLI and/or SDK for the build.",
    "D8 - Configuration decisions: states and justifies the implementation decisions (IaC service, datastore, "
    "decoupling, compute, environment parameterisation).",
    "D9 - Testing and validation: documents the tests run - operate-template (deploy/update/delete), own-"
    "template parameterise/redeploy, microservice functional, and the alarm test.",
    "D10 - IaC user documentation: produces user documentation for operating the stack (deploy/update/delete "
    "with parameters).",
    "D11 - Feedback and sign-off: confirms/seeks/responds to feedback and obtains final sign-off from the "
    "required personnel.",
    "D12 - Knowledge Evidence appendix: written contextual responses, in a Knowledge Evidence appendix, "
    "linking the build to the underlying IaC and microservice concepts and the shared cloud foundations.",
    "D13 - Document quality: the Deployment Report uses the YAT template, plain professional English, and is "
    "complete with cross-referenced evidence appendices.",
]

# Marking benchmark: UoC traceability + an assessor note per criterion (what 'satisfactory' looks like).
BENCHMARK = [
    ("Deployment Report - Cloud Microservice & IaC Implementation (ICTCLD503 build + ICTCLD505)", [
        ("D1", "[ICTCLD503 AC 5] / [ICTCLD505 AC 4] scenario data sources; scope of the implemented build. "
               "Assessor note: scope should name the microservice + its IaC as the build, and exclude the "
               "web-tier scaling and DR (those are AT1 design/plan, not built here)."),
        ("D2", "[ICTCLD505 PC 1.1-1.4] benefits/automation/issues/service selection; [PC 2.1-2.6] template "
               "syntax, review, deploy/update/delete, confirm, remove, troubleshoot; [PE 1]. Assessor note: "
               "must show the FULL lifecycle on the PROVIDED template, not just a deploy - "
               "a review of its resources/dependencies, an update + redeploy, a delete, and - if the supplied "
               "template does not deploy - the fix as well."),
        ("D3", "[ICTCLD505 PC 3.1-3.7] author/deploy/update/parameterise/remove/troubleshoot own template; "
               "[PE 2] create, run and update at least one own template. Assessor note: the "
               "student's OWN template is the microservice (API + queue + function), parameterised (e.g. EnvName; "
               "the region is the deploy target, not a parameter), and references the provided table "
               "(CloudFormation import or a parameter)."),
        ("D4", "[ICTCLD503 PC 3.1-3.3] review design/code, deploy and configure, test and confirm functioning; "
               "[PE 3] deploy a microservice using serverless technologies. Assessor note: the deployed service "
               "must ACTUALLY write a record to the Appendix-A table from a test event - not just stand up."),
        ("D5", "[ICTCLD503 PC 3.4] troubleshoot and fix errors; [ICTCLD505 PC 2.6, 3.7] test/troubleshoot "
               "template errors. Assessor note: the provided template "
               "(Appendix A) does not deploy as supplied; the fault is diagnosed from the error and fixed; the "
               "log shows symptom -> cause -> fix."),
        ("D6", "[ICTCLD503 PC 4.1] set up metrics and trigger scaling alarms. Assessor note: at least one real alarm wired to a "
               "notification; a scaling-relevant metric (queue depth, concurrency, or provisioned-capacity "
               "utilisation) is the strongest evidence."),
        ("D7", "[ICTCLD503 PE 4] / [ICTCLD505 PE 3] use cloud management consoles, SDKs or command line tools."),
        ("D8", "Configuration decisions documented and justified in the report (§5); the underlying "
               "methods/metrics knowledge ([ICTCLD505 KE 10, 11]) is evidenced in the Knowledge Evidence "
               "appendix - see D12."),
        ("D9", "[ICTCLD503 PC 3.3] test and confirm the application functions; [ICTCLD505 PC 2.6, 3.7] test and "
               "troubleshoot. Assessor note: tests cover BOTH templates (operate + own) plus the microservice "
               "function and the alarm."),
        ("D10", "[ICTCLD505 PC 4.1] create user documentation including IaC templates; [PE 4] create user "
                "documentation; [ICTCLD505 FS Writing]."),
        ("D11", "[ICTCLD503 PC 4.2] confirm, seek and respond to feedback; [PC 4.3] obtain final sign-off; "
                "[ICTCLD505 PC 4.2] obtain final sign-off; [ICTCLD505 FS Oral communication]."),
        ("D12", "Written contextual responses in the Knowledge Evidence appendix, each tied to the student's "
                "own build: [ICTCLD505 KE 1-11] (IaC concepts, methods/metrics, testing/debugging) + "
                "[ICTCLD503 KE 1, 2, 5] (industry standards, standard products, testing/debugging)."),
        ("D13", "[ICTCLD505 FS Writing] prepares user documentation in a logical manner using required syntax "
                "and language; [ICTCLD503 FS Writing]."),
    ]),
]

STUDENT_INTRO = [
    ("The engagement and your role", "Heading 2"),
    ("This is the implementation phase of the YAT Website Global Expansion engagement. The design and disaster "
     "recovery plan were produced and approved in AT1; in AT2 you build the audit-log microservice from that "
     "design - operating a provided data-store template and authoring your own template for the service - in "
     "the AWS Academy lab, and document it in a Deployment Report.", "Assessor text"),
    ("You remain an MP Tech Solutions (MTS) consultant reporting to Pat Lin (MTS Senior Consultant), with "
     "Sam Walker (YAT ICT Manager) as the accepting authority for the build.", "Assessor text"),
]

STUDENT_TASK = [
    ("The Deployment Report", "Heading 2"),
    ("Using the YAT Deployment Report template (download from the intranet's Templates section), document the "
     "build you perform in the lab. Your report must cover four streams:", "Assessor text"),
    ("• Operate the provided data-store template (Appendix A) - review it and deploy it (diagnosing and "
     "fixing it first if it does not deploy as supplied), update a setting and redeploy, then delete it.", "Assessor text"),
    ("• Author your own infrastructure-as-code template for the microservice - provision the API, queue "
     "and function, parameterise it (e.g. EnvName), wire it to the data-store table, then update/"
     "redeploy and remove it.", "Assessor text"),
    ("• Build the microservice (Appendix B) - deploy the supplied code behind the API and queue, and "
     "confirm a test event is written to the data-store table.", "Assessor text"),
    ("• Monitor, test, document and hand over - configure a metric and alarm, run and record your tests, "
     "produce the infrastructure-as-code user documentation, and obtain sign-off.", "Assessor text"),
    ("Lab environment: the AWS Academy Learner Lab (and the deployment region) will be advised before you "
     "start. The AT1 design places the store in India (ap-south-1) for residency; the deployment region is the "
     "deploy target (the --region flag), so the same template deploys to the lab region or production - "
     "parameterise environment-specific values (e.g. EnvName), not the region.", "Assessor text"),
    ("Knowledge Evidence (required). In a Knowledge Evidence appendix of your Deployment Report, answer "
     "the following questions in your own words, about your own build:", "Assessor text"),
    ("• Why did you use infrastructure as code rather than provisioning by hand?", "Assessor text"),
    ("• Which infrastructure-as-code service did you use, what are the main sections of one of your templates, "
     "and what tooling deploys it?", "Assessor text"),
    ("• How did you parameterise your template so it can be reused across environments, and how does it "
     "connect to the provided data store?", "Assessor text"),
    ("• How did you test and debug your work - both the templates and the microservice?", "Assessor text"),
    ("• Which industry infrastructure-as-code practices did you apply, and what signals or metrics are used to "
     "manage templates?", "Assessor text"),
    ("• Which industry standards and standard cloud products does your build rely on, and why is your "
     "data-store choice appropriate for an append-only audit log?", "Assessor text"),
    ("Attach your evidence - screenshots, the templates, and the test/troubleshooting evidence - in the "
     "appendices.", "Assessor text"),
    ("AT2 is submitted as the populated Deployment Report (.docx) with the templates and evidence appendices "
     "completed.", "Assessor text"),
]

TIPS = [
    ("Tips for success", "Heading 2"),
    ("Build what AT1 approved. Your report implements the microservice from the AT1 Solution Design - keep it "
     "consistent with that design.", "Assessor text"),
    ("Operate the provided template fully. The data-store template (Appendix A) is one you operate, not author "
     "- show the whole lifecycle: review, deploy, update, delete - and, if it does not deploy as supplied, the "
     "fix too.", "Assessor text"),
    ("Author your own template for the service. The microservice template is yours - parameterise it "
     "(e.g. EnvName) and reference the provided table, so it is clearly your own reusable work.", "Assessor text"),
    ("Capture evidence as you go. Screenshot each stack state, the data-store item written by a test event, "
     "the test request/response and the alarm - rather than reconstructing them afterwards.", "Assessor text"),
    ("Sign-off is part of the assessment. You cannot achieve Satisfactory without obtaining the manager's "
     "sign-off on the build.", "Assessor text"),
]

# ---- Appendix A: the PROVIDED data-store template (the student operates this) ----
DATASTORE_YAML = r"""AWSTemplateFormatVersion: '2010-09-09'

Description: >-
  YAT audit-log data store (PROVIDED - you operate this, you do not author it).
  A single DynamoDB table that holds the append-only access-log records your
  microservice writes. Operate it: review, deploy, update a parameter, and delete.
  It may not deploy as supplied - if so, diagnose the error, fix it, and redeploy.

Parameters:
  EnvName:
    Type: String
    Default: dev
    AllowedValues: [dev, prod]
    Description: Environment suffix used in the table name.

Resources:
  AuditTable:
    Type: AWS::DynamoDB::Table
    Properties:
      TableName: !Sub 'yat-audit-${EnvName}'
      BillingMode: PAY_PER_REQUEST
      AttributeDefinitions:
        - { AttributeName: id, AttributeType: S }
      KeySchema:
        - { AttributeName: event_id, KeyType: HASH }
      Tags:
        - { Key: Project, Value: YAT-Website }
        - { Key: Environment, Value: !Ref EnvName }
        - { Key: DataClassification, Value: audit-log }

Outputs:
  AuditTableName:
    Description: The audit table name - your microservice stack reads this (import or parameter).
    Value: !Ref AuditTable
    Export:
      Name: !Sub 'yat-audit-table-${EnvName}'""".splitlines()

APPENDIX_A = [
    ("Appendix A - Provided data-store template", "Heading 1"),
    ("This infrastructure-as-code template is PROVIDED to you. You OPERATE it - you do not author it: review "
     "it to see what it creates and its dependencies, deploy it, update a parameter and redeploy, and delete "
     "it. It creates the DynamoDB table your microservice writes to.", "Assessor text"),
    ("Note: this template may contain a fault and fail to deploy as supplied. If it does, read the error, "
     "diagnose and fix the problem, then redeploy - and record the symptom, cause and fix in your "
     "troubleshooting log.", "Assessor text"),
    ("datastore.yaml", "Heading 2"),
]
APPENDIX_A_AFTER = [
    ("How to operate it (the region will be advised - substitute it for <REGION>):", "Heading 2"),
]
DATASTORE_OPS = [
    "# Review it first: open the file and note the resources it creates and any dependencies.",
    "",
    "# Deploy:",
    "aws cloudformation deploy --region <REGION> \\",
    "  --stack-name yat-audit-store --template-file datastore.yaml \\",
    "  --parameter-overrides EnvName=dev",
    "",
    "# Update (change a parameter, e.g. EnvName=prod) and redeploy the same way; observe the change.",
    "",
    "# Confirm in the console or with the CLI, then delete (clean teardown):",
    "aws cloudformation delete-stack --region <REGION> --stack-name yat-audit-store",
]

# ---- Appendix B: the PROVIDED microservice code + contract (the student authors the IaC around this) ----
APPENDIX_B_INTRO = [
    ("Appendix B - Provided microservice code and contract", "Heading 1"),
    ("The microservice CODE is PROVIDED to you (you do not write it). You AUTHOR the infrastructure-as-code "
     "template that deploys it - the API, the queue, and this function - and wire it to the data-store table "
     "from Appendix A. The flow is: HTTP API -> SQS queue -> this function -> the Appendix-A DynamoDB table.",
     "Assessor text"),
    ("B.1  Writer function (provided) - handler.py", "Heading 2"),
]
APPENDIX_B_MID = [
    ("B.2  Webhook contract", "Heading 2"),
    ("The single integration point. The website (the event producer) calls this; your microservice receives it. "
     "All six fields are required; a re-sent event with the same event_id is not duplicated.", "Assessor text"),
]
CONTRACT_LINES = r"""POST  {ApiEndpoint}              # {ApiEndpoint} is an Output of your deployed stack
Content-Type: application/json

{
  "event_id":    "3f9a1c2e-...",          # unique id (UUID) - the idempotency key
  "occurred_at": "2026-06-07T01:23:45Z",   # ISO-8601 UTC timestamp
  "user_ref":    "u-48217",                # opaque user reference (not the user's name)
  "cohort":      "IN",                     # "AU" or "IN"  (IN = the India cohort)
  "event_type":  "login",                  # "login" | "course_access" | "assessment_view"
  "source_ip":   "203.0.113.7"             # the client IP the event came from
}

Responses:   200 = accepted and queued for writing      400 = malformed request""".splitlines()
APPENDIX_B_USE = [
    ("How to use it:", "Heading 2"),
    ("• Author a CloudFormation template that provisions the API (HTTP API), an SQS queue, and a Lambda "
     "function running this handler, plus the permissions to read the queue and write the table.", "Assessor text"),
    ("• Set the function's AUDIT_TABLE environment variable to the table from Appendix A (import the "
     "exported value, or pass the table name as a parameter).", "Assessor text"),
    ("• Deploy your template, then POST a sample event (per the contract) and confirm the record appears "
     "in the data-store table.", "Assessor text"),
]

# Provided writer function source (single source of truth for Appendix B.1).
HANDLER_SRC = r'''"""YAT audit-log microservice - SQS-triggered writer Lambda.

Reads access-event messages from the SQS queue and appends each as an immutable
audit record to the DynamoDB table. Idempotent: event_id is the partition key and a
conditional put means a re-delivered message is skipped rather than duplicated.

Data flow:  HTTP API  ->  SQS  ->  (this Lambda)  ->  DynamoDB

Runtime: python3.12 (boto3 is provided by the Lambda runtime - no dependencies to package).
Environment:
    AUDIT_TABLE   the DynamoDB table name (set by the CloudFormation stack).
"""
import json
import os

import boto3
from botocore.exceptions import ClientError

VALID_COHORTS = {"AU", "IN"}
VALID_EVENT_TYPES = {"login", "course_access", "assessment_view"}
REQUIRED_FIELDS = ("event_id", "occurred_at", "user_ref", "cohort", "event_type", "source_ip")


def _validate(record):
    """Return (ok, reason). Rejects anything that isn't a well-formed access event."""
    for field in REQUIRED_FIELDS:
        if not record.get(field):
            return False, f"missing required field: {field}"
    if record["cohort"] not in VALID_COHORTS:
        return False, f"invalid cohort: {record['cohort']}"
    if record["event_type"] not in VALID_EVENT_TYPES:
        return False, f"invalid event_type: {record['event_type']}"
    return True, ""


def handler(event, context=None):
    """SQS event handler. Returns a small summary (and logs it) for observability."""
    table = boto3.resource("dynamodb").Table(os.environ.get("AUDIT_TABLE", "yat-audit-log"))
    processed = skipped = rejected = 0

    for message in event.get("Records", []):
        try:
            body = json.loads(message["body"])
        except (KeyError, json.JSONDecodeError):
            rejected += 1
            print("reject: message body is not valid JSON")
            continue

        ok, reason = _validate(body)
        if not ok:
            rejected += 1
            print(f"reject: {reason}")
            continue

        item = {field: body[field] for field in REQUIRED_FIELDS}
        try:
            table.put_item(Item=item, ConditionExpression="attribute_not_exists(event_id)")
            processed += 1
        except ClientError as exc:
            if exc.response["Error"]["Code"] == "ConditionalCheckFailedException":
                skipped += 1  # already written - idempotent re-delivery
            else:
                raise  # a real failure: let SQS retry / send to the DLQ

    summary = {"processed": processed, "skipped": skipped, "rejected": rejected}
    print(json.dumps(summary))
    return summary
'''
HANDLER_LINES = HANDLER_SRC.splitlines()


def code_block(doc, lines):
    """Render a monospaced code/listing block, one paragraph per line."""
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(ln if ln else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)


def build(path):
    doc = Document(TEMPLATE)

    # ---- Details ----
    t = doc.tables[0]
    a1.set_cell(t.rows[1].cells[1], DETAILS["qualification"])
    a1.set_cell(t.rows[2].cells[1], DETAILS["units"])
    a1.set_cell(t.rows[3].cells[1], DETAILS["task_title"])
    a1.set_cell(t.rows[4].cells[1], DETAILS["task_number"])

    # ---- Instructions ----
    ti = doc.tables[1]
    a1.set_cell(a1.instr_row(ti, "Assessment overview"), OVERVIEW)
    a1.set_cell(a1.instr_row(ti, "Task"), TASKS)
    a1.set_cell(a1.instr_row(ti, "Resources required"), RESOURCES)
    a1.set_cell(a1.instr_row(ti, "Assessment criteria"), CRITERIA_STATEMENT)
    cond = ti.add_row()
    a1.set_cell(cond.cells[0], "Assessment Conditions & Setup Requirements")
    for r in cond.cells[0].paragraphs[0].runs:
        r.bold = True
    a1.set_cell(cond.cells[1], CONDITIONS)

    # ---- Marking guide ----
    tm = doc.tables[2]
    a1.clear_table_rows(tm, 2)
    a1.add_section_row(tm, "Deployment Report")
    for c in CRITERIA:
        a1.add_criterion_row(tm, c)

    # ---- Instructions to Student ----
    doc.add_paragraph("Instructions to Student", style="Heading 1")
    for text, style in (STUDENT_INTRO + STUDENT_TASK + TIPS):
        doc.add_paragraph(text, style=style)

    # ---- Benchmark (UoC traceability + assessor notes) ----
    doc.add_paragraph("Marking Benchmark - UoC traceability and assessor notes", style="Heading 1")
    doc.add_paragraph(
        "For each criterion, the unit-of-competency items it evidences and a short assessor note on what a "
        "satisfactory response looks like. Every PC, PE, KE and FS item allocated to AT2 in the cluster "
        "assessment plan is covered below.", style="Assessor text")
    for part_title, rows in BENCHMARK:
        doc.add_paragraph(part_title, style="Heading 2")
        for cid, uoc in rows:
            p = doc.add_paragraph(style="Assessor text")
            run = p.add_run(f"{cid}  -  "); run.bold = True
            p.add_run(uoc)

    # ---- Appendix A: provided data-store template ----
    for text, style in APPENDIX_A:
        doc.add_paragraph(text, style=style)
    code_block(doc, DATASTORE_YAML)
    for text, style in APPENDIX_A_AFTER:
        doc.add_paragraph(text, style=style)
    code_block(doc, DATASTORE_OPS)

    # ---- Appendix B: provided microservice code + contract ----
    for text, style in APPENDIX_B_INTRO:
        doc.add_paragraph(text, style=style)
    code_block(doc, HANDLER_LINES)
    for text, style in APPENDIX_B_MID:
        doc.add_paragraph(text, style=style)
    code_block(doc, CONTRACT_LINES)
    for text, style in APPENDIX_B_USE:
        doc.add_paragraph(text, style=style)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL2-Cloud-Disaster-Recovery/assessments/AT2/AT2-Microservice-IaC-Implementation-Assessor.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
