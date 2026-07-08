"""Document-level YAT/MTS brand scaffolding for python-docx.

The pieces that turn a blank Document into a branded one: paragraph/heading styles,
the per-section header (ochre disclosure banner) and footer (lockup + page numbers),
and the cover wordmark. Builds on the palette and the low-level styling helpers.
"""
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

from brand import (
    FONT, CHARCOAL, TEAL, TERRACOTTA, OCHRE, GREY, DISCLOSURE,
)
from helpers.docx_styling import shade_paragraph, paragraph_bottom_rule, add_field


def configure_styles(doc):
    """Set the Normal + Title/Heading 1-3 styles to the YAT brand font and colours."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(CHARCOAL)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, bold in (
        ("Title", 30, TEAL, True),
        ("Heading 1", 17, TEAL, True),
        ("Heading 2", 13, TEAL, True),
        ("Heading 3", 11.5, CHARCOAL, True),
    ):
        st = doc.styles[name]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = RGBColor.from_string(color)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def build_header_footer(section):
    """Apply the ochre disclosure-banner header and lockup/page-number footer to a section."""
    # Header: ochre disclosure banner (brand-pack §5.3) on every page.
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(hp, OCHRE)
    r = hp.add_run(DISCLOSURE)
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(CHARCOAL)
    r.bold = True

    # Footer: YAT lockup (left) + page x of y (right), thin teal rule above.
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    paragraph_bottom_rule(fp, TEAL, sz=6)
    fp.text = ""
    tabs = fp.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
    lock = fp.add_run("YAT College")
    lock.bold = True
    lock.font.size = Pt(8.5)
    lock.font.color.rgb = RGBColor.from_string(TEAL)
    fp.add_run("\t")
    pg = fp.add_run("Page ")
    pg.font.size = Pt(8.5)
    pg.font.color.rgb = RGBColor.from_string(GREY)
    add_field(fp, "PAGE")
    of = fp.add_run(" of ")
    of.font.size = Pt(8.5)
    of.font.color.rgb = RGBColor.from_string(GREY)
    add_field(fp, "NUMPAGES")


def wordmark(paragraph):
    """Render the 'YAT COLLEGE' cover wordmark into ``paragraph``."""
    yat = paragraph.add_run("YAT")
    yat.bold = True
    yat.font.size = Pt(34)
    yat.font.name = FONT
    yat.font.color.rgb = RGBColor.from_string(TEAL)
    coll = paragraph.add_run("  COLLEGE")
    coll.bold = True
    coll.font.size = Pt(16)
    coll.font.color.rgb = RGBColor.from_string(TERRACOTTA)
