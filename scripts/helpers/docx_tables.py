"""Table helpers for python-docx documents — branded tables and cell ops.

Two table builders (a blank *template* table with grey placeholder styling, and a
*data* table with filled values) plus low-level cell/row operations. Branded tables
use the YAT palette and the shared shading/border helpers.
"""
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor

from helpers.docx_styling import shade_cell, set_cell_borders
from brand import TEAL, WHITE, GREY


def add_template_table(doc, headers, rows, widths=None):
    """A branded *template* table: teal header, grey-italic placeholder cells.

    Cells beginning with '[' or '$' are rendered as italic grey placeholders.
    ``widths`` (cm, per column) is optional. Returns the table.
    """
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], TEAL)
        set_cell_borders(hdr[i])
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_borders(cells[i])
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
            if val.startswith("[") or val.startswith("$"):
                run.italic = True
                run.font.color.rgb = RGBColor.from_string(GREY)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def add_data_table(doc, headers, rows, widths=None):
    """A branded *data* table: teal header, filled values, money/units right-aligned.

    First column is bold; cells that are money ('$', '+$', '−$') or end in 'hr'/'%'
    are right-aligned. ``widths`` (cm, per column) is optional. Returns the table.
    """
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], TEAL)
        set_cell_borders(hdr[i])
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_borders(cells[i])
            para_c = cells[i].paragraphs[0]
            run = para_c.add_run(val)
            run.font.size = Pt(9)
            money = val.startswith("$") or val.startswith("+$") or val.startswith("−$")
            if money or val.endswith("hr") or val.endswith("%"):
                para_c.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if i == 0:
                run.bold = True
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def set_cell_content(cell, lines):
    """Replace a cell's content with one paragraph per line, keeping its style.

    ``lines`` may be a single string or a list of strings.
    """
    if isinstance(lines, str):
        lines = [lines]
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    for r in list(first.runs):
        r._element.getparent().remove(r._element)
    first.add_run(lines[0])
    style = first.style
    for line in lines[1:]:
        np = cell.add_paragraph()
        np.style = style
        np.add_run(line)


def find_instruction_row(table, label):
    """Return the column-1 cell of the row whose column-0 starts with ``label``.

    Raises KeyError if no row matches.
    """
    for row in table.rows:
        if row.cells[0].text.strip().lower().startswith(label.lower()):
            return row.cells[1]
    raise KeyError(label)


def clear_table_rows(table, keep):
    """Remove all rows after the first ``keep`` rows."""
    for row in table.rows[keep:]:
        row._element.getparent().remove(row._element)


def add_section_row(table, text):
    """Add a 2-column row: bold section label in col 0, empty col 1. Returns the row."""
    row = table.add_row()
    set_cell_content(row.cells[0], text)
    for r in row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell_content(row.cells[1], "")
    return row


def add_criterion_row(table, text):
    """Add a 2-column marking row: criterion in col 0, 'Yes / No' check in col 1. Returns the row."""
    row = table.add_row()
    set_cell_content(row.cells[0], text)
    set_cell_content(row.cells[1], "Yes        No")
    return row
