"""Callout helpers for python-docx templates — applicability notes and convention boxes.

Used by the document *templates* to flag conditional sections ("Applicability: complete
for ...") and to render shaded convention/instruction boxes.
"""
from docx.shared import Pt, RGBColor

from helpers.yat_brand import TERRACOTTA, CREAM, OCHRE, TEAL, CHARCOAL
from helpers.docx_styling import shade_cell, set_cell_borders

#: Run size (points) for the applicability marker.
APPLICABILITY_FONT_SIZE_PT = 9


def add_applicability_note(doc, hint):
    """Append a small terracotta applicability marker for a conditional section.

    Returns the created paragraph.
    """
    p = doc.add_paragraph()
    r = p.add_run(f"Applicability: complete for {hint}; otherwise mark this section "
                  "“Not applicable — [reason]”.")
    r.italic = True
    r.font.size = Pt(APPLICABILITY_FONT_SIZE_PT)
    r.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_convention_box(doc, lines):
    """Append a shaded convention box (cream fill, ochre border).

    ``lines`` is an iterable of (head, body) pairs; ``head`` is bold teal, ``body`` (if any)
    follows in charcoal. Returns the table.
    """
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade_cell(cell, CREAM)
    set_cell_borders(cell, OCHRE, sz=8)
    for i, (head, body) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        hr = p.add_run(head)
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.color.rgb = RGBColor.from_string(TEAL)
        if body:
            br = p.add_run("  " + body)
            br.font.size = Pt(10)
            br.font.color.rgb = RGBColor.from_string(CHARCOAL)
    doc.add_paragraph()
    return t
