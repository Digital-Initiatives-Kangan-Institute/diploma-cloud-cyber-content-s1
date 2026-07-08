"""Body-text helpers for python-docx documents — paragraphs and bullet lists.

Generic building blocks shared across the YAT / MTS document builders. Pure
python-docx: no dependency on a brand palette or on any end-product builder.
"""
from docx.shared import Pt, RGBColor

from brand import GREY

#: Standard body-text run size (points) used across YAT/MTS documents.
BODY_FONT_SIZE_PT = 10.5
#: Run size (points) for italic guidance / placeholder text.
GUIDANCE_FONT_SIZE_PT = 9.5


def add_body_paragraph(doc, text):
    """Append ``text`` as a single body paragraph at body size.

    Args:
        doc:  a python-docx ``Document`` (anything exposing ``add_paragraph``).
        text: the paragraph text.

    Returns:
        The created paragraph.
    """
    p = doc.add_paragraph()
    p.add_run(text).font.size = Pt(BODY_FONT_SIZE_PT)
    return p


def add_guidance_text(doc, text):
    """Append italic grey guidance text (template authoring instructions).

    Returns the created paragraph.
    """
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = RGBColor.from_string(GREY)
    r.font.size = Pt(GUIDANCE_FONT_SIZE_PT)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_response_placeholder(doc, label="[ Write your response here ]"):
    """Append an italic grey response placeholder paragraph.

    Returns the created paragraph.
    """
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.italic = True
    r.font.color.rgb = RGBColor.from_string(GREY)
    return p


def add_bullet_list(doc, items):
    """Append each string in ``items`` as a 'List Bullet' paragraph at body size.

    Args:
        doc:   a python-docx ``Document`` (anything exposing ``add_paragraph``).
        items: an iterable of strings — one bullet paragraph per item.

    Returns:
        list: the created paragraphs, in order (empty list if ``items`` is empty).
    """
    paragraphs = []
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item).font.size = Pt(BODY_FONT_SIZE_PT)
        paragraphs.append(p)
    return paragraphs
