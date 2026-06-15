"""Low-level python-docx / OOXML styling helpers — shading, borders, fields.

These manipulate the underlying OOXML (``w:shd``, ``w:tcBorders``, ``w:pBdr``,
``w:fldChar``) that python-docx doesn't expose directly. Shared by the document
and table builders. Colour defaults come from the YAT brand palette.
"""
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from helpers.yat_brand import STONE, TEAL, GREY


def _shade(element_pr, hex_fill):
    """Apply a solid ``w:shd`` fill to a paragraph- or cell-properties element."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    element_pr.append(shd)


def shade_paragraph(paragraph, hex_fill):
    """Fill a paragraph's background with ``hex_fill`` (6-digit RGB)."""
    _shade(paragraph._p.get_or_add_pPr(), hex_fill)


def shade_cell(cell, hex_fill):
    """Fill a table cell's background with ``hex_fill`` (6-digit RGB)."""
    _shade(cell._tc.get_or_add_tcPr(), hex_fill)


def set_cell_borders(cell, hex_color=STONE, sz=4):
    """Draw single borders on all four edges of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), hex_color)
        borders.append(e)
    tcPr.append(borders)


def paragraph_bottom_rule(paragraph, hex_color=TEAL, sz=12):
    """Add a single bottom border (a horizontal rule) under a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_field(paragraph, instr, placeholder=""):
    """Insert a Word field code (e.g. ``TOC``, ``PAGE``, ``NUMPAGES``).

    ``placeholder`` text (grey) is shown until the user updates the field in Word.
    """
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(it); run._r.append(sep)
    if placeholder:
        ph = paragraph.add_run(placeholder)
        ph.font.color.rgb = RGBColor.from_string(GREY)
    run2 = paragraph.add_run()
    run2._r.append(end)
