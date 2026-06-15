"""Assessor-exemplar evidence markers for python-docx.

Two helpers used in worked exemplars / instruments: one marks a section deferred
("Not applicable — <reason>"), the other describes evidence that would be captured
("[SCREENSHOT] — should show ..."). Assessment artefacts only.
"""
from docx.shared import Pt, RGBColor

from helpers.yat_brand import TERRACOTTA, TEAL, GREY

#: Run size (points) for the "Not applicable" marker.
NOT_APPLICABLE_FONT_SIZE_PT = 10.5
#: Run size (points) for described-evidence placeholders.
EVIDENCE_FONT_SIZE_PT = 9.5


def add_not_applicable(doc, reason):
    """Append a bold terracotta 'Not applicable — <reason>' paragraph. Returns it."""
    p = doc.add_paragraph()
    r = p.add_run(f"Not applicable — {reason}")
    r.bold = True
    r.font.size = Pt(NOT_APPLICABLE_FONT_SIZE_PT)
    r.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_described_evidence(doc, kind, desc):
    """Append a described-evidence placeholder: teal '[kind]' tag + grey italic desc.

    e.g. add_described_evidence(doc, "SCREENSHOT", "should show CREATE_COMPLETE").
    Returns the paragraph.
    """
    p = doc.add_paragraph()
    tag = p.add_run(f"[{kind}] ")
    tag.bold = True
    tag.font.size = Pt(EVIDENCE_FONT_SIZE_PT)
    tag.font.color.rgb = RGBColor.from_string(TEAL)
    d = p.add_run("— " + desc)
    d.font.size = Pt(EVIDENCE_FONT_SIZE_PT)
    d.italic = True
    d.font.color.rgb = RGBColor.from_string(GREY)
    p.paragraph_format.space_after = Pt(6)
    return p
