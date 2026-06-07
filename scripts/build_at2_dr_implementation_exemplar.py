#!/usr/bin/env python3
"""Build the S1-CL2 AT2 Deployment Report EXEMPLAR (.docx) — assessor marking reference.

A worked model answer for the YAT LMS Global Expansion *DR Implementation* (S1-CL2 AT2),
implementing the AT1-approved design: the audit-log microservice provisioned as
infrastructure-as-code, with monitoring, built and tested in the AWS Academy lab. Covers
ICTCLD503 element 3 (microservice build), ICTCLD505 (IaC — operate predefined templates +
author own), the monitoring strand, and the build documentation + sign-off.

Reuses the Deployment Report shell (as CL1 AT2) but with a serverless + IaC build narrative
(§4) and testing (§6) instead of the EC2/RDS infrastructure sub-sections. Evidence is
described, not captured: "[SCREENSHOT — should show ...]". Retains UoC `Evidences:` tags and
the §8 Knowledge Evidence responses (assessor layers an org template omits).

Reuses the docx brand helpers (build_bc_template) and prose/table helpers (build_bc_exemplar).

Usage:  python scripts/build_at2_dr_implementation_exemplar.py [output.docx]
Default: S1-CL2-Cloud-Disaster-Recovery/assessments/AT2/AT2-exemplar-deployment-report.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_bc_template as bc   # noqa: E402
import build_bc_exemplar as ex   # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def na(doc, reason):
    p = doc.add_paragraph()
    r = p.add_run(f"Not applicable — {reason}")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    p.paragraph_format.space_after = Pt(6)
    return p


def ev(doc, kind, desc):
    """Described-evidence placeholder, e.g. [SCREENSHOT — should show ...]."""
    p = doc.add_paragraph()
    tag = p.add_run(f"[{kind}] ")
    tag.bold = True
    tag.font.size = Pt(9.5)
    tag.font.color.rgb = RGBColor.from_string(bc.TEAL)
    d = p.add_run("— " + desc)
    d.font.size = Pt(9.5)
    d.italic = True
    d.font.color.rgb = RGBColor.from_string(bc.GREY)
    p.paragraph_format.space_after = Pt(6)
    return p


def code(doc, lines):
    """A short monospaced code/snippet block."""
    for ln in lines:
        p = doc.add_paragraph()
        r = p.add_run(ln)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(bc.CHARCOAL)
        p.paragraph_format.space_after = Pt(0)


def build(path):
    doc = Document()
    bc.configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    bc.build_header_footer(sec)

    # ---- COVER ----
    bc.wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(bc.ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(bc.GREY)
    bc.paragraph_bottom_rule(doc.add_paragraph(), bc.TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Deployment Report")
    sub = doc.add_paragraph().add_run("YAT LMS Global Expansion — DR Implementation")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    note = doc.add_paragraph().add_run("Assessor exemplar — internal marking reference (not for distribution to students)")
    note.italic = True; note.font.size = Pt(10); note.font.color.rgb = RGBColor.from_string(bc.GREY)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT LMS Global Expansion & Disaster Recovery — Implementation"),
        ("Engagement reference", "YAT-MTS-2026-002"),
        ("Report version", "v1.0"),
        ("Prepared by", "MTS Consultant"),
        ("Date submitted", "[DD/MM/YYYY]"),
        ("Submitted to", "Sam Walker (YAT ICT Manager) · Pat Lin (MTS Senior Consultant)"),
        ("Related documents", "LMS Global Expansion Solution Design (the approved design); the Disaster "
                              "Recovery Plan"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        bc.set_cell_borders(cells[0]); bc.set_cell_borders(cells[1]); bc.shade_cell(cells[0], bc.CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    bc.add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    # 1 Executive Summary
    h1("1. Executive Summary")
    ex.uoc(doc, "[ICTCLD503 PE 3] · [ICTCLD505 PE 1, 2] — overview of the implemented build")
    ex.para(doc, "This Deployment Report documents the Phase 2 implementation of the YAT LMS Global Expansion: "
                 "the audit-log microservice, provisioned as infrastructure-as-code, with monitoring, built and "
                 "tested in the AWS Academy lab. It implements the Solution Design approved at AT1.")
    ex.para(doc, "The build was delivered in three streams: operating a set of predefined infrastructure-as-code "
                 "templates to stand up the supporting resources; authoring a parameterised template that "
                 "provisions the microservice itself; and deploying and testing the serverless microservice "
                 "(API Gateway → SQS → Lambda → DynamoDB). Monitoring alarms were configured, the build was "
                 "tested and troubleshooted, documented for operational handover, and signed off.")

    # 2 Engagement Context
    h1("2. Engagement Context")
    ex.uoc(doc, "[ICTCLD503 AC 5] · [ICTCLD505 AC 4] (scenario data sources used)")
    ex.para(doc, "Following approval of the design at AT1, MTS implemented the audit-log microservice for the "
                 "India-cohort access logs. The implementation was built in the AWS Academy lab environment "
                 "using the services, code elements and reference data supplied for the engagement.")
    ex.para(doc, "Lab / region note: the design places the audit log in the India region (ap-south-1). The AWS "
                 "Academy lab does not offer ap-south-1, so us-west-2 stands in for the India region and "
                 "us-east-1 for the primary; the mechanics (cross-region separation, IaC, the serverless stack) "
                 "are real, only the geography is simulated. Region is a template parameter, so the same "
                 "templates target ap-south-1 in production.")

    # 3 Scope of Deployment
    h1("3. Scope of Deployment")
    ex.para(doc, "In scope: the audit-log microservice and the infrastructure-as-code that provisions it — "
                 "operating the predefined templates, authoring the student's own template, the serverless build "
                 "(API Gateway, SQS, Lambda, DynamoDB), the monitoring alarms, the IaC user documentation, and "
                 "the build sign-off.")
    ex.para(doc, "Out of scope: the LMS web-tier scaling changes (CloudFront / caching — a separate LMS "
                 "workstream, designed but not built here); the disaster-recovery drill itself (covered by the "
                 "DR Plan); and production cutover. The microservice integrates with the LMS only through the "
                 "webhook contract, so no change to the live LMS was required.")

    # 4 Build / Change Narrative
    h1("4. Build and Implementation Narrative")
    h3("4.1 Lab environment and access")
    ex.uoc(doc, "[ICTCLD505 AC 1, 5, 7, 8] (vendor platform, IDE, CLI/SSH, console)")
    ex.para(doc, "Work was performed in the AWS Academy lab via the console, the AWS CLI, and an IDE. A least-"
                 "privilege execution role was used for the Lambda function (write to the one DynamoDB table and "
                 "read from the queue only); no long-lived credentials were created.")
    ev(doc, "SCREENSHOT", "should show the AWS Academy lab session and the IAM execution role's attached policy.")

    h3("4.2 Operating the predefined IaC templates")
    ex.uoc(doc, "[ICTCLD505 PC 1.1–1.4] · [ICTCLD505 PC 2.1–2.6] · [ICTCLD505 PE 1, 3] · [ICTCLD505 KE 3, 4, 5, 6, 7, 10]")
    ex.para(doc, "Benefits and service selection (PC 1.1–1.4): infrastructure-as-code was chosen over manual "
                 "console provisioning because it makes the stack repeatable, reviewable and parameterised by "
                 "region — exactly what a cross-region recovery design needs. AWS CloudFormation was selected as "
                 "the IaC service: it is native to the platform, needs no extra tooling, and the lab supports it.")
    ex.para(doc, "Template syntax and review (PC 2.1, 2.2, KE 5): the supplied templates were read to confirm "
                 "the resources they create and their dependencies — a DynamoDB table, an SQS queue, and the "
                 "IAM roles — before deploying. The CloudFormation template language (Resources, Parameters, "
                 "Outputs, intrinsic functions such as !Ref and !GetAtt) is summarised in §8 Q2.")
    ex.para(doc, "Deploy, update, delete (PC 2.3–2.5, PE 1): the predefined templates were used to deploy the "
                 "resources, confirm them in the console, update a stack (changing a parameter and observing the "
                 "change set), and then delete a stack to confirm clean teardown.")
    ex.etable(doc, ["Action", "Command / method", "Result"],
              [["Deploy", "aws cloudformation deploy --template-file base.yaml --stack-name yat-audit-base",
                "CREATE_COMPLETE; resources visible in the console"],
               ["Update", "change a parameter; aws cloudformation deploy (change set reviewed)",
                "UPDATE_COMPLETE; only the intended resource changed"],
               ["Delete", "aws cloudformation delete-stack --stack-name yat-audit-base",
                "DELETE_COMPLETE; all resources removed"]],
              widths=[2.6, 8.4, 5.0])
    ex.para(doc, "Confirm and troubleshoot (PC 2.4, 2.6, KE 7): deployments were confirmed in the console and "
                 "with describe-stacks; template errors (an invalid attribute and a circular dependency) were "
                 "diagnosed from the stack events and fixed — see the troubleshooting log in §6.6.")
    ev(doc, "SCREENSHOT", "should show a CloudFormation stack at CREATE_COMPLETE with its resources tab.")

    h3("4.3 Authoring the infrastructure-as-code template")
    ex.uoc(doc, "[ICTCLD505 PC 3.1–3.7] · [ICTCLD505 PE 2] · [ICTCLD505 KE 8, 9]")
    ex.para(doc, "A template was authored from scratch to provision the microservice as a related set of "
                 "resources (PC 3.2): the DynamoDB audit table, the SQS queue, the Lambda function, the API "
                 "Gateway, and the execution role. It was deployed, then updated to add the API Gateway stage "
                 "(PC 3.3), and parameterised so the same template deploys to either region and to a named "
                 "environment without editing the body (PC 3.5) — the key code-reuse outcome (KE 8).")
    code(doc, [
        "Parameters:",
        "  TargetRegion:   { Type: String, Default: us-west-2 }   # ap-south-1 in production",
        "  EnvName:        { Type: String, Default: dev }",
        "  RetentionDays:  { Type: Number, Default: 180 }         # CERT-In retention",
        "Resources:",
        "  AuditTable:  { Type: AWS::DynamoDB::Table, Properties: { BillingMode: PAY_PER_REQUEST, ... } }",
        "  EventQueue:  { Type: AWS::SQS::Queue }",
        "  WriterFn:    { Type: AWS::Lambda::Function, Properties: { Runtime: python3.12, ... } }",
    ])
    ex.para(doc, "The template follows industry practice (KE 9): parameters not hard-coding, least-privilege "
                 "roles, tagging (Project/Environment/Owner/DataClassification), and Outputs that export the "
                 "API endpoint. It was redeployed with a different parameter set to confirm reuse, then the "
                 "stack was removed cleanly (PC 3.6); template errors were tested and fixed (PC 3.7).")
    ev(doc, "FILE", "the authored template is attached in Appendix B (audit-service.yaml).")

    h3("4.4 The audit-log microservice")
    ex.uoc(doc, "[ICTCLD503 PC 3.1–3.4] · [ICTCLD503 PE 3, 4] · [ICTCLD503 KE 5]")
    ex.para(doc, "Reviewing the design and code (PC 3.1): the supplied microservice code elements (the Lambda "
                 "handler and the webhook payload schema) were reviewed against the AT1 design before "
                 "deployment. The data flow is: an event webhook → API Gateway → SQS → Lambda → DynamoDB.")
    ex.para(doc, "Deploy and configure (PC 3.2): the serverless services were deployed and configured via the "
                 "authored template and the console — the API Gateway POST /events route, the queue as the "
                 "Lambda event source, and the Lambda writing append-only items to DynamoDB. The webhook "
                 "contract is the single integration point:")
    code(doc, [
        "POST /events",
        "{ event_id, occurred_at, user_ref, cohort, event_type, source_ip }   # signed with a shared key",
    ])
    ex.para(doc, "Test and confirm functioning (PC 3.3): sample events were posted and confirmed persisted "
                 "exactly once in DynamoDB (see §6.3). Troubleshoot (PC 3.4): an IAM permission error and a "
                 "payload-validation bug were diagnosed from the Lambda logs and fixed (§6.6). Tooling (PE 4): "
                 "the build used the console, the AWS CLI and the SDK.")
    ev(doc, "SCREENSHOT", "should show a DynamoDB item written by the Lambda after a test POST to /events.")

    h3("4.5 Monitoring and alarms")
    ex.uoc(doc, "[ICTCLD503 PC 4.1] · [ICTCLD501 KE 6]")
    ex.para(doc, "Metrics and alarms (PC 4.1): CloudWatch alarms were configured on the queue depth "
                 "(ApproximateNumberOfMessagesVisible — a scaling/backlog signal), the Lambda error rate and "
                 "throttles, and the DynamoDB write throttles, each notifying an SNS topic. The queue-depth "
                 "alarm is the scaling signal — a sustained backlog indicates the writer is not keeping up. This "
                 "is also the detection technique the DR plan relies on for the audit path (KE 6).")
    ev(doc, "SCREENSHOT", "should show the CloudWatch alarms list with the queue-depth alarm in OK state.")

    # 5 Configuration Decisions
    h1("5. Configuration Decisions")
    ex.uoc(doc, "[ICTCLD505 KE 10, 11] · [ICTCLD503 KE 1, 2] · [ICTCLD505 KE 1, 2] (justified choices, standards)")
    ex.etable(doc, ["Decision", "Choice", "Why"],
              [["IaC service", "AWS CloudFormation", "Native to the platform; no extra tooling; supported in the lab"],
               ["Datastore", "DynamoDB, on-demand (PAY_PER_REQUEST)", "Append-only audit log; scales to zero and to spikes; no capacity to manage"],
               ["Decoupling", "SQS between API Gateway and Lambda", "Buffers spikes so events queue rather than drop"],
               ["Compute", "Lambda (python3.12)", "Serverless; scales with event volume; no servers to run"],
               ["Region", "Template parameter (us-west-2 / ap-south-1)", "Same template deploys to the lab stand-in or production India region"]],
              widths=[3.4, 5.0, 7.6])
    ex.para(doc, "Industry standards (KE 1, 2 / 503 + 505): the build uses standard managed cloud services and "
                 "the JSON/HTTPS and IAM standards common to the platform; storage is the managed NoSQL "
                 "key-value service (DynamoDB) appropriate to an append-only log.")

    # 6 Testing and Validation
    h1("6. Testing and Validation")
    ex.uoc(doc, "[ICTCLD503 PC 3.3, 3.4] · [ICTCLD505 PC 2.6, 3.7]")
    h3("6.1 IaC deploy / update / delete verification")
    ex.para(doc, "Confirmed each predefined template deploys, updates (via a reviewed change set) and deletes "
                 "cleanly, with the console and describe-stacks as evidence (PC 2.4).")
    ev(doc, "SCREENSHOT", "stack events showing CREATE_COMPLETE → UPDATE_COMPLETE → DELETE_COMPLETE.")
    h3("6.2 Own-template parameterise and redeploy")
    ex.para(doc, "Deployed the authored template, then redeployed it with a second parameter set (different "
                 "EnvName and region) to confirm configuration reuse without editing the template (PC 3.4, 3.5).")
    h3("6.3 Microservice functional test")
    ex.para(doc, "Posted sample events to /events and confirmed each was written exactly once to DynamoDB, with "
                 "the expected attributes, and that an unsigned request was rejected (PC 3.3).")
    ev(doc, "SCREENSHOT", "the test request, the 200 response, and the resulting DynamoDB item.")
    h3("6.4 Durability test")
    ex.para(doc, "Forced a downstream stall (throttled the writer) and confirmed events queued in SQS and then "
                 "drained without loss once the writer recovered — the decoupling works as designed.")
    h3("6.5 Monitoring / alarm test")
    ex.para(doc, "Drove the queue depth above the threshold and confirmed the alarm transitioned to ALARM and "
                 "notified the SNS topic, then returned to OK as the backlog cleared (PC 4.1).")
    h3("6.6 Troubleshooting log")
    ex.etable(doc, ["Symptom", "Cause", "Fix"],
              [["Stack CREATE_FAILED — circular dependency", "Role referenced the queue and the queue policy referenced the role",
                "Split the policy into a separate AWS::IAM::Policy resource"],
               ["Lambda AccessDenied on PutItem", "Execution role lacked dynamodb:PutItem on the table",
                "Scoped a PutItem statement to the table ARN"],
               ["Events accepted but not stored", "Payload-validation bug rejected valid cohort values",
                "Fixed the schema check; re-tested (§6.3)"]],
              widths=[5.2, 5.4, 5.4])

    # 7 Operational Handover
    h1("7. Operational Handover")
    h3("7.1 Infrastructure-as-code user documentation")
    ex.uoc(doc, "[ICTCLD505 PC 4.1] · [ICTCLD505 PE 4] · [ICTCLD505 FS Writing]")
    ex.para(doc, "User documentation was produced so YAT ICT can operate the stack: prerequisites; how to deploy, "
                 "update and delete the stack with parameters; the parameters and their meaning; the Outputs "
                 "(the API endpoint); and how to roll back. It is included as Appendix B with the templates.")
    h3("7.2 Access and runbook references")
    ex.para(doc, "Access is via the YAT AWS account roles; the operational runbook references the alarms (§4.5) "
                 "and the redeploy procedure (§7.1).")
    h3("7.3 Known limitations and what's next")
    ex.bullets(doc, [
        "Built in the AWS Academy lab with us-west-2 standing in for ap-south-1; production deploys to "
        "ap-south-1 by changing the region parameter.",
        "The LMS web-tier scaling (CloudFront / caching) is designed but not built here — a separate LMS "
        "workstream.",
    ])
    h3("7.4 Documentation filing")
    ex.para(doc, "The report, the templates and the evidence are filed in YAT ICT shared documentation under the "
                 "Records Management Policy.")
    h3("7.5 Feedback record")
    ex.uoc(doc, "[ICTCLD503 PC 4.2] confirm, seek and respond to feedback with required personnel")
    ex.etable(doc, ["Feedback received", "From", "Response", "Action"],
              [["Confirm the audit data stays in the India-region store", "Sam Walker (YAT ICT Manager)",
                "Confirmed — the table is provisioned in the India region (us-west-2 in the lab); region is a parameter",
                "Noted in §2 and §4.3"],
               ["Confirm events can't be lost under load", "Pat Lin (MTS Senior Consultant)",
                "The SQS buffer queues spikes; demonstrated in §6.4", "Recorded in §6.4"]],
              widths=[5.0, 3.0, 4.6, 2.8])
    h3("7.6 Sign-off")
    ex.uoc(doc, "[ICTCLD503 PC 4.3] · [ICTCLD505 PC 4.2] · [ICTCLD505 FS Oral communication]")
    ex.etable(doc, ["Role", "Name", "Date", "Acceptance"],
              [["Prepared by", "MTS Consultant", "[date]", "Submitted"],
               ["Reviewed by", "Pat Lin (MTS Senior Consultant)", "[date]", "Approved for submission"],
               ["Approved by", "Sam Walker (YAT ICT Manager)", "[date]",
                "Build accepted — microservice and IaC implemented, tested and documented; cleared for production deployment"]],
              widths=[4.0, 5.0, 2.5, 4.0])

    # 8 Knowledge Evidence Responses
    h1("8. Knowledge Evidence Responses")
    ex.uoc(doc, "[ICTCLD505 KE 3–11] · [ICTCLD503 KE 1, 2, 5] · [ICTCLD505 KE 1, 2] (written contextual KE)")
    h3("Q1 — Why infrastructure as code, not manual provisioning?")
    ex.para(doc, "[KE 3] In my build, IaC made the stack repeatable and reviewable, let me tear down and "
                 "rebuild identically, and — most importantly for this design — parameterised the region so the "
                 "same template deploys to the lab stand-in and to the production India region. Manual console "
                 "provisioning would not be reproducible across regions or auditable.")
    h3("Q2 — IaC services, template syntax, and tooling")
    ex.para(doc, "[KE 4, 5, 6] I used AWS CloudFormation (other options include Terraform and the AWS CDK). A "
                 "CloudFormation template declares Parameters, Resources and Outputs in YAML, with intrinsic "
                 "functions such as !Ref and !GetAtt for dependencies. The tooling to execute it is the AWS CLI "
                 "(cloudformation deploy / delete-stack) and the console.")
    h3("Q3 — Parameterisation for reuse")
    ex.para(doc, "[KE 8] My template exposes TargetRegion, EnvName and RetentionDays as parameters, so one "
                 "template deploys dev and prod, and the lab region or the India region, without editing the "
                 "body — the configuration is supplied at deploy time.")
    h3("Q4 — Testing and debugging")
    ex.para(doc, "[KE 7 / 503 KE 5] I debugged template errors from the CloudFormation stack events (a circular "
                 "dependency, an invalid attribute) and microservice errors from the Lambda CloudWatch logs (an "
                 "IAM AccessDenied and a payload-validation bug) — see the §6.6 log.")
    h3("Q5 — Industry IaC practices and metrics")
    ex.para(doc, "[KE 9, 10, 11] Industry practice I applied: parameters over hard-coding, least-privilege "
                 "roles, consistent tagging, Outputs for integration values, and change sets to preview updates. "
                 "Stack status and drift detection are the standard signals for managing templates.")
    h3("Q6 — Industry standards and standard products")
    ex.para(doc, "[503 KE 1, 2 / 505 KE 1, 2] The build uses standard cloud-industry technologies — HTTPS/JSON "
                 "for the API, IAM for access, and the managed serverless and NoSQL services (Lambda, DynamoDB, "
                 "SQS, API Gateway). DynamoDB is the standard managed key-value store, appropriate for an "
                 "append-only audit log.")

    # Appendices
    h1("Appendix A — Build evidence (screenshots)")
    ex.para(doc, "Collected console screenshots referenced through §4–§6 (lab session, stack CREATE_COMPLETE, "
                 "DynamoDB item, alarms, test request/response).")
    h1("Appendix B — Infrastructure-as-code templates and user documentation")
    ex.para(doc, "The authored template (audit-service.yaml) and the predefined templates operated, plus the "
                 "IaC user documentation (deploy/update/delete with parameters).")
    h1("Appendix C — Test and troubleshooting evidence")
    ex.para(doc, "The functional, durability and alarm test evidence, and the troubleshooting log (§6.6).")
    h1("Appendix D — Reflections")
    ex.para(doc, "Short reflections on what the student would do differently and what the build confirmed about "
                 "the AT1 design.")

    # Document control
    h1("Document control")
    ex.etable(doc, ["Field", "Value"],
              [["Document version", "v1.0 — Deployment Report (DR Implementation)"],
               ["Author", "MTS Consultant"],
               ["Engagement", "YAT LMS Global Expansion & Disaster Recovery"],
               ["Implements", "LMS Global Expansion Solution Design (AT1)"],
               ["Companion document", "Disaster Recovery Plan"]],
              widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL2-Cloud-Disaster-Recovery/assessments/AT2/AT2-exemplar-deployment-report.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
