#!/usr/bin/env python3
"""Build the YAT / MTS Disaster Recovery Plan template (.docx) from the brand pack.

ONE generic, in-world DR Plan template — the student-facing superset, derived from the AT1
DR Plan exemplar by keeping the exact section skeleton (so it covers the full ICTCLD501
element sequence) while stripping the worked answers and the assessor "Evidences:" UoC tags,
and swapping in guidance + response space and empty tables. Sections a particular plan
doesn't need carry the "Not applicable — reason" convention rather than being deleted.

Usage:  python scripts/build_dr_plan_template.py [output.docx]
Default: ../diploma-cloud-cyber-website/public/templates/YAT-DR-Plan-Template.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_guidance_text, add_response_placeholder  # noqa: E402
from helpers.docx_callouts import add_convention_box  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from helpers.docx_tables import add_template_table  # noqa: E402
from helpers.yat_brand import ADDRESS, CHARCOAL, CREAM, GREY, OCHRE, TEAL  # noqa: E402
from helpers.yat_docx_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def add_convention_box(doc, lines):
    """Shaded convention box."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade_cell(cell, CREAM)
    set_cell_borders(cell, OCHRE, sz=8)
    for i, (head, body) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        hr = p.add_run(head)
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.color.rgb = RGBColor.from_string(TEAL)
        if body:
            br = p.add_run("  " + body)
            br.font.size = Pt(10)
            br.font.color.rgb = RGBColor.from_string(CHARCOAL)
    doc.add_paragraph()


def build(path):
    doc = Document()
    configure_styles(doc)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Disaster Recovery Plan")
    sub = doc.add_paragraph().add_run("[ Engagement / system name ]")
    sub.font.size = Pt(15); sub.italic = True; sub.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()

    cover_rows = [
        ("Engagement", "[ Engagement name ]"),
        ("Engagement reference", "[ Reference ]"),
        ("Plan version", "[ e.g. v1.0 ]"),
        ("Prepared by", "[ Consultant name ]"),
        ("Consultant role", "[ Role, e.g. MTS Consultant ]"),
        ("Date submitted", "[ DD/MM/YYYY ]"),
        ("Submitted to", "[ Acceptance authority / sponsor ]"),
        ("Related documents", "[ e.g. the companion solution design ]"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in cover_rows:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        vr = cells[1].paragraphs[0].add_run(v); vr.font.size = Pt(10); vr.italic = True
        vr.font.color.rgb = RGBColor.from_string(GREY)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- HOW TO USE + CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("How to use this template", style="Heading 1")
    add_convention_box(doc, [
        ("One plan format for any disaster recovery plan.", "This is the standard YAT disaster recovery "
         "plan — complete each section for your engagement."),
        ("Complete every section.", "Where a section does not apply, mark it “Not applicable — [reason]” "
         "rather than deleting it, so the plan proves nothing was overlooked."),
        ("The impact analysis is the heart.", "Your risk events, recovery objectives, and the recovery "
         "strategy that follows are what the plan is judged on — make them specific to your system."),
        ("Validate before you finalise.", "Walk the plan through with the required personnel, record their "
         "feedback and your response, and obtain sign-off before lodging it."),
    ])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Executive Summary")
    add_guidance_text(doc, "Write this last. A ≤ 1-page summary the reader sees first: the system this plan protects "
                     "and why disaster recovery is needed now; the major risk events you cover; your recovery "
                     "objectives (RTO and RPO); and, in a line, your recommended recovery strategy. ~250–350 words.")
    add_response_placeholder(doc)

    h1("2. Engagement Context and Scope")
    add_guidance_text(doc, "Set the context (≤ ½ page): the system and the prior work this plan builds on; why "
                     "region-level disaster recovery is now required; and the DR plan requirements set by the "
                     "business — your RTO and RPO targets and any data-residency constraint. State clearly what "
                     "is in scope and what is out of scope (e.g. in-region failure already handled, "
                     "application-layer recovery owned by others).")
    add_response_placeholder(doc)

    h1("3. Current Recovery Position")
    h3("3.1 Existing recovery arrangements")
    add_guidance_text(doc, "Describe the recovery arrangements already in place (e.g. in-region high availability, "
                     "automated backups, the organisation's backup/retention and recovery procedures) — and, "
                     "importantly, what they do not protect against. Note whether any organisational recovery "
                     "plan already exists.")
    add_response_placeholder(doc)
    h3("3.2 Vendor (cloud provider) provisions and SLAs")
    add_guidance_text(doc, "Describe the cloud provider's relevant disaster-recovery provisions and service levels you "
                     "will rely on: the regional / Availability Zone architecture, the native backup and "
                     "recovery services, the support-response commitment, and the shared-responsibility split "
                     "between the provider and the organisation.")
    add_response_placeholder(doc)

    h1("4. Impact Analysis")
    h3("4.1 Recovery objectives (RTO / RPO)")
    add_guidance_text(doc, "State your Recovery Time Objective (RTO) and Recovery Point Objective (RPO) and explain "
                     "what each means for this system. These targets drive your backup cadence and recovery "
                     "strategy, so justify them against the business's tolerance for downtime and data loss.")
    add_response_placeholder(doc)
    h3("4.2 Data managed")
    add_guidance_text(doc, "Estimate the amount of data managed and describe the classes of data and their "
                     "sensitivity / security level (e.g. personal information, records, logs), including any "
                     "data subject to residency or other legal obligations.")
    add_response_placeholder(doc)
    h3("4.3 Risk assessment")
    add_guidance_text(doc, "Identify at least three major risk events to the system, with their likelihood, impact, "
                     "and resulting severity, and note what each risk drives in your recovery design. Briefly "
                     "state the method you used to identify and rate them. Complete the register below (add rows "
                     "as needed).")
    add_template_table(doc, ["#", "Risk event", "Likelihood", "Impact", "Severity", "What it drives"],
             [["RE1", "[ … ]", "[ ]", "[ ]", "[ ]", "[ … ]"],
              ["RE2", "[ … ]", "[ ]", "[ ]", "[ ]", "[ … ]"],
              ["RE3", "[ … ]", "[ ]", "[ ]", "[ ]", "[ … ]"],
              ["…", "[ add further risks ]", "[ ]", "[ ]", "[ ]", "[ … ]"]],
             widths=[0.9, 5.0, 2.0, 1.8, 1.8, 4.0])
    h3("4.4 Plan exclusions")
    add_guidance_text(doc, "State what this plan deliberately excludes and why (e.g. single-component / AZ failures "
                     "handled by existing HA; application-layer faults owned by another team; assumptions about "
                     "what remains available).")
    add_response_placeholder(doc)
    h3("4.5 Recording the analysis")
    add_guidance_text(doc, "State where the outcomes of this impact analysis are recorded and under what "
                     "organisational policy (e.g. records-management / documentation procedures), so the "
                     "analysis is retained, auditable, and available to the recovery team.")
    add_response_placeholder(doc)

    h1("5. Recovery Strategy")
    h3("5.1 Options evaluated")
    add_guidance_text(doc, "Evaluate the range of recovery strategies available for your platform against your "
                     "RTO/RPO targets and cost. Complete the comparison (the standard cloud DR tiers are a "
                     "useful starting point).")
    add_template_table(doc, ["Option", "RTO", "RPO", "Relative cost", "Fit against the objectives"],
             [["[ … ]", "[ ]", "[ ]", "[ ]", "[ … ]"],
              ["[ … ]", "[ ]", "[ ]", "[ ]", "[ … ]"],
              ["[ … ]", "[ ]", "[ ]", "[ ]", "[ … ]"],
              ["[ … ]", "[ ]", "[ ]", "[ ]", "[ … ]"]],
             widths=[4.4, 2.0, 2.4, 2.4, 4.3])
    h3("5.2 Recommended approach")
    add_guidance_text(doc, "State and justify your recommended recovery strategy — how it meets the RTO and RPO, how "
                     "it uses backups / replication, and why it is the most appropriate (and cost-proportionate) "
                     "option.")
    add_response_placeholder(doc)
    h3("5.3 Vendor protections and risk prioritisation")
    add_guidance_text(doc, "Describe the provider-native protections your strategy relies on, and prioritise the risks "
                     "from §4 — which risk your strategy addresses first, and how.")
    add_response_placeholder(doc)
    h3("5.4 Insurance")
    add_guidance_text(doc, "Assess any external insurance protection (e.g. cyber / business-interruption) and its "
                     "suitability — as a complement to, not a substitute for, technical recovery. Note what "
                     "residual risk it transfers and any notification / evidence obligations. If none applies, "
                     "mark Not applicable.")
    add_response_placeholder(doc)
    h3("5.5 Other recovery components")
    add_guidance_text(doc, "Identify the other components your plan needs beyond the technical recovery — e.g. a "
                     "recovery runbook, a declaration / escalation path, a communications plan, DNS failover, "
                     "the order of recovery, and any relevant continuity standards.")
    add_response_placeholder(doc)

    h1("6. The Disaster Recovery Plan")
    h3("6.1 Detection and alerting")
    add_guidance_text(doc, "Describe how a disaster is detected and how an alert leads to a declaration — the "
                     "monitoring / alarms used, the threshold that triggers, and who is notified.")
    add_response_placeholder(doc)
    h3("6.2 Recovery steps")
    add_guidance_text(doc, "Set out the steps of the recovery plan in order, with the owner and a target (cumulative) "
                     "time for each, including key features and the service providers involved. The total must "
                     "sit within your RTO. Make sure the steps visibly address the risks you prioritised in §4 "
                     "— your highest-severity risk should be answered first. Complete the runbook below (add "
                     "rows as needed).")
    add_template_table(doc, ["Step", "Action", "Owner", "Target (cumulative)"],
             [["1", "[ … ]", "[ … ]", "[ ]"],
              ["2", "[ … ]", "[ … ]", "[ ]"],
              ["3", "[ … ]", "[ … ]", "[ ]"],
              ["…", "[ add further steps ]", "[ … ]", "[ ]"]],
             widths=[1.0, 8.3, 3.0, 3.2])
    h3("6.3 Meeting the recovery objectives")
    add_guidance_text(doc, "Explain how the plan meets your RTO and RPO — reference the backup cadence for the RPO and "
                     "the step timeline for the RTO.")
    add_response_placeholder(doc)

    h1("7. Plan Validation and Approval")
    h3("7.1 Walkthrough")
    add_guidance_text(doc, "Describe the verbal walkthrough of this plan with the required personnel — who attended, "
                     "and what scenario you walked through to confirm the steps and timings.")
    add_response_placeholder(doc)
    h3("7.2 Feedback record")
    add_guidance_text(doc, "Record the feedback you sought and received, your response, and the resulting action.")
    add_template_table(doc, ["Feedback received", "From", "Your response", "Resulting action"],
             [["[ … ]", "[ … ]", "[ … ]", "[ … ]"],
              ["[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[5.2, 3.0, 4.3, 3.0])
    h3("7.3 Lodgement")
    add_guidance_text(doc, "State where the approved plan is lodged and under what protocol (records management / "
                     "change management), and the review cadence.")
    add_response_placeholder(doc)
    h3("7.4 Sign-off")
    add_guidance_text(doc, "Obtain final sign-off from the required personnel.")
    add_template_table(doc, ["Role", "Name", "Date", "Acceptance"],
             [["Prepared by", "[ … ]", "", ""],
              ["Reviewed by", "[ … ]", "", ""],
              ["Approved by (acceptance authority)", "[ … ]", "", ""]],
             widths=[5.5, 4.5, 2.5, 3.0])

    h1("Document control")
    add_template_table(doc, ["Field", "Value"],
             [["Document version", "[ v1.0 ]"],
              ["Author", "[ Name, role ]"],
              ["Engagement", "[ Engagement name ]"],
              ["Date submitted", "[ DD/MM/YYYY ]"],
              ["Companion document", "[ the solution design this plan accompanies ]"],
              ["Successor document", "[ the deployment report that implements the plan ]"]],
             widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website/public/templates/YAT-DR-Plan-Template.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
