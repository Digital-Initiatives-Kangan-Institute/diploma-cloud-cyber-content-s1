#!/usr/bin/env python3
"""Build the YAT / MTS Solution Design template(s) (.docx) from the brand pack.

The third document type alongside Business Case and Deployment Report — the cloud
solution/architecture design that a Deployment Report then implements. Generalised from
the AT3 HA Design outline: the architecture-design core is generic; the change-specific
sections (review of an existing baseline, implementation sequencing/rollback, simulation
plan) carry the "Not applicable — reason" convention so a greenfield design can omit them.
Assessment scaffolding (UoC notes, student ID) stripped.

Two scoped variants from ONE definition, selected by the `web_scale` flag:
  * web_scale=False — the base template (CL1 shape: sections 1-9 + Document control).
    This output is byte-for-byte the long-standing template; CL1 relies on its shape.
  * web_scale=True  — adds two sections after Architecture Design: §5 Web-scale Design
    and §6 Microservice and Serverless Design (the rest shift down to §7-§11). Used by
    CL2 AT1 Part A, where the design includes web-scaling + a microservice. The base
    sections are identical between the two — only the two extra sections differ — so there
    is a single source of truth and no drift.

Usage:  python scripts/build_solution_design_template.py   (builds BOTH variants)
"""
import sys
from pathlib import Path

sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_guidance_text, add_response_placeholder  # noqa: E402
from helpers.docx_callouts import add_applicability_note, add_convention_box  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from helpers.docx_tables import add_template_table  # noqa: E402
from brand import ADDRESS, CREAM, GREY, TEAL  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path, web_scale=False):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Solution Design")
    sub = doc.add_paragraph().add_run("[ Solution / initiative name ]")
    sub.font.size = Pt(15); sub.italic = True; sub.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()
    cover = [
        ("Engagement", "[ Engagement name ]"),
        ("Engagement reference", "[ Reference ]"),
        ("Document version", "[ e.g. v1.0 ]"),
        ("Prepared by", "[ Consultant name, role ]"),
        ("Date", "[ DD/MM/YYYY ]"),
        ("Submitted to", "[ Acceptance authority / sponsor ]"),
        ("Related documents", "[ requirements + any baseline design this supersedes ]"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        vr = cells[1].paragraphs[0].add_run(v); vr.font.size = Pt(10); vr.italic = True
        vr.font.color.rgb = RGBColor.from_string(GREY)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS + convention ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("How to use this template", style="Heading 1")
    applic_lead = "a given design may mark" if web_scale else "a greenfield design usually marks"
    applic_examples = ("reviewing an existing baseline, "
                       + ("web-scale design, microservice / serverless design, " if web_scale else "")
                       + "implementation sequencing on a live system, and the simulation plan.")
    add_convention_box(doc, [
        ("One design for any cloud solution.", "Use this for a greenfield architecture or for hardening / "
         "changing an existing one."),
        ("Complete every section.", "Where a section does not apply, mark it “Not applicable — [reason]” "
         "rather than deleting it."),
        ("Sections flagged with an applicability note", f"(terracotta) are the ones {applic_lead} "
         "Not applicable — " + applic_examples),
        ("This design is implemented by a Deployment Report.", "Business Case (why) → Solution Design "
         "(what / how) → Deployment Report (what was built)."),
    ])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Purpose and Scope")
    add_guidance_text(doc, "What this design covers and what it doesn't (≤ ½ page): the solution being designed, "
                     "the objective it serves, and what is in / out of scope.")
    add_response_placeholder(doc)

    h1("2. Design Inputs and Requirements")
    h3("2.1 Inputs")
    add_guidance_text(doc, "The source documents this design is built from — the approved business case, the "
                     "requirements, the workload/application specification, and (if hardening an existing "
                     "system) the baseline design and the report of what was actually built.")
    add_response_placeholder(doc)
    h3("2.2 Requirements the design must meet")
    add_guidance_text(doc, "State the targets the design is held to — availability, recovery, performance, "
                     "security, data residency. Add rows as needed.")
    add_template_table(doc, ["Requirement / metric", "Target"],
             [["Availability", "[ e.g. 99.9% ]"],
              ["RPO (acceptable data loss)", "[ e.g. ≤ 1 hour ]"],
              ["RTO (time to recover)", "[ e.g. ≤ 4 hours ]"],
              ["Support response (severity-1)", "[ e.g. ≤ 1 hour ]"],
              ["Data residency", "[ e.g. Australia ]"],
              ["[ add others ]", "[ … ]"]],
             widths=[9.0, 7.0])

    h1("3. Review of the Existing Architecture")
    add_applicability_note(doc, "designs that harden or change an existing system (a greenfield design has no baseline to review)")
    h3("3.1 Architecture review")
    add_guidance_text(doc, "Review the existing architecture against the requirements, layer by layer (compute, "
                     "load balancing, database, network, monitoring) — where each currently meets or fails "
                     "the targets.")
    add_response_placeholder(doc)
    h3("3.2 Single points of failure")
    add_template_table(doc, ["Component", "Failure mode", "Consequence"],
             [["[ … ]", "[ … ]", "[ … ]"], ["[ … ]", "[ … ]", "[ … ]"]],
             widths=[5.0, 5.0, 6.0])
    h3("3.3 Recovery objectives — current state")
    add_template_table(doc, ["Component", "Current RPO", "Current RTO", "Meets target?"],
             [["[ … ]", "[ … ]", "[ … ]", "[ Yes/No ]"], ["[ … ]", "[ … ]", "[ … ]", "[ Yes/No ]"]],
             widths=[5.5, 3.5, 3.5, 3.5])
    h3("3.4 Components requiring vertical scaling")
    add_template_table(doc, ["Component", "Vertical scale required for", "Availability impact during scale"],
             [["[ … ]", "[ … ]", "[ … ]"]],
             widths=[5.0, 5.0, 6.0])
    h3("3.5 Review findings summary")
    add_guidance_text(doc, "Summarise the gap between the current architecture and the requirements: what's met, "
                     "what isn't, and which components drive the gap (≤ 250 words).")
    add_response_placeholder(doc)

    h1("4. Architecture Design")
    add_guidance_text(doc, "The design proper. For a greenfield design this is the full architecture; for a change, "
                     "describe the design relative to the baseline.")
    h3("4.1 Assumptions and constraints")
    add_response_placeholder(doc)
    h3("4.2 AWS account and region")
    add_guidance_text(doc, "Region(s) and account; note any data-residency constraint and any new region (e.g. a "
                     "cross-Region backup destination).")
    add_response_placeholder(doc)
    for n, title, hint in [
        ("4.3", "Identity and access management (IAM)", "groups/roles/users, MFA, instance profiles — or, for a change, any additions"),
        ("4.4", "Network topology", "VPC, subnets, AZ distribution, gateways, route tables (use the subnet table below)"),
        ("4.5", "Compute (EC2 + Auto Scaling)", "instance type, ASG capacity + AZ spread, scaling policy, health checks"),
        ("4.6", "Load balancing (ALB)", "the load balancer, target group, listener, AZ coverage"),
        ("4.7", "Database (RDS)", "engine + version, instance class, storage, encryption, backups, Multi-AZ (if designed)"),
        ("4.8", "Storage (EBS + S3)", "volumes, buckets, encryption, public-access settings, cross-Region copy (if designed)"),
        ("4.9", "Security", "the tiered security-group model, encryption in transit + at rest, any HA-related adjustments"),
        ("4.10", "Monitoring", "the alarms and the service-level/availability tracking (use the alarm table below)"),
        ("4.11", "Naming and tagging conventions", "the tagging scheme (Project, Environment, Owner, CostCentre, DataClassification, AZ)"),
        ("4.12", "Backup", "the backup baseline and any cross-Region copy / retention design"),
    ]:
        h3(f"{n} {title}")
        add_guidance_text(doc, f"Cover: {hint}.")
        if n == "4.4":
            add_template_table(doc, ["Subnet", "CIDR", "AZ", "Purpose"],
                     [["[ … ]", "[ … ]", "[ … ]", "[ … ]"], ["[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
                     widths=[4.0, 3.5, 4.0, 4.5])
        elif n == "4.10":
            add_template_table(doc, ["Alarm", "Metric", "Threshold", "Triggers"],
                     [["[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
                     widths=[4.0, 4.0, 3.5, 4.5])
        else:
            add_response_placeholder(doc)
    h3("4.13 Recovery objectives — designed state")
    add_template_table(doc, ["Component", "Designed RPO", "Designed RTO", "Notes"],
             [["[ … ]", "[ … ]", "[ … ]", "[ … ]"], ["Overall service", "[ … ]", "[ … ]", "[ meets target ]"]],
             widths=[5.0, 3.5, 3.5, 4.0])
    h3("4.14 Components requiring vertical scaling — designed state")
    add_template_table(doc, ["Component", "Vertical scale required for", "Availability impact (designed)"],
             [["[ … ]", "[ … ]", "[ … ]"]],
             widths=[5.0, 5.0, 6.0])
    h3("4.15 Single points of failure removed")
    add_applicability_note(doc, "designs that hardened an existing system (cross-reference §3.2)")
    add_template_table(doc, ["SPOF (from §3.2)", "Mitigation in this design"],
             [["[ … ]", "[ … ]"], ["[ … ]", "[ … ]"]],
             widths=[7.0, 9.0])

    # ---- Optional web-scale + microservice sections (CL2 variant) ----
    if web_scale:
        h1("5. Web-scale Design")
        add_applicability_note(doc, "designs that scale a web application for growth or a global / multi-region user base (a fixed-capacity design marks this Not applicable)")
        h3("5.1 Web-scaling needs")
        add_guidance_text(doc, "The scaling drivers this design must meet — load growth, a global / multi-region user "
                         "base, and the peak-demand profile — confirmed against the business needs.")
        add_response_placeholder(doc)
        h3("5.2 Scaling by layer")
        add_guidance_text(doc, "For each layer, the cloud service used and how it scales network, compute and storage "
                         "as utilisation increases.")
        add_template_table(doc, ["Layer", "Service", "How it scales"],
                 [["Edge / network", "[ e.g. CDN, latency routing ]", "[ … ]"],
                  ["Compute", "[ … ]", "[ … ]"],
                  ["Caching / read path", "[ … ]", "[ … ]"],
                  ["Storage", "[ … ]", "[ … ]"],
                  ["Database", "[ … ]", "[ … ]"]],
                 widths=[4.0, 5.5, 6.5])
        h3("5.3 Global delivery")
        add_guidance_text(doc, "How a global / multi-region user base is served with acceptable latency — edge "
                         "delivery, routing, and what is served from the edge versus the origin.")
        add_response_placeholder(doc)
        h3("5.4 Web-scale component choices")
        add_guidance_text(doc, "Justify the component choices across the web-scale options: SQL vs NoSQL databases, "
                         "monolithic vs microservice architecture, virtual / container / serverless compute, and "
                         "content delivery networks and in-memory data stores — each chosen where it fits.")
        add_response_placeholder(doc)
        h3("5.5 Availability and security maintained")
        add_guidance_text(doc, "Confirm the scaling changes preserve availability and security, and note the review "
                         "that checked this against the requirements.")
        add_response_placeholder(doc)

        h1("6. Microservice and Serverless Design")
        add_applicability_note(doc, "designs that include a microservice or serverless component (an infrastructure-only design marks this Not applicable)")
        h3("6.1 Microservices and data transactions")
        add_guidance_text(doc, "Identify the microservice(s) and the data transaction(s) each handles, according to "
                         "business needs — one responsibility per service.")
        add_response_placeholder(doc)
        h3("6.2 Supporting cloud services")
        add_guidance_text(doc, "The cloud services that support the microservice architecture, by concern.")
        add_template_table(doc, ["Concern", "Service", "Why"],
                 [["Ingress / API", "[ … ]", "[ … ]"],
                  ["Messaging / queuing", "[ … ]", "[ … ]"],
                  ["Compute", "[ … ]", "[ … ]"],
                  ["Persistent storage", "[ … ]", "[ … ]"]],
                 widths=[4.0, 5.5, 6.5])
        h3("6.3 Microservice architecture")
        add_guidance_text(doc, "The architecture and event / data flow — how the services are composed, and how the "
                         "design keeps them highly cohesive and loosely coupled.")
        add_response_placeholder(doc)
        h3("6.4 Interface / integration contract")
        add_guidance_text(doc, "The API / webhook contract the rest of the system integrates through — endpoint, "
                         "method and payload schema — the single coupling point between the service and its "
                         "callers.")
        add_response_placeholder(doc)

    # ---- Sections after the design core (numbers shift when web_scale) ----
    seq = 7 if web_scale else 5
    sim, oos, ref, rev = seq + 1, seq + 2, seq + 3, seq + 4

    h1(f"{seq}. Implementation Sequencing")
    add_applicability_note(doc, "deployments into a running system, where order and rollback matter")
    add_guidance_text(doc, "The order the changes are applied, with per-change duration, expected impact, a "
                     "verification step, and a rollback if it fails. State the total window and buffer.")
    add_template_table(doc, ["#", "Change", "Duration", "Expected impact", "Verification", "Rollback"],
             [["1", "[ … ]", "[ … ]", "[ none ]", "[ … ]", "[ … ]"],
              ["2", "[ … ]", "[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[0.8, 4.0, 2.0, 3.0, 3.0, 3.0])

    h1(f"{sim}. Simulation and Verification Plan")
    add_applicability_note(doc, "resilience / high-availability designs (the plan the Deployment Report then executes)")
    h3(f"{sim}.1 Failure simulation plan")
    add_template_table(doc, ["#", "Simulation", "Method", "Expected outcome", "Verification"],
             [["F1", "[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[0.8, 3.0, 3.5, 4.5, 4.0])
    h3(f"{sim}.2 Resize simulation plan")
    add_template_table(doc, ["#", "Simulation", "Method", "Expected outcome", "Verification"],
             [["R1", "[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[0.8, 3.0, 3.5, 4.5, 4.0])
    h3(f"{sim}.3 Verification criteria")
    add_guidance_text(doc, "The success criteria — what evidence will demonstrate the design works as intended.")
    add_response_placeholder(doc)

    h1(f"{oos}. Out of Scope")
    add_guidance_text(doc, "What this design deliberately does not address, and why (e.g. cross-Region active-active "
                     "DR, application-layer HA, items owned by another team).")
    add_response_placeholder(doc)

    h1(f"{ref}. References")
    add_guidance_text(doc, "The source documents and standards informing this design (requirements, baseline "
                     "design, application specification, reference architectures, industry standards), with "
                     "external sources cited with access dates.")
    add_response_placeholder(doc)

    h1(f"{rev}. Review and Approval")
    add_guidance_text(doc, "The completed design is submitted to the accepting authority / superior for review. "
                     "The reviewer records their feedback below; the author records how each point was "
                     "addressed; the reviewer then signs off on the design before it is implemented.")
    h3(f"{rev}.1 Reviewer feedback and author response")
    add_template_table(doc, ["#", "Reviewer feedback / comment", "Author response", "Resulting change"],
             [["1", "[ … ]", "[ … ]", "[ … ]"],
              ["2", "[ … ]", "[ … ]", "[ … ]"],
              ["3", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[0.8, 6.0, 5.2, 4.0])
    h3(f"{rev}.2 Sign-off")
    add_guidance_text(doc, "The accepting authority records their decision. “Approved with comments” means the "
                     "design proceeds subject to the changes recorded above.")
    add_template_table(doc, ["Role", "Name", "Decision", "Date", "Signature"],
             [["Prepared by (author)", "[ name, role ]", "—", "[ DD/MM/YYYY ]", ""],
              ["Reviewed and approved by (accepting authority)", "[ name, role ]",
               "[ Approved / Approved with comments / Rejected ]", "[ DD/MM/YYYY ]", ""]],
             widths=[4.6, 3.2, 3.6, 2.4, 2.2])

    h1("Document control")
    add_template_table(doc, ["Field", "Value"],
             [["Document version", "[ v1.0 ]"],
              ["Author", "[ Name, role ]"],
              ["Engagement", "[ Engagement name ]"],
              ["Supersedes", "[ any baseline design this replaces ]"],
              ["Implemented by", "[ the Deployment Report that builds this design ]"],
              ["Approval status", "[ Pending / Approved / Rejected with comments ]"]],
             widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    base_dir = Path("../diploma-cloud-cyber-website-s1/public/templates")
    build(base_dir / "YAT-Solution-Design-Template.docx", web_scale=False)
    build(base_dir / "YAT-Solution-Design-Template-Webscale.docx", web_scale=True)
