#!/usr/bin/env python3
"""Build the YAT / MTS Deployment Report template(s) (.docx) from the brand pack.

ONE generic, in-world Deployment Report template — the superset that serves any
deployment, from a greenfield foundation build to a full HA hardening. Built from the
HA report outline (the larger of the two), with AT2's foundation-only sections folded
back in, and assessment-only scaffolding (KE Q&A, reflections, student ID, UoC tags)
stripped — the same generic-template principle applied to the Business Case.

Two scoped variants from ONE definition, selected by the `serverless` flag:
  * serverless=False — the base template (EC2/RDS infrastructure build/HA hardening). This
    output is byte-for-byte the long-standing template; CL1 AT2/AT3 rely on its shape.
  * serverless=True  — a serverless + infrastructure-as-code build (§4/§5/§6 swapped for
    operating predefined IaC templates, authoring your own, a microservice deploy and
    monitoring). Used by S1-CL2 AT2. Top-level sections (§1-§7 + appendices) are identical
    between the two — only §4/§5/§6 content differs.

Sections a simpler deployment doesn't need carry an explicit "Not applicable — reason"
convention rather than being deleted, so the report always proves completeness.

Usage:  python scripts/build_deployment_report_template.py   (builds BOTH variants)
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_body_text import add_guidance_text, add_response_placeholder  # noqa: E402
from helpers.docx_callouts import add_applicability_note, add_convention_box  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from helpers.docx_tables import add_template_table  # noqa: E402
from brand import ADDRESS, CHARCOAL, CREAM, GREY, OCHRE, TEAL, TERRACOTTA  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path, serverless=False):
    doc = Document()
    configure_styles(doc)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Deployment Report")
    sub = doc.add_paragraph().add_run("[ Deployment / engagement name ]")
    sub.font.size = Pt(15); sub.italic = True; sub.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()

    cover_rows = [
        ("Engagement", "[ Engagement name ]"),
        ("Engagement reference", "[ Reference ]"),
        ("Report version", "[ e.g. v1.0 ]"),
        ("Prepared by", "[ Consultant name ]"),
        ("Consultant role", "[ Role, e.g. MTS Consultant ]"),
        ("Date submitted", "[ DD/MM/YYYY ]"),
        ("Submitted to", "[ Acceptance authority / sponsor ]"),
        ("Related documents", "[ e.g. the approved design this report implements ]"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in cover_rows:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        vr = cells[1].paragraphs[0].add_run(v); vr.font.size = Pt(10); vr.italic = True
        vr.font.color.rgb = RGBColor.from_string(GREY)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS + convention ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("How to use this template", style="Heading 1")
    intro = ("This is the standard YAT deployment report — it covers everything from a greenfield foundation "
             "build to a full high-availability hardening.")
    applic_eg = ("the maintenance-window context for a non-production build" if serverless
                 else "the HA failure/resize simulations")
    if serverless:
        intro = ("This is the standard YAT deployment report, set up for a serverless + infrastructure-as-code "
                 "deployment (operating predefined templates, authoring your own, and a microservice build).")
    add_convention_box(doc, [
        ("One report for any deployment.", intro),
        ("Complete every section.", "Where a section does not apply to your deployment, mark it "
         "“Not applicable — [reason]” rather than deleting it, so the report proves "
         "nothing was overlooked."),
        ("Sections flagged with an applicability note", "(in terracotta) are the ones a simpler "
         "deployment will often mark Not applicable — e.g. " + applic_eg + "."),
        ("Cross-reference your evidence.", "The build narrative and testing sections reference the "
         "screenshots and configuration exports captured in the appendices."),
    ])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h2 = lambda t: doc.add_paragraph(t, style="Heading 2")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Executive Summary")
    add_guidance_text(doc, "Write this last. A ≤ 1-page summary the reader sees first: what was deployed (or "
                     "changed); the region/AZ footprint; the 2–3 highlights; any limitations or items "
                     "deferred to a later phase; and — for a resilience deployment — whether the "
                     "availability target is now met and the headline simulation outcomes. ~250–400 words.")
    add_response_placeholder(doc)

    h1("2. Engagement Context")
    add_guidance_text(doc, "Brief context for the reader (≤ ½ page): the strategic/prior work this deployment "
                     "builds on (the approved business case and the design being implemented), your role, "
                     "and the scope hand-off to any later phase.")
    add_response_placeholder(doc)

    h1("3. Scope of Deployment")
    add_guidance_text(doc, "What is included in this deployment and what is deferred (≤ ½ page). Restate from the "
                     "approved design in your own words: the in-scope components, and what is out of scope / "
                     "deferred to a later phase.")
    add_response_placeholder(doc)
    h3("3.1 Maintenance-window context")
    add_applicability_note(doc, "a change to a running production system")
    add_guidance_text(doc, "If this work was performed on a live system within a maintenance window, describe the "
                     "window (duration, timing) and the requirement to finish either with the work complete "
                     "or cleanly rolled back. For a greenfield build, mark Not applicable.")
    add_response_placeholder(doc)

    # ===================== §4 / §5 / §6 — variant-specific =====================
    if serverless:
        _serverless_body(doc, h1, h3)
    else:
        _infra_body(doc, h1, h3)

    # ---- §7 Operational Handover (shared) ----
    h1("7. Operational Handover")
    add_guidance_text(doc, "Hand-over information for the team taking over the infrastructure.")
    h3("7.1 Access")
    add_guidance_text(doc, "Who has what access post-handover, MFA enforcement, any IAM group changes.")
    add_response_placeholder(doc)
    h3("7.2 Runbook references")
    add_guidance_text(doc, "Pointers to the design document, naming/tagging conventions, backup arrangements, and "
                     "the alarms list + notification destinations.")
    add_response_placeholder(doc)
    h3("7.3 Known limitations and what's next")
    add_guidance_text(doc, "Be explicit about what is not covered today and what a later phase would add.")
    add_response_placeholder(doc)
    h3("7.4 Documentation filing")
    add_template_table(doc, ["Item", "Filed in", "Reference"],
             [["This Deployment Report", "[ YAT ICT shared documentation ]", "[ ref ]"],
              ["Configuration exports (Appendix B)", "[ … ]", "[ ref ]"],
              ["Test / simulation evidence (Appendix C)", "[ … ]", "[ ref ]"]],
             widths=[6.5, 5.0, 4.0])
    h3("7.5 Feedback record")
    add_applicability_note(doc, "deployments where stakeholder feedback is captured")
    add_template_table(doc, ["Feedback received", "From", "Your response", "Resulting action"],
             [["[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
             widths=[5.0, 3.0, 4.0, 4.0])
    h3("7.6 Sign-off")
    add_template_table(doc, ["Role", "Name", "Date", "Signature"],
             [["Prepared by", "[ … ]", "", ""],
              ["Reviewed by", "[ … ]", "", ""],
              ["Approved by (acceptance authority)", "[ … ]", "", ""]],
             widths=[5.5, 4.5, 2.5, 3.0])

    # ---- APPENDICES ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    if serverless:
        h1("Appendix A — Build evidence (screenshots)")
        add_guidance_text(doc, "Capture a console screenshot evidencing each part you built, with the region "
                         "indicator visible. List each below and cross-reference it from §4 / §6. Examples:")
        add_template_table(doc, ["#", "Screenshot", "What must be visible"],
                 [["A1", "Lab session + execution role", "[ the session; the least-privilege role policy ]"],
                  ["A2", "CloudFormation stack", "[ a stack at CREATE_COMPLETE with its resources ]"],
                  ["A3", "The datastore", "[ an item written by the function ]"],
                  ["A4", "The API / microservice test", "[ the request + a 200 response ]"],
                  ["A5", "CloudWatch alarms", "[ the configured alarms and their state ]"],
                  ["…", "[ add as your build requires ]", "[ … ]"]],
                 widths=[1.0, 5.5, 9.0])
        h1("Appendix B — Infrastructure-as-code templates and user documentation")
        add_guidance_text(doc, "Attach the template you authored and the predefined templates you operated, plus the "
                         "IaC user documentation (how to deploy/update/delete the stack with parameters).")
        add_response_placeholder(doc, "[ Templates + user documentation ]")
        h1("Appendix C — Test and troubleshooting evidence")
        add_guidance_text(doc, "Attach the evidence supporting §6 — the deploy/update/delete captures, the "
                         "microservice test request/response and datastore item, the alarm test, and the "
                         "troubleshooting log.")
        add_response_placeholder(doc, "[ Test and troubleshooting evidence ]")
    else:
        h1("Appendix A — Build evidence (screenshots)")
        add_guidance_text(doc, "Capture a console screenshot evidencing each component you built or changed, with the "
                         "region indicator visible. List each below and cross-reference it from §4 / §6. Examples:")
        add_template_table(doc, ["#", "Screenshot", "What must be visible"],
                 [["A1", "IAM groups / MFA", "[ created groups; MFA enabled ]"],
                  ["A2", "VPC subnets", "[ subnets + AZs ]"],
                  ["A3", "EC2 instances + ASG", "[ running instances across AZs; ASG min/desired/max ]"],
                  ["A4", "ALB target group health", "[ healthy targets ]"],
                  ["A5", "RDS database", "[ available; Multi-AZ status; encryption ]"],
                  ["A6", "CloudWatch alarms / dashboard", "[ the alarms / service-level dashboard ]"],
                  ["…", "[ add as your deployment requires ]", "[ … ]"]],
                 widths=[1.0, 5.5, 9.0])
        h1("Appendix B — Configuration exports")
        add_guidance_text(doc, "Export each configuration (AWS CLI or console) and attach as a code block or file. "
                         "Examples: IAM policies; security-group rules; VPC/subnet/route tables; launch template "
                         "+ ASG; ALB + target groups; RDS instance; S3 bucket policy/encryption; CloudWatch alarms.")
        add_response_placeholder(doc, "[ Configuration exports ]")
        h1("Appendix C — Test and simulation evidence")
        add_guidance_text(doc, "Attach the evidence supporting the results in §6 — screenshots, terminal/log excerpts, "
                         "metric graphs, and (for HA work) failure/resize simulation captures and the computed "
                         "availability over the window.")
        add_response_placeholder(doc, "[ Test and simulation evidence ]")

    h1("Document control")
    add_template_table(doc, ["Field", "Value"],
             [["Document version", "[ v1.0 ]"],
              ["Author", "[ Name, role ]"],
              ["Engagement", "[ Engagement name ]"],
              ["Date submitted", "[ DD/MM/YYYY ]"],
              ["Distribution", "[ … ]"],
              ["Related documents", "[ the design implemented; predecessor/successor reports ]"]],
             widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


def _infra_body(doc, h1, h3):
    """§4/§5/§6 for an EC2/RDS infrastructure build or HA hardening (the base template)."""
    h1("4. Build / Change Narrative")
    add_guidance_text(doc, "A layer-by-layer account of what was built or changed. For each layer, write a short "
                     "narrative — for a fresh build, what you stood up; for a change to an existing build, "
                     "what changed from the baseline — and cross-reference the Appendix A screenshots and "
                     "Appendix B configuration exports.")
    for n, title, hint in [
        ("4.1", "Identity and access management (IAM)", "account access, groups/users/roles, MFA, instance profiles — or, for a change, any IAM additions (state if none)"),
        ("4.2", "Network topology", "VPC, subnets, gateways, route tables — or the subnets/routes added (e.g. a second AZ)"),
        ("4.3", "Compute (EC2 + Auto Scaling)", "launch template, instance type + why, ASG min/desired/max + scaling policy — or capacity/AZ-spread changes"),
        ("4.4", "Load balancing (ALB)", "the load balancer, target group, listener, health check — or cross-AZ target changes"),
        ("4.5", "Database (RDS)", "instance class + why, engine version, storage, encryption, backups — or enabling Multi-AZ"),
        ("4.6", "Storage (EBS + S3)", "EBS volumes, S3 buckets, encryption, public-access settings"),
        ("4.7", "Security (security groups + encryption)", "the tiered security-group model, encryption in transit + at rest — or HA-related adjustments"),
        ("4.8", "Monitoring", "the CloudWatch alarms and thresholds — baseline, or the HA-tuned set (per-AZ counts, failover detection, replica lag, service-level dashboard)"),
    ]:
        h3(f"{n} {title}")
        add_guidance_text(doc, f"Cover: {hint}. Cross-reference the relevant Appendix A screenshots and Appendix B exports.")
        add_response_placeholder(doc)
    h3("4.9 Cross-Region backup / replication")
    add_applicability_note(doc, "deployments with cross-Region resilience in scope")
    add_guidance_text(doc, "Cross-Region snapshot copies, AWS Backup, or S3 cross-Region replication, if in scope. "
                     "Otherwise mark Not applicable.")
    add_response_placeholder(doc)

    h1("5. Configuration Decisions")
    add_guidance_text(doc, "The approved design leaves specific decisions to the implementer. For each decision you "
                     "made, state your choice and justify it against the workload and requirements. Add rows "
                     "as needed; examples below span foundation-build and HA decisions.")
    add_template_table(doc, ["#", "Decision point", "Your decision", "Rationale"],
             [["C1", "EC2 instance type (within the chosen family)", "[ … ]", "[ workload + cost ]"],
              ["C2", "RDS instance class", "[ … ]", "[ … ]"],
              ["C3", "Storage sizing (EBS data volume / RDS storage)", "[ show the calc ]", "[ data footprint + growth ]"],
              ["C4", "ASG scaling threshold", "[ … ]", "[ expected CPU profile ]"],
              ["C5", "RDS Multi-AZ apply timing (HA)", "[ … ]", "[ maintenance-window rationale ]"],
              ["C6", "Post-HA ASG capacity (min ≥ 2 across AZs)", "[ … ]", "[ AZ-failure resilience ]"],
              ["C7", "Cross-Region backup destination (HA)", "[ … ]", "[ … ]"],
              ["…", "[ add further decisions ]", "[ … ]", "[ … ]"]],
             widths=[1.0, 6.0, 4.0, 4.5])

    h1("6. Testing, Simulation and Validation")
    add_guidance_text(doc, "Document the tests run to verify the deployment. For each test, state the test, the "
                     "result, and reference the supporting evidence in Appendix C.")
    h3("6.1 Connectivity tests")
    add_template_table(doc, ["Test", "Outcome (Pass/Fail)", "Notes"],
             [["ALB → EC2 health check", "[ ]", "[ ]"],
              ["EC2 → RDS connection (private)", "[ ]", "[ ]"],
              ["EC2 → internet via NAT", "[ ]", "[ ]"],
              ["RDS not publicly reachable (negative test)", "[ ]", "[ ]"]],
             widths=[7.0, 4.0, 4.5])
    h3("6.2 Autoscaling test")
    add_guidance_text(doc, "Trigger a scaling event (e.g. load against the ALB) and confirm the ASG scales out and "
                     "back in. Cross-reference Appendix C.")
    add_response_placeholder(doc)
    h3("6.3 Database connectivity and basic operations")
    add_guidance_text(doc, "Confirm the database tier is reachable from the app tier over the private network, the "
                     "engine version meets the application requirement, and encryption-in-transit is in place.")
    add_response_placeholder(doc)
    h3("6.4 Infrastructure end-to-end smoke test")
    add_guidance_text(doc, "Confirm the infrastructure is ready to serve traffic (e.g. a placeholder page via the "
                     "ALB DNS returns HTTP 200/302 from a backend instance that can reach the database).")
    add_response_placeholder(doc)
    h3("6.5 Failure simulation")
    add_applicability_note(doc, "high-availability / resilience deployments")
    add_guidance_text(doc, "Execute each failure simulation from the design (e.g. EC2 termination, RDS Multi-AZ "
                     "failover, AZ partition). Record method, observed outcome, and whether the service "
                     "stayed reachable. Otherwise mark Not applicable.")
    add_template_table(doc, ["#", "Simulation", "Method", "Observed outcome", "Reachable throughout?", "Evidence"],
             [["F1", "[ e.g. EC2 termination ]", "[ … ]", "[ … ]", "[ Yes/No ]", "C—"],
              ["F2", "[ e.g. RDS failover ]", "[ … ]", "[ … ]", "[ Yes (brief blip) ]", "C—"]],
             widths=[0.8, 3.0, 3.0, 3.6, 2.6, 1.5])
    h3("6.6 Resize simulation")
    add_applicability_note(doc, "high-availability / resilience deployments")
    add_guidance_text(doc, "Execute each resize simulation from the design (e.g. ASG capacity increase, RDS instance "
                     "class change) and record the availability impact. Otherwise mark Not applicable.")
    add_response_placeholder(doc)
    h3("6.7 Availability measurement")
    add_applicability_note(doc, "deployments with an availability target to demonstrate")
    add_guidance_text(doc, "Describe how availability was measured (e.g. CloudWatch dashboard, a curl-loop against "
                     "the ALB during the window) and report the measured availability across the window.")
    add_response_placeholder(doc)
    h3("6.8 Simulation findings vs the design")
    add_applicability_note(doc, "high-availability / resilience deployments")
    add_guidance_text(doc, "For each simulation, compare the observed outcome against the expected outcome in the "
                     "design; document any divergence and why.")
    add_response_placeholder(doc)
    h3("6.9 Adjustments made per simulation outcomes")
    add_applicability_note(doc, "high-availability / resilience deployments")
    add_guidance_text(doc, "Any changes made to the architecture, configuration, or monitoring based on what the "
                     "simulations revealed — or an explicit statement that none were needed, with reasoning.")
    add_response_placeholder(doc)


def _serverless_body(doc, h1, h3):
    """§4/§5/§6 for a serverless + infrastructure-as-code build (CL2 AT2)."""
    h1("4. Build / Change Narrative")
    add_guidance_text(doc, "An account of what you built, stream by stream — operating the predefined templates, "
                     "authoring your own, the microservice, and monitoring. Cross-reference the Appendix A "
                     "screenshots and the Appendix B templates/exports.")
    h3("4.1 Lab environment and access")
    add_guidance_text(doc, "The cloud platform and region used; the least-privilege execution role for the function "
                     "(what it can do, and nothing more); and how access was obtained. Note any region "
                     "simulation (e.g. a lab region standing in for the production region).")
    add_response_placeholder(doc)
    h3("4.2 Operating the predefined infrastructure-as-code templates")
    add_guidance_text(doc, "The IaC service you chose and why; reading the supplied templates to determine the "
                     "resources they create and their dependencies; and deploying, confirming, updating and "
                     "deleting resources with them. Record troubleshooting in §6.")
    add_template_table(doc, ["Action", "Command / method", "Result"],
             [["Deploy", "[ e.g. aws cloudformation deploy … ]", "[ CREATE_COMPLETE ]"],
              ["Update", "[ change a parameter; reviewed change set ]", "[ UPDATE_COMPLETE ]"],
              ["Delete", "[ delete-stack ]", "[ DELETE_COMPLETE ]"]],
             widths=[2.6, 8.4, 5.0])
    h3("4.3 Authoring the infrastructure-as-code template")
    add_guidance_text(doc, "The template you wrote to provision a related set of resources: the resources, how you "
                     "parameterised it for configuration and code reuse, how you updated/redeployed it, and how "
                     "you removed it cleanly. Attach the template in Appendix B.")
    add_response_placeholder(doc)
    h3("4.4 The microservice")
    add_guidance_text(doc, "Reviewing the design and supplied code; deploying and configuring the serverless services; "
                     "the interface/integration contract (endpoint + payload); and confirming the application "
                     "functions. Record troubleshooting in §6.")
    add_response_placeholder(doc)
    h3("4.5 Monitoring and alarms")
    add_guidance_text(doc, "The metrics and alarms you configured (e.g. queue depth, function errors/throttles, "
                     "datastore throttles), their thresholds, and where they notify.")
    add_response_placeholder(doc)

    h1("5. Configuration Decisions")
    add_guidance_text(doc, "The approved design leaves specific decisions to the implementer. For each decision you "
                     "made, state your choice and justify it against the workload and requirements. Add rows "
                     "as needed.")
    add_template_table(doc, ["#", "Decision point", "Your decision", "Rationale"],
             [["C1", "Infrastructure-as-code service", "[ … ]", "[ native / tooling / supported ]"],
              ["C2", "Datastore (and capacity mode)", "[ … ]", "[ workload shape ]"],
              ["C3", "Decoupling (queue / messaging)", "[ … ]", "[ durability under load ]"],
              ["C4", "Compute (function runtime / memory)", "[ … ]", "[ … ]"],
              ["C5", "Environment parameterisation", "[ … ]", "[ env values parameterised; the region is the deploy target, not a template parameter ]"],
              ["…", "[ add further decisions ]", "[ … ]", "[ … ]"]],
             widths=[1.0, 6.0, 4.0, 4.5])

    h1("6. Testing, Simulation and Validation")
    add_guidance_text(doc, "Document the tests run to verify the build. For each test, state the test, the result, "
                     "and reference the supporting evidence in Appendix C.")
    h3("6.1 IaC deploy / update / delete verification")
    add_guidance_text(doc, "Confirm each template deploys, updates (via a reviewed change set) and deletes cleanly, "
                     "with the console / describe-stacks as evidence.")
    add_response_placeholder(doc)
    h3("6.2 Own-template parameterise and redeploy")
    add_guidance_text(doc, "Deploy your authored template, then redeploy it with a different parameter set to confirm "
                     "configuration reuse without editing the template body.")
    add_response_placeholder(doc)
    h3("6.3 Microservice functional test")
    add_guidance_text(doc, "Exercise the microservice end to end (e.g. post a sample event) and confirm the expected "
                     "result is produced and persisted, including a negative test.")
    add_response_placeholder(doc)
    h3("6.4 Durability test")
    add_applicability_note(doc, "builds with a queue / decoupling component")
    add_guidance_text(doc, "Force a downstream stall and confirm events queue and then drain without loss. Otherwise "
                     "mark Not applicable.")
    add_response_placeholder(doc)
    h3("6.5 Monitoring / alarm test")
    add_guidance_text(doc, "Drive a metric past its threshold and confirm the alarm transitions and notifies, then "
                     "returns to OK.")
    add_response_placeholder(doc)
    h3("6.6 Troubleshooting log")
    add_template_table(doc, ["Symptom", "Cause", "Fix"],
             [["[ … ]", "[ … ]", "[ … ]"],
              ["[ … ]", "[ … ]", "[ … ]"]],
             widths=[5.2, 5.4, 5.4])


if __name__ == "__main__":
    base_dir = Path("../diploma-cloud-cyber-website-s1/public/templates")
    build(base_dir / "YAT-Deployment-Report-Template.docx", serverless=False)
    build(base_dir / "YAT-Deployment-Report-Template-Serverless.docx", serverless=True)
