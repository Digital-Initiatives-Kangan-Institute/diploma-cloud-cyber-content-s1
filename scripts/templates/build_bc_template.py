#!/usr/bin/env python3
"""Build the YAT / MTS Business Case template (.docx) from the brand pack.

Generates a professionally branded, system-agnostic Business Case document template:
cover sheet, table-of-contents field, all narrative sections, and the standard tables
(gap analysis, CBA summaries, risk register, action plan, sign-off). Branding follows
scenario/branding/brand-pack.md (§4 visual identity, §5.2 document templates,
§5.3 disclosure). Source of structure: the BC section outline used across S1-CL1.

This is the in-world template (no assessment/UoC scaffolding) — the canonical shape of a
YAT/MTS Business Case, suitable for handing students per-section snippets for exercises
and as the basis for the AT1 deliverable variant.

Usage:  python scripts/build_bc_template.py [output.docx]
Default output: scenario/templates/YAT-Business-Case-Template.docx
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.yat_brand import ADDRESS, CHARCOAL, CREAM, DISCLOSURE, FONT, GREY, OCHRE, STONE, TEAL, TERRACOTTA, WHITE  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell, shade_paragraph  # noqa: E402
from helpers.docx_body_text import add_guidance_text, add_response_placeholder  # noqa: E402
from helpers.docx_tables import add_template_table  # noqa: E402
from helpers.yat_docx_document import build_header_footer, configure_styles, wordmark  # noqa: E402


def build(path):
    doc = Document()
    configure_styles(doc)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    lock_p = doc.add_paragraph()
    wordmark(lock_p)
    addr = doc.add_paragraph()
    ar = addr.add_run(ADDRESS)
    ar.font.size = Pt(9)
    ar.font.color.rgb = RGBColor.from_string(GREY)
    rule = doc.add_paragraph()
    paragraph_bottom_rule(rule, TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph(style="Title")
    title.add_run("Business Case")
    sub = doc.add_paragraph()
    sr = sub.add_run("[ Project / initiative name ]")
    sr.font.size = Pt(15)
    sr.italic = True
    sr.font.color.rgb = RGBColor.from_string(GREY)
    for _ in range(2):
        doc.add_paragraph()

    cover_rows = [
        ("Prepared for", "[ Client / board — e.g. YAT board, via the engagement sponsor ]"),
        ("Prepared by", "[ Consultant name, MP Tech Solutions (MTS) ]"),
        ("Engagement", "[ Engagement name ]"),
        ("Document version", "[ e.g. v0.1 ]"),
        ("Date", "[ Date ]"),
        ("Analysis period", "[ e.g. 5 years ]"),
        ("Currency", "AUD ex GST"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in cover_rows:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1])
        shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k)
        kr.bold = True
        kr.font.size = Pt(10)
        vr = cells[1].paragraphs[0].add_run(v)
        vr.font.size = Pt(10)
        vr.italic = True
        vr.font.color.rgb = RGBColor.from_string(GREY)
        cells[0].width = Cm(4.5)
        cells[1].width = Cm(12.0)

    # ---- TABLE OF CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE)
    build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    toc_p = doc.add_paragraph()
    add_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u',
              placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE)
    build_header_footer(doc.sections[-1])

    def h1(text):
        doc.add_paragraph(text, style="Heading 1")

    def h2(text):
        doc.add_paragraph(text, style="Heading 2")

    h1("1. Executive Summary")
    add_guidance_text(doc, "Write this section last. One page the board reads first: the problem in a "
                  "sentence; the options considered; your recommendation; the high-level 5-year "
                  "cost position for each option; the top 2–3 risks of the recommended option and "
                  "how you'll manage them; and the decision you are asking the board to take today.")
    add_response_placeholder(doc)

    h1("2. Engagement Context")
    add_guidance_text(doc, "Establish who you are, the engagement, who you report to, and what this "
                  "Business Case asks the board to decide. Approx. one short paragraph.")
    add_response_placeholder(doc)

    h1("3. Strategic Alignment Analysis")
    add_guidance_text(doc, "Analyse the organisation's ICT Strategic Plan against the broader industry "
                  "environment and the organisation's objectives. Identify the objectives material "
                  "to this initiative (cite them); characterise the relevant industry direction; "
                  "state where the plan aligns with industry and where it diverges; draw out the "
                  "implications for this initiative. Approx. 250–400 words.")
    add_response_placeholder(doc)

    h1("4. Current State")
    add_guidance_text(doc, "Synthesise the current-state material into a board-level summary — do not "
                  "reproduce source documents. Decide what is material to the decision, describe it "
                  "in your own words, and focus the reader on what drives the gap analysis. Cover "
                  "topology and security posture; the system's current status (what it runs, "
                  "availability, recovery objectives, condition); supporting infrastructure and "
                  "integrations; and staff capability and external dependencies. Approx. 150–250 words.")
    add_response_placeholder(doc)

    h1("5. Gap Analysis")
    add_guidance_text(doc, "Compare each in-scope strategic objective against the current state. Identify the "
                  "gap, the improvement opportunity, and the proposed change. Add rows as needed.")
    add_template_table(doc,
          ["Strategic objective", "Current state", "Desired future state", "Gap",
           "Improvement opportunity", "Proposed change"],
          [["[ Objective ]", "[ … ]", "[ … ]", "[ … ]", "[ … ]", "[ … ]"]] * 3,
          widths=[3.0, 2.6, 2.6, 2.6, 2.6, 2.6])
    add_guidance_text(doc, "Narrative summary tying the table together (100–200 words).")
    add_response_placeholder(doc)

    h1("6. Options Considered and Evaluation")
    h2("6.1 Workload definition")
    add_guidance_text(doc, "Define the workload the chosen option must support: user load, peak patterns, "
                  "data volumes, integration requirements, availability target.")
    add_response_placeholder(doc)
    h2("6.2 Options considered")
    add_guidance_text(doc, "State the options you assessed (e.g. Option A — renew on-premises; "
                  "Option B — migrate to AWS). Note any other options and why they were / weren't assessed.")
    add_response_placeholder(doc)
    h2("6.3 Evaluation method")
    add_guidance_text(doc, "State the evaluation method applied (CBA + intangibles + risk register, and any "
                  "weighted decision matrix). Justify why it suits this decision.")
    add_response_placeholder(doc)
    h2("6.4 Initial impact and difficulty assessment")
    add_template_table(doc, ["", "Option A", "Option B"],
          [["Strategic impact", "[ … ]", "[ … ]"],
           ["Implementation difficulty", "[ … ]", "[ … ]"],
           ["Headline pros", "[ … ]", "[ … ]"],
           ["Headline cons", "[ … ]", "[ … ]"]],
          widths=[4.5, 6.0, 6.0])

    h1("7. Cost-Benefit Analysis")
    add_guidance_text(doc, "A full 5-year CBA comparing the options. Detailed line items live in Appendix 1; "
                  "the summary tables here are the headline view the board reads. All figures AUD ex GST.")
    h2("7.1 Assumptions")
    add_template_table(doc, ["Assumption", "Value", "Used as-is?"],
          [["[ Assumption ]", "[ Value ]", "[ ] yes  /  [ ] adjusted — see note"]] * 4,
          widths=[7.0, 5.0, 4.5])
    h2("7.2 Option A — 5-year summary")
    yr = ["", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5", "5-year"]
    add_template_table(doc, yr,
          [["One-off capital", "$_____", "—", "—", "—", "—", "$_____"],
           ["Recurring", "$_____", "$_____", "$_____", "$_____", "$_____", "$_____"],
           ["Annual total", "$_____", "$_____", "$_____", "$_____", "$_____", "$_____"]])
    h2("7.3 Option B — 5-year summary")
    add_template_table(doc, yr,
          [["One-off project", "$_____", "—", "—", "—", "—", "$_____"],
           ["Cloud / direct", "$_____", "$_____", "$_____", "$_____", "$_____", "$_____"],
           ["Staff + external", "$_____", "$_____", "$_____", "$_____", "$_____", "$_____"],
           ["Annual total", "$_____", "$_____", "$_____", "$_____", "$_____", "$_____"]])
    h2("7.4 Quantified benefit (e.g. avoided downtime)")
    add_guidance_text(doc, "Quantify the material benefit of the recommended option (for an availability "
                  "uplift, the avoided-downtime cost; otherwise the relevant quantified benefit).")
    add_response_placeholder(doc)
    h2("7.5 Comparison summary")
    add_template_table(doc, ["", "Option A", "Option B", "Delta (B − A)"],
          [["5-year direct cost", "$_____", "$_____", "$_____"],
           ["Quantified benefit", "—", "$_____", "$_____"],
           ["Net 5-year position", "$_____", "$_____", "$_____"]],
          widths=[5.0, 3.8, 3.8, 3.8])
    h2("7.6 Sensitivity analysis")
    add_guidance_text(doc, "Pick at least two assumptions where a reasonable change would materially affect "
                  "the outcome. Show the recalculation and the effect on the recommendation.")
    add_response_placeholder(doc)

    h1("8. Risk and Impact Assessment")
    h2("8.1 Intangibles comparison")
    add_guidance_text(doc, "One to two sentences per factor. Be honest about trade-offs.")
    add_template_table(doc, ["Factor", "Option A", "Option B"],
          [["[ Availability target ]", "[ … ]", "[ … ]"],
           ["[ Recovery objectives ]", "[ … ]", "[ … ]"],
           ["[ Capacity for growth / peaks ]", "[ … ]", "[ … ]"],
           ["[ Staff capability impact ]", "[ … ]", "[ … ]"],
           ["[ Vendor lock-in ]", "[ … ]", "[ … ]"],
           ["[ Security posture ]", "[ … ]", "[ … ]"]],
          widths=[5.0, 5.75, 5.75])
    h2("8.2 Risk register (recommended option)")
    add_template_table(doc, ["Risk", "Likelihood", "Impact", "Mitigation"],
          [["[ Risk ]", "[ H/M/L ]", "[ H/M/L ]", "[ … ]"]] * 3,
          widths=[6.5, 2.5, 2.5, 5.0])

    h1("9. Recommendation")
    add_guidance_text(doc, "State the recommended option and the headline rationale (3–5 sentences) — connect "
                  "it to the CBA findings, the intangibles, and the risk register. This drives the "
                  "prioritisation and action plan in §10.")
    add_response_placeholder(doc)

    h1("10. Action Plan")
    h2("10.1 Prioritised changes")
    add_template_table(doc, ["#", "Change", "Priority rationale"],
          [["1", "[ … ]", "[ … ]"], ["2", "[ … ]", "[ … ]"], ["3", "[ … ]", "[ … ]"]],
          widths=[1.2, 7.0, 8.0])
    h2("10.2 Implementation schedule")
    add_guidance_text(doc, "Indicative schedule — phase the work; durations and dependencies. Do not propose "
                  "a single all-at-once cutover unless the risk register supports it.")
    add_template_table(doc, ["Phase", "Activities", "Start", "Duration", "Dependencies", "Owner"],
          [["1", "[ … ]", "[ … ]", "[ … ]", "—", "[ … ]"],
           ["2", "[ … ]", "[ … ]", "[ … ]", "[ … ]", "[ … ]"],
           ["3", "[ … ]", "[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
          widths=[1.5, 5.0, 2.2, 2.2, 3.0, 2.4])
    h2("10.3 Standards, targets, and success metrics")
    add_template_table(doc, ["Aspect", "Standard / target", "How measured"],
          [["[ Availability target ]", "[ … ]", "[ … ]"],
           ["[ Recovery (RPO/RTO) ]", "[ … ]", "[ … ]"],
           ["[ Cost envelope ]", "[ … ]", "[ … ]"],
           ["[ Data residency ]", "[ … ]", "[ … ]"]],
          widths=[5.0, 5.0, 5.5])
    h2("10.4 Implementation methods")
    add_guidance_text(doc, "For each phase, identify the methods, tools, and standards applied (e.g. a "
                  "Well-Architected review at design milestones; the change-management procedure "
                  "for production changes).")
    add_response_placeholder(doc)
    h2("10.5 Alignment with change-management procedure")
    add_guidance_text(doc, "State how the action plan aligns with the organisation's change-management "
                  "procedure: which phases require change requests, approval, and sign-off.")
    add_response_placeholder(doc)

    h1("11. Next Steps and Decision Requested")
    add_guidance_text(doc, "State precisely what you are asking the board to decide today, and list any items "
                  "deferred to later sign-off gates.")
    add_response_placeholder(doc)

    h1("Sign-off")
    add_template_table(doc, ["Role", "Name", "Date", "Signature"],
          [["Prepared by", "", "", ""],
           ["Reviewed by", "", "", ""],
           ["Approved by", "", "", ""]],
          widths=[5.0, 4.5, 3.0, 3.5])

    # ---- APPENDICES ----
    doc.add_section(WD_SECTION.NEW_PAGE)
    build_header_footer(doc.sections[-1])
    h1("Appendix 1 — Cost-Benefit Analysis: Detailed Line Items")
    add_guidance_text(doc, "The detailed line items underpinning the §7 summary tables — one-off and "
                  "recurring costs per option, and the year-by-year projection whose totals feed §7. "
                  "Build out the line-item tables here per option.")
    add_response_placeholder(doc, "[ Detailed CBA line-item tables ]")

    h1("Appendix 2 — Supporting Research")
    add_guidance_text(doc, "Attach or link the supporting research: pricing-calculator exports for the "
                  "Appendix 1 figures, industry-environment sources cited in §3 (with access dates), "
                  "and any other supporting material.")
    add_response_placeholder(doc, "[ Supporting research ]")

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website/public/templates/YAT-Business-Case-Template.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
