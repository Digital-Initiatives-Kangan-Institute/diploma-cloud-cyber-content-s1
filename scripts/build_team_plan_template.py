#!/usr/bin/env python3
"""Build the YAT / MTS Team Plan template (.docx) — the in-world student-facing template.

Derived from the AT1 Team Plan exemplar (build_s1_cl3_at1_team_plan_exemplar) by stripping the worked
content and the UoC `Evidences:` tags and replacing them with section guidance and blank
response areas / tables. Branding follows scenario/branding/brand-pack.md via build_bc_template.

This is the in-world template (no assessment/UoC scaffolding) — the canonical shape of a YAT / MTS
team plan, used as the basis for the AT1 Part A deliverable (BSBXTW401 element 1 + task
allocation). The Knowledge Evidence appendix is kept as a prompt because the team plan is where
the BSBXTW401 element-1 knowledge is evidenced; the specific questions live in the assessment task.

Usage:  python scripts/build_team_plan_template.py [output.docx]
Default: ../diploma-cloud-cyber-website/public/templates/YAT-Team-Plan-Template.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_bc_template as bc  # noqa: E402  (branding + guidance/response/table helpers)

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path):
    doc = Document()
    bc.configure_styles(doc)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    bc.build_header_footer(sec)

    # ---- COVER ----
    bc.wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(bc.ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(bc.GREY)
    bc.paragraph_bottom_rule(doc.add_paragraph(), bc.TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Team Plan")
    sub = doc.add_paragraph().add_run("[ Engagement name ]")
    sub.font.size = Pt(15); sub.italic = True; sub.font.color.rgb = RGBColor.from_string(bc.GREY)
    for _ in range(2):
        doc.add_paragraph()

    cover_rows = [
        ("Prepared for", "[ MTS Senior Consultant — engagement supervision ]"),
        ("Prepared by", "[ Your name — team lead ]"),
        ("Engagement", "[ Engagement name ]"),
        ("Team", "[ Your team — list each member and the dimension they own ]"),
        ("Document version", "[ e.g. v0.1 ]"),
        ("Date", "[ Date ]"),
        ("Classification", "Internal — MTS / YAT"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in cover_rows:
        cells = ct.add_row().cells
        bc.set_cell_borders(cells[0]); bc.set_cell_borders(cells[1]); bc.shade_cell(cells[0], bc.CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        vr = cells[1].paragraphs[0].add_run(v); vr.font.size = Pt(10)
        vr.italic = True; vr.font.color.rgb = RGBColor.from_string(bc.GREY)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    bc.add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")

    h1("1. Purpose and team objectives")
    bc.guidance(doc, "State the engagement, your team's common objective, the responsibilities (who "
                     "owns what), and the required outcomes (the sign-off gates). A few sentences plus "
                     "the responsibilities — keep it focused on what the team is set up to deliver.")
    bc.response(doc)

    h1("2. Team structure and dimension ownership")
    bc.guidance(doc, "List your team members and the improvement dimension each owns (security, "
                     "reliability, scalability, cost) so the team covers all four. Note who is leading "
                     "this phase — the lead role rotates. Add rows as needed.")
    bc.table(doc, ["Member", "Dimension owned", "Focus of the dimension"],
             [["[ Name ]", "[ Security / Reliability / Scalability / Cost ]", "[ … ]"]] * 4,
             widths=[4.2, 3.5, 8.0])

    h1("3. Performance expectations")
    bc.guidance(doc, "For each member, set the expected outcomes, goals and behaviours — in line with "
                     "the team objective and YAT's policies (Change Management Procedure, Security and "
                     "Incident Response Policy, Privacy / Data Handling Policy). One row per member.")
    bc.table(doc, ["Member (dimension)", "Expected outcomes", "Goals", "Expected behaviours"],
             [["[ Member ]", "[ … ]", "[ … ]", "[ … ]"]] * 4,
             widths=[3.4, 4.0, 4.0, 4.1])

    h1("4. Accountability strategies")
    bc.guidance(doc, "How will you make sure each member is accountable for their role and "
                     "responsibilities? Consider owned-dimension deliverables, the sign-off gates, the "
                     "led-meeting record, peer review at integration, and the escalation path.")
    bc.response(doc)

    h1("5. Task allocation")
    bc.guidance(doc, "Allocate the improvement work across the team, by dimension, with the instruction "
                     "each member needs and an allowance for your contingencies (§6). Each member carries "
                     "their dimension across analysis (AT1), design (AT2) and implementation (AT3).")
    bc.table(doc, ["Member", "Allocated tasks (AT1 → AT2 → AT3)", "Instruction / notes"],
             [["[ Member ]", "[ … ]", "[ … ]"]] * 4,
             widths=[3.0, 7.5, 5.0])

    h1("6. Contingency planning")
    bc.guidance(doc, "Identify the contingencies most likely to impact your team — for example unplanned "
                     "leave or absence, re-allocation of work tasks, and succession for an important role "
                     "(such as the rotating lead) — and your plan for each.")
    bc.table(doc, ["Contingency", "Likelihood", "Plan"],
             [["[ Contingency ]", "[ H/M/L ]", "[ … ]"]] * 4,
             widths=[5.0, 2.3, 8.2])

    h1("7. Team operating norms")
    bc.guidance(doc, "How will the team actually work — the meeting cadence (rotating chair), how "
                     "decisions are recorded, the change discipline (avoid the accounting close and EOFY "
                     "for production change), and how you communicate with YAT.")
    bc.response(doc)

    # ---- KNOWLEDGE EVIDENCE ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1("Appendix — Knowledge Evidence")
    bc.guidance(doc, "Answer the Knowledge Evidence questions from your assessment task here, in your "
                     "own words and about your own plan — the organisational requirements that shape how "
                     "your team operates, the legislative requirements relevant to its work, and the "
                     "typical workplace contingencies your team faces and how you handle them.")
    bc.table(doc, ["Question", "Response"],
             [["[ Knowledge Evidence question from the task ]", "[ Your response ]"]] * 3,
             widths=[6.5, 9.0])

    # ---- SIGN-OFF ----
    h1("Sign-off")
    bc.table(doc, ["Role", "Name", "Date", "Signature"],
             [["Prepared by (team lead)", "", "", ""],
              ["Reviewed by", "", "", ""],
              ["Noted by", "", "", ""]],
             widths=[5.0, 5.5, 2.5, 3.0])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website/public/templates/YAT-Team-Plan-Template.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
