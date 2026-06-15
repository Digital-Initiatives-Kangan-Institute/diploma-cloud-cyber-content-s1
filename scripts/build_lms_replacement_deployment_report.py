#!/usr/bin/env python3
"""Build the AccentLoitte LMS Replacement Deployment Report (.docx) — student model.

The completed on-premises deployment report for the 2022 GrayBoard -> DOODLE replacement,
on the generic Deployment Report template. Student-facing model (in-world, no UoC tags, no
KE/reflections — those are assessment-only), companion to the AccentLoitte Business Case,
presentation and Solution Design. A FINISHED deliverable: real Pass outcomes and a closure
sign-off, evidence described as filed (not placeholder prompts). On-prem, so cloud/HA
sections are "Not applicable"; AccentLoitte's end-to-end scope means data migration,
cutover and training ARE included. Output to public/documents/ for printing to PDF.

Usage:  python scripts/build_lms_replacement_deployment_report.py [output.docx]
Default: ../diploma-cloud-cyber-website/public/documents/YAT-LMS-Replacement-Deployment-Report.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_bullet_list  # noqa: E402
import build_bc_template as bc   # noqa: E402
import build_s1_cl1_at1_bc_exemplar as ex   # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def na(doc, reason):
    p = doc.add_paragraph()
    r = p.add_run(f"Not applicable — {reason}")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    p.paragraph_format.space_after = Pt(6)
    return p


def build(path):
    doc = Document()
    bc.configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    bc.build_header_footer(sec)

    # ---- COVER ----
    bc.wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(bc.ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(bc.GREY)
    bc.paragraph_bottom_rule(doc.add_paragraph(), bc.TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Deployment Report")
    sub = doc.add_paragraph().add_run("LMS Replacement — GrayBoard to DOODLE (on-premises)")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT LMS Replacement Project — Phases 2–3"),
        ("Report version", "v1.0 — Final"),
        ("Prepared by", "AccentLoitte Consulting"),
        ("Reviewed by", "Jamie O'Donnell (AccentLoitte Senior Consultant)"),
        ("Date", "December 2022"),
        ("Submitted to", "Sam Walker (YAT ICT Manager)"),
        ("Related documents", "LMS Replacement Solution Design; LMS Replacement Business Case"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        bc.set_cell_borders(cells[0]); bc.set_cell_borders(cells[1]); bc.shade_cell(cells[0], bc.CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    bc.add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Executive Summary")
    ex.para(doc, "AccentLoitte deployed the DOODLE LMS on-premises at the YAT Cremorne campus, replacing the "
                 "end-of-life GrayBoard platform, and delivered the engagement end-to-end through deployment, "
                 "data migration, cutover and stabilisation. All student records were migrated from GrayBoard "
                 "with the post-migration counts reconciling against the pre-migration totals and sampled "
                 "records verified for fidelity. Cutover was performed during the December–January teaching "
                 "break, with GrayBoard retained read-only as a rollback position until acceptance. Staff "
                 "training was delivered before cutover.")
    ex.para(doc, "Acceptance testing confirmed DOODLE meets the ≥ 99% availability target, matches or "
                 "improves on GrayBoard's performance at peak load, and conforms to WCAG 2.1 Level AA. The "
                 "infrastructure, as-built documentation and operational runbooks were handed to YAT ICT, and "
                 "the engagement closed on 16 December 2022.")

    h1("2. Engagement Context")
    ex.para(doc, "This report covers Phases 2–3 of the LMS Replacement engagement — deployment and data "
                 "migration, then stabilisation and closure — implementing the approved LMS Replacement "
                 "Solution Design. Unlike a cloud-infrastructure engagement, AccentLoitte's scope was "
                 "end-to-end, so the application installation, the data migration, the cutover and the staff "
                 "training were all AccentLoitte's responsibility.")

    h1("3. Scope of Deployment")
    ex.para(doc, "In scope: provisioning the on-premises server, installing and configuring DOODLE, "
                 "integrating it with Active Directory / Office 365 / AVETMISS, migrating the GrayBoard "
                 "student records, cutover, staff training, stabilisation support, and the closure pack. Out "
                 "of scope: ongoing operational support after acceptance (separately contracted) and cloud "
                 "hosting (a future consideration).")
    h3("3.1 Maintenance-window context")
    ex.para(doc, "Cutover was a change to the live LMS, performed during the December–January teaching break "
                 "to minimise disruption. GrayBoard was retained read-only throughout as the rollback "
                 "position; the cutover completed within the break and no rollback was required.")

    h1("4. Build / Change Narrative")
    h3("4.1 Identity and access")
    ex.para(doc, "DOODLE was integrated with Active Directory via LDAP bind for single sign-on; the existing "
                 "LDAP user-group structures were preserved, service accounts created, and MFA applied to "
                 "staff with grading and course-management roles.")
    h3("4.2 Network")
    ex.para(doc, "The server was placed in the staff network zone, with student-zone access to the LMS "
                 "governed by AD permissions, on the redundant (dual-NIC) campus network. Office 365 SMTP "
                 "relay was configured for LMS notifications and firewall rules opened for vendor updates.")
    h3("4.3 Compute (server)")
    ex.para(doc, "A new mid-range enterprise server (dual PSU, RAID-protected disks) was provisioned with "
                 "Windows Server 2016 in the UPS-protected campus computer room, sized to the GrayBoard "
                 "performance baseline plus the end-of-term peak load.")
    h3("4.4 Load balancing")
    na(doc, "single-server on-premises deployment — no load balancer.")
    h3("4.5 Database and data migration")
    ex.para(doc, "MySQL was installed on the server and the DOODLE schema created. The GrayBoard student "
                 "records were migrated via a full export and import during the cutover window; record "
                 "counts were reconciled against the source and a sample of records verified for content "
                 "fidelity (see §6.3).")
    h3("4.6 Storage")
    ex.para(doc, "RAID-protected local disk holds the OS, the DOODLE application and the MySQL database; the "
                 "campus NAS (RAID-5) holds course attachments, student submissions and backup copies.")
    h3("4.7 Security")
    ex.para(doc, "Access is controlled by Active Directory permissions and the two-zone separation, with "
                 "antivirus/EDR on the server, HTTPS for application traffic and TLS for the database "
                 "connection. WCAG 2.1 Level AA conformance was confirmed (see §6.4).")
    h3("4.8 Monitoring")
    ex.para(doc, "DOODLE was brought under the existing YAT system-management and monitoring server, with "
                 "nightly backups configured and monitored per the Backup and Retention Policy.")
    h3("4.9 Cross-Region backup / replication")
    na(doc, "on-premises deployment — cross-Region cloud backup does not apply; backup is to the "
            "system-management server with offsite tape rotation.")

    h1("5. Configuration Decisions")
    ex.etable(doc, ["#", "Decision", "Choice", "Rationale"],
              [["C1", "Server specification", "Mid-range enterprise (dual PSU, RAID-10 SSD)",
                "Meets the GrayBoard baseline + ~3× end-of-term peak with headroom"],
               ["C2", "MySQL engine version", "DOODLE-supported MySQL release",
                "Confirmed against the DOODLE compatibility matrix"],
               ["C3", "Storage sizing", "Local RAID for app/DB; NAS for attachments + backups",
                "Sized from the migrated data footprint plus growth"],
               ["C4", "Backup schedule", "Nightly to the management server + offsite tape",
                "Aligned to the existing backup windows (per the Backup and Retention Policy)"]],
              widths=[1.0, 4.2, 5.0, 5.3])

    h1("6. Testing, Simulation and Validation")
    h3("6.1 Connectivity and integration tests")
    ex.etable(doc, ["Test", "Outcome", "Notes"],
              [["Active Directory single sign-on", "Pass", "Staff + student login via AD"],
               ["Office 365 SMTP notifications", "Pass", "LMS notifications delivered"],
               ["AVETMISS / ASQA export", "Pass", "Statutory export produced correctly"]],
              widths=[7.0, 3.0, 5.5])
    h3("6.2 Autoscaling test")
    na(doc, "single-server on-premises deployment — no autoscaling.")
    h3("6.3 Database connectivity and data-migration reconciliation")
    ex.para(doc, "The application connected to MySQL successfully. Post-migration record counts reconciled "
                 "against the pre-migration GrayBoard totals, and a sample of student records (enrolments, "
                 "submissions, grades) was verified for content fidelity — all matched.")
    h3("6.4 End-to-end acceptance test")
    ex.para(doc, "DOODLE was reachable and login succeeded via AD. Load testing at ~3× typical concurrent "
                 "users confirmed response times met or improved on the GrayBoard baseline, and an "
                 "independent WCAG 2.1 Level AA accessibility audit confirmed conformance.")
    h3("6.5 Failure simulation")
    na(doc, "no high-availability design to simulate against — this is a single-server on-premises platform.")
    h3("6.6 Resize simulation")
    na(doc, "not applicable to this deployment.")
    h3("6.7 Availability measurement")
    ex.para(doc, "Availability was tracked through the stabilisation period via the system-management "
                 "monitoring server and met the ≥ 99% target.")
    h3("6.8 Simulation findings vs the design")
    na(doc, "no simulations were in scope (see §6.5–6.6).")
    h3("6.9 Adjustments made per simulation outcomes")
    na(doc, "no simulations were in scope.")

    h1("7. Operational Handover")
    h3("7.1 Access")
    ex.para(doc, "Administrative access was handed to YAT ICT; AD-based access for staff and students "
                 "continued unchanged, with MFA retained for grading roles.")
    h3("7.2 Runbook references")
    add_bullet_list(doc, [
        "The LMS Replacement Solution Design — the operational reference for the deployed architecture.",
        "The as-built documentation set and operational runbooks produced at closure.",
        "The Backup and Retention Policy and the monitoring on the system-management server.",
    ])
    h3("7.3 Known limitations and what's next")
    add_bullet_list(doc, [
        "Single-server on-premises platform — not highly available; cloud hosting noted as a future consideration.",
        "Ongoing operational support after acceptance is separately contracted.",
    ])
    h3("7.4 Documentation filing")
    ex.etable(doc, ["Item", "Filed in", "Reference"],
              [["This Deployment Report + Closure Pack", "YAT ICT shared documentation", "[ref]"],
               ["As-built documentation + runbooks", "YAT ICT shared documentation", "[ref]"],
               ["Migration reconciliation + acceptance evidence", "YAT ICT shared documentation", "[ref]"]],
              widths=[6.5, 5.0, 4.0])
    h3("7.5 Feedback record")
    ex.etable(doc, ["Feedback received", "From", "Response", "Resulting action"],
              [["Request additional admin-staff training on the new reporting screens",
                "Sam Walker (YAT ICT Manager)", "Scheduled a follow-up admin session in the stabilisation period",
                "Session delivered; covered in the closure pack"]],
              widths=[5.0, 3.2, 4.0, 3.8])
    h3("7.6 Sign-off")
    ex.etable(doc, ["Role", "Name", "Date", "Acceptance"],
              [["Prepared by", "AccentLoitte Consultant", "Dec 2022", "Submitted"],
               ["Reviewed by", "Jamie O'Donnell (AccentLoitte)", "Dec 2022", "Approved for submission"],
               ["Approved by", "Sam Walker (YAT ICT Manager)", "16 Dec 2022",
                "Replacement accepted; Closure Pack signed off; engagement closed"]],
              widths=[4.0, 5.0, 2.5, 4.0])

    # ---- APPENDICES ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1("Appendix A — Build evidence")
    ex.para(doc, "The build evidence — server configuration, network placement, the DOODLE installation, and "
                 "the AD / Office 365 / AVETMISS integration configuration — is captured in the as-built "
                 "documentation filed with this report in the YAT ICT documentation repository.")
    h1("Appendix B — Configuration records")
    ex.para(doc, "Server, database, storage and backup configuration records are filed with the as-built "
                 "documentation.")
    h1("Appendix C — Test and migration evidence")
    ex.para(doc, "The data-migration reconciliation report, the performance test results, and the WCAG 2.1 "
                 "AA accessibility audit report are filed with this report.")

    h1("Document control")
    ex.etable(doc, ["Field", "Value"],
              [["Document version", "v1.0 — Final"],
               ["Prepared by", "AccentLoitte Consulting"],
               ["Engagement", "YAT LMS Replacement Project — Phases 2–3"],
               ["Related documents", "LMS Replacement Solution Design; LMS Replacement Business Case"],
               ["Engagement status", "Closed 16 December 2022"]],
              widths=[5.0, 11.0])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website/public/documents/YAT-LMS-Replacement-Deployment-Report.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
