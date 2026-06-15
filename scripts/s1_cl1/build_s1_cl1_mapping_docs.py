"""Populate the S1-CL1 per-UoC Assessment Mapping docx files.

For each of the 3 unit mapping docx files (ICTCLD401, ICTCLD502, ICTICT517),
write the criterion code(s) (e.g. "A4", "A4, B3") into the AT1/AT2/AT3 columns
for each UoC item row in:

- Table 3 — Critical Aspects / Assessment Conditions
    (special: original block row preserved; per-AC rows inserted below)
- Table 4 — Elements / Performance Criteria (PCs)
    (match by parsing leading "X.Y" from col 1)
- Table 5 — Required Skills / Performance Evidence (PEs)
    (positional: data rows in source-UoC order)
- Table 6 — Required Knowledge / Knowledge Evidence (KEs)
    (positional: data rows in source-UoC order)
- Table 7 — Foundation Skills (FSs)
    (match by skill name in col 0)

Data is sourced from the AT1, AT2, AT3 assessor docx reverse-maps + the
assessment_plan.md group coverage map. Reflects the 2026-05-25 502 PC
reassignment (PCs 1.3, 4.1, 4.2, 4.3 moved AT3 → AT2) and the AT3 v1.0
simplified shape (no closure pack; no SRM; HA Deployment Report carries
PCs 5.1, 5.2, 5.3).

USAGE:
    python populate_mapping_docs.py 401     # do just ICTCLD401
    python populate_mapping_docs.py 502     # do just ICTCLD502
    python populate_mapping_docs.py 517     # do just ICTICT517
    python populate_mapping_docs.py all     # do all three

Each run backs up the docx to <name>.docx.bak before editing.
"""

import copy
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

REPO = Path(__file__).resolve().parents[2]  # scripts/s1_cl1/ -> content repo root
MAPPINGS = REPO / "S1-CL1-Cloud-Design-Build" / "mappings"

# Column indices in the mapping tables (0-indexed)
# Table 3 (AC): col 0 = description; cols 1,2,3,4,5 = AT1..AT5
# Table 4 (PC): col 0 = element; col 1 = PC; cols 2,3,4,5,6 = AT1..AT5
# Table 5 (PE): col 0 = PE description; cols 1,2,3,4,5 = AT1..AT5
# Table 6 (KE): col 0 = KE description; cols 1,2,3,4,5 = AT1..AT5
# Table 7 (FS): col 0 = skill; col 1 = PC (often empty); col 2 = description;
#                cols 3,4,5,6,7 = AT1..AT5

# -----------------------------------------------------------------------------
# Mapping data: UoC item → {AT1, AT2, AT3} criterion code(s)
# Empty string means "no criterion evidences this in that AT" (cell stays blank).
# -----------------------------------------------------------------------------

DATA_401 = {
    # PC text starts with "X.Y" — match by leading prefix
    "pcs": {
        "1.1": {"AT1": "", "AT2": "A5", "AT3": ""},
        "1.2": {"AT1": "", "AT2": "A4", "AT3": ""},
        "1.3": {"AT1": "", "AT2": "A5", "AT3": ""},
        "1.4": {"AT1": "", "AT2": "A2", "AT3": ""},
        "1.5": {"AT1": "", "AT2": "A2", "AT3": ""},
        "1.6": {"AT1": "", "AT2": "A2", "AT3": ""},
        "1.7": {"AT1": "", "AT2": "A2", "AT3": ""},
        "1.8": {"AT1": "A4", "AT2": "", "AT3": ""},
        "2.1": {"AT1": "", "AT2": "A2", "AT3": ""},
        "2.2": {"AT1": "", "AT2": "A2", "AT3": ""},
        "2.3": {"AT1": "", "AT2": "A3", "AT3": ""},
        "2.4": {"AT1": "", "AT2": "A3", "AT3": ""},
        "2.5": {"AT1": "", "AT2": "A3", "AT3": ""},
        "2.6": {"AT1": "", "AT2": "A6", "AT3": ""},
        "3.1": {"AT1": "", "AT2": "A4", "AT3": ""},
        "3.2": {"AT1": "", "AT2": "A6", "AT3": ""},
        "4.1": {"AT1": "A14, B4", "AT2": "", "AT3": "B1"},
        "4.2": {"AT1": "B5", "AT2": "", "AT3": ""},
        "4.3": {"AT1": "", "AT2": "A7", "AT3": "B9"},
    },
    # PEs in source order
    "pes": [
        # PE 1 — build at least one simple virtual network
        {"AT1": "", "AT2": "A9", "AT3": ""},
        # PE 2 — configure compute, storage, database, autoscaling
        {"AT1": "", "AT2": "A9, A10", "AT3": ""},
        # PE 3 — conduct simple tests to confirm access
        {"AT1": "", "AT2": "A11", "AT3": ""},
    ],
    # KEs in source order
    "kes": [
        # KE 1 — industry technology standards
        {"AT1": "A11", "AT2": "", "AT3": ""},
        # KE 2 — industry standard hardware/software products
        {"AT1": "A11", "AT2": "", "AT3": ""},
        # KE 3 — IaaS/PaaS/SaaS
        {"AT1": "A11", "AT2": "", "AT3": ""},
        # KE 4 — cost models
        {"AT1": "A5, A11", "AT2": "", "AT3": ""},
        # KE 5 — VM/networking/scaling features
        {"AT1": "", "AT2": "A8", "AT3": ""},
        # KE 6 — vertical/horizontal scaling, db types, storage options
        {"AT1": "", "AT2": "A8", "AT3": ""},
        # KE 7 — shared security responsibility
        {"AT1": "", "AT2": "A8", "AT3": ""},
        # KE 8 — user access protocols
        {"AT1": "", "AT2": "A8", "AT3": ""},
        # KE 9 — security policies and protocols
        {"AT1": "", "AT2": "A8", "AT3": ""},
        # KE 10 — DNS
        {"AT1": "", "AT2": "A8", "AT3": ""},
        # KE 11 — cloud models (on-prem/private/hybrid/public)
        {"AT1": "A11", "AT2": "", "AT3": ""},
    ],
    # FSs keyed by skill name (matches col 0 of Table 7)
    "fss": {
        "Reading": {"AT1": "", "AT2": "A13", "AT3": "B16"},
        "Writing": {"AT1": "", "AT2": "A13", "AT3": "A7, B16"},
        "Learning": {"AT1": "", "AT2": "A12", "AT3": "B15"},
        "Planning and organising": {"AT1": "", "AT2": "A12", "AT3": "B15"},
        "Self-management skills": {"AT1": "", "AT2": "A12", "AT3": "B15"},
    },
    # ACs in source order — short label + per-AT codes; rows inserted into Table 3
    "acs": [
        ("cloud vendor service provider",
         {"AT1": "", "AT2": "C1", "AT3": "C1"}),
        ("cloud managed database service",
         {"AT1": "", "AT2": "C1", "AT3": "C1"}),
        ("internet and web browser",
         {"AT1": "", "AT2": "C1", "AT3": "C1"}),
        ("data to gather information from to determine output and user requirements, including user access and business protocols",
         {"AT1": "C1", "AT2": "C2", "AT3": "C2"}),
        ("Assessors of this unit must satisfy the requirements for assessors in applicable vocational education and training legislation, frameworks and/or standards",
         {"AT1": "✓", "AT2": "✓", "AT3": "✓"}),
    ],
}

DATA_502 = {
    "pcs": {
        "1.1": {"AT1": "", "AT2": "", "AT3": "A1"},
        "1.2": {"AT1": "A4", "AT2": "", "AT3": ""},
        "1.3": {"AT1": "", "AT2": "A2", "AT3": ""},
        "2.1": {"AT1": "", "AT2": "", "AT3": "A2"},
        "2.2": {"AT1": "", "AT2": "", "AT3": "A2"},
        "2.3": {"AT1": "", "AT2": "", "AT3": "A2"},
        "2.4": {"AT1": "", "AT2": "", "AT3": "A2"},
        "2.5": {"AT1": "", "AT2": "", "AT3": "A3"},
        "3.1": {"AT1": "", "AT2": "", "AT3": "A4"},
        "3.2": {"AT1": "", "AT2": "", "AT3": "A5"},
        "3.3": {"AT1": "", "AT2": "", "AT3": "A5"},
        "3.4": {"AT1": "", "AT2": "", "AT3": "A5"},
        "3.5": {"AT1": "", "AT2": "", "AT3": "A6"},
        "4.1": {"AT1": "", "AT2": "A3", "AT3": "B2"},
        "4.2": {"AT1": "", "AT2": "A6, A11", "AT3": "B3"},
        "4.3": {"AT1": "", "AT2": "A4, A11", "AT3": "B6"},
        "4.4": {"AT1": "", "AT2": "", "AT3": "B4"},
        "4.5": {"AT1": "", "AT2": "", "AT3": "B5"},
        "4.6": {"AT1": "", "AT2": "", "AT3": "B7"},
        "5.1": {"AT1": "", "AT2": "", "AT3": "B8"},
        "5.2": {"AT1": "B6", "AT2": "", "AT3": "B10"},
        "5.3": {"AT1": "", "AT2": "", "AT3": "B11"},
    },
    "pes": [
        # PE 1 — design + implement at least one fault tolerant cloud infra
        {"AT1": "", "AT2": "", "AT3": "A4, B2, B13"},
        # PE 2 — design + deploy automated infrastructure scaling
        {"AT1": "", "AT2": "", "AT3": "A4, B2, B13"},
        # PE 3 — simulate failures + demonstrate fault tolerance
        {"AT1": "", "AT2": "", "AT3": "B4, B14"},
        # PE 4 — use cloud management console / SDK / CLI
        {"AT1": "", "AT2": "", "AT3": "B2, B13"},
        # PE 5 — define, monitor and record resource availability
        {"AT1": "", "AT2": "", "AT3": "B6, B14"},
    ],
    "kes": [
        # KE 1 — industry technology standards
        {"AT1": "A11", "AT2": "", "AT3": ""},
        # KE 2 — industry hardware/software products
        {"AT1": "A11", "AT2": "", "AT3": ""},
        # KE 3 — cloud cost models re scalability
        {"AT1": "A5, A11", "AT2": "", "AT3": ""},
        # KE 4 — HA cloud infra concepts (FT, SPOFs, MTTF/MTTR/MTBF, RTO/RPO, SLAs, scalability)
        {"AT1": "", "AT2": "", "AT3": "B12"},
        # KE 5 — testing and debugging techniques
        {"AT1": "", "AT2": "", "AT3": "B12"},
        # KE 6 — tools and techniques to measure availability impact
        {"AT1": "", "AT2": "", "AT3": "B12"},
        # KE 7 — built-in vs designed fault tolerance
        {"AT1": "", "AT2": "", "AT3": "B12"},
        # KE 8 — load balancing + autoscaling for availability
        {"AT1": "", "AT2": "", "AT3": "B12"},
        # KE 9 — performance monitoring techniques and metrics
        {"AT1": "", "AT2": "", "AT3": "B12"},
    ],
    "fss": {
        "Oral communication": {"AT1": "B9", "AT2": "", "AT3": ""},
        "Reading": {"AT1": "", "AT2": "", "AT3": "A7, B16"},
        "Problem solving": {"AT1": "", "AT2": "", "AT3": "B15"},
        "Self-management": {"AT1": "", "AT2": "", "AT3": "B15"},
    },
    "acs": [
        ("cloud vendor service provider",
         {"AT1": "", "AT2": "C1", "AT3": "C1"}),
        ("cloud managed database service",
         {"AT1": "", "AT2": "C1", "AT3": "C1"}),
        ("information and data sources required to design and implement cloud infrastructure",
         {"AT1": "C1", "AT2": "C2", "AT3": "C2"}),
        ("integrated development environment (IDE)",
         {"AT1": "", "AT2": "C1", "AT3": "C1"}),
        ("specific requirements and industry standards, organisational procedures and legislative requirements, including business and functionality requirements, as required",
         {"AT1": "C1", "AT2": "C2", "AT3": "C2"}),
        ("internet and web browser",
         {"AT1": "", "AT2": "C1", "AT3": "C1"}),
        ("secure shell (SSH) or remote desktop protocol (RDP) client to connect to cloud-hosted instances",
         {"AT1": "", "AT2": "C1", "AT3": "C1"}),
        ("data to gather information from to determine output and user requirements, including user access and business protocols",
         {"AT1": "C1", "AT2": "C2", "AT3": "C2"}),
        ("Assessors of this unit must satisfy the requirements for assessors in applicable vocational education and training legislation, frameworks and/or standards",
         {"AT1": "✓", "AT2": "✓", "AT3": "✓"}),
    ],
}

DATA_517 = {
    "pcs": {
        "1.1": {"AT1": "A1", "AT2": "", "AT3": ""},
        "1.2": {"AT1": "A2", "AT2": "", "AT3": ""},
        "1.3": {"AT1": "A3", "AT2": "", "AT3": ""},
        "1.4": {"AT1": "A10, B1", "AT2": "", "AT3": ""},
        "2.1": {"AT1": "A4, A5, A6", "AT2": "", "AT3": ""},
        "2.2": {"AT1": "A4, A6", "AT2": "", "AT3": ""},
        "2.3": {"AT1": "A7", "AT2": "", "AT3": ""},
        "2.4": {"AT1": "A13, B2", "AT2": "", "AT3": ""},
        "3.1": {"AT1": "A8", "AT2": "", "AT3": ""},
        "3.2": {"AT1": "A8", "AT2": "", "AT3": ""},
        "3.3": {"AT1": "A9, A14, B3", "AT2": "", "AT3": ""},
    },
    # NOTE: 517 PE table is structurally different — the source UoC has one
    # top-level bullet "For one organisation:" with 6 nested sub-bullets,
    # and the mapping docx has a single row containing all 6 sub-bullets
    # joined together. We need to split that single row into 6 per-PE rows
    # (like the AC table). Use the (label, codes) tuple format so the splitter
    # can populate col 0 with the PE text.
    "pes_split": [
        ("interpret a strategic plan and objectives of the organisation",
         {"AT1": "A1", "AT2": "", "AT3": ""}),
        ("evaluate the current state of ICT in the organisation",
         {"AT1": "A2", "AT2": "", "AT3": ""}),
        ("identify possible gaps and opportunities in ICT and evaluate organisational impact with reference to the strategic plan and the objectives",
         {"AT1": "A3", "AT2": "", "AT3": ""}),
        ("determine and prioritise proposed changes to meet organisational needs",
         {"AT1": "A7", "AT2": "", "AT3": ""}),
        ("evaluate the difficulty of implementing proposed changes",
         {"AT1": "A6", "AT2": "", "AT3": ""}),
        ("develop an action plan for implementation",
         {"AT1": "A8", "AT2": "", "AT3": ""}),
    ],
    "kes": [
        # KE 1 — Key sections of an action plan
        {"AT1": "A8, A11", "AT2": "", "AT3": ""},
        # KE 2 — Methods of evaluation and planning approaches
        {"AT1": "A5, A11", "AT2": "", "AT3": ""},
        # KE 3 — Methods of evaluation of competing ICT systems and products
        {"AT1": "A4, A5, A11", "AT2": "", "AT3": ""},
        # KE 4 — Current and emerging trends
        {"AT1": "A11", "AT2": "", "AT3": ""},
    ],
    "fss": {
        "Reading": {"AT1": "A12", "AT2": "", "AT3": ""},
        "Writing": {"AT1": "A15", "AT2": "", "AT3": ""},
        "Oral Communication": {"AT1": "B7, B8", "AT2": "", "AT3": ""},
        "Numeracy": {"AT1": "A5", "AT2": "", "AT3": ""},
        # FSs below were originally unmapped in the AT1 marking guide; mapped to
        # closest-fit AT1 criteria 2026-05-26 per Tim's "avoid changing the
        # assessments at this stage if we can manage it". The AT1 assessor docx
        # criterion UoC columns are NOT updated to claim these — known
        # inconsistency to address in a later pass if needed for audit.
        "Navigate the world of work": {"AT1": "A8", "AT2": "", "AT3": ""},
        "Interact with others": {"AT1": "B5, B6", "AT2": "", "AT3": ""},
        "Get the work done": {"AT1": "A4, A8, A12", "AT2": "", "AT3": ""},
    },
    "acs": [
        ("A site where ICT needs and strategic directions of the organisation are coordinated",
         {"AT1": "C1", "AT2": "", "AT3": ""}),
        ("Detailed information relating to a strategic organisation plan, objectives, and direction",
         {"AT1": "C1", "AT2": "", "AT3": ""}),
        ("Organisational policies and procedures relating to the implementation of ICT changes",
         {"AT1": "C1", "AT2": "", "AT3": ""}),
        ("Individual superior in the organisation",
         {"AT1": "C2", "AT2": "", "AT3": ""}),
        ("Information on current ICT systems and practices in the organisation including operating systems, hardware, and security",
         {"AT1": "C1", "AT2": "", "AT3": ""}),
        ("Assessors of this unit must satisfy the assessor requirements in applicable vocational education and training legislation, frameworks and/or standards",
         {"AT1": "✓", "AT2": "✓", "AT3": "✓"}),
    ],
}

UNIT_DATA = {
    "401": ("ICTCLD401_Assessment_Mapping.docx", DATA_401),
    "502": ("ICTCLD502_Assessment_Mapping.docx", DATA_502),
    "517": ("ICTICT517_Assessment_Mapping.docx", DATA_517),
}

# Table indices (zero-based) — confirmed by inspection of the three docx files
TABLE_AC = 3
TABLE_PC = 4
TABLE_PE = 5
TABLE_KE = 6
TABLE_FS = 7

# Column indices for the AT columns
# Tables 3, 5, 6: cols 1,2,3,4,5 = AT1..AT5 (col 0 = description)
# Table 4 (PC):   cols 2,3,4,5,6 = AT1..AT5 (col 0 = element, col 1 = PC)
# Table 7 (FS):   cols 3,4,5,6,7 = AT1..AT5 (cols 0/1/2 = skill / PC / description)
AT_COLS_AC = (1, 2, 3)  # AT1, AT2, AT3 in Table 3
AT_COLS_PC = (2, 3, 4)
AT_COLS_PE = (1, 2, 3)
AT_COLS_KE = (1, 2, 3)
AT_COLS_FS = (3, 4, 5)


def set_cell_text(cell, value: str) -> None:
    """Replace cell content with `value` while preserving the first paragraph
    and run structure (so existing formatting carries). If the cell already has
    non-empty text, leave it untouched (don't overwrite Tim's prior edits).
    """
    if cell.text.strip():
        return  # respect prior content
    # Clear paragraphs, then write into the first one
    for p in cell.paragraphs:
        for r in list(p.runs):
            r.text = ""
    if cell.paragraphs:
        first = cell.paragraphs[0]
        if first.runs:
            first.runs[0].text = value
        else:
            first.add_run(value)
    else:
        cell.add_paragraph(value)


def populate_pcs(table, pcs_data: dict, results: list) -> None:
    """For each data row in Table 4, parse the leading 'X.Y ' from col 1 and
    look up the mapping. Skip header rows."""
    for ri, row in enumerate(table.rows):
        c1 = row.cells[1].text.strip()
        # The PC text starts with "X.Y " — extract that prefix
        import re
        m = re.match(r"^(\d+\.\d+)\b", c1)
        if not m:
            continue  # header row or untextual
        pc_num = m.group(1)
        if pc_num not in pcs_data:
            results.append(("PC", pc_num, "NO DATA", None))
            continue
        codes = pcs_data[pc_num]
        for col_idx, at_key in zip(AT_COLS_PC, ("AT1", "AT2", "AT3")):
            value = codes.get(at_key, "")
            if value:
                set_cell_text(row.cells[col_idx], value)
        results.append(("PC", pc_num, "OK", codes))


def populate_positional(table, data_list: list, at_cols, label: str, results: list) -> None:
    """For tables where each data row maps to the next entry in data_list.
    Skip the first two rows (header rows) and apply data in order."""
    data_rows = list(table.rows)[2:]  # rows 0 and 1 are headers
    if len(data_rows) != len(data_list):
        results.append((label, "ROW COUNT MISMATCH", f"docx has {len(data_rows)}, data has {len(data_list)}", None))
        return
    for i, (row, codes) in enumerate(zip(data_rows, data_list)):
        for col_idx, at_key in zip(at_cols, ("AT1", "AT2", "AT3")):
            value = codes.get(at_key, "")
            if value:
                set_cell_text(row.cells[col_idx], value)
        results.append((label, f"#{i+1}", "OK", codes))


def populate_fss(table, fss_data: dict, results: list) -> None:
    """For each data row in Table 7, match col 0 (skill name) against fss_data."""
    for ri, row in enumerate(table.rows):
        c0 = row.cells[0].text.strip()
        if c0 in ("", "Foundation Skills (If specified in the UOC)",
                  "Skill (Copy and paste these from the UOC)"):
            continue
        if c0 not in fss_data:
            results.append(("FS", c0, "NO DATA", None))
            continue
        codes = fss_data[c0]
        for col_idx, at_key in zip(AT_COLS_FS, ("AT1", "AT2", "AT3")):
            value = codes.get(at_key, "")
            if value:
                set_cell_text(row.cells[col_idx], value)
        results.append(("FS", c0, "OK", codes))


def insert_ac_rows(table, acs_data: list, results: list) -> None:
    """Populate the AC table (or any split-style table).

    Two cases are handled:

    (a) **Single-block source row** — the table has exactly one data row
        containing the full AC block joined together (e.g. 401's AC table as
        shipped from the source UoC). Action: preserve the block row as
        context; insert one new per-AC row below it for each item in
        acs_data.

    (b) **Pre-split source rows** — the table already has multiple data rows,
        one per AC item, with the AT columns empty (typically Tim's prior
        in-flight editing for 502 and 517). Action: match each acs_data
        item to an existing row by substring (case-insensitive). Populate
        the AT columns of the matched row. For any items not matched,
        append a new row at the end.

    Cell-text preservation: if a cell already has content (e.g. from a
    prior partial edit), leave it alone — don't overwrite.
    """
    data_rows = list(table.rows)[2:]  # skip 2 header rows
    pre_split = len(data_rows) > 1

    # The "template" row carries the styling for any new rows we insert.
    # When pre-split, use the last data row as template; when single-block,
    # use the only data row.
    template_xml = data_rows[-1]._tr
    insert_after = template_xml

    matched_row_ids = set()

    for label, codes in acs_data:
        # When pre-split, first try to match the label against an existing row
        if pre_split:
            matched = None
            for row in data_rows:
                if id(row._tr) in matched_row_ids:
                    continue
                row_text = row.cells[0].text.strip().lower()
                if label.lower() in row_text:
                    matched = row
                    break
            if matched is not None:
                matched_row_ids.add(id(matched._tr))
                for col_idx, at_key in zip(AT_COLS_AC, ("AT1", "AT2", "AT3")):
                    value = codes.get(at_key, "")
                    if value:
                        set_cell_text(matched.cells[col_idx], value)
                results.append(("AC", label[:60], "MATCHED EXISTING", codes))
                continue
            # Fall through to insert if no match

        # Insert a new row using the template
        new_tr = copy.deepcopy(template_xml)
        insert_after.addnext(new_tr)
        insert_after = new_tr  # next row will go after this one
        # Find the new row by matching XML element identity
        new_row = None
        for r in table.rows:
            if r._tr is new_tr:
                new_row = r
                break
        if new_row is None:
            results.append(("AC", label[:40], "ROW LOOKUP FAILED", None))
            continue
        # Clear col 0 and set to the AC label
        cell0 = new_row.cells[0]
        for p in cell0.paragraphs:
            for r in list(p.runs):
                r.text = ""
        # Remove all paragraphs except one, then set its text
        if cell0.paragraphs:
            first_p = cell0.paragraphs[0]
            if first_p.runs:
                first_p.runs[0].text = label
            else:
                first_p.add_run(label)
            # Remove any additional paragraphs (the AC block had many)
            for p in cell0.paragraphs[1:]:
                p._element.getparent().remove(p._element)
        # Fill AT columns
        for col_idx, at_key in zip(AT_COLS_AC, ("AT1", "AT2", "AT3")):
            value = codes.get(at_key, "")
            if value:
                # The template row's AT columns may have inherited content;
                # explicitly clear then set
                cell = new_row.cells[col_idx]
                for p in cell.paragraphs:
                    for r in list(p.runs):
                        r.text = ""
                if cell.paragraphs:
                    first_p = cell.paragraphs[0]
                    if first_p.runs:
                        first_p.runs[0].text = value
                    else:
                        first_p.add_run(value)
                for p in cell.paragraphs[1:]:
                    p._element.getparent().remove(p._element)
            else:
                # Clear the cell completely
                cell = new_row.cells[col_idx]
                for p in cell.paragraphs:
                    for r in list(p.runs):
                        r.text = ""
                for p in cell.paragraphs[1:]:
                    p._element.getparent().remove(p._element)
        results.append(("AC", label[:60], "OK", codes))


def process_unit(unit_key: str) -> None:
    fname, data = UNIT_DATA[unit_key]
    docx_path = MAPPINGS / fname
    backup_path = docx_path.with_suffix(".docx.bak")

    # Back up
    shutil.copy2(docx_path, backup_path)
    print(f"Backup: {backup_path}")

    doc = Document(str(docx_path))

    results = []
    populate_pcs(doc.tables[TABLE_PC], data["pcs"], results)
    if "pes" in data:
        populate_positional(doc.tables[TABLE_PE], data["pes"], AT_COLS_PE, "PE", results)
    if "pes_split" in data:
        # 517 special case: split single-row PE block into per-PE rows
        insert_ac_rows(doc.tables[TABLE_PE], data["pes_split"], results)
    populate_positional(doc.tables[TABLE_KE], data["kes"], AT_COLS_KE, "KE", results)
    populate_fss(doc.tables[TABLE_FS], data["fss"], results)
    insert_ac_rows(doc.tables[TABLE_AC], data["acs"], results)

    doc.save(str(docx_path))

    # Report
    print(f"\n--- {fname} ---")
    ok = sum(1 for r in results if r[2] == "OK")
    issues = [r for r in results if r[2] != "OK"]
    print(f"  {ok} rows populated successfully")
    if issues:
        print(f"  {len(issues)} issues:")
        for r in issues:
            print(f"    {r}")


def main():
    if len(sys.argv) != 2 or sys.argv[1] not in ("401", "502", "517", "all"):
        print("Usage: populate_mapping_docs.py {401|502|517|all}")
        sys.exit(2)
    targets = ["401", "502", "517"] if sys.argv[1] == "all" else [sys.argv[1]]
    for t in targets:
        process_unit(t)


if __name__ == "__main__":
    main()
