#!/usr/bin/env python3
"""Build the AT3 HA Deployment Report EXEMPLAR (.docx) — assessor marking reference.

A worked model answer for the YAT LMS *HA hardening* phase (S1-CL1 AT3 Part B, ICTCLD502),
on the generic Deployment Report superset template. The HA sections AT2 marked "Not
applicable" are now completed: Multi-AZ database, cross-AZ compute/ALB, failure & resize
simulations, availability measurement, DR/backup, feedback and sign-off. Together with the
AT2 exemplar this shows the template's full range — AT2 marks the HA sections N/A; AT3
completes them. Evidence is described, not captured. References its companion HA Design.

Reuses the docx brand helpers + the AT2 exemplar's na()/ev() helpers.

Usage:  python scripts/s1_cl1/build_s1_cl1_at3_ha_deployment_exemplar.py [output.docx]
Default: S1-CL1-Cloud-Design-Build/assessments/AT3/AT3-exemplar-ha-deployment-report.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_body_paragraph, add_bullet_list  # noqa: E402
from helpers.docx_evidence import add_described_evidence  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from helpers.docx_tables import add_data_table  # noqa: E402
from helpers.uoc_tags import add_uoc_evidence_tag  # noqa: E402
from helpers.yat_brand import ADDRESS, CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402
from helpers.yat_docx_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Deployment Report")
    sub = doc.add_paragraph().add_run("YAT LMS Cloud Migration — HA Hardening Phase")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    note = doc.add_paragraph().add_run("Assessor exemplar — internal marking reference (not for distribution to students)")
    note.italic = True; note.font.size = Pt(10); note.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT LMS Cloud Migration — HA Hardening"),
        ("Engagement reference", "YAT-LMS-MIG-2026"),
        ("Report version", "v1.0"),
        ("Prepared by", "MTS Consultant"),
        ("Date submitted", "[DD/MM/YYYY]"),
        ("Submitted to", "Sam Walker (YAT ICT Manager) · Pat Lin (MTS Senior Consultant)"),
        ("Related documents", "YAT LMS Cloud Architecture — HA Design v1.0 (Part A); AT2 Deployment Report"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Executive Summary")
    add_uoc_evidence_tag(doc, "[ICTCLD401 PC 4.1] — partial (document and communicate the work)")
    add_body_paragraph(doc, "This report documents the high-availability hardening of the YAT LMS, applied to the "
                 "single-AZ foundation delivered in the previous phase. Working within a single Saturday "
                 "late-night maintenance window (~3.5 hours), I enabled RDS Multi-AZ, extended the Auto "
                 "Scaling Group and Application Load Balancer across two Availability Zones (ap-southeast-2a "
                 "and 2b), added HA-tuned monitoring, and configured a cross-Region backup copy to "
                 "ap-southeast-4 (Melbourne). Failure and resize simulations confirmed the LMS now survives "
                 "an instance loss, an AZ impairment, and an RDS failover with only brief, sub-90-second "
                 "blips. Measured availability across the maintenance window was 99.96%. The 99.9% strategic "
                 "target is now achievable.")
    add_body_paragraph(doc, "Cross-Region active-active DR and application-layer clustering remain out of scope (noted "
                 "in §7.3). This report, with the companion HA Design (Part A), closes out MTS's role on the "
                 "engagement.")

    h1("2. Engagement Context")
    add_body_paragraph(doc, "This phase implements the YAT LMS Cloud Architecture — HA Design (Part A of this AT3 "
                 "submission), hardening the AT2 foundation for high availability. As the MTS consultant I "
                 "executed the change within a planned maintenance window, with the requirement to finish "
                 "either HA-complete or cleanly rolled back. Organisational change management around any "
                 "further LMS work remains YAT IT's responsibility.")

    h1("3. Scope of Deployment")
    add_body_paragraph(doc, "In scope: Multi-AZ database, cross-AZ compute and load-balancer resilience, HA-tuned "
                 "monitoring, and a cross-Region backup copy. Out of scope, carried to a future engagement: "
                 "cross-Region active-active or hot-standby DR, and application-layer HA (DOODLE clustering, "
                 "a YAT IT responsibility).")
    h3("3.1 Maintenance-window context")
    add_body_paragraph(doc, "The hardening was performed in a single simulated Saturday late-night maintenance window of "
                 "approximately 3.5 hours. The standing requirement was that the LMS be back online at the "
                 "end of the window either with HA hardening complete or rolled back to the AT2 baseline. The "
                 "work completed with brief, expected failover blips and no rollback.")

    h1("4. Build / Change Narrative")
    add_uoc_evidence_tag(doc, "[ICTCLD502 PC 4.1] (implement the design) · [ICTCLD502 PE 1, PE 2, PE 4]")
    h3("4.1 Identity and access management (changes from baseline)")
    add_body_paragraph(doc, "Minimal change: I added an AWS Backup service role and a cross-Region copy role for the "
                 "backup plan. No change to the human IAM groups or MFA enforcement from the AT2 baseline.")
    add_described_evidence(doc, "SCREENSHOT", "should show the AWS Backup service role / cross-Region copy role added; "
                            "region ap-southeast-2.")
    h3("4.2 Network topology (cross-AZ subnets)")
    add_body_paragraph(doc, "I added three subnets in ap-southeast-2b — public-web-b, private-app-b, private-data-b — "
                 "mirroring the AZ-a CIDR plan, with route-table associations. A second NAT Gateway in "
                 "public-web-b gives the AZ-b app subnet AZ-independent outbound access, so a single NAT or "
                 "AZ failure can't sever egress.")
    add_described_evidence(doc, "SCREENSHOT", "should show all six subnets across ap-southeast-2a and 2b with CIDRs/AZ, and "
                            "the second NAT Gateway.")
    h3("4.3 Compute (cross-AZ Auto Scaling)")
    add_body_paragraph(doc, "I extended the ASG across both AZ subnets and raised capacity to min 2 / desired 2 / max 4, "
                 "so the loss of one AZ still leaves a healthy instance serving while the ASG rebuilds "
                 "capacity in the survivor.")
    add_described_evidence(doc, "SCREENSHOT", "should show the ASG spanning both AZs with min/desired/max ≥ 2 and instances "
                            "healthy in each AZ.")
    h3("4.4 Application Load Balancer (cross-AZ targets)")
    add_body_paragraph(doc, "I extended the ALB's subnet group to both AZs and confirmed the target group reports "
                 "healthy targets in each AZ, so the ALB can route around an AZ-level failure.")
    add_described_evidence(doc, "SCREENSHOT", "should show the ALB target group with healthy targets in both ap-southeast-2a "
                            "and 2b.")
    h3("4.5 Database (Multi-AZ)")
    add_body_paragraph(doc, "I enabled Multi-AZ on the RDS instance; AWS provisioned a synchronous standby in "
                 "ap-southeast-2b. The conversion caused one brief failover blip, after which the primary "
                 "endpoint resolved to the new active instance. Multi-AZ removes the database single point of "
                 "failure that the AT2 baseline carried.")
    add_described_evidence(doc, "SCREENSHOT", "should show the RDS instance with Multi-AZ = Yes and the standby AZ; and the "
                            "RDS event log entry for the Multi-AZ conversion.")
    h3("4.6 Storage (cross-Region backup)")
    add_body_paragraph(doc, "I configured an AWS Backup plan with a cross-Region copy of the RDS and EBS snapshots to "
                 "ap-southeast-4 (Melbourne) — keeping the data onshore per the data-residency requirement "
                 "while protecting against a regional event.")
    add_described_evidence(doc, "SCREENSHOT", "should show the AWS Backup plan with the cross-Region copy rule to "
                            "ap-southeast-4.")
    h3("4.7 Security (HA-related adjustments)")
    add_body_paragraph(doc, "No security-group changes were required — cross-AZ traffic already flows within the VPC "
                 "under the existing tiered rules, and the new subnets inherited the tier-appropriate groups.")
    h3("4.8 Monitoring (HA-tuned)")
    add_body_paragraph(doc, "I added HA-tuned CloudWatch alarms beyond the AT2 baseline: a per-AZ healthy-host count, "
                 "RDS failover detection, replica lag, and cross-AZ latency, plus a service-level dashboard "
                 "tracking LMS availability. These detect a degradation of resilience even when the service "
                 "as a whole is still up.")
    add_described_evidence(doc, "SCREENSHOT", "should show the HA-tuned alarms list and the service-level availability "
                            "dashboard.")
    h3("4.9 Cross-Region backup / replication")
    add_body_paragraph(doc, "Completed — see §4.6 (AWS Backup cross-Region copy to ap-southeast-4).")

    h1("5. Configuration Decisions")
    add_data_table(doc, ["#", "Decision", "Choice", "Rationale"],
              [["C1", "RDS Multi-AZ apply timing", "Apply immediately (in the window)",
                "The planned maintenance window absorbs the single conversion blip; no need to wait."],
               ["C2", "Post-HA ASG capacity", "min 2 / desired 2 / max 4 across AZs",
                "Survives an AZ loss with a healthy instance still serving; max 4 covers the peak."],
               ["C3", "Cross-Region backup destination", "ap-southeast-4 (Melbourne)",
                "Regional protection while keeping financial/personal data onshore (data residency)."],
               ["C4", "HA alarm thresholds", "Healthy host < 1 per AZ (1 min); replica lag > 30 s",
                "Detects loss of AZ capacity or replication health before it becomes an outage."]],
              widths=[1.0, 4.2, 4.3, 6.0])

    h1("6. Testing, Simulation and Validation")
    add_uoc_evidence_tag(doc, "[ICTCLD502 PC 4.2–4.6, 5.1] · [ICTCLD502 PE 3, PE 5]")
    h3("6.1 Connectivity tests (regression)")
    add_body_paragraph(doc, "I re-ran the AT2 connectivity, autoscaling, database and smoke tests after the HA changes; "
                 "all passed across both AZs, confirming the hardening introduced no regression.")
    h3("6.2–6.4 Foundation tests (regression-confirmed)")
    add_body_paragraph(doc, "Autoscaling, database connectivity, and the end-to-end smoke test from the foundation phase "
                 "were re-verified and passed; detailed below is the HA-specific simulation work that is the "
                 "core of this phase.")
    h3("6.5 Failure simulation")
    add_data_table(doc, ["#", "Simulation", "Method", "Observed outcome", "Reachable?", "Evidence"],
              [["F1", "EC2 instance termination", "Terminated an instance via console",
                "ASG launched a replacement; ALB de-registered the dead target; new target healthy ~3 min",
                "Yes", "C2"],
               ["F2", "RDS Multi-AZ failover", "Reboot with failover",
                "Failover to standby in ~75 s; app reconnected on retry", "Yes (≈75 s blip)", "C3"],
               ["F3", "AZ impairment", "Terminated all instances in AZ-a",
                "ALB served from AZ-b throughout; ASG rebuilt capacity in the survivor", "Yes", "C4"]],
              widths=[0.8, 2.8, 2.8, 4.2, 2.0, 1.2])
    h3("6.6 Resize simulation")
    add_data_table(doc, ["#", "Simulation", "Method", "Observed outcome", "Availability impact"],
              [["R1", "ASG capacity increase", "Set desired = 4",
                "2 new instances healthy in ALB within ~5 min", "No downtime"],
               ["R2", "RDS instance class change", "db.m6i.large → db.m6i.xlarge, apply immediately",
                "~45 s failover blip during apply", "≤ 1 min; reconnect on retry"]],
              widths=[0.8, 3.2, 3.6, 4.0, 3.2])
    h3("6.7 Availability measurement")
    add_body_paragraph(doc, "I measured availability two ways during the window: a CloudWatch service-level dashboard, "
                 "and a curl loop against the ALB DNS name at a 1-second interval, counting failed responses. "
                 "Across the ~3.5-hour window the measured availability was 99.96%, with the only dips during "
                 "the RDS failover simulation (~75 s) and the RDS resize (~45 s).")
    add_described_evidence(doc, "TEST EVIDENCE", "should show the CloudWatch availability graph over the window and the "
                               "curl-loop summary (failed/total requests) used to compute the dips.")
    h3("6.8 Simulation findings vs the design")
    add_data_table(doc, ["Simulation", "Expected (HA Design)", "Actual", "Divergence?"],
              [["F1 EC2 termination", "Replacement, no outage", "As expected (~3 min to healthy)", "None"],
               ["F2 RDS failover", "≈ 60 s blip", "≈ 75 s blip", "Minor — within RTO; see §6.9"],
               ["F3 AZ impairment", "Survivor AZ serves", "As expected", "None"]],
              widths=[3.5, 4.0, 4.0, 4.0])
    h3("6.9 Adjustments made per simulation outcomes")
    add_body_paragraph(doc, "The RDS failover ran ~15 s longer than the design predicted. Investigation (correlating the "
                 "RDS event log against the application connection-retry log) showed the extra time was "
                 "app-side retry backoff, not the database. I shortened the connection-retry interval and "
                 "added the RDS-failover-detection alarm so the next failover is observed immediately. No "
                 "architectural change was required; the result remained within the RTO ≤ 4 h target.")

    h1("7. Operational Handover")
    add_uoc_evidence_tag(doc, "[ICTCLD401 PC 4.3] · [ICTCLD502 PC 5.2, 5.3]")
    h3("7.1 Access")
    add_body_paragraph(doc, "Unchanged from the AT2 handover: YAT-ICT-Admins full administration, MTS-Consultants "
                 "within the permission boundary, MFA enforced, instance access via SSM Session Manager.")
    h3("7.2 Runbook references")
    add_bullet_list(doc, [
        "The HA Design document (Part A) — the operational reference for the resilient architecture.",
        "The HA-tuned monitoring and alarms (§4.8) and the SNS notification destinations.",
        "The simulation outcomes (§6) — the expected behaviour for future HA testing.",
        "The cross-Region backup schedule and the restore procedure (ap-southeast-4).",
    ])
    h3("7.3 Known limitations and what's next")
    add_bullet_list(doc, [
        "Cross-Region active-active / hot-standby DR — not in scope; a future engagement.",
        "Application-layer HA (DOODLE clustering, session affinity) — YAT IT scope.",
        "Cost impact: Multi-AZ roughly doubles the database cost line — documented for YAT planning.",
    ])
    h3("7.4 Documentation filing")
    add_data_table(doc, ["Item", "Filed in", "Reference"],
              [["HA Design (Part A)", "YAT ICT shared documentation", "[ref]"],
               ["This HA Deployment Report (v1.0)", "YAT ICT shared documentation", "[ref]"],
               ["Simulation evidence (Appendix C)", "YAT ICT shared documentation", "[ref]"]],
              widths=[6.5, 5.0, 4.0])
    h3("7.5 Feedback record")
    add_data_table(doc, ["Feedback received", "From", "Response", "Resulting action"],
              [["Cost impact of Multi-AZ doubling the DB line wasn't called out in the design",
                "Sam Walker (YAT ICT Manager)", "Acknowledged; documented the cost impact in §4.5 and §7.3",
                "HA Design re-issued v1.1 with the cost note"],
               ["Confirm the cross-Region copy keeps data onshore",
                "Sam Walker", "Confirmed ap-southeast-4 (Melbourne) is an Australian region", "Noted in §4.6"]],
              widths=[5.0, 3.2, 4.0, 3.8])
    h3("7.6 Sign-off")
    add_data_table(doc, ["Role", "Name", "Date", "Acceptance"],
              [["Prepared by", "MTS Consultant", "[date]", "Submitted"],
               ["Reviewed by", "Pat Lin (MTS Senior Consultant)", "[date]", "Approved for submission"],
               ["Approved by", "Sam Walker (YAT ICT Manager)", "[date]",
                "HA hardening accepted — Multi-AZ DB, cross-AZ compute, HA monitoring confirmed operational"]],
              widths=[4.0, 5.0, 2.5, 4.0])

    h1("8. Knowledge Evidence Responses")
    add_uoc_evidence_tag(doc, "[ICTCLD502 KE 4, 5, 6, 7, 8, 9]")
    h3("Q1 — HA concepts in your design")
    add_body_paragraph(doc, "Single point of failure: the AT2 baseline's single RDS instance and single AZ were SPOFs; "
                 "the hardening removed both via Multi-AZ (§4.5) and cross-AZ compute (§4.3). Fault tolerance: "
                 "the ASG across two AZs (§4.3) tolerates an instance or AZ loss with no service outage, as "
                 "F1/F3 demonstrated. RTO/RPO: the targets (≤ 4 h / ≤ 1 h) drove the Multi-AZ choice — its "
                 "~75 s failover and the automated + cross-Region backups sit comfortably inside both, which "
                 "is why Multi-AZ (not just nightly snapshots) was the right call.")
    h3("Q2 — Testing and debugging techniques")
    add_body_paragraph(doc, "To avoid causing a single point of failure during testing, before terminating all "
                 "instances in AZ-a (F3) I confirmed AZ-b had a healthy instance in the ALB, so traffic kept "
                 "flowing while I tested the AZ loss. For debugging: when the F2 failover ran longer than "
                 "expected, I isolated the cause by lining up the RDS event-log failover timestamp, the "
                 "application connection-retry log, and the curl-loop failure window — which showed the extra "
                 "time was app-side retry backoff, not a slow database failover.")
    h3("Q3 — Tools and techniques to measure availability impact")
    add_body_paragraph(doc, "I used CloudWatch metrics and alarms, the ALB and RDS event logs, and a curl loop against "
                 "the ALB DNS name at a 1-second interval. For the RDS failover sim I computed the impact as "
                 "(failed responses × 1 s) ≈ the outage window (~75 s). The limitation is that 1-second "
                 "sampling can't see sub-second dips and the synthetic probe doesn't capture real per-user "
                 "session impact.")
    h3("Q4 — Built-in vs designed fault tolerance")
    add_body_paragraph(doc, "S3: fault tolerance is built in by AWS — objects are stored redundantly across AZs with no "
                 "action from me. RDS: a single instance has limited built-in fault tolerance (AWS handles "
                 "host/hardware, but the instance is a SPOF); Multi-AZ is fault tolerance I designed in by "
                 "enabling it. ALB: the load balancer is itself a built-in multi-AZ AWS service, but routing "
                 "across healthy targets in two AZs was a design choice I made by registering cross-AZ "
                 "targets.")
    h3("Q5 — Load balancing and autoscaling for availability")
    add_body_paragraph(doc, "The combination delivers more than either alone: the ALB stops routing to an unhealthy "
                 "target immediately, while the ASG replaces it and keeps minimum capacity — so a failed "
                 "instance is both routed around and rebuilt, and an AZ loss is survived because the ASG "
                 "holds capacity in the surviving AZ. Load balancing alone wouldn't replace the instance; "
                 "autoscaling alone wouldn't distribute traffic or health-check it. In F1, ALB de-registered "
                 "the terminated instance instantly and the ASG launched a replacement — users saw no outage.")
    h3("Q6 — Performance monitoring techniques and metrics")
    add_body_paragraph(doc, "I set up a per-AZ healthy-host-count alarm that didn't exist in the AT2 baseline. It "
                 "measures the number of healthy ALB targets in each AZ; the threshold is fewer than one "
                 "healthy host in any AZ for one minute. It detects an AZ losing all of its capacity — a "
                 "resilience degradation that the overall service-up metric would miss while the other AZ is "
                 "still serving.")

    # ---- APPENDICES ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1("Appendix A — Build evidence (screenshots)")
    add_body_paragraph(doc, "Evidence captures are described below in lieu of live screenshots for this exemplar.")
    for kind, d in [
        ("SCREENSHOT A1", "VPC subnets — all six across ap-southeast-2a and 2b with CIDRs/AZ; second NAT Gateway."),
        ("SCREENSHOT A2", "ASG spanning both AZs, min/desired/max ≥ 2, instances healthy in each AZ."),
        ("SCREENSHOT A3", "ALB target group — healthy targets in both AZs."),
        ("SCREENSHOT A4", "RDS — Multi-AZ = Yes, standby AZ; RDS event log Multi-AZ conversion."),
        ("SCREENSHOT A5", "HA-tuned CloudWatch alarms list; service-level availability dashboard."),
        ("SCREENSHOT A6", "ASG activity history — replacement event from the failure simulation."),
        ("SCREENSHOT A7", "RDS event log — failover event from the failover simulation."),
        ("SCREENSHOT A8", "CloudWatch availability metric across the 3.5-hour window, dips visible."),
        ("SCREENSHOT A9", "AWS Backup plan with cross-Region copy to ap-southeast-4."),
    ]:
        add_described_evidence(doc, kind, d)
    h1("Appendix B — Configuration exports")
    for kind, d in [
        ("CONFIG EXPORT B1", "Updated VPC/subnet/route-table config (six subnets, second NAT)."),
        ("CONFIG EXPORT B2", "Updated ASG + launch template (cross-AZ, min ≥ 2)."),
        ("CONFIG EXPORT B3", "Updated ALB + target group (both AZs)."),
        ("CONFIG EXPORT B4", "RDS instance config showing Multi-AZ."),
        ("CONFIG EXPORT B5", "HA-tuned CloudWatch alarm configurations."),
        ("CONFIG EXPORT B6", "AWS Backup plan / cross-Region copy configuration."),
    ]:
        add_described_evidence(doc, kind, d)
    h1("Appendix C — Test and simulation evidence")
    for kind, d in [
        ("TEST EVIDENCE C1", "Connectivity regression — ALB/EC2/RDS reachable across both AZs post-HA."),
        ("TEST EVIDENCE C2", "F1 — ASG activity history + CloudWatch metric showing replacement; curl log (LMS reachable throughout)."),
        ("TEST EVIDENCE C3", "F2 — RDS event log failover; app reconnect log; curl log measuring the ~75 s blip."),
        ("TEST EVIDENCE C4", "F3 — AZ impairment; ALB serving from the surviving AZ; ASG rebuild."),
        ("TEST EVIDENCE C5", "R1/R2 — ASG capacity increase and RDS resize captures."),
        ("TEST EVIDENCE C6", "Service-level availability over the window with computed % (99.96%)."),
    ]:
        add_described_evidence(doc, kind, d)

    h1("Appendix D — Reflections")
    add_uoc_evidence_tag(doc, "[ICTCLD502 FS Self-management] · [ICTCLD401 FS Learning, Planning and organising]")
    h3("R1 — Decisions in hindsight")
    add_body_paragraph(doc, "Applying Multi-AZ immediately inside the maintenance window was the right call — it cost a "
                 "single, planned failover blip rather than a second window later. If I did it again I would "
                 "not stack the RDS instance-class resize (R2) in the same window as the Multi-AZ "
                 "conversion: two failover blips close together made the availability measurement noisier "
                 "than it needed to be. I'd schedule the resize for a separate, quieter slot.")
    h3("R2 — Problem solving under time pressure")
    add_body_paragraph(doc, "When the F2 failover ran ~15 s longer than the design predicted, with the application "
                 "throwing connection errors past the point I expected, the time-boxed window meant I had to "
                 "decide quickly whether it was a database problem (which might threaten the rollback "
                 "deadline) or an application retry artefact. I correlated the three logs, saw the database "
                 "was back in ~75 s while the app kept retrying on a long backoff, and judged it safe to "
                 "proceed and tune the retry afterwards rather than roll back. Next time I'd pre-tune the "
                 "connection-retry settings before the window so the symptom doesn't appear at all.")

    h1("Document control")
    add_data_table(doc, ["Field", "Value"],
              [["Document version", "v1.0 — HA Deployment Report (post-implementation)"],
               ["Author", "MTS Consultant"],
               ["Engagement", "YAT LMS Cloud Migration — HA Hardening Phase"],
               ["Companion document", "YAT LMS Cloud Architecture — HA Design v1.0 (Part A)"],
               ["Successor document", "n/a — closes out MTS's role on the engagement"]],
              widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/assessments/AT3/AT3-exemplar-ha-deployment-report.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
