#!/usr/bin/env python3
"""Build the S1-CL1 AT1 ASSESSOR instrument (.docx) by populating the Kangan template.

This is an institutional compliance document, NOT a YAT-branded artefact: it loads the
official Kangan 'Project Assessment - Assessor' template and fills it in, preserving the
Kangan structure and styles (Details, Teacher/Assessor instructions, Marking Guide,
Instructions to Student, Benchmark). It retro-fits a generator to the approved, hand-authored
CL1 AT1 instrument, reproducing it exactly at the content level except that the literal
scenario-site URL is dropped (evergreen — students are given the live URL out-of-band, so the
instrument never goes stale; the intranet is referred to by name, as in CL2/CL3). Mirrors the
CL2 AT1 generator's mechanics.

AT1 = Business Case: YAT LMS Cloud Migration (ICTICT517 + ICTCLD401 + ICTCLD502) — two parts:
  Part A  Business Case      (written; sections 1-11 + Sign-off + Appendices 1-4)
  Part B  Presentation       (observed; presented to the role-played YAT board)

The assessor instrument carries the task instructions, the Marking Guide, the Business Case
Benchmark (model answers + worked CBA example), the Presentation Observation Process, and the
UoC coverage reverse-map. The student instrument (build_s1_cl1_at1_student) carries only the
student-facing content.

Usage:  python scripts/s1_cl1/build_s1_cl1_at1_assessor.py [output.docx]
Default: S1-CL1-Cloud-Design-Build/assessments/AT1/AT1-BusinessCase-Assessor.docx
"""
import sys
from pathlib import Path

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Assessor.docx")

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_tables import add_section_row, clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402


# ---------- content ----------

CHECK = "☐ Yes  ☐ No"  # marking-guide Satisfactory? cell (matches the Kangan template)

DETAILS = {
    "qualification": "ICT50220 Diploma of Information Technology",
    "units": [
        "ICTICT517 Match ICT needs with the strategic direction of the organisation",
        "ICTCLD401 Configure cloud services",
        "ICTCLD502 Design and implement highly-available cloud infrastructure",
    ],
    "task_title": "AT1 — Business Case: YAT LMS Cloud Migration",
    "task_number": "1 of 3",
}

OVERVIEW = [
    'Students are assessed on the production and presentation of a Business Case for the YAT LMS Cloud Migration. The assessment is split into two parts (Parts A and B — see Task(s) to be assessed below):',
    'Part A (written): the Business Case document, including all four appendices (CBA detail, Knowledge Evidence responses, supporting research, Feedback Record)',
    'Part B (observed): the Business Case presentation delivered to the role-played YAT board, with sign-off captured from the role-played YAT ICT Manager',
    'This is an open-book assessment. Students may use the YAT intranet, the AWS Pricing Calculator, AWS Academy lab environments, course reference materials, and external research (which must be cited) throughout. Where the Business Case requires synthesis (e.g. the current-state summary in §4), students must produce content in their own words — verbatim reproduction of intranet material is not satisfactory.',
    'Reasonable adjustment for this assessment may include extending the time allowed, varying the location or format of the presentation event (e.g. in-person vs video conference), allowing the student to present to a recorded audience for asynchronous assessment where appropriate, or providing one-on-one verbal explanation of board questions where needed.',
    'Teacher/assessor support level: the teacher/assessor may clarify task requirements and the scenario but must not guide the student to specific recommendations or correct answers. During the presentation, the teacher/assessor (in the role of the YAT ICT Manager) may ask substantive questions but must not coach the student through their responses.',
    'Submission: Part A (Business Case) is submitted via the LMS as a .docx with all four appendices populated. Part B (presentation deck) is submitted via the LMS as a .pptx. The completed Feedback Record and the signed Sign-off block are attached to Part A.',
    'The assessment will not proceed if for any reason it is not safe to do so. You must advise the student of the reason for suspending the assessment, and what safety action should be taken. Advise the student of revised arrangements for the assessment when it is safe to do so.',
    'There is a zero tolerance for plagiarism, cheating and collusion. Students will be expected to make a declaration that all work is their own prior to submission. Refer to the Training and Assessment Policy for further information.',
]

TASKS = [
    'In this project, students take on the role of a consultant from MP Tech Solutions (MTS) engaged by YAT College to lead a migration of their mission-critical Learning Management System (LMS) to the cloud. AT1 is the analysis and planning phase of that engagement — students produce a Business Case for the migration and present it to the YAT board for approval of the action plan.',
    'There are two parts to this assessment (Parts A and B). Each part is individually assessed as Satisfactory / Not Yet Satisfactory.',
    'Part A — Business Case',
    "Students produce a written Business Case for the YAT LMS Cloud Migration using YAT's standard Business Case template.",
    'Part B — Presentation to the YAT board',
    'Students prepare a slide deck and present their Business Case to the YAT board to seek approval of the action plan in Business Case §10.',
]

RESOURCES = [
    'Teacher/assessor supplied resources',
    "Access to the YAT intranet — provided as the engagement's reference site",
    'AWS Academy lab access — Cloud Foundations [104469] + Cloud Architecting [172221]',
    'Role-play preparation for the presentation event (board roles: Sam Walker — YAT ICT Manager; Pat Lin — MTS Senior Consultant; one peer as a board member)',
    'Presentation observation rubric (later in this document) for marking the presentation event',
    'Student supplied resources',
    'Computer with web browser',
    'Word-processing and presentation software (e.g. Microsoft Word + PowerPoint, or equivalents)',
]

CRITERIA = [
    'To receive a Satisfactory outcome for this assessment the student must:',
    'Achieve Satisfactory on every criterion in the Part A Marking Guide (Business Case — covers Business Case sections §1–§11, Sign-off, and Appendices 1–4)',
    'Achieve Satisfactory on every criterion in the Part B Marking Guide (Presentation — covers the observed presentation event including Q&A, feedback capture, and sign-off)',
    'Submit all required artefacts:',
    'Completed Business Case (.docx) with all four appendices populated',
    'Presentation deck (.pptx) with populated speaker notes',
    'Completed Feedback Record (attached as Business Case Appendix 4)',
    'Signed Sign-off block from the role-played YAT ICT Manager (within the Business Case)',
]

CONDITIONS = [
    'C1 - The YAT scenario site / intranet is accessible to the student throughout the assessment — supplying the strategic plan, current ICT environment, organisational policies, templates, reference documents, and Document Archive [ICTICT517 AC 1] · [ICTICT517 AC 2] · [ICTICT517 AC 3] · [ICTICT517 AC 5] · [ICTCLD401 AC 4] · [ICTCLD502 AC 3] · [ICTCLD502 AC 5] · [ICTCLD502 AC 8]',
    'C2 - The teacher/assessor is present in the role of YAT ICT Manager (Sam Walker) and/or MTS Senior Consultant (Pat Lin) for the Part B presentation event [ICTICT517 AC 4]',
]

# Marking guide — Part A (A1-A15) and Part B (B1-B9); each criterion is a list of cell lines
# (the criterion statement plus, where present, its UoC-traceability line).
MARKING_A = [
    ["A1 - Strategic Alignment Analysis — student analyses YAT's ICT Strategic Plan against the industry environment and organisational objectives, with citations ", '[ICTICT517 PC 1.1] · [ICTICT517 PE 1]'],
    ["A2 -Current State of YAT's ICT — student produces a synthesised summary of YAT's current ICT systems and practices in their own words (not verbatim from the intranet) ", '[ICTICT517 PC 1.2] · [ICTICT517 PE 2]'],
    ["A3 - Gap Analysis — student compares YAT's strategic plan objectives against the current state and identifies gaps, improvement opportunities, and proposed changes ", '[ICTICT517 PC 1.3] · [ICTICT517 PE 3]'],
    ['A4 - Options Considered and Evaluation — student defines the LMS workload, considers both options (in-house renewal vs cloud migration to AWS), names and justifies the evaluation method, and produces an initial impact + difficulty assessment for both options ', '[ICTICT517 KE 3] · [ICTICT517 PC 2.1] (initial) · [ICTICT517 PC 2.2] (initial) · [ICTCLD401 PC 1.8] · [ICTCLD502 PC 1.2] · [ICTICT517 FS Get the work done]'],
    ['A5 - Appendix 1 Cost-Benefit Analysis — student completes the 5-year CBA covering both options with line-item detail in Appendix 1; assumptions stated; year-by-year projections; comparison summary; avoided-downtime benefit; Year-1 cash-flow comparison; sensitivity analysis. AWS Pricing Calculator export attached in Appendix 3. Numbers reconcile between Appendix 1 detail, §7 summary, and §1 Executive Summary ', '[ICTICT517 PC 2.1] (financial impact) · [ICTICT517 KE 2] (sensitivity as evaluation method) · [ICTICT517 KE 3] (evaluation of competing options) · [ICTICT517 FS Numeracy] · [ICTCLD401 KE 1] (industry standards in selecting AWS services) · [ICTCLD401 KE 4] (cost models) · [ICTCLD502 KE 3] (cloud cost models and scalability)'],
    ['A6 - Risk and Impact Assessment — student populates intangibles comparison for both options (honest about trade-offs of the recommended option) and produces a risk register for the recommended option', '[ICTICT517 PC 2.1] (non-financial impact) · [ICTICT517 PC 2.2] (difficulty) · [ICTICT517 PE 5]'],
    ['A7 - Recommendation — student states a clear recommendation with rationale connecting back to the CBA findings, intangibles, and risk register ', '[ICTICT517 PC 2.3] · [ICTICT517 PE 4]'],
    ["A8 - Action Plan — student develops a phased action plan with prioritised changes, implementation schedule with dependencies, standards / targets / success metrics, implementation methods, alignment with YAT's change-management procedure, and an action-plan risk register ", '[ICTICT517 PC 3.1] · [ICTICT517 PC 3.2] · [ICTICT517 PE 6] · [ICTICT517 KE 1] · [ICTICT517 FS Navigate the world of work] · [ICTICT517 FS Get the work done]'],
    ['A9 - Next Steps and Decision Requested — student explicitly names what the board is being asked to approve today and what is deferred to later sign-off gates', ' [ICTICT517 PC 3.3] (the asking; granted approval is at A14)'],
    ['A10 - Executive Summary — one-page summary covering the recommendation, 5-year cost position, top 2–3 risks of the recommended option, and the asked-for decision. Numbers reconcile with A5 ', '[ICTICT517 PC 1.4] (written report-out to superior)'],
    ['A11 - Appendix 2 Knowledge Evidence — student answers all 9 KE questions (5 selection-style + 4 demonstration-style) with reference to specific sections of their own Business Case (not generic textbook answers)', '[ICTCLD401 KE 1] · [ICTCLD401 KE 2] · [ICTCLD401 KE 3] · [ICTCLD401 KE 4] · [ICTCLD401 KE 11] · [ICTCLD502 KE 1] · [ICTCLD502 KE 2] · [ICTCLD502 KE 3] · [ICTICT517 KE 1] · [ICTICT517 KE 2] · [ICTICT517 KE 3] · [ICTICT517 KE 4]'],
    ['A12 - Appendix 3 Supporting Research — AWS Pricing Calculator export attached; external research sources cited with URLs and access dates ', '[ICTICT517 FS Reading] · [ICTICT517 FS Get the work done]'],
    ['A13 - Appendix 4 Feedback Record — completed during/after the Part B presentation event using the YAT Feedback Record template; captures actual board feedback ', '[ICTICT517 PC 2.4] (documented evaluation provided to superior for feedback)'],
    ['A14 - Sign-off block — signed by the role-played YAT ICT Manager at the end of the Part B presentation event [ICTICT517 PC 3.3] (granted approval) ·', '[ICTCLD401 PC 4.1] (documented + communicated work formally signed off)'],
    ['A15 - Document quality — Business Case uses plain English, appropriate vocabulary, terminology, diagrams, numerical information, formatting and structure relevant to a professional consulting deliverable for an RTO client ', '[ICTICT517 FS Writing]'],
]

MARKING_B = [
    ['B1 - Student reports on proposed changes, gaps and improvement opportunities to the role-played superior during the meeting (walks the board through the relevant Business Case content) ', '[ICTICT517 PC 1.4] (verbal report-out — complements A10 written report-out)'],
    ['B2 - The documented evaluation process (CBA + analysis from §6–§8) is provided to the superior at/before the meeting; student walks the superior through it for feedback during the meeting ', '[ICTICT517 PC 2.4] (the act of providing + walking through — A13 covers the captured feedback record afterwards)'],
    ["B3 - Action plan is provided to the superior during the meeting; student explicitly seeks the superior's feedback and approval on it ", '[ICTICT517 PC 3.3] (verbal seeking of approval — A9 covers the written ask, A14 covers the granted sign-off)'],
    ['B4 - Student documents and communicates the work through the presentation event (deck + verbal delivery to the role-played stakeholder) [ICTCLD401 PC 4.1]'],
    ['B5 - Student seeks feedback during Q&A and responds substantively to feedback received [ICTCLD401 PC 4.2] · [ICTICT517 FS Interact with others]'],
    ['B6 - Student confirms, seeks, and responds to feedback with the role-played superior [ICTCLD502 PC 5.2] · [ICTICT517 FS Interact with others]'],
    ['B7 - Student uses plain English and translates technical terminology when necessary, communicating with the role-played superior to articulate ideas, requirements, and plans [ICTICT517 FS Oral Communication] ("Uses plain English, translating technical terminology when necessary, to communicate with a range of personnel and determine objectives, articulate ideas and requirements, and develop plans")'],
    ['B8 - Student elicits information using effective listening and questioning techniques during the Q&A ', '[ICTICT517 FS Oral Communication] ("Elicits information using effective listening and questioning techniques")'],
    ['B9 - Student uses listening and questioning techniques to articulate complex concepts and requirements using industry language for the intended audience', '[ICTCLD502 FS Oral communication]'],
]

# Shared prose — the 'Instructions to Student' body (single-sourced; the student builder imports it).
PROSE = [
    ('The engagement and your role', 'Assessor text'),
    ("YAT College is a Registered Training Organisation (RTO) based at 175 Cremorne Street, Cremorne VIC. YAT's mission-critical Learning Management System (LMS) — running on a single on-premises Windows Server 2016 / MySQL / DOODLE stack — is at the end of its hardware life and falls short of YAT's strategic availability target of 99.9%. YAT's ICT staff are highly capable on their current environment but lack cloud experience, which is why MTS has been engaged.", 'Assessor text'),
    ('You are an MP Tech Solutions (MTS) consultant reporting to Pat Lin (MTS Senior Consultant). Pat liaises with Sam Walker (YAT ICT Manager) for sign-offs. Sam will receive your Business Case and is your primary YAT-side stakeholder.', 'Assessor text'),
    ('AT1 is the analysis and planning phase of the engagement. The work you produce in AT1 flows into AT2 and AT3 of this cluster — the action plan approved at the end of AT1 becomes the brief for the build work in AT2.', 'Assessor text'),
    ('Part A — Business Case', 'Heading 2'),
    ("Using the YAT Business Case template (download from the intranet's Templates section), populate every section. The Business Case is the primary written deliverable and must contain:", 'Assessor text'),
    ('§1 Executive Summary — a one-page summary of your recommendation (write last)', 'Assessor text'),
    ('§2 Engagement context — who you are, the engagement, the decision being asked of the board', 'Assessor text'),
    ("§3 Strategic Alignment Analysis — analyse YAT's ICT Strategic Plan against the industry environment and organisational objectives", 'Assessor text'),
    ("§4 Current State of YAT's ICT — synthesised summary from the intranet materials (synthesis, not verbatim reproduction)", 'Assessor text'),
    ('§5 Gap Analysis — gaps, improvement opportunities, proposed changes', 'Assessor text'),
    ('§6 Options Considered and Evaluation — workload definition, options assessed (in-house renewal vs cloud migration to AWS), initial impact and difficulty assessment', 'Assessor text'),
    ('§7 Cost-Benefit Analysis — assumptions, summary tables for both options, avoided-downtime benefit, comparison summary, Year-1 cash-flow comparison, sensitivity analysis. The detailed line items live in Appendix 1 of the same document.', 'Assessor text'),
    ('§8 Risk and Impact Assessment — intangibles comparison and risk register for the recommended option', 'Assessor text'),
    ('§9 Recommendation — your chosen option with rationale', 'Assessor text'),
    ("§10 Action Plan — prioritised changes, implementation schedule, standards / targets / success metrics, implementation methods, alignment with YAT's change-management procedure, risk register", 'Assessor text'),
    ('§11 Next Steps and Decision Requested — what you are asking the board to decide today', 'Assessor text'),
    ('Sign-off block — completed during/after your presentation', 'Assessor text'),
    ('Appendix 1 — CBA detailed line items — every cost line you researched and computed', 'Assessor text'),
    ('Appendix 2 — Knowledge Evidence — short-answer responses linking your work to underlying cloud and evaluation principles', 'Assessor text'),
    ('Appendix 3 — Supporting Research — AWS Pricing Calculator export, citations', 'Assessor text'),
    ('Appendix 4 — Feedback Record — completed using the YAT Feedback Record template during/after the presentation', 'Assessor text'),
    ('You will need to research AWS pricing yourself using the AWS Pricing Calculator (see Appendix 1 of the Business Case template for the structure your AWS line items must follow). YAT-internal cost figures for the on-prem option are provided in the scenario materials on the intranet.', 'Assessor text'),
    ('Resources for Part A — available on the YAT intranet', 'Assessor text'),
    ("The YAT Business Case template (download from the intranet's Templates section)", 'Assessor text'),
    ("The YAT Feedback Record template (download from the intranet's Templates section)", 'Assessor text'),
    ('The YAT scenario materials — strategic plan, current ICT environment description, organisational policies, role brief, ICT manager consultation notes, reference materials', 'Assessor text'),
    ('The YAT Document Archive containing examples of previous business cases produced by MTS for YAT — review at least one example before starting your own as a reference for what a YAT business case looks like', 'Assessor text'),
    ('Resources for Part A — external', 'Assessor text'),
    ('AWS Pricing Calculator (calculator.aws) — to research the AWS cost line items for the cloud option', 'Assessor text'),
    ('AWS Academy Cloud Foundations [104469] + AWS Academy Cloud Architecting [172221] — authorised lab environments', 'Assessor text'),
    ('Part A is submitted to the LMS as the populated Business Case (.docx) with all four appendices completed.', 'Assessor text'),
    ('Part B — Presentation to the YAT board', 'Heading 2'),
    ("Using the YAT Board Presentation Deck template (download from the intranet's Templates section), prepare a deck that walks the YAT board through your Business Case. The deck is 8–10 slides plus speaker notes; each slide walks the board through one section of the Business Case.", 'Assessor text'),
    ('Review at least one example of a previous board presentation deck from the YAT Document Archive on the intranet — these accompany the previous business cases referenced for Part A — as a reference for what a YAT board presentation looks like in style, depth, and length.', 'Assessor text'),
    ('You will deliver the presentation to the YAT board, role-played by your assessor (as Sam Walker, YAT ICT Manager, and/or Pat Lin, MTS Senior Consultant) and a peer (as a board member).', 'Assessor text'),
    ('The presentation event', 'Assessor text'),
    ('Duration: 10–15 minutes of presenting + 5 minutes of board questions, feedback, and sign-off', 'Assessor text'),
    ('Format: in-person on campus or via video conference (confirm with your assessor)', 'Assessor text'),
    ('Required artefacts at the event:', 'Assessor text'),
    ('The completed Business Case submitted to the board at least 48 hours in advance', 'Assessor text'),
    ('The deck with populated speaker notes', 'Assessor text'),
    ('A printed or shared Feedback Record ready to be populated during/after the meeting (using the YAT Feedback Record template)', 'Assessor text'),
    ('A printed or shared sign-off sheet for the board to sign at the end (the Sign-off block from the Business Case)', 'Assessor text'),
    ('During the event you will:', 'Assessor text'),
    ('Walk the board through your Business Case using the deck', 'Assessor text'),
    ("Answer the board's questions during Q&A (the assessor's questions probe both the analysis and your underlying knowledge of cloud and evaluation principles)", 'Assessor text'),
    ('Capture board feedback in the Feedback Record', 'Assessor text'),
    ("Obtain the board's sign-off on the action plan in Business Case §10", 'Assessor text'),
    ('Part B is submitted to the LMS as the presentation deck (.pptx) with populated speaker notes. The completed Feedback Record and the signed Sign-off block are attached to the Business Case (as Appendix 4 and the Sign-off block respectively).', 'Assessor text'),
    ('Tips for success', 'Heading 2'),
    ('Start with the scenario. Read the YAT intranet end-to-end before you start writing — particularly the ICT Strategic Plan, the current ICT environment, the role brief, and the ICT manager consultation notes.', 'Assessor text'),
    ("Synthesise, don't copy. §4 (Current State) is the section most likely to drift into verbatim reproduction. Write it in your own words and focus on what's material to the migration.", 'Assessor text'),
    ('Show your AWS pricing working. Generate an AWS Pricing Calculator estimate, export it, and attach it as Appendix 3. Numbers without evidence are not credible.', 'Assessor text'),
    ('Sense-check your CBA totals. Year-1 outlay differences matter to senior management even when 5-year totals favour one option — make the cash-flow position visible.', 'Assessor text'),
    ('Be honest in §8 Intangibles. A one-sided intangibles table — only positives for your preferred option — is less credible than an honest trade-off discussion.', 'Assessor text'),
    ("The action plan is what the board approves. §10 is the most scrutinised section. Spend time on prioritisation, scheduling, and alignment with YAT's change-management procedure.", 'Assessor text'),
    ('Prepare for the Q&A. The assessor will probe both the analysis and your underlying knowledge of cloud and evaluation principles. The Q&A is when your speaker preparation pays off.', 'Assessor text'),
    ('Bring the sign-off sheet to the meeting. Sign-off is part of the assessment — you cannot achieve Satisfactory without it.', 'Assessor text'),
]

# Assessor-only body — Business Case Benchmark + Presentation Observation Process + UoC reverse-map.
# Each block is ('h1'|'h2'|'p', text) or ('tbl', [[row cells], ...]).
ASSESSOR_BODY = [
    ('h1', 'Business Case Benchmark'),
    ('p', '1. Marking §1 Executive Summary'),
    ('p', 'UoC evidenced: contributes to [ICTICT517 PC 1.4] Report on proposed changes to superior.'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'One page; covers all 6 prompt bullets'),
    ('p', 'The recommendation is stated (not deferred until later)'),
    ('p', '5-year cost numbers are present and reconcile with §7'),
    ('p', 'Top 2–3 risks named'),
    ('p', 'The decision asked of the board is explicit'),
    ('p', 'Not Yet Satisfactory:'),
    ('p', 'Missing or absent'),
    ('p', 'Reads as background-only with no recommendation'),
    ('p', 'Numbers do not reconcile with §7'),
    ('p', '2. Marking §2 Engagement Context'),
    ('p', 'UoC evidenced: [ICTICT517 AC 4] Individual superior in the organisation.'),
    ('p', 'Satisfactory looks like:'),
    ('p', "Names the student's MTS role + the YAT superior + the engagement scope"),
    ('p', 'States what the Business Case asks the board to decide'),
    ('p', 'NYS: missing or vague about who is engaged and who decides.'),
    ('p', '3. Marking §3 Strategic Alignment Analysis'),
    ('p', 'UoC evidenced: [ICTICT517 PC 1.1] Analyse and document current strategic plan against industry environment and organisational objectives · [ICTICT517 PE 1] interpret a strategic plan and objectives.'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'Cites the YAT ICT Strategic Plan objectives material to the engagement (verbatim or paraphrased with reference)'),
    ('p', 'Includes external industry context — RTO sector cloud adoption, availability expectations, data residency obligations — with cited sources'),
    ('p', "Identifies where YAT's plan aligns with industry direction and where it is more ambitious / more cautious"),
    ('p', 'Connects the analysis to the LMS migration'),
    ('p', 'NYS:'),
    ('p', 'No external industry research (analysis is internal-only)'),
    ('p', 'No citations'),
    ('p', 'Strategic Plan content reproduced without analysis'),
    ('p', "4. Marking §4 Current State of YAT's ICT"),
    ('p', 'UoC evidenced: [ICTICT517 PC 1.2] Determine and document current state · [ICTICT517 PE 2] evaluate the current state.'),
    ('p', 'Satisfactory looks like:'),
    ('p', "Synthesised summary — in the student's own words — covering at minimum the four prompted areas (network/security topology, LMS server status, supporting infrastructure, staff capability + dependencies)"),
    ('p', 'Material facts are accurate to the YAT scenario'),
    ('p', "Focus is on what's relevant to the migration; non-material detail is omitted"),
    ('p', 'NYS:'),
    ('p', 'Verbatim copy-paste of the intranet material'),
    ('p', 'Material inaccuracies (e.g. wrong OS version, wrong availability baseline)'),
    ('p', 'Missing one or more of the prompted areas'),
    ('p', "Note: synthesis is the assessable skill here, per the YAT case study tradition (517 AT2 Q2 asks for 100–200 words synthesis). Don't penalise a student for not reproducing the full intranet detail — penalise verbatim reproduction."),
    ('p', '5. Marking §5 Gap Analysis'),
    ('p', 'UoC evidenced: [ICTICT517 PC 1.3] Compare strategic plan vs current state · [ICTICT517 PE 3] identify gaps and opportunities.'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'At least 3 rows in the gap table, each covering all 6 columns (Strategic Plan objective, Current state, Desired future state, Gap, Improvement opportunity, Proposed change)'),
    ('p', 'Gaps are real and supported by §3 + §4'),
    ('p', 'Narrative summary ties the table together'),
    ('p', 'Expected gap items the student should identify:'),
    ('p', 'Availability gap (99.2% current vs 99.9% target)'),
    ('p', 'Recovery objectives gap (~24h RPO / 7–11h RTO current vs ≤1h RPO / ≤4h RTO target)'),
    ('p', 'LMS server end-of-life vs the "Deploy LMS infrastructure in cloud" objective'),
    ('p', 'Scalability gap for assessment-week peaks'),
    ('p', 'Optional: cloud-adoption gap aligned to the "Reduce dependency on in-house server infrastructure" objective'),
    ('p', 'NYS:'),
    ('p', 'Fewer than 3 gaps'),
    ('p', 'Gaps not connected to a stated Strategic Plan objective'),
    ('p', '"Proposed change" column missing or generic'),
    ('p', '6. Marking §6 Options Considered and Evaluation'),
    ('p', 'UoC evidenced: [ICTICT517 PC 2.1] Evaluate impact · [ICTICT517 PC 2.2] Evaluate difficulty · [ICTICT517 KE 3] Methods of evaluation of competing ICT systems and products · [ICTCLD401 PC 1.8] Define workload according to business requirements · [ICTCLD502 PC 1.2] Determine cloud infrastructure according to business needs.'),
    ('p', 'Satisfactory looks like:'),
    ('p', '§6.1 Workload definition references the LMS application spec'),
    ('p', '§6.2 names the two options (Option A in-house renewal, Option B cloud migration to AWS) and explains why these two'),
    ('p', '§6.3 names the evaluation method (CBA + intangibles + risk register) and justifies its choice'),
    ('p', '§6.4 initial impact + difficulty table populated for both options'),
    ('p', 'NYS:'),
    ('p', 'Only one option assessed'),
    ('p', 'No evaluation method named'),
    ('p', 'Workload definition missing'),
    ('p', '7. Marking §7 Cost-Benefit Analysis (and Appendix 1)'),
    ('p', 'UoC evidenced: [ICTICT517 PC 2.1] Evaluate impact · [ICTICT517 KE 3] · [ICTICT517 FS Numeracy] · [ICTCLD401 KE 4] · [ICTCLD502 KE 3] · [ICTCLD401 KE 1].'),
    ('p', 'This is the most numerical section. Use the worked example below as the benchmark.'),
    ('p', '7.1 Benchmark assumptions'),
    ('p', 'LMS user population: 800 students + 60 staff'),
    ('p', '5-year analysis period at 2026 price levels (AUD ex GST)'),
    ('p', 'ICT FTE fully-loaded cost: $115k/year'),
    ('p', 'Cost of downtime during teaching hours: $750/hour'),
    ('p', 'Teaching hours per year (indicative): 1,820 hours'),
    ('p', 'AWS pricing region: ap-southeast-2 (Sydney)'),
    ('p', 'AWS compute and database baseline assumed on 1-year Reserved Instance (~30% discount applied)'),
    ('p', '7.2 Option A — In-house renewal (5-year worked example)'),
    ('p', 'One-off capital (Year 1):'),
    ('tbl', [
        ['Item', 'Cost'],
        ['Replacement LMS server', '$25,000'],
        ['Backup tape library refresh + drive', '$8,000'],
        ['UPS upgrade', '$3,000'],
        ['One-off internal migration labour', '$15,000'],
        ['Year-1 capital', '$51,000'],
    ]),
    ('p', 'Recurring operational (per year):'),
    ('tbl', [
        ['Category', 'Item', 'Annual cost'],
        ['Software licensing', 'Windows Server Standard', '$1,500'],
        ['', 'Antivirus / EDR', '$300'],
        ['', 'Monitoring / management', '$2,000'],
        ['Power and facilities', 'Electricity + cooling', '$1,200'],
        ['', 'Server-room rent allocation', '$5,000'],
        ['', 'UPS battery (amortised)', '$500'],
        ['Backup', 'Tape media', '$1,500'],
        ['', 'Offsite tape storage', '$2,400'],
        ['', 'Backup software maintenance', '$1,500'],
        ['Staff time', 'LMS admin (0.20 FTE × $115k)', '$23,000'],
        ['', 'Incident response (0.05 FTE × $115k)', '$5,750'],
        ['External support', 'MTS application support', '$30,000'],
        ['Recurring per year', '', '$74,650'],
    ]),
    ('p', '5-year Option A total:'),
    ('tbl', [
        ['Component', '5-year'],
        ['One-off capital', '$51,000'],
        ['Recurring (5 × $74,650)', '$373,250'],
        ['Option A — 5 years', '$424,250'],
    ]),
    ('p', '7.3 Option B — Cloud migration to AWS (5-year worked example)'),
    ('p', 'One-off project costs (Year 1):'),
    ('tbl', [
        ['Item', 'Cost'],
        ['MTS migration labour (architecture → cutover → handover)', '$80,000'],
        ['YAT ICT staff time (~0.30 FTE × 6 months × $115k)', '$17,250'],
        ['Parallel running of on-prem and cloud during cutover (~2 months)', '$12,500'],
        ['Decommissioning on-prem LMS + secure data destruction', '$2,500'],
        ['Year-1 project', '$112,250'],
    ]),
    ('p', 'AWS direct recurring (per year, post-migration):'),
    ('p', 'Sized for the workload in the LMS application specification. RI discount of ~30% on compute and managed-DB baseline.'),
    ('tbl', [
        ['Category', 'Item', 'Annual cost'],
        ['Compute', '2× m6i.xlarge EC2 (HA baseline, 1-yr RI)', '$6,600'],
        ['', 'Auto-scaling peaks (~10 instance-weeks/year)', '$1,500'],
        ['Database', 'RDS db.m6i.large Multi-AZ MySQL (1-yr RI)', '$2,000'],
        ['', 'RDS storage + automated backups', '$400'],
        ['Network', 'Application Load Balancer', '$400'],
        ['', 'NAT gateway', '$400'],
        ['', 'Data transfer out (~50 GB/mo)', '$50'],
        ['Storage', 'EBS gp3 (~500 GB)', '$600'],
        ['', 'S3 attachments + snapshots (~200 GB)', '$50'],
        ['', 'AWS Backup cross-Region copy', '$300'],
        ['Monitoring', 'CloudWatch metrics + logs + alarms', '$500'],
        ['Support', 'AWS Business Support', '$1,800'],
        ['AWS direct per year', '', '$14,600'],
    ]),
    ('p', 'Staff time + ongoing external support (per year, Years 2–5):'),
    ('tbl', [
        ['Category', 'Item', 'Annual cost'],
        ['Staff time', 'LMS admin (0.10 FTE × $115k)', '$11,500'],
        ['', 'Incident response (0.03 FTE × $115k)', '$3,450'],
        ['External support', 'MTS ongoing support (reduced — managed services)', '$20,000'],
        ['Staff + external per year', '', '$34,950'],
    ]),
    ('p', '5-year Option B total:'),
    ('tbl', [
        ['Component', '5-year'],
        ['Year-1 project', '$112,250'],
        ['AWS direct (5 × $14,600)', '$73,000'],
        ['Staff + external (4 × $34,950 — Year 1 covered by MTS migration project)', '$139,800'],
        ['Option B — 5 years', '$325,050'],
    ]),
    ('p', '7.4 Avoided-downtime benefit'),
    ('p', 'Current availability: 99.2%. Target: 99.9%. Improvement: 0.7 percentage points = ~61 hours/year avoided unavailability.'),
    ('p', 'Assume ~30% of avoided hours fall during teaching hours:'),
    ('p', '61 hr × 0.30 × $750/hr = ~$13,725/year avoided downtime cost'),
    ('p', 'Over 5 years: ~$68,625'),
    ('p', '7.5 Comparison summary'),
    ('tbl', [
        ['', 'Option A — In-house', 'Option B — Cloud (AWS)'],
        ['One-off (Year 1)', '$51,000', '$112,250'],
        ['Recurring 5-year', '$373,250', '$212,800'],
        ['5-year total', '$424,250', '$325,050'],
        ['5-year saving (B − A)', '—', '$99,200'],
        ['Avoided-downtime benefit (B vs A, 5-year)', '—', '~$68,625'],
        ['Net 5-year delta in favour of cloud', '—', '~$167,825'],
    ]),
    ('p', '7.6 Year-1 cash-flow profile'),
    ('p', "Important for sign-off conversation — both options have similar Year-1 cash impact; cloud's savings appear in Years 2–5:"),
    ('tbl', [
        ['', 'Year 1 outlay', 'Year 2 outlay (recurring)'],
        ['Option A', '$51,000 + $74,650 = $125,650', '$74,650'],
        ['Option B', '$112,250 + $14,600 = $126,850', '$14,600 + $34,950 = $49,550'],
    ]),
    ('p', '7.7 Acceptable ranges for student totals'),
    ('p', 'A student arriving at a 5-year cloud total in the range $280k–$370k has done credible AWS sizing and pricing work.'),
    ('p', 'A student arriving at a 5-year on-prem total in the range $400k–$450k has used the inputs correctly and applied reasonable inflation/staff-time treatment.'),
    ('p', 'Students sharply outside these ranges either have a different sizing rationale (in which case marking checks the rationale) or have miscalculated (in which case flag the specific line items).'),
    ('p', '7.8 Marking checks for §7 + Appendix 1'),
    ('p', 'Award credit for:'),
    ('p', 'All cost categories present and reasonably populated in both options'),
    ('p', 'AWS Pricing Calculator (or equivalent) evidence attached as Appendix 3'),
    ('p', '5-year projection per option'),
    ('p', 'Stated sizing assumption for AWS option'),
    ('p', 'Sensitivity analysis (§7.6) around at least one key assumption'),
    ('p', 'Year-1 cash-flow point (§7.5.1) made explicitly'),
    ('p', 'Quantified avoided-downtime calculation (§7.4)'),
    ('p', 'Detailed line items in Appendix 1 roll up correctly to §7 summary tables'),
    ('p', 'Common deductions:'),
    ('p', 'Cloud option missing entire categories (e.g. no Multi-AZ DB pricing, no support tier)'),
    ('p', 'Pricing values lifted without source — no evidence of AWS Pricing Calculator use'),
    ('p', 'Missing or absent on-prem option (just doing AWS pricing and skipping the comparison)'),
    ('p', 'Appendix 1 totals do not reconcile with §7 summary tables'),
    ('p', '8. Marking §8 Risk and Impact Assessment'),
    ('p', 'UoC evidenced: [ICTICT517 PC 2.1] non-financial impact · [ICTICT517 PC 2.2] difficulty · [ICTICT517 PE 5] evaluate difficulty.'),
    ('p', 'Satisfactory looks like:'),
    ('p', '§8.1 intangibles table populated for both options across all 9 factors'),
    ('p', 'Treatment is honest — both options have pros and cons listed; no one-sided cheerleading for the preferred option'),
    ('p', '§8.2 risk register has at least 3 risks for the recommended option, each with likelihood / impact / mitigation'),
    ('p', 'Expected intangibles coverage (student should address most):'),
    ('p', 'Availability — Option B can achieve 99.9% target; Option A would require significant additional capital for HA (probably out of scope of the renewal cost).'),
    ('p', "Recovery objectives — Option B meets RTO ≤ 4h / RPO ≤ 1h targets; Option A's current 7–11h RTO and 24h RPO do not."),
    ('p', 'Growth capacity — Option B autoscales; Option A would require server resizing with downtime.'),
    ('p', 'Peak handling — Option B autoscales for assessment-week spikes; Option A must be statically sized.'),
    ('p', 'Staff capability — Option B exposes YAT ICT to cloud (positive over time; gap during transition).'),
    ('p', 'Vendor lock-in — Option A no lock-in; Option B moderate AWS lock-in via managed services.'),
    ('p', 'Sustainability — Option A 24×7 server-room cooling; Option B leverages AWS sustainability profile.'),
    ('p', "Cyber security — Option B richer monitoring + threat intelligence; Option A relies on YAT's existing controls."),
    ('p', 'Disaster recovery — Option B Multi-AZ + cross-Region copy; Option A tape-restore from offsite, ~24h.'),
    ('p', 'NYS:'),
    ('p', 'One-sided intangibles (only positives for cloud)'),
    ('p', 'Fewer than 4 intangibles addressed substantively'),
    ('p', 'Risk register absent or has fewer than 3 risks'),
    ('p', '9. Marking §9 Recommendation'),
    ('p', 'UoC evidenced: [ICTICT517 PC 2.3] Prioritise proposed changes · [ICTICT517 PE 4] determine and prioritise · contributes to [ICTICT517 PC 3.1] Develop action plan.'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'Recommended option clearly stated (Option A or Option B)'),
    ('p', '3–5 sentence rationale connecting back to the CBA findings, intangibles, and risk register'),
    ('p', 'Recommendation flows logically into §10 Action Plan'),
    ('p', 'Expected recommendation: most credible analyses arrive at Option B (cloud migration to AWS). A student recommending Option A is acceptable as long as the recommendation is logically connected to their analysis — e.g. they have identified a particular blocker to cloud that makes in-house renewal the better choice. Marking does not penalise the choice of option; it marks whether the choice is justified by the work.'),
    ('p', 'NYS:'),
    ('p', 'Recommendation not connected to the analysis (e.g. dollar comparison shows clear winner but student recommends the other option without justification)'),
    ('p', 'No rationale given'),
    ('p', 'Recommendation hedged ("either could work") without committing'),
    ('p', '10. Marking §10 Action Plan'),
    ('p', 'UoC evidenced: [ICTICT517 PC 3.1] Develop action plan including prioritised schedule and consistency with org policy and procedures · [ICTICT517 PC 3.2] Detail standards, targets and implementation methods · [ICTICT517 PE 6] develop action plan · [ICTICT517 KE 1] Key sections of an action plan.'),
    ('p', 'This is one of the critical sections — NYS here forces overall NYS.'),
    ('p', 'Satisfactory looks like:'),
    ('p', '§10.1 prioritised changes — at least 5 items with priority rationale'),
    ('p', '§10.2 implementation schedule — phased; dependencies named; no all-at-once cutover proposed'),
    ('p', '§10.3 standards / targets table populated, with measurable targets (99.9% availability, RPO ≤ 1h, RTO ≤ 4h, data residency Australia, cost envelopes)'),
    ('p', '§10.4 implementation methods — names specific methods, tools, standards (Well-Architected, AWS managed services, ITIL change management, etc.)'),
    ('p', '§10.5 alignment with YAT change-management procedure — explicit'),
    ('p', '§10.6 risk register cross-references §8.2'),
    ('p', 'Common phasing the student should produce (any of):'),
    ('p', 'Phase 1: design + foundation build (AWS account setup, IAM, VPC, monitoring baseline)'),
    ('p', 'Phase 2: application migration (LMS server build, app deployment, parallel running)'),
    ('p', 'Phase 3: HA hardening (Multi-AZ DB, autoscaling, monitoring + alerting)'),
    ('p', 'Phase 4: cutover (high-risk change requiring full change-management sign-off)'),
    ('p', 'Phase 5: stabilisation + handover to YAT ICT operations'),
    ('p', 'NYS:'),
    ('p', 'Single-phase all-at-once cutover with no risk treatment'),
    ('p', 'Missing standards / targets / success metrics'),
    ('p', 'No alignment with change-management procedure'),
    ('p', 'Phases have no dependencies or owners'),
    ('p', '11. Marking §11 Next Steps and Decision Requested'),
    ('p', 'UoC evidenced: [ICTICT517 PC 3.3] Provide action plan to superior for feedback and approval.'),
    ('p', 'UoC evidenced (foundation skills & additional criteria — closest-fit to the noted marking criteria): [ICTICT517 FS Reading] (A12) · [ICTICT517 FS Writing] (A15) · [ICTICT517 FS Get the work done] (A4/A8/A12) · [ICTICT517 FS Oral Communication] (B7/B8) · [ICTICT517 FS Navigate the world of work] (A8) · [ICTICT517 FS Interact with others] (B5/B6) · [ICTCLD502 FS Oral communication] (B9) · [ICTCLD401 PC 4.2] (B5) · [ICTICT517 PC 2.4] (A13/B2).'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'States precisely what the board is being asked to approve today'),
    ('p', 'Names what is deferred to later sign-off gates'),
    ('p', 'NYS: vague or missing.'),
    ('p', '12. Marking Appendix 2 — Knowledge Evidence'),
    ('p', 'This is a critical section — NYS here forces overall NYS.'),
    ('p', 'Below are model answers for each question. Student answers may differ in detail but must demonstrate the underlying knowledge applied to their own Business Case.'),
    ('p', '12.1 Selection-style questions'),
    ('p', 'Q1 — IaaS / PaaS / SaaS classification of solution components [ICTCLD401 KE 3]'),
    ('p', 'Model coverage: student identifies which AWS services are IaaS (e.g. EC2, EBS — they manage the instances), PaaS (e.g. RDS Multi-AZ, ALB — managed services where AWS runs the platform), SaaS (none in their solution typically — they may note Office 365 as a YAT-existing SaaS but not part of their migration). Rationale: PaaS chosen for DB because YAT ICT lacks DB-admin cloud experience; IaaS chosen for the LMS application tier because the existing Windows Server / DOODLE stack is preserved.'),
    ('p', 'Q2 — Cloud deployment model [ICTCLD401 KE 11]'),
    ('p', "Model coverage: Public cloud (AWS). Rationale: YAT's strategic plan explicitly targets cloud migration to reduce in-house infrastructure; private cloud doesn't address the cost reduction objective; hybrid is what they have today (Office 365 SaaS + on-prem LMS) and the migration shifts the on-prem LMS into the public cloud, leaving Office 365 SaaS unchanged."),
    ('p', 'Q3 — Industry technology standards [ICTCLD401 KE 1 · ICTCLD502 KE 1]'),
    ('p', 'Model coverage: At least three named, with material role explained. Acceptable: NIST SP 800-145 (definitions of IaaS/PaaS/SaaS) — informs §6 service-type classification; AWS Well-Architected Framework — informs §10 implementation methods; ISO/IEC 27017 — informs §8 security treatment; ACSC Essential Eight — informs §10 security controls; ITIL 4 — informs alignment with YAT change management procedure.'),
    ('p', 'Q4 — Industry-standard hardware/software products [ICTCLD401 KE 2 · ICTCLD502 KE 2]'),
    ('p', 'Model coverage: Student identifies the specific AWS services in their solution + their general features + role. Example: Amazon EC2 (general-purpose virtual machine; underpins the LMS application tier); Amazon RDS for MySQL Multi-AZ (managed relational database with automatic failover; underpins the LMS data tier); Application Load Balancer (layer-7 load balancer with health checks; distributes traffic across HA EC2 instances); etc.'),
    ('p', 'Q5 — Cloud cost models and scalability [ICTCLD401 KE 4 · ICTCLD502 KE 3]'),
    ('p', 'Model coverage: Student explains which cost model applies to each AWS service in their CBA. Example: 1-year Reserved Instance for EC2 baseline + RDS baseline (~30% discount — appropriate because demand is predictable); on-demand or autoscaling-spot for EC2 peak capacity (cost scales with assessment-week demand); S3 pay-as-you-go (cost scales with storage growth); Business Support 10% of usage (cost scales with overall AWS spend). Scaling consequence: total AWS cost grows roughly linearly with student-population growth (the 15%/year target).'),
    ('p', '12.2 Demonstration-style questions'),
    ('p', 'Q6 — Key sections of an action plan [ICTICT517 KE 1]'),
    ('p', 'Model coverage: Student identifies the key sections of an ICT implementation action plan (prioritised list of changes, implementation schedule with dependencies, standards/targets/success metrics, implementation methods, alignment with change-management procedure, risk register, owners) and cites where each appears in their §10. Demonstrates §10 has all key sections.'),
    ('p', 'Q7 — Methods of evaluation and planning approaches [ICTICT517 KE 2]'),
    ('p', 'Model coverage: Student names the methods applied (CBA, sensitivity analysis, intangibles assessment, risk register) and cites the §7 / §8 sections where each is applied. May also cite their use of structured gap-analysis methodology in §5.'),
    ('p', 'Q8 — Evaluation of competing ICT systems and products [ICTICT517 KE 3]'),
    ('p', 'Model coverage: Student names the evaluation method — typically the CBA combined with intangibles — and cites §6.3 (where the method is named) + §7 (where the CBA evaluates) + §8.1 (where intangibles compare). May also mention weighted-criteria evaluation if the student used one.'),
    ('p', 'Q9 — Current and emerging trends [ICTICT517 KE 4]'),
    ('p', 'Model coverage: Student identifies cloud migration itself as the current/emerging trend that informed the recommendation; possibly also Multi-AZ HA design patterns, managed services as a reducing-burden trend, autoscaling-for-cost-optimisation, etc. Cites where in their reasoning each trend was material.'),
    ('p', 'Marking checks for Appendix 2'),
    ('p', 'Award credit for:'),
    ('p', 'All 9 questions answered'),
    ('p', "Each answer references specific sections of the student's own Business Case"),
    ('p', "The cited content is actually present in the student's Business Case (not invented)"),
    ('p', 'Common deductions:'),
    ('p', 'Generic textbook answers not tied to the student\'s own work ("IaaS is infrastructure-as-a-service..." without saying which AWS service in their solution is IaaS)'),
    ('p', "Citations to Business Case sections that don't contain what the answer claims they do"),
    ('p', '13. Marking Appendix 3 — Supporting Research'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'AWS Pricing Calculator export attached (screenshot or shared link)'),
    ('p', 'Industry-research sources from §3 cited with URLs and access dates'),
    ('p', 'Any other source documents used for cost figures cited'),
    ('p', "NYS: AWS Pricing Calculator evidence absent — implies the student didn't research their AWS figures and is fabricating."),
    ('p', '14. Marking Appendix 4 — Feedback Record'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'Completed using the Feedback Record template'),
    ('p', 'Captures actual board feedback received during the presentation event'),
    ('p', "Includes student's response or follow-up action where applicable"),
    ('p', 'NYS: absent, blank, or evidently not completed at the event.'),
    ('p', '15. Marking the Sign-off block'),
    ('p', 'This is a critical artefact — NYS here forces overall NYS.'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'Sign-off block populated with names, dates, signatures'),
    ('p', 'Approval line completed by the role-played YAT ICT Manager (the assessor in role)'),
    ('p', 'NYS: sign-off block left blank — implies the student did not seek or obtain approval per [ICTICT517 PC 3.3].'),
    ('h1', 'Presentation Observation Process'),
    ('h2', 'Main Presentation '),
    ('p', 'To be assessed against the criteria established in the Assessment – Marking Guide Part B'),
    ('h2', 'Assessor-led questions (Suggested)'),
    ('p', 'Each suggested question is tagged with the UoC requirement(s) it probes. The assessor selects 3–5 to ask the student during Q&A, varying across the listed dimensions.'),
    ('tbl', [
        ['Question', 'UoC requirement(s) probed'],
        ['"Your cloud-side compute figure is built on 1-year Reserved Instances. What happens to the comparison if YAT prefers on-demand pricing because they want migration flexibility?"', '[ICTICT517 PC 2.1] Evaluate impact of proposed changes · [ICTICT517 FS Numeracy]'],
        ['"You\'ve assumed YAT ICT staff time drops by N FTE post-migration. What\'s that based on, and what\'s the risk if that drop doesn\'t materialise?"', '[ICTICT517 PC 2.1] · [ICTICT517 PC 2.2] Evaluate the difficulty of implementing proposed changes'],
        ['"The Strategic Plan talks about expanding to other locations. How does your recommendation hold up if YAT opens a second campus in Year 3?"', '[ICTICT517 PC 1.1] Analyse and document current strategic plan · [ICTICT517 PC 2.1]'],
        ['"You\'ve recommended the cloud option. What\'s the single biggest reason senior management might say no?"', '[ICTICT517 PC 2.2] · [ICTICT517 PC 2.3] Prioritise proposed changes'],
        ['"Of the prioritised changes in your action plan, which one carries the most implementation risk and how are you proposing to mitigate it?"', '[ICTICT517 PC 2.3] · [ICTICT517 PC 3.1] Develop action plan · [ICTICT517 PC 3.2] Detail standards, targets and implementation methods'],
        ['"What in your AWS sizing or configuration choices is driven by the data-residency requirement?"', '[ICTCLD502 PC 1.1] Determine reliability, recoverability and service levels required for application · [ICTCLD502 PC 1.2] Determine cloud infrastructure according to business needs'],
        ['"For each layer of your proposed YAT solution, identify whether it is IaaS / PaaS / SaaS and explain why you chose that service type."', '[ICTCLD401 KE 3] principles and functions of cloud computing solutions and technologies, including IaaS / PaaS / SaaS — contextual reflective question'],
        ['"Which industry standards informed your migration recommendation, and how?"', '[ICTCLD401 KE 1] · [ICTCLD502 KE 1] industry technology standards · [ICTCLD502 AC 5] industry standards'],
    ]),
    ('h1', 'UoC coverage verification (reverse map)'),
    ('p', 'This table closes the loop on bidirectional traceability: every UoC requirement AT1 claims to evidence is named below with the marking criterion (or criteria) that evidence it. No UoC requirement claimed by AT1 is left without a criterion.'),
    ('p', 'ICTICT517 — Match ICT needs with the strategic direction of the organisation'),
    ('tbl', [
        ['UoC item', 'Evidenced by criterion(ia)'],
        ['PC 1.1 — Analyse and document current strategic plan', 'A1'],
        ['PC 1.2 — Determine and document current state of ICT', 'A2'],
        ['PC 1.3 — Compare strategic plan vs current state; determine gaps', 'A3'],
        ['PC 1.4 — Report on proposed changes to superior', 'A10 (written) · B1 (verbal)'],
        ['PC 2.1 — Evaluate impact of proposed changes', 'A4 (initial) · A5 (financial) · A6 (non-financial)'],
        ['PC 2.2 — Evaluate difficulty of implementing changes', 'A4 (initial) · A6 (deep)'],
        ['PC 2.3 — Prioritise proposed changes', 'A7'],
        ['PC 2.4 — Document evaluation process and provide to superior for feedback', 'A13 (feedback captured) · B2 (provided + walked through in meeting)'],
        ['PC 3.1 — Develop action plan including prioritised schedule', 'A8'],
        ['PC 3.2 — Detail standards, targets and implementation methods', 'A8'],
        ['PC 3.3 — Provide action plan to superior for feedback and approval', 'A9 (written ask) · A14 (granted approval) · B3 (verbal seeking)'],
        ['PE 1 — Interpret a strategic plan and objectives', 'A1'],
        ['PE 2 — Evaluate current state of ICT', 'A2'],
        ['PE 3 — Identify gaps and opportunities; evaluate organisational impact', 'A3'],
        ['PE 4 — Determine and prioritise proposed changes', 'A7'],
        ['PE 5 — Evaluate difficulty of implementing changes', 'A6'],
        ['PE 6 — Develop action plan for implementation', 'A8'],
        ['KE 1 — Key sections of an action plan', 'A8 · A11'],
        ['KE 2 — Methods of evaluation and planning approaches', 'A5 · A11'],
        ['KE 3 — Methods of evaluation of competing ICT systems and products', 'A4 · A5 · A11'],
        ['KE 4 — Current and emerging trends', 'A11'],
        ['FS Reading', 'A12'],
        ['FS Writing', 'A15'],
        ['FS Oral Communication', 'B7 · B8'],
        ['FS Numeracy', 'A5'],
        ['AC 1 — Coordination site', 'C1 (pre-condition)'],
        ['AC 2 — Strategic plan information', 'C1 (pre-condition)'],
        ['AC 3 — Organisational policies and procedures for ICT changes', 'C1 (pre-condition — YAT change-management procedure on intranet)'],
        ['AC 4 — Individual superior', 'C2 (pre-condition)'],
    ]),
    ('p', 'ICTCLD401 — Configure cloud services (AT1-relevant items only)'),
    ('tbl', [
        ['UoC item', 'Evidenced by criterion(ia)'],
        ['PC 1.8 — Define workload according to business requirements', 'A4'],
        ['PC 4.1 — Document and communicate work to required personnel', 'A14 · B4'],
        ['PC 4.2 — Seek and respond to feedback as required', 'B5'],
        ['KE 1 — Industry technology standards', 'A5 · A11'],
        ['KE 2 — Industry standard hardware and software products', 'A11'],
        ['KE 3 — IaaS / PaaS / SaaS', 'A11'],
        ['KE 4 — Cost models and cloud economic theories', 'A5 · A11'],
        ['KE 11 — Cloud models (on-prem / private / hybrid / public)', 'A11'],
        ['AC 4 — Data to gather information from to determine output and user requirements', 'C1 (pre-condition — YAT scenario materials on intranet)'],
    ]),
    ('p', '(Note: 401 PCs not listed above (PC 1.1–1.7, 2.1–2.6, 3.1–3.2, 4.3) are evidenced by AT2 Cloud Foundation Build. 401 KEs not listed (KE 5–10) are evidenced contextually in AT2/AT3.)'),
    ('p', 'ICTCLD502 — Design and implement highly-available cloud infrastructure (AT1-relevant items only)'),
    ('tbl', [
        ['UoC item', 'Evidenced by criterion(ia)'],
        ['PC 1.2 — Determine cloud infrastructure according to business needs', 'A4'],
        ['PC 5.2 — Confirm, seek and respond to feedback with required personnel', 'B6'],
        ['KE 1 — Industry technology standards', 'A11'],
        ['KE 2 — Industry standard hardware and software products', 'A11'],
        ['KE 3 — Cloud cost models and scalability', 'A5 · A11'],
        ['FS Oral communication', 'B9'],
        ['AC 3 — Information and data sources required to design and implement cloud infrastructure', 'C1 (pre-condition)'],
        ['AC 5 — Specific requirements + industry standards + organisational procedures + legislative requirements + business and functionality requirements', 'C1 (pre-condition — supplied via YAT intranet including industry standards reference, legislative requirements reference, LMS application spec, LMS cloud migration requirements)'],
        ['AC 8 — Data to gather information from to determine output and user requirements', 'C1 (pre-condition)'],
    ]),
    ('p', '(Note: 502 PCs not listed above (PC 1.1, 1.3, 2.1–2.5, 3.1–3.5, 4.1–4.6, 5.1, 5.3) are evidenced by AT3 HA Design + Project Closure. 502 KEs not listed (KE 4–9) are evidenced in AT3.)'),
]


# ---------- build helpers ----------

STYLE = {"h1": "Heading 1", "h2": "Heading 2", "p": "Normal"}


def add_marking_row(table, lines):
    """A marking-guide criterion row: criterion (+ UoC line) in col 0, the check box in col 1."""
    row = table.add_row()
    set_cell_content(row.cells[0], lines)
    set_cell_content(row.cells[1], CHECK)


def set_body_paragraph_text(doc, old, new):
    for p in doc.paragraphs:
        if p.text.strip() == old:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            p.add_run(new)
            return
    raise KeyError(old)


def delete_body_paragraph(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            p._element.getparent().remove(p._element)
            return


def render_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    for r, cells in enumerate(rows):
        for c, val in enumerate(cells):
            set_cell_content(t.rows[r].cells[c], val)
    doc.add_paragraph()


def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details ----
    t_details = doc.tables[0]
    set_cell_content(t_details.rows[1].cells[1], DETAILS["qualification"])
    set_cell_content(t_details.rows[2].cells[1], DETAILS["units"])
    set_cell_content(t_details.rows[3].cells[1], DETAILS["task_title"])
    set_cell_content(t_details.rows[4].cells[1], DETAILS["task_number"])

    # ---- Table 1: Teacher/Assessor instructions ----
    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Time allowed"), "")
    set_cell_content(find_instruction_row(t_instr, "Location"), "")
    set_cell_content(find_instruction_row(t_instr, "Resources required"), RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA)
    # add a Conditions row at the end (matching the table style), as CL1/CL2 do
    cond_row = t_instr.add_row()
    set_cell_content(cond_row.cells[0], "Assessment Conditions & Setup Requirements")
    for r in cond_row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell_content(cond_row.cells[1], CONDITIONS)

    # ---- Marking Guide heading (the authored instrument drops the 'Project ' prefix) ----
    set_body_paragraph_text(doc, "Project Assessment – Marking Guide", "Assessment – Marking Guide")

    # ---- Table 2: Marking Guide ----
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)  # keep 'Assessment criteria' + 'Criteria | Satisfactory?' header rows
    add_section_row(t_mark, "Part A – Business Case")
    for lines in MARKING_A:
        add_marking_row(t_mark, lines)
    add_section_row(t_mark, "Part B - Presentation")
    for lines in MARKING_B:
        add_marking_row(t_mark, lines)

    # ---- drop the template's trailing marking-guide boilerplate ----
    delete_body_paragraph(doc, "Add or delete rows as required")
    delete_body_paragraph(doc, "If questioning or observation is incorporated into this assessment task, "
                                "you can incorporate a Practical Observation Checklist.")

    # ---- Instructions to Student (shared prose) ----
    doc.add_paragraph("Instructions to Student", style="Heading 1")
    for text, style in PROSE:
        doc.add_paragraph(text, style=style)

    # ---- Assessor-only body (Benchmark + Observation + reverse-map) ----
    for kind, payload in ASSESSOR_BODY:
        if kind == "tbl":
            render_table(doc, payload)
        else:
            doc.add_paragraph(payload, style=STYLE[kind])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/assessments/AT1/AT1-BusinessCase-Assessor.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
