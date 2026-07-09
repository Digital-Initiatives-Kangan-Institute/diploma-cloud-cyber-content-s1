#!/usr/bin/env python3
"""Build the S1-CL1 AT3 ASSESSOR instrument (.docx) by populating the Kangan template.

This is an institutional compliance document, NOT a YAT-branded artefact: it loads the
official Kangan 'Project Assessment - Assessor' template and fills it in, preserving the
Kangan structure and styles (Details, Teacher/Assessor instructions, Marking Guide,
Instructions to Student, Benchmark, UoC reverse-map). It retro-fits a generator to the
approved, hand-authored CL1 AT3 instrument, reproducing it exactly at the content level.
Mirrors the CL1 AT1 generator's mechanics.

EVERGREEN: the generator emits NO literal scenario-site URL. Where the committed instrument
carried the placeholder intranet URL (a parenthetical after "the YAT intranet", and an
"... at <url>" tail), the generator drops only the URL + its immediate wrapper and keeps
every other word verbatim (proven by validate_instrument_reproduction with url-only --subs).

AT3 = High Availability: YAT LMS Migration (ICTICT517 + ICTCLD401 + ICTCLD502) — two written
parts submitted as one bundle:
  Part A  HA Design            (student-authored HA-equivalent architecture; A1-A7)
  Part B  HA Deployment Report (implementation + failure/resize simulation + sign-off; B1-B16)

The assessor instrument carries the task instructions, the Marking Guide (criteria + inline UoC
traceability), the HA Design + HA Deployment Report Benchmarks (per-section guidance + KE model
answers), and the UoC coverage reverse-map. The student instrument (build_s1_cl1_at3_student)
carries only the student-facing content.

Usage:  python scripts/s1_cl1/build_s1_cl1_at3_assessor.py [output.docx]
Default: S1-CL1-Cloud-Design-Build/assessments/AT3/AT3-High-Availability-Assessor.docx
"""
import sys
from pathlib import Path

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Assessor.docx")

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_tables import add_section_row, clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402


# ---------- content ----------

CHECK = "☐ Yes  ☐ No"  # marking-guide Satisfactory? cell (matches the Kangan template)

DETAILS = {
    "qualification": "ICT50220 Diploma of Information Technology",
    "units": [
        "ICTICT517 Match ICT needs with the strategic direction of the organisation",
        "ICTCLD401 Configure cloud services",
        "ICTCLD502 Design and implement highly-available cloud infrastructure",
    ],
    "task_title": "AT3 — High Availability: YAT LMS Migration",
    "task_number": "3 of 3",
}

OVERVIEW = [
    'Students are assessed on the design and implementation of HA improvements to the cloud infrastructure deployed in AT2. AT3 is the third and final assessment task in the S1-CL1 Cloud Design and Build cluster — it closes out the YAT LMS Cloud Migration engagement.',
    'The assessment has two parts, both written deliverables:',
    'Part A (written design): the HA Design document — student-produced, specifying the HA-equivalent architecture that supersedes the AT2 baseline. Includes the architecture review of AT2, HA requirements, designed HA architecture, implementation sequencing, and simulation plan.',
    'Part B (written deployment report): the HA Deployment Report — documents the implementation of the HA design (during a simulated Saturday late-night maintenance window), the failure and resize simulation outcomes, post-simulation adjustments, the feedback received, and the final sign-off.',
    'Both parts are submitted together as one bundle. There is no presentation or observation event — all evidence is captured in the two documents and their appendices.',
    'This is an open-book practical assessment. Students may use the YAT intranet, AWS Academy lab environments, AWS documentation, course reference materials, and external research (which must be cited) throughout.',
    "Maintenance-window framing for the implementation work: Part B's HA implementation is performed within a simulated Saturday late-night maintenance window of approximately 3.5 hours. Students plan their changes to minimise service disruption but acceptably may have brief outages within the window as long as the LMS infrastructure is back online (either with HA hardening complete, or fully rolled back to the pre-window state) by the end. This is a deliberate framing — at this qualification level, a new graduate working on truly mission-critical infrastructure would be supervised; the simulated maintenance window is a fair proxy for the real-world risk envelope they would experience.",
    'AT2 baseline starting state: to ensure students begin AT3 from a consistent baseline regardless of what they personally built in AT2, the assessor distributes the AT2 baseline CloudFormation template (see scenario/assessor-resources/at2-baseline-cloudformation.{md,yaml}) at the start of the assessment day. Students deploy this in their AWS Academy lab to instantiate the AT2 baseline as their AT3 starting state. Recommended schedule: morning prep slot for deployment + smoke-test verification; afternoon for the maintenance window + report writing.',
    'Reasonable adjustment for this assessment may include extending the maintenance window beyond 3.5 hours, providing one-on-one verbal explanation of the AT2 baseline design, allowing alternative screenshot-evidence formats where assistive technology requires it, or scheduling the work across multiple sittings rather than a single window.',
    "Teacher/assessor support level: the teacher/assessor may clarify task requirements, scenario context, or the AT2 baseline design but must not guide students to specific HA design choices, simulation outcomes, or knowledge-evidence answers. The HA design is the student's own work.",
    'Submission: the completed HA Design (Part A, .docx) AND HA Deployment Report (Part B, .docx) bundle submitted via the LMS.',
    'The assessment will not proceed if for any reason it is not safe to do so. You must advise the student of the reason for suspending the assessment, and what safety action should be taken. Advise the student of revised arrangements for the assessment when it is safe to do so.',
    'There is a zero tolerance for plagiarism, cheating and collusion. Students will be expected to make a declaration that all work is their own prior to submission. Refer to the Training and Assessment Policy for further information.',
]

# EVERGREEN: the committed source read "... available on the YAT intranet at <url> — students ...";
# the URL + its " at " wrapper are dropped, every other word verbatim.
TASKS = [
    "Following the AT2 foundation-build phase, MTS returns to YAT for the HA hardening phase of the engagement. The cloud infrastructure built in AT2 is single-AZ and non-HA; YAT's strategic target of 99.9% availability has not yet been achieved. AT3's HA work closes that gap.",
    'The task has two parts that combine into a single submission:',
    "Part A — HA Design. Student produces an HA Design document using the YAT-provided HA Design template (mirrors the AT2 supplied design's section structure but is student-authored). Covers HA requirements, review of the AT2 baseline (architecture review, SPOFs identified, recovery objectives estimated, vertical-scaling impact), HA-equivalent architecture design (with recovery objectives, SPOFs removed, scaling considerations), implementation sequencing, and a simulation plan.",
    'Part B — HA Deployment Report. Student implements the HA design in the AWS Academy lab during a simulated maintenance window, runs the failure and resize simulations specified in the HA Design §6, captures the outcomes and any adjustments needed, secures feedback and final sign-off from the role-played YAT ICT Manager, and produces the HA Deployment Report using the YAT-provided HA Deployment Report template. The report mirrors the AT2 Deployment Report shape (same sections in the same places) with HA-specific content within each section.',
    'MTS scope at AT3: in-place HA hardening of the running infrastructure. The LMS application is already deployed (by YAT IT after the AT2 handover) and is running in production on the infrastructure students built in AT2. AT3 hardens that running infrastructure for HA — there is no application re-deployment, data migration, or cutover, because those happened in earlier phases. The HA changes are made in-place during the simulated maintenance window. Post-handover, YAT IT re-validates the LMS application against the HA-hardened infrastructure (verifying DOODLE behaves correctly during Multi-AZ failover, session-affinity tuning, etc.) — that re-validation is out of MTS scope per the LMS Migration Role Brief.',
    'This is the third and final assessment task in the cluster. Together with AT1 (Business Case) and AT2 (Foundation Build), AT3 closes out the YAT LMS Cloud Migration engagement.',
    'All scenario materials, organisational policies, the AT2 supplied design, templates, and previous-project examples for YAT College are available on the YAT intranet — students sign in to the intranet at the start of the engagement and refer to it throughout.',
]

TIME = ["Design Phase - ", "Implementation Window – 3.5 hours"]

# EVERGREEN: the committed source read "Access to the YAT intranet (<url>) — supplying ...";
# the URL parenthetical is dropped, every other word verbatim.
RESOURCES = [
    'Teacher/assessor supplied resources',
    'Access to the YAT intranet — supplying the AT2 baseline design, the HA Design template, the HA Deployment Report template, the example past closure exemplar (when authored), the HA database requirements, the LMS application spec, organisational policies, and reference documents',
    'AWS Academy lab access — Cloud Foundations [104469] + Cloud Architecting [172221]',
    'AT2 baseline CloudFormation template — distributed to students at the start of the assessment day for deployment in their AWS Academy lab (scenario/assessor-resources/at2-baseline-cloudformation.{md,yaml})',
    'The AT2 supplied design (scenario/internal-lms-cloud-architecture-design-S1-CL1-AT2.md) — operational reference for the HA design work',
    'HA Design Benchmark + HA Deployment Report Benchmark (later in this document) for marking the two parts',
    'Student supplied resources',
    'Computer with web browser',
    'Word-processing software (Microsoft Word or equivalent)',
    'A screenshot tool',
]

CRITERIA = [
    'To receive a Satisfactory outcome for this assessment the student must:',
    'Achieve Satisfactory on every criterion in the Part A Marking Guide (HA Design)',
    'Achieve Satisfactory on every criterion in the Part B Marking Guide (HA Deployment Report)',
    'Submit a completed HA Design (.docx) and a completed HA Deployment Report (.docx) using the YAT-supplied templates, with every section and every appendix populated',
]

CONDITIONS = 'These are conditions the assessor verifies as present before marking begins. They are not student-performance criteria — they are the conditions under which the assessment can validly be conducted.'

# Marking guide — Part A (A1-A7 HA Design) and Part B (B1-B16 HA Deployment Report). Each criterion is a
# list of cell lines (criterion statement + its inline UoC-traceability line; A4 concatenates the two).
MARKING_A = [
    ['A1 - §2 HA requirements — student documents the reliability, recoverability and service-level requirements the HA design must meet, derived from the LMS application specification and the LMS cloud migration requirements', '[ICTCLD502 PC 1.1]'],
    ['A2 - §3.1–§3.4 Review of AT2 baseline — student reviews the AT2 baseline architecture, identifies all single points of failure, estimates current-state recovery objectives per component, and identifies components requiring vertical scaling with their availability impact', '[ICTCLD502 PC 2.1] · [ICTCLD502 PC 2.2] · [ICTCLD502 PC 2.3] · [ICTCLD502 PC 2.4]'],
    ['A3 - §3.5 Review findings documented — student produces a clear summary of the review findings (the gap between AT2 baseline and the §2 HA requirements) suitable for in-scenario stakeholder review', '[ICTCLD502 PC 2.5]'],
    ['A4 - §4.1–§4.12 HA architecture designed — student produces an HA-equivalent architecture covering all layers (IAM changes if any, cross-AZ network, cross-AZ compute + ASG, cross-AZ ALB, Multi-AZ database, HA-aware storage, security adjustments, HA-tuned monitoring, naming, backup baseline)[ICTCLD502 PC 3.1] · [ICTCLD502 PE 1] (design half) · [ICTCLD502 PE 2] (design half)'],
    ['A5 - §4.13–§4.15 HA design — recovery objectives, vertical-scaling considerations, SPOFs removed — student documents the recovery objectives the HA design achieves per component, the components still requiring vertical scaling (with reduced availability impact), and the SPOFs from the baseline that the HA design removes', '[ICTCLD502 PC 3.2] · [ICTCLD502 PC 3.3] · [ICTCLD502 PC 3.4]'],
    ["A6 - §4 (overall) Design documented per business needs — student produces a complete, internally consistent HA design document that addresses YAT's business needs (99.9% availability target, ≤1h RPO, ≤4h RTO, data residency, etc.) and is suitable for in-scenario stakeholder approval", '[ICTCLD502 PC 3.5]'],
    ['A7 - HA Design document quality — the design uses plain English, appropriate technical vocabulary, structure, formatting, and depth relevant to a professional consulting design deliverable for an RTO client; student has correctly interpreted the supplied AT2 design + scenario reference documents', '[ICTCLD401 FS Writing] · [ICTCLD502 FS Reading]'],
]

MARKING_B = [
    ['B1 - §1 Executive Summary — student produces a concise (≤ 1 page) summary of the HA hardening delivered, including the simulation outcomes in headline form', '[ICTCLD401 PC 4.1] (document and communicate work — written executive-level communication evidence)'],
    ['B2 - §4 Build narrative — HA implementation across all tiers (cross-AZ network, ASG cross-AZ expansion, Multi-AZ database, ALB cross-AZ targets, monitoring HA-tuning, storage adjustments) — student documents the implementation work that translated the HA Design into running infrastructure', '[ICTCLD502 PC 4.1] · [ICTCLD502 PE 1] (implement half) · [ICTCLD502 PE 2] (deploy half — autoscaling) · [ICTCLD502 PE 4] (use cloud management console / SDK / CLI)'],
    ['B3 - §6.1 Connectivity tests — student verifies post-implementation connectivity between resources at all tiers (ALB → both AZs, EC2 → RDS, etc.)', '[ICTCLD502 PC 4.2]'],
    ['B4 - §6.2 Failure simulation — student executes the failure simulations planned in HA Design §6.1, observes outcomes, captures evidence', '[ICTCLD502 PC 4.4] · [ICTCLD502 PE 3] (simulate failures and demonstrate fault tolerance)'],
    ['B5 - §6.3 Resize simulation — student executes the resize simulations planned in HA Design §6.2, measures availability impact, captures evidence', '[ICTCLD502 PC 4.5]'],
    ['B6 - §6.4 Availability measurement — student sets up and uses CloudWatch (or equivalent) to define, monitor, and record resource availability across the maintenance window', '[ICTCLD502 PC 4.3] · [ICTCLD502 PE 5] (define, monitor, record availability — reliability, recoverability, service levels, scalability)'],
    ['B7 - §6.5 Simulation findings vs documented design — student compares the actual simulation outcomes against the expected outcomes specified in the HA Design §6 simulation plan and documents the comparison', '[ICTCLD502 PC 4.6]'],
    ['B8 - §6.6 Adjustments per simulations — student documents any adjustments made to the architecture, configuration, or monitoring based on simulation outcomes (or explicitly justifies no adjustments needed)', '[ICTCLD502 PC 5.1]'],
    ["B9 - §7.4 Documentation filed per organisational policies — student documents the filing of the HA Design + HA Deployment Report + appendices per YAT's Records Management Policy / Backup and Retention Policy", '[ICTCLD401 PC 4.3]'],
    ["B10 - §7.5 Feedback record — student captures the feedback received from the role-played YAT ICT Manager / MTS Senior Consultant on the HA Design and HA implementation, including the student's response", '[ICTCLD502 PC 5.2]'],
    ['B11 - §7.6 Final sign-off obtained — the role-played YAT ICT Manager has signed off on the HA hardening; the sign-off block in §7.6 is populated and signed', '[ICTCLD502 PC 5.3]'],
    ['B12 - §8 Knowledge Evidence Responses — student answers all 6 KE questions with specific reference to their own HA design and implementation, not generic textbook answers', '[ICTCLD502 KE 4] · [ICTCLD502 KE 5] · [ICTCLD502 KE 6] · [ICTCLD502 KE 7] · [ICTCLD502 KE 8] · [ICTCLD502 KE 9]'],
    ["B13 - Appendix A Build Evidence — all named screenshots present (cross-AZ subnets, Multi-AZ RDS, cross-AZ ASG instances, HA-tuned alarms, etc.), correct AWS region visible, named items visible per the template's requirements", '[ICTCLD502 PE 1] · [ICTCLD502 PE 2] (the built HA artefacts) · [ICTCLD502 PE 4] (console usage evidence)'],
    ['B14 - Appendix C Simulation Evidence — all simulation evidence items (failure sim, resize sim, availability measurement) present and consistent with the §6 narrative outcomes', '[ICTCLD502 PE 3] (failure simulation evidence) · [ICTCLD502 PE 5] (availability measurement evidence)'],
    ["B15 - Appendix D Reflections — two honest reflective responses present (decisions in hindsight + problem solving under time pressure), both specific to the student's own HA work", '[ICTCLD401 FS Learning] · [ICTCLD401 FS Planning and organising] · [ICTCLD401 FS Self-management skills] · [ICTCLD502 FS Problem solving] · [ICTCLD502 FS Self-management]'],
    ['B16 - HA Deployment Report document quality — report uses plain English, appropriate technical vocabulary, structure, formatting, and depth relevant to a professional consulting deliverable for an RTO client; student has correctly interpreted the HA Design and other complex technical documentation in the implementation', '[ICTCLD401 FS Writing] · [ICTCLD401 FS Reading] · [ICTCLD502 FS Reading]'],
]

# Shared prose — the 'Instructions to Student' body. NOTE: unlike CL1 AT1, the AT3 student copy carries a
# DIFFERENT (fuller, procedural) instructions body, so the prose is NOT single-sourced across the two
# instruments; each authors its own. Only DETAILS + CHECK are shared (imported by the student builder).
INSTRUCTIONS = [
    'The engagement — picking up where AT2 left off',
    'YAT College is migrating their mission-critical Learning Management System (LMS) from on-premises to AWS. You are an MTS Consultant on this engagement, reporting to Pat Lin (MTS Senior Consultant). Sam Walker (YAT ICT Manager) is your primary YAT-side stakeholder.',
    "At the end of AT2 you handed YAT the cloud foundation — single-AZ, non-HA, ready to operate but not yet meeting YAT's strategic 99.9% availability target. The board, after reviewing the AT2 Deployment Report, has approved the next phase: HA hardening. You are returning to MTS to lead this final phase of the engagement.",
    'Scope of your work in this assessment',
    'The LMS application is already deployed (by YAT IT after the AT2 handover) and is running in production on the AWS infrastructure you built in AT2. Your AT3 work is HA hardening of the existing running infrastructure — Multi-AZ database, cross-AZ compute, HA-tuned monitoring, and so on — performed in-place during the simulated maintenance window. No application re-deployment, data migration, or cutover is required: those happened in earlier phases.',
    "What remains YAT IT's responsibility post-handover: re-validating the LMS application against your HA-hardened infrastructure (e.g. verifying DOODLE behaves correctly during a Multi-AZ database failover, tuning session-affinity at the application layer if needed). That re-validation is out of MTS scope for this assessment per the LMS Migration Role Brief on the YAT intranet.",
    'Your AT3 deliverable stops at HA-hardened infrastructure handed over to YAT IT. You design the HA improvements, implement them in-place, run the simulations, capture feedback, obtain sign-off, file the documentation — and the engagement closes.',
    'The two parts of this assessment',
    'Part A — HA Design',
    "Produce an HA Design document using the YAT-provided HA Design template (available on the YAT intranet's Templates section). Cover:",
    "The HA requirements you're designing against (from the LMS application spec + cloud migration requirements)",
    'A review of the AT2 baseline architecture — identifying single points of failure, estimating current recovery objectives, identifying vertical-scaling components and their availability impact',
    'The HA-equivalent architecture you propose — across all layers (network cross-AZ, compute cross-AZ + ASG expansion, ALB cross-AZ, RDS Multi-AZ, monitoring HA-tuned, storage if applicable)',
    'The recovery objectives, scaling considerations, and SPOFs-removed in your designed state',
    "An implementation sequencing plan (the order you'll apply the changes during the maintenance window)",
    "A simulation plan (the failure and resize simulations you'll run to verify the design works)",
    "The HA Design uses the same shape as the AT2 supplied design — but this time you're the author.",
    'Part B — HA Deployment Report',
    "After producing the HA Design, you implement it in the AWS Academy lab during a simulated Saturday late-night maintenance window of approximately 3.5 hours. You then document everything in the HA Deployment Report using the YAT-provided HA Deployment Report template (available on the YAT intranet's Templates section).",
    'The report mirrors the same shape as the AT2 Deployment Report — but the content within each section is HA-specific:',
    'Build narrative documents the HA changes (not full builds)',
    'Testing section is dominated by the failure simulation + resize simulation (the core of HA verification)',
    'Includes feedback record + final sign-off sections (closing the engagement)',
    'Knowledge Evidence questions cover HA concepts (RTO/RPO, SPOFs, MTTF/MTTR, FT, load balancing for availability, monitoring metrics)',
    'Appendices include HA-specific screenshots (Multi-AZ status, cross-AZ targets, simulation events) and simulation evidence',
    'Maintenance window framing',
    "For Part B's implementation work, you have a simulated Saturday late-night maintenance window of approximately 3.5 hours. Plan your changes to minimise service disruption — most HA changes on AWS are non-disruptive, but the RDS Single-AZ → Multi-AZ conversion incurs a brief failover blip. Brief service blips within the window are acceptable. At the end of the window the LMS infrastructure must be back online either with HA hardening complete or rolled back to the pre-window state.",
    'This is a deliberate framing — at this qualification level, a student working on truly mission-critical infrastructure would be supervised by a senior engineer; the simulated maintenance window is a fair proxy for the real-world risk envelope you will experience.',
    'AT2 baseline starting state',
    'You begin AT3 from a known, consistent AT2 baseline state. The assessor will provide a CloudFormation template at the start of the assessment day; deploy it in your AWS Academy lab during the morning prep slot. The afternoon is then your HA work — design, implement, simulate, report.',
    'Before you start your own HA Design',
    'Read the AT2 supplied design (internal-lms-cloud-architecture-design-S1-CL1-AT2.md on the YAT intranet) end-to-end. Your HA Design follows the same shape but addresses the HA gap.',
    'Submit',
    'Submit both deliverables together via the LMS:',
    'HA Design document (.docx)',
    'HA Deployment Report (.docx) — including all appendices populated',
    'Tips for success',
    "Read the AT2 supplied design + your own AT2 Deployment Report first. They are the inputs to your HA Design — the baseline you're hardening",
    'The architecture review in HA Design §3 is the hardest section to do well. Take time to identify SPOFs carefully — every one you miss is one your simulation will fail to catch',
    'Recovery objectives must be quantified. "Better than before" is not an answer. Cite the specific RPO/RTO targets and what your design delivers',
    "The implementation sequence matters. Plan additive changes first (new subnets, new ASG capacity); leave the RDS Multi-AZ conversion until after you've added the new AZ infrastructure",
    "Simulate honestly. A simulation that confirms the design works is great; a simulation that exposes a gap is also great — both demonstrate competent engineering. Don't fake outcomes",
    "Watch the clock during the maintenance window. If you're 2.5 hours in and not yet stable, start the rollback rather than push to completion",
    "Knowledge Evidence responses (§8) are about your own work. Generic textbook answers don't pass",
    "File the report per YAT's records procedures (§7.4) — this is part of what's being assessed",
    'Get the sign-off block (§7.6) signed at the end of the engagement — without it Part B cannot be Satisfactory',
]

# Assessor-only body — HA Design Benchmark + HA Deployment Report Benchmark + UoC reverse-map.
# Each block is ('h1'|'h2'|'p'|'pn', text) or ('tbl', [[row cells], ...]).
BENCHMARK_A = [
    ('h1', 'HA Design Benchmark'),
    ('p', 'Per-section marking guidance for the assessor — Part A. What "Satisfactory" looks like for each section of the HA Design document.'),
    ('p', '1. Marking framework — Part A'),
    ('p', 'The HA Design is marked Satisfactory / Not Yet Satisfactory as a whole, with per-section judgements feeding the overall.'),
    ('p', "Students may arrive at different specific design choices (e.g. add a NAT Gateway in the second AZ vs share the existing one; min=2 vs min=3; what to do about cross-Region backup) — that's acceptable as long as the choices are internally consistent and tied to the HA requirements."),
    ('p', 'Critical sections (where NYS forces overall NYS): §3 Review of AT2 baseline, §4.1–§4.12 HA architecture design, §4.13–§4.15 designed-state recovery + SPOFs.'),
    ('p', 'Supporting sections: §1, §2, §5, §6, §7, §8.'),
    ('p', '2. Marking §2 HA requirements'),
    ('p', 'UoC evidenced: [ICTCLD502 PC 1.1].'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'The 99.9% availability target stated'),
    ('p', 'RPO ≤ 1 hr and RTO ≤ 4 hr targets stated'),
    ('p', 'Source documents named (LMS application spec, LMS cloud migration requirements)'),
    ('p', 'Any other reliability/recoverability/service-level requirements identified from the brief'),
    ('p', "NYS: targets missing or vague; doesn't cite source documents; misinterprets target values."),
    ('p', '3. Marking §3 Review of AT2 baseline (critical section)'),
    ('p', 'UoC evidenced: [ICTCLD502 PC 2.1, 2.2, 2.3, 2.4, 2.5].'),
    ('p', 'Satisfactory looks like:'),
    ('p', '§3.1 Architecture review: all layers reviewed (network, compute, ALB, DB, storage, monitoring) against HA expectations'),
    ('p', '§3.2 SPOFs: the major SPOFs identified — at minimum: single-AZ deployment (whole AZ failure kills LMS), single RDS instance (instance failure kills DB), ASG min=1 (single instance failure kills compute capacity), single NAT Gateway (NAT failure kills outbound from private subnet)'),
    ('p', '§3.3 Recovery objectives current state: quantified estimates per component — e.g. "RDS Single-AZ: RPO ~24h (last backup), RTO 2–6h (restore time)"'),
    ('p', '§3.4 Vertical scaling components: identified and impact stated'),
    ('p', '§3.5 Summary: clear statement of the gap between AT2 baseline and the §2 requirements'),
    ('p', 'NYS: SPOFs missing (especially the obvious AZ-failure one); recovery objectives not quantified; review reads as generic architecture description rather than HA gap analysis.'),
    ('p', '4. Marking §4 HA architecture designed (critical section)'),
    ('p', 'UoC evidenced: [ICTCLD502 PC 3.1] · [ICTCLD502 PE 1] (design half) · [ICTCLD502 PE 2] (design half).'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'Cross-AZ subnets designed in (-b subnets added in ap-southeast-2b)'),
    ('p', 'ASG cross-AZ expansion designed (min ≥ 2, distributed across both AZs)'),
    ('p', 'ALB extended to include the new AZ subnets'),
    ('p', 'RDS Multi-AZ designed (standby in second AZ, automatic failover, single endpoint)'),
    ('p', 'HA-tuned monitoring designed (per-AZ instance counts, RDS failover detection, service-level metric)'),
    ('p', 'Optional but encouraged: cross-Region backup (snapshot copies or AWS Backup)'),
    ('p', 'Each design decision tied back to a §3 finding (e.g. "the design removes the ASG-min=1 SPOF identified in §3.2 by expanding to min=2 across two AZs")'),
    ('p', 'NYS: any major layer left non-HA (e.g. RDS still Single-AZ in the design); design decisions not tied to identified gaps; design contradicts the supplied AT2 baseline assumptions (e.g. proposes changing region or application stack — both out of scope).'),
    ('p', '5. Marking §4.13–§4.15 designed-state recovery + SPOFs removed (critical section)'),
    ('p', 'UoC evidenced: [ICTCLD502 PC 3.2, 3.3, 3.4].'),
    ('p', 'Satisfactory looks like:'),
    ('p', '§4.13: designed recovery objectives quantified per component; LMS overall meets the §2 targets (RPO ≤ 1h, RTO ≤ 4h)'),
    ('p', '§4.14: vertical-scaling components identified post-HA, with the (typically much lower) availability impact stated'),
    ('p', '§4.15: each SPOF from §3.2 cross-referenced with its mitigation in the design — every baseline SPOF accounted for'),
    ('p', "NYS: designed recovery objectives don't meet §2 targets without explanation; SPOF mitigations missing for SPOFs identified in §3.2."),
    ('p', '6. Marking §4 Documented per business needs'),
    ('p', 'UoC evidenced: [ICTCLD502 PC 3.5].'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'Design is complete and internally consistent'),
    ('p', 'Suitable for in-scenario stakeholder review (Pat Lin / Sam Walker could read it and approve)'),
    ('p', 'All sections of the template populated'),
    ('p', 'NYS: template sections left empty; design described at a level too abstract to act on; internal contradictions (e.g. §4.1 says no IAM changes but §4.6 references a new IAM role).'),
    ('p', '7. Marking HA Design document quality'),
    ('p', 'UoC evidenced: [ICTCLD401 FS Writing] · [ICTCLD502 FS Reading].'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'Plain English, appropriate technical vocabulary'),
    ('p', 'Diagrams or tables used where they aid clarity (HA topology, recovery objectives table, SPOF mitigation table)'),
    ('p', 'Source documents correctly interpreted and cited'),
    ('p', 'NYS: unintelligible prose; design references contradicted by the source documents; no diagrams or tables in a document that needs them.'),
]

BENCHMARK_B = [
    ('h1', 'HA Deployment Report Benchmark'),
    ('p', 'Per-section marking guidance for the assessor — Part B. What "Satisfactory" looks like for each section of the HA Deployment Report. Model answers for the KE questions. Common deductions.'),
    ('p', '8. Marking framework — Part B'),
    ('p', 'The HA Deployment Report is marked Satisfactory / Not Yet Satisfactory as a whole, with per-section judgements feeding the overall.'),
    ('p', 'Specific implementation choices can vary (Multi-AZ apply-immediately vs maintenance-window; min=2 vs min=3 ASG capacity) — acceptable as long as outcomes match the HA Design and the simulations confirm fault tolerance.'),
    ('p', 'Critical sections (where NYS forces overall NYS): §4 Build narrative, §6.2 Failure simulation, §6.3 Resize simulation, §6.5 Simulation findings, §7.6 Sign-off, §8 KE Responses, Appendix A Build evidence, Appendix C Simulation evidence.'),
    ('p', 'Supporting sections: §1, §2, §3, §5, §7.1–§7.3, Appendix B, Appendix D.'),
    ('p', '9. Marking §1 Executive Summary'),
    ('p', 'UoC evidenced: [ICTCLD401 PC 4.1] (partial).'),
    ('p', 'Satisfactory: one page or less; names what was delivered; states whether 99.9% target is achievable; simulation outcomes in headline form; any deferred items called out.'),
    ('p', "NYS: missing; reads as full report rather than summary; doesn't reconcile with §6 outcomes."),
    ('p', '10. Marking §4 Build narrative (critical section)'),
    ('p', 'UoC evidenced: [ICTCLD502 PC 4.1] · [ICTCLD502 PE 1, 2, 4].'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'All HA changes documented layer-by-layer'),
    ('p', "What was changed vs the AT2 baseline is explicit (don't just describe the new state — describe the change)"),
    ('p', 'Cross-references to Appendix A screenshots and Appendix B configuration exports'),
    ('p', 'NYS: layers missing (e.g. no Multi-AZ DB section); narrative contradicts the screenshot evidence; reads as generic AWS HA documentation rather than what was actually changed.'),
    ('p', '11. Marking §6 Testing, Simulation and Validation (critical section)'),
    ('p', 'UoC evidenced: [ICTCLD502 PC 4.2, 4.3, 4.4, 4.5, 4.6, 5.1] · [ICTCLD502 PE 3, 5].'),
    ('p', '11.1 §6.1 Connectivity tests (criterion B3)'),
    ('p', 'All four post-implementation tests recorded with Pass/Fail. Negative — ensure RDS still not publicly reachable.'),
    ('p', '11.2 §6.2 Failure simulation (criterion B4 — critical)'),
    ('p', 'At least one failure simulation executed (EC2 instance termination, or RDS failover, or AZ network partition)'),
    ('p', 'Outcome captured with timing data (recovery time)'),
    ('p', 'LMS reachability status recorded'),
    ('p', 'Appendix C evidence attached'),
    ('p', 'NYS: no simulation executed; "I would have tested..." rather than "I tested..."; evidence fabricated (no supporting Appendix C entry)'),
    ('p', '11.3 §6.3 Resize simulation (criterion B5 — critical)'),
    ('p', 'At least one resize simulation executed'),
    ('p', 'Outcome captured'),
    ('p', 'Availability impact measured'),
    ('p', 'Appendix C evidence attached'),
    ('p', 'NYS: as above'),
    ('p', '11.4 §6.4 Availability measurement (criterion B6)'),
    ('p', 'CloudWatch dashboard / metric set up to measure availability'),
    ('p', 'Window-wide availability percentage computed (typically 99.5–99.9% across the window depending on simulation blip durations)'),
    ('p', 'NYS: no measurement infrastructure set up; availability stated as a single number with no supporting data'),
    ('p', '11.5 §6.5 Simulation findings (criterion B7 — critical)'),
    ('p', 'Actual outcomes compared against the expected outcomes from HA Design §6'),
    ('p', 'Divergences documented with reasons'),
    ('p', 'NYS: no comparison; "all worked as expected" without specifics'),
    ('p', '11.6 §6.6 Adjustments per simulations (criterion B8)'),
    ('p', 'Adjustments made are documented, OR a justified statement that no adjustments were needed'),
    ('p', 'NYS: simulations exposed a gap but no adjustment recorded'),
    ('p', '12. Marking §7 Operational Handover'),
    ('p', '12.1 §7.4 Documentation filing (criterion B9)'),
    ('p', '[ICTCLD401 PC 4.3]. Student names where the HA Design + HA Deployment Report + appendices are filed per internal-records-management-policy.md / internal-backup-retention-policy.md.'),
    ('p', 'NYS: filing not addressed.'),
    ('p', '12.2 §7.5 Feedback record (criterion B10)'),
    ('p', '[ICTCLD502 PC 5.2]. Feedback captured from at least one stakeholder; student response documented; resulting actions noted (or "no action required" stated with reason).'),
    ('p', 'NYS: feedback absent.'),
    ('p', '12.3 §7.6 Sign-off (criterion B11 — critical)'),
    ('p', '[ICTCLD502 PC 5.3]. Sign-off block populated and signed by the role-played YAT ICT Manager.'),
    ('p', 'NYS: sign-off block blank — implies the engagement was not formally closed.'),
    ('p', '13. Marking §8 Knowledge Evidence Responses (critical section)'),
    ('p', 'UoC evidenced: [ICTCLD502 KE 4, 5, 6, 7, 8, 9].'),
    ('p', 'Below are model answers for each question. Student answers may differ in detail but must demonstrate the underlying knowledge applied to their own HA design + implementation.'),
    ('p', 'Q1 — HA concepts in your design [ICTCLD502 KE 4]'),
    ('p', 'Model coverage: Student picks ≥ 3 of: fault tolerance, SPOFs, RPO/RTO, SLAs, vertical vs horizontal scaling, MTTF/MTTR/MTBF. For each: definition in own words, named application in their work, and design impact.'),
    ('p', 'Example: "Fault tolerance: the system continues operating in the presence of component failures. In my work, fault tolerance shows up in the RDS Multi-AZ design (§4.5 of my HA Design) — when the primary AZ fails, the standby takes over automatically. This affected my implementation choice: I converted RDS to Multi-AZ before expanding the ASG, because the database failover was the riskier change and needed validating first."'),
    ('p', 'NYS: generic textbook definitions without reference to own work; covers fewer than 3 concepts.'),
    ('p', 'Q2 — Testing and debugging techniques [ICTCLD502 KE 5]'),
    ('p', 'Model coverage:'),
    ('p', '(a) A technique used to avoid causing a SPOF during testing — e.g. "I generated traffic from the surviving AZ while testing failure in the other AZ, so I always had a known-good baseline to compare against"'),
    ('p', '(b) A debugging technique that could isolate cause of unexpected outcome — e.g. checking RDS event log + CloudWatch + application connection retry timing to isolate whether a long failover time was DB-side or application-side'),
    ('p', "NYS: generic testing definitions; doesn't address either sub-part."),
    ('p', 'Q3 — Tools and techniques to measure availability impact [ICTCLD502 KE 6]'),
    ('p', 'Model coverage:'),
    ('p', '(a) Specific AWS tools used (CloudWatch metrics, ALB access logs, RDS event log)'),
    ('p', '(b) How availability impact was computed for one specific simulation (with timing data)'),
    ('p', '(c) Limitation acknowledged (e.g. "my 1-second poll missed any sub-second blips")'),
    ('p', "NYS: doesn't name specific tools used; can't show their computation; no limitation acknowledged."),
    ('p', 'Q4 — Built-in vs designed fault tolerance [ICTCLD502 KE 7]'),
    ('p', 'Model coverage:'),
    ('p', 'S3: built-in (multi-AZ within a region by default; cross-region needs replication setup)'),
    ('p', 'RDS: hybrid — instance is single-AZ by default (you design Multi-AZ); backups are managed by AWS'),
    ('p', 'ALB: built-in (multi-AZ-capable from creation; you configure which AZs to use)'),
    ('p', 'NYS: student gets the responsibility line wrong (e.g. claims ALB needs you to design failover); generic statements without service-specific detail.'),
    ('p', 'Q5 — Load balancing + autoscaling for availability [ICTCLD502 KE 8]'),
    ('p', 'Model coverage:'),
    ('p', '(a) Combination delivers availability above what either provides alone — load balancer routes traffic to healthy targets; autoscaler maintains target count to ensure healthy targets exist'),
    ('p', '(b) Specific scenario from their work — e.g. failure simulation F1 (EC2 termination) was survivable because ASG replaced the instance AND ALB de-registered the failed target'),
    ('p', "NYS: generic textbook description; can't tie to specific scenario in their work."),
    ('p', 'Q6 — Monitoring techniques and metrics [ICTCLD502 KE 9]'),
    ('p', 'Model coverage:'),
    ('p', 'One HA-specific metric named (e.g. per-AZ EC2 healthy target count, RDS replica lag, ALB 5XX rate)'),
    ('p', 'What it measures, threshold, why'),
    ('p', 'Failure mode it detects'),
    ('p', 'NYS: generic CloudWatch description; no specific metric / threshold / failure mode named.'),
    ('p', 'Marking checks for §8 overall'),
    ('p', 'Award credit for:'),
    ('p', 'All 6 questions answered'),
    ('p', "Each answer references specific sections of the student's own HA Design and HA Deployment Report"),
    ('p', 'Cited content is actually present in their documents'),
    ('p', 'Common deductions:'),
    ('p', 'Generic textbook answers'),
    ('p', "Citations to sections that don't contain what the answer claims"),
    ('p', "Repeating supplied content (e.g. AT2 design's shared-responsibility table) instead of synthesising"),
    ('p', '14. Marking Appendix A — Build evidence (critical section)'),
    ('p', 'UoC evidenced: [ICTCLD502 PE 1, 2, 4].'),
    ('p', 'Satisfactory: all required screenshots present; correct region indicator visible; named items visible.'),
    ('p', "Common deductions: region indicator missing (can't verify ap-southeast-2); screenshot doesn't show the named items (e.g. RDS screenshot doesn't show Multi-AZ status); cross-AZ ASG screenshot only shows one AZ."),
    ('p', '15. Marking Appendix B — Configuration exports'),
    ('p', 'UoC evidenced: [ICTCLD502 PE 4].'),
    ('p', "Satisfactory: all named exports present; consistent with Appendix A screenshots; recognisably from the student's own deployed environment."),
    ('p', '16. Marking Appendix C — Simulation evidence (critical section)'),
    ('p', 'UoC evidenced: [ICTCLD502 PE 3, 5].'),
    ('p', 'Satisfactory: evidence supports the §6 narrative; failure simulation evidence shows the actual failure-and-recovery event (not a screenshot of a tutorial); availability measurement data is from the actual maintenance window.'),
    ('p', 'Common deductions: evidence inconsistent with the recorded §6 outcome; missing negative test (RDS still not publicly reachable post-HA); evidence from a fictional / earlier deployment.'),
    ('p', '17. Marking Appendix D — Reflections'),
    ('p', 'UoC evidenced: [ICTCLD401 FS Learning, Planning and organising, Self-management skills] · [ICTCLD502 FS Problem solving, Self-management].'),
    ('p', 'Satisfactory:'),
    ('p', "R1 names a decision validated AND a decision to revise — both specific to the student's HA work"),
    ('p', 'R2 describes a moment of problem-solving under time pressure (or honestly states no such moment, with a likely-gotcha analysis)'),
    ('p', 'NYS: reflections absent; all positive (no honest "would revise"); generic content not tied to actual HA work.'),
    ('p', '18. Common deductions across the HA Deployment Report'),
    ('p', 'Build narrative inconsistent with screenshots — §4 describes one configuration; appendices show another'),
    ('p', 'Simulation outcomes recorded without supporting evidence — §6 claims a Pass without Appendix C'),
    ('p', 'Sign-off block left blank — engagement not formally closed'),
    ('p', 'KE responses divorced from own work — answers could be cut-pasted into any HA report'),
    ('p', 'Maintenance-window constraints not respected — student documents work that extends well beyond a 3.5-hour window without rollback acknowledgement'),
]

# UoC coverage verification (reverse map): heading + prose + three unit tables (with in-between notes).
REVMAP_502 = [
    ['UoC item', 'Evidenced by criterion(ia)'],
    ['PC 1.1 — Determine reliability, recoverability and service levels required for application', 'A1'],
    ['PC 2.1 — Review architecture of traditional multi-tier web application + identify HA requirements', 'A2'],
    ['PC 2.2 — Identify any single points of failure', 'A2'],
    ['PC 2.3 — Estimate recovery objectives for multi-tier web components and overall architecture', 'A2'],
    ['PC 2.4 — Determine components that must scale vertically and the potential impact on system availability', 'A2'],
    ['PC 2.5 — Document architecture review findings according to business needs', 'A3'],
    ['PC 3.1 — Design equivalent architecture for high availability using cloud services', 'A4'],
    ['PC 3.2 — Identify and remove single points of failure as required', 'A5'],
    ['PC 3.3 — Estimate recovery objectives for each component and overall architecture (designed state)', 'A5'],
    ['PC 3.4 — Determine components that must scale vertically and the potential impact on system availability (designed state)', 'A5'],
    ['PC 3.5 — Document architecture design according to business needs', 'A6'],
    ['PC 4.1 — Implement architecture design in cloud environment', 'B2'],
    ['PC 4.2 — Demonstrate connectivity between resources at all tiers', 'B3'],
    ['PC 4.3 — Monitor and measure availability of resources', 'B6 (full HA-tuned monitoring; AT2 carried baseline-only)'],
    ['PC 4.4 — Simulate failures of component and confirm that infrastructure is fault tolerant', 'B4'],
    ['PC 4.5 — Simulate resizing components likely to impact performance and measure availability impact', 'B5'],
    ['PC 4.6 — Compare and document simulation findings according to documented design', 'B7'],
    ['PC 5.1 — Adjust and improve availability of architecture according to simulations as required', 'B8'],
    ['PC 5.2 — Confirm, seek and respond to feedback with required personnel', 'B10'],
    ['PC 5.3 — Obtain final sign off from required personnel', 'B11'],
    ['PE 1 — Design and implement at least one fault tolerant cloud infrastructure', 'A4 (design half) · B2 (implement half) · B13 (built-artefact evidence)'],
    ['PE 2 — Design and deploy automated infrastructure scaling for at least one business need', 'A4 (design half) · B2 (deploy half) · B13 (built-artefact evidence)'],
    ['PE 3 — Simulate failures of at least one component and demonstrate fault tolerance', 'B4 · B14 (evidence)'],
    ['PE 4 — Use cloud management console, software development kits or command line tools', 'B2 · B13 (screenshots evidence console usage)'],
    ['PE 5 — Define, monitor and record resource availability', 'B6 · B14 (availability measurement evidence)'],
    ['KE 4 — HA cloud-infrastructure concepts (fault tolerance, SPOFs, MTTF/MTTR/MTBF, RTO/RPO, SLAs, vertical/horizontal scalability)', 'B12 (Q1)'],
    ['KE 5 — Testing and debugging techniques (including techniques to avoid SPOFs)', 'B12 (Q2)'],
    ['KE 6 — Tools and techniques to measure availability impact', 'B12 (Q3)'],
    ['KE 7 — Features of cloud services — built-in vs designed fault tolerance', 'B12 (Q4)'],
    ['KE 8 — Load balancing + autoscaling for availability', 'B12 (Q5)'],
    ['KE 9 — Performance monitoring techniques and metrics', 'B12 (Q6)'],
    ['FS Reading', 'A7 · B16'],
    ['FS Problem solving', 'B15 (Appendix D Reflections)'],
    ['FS Self-management', 'B15 (Appendix D Reflections)'],
    ['AC 1 — Cloud vendor service provider', 'C1 (pre-condition)'],
    ['AC 2 — Cloud managed database service', 'C1 (pre-condition)'],
    ['AC 3 — Information and data sources required to design and implement cloud infrastructure', 'C2 (pre-condition)'],
    ['AC 4 — Integrated development environment (IDE)', 'C1 (pre-condition)'],
    ['AC 5 — Specific requirements + industry standards + organisational procedures + legislative requirements + business and functionality requirements', 'C2 (pre-condition)'],
    ['AC 6 — Internet and web browser', 'C1 (pre-condition)'],
    ['AC 7 — Secure shell (SSH) or remote desktop protocol (RDP) client', 'C1 (pre-condition)'],
    ['AC 8 — Data to gather information from to determine output and user requirements', 'C2 (pre-condition)'],
]

REVMAP_401 = [
    ['UoC item', 'Evidenced by criterion(ia)'],
    ['PC 4.1 — Document and communicate work to required personnel', 'B1 (partial — written executive-level communication; fuller PC 4.1 was AT1)'],
    ['PC 4.3 — Save and store user documentation according to organisational policies', 'B9'],
    ['FS Reading', 'B16 (partial)'],
    ['FS Writing', 'A7 · B16'],
    ['FS Learning', 'B15 (Appendix D Reflections)'],
    ['FS Planning and organising', 'B15 (Appendix D Reflections)'],
    ['FS Self-management skills', 'B15 (Appendix D Reflections)'],
    ['AC 4 — Data to gather information from to determine output and user requirements', 'C2 (pre-condition)'],
]

REVMAP_517 = [
    ['UoC item', 'Evidenced by criterion(ia)'],
    ['—', "AT3 does not evidence any 517 PCs/PEs/KEs/FSs/ACs directly. 517 is primarily AT1 (Business Case + presentation event); the AT3 work is a downstream consequence of AT1's recommendations but does not re-evidence 517 items."],
]

REVERSE_MAP = [
    ('h1', 'UoC coverage verification (reverse map)'),
    ('pn', 'This table closes the loop on bidirectional traceability: every UoC requirement AT3 claims to evidence is named below with the marking criterion (or criteria) that evidence it. No UoC requirement claimed by AT3 is left without a criterion.'),
    ('pn', 'ICTCLD502 — Design and implement highly-available cloud infrastructure (AT3-evidenced items)'),
    ('tbl', REVMAP_502),
    ('pn', 'ICTCLD401 — Configure cloud services (AT3-evidenced items only)'),
    ('tbl', REVMAP_401),
    ('pn', '(Note: 401 PCs not listed above (PC 1.1–1.7, 1.8, 2.1–2.6, 3.1–3.2, 4.2) are evidenced in AT1 (PC 1.8) or AT2 (the rest). 401 KEs are not directly evidenced in AT3 — 401 KE 1–4, 11 evidenced in AT1; 401 KE 5–10 evidenced in AT2.)'),
    ('pn', 'ICTICT517 — Match ICT needs with the strategic direction of the organisation (AT3-evidenced items)'),
    ('tbl', REVMAP_517),
]

ASSESSOR_BODY = BENCHMARK_A + BENCHMARK_B + REVERSE_MAP


# ---------- build helpers ----------

STYLE = {"h1": "Heading 1", "h2": "Heading 2", "p": "Assessor text", "pn": "Normal"}


def add_marking_row(table, lines, check=CHECK):
    """A marking-guide criterion row: criterion (+ optional UoC line) in col 0, the check box in col 1."""
    row = table.add_row()
    set_cell_content(row.cells[0], lines)
    set_cell_content(row.cells[1], check)


def add_partb_section_row(table, text):
    """The 'Part B' section row — bold label in col 0, and (per the authored instrument) a check in col 1."""
    row = table.add_row()
    set_cell_content(row.cells[0], text)
    for r in row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell_content(row.cells[1], CHECK)


def delete_body_paragraph(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            p._element.getparent().remove(p._element)
            return


def render_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    for r, cells in enumerate(rows):
        for c, val in enumerate(cells):
            set_cell_content(t.rows[r].cells[c], val)
    doc.add_paragraph()


def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details ----
    # NOTE: the authored instrument leaves the task-number row at the template default ('__ of __')
    # and instead carries '3 of 3' on the Teacher/Assessor-name row — reproduced faithfully here.
    t_details = doc.tables[0]
    set_cell_content(t_details.rows[1].cells[1], DETAILS["qualification"])
    set_cell_content(t_details.rows[2].cells[1], DETAILS["units"])
    set_cell_content(t_details.rows[3].cells[1], DETAILS["task_title"])
    set_cell_content(t_details.rows[5].cells[1], DETAILS["task_number"])

    # ---- Table 1: Teacher/Assessor instructions ----
    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Time allowed"), TIME)
    set_cell_content(find_instruction_row(t_instr, "Location"), "")
    set_cell_content(find_instruction_row(t_instr, "Resources required"), RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA)
    # 'Second attempt' and 'Assessment retention' keep the template's standard text.
    # add a Conditions row at the end (matching the table style), as CL1/CL2/CL3 do
    cond_row = t_instr.add_row()
    set_cell_content(cond_row.cells[0], "Assessment Conditions & Setup Requirements")
    for r in cond_row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell_content(cond_row.cells[1], CONDITIONS)

    # ---- Table 2: Marking Guide (criteria + inline UoC traceability) ----
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)  # keep 'Assessment criteria' + 'Criteria | Satisfactory?' header rows
    add_section_row(t_mark, "Part A")
    for lines in MARKING_A:
        add_marking_row(t_mark, lines)
    add_partb_section_row(t_mark, "Part B")
    for lines in MARKING_B:
        add_marking_row(t_mark, lines)

    # ---- drop the template's trailing marking-guide boilerplate ----
    delete_body_paragraph(doc, "Add or delete rows as required")
    delete_body_paragraph(doc, "If questioning or observation is incorporated into this assessment task, "
                                "you can incorporate a Practical Observation Checklist.")

    # ---- Instructions to Student (shared prose) ----
    doc.add_paragraph("Instructions to Student", style="Heading 1")
    for text in INSTRUCTIONS:
        doc.add_paragraph(text, style="Assessor text")

    # ---- Assessor-only body (Benchmarks + reverse-map) ----
    for kind, payload in ASSESSOR_BODY:
        if kind == "tbl":
            render_table(doc, payload)
        else:
            doc.add_paragraph(payload, style=STYLE[kind])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/assessments/AT3/AT3-High-Availability-Assessor.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
