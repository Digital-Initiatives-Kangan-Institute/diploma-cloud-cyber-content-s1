#!/usr/bin/env python3
"""Build the worked LMS Business Case EXEMPLAR (.docx) — assessor marking reference.

A complete, worked model answer for the YAT LMS Cloud Migration Business Case (S1-CL1 AT1).
Reuses the branding helpers from the helpers/ package so it carries identical YAT brand
styling, and retains the UoC `Evidences:` tags (this artefact is AT1-linked and internal
to assessors). Figures are indicative and internally consistent (AUD ex GST,
ap-southeast-2); the AWS line items would be backed by a Pricing Calculator export.

Usage:  python scripts/s1_cl1/build_s1_cl1_at1_bc_exemplar.py [output.docx]
Default output: scenario/exemplars/internal-document-exemplar-business-case.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_body_text import add_body_paragraph, add_bullet_list  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from helpers.docx_tables import add_data_table  # noqa: E402
from helpers.uoc_tags import add_uoc_evidence_tag  # noqa: E402
from brand import ADDRESS, CREAM, GREY, TEAL, TERRACOTTA, WHITE  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path):
    doc = Document()
    configure_styles(doc)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Business Case")
    sub = doc.add_paragraph().add_run("YAT LMS Cloud Migration")
    sub.font.size = Pt(15); sub.font.color.rgb = RGBColor.from_string(TERRACOTTA); sub.bold = True
    note = doc.add_paragraph().add_run("Assessor exemplar — internal marking reference (not for distribution to students)")
    note.italic = True; note.font.size = Pt(10); note.font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph()

    cover = [
        ("Prepared for", "YAT Board, via Sam Walker (YAT ICT Manager)"),
        ("Prepared by", "MTS Consultant, MP Tech Solutions (MTS)"),
        ("Reviewed by", "Pat Lin, MTS Senior Consultant"),
        ("Engagement", "YAT LMS Cloud Migration — S1-CL1"),
        ("Document version", "v1.0"),
        ("Analysis period", "5 years"),
        ("Currency", "AUD, exclusive of GST"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h2 = lambda t: doc.add_paragraph(t, style="Heading 2")

    h1("1. Executive Summary")
    add_uoc_evidence_tag(doc, "contributes to [ICTICT517 PC 1.4]")
    add_body_paragraph(doc, "YAT's Learning Management System (DOODLE) runs on a single, end-of-supportable-life "
              "on-premises server delivering ~99.2% availability — short of the 99.9% target the ICT "
              "Strategic Plan sets for critical systems, and unable to meet the recovery and capacity "
              "expectations of a growing RTO. This Business Case evaluates two options over a five-year "
              "horizon: Option A — renew the on-premises server; Option B — migrate the LMS to Amazon Web "
              "Services (AWS) in the Sydney region (ap-southeast-2), preserving the existing Windows "
              "Server 2016 / DOODLE / MySQL stack.")
    add_body_paragraph(doc, "The five-year direct cost of the two options is comparable (~$447,300 on-premises vs "
              "~$441,400 cloud). The decision therefore turns on the non-cost factors: only Option B "
              "credibly achieves 99.9% availability and the RTO ≤ 4h / RPO ≤ 1h recovery objectives, and "
              "only Option B provides the elastic capacity for 15% annual student growth and the ~3× "
              "assessment-window peaks. Quantifying the avoided downtime adds ~$69,000 of benefit to the "
              "cloud option over five years, giving it the stronger net position (~$372,400 vs ~$447,300).")
    add_body_paragraph(doc, "The principal risks of the cloud option — staff upskilling, cost control, and a safe "
              "cutover — are manageable through continued MTS support, AWS cost guardrails, and phased "
              "migration outside assessment windows. We recommend Option B and ask the board to approve "
              "the migration, the phased action plan in §10, and the Year-1 budget envelope.")

    h1("2. Engagement Context")
    add_uoc_evidence_tag(doc, "[ICTICT517 AC 4]")
    add_body_paragraph(doc, "This Business Case is prepared by MP Tech Solutions (MTS) for the YAT Board, via the YAT "
              "ICT Manager (Sam Walker), under the MTS–YAT Master Services Agreement, and reviewed by the "
              "MTS Senior Consultant (Pat Lin) before submission. MTS has been engaged to analyse YAT's "
              "LMS operating-model options and recommend a way forward. It asks the board to decide "
              "whether to renew the LMS on-premises or migrate it to AWS, and to approve the associated "
              "action plan and Year-1 budget.")

    h1("3. Strategic Alignment Analysis")
    add_uoc_evidence_tag(doc, "[ICTICT517 PC 1.1] · [ICTICT517 PE 1]")
    add_body_paragraph(doc, "YAT's ICT Strategic Plan sets goals to ensure reliable ICT services, respond quickly to "
              "change, reduce dependency on in-house server infrastructure, and adopt sustainable "
              "practices. Three objectives are material to this engagement: investigate transitioning "
              "critical on-premises application services to the cloud (by mid-2027); deploy the LMS in the "
              "cloud (by end-2027); and achieve 99.9% availability for critical systems (by end-2027). The "
              "plan names the LMS as the first system to transition, making this engagement the lead "
              "initiative of the strategy. These ICT objectives in turn serve YAT's business objectives — "
              "location-independent learning, national expansion, and 15% annual student growth — for "
              "which the plan identifies cloud-hosted infrastructure as the principal ICT enabler.")
    add_body_paragraph(doc, "Against the industry environment, the direction is well aligned: Australian RTOs and the "
              "wider education sector are moving steadily to cloud and managed services, shifting from "
              "capital refresh cycles (CAPEX) to consumption-based operating costs (OPEX), and treating "
              "resilience and elastic capacity as baseline expectations rather than premium features. "
              "Data-residency obligations under the Privacy Act 1988 (APP 8) and the Standards for RTOs "
              "2015 are equally part of that environment, and are met by deploying in an Australian region.")
    add_body_paragraph(doc, "Where YAT's plan is more ambitious than its current posture is the 99.9% availability "
              "target — a meaningful step up from the measured 99.2% and one a single on-premises server "
              "cannot credibly deliver. That ambition is appropriate, not excessive: it reflects the LMS's "
              "mission-critical role. The implication is clear — the LMS migration is the proving ground "
              "for the broader strategy, and delivering it well sets the pattern for moving YAT's other "
              "suitable on-site systems to the cloud.")

    h1("4. Current State of YAT's ICT")
    add_uoc_evidence_tag(doc, "[ICTICT517 PC 1.2] · [ICTICT517 PE 2]")
    add_body_paragraph(doc, "YAT's network is resilient at the infrastructure layer (no single point of failure, "
              "dual-NIC servers, two AD-controlled security zones), but the LMS itself is not. DOODLE runs "
              "on a single Windows Server 2016 host with a MySQL database and ~178 GB of data (of ~250 GB "
              "allocated), and the server is at end of supportable life. Measured availability is 99.2%; "
              "recovery rests on nightly tape backups, giving an RPO of ~24 hours and an RTO of ~7–11 "
              "hours — both well short of the migration targets. The LMS integrates with Active Directory "
              "(authentication), Office 365 (notifications), the Accounting system (fee-status), and ASQA/"
              "DET reporting. YAT ICT staff are highly capable on-premises but have limited cloud "
              "experience; MTS provides application support under a standing contract. The material "
              "limitations for this decision are the single point of failure, the end-of-life hardware, "
              "recovery objectives far from target, and fixed capacity that cannot absorb the ~3× "
              "assessment-window peaks or 15% annual growth without over-provisioning.")

    h1("5. Gap Analysis")
    add_uoc_evidence_tag(doc, "[ICTICT517 PC 1.3] · [ICTICT517 PE 3]")
    add_data_table(doc,
           ["Strategic objective", "Current state", "Desired future state", "Gap",
            "Improvement opportunity", "Proposed change"],
           [["99.9% availability for critical systems",
             "99.2% on a single server", "99.9% (Multi-AZ resilient)",
             "~0.7 pp (~61 hr/yr); single server cannot deliver", "Highly available architecture",
             "Migrate to AWS across two Availability Zones"],
            ["Reduce dependency on in-house infrastructure",
             "Owns/runs end-of-life server + room", "No in-house server to own or run",
             "Full reliance on self-run hardware", "Managed cloud operating model",
             "Migrate LMS to AWS (managed services)"],
            ["Recover quickly (RTO ≤ 4h / RPO ≤ 1h)",
             "RTO ~7–11h, RPO ~24h (nightly tape)", "RTO ≤ 4h, RPO ≤ 1h",
             "Hours-to-a-day short on both", "Automated backups + failover",
             "RDS Multi-AZ + AWS Backup"],
            ["Scale for growth and assessment peaks",
             "Fixed server, no elasticity", "Auto-scale to ~3× peaks; absorb 15%/yr",
             "No elasticity; over-provision or degrade", "Elastic compute",
             "EC2 Auto Scaling; scale in off-peak"]],
           widths=[2.9, 2.6, 2.6, 2.6, 2.5, 2.6])
    add_body_paragraph(doc, "The gaps cluster around resilience, recovery, capacity, and ownership. An on-premises "
              "refresh (Option A) addresses end-of-life hardware but leaves the single point of failure, "
              "the recovery shortfall, and the fixed capacity largely intact. A cloud operating model "
              "(Option B) addresses all four. This is the central finding the options analysis tests.")

    h1("6. Options Considered and Evaluation")
    add_uoc_evidence_tag(doc, "[ICTICT517 PC 2.1] · [ICTICT517 PC 2.2] · [ICTICT517 KE 3] · [ICTCLD401 PC 1.8] · [ICTCLD502 PC 1.2]")
    h2("6.1 Workload definition")
    add_body_paragraph(doc, "The LMS serves ~860 users (≈800 students + 60 staff), with 200–300 typical concurrent "
              "users, 500–700 at peak (~3× for the ~10–14-day assessment window each term), and very low "
              "overnight/term-break load. Data is ~178 GB, growing ~25 GB/year. It must preserve the "
              "Windows Server 2016 / DOODLE / MySQL stack with full OS control, keep its AD, Office 365, "
              "Accounting and ASQA integrations, hold all data in Australia, and meet 99.9% availability, "
              "RTO ≤ 4h and RPO ≤ 1h.")
    h2("6.2 Options considered")
    add_bullet_list(doc, [
        "Option A — In-house renewal: replace the end-of-life server and ancillary hardware; retain the "
        "single-site on-premises operating model.",
        "Option B — Cloud migration to AWS: redeploy the LMS onto a highly available AWS architecture "
        "(EC2 Auto Scaling across two AZs, RDS Multi-AZ for MySQL, ALB), preserving the OS and application stack.",
        "A standalone managed-hosting option was considered but not carried forward separately: its "
        "benefits (no in-house server, managed resilience) are a subset of Option B's, while AWS adds the "
        "elasticity and recovery posture the requirements demand.",
    ])
    h2("6.3 Evaluation method")
    add_body_paragraph(doc, "Each option is evaluated with a five-year cost-benefit analysis (§7), a quantified "
              "avoided-downtime benefit (§7.4), an intangibles comparison (§8.1), and a risk register "
              "(§8.2). This combination is appropriate because the decision turns on both cost and "
              "non-cost factors — resilience, capacity, and strategic alignment — that a cost comparison "
              "alone would not capture.")
    h2("6.4 Initial impact and difficulty assessment")
    add_data_table(doc, ["", "Option A — In-house", "Option B — Cloud (AWS)"],
           [["Strategic impact", "Refreshes hardware but leaves the 99.9% target out of reach",
             "Directly delivers the strategy's lead initiative and availability target"],
            ["Implementation difficulty", "Low — familiar in-house refresh",
             "Moderate — new platform; mitigated by MTS and a phased plan"],
            ["Headline pros", "Cheap upfront; uses existing skills",
             "Meets availability/recovery/scaling targets; reduces in-house burden"],
            ["Headline cons", "Single point of failure persists; fixed capacity",
             "Higher Year-1 outlay; cloud upskilling; vendor dependency"]],
           widths=[4.0, 6.25, 6.25])

    h1("7. Cost-Benefit Analysis")
    add_uoc_evidence_tag(doc, "[ICTICT517 PC 2.1] · [ICTICT517 KE 3] · [ICTICT517 FS Numeracy] · [ICTCLD401 KE 4] · [ICTCLD502 KE 3] · [ICTCLD401 KE 1]")
    add_body_paragraph(doc, "A five-year CBA comparing the two options (AUD ex GST). Detailed line items are in "
              "Appendix 1; AWS figures are indicative, sized per §6.1 and backed by a Pricing Calculator "
              "export. A 3% annual inflation factor is applied to recurring costs from Year 2.")
    h2("7.1 Assumptions")
    add_data_table(doc, ["Assumption", "Value", "Used as-is?"],
           [["LMS user population", "800 students + 60 staff", "Yes"],
            ["Annual student growth", "+15% per year", "Yes"],
            ["Current / target availability", "99.2% / 99.9%", "Yes"],
            ["ICT FTE fully-loaded cost", "$115,000 / year (≈$63.19/hr)", "Yes"],
            ["Cost of downtime (teaching hours)", "$750 / hour", "Yes"],
            ["Inflation applied to recurring", "3% per year from Year 2", "Adjusted — see note"],
            ["AWS region", "ap-southeast-2 (Sydney)", "Yes (data residency)"]],
           widths=[6.5, 5.5, 4.5])
    yr = ["", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5", "5-year"]
    h2("7.2 Option A — In-house renewal (5-year summary)")
    add_data_table(doc, yr,
           [["One-off capital", "$51,000", "—", "—", "—", "—", "$51,000"],
            ["Recurring", "$74,650", "$76,890", "$79,196", "$81,572", "$84,019", "$396,327"],
            ["Annual total", "$125,650", "$76,890", "$79,196", "$81,572", "$84,019", "$447,327"]])
    h2("7.3 Option B — Cloud migration to AWS (5-year summary)")
    add_data_table(doc, yr,
           [["One-off project", "$108,500", "—", "—", "—", "—", "$108,500"],
            ["AWS direct", "$27,250", "$28,068", "$28,910", "$29,777", "$30,670", "$144,675"],
            ["Staff + external", "(in project)", "$45,000", "$46,350", "$47,741", "$49,173", "$188,264"],
            ["Annual total", "$135,750", "$73,068", "$75,260", "$77,518", "$79,843", "$441,439"]])
    h2("7.4 Avoided-downtime benefit (Option B)")
    add_data_table(doc, ["", "Calculation", "Value"],
           [["Avoided unavailability / year", "(0.008 − 0.001) × 8,760 hr", "61.3 hr"],
            ["Proportion during teaching hours", "assumption", "30%"],
            ["Avoided downtime cost / year", "61.3 × 30% × $750", "$13,800"],
            ["5-year avoided-downtime benefit", "× 5", "$69,000"]],
           widths=[6.0, 6.5, 4.0])
    h2("7.5 Comparison summary")
    add_data_table(doc, ["", "Option A", "Option B", "Delta (B − A)"],
           [["One-off (Year 1)", "$51,000", "$108,500", "+$57,500"],
            ["Recurring (5-year)", "$396,327", "$332,939", "−$63,388"],
            ["5-year direct cost", "$447,327", "$441,439", "−$5,888"],
            ["Avoided-downtime benefit", "—", "$69,000", "−$69,000"],
            ["Net 5-year position", "$447,327", "$372,439", "−$74,888"]],
           widths=[5.0, 3.8, 3.8, 3.8])
    add_body_paragraph(doc, "On direct cost the options are within ~1.3% over five years. Option B carries a higher "
              "Year-1 outlay but lower recurring cost, and once the avoided-downtime benefit is included "
              "it is ~$74,900 ahead on net position — before the unpriced strategic and capacity benefits.")
    h2("7.6 Sensitivity analysis")
    add_body_paragraph(doc, "Sensitivity 1 — teaching-hours downtime proportion 30% → 50%: the avoided-downtime "
              "benefit rises from ~$69,000 to ~$115,000 over five years, widening Option B's net "
              "advantage. The recommendation is reinforced.")
    add_body_paragraph(doc, "Sensitivity 2 — MTS migration labour $60,000 → $80,000: Option B's five-year direct cost "
              "rises to ~$461,400, slightly above Option A; but after the avoided-downtime benefit its net "
              "position (~$392,400) remains well below Option A's $447,327. The recommendation holds.")

    h1("8. Risk and Impact Assessment")
    add_uoc_evidence_tag(doc, "[ICTICT517 PC 2.1] · [ICTICT517 PC 2.2] · [ICTICT517 PE 5]")
    h2("8.1 Intangibles comparison")
    add_data_table(doc, ["Factor", "Option A — In-house", "Option B — Cloud (AWS)"],
           [["99.9% availability target", "Not credibly achievable on a single server",
             "Achievable via Multi-AZ deployment"],
            ["Recovery (RTO ≤ 4h / RPO ≤ 1h)", "Improves but stays single-site",
             "Met via RDS Multi-AZ + automated backups"],
            ["Capacity for 15% growth", "Needs over-provisioning / re-refresh", "Scales on demand"],
            ["Assessment-week peaks (~3×)", "Must size for peak year-round", "Auto-scales for peaks"],
            ["Staff capability impact", "Uses existing skills", "Requires upskilling (MTS-supported)"],
            ["Vendor lock-in", "None new", "AWS dependency; standard stack aids portability"],
            ["Sustainability", "Runs own server + room", "Shared, efficient cloud infrastructure"],
            ["DR (whole-campus loss)", "Slow tape restore offsite", "Cross-Region backup copy"]],
           widths=[4.5, 5.75, 6.0])
    h2("8.2 Risk register (recommended option — cloud)")
    add_data_table(doc, ["Risk", "Likelihood", "Impact", "Mitigation"],
           [["Cloud cost overrun / bill shock", "Medium", "Medium",
             "Budgets + CloudWatch alarms; tagging; Savings Plans for baseline"],
            ["Staff cloud-skills gap", "High", "Medium",
             "MTS support contract; AWS training; documented runbooks"],
            ["Data loss/corruption during migration", "Low", "High",
             "Parallel running; verified backups; tested rollback"],
            ["Cutover disrupts teaching", "Medium", "High",
             "Schedule outside assessment windows; maintenance window; comms plan"],
            ["Vendor dependency", "Low", "Medium",
             "Standard Windows/MySQL stack; infrastructure-as-code; documented exit plan"]],
           widths=[5.0, 2.3, 2.0, 6.9])

    h1("9. Recommendation")
    add_uoc_evidence_tag(doc, "[ICTICT517 PC 2.3] · [ICTICT517 PE 4] · contributes to [ICTICT517 PC 3.1]")
    add_body_paragraph(doc, "Recommended option: Option B — Cloud migration to AWS. Over five years the two options "
              "are effectively cost-neutral on direct spend, but only the cloud option meets the 99.9% "
              "availability target and the RTO ≤ 4h / RPO ≤ 1h recovery objectives, scales for 15% growth "
              "and the assessment-window peaks, and materially improves disaster recovery. The "
              "avoided-downtime benefit puts it ahead on net cost as well. The main risks — upskilling, "
              "cost control, and a safe cutover — are manageable and addressed in §8.2 and §10. This "
              "recommendation drives the prioritisation and action plan below.")

    h1("10. Action Plan")
    add_uoc_evidence_tag(doc, "[ICTICT517 PC 3.1] · [ICTICT517 PC 3.2] · [ICTICT517 PE 6] · [ICTICT517 KE 1]")
    h2("10.1 Prioritised changes")
    add_data_table(doc, ["#", "Change", "Priority rationale"],
           [["1", "Stand up the AWS foundation (network, IAM, guardrails)",
             "Everything else depends on a secure, governed landing zone"],
            ["2", "Build the highly available LMS infrastructure (EC2 ASG Multi-AZ, RDS Multi-AZ, ALB)",
             "Delivers the availability/recovery/scaling targets before any data moves"],
            ["3", "Deploy DOODLE and migrate the MySQL data; run in parallel",
             "Validate the new platform against the live system before commitment"],
            ["4", "Cut over and decommission the on-premises server",
             "Final, highest-risk step; only after parallel-run acceptance"]],
           widths=[1.2, 7.3, 7.5])
    h2("10.2 Implementation schedule")
    add_data_table(doc, ["Phase", "Activities", "Start", "Duration", "Dependencies", "Owner"],
           [["1", "AWS foundation + design review", "Wk 1", "2 wks", "—", "MTS"],
            ["2", "HA infrastructure build", "Wk 3", "3 wks", "Phase 1", "MTS"],
            ["3", "App deploy + data migration + parallel run", "Wk 6", "3 wks", "Phase 2", "MTS + YAT ICT"],
            ["4", "Cutover + decommission", "Wk 9", "1 wk", "Phase 3 acceptance", "YAT ICT + MTS"]],
           widths=[1.5, 5.2, 2.0, 2.0, 3.0, 2.6])
    add_body_paragraph(doc, "Scheduling avoids month-end and the two-week assessment windows; the cutover runs in an "
              "approved maintenance window with a tested rollback.")
    h2("10.3 Standards, targets, and success metrics")
    add_data_table(doc, ["Aspect", "Standard / target", "How measured"],
           [["Availability", "99.9%", "CloudWatch uptime over rolling 12 months"],
            ["Recovery", "RPO ≤ 1h / RTO ≤ 4h", "DR test results; backup frequency"],
            ["Data residency", "Australia (ap-southeast-2)", "AWS Config region check"],
            ["Cost envelope (Year 1)", "≈ $135,750", "Actuals vs budget in Cost Explorer"]],
           widths=[4.5, 5.5, 5.5])
    h2("10.4 Implementation methods")
    add_body_paragraph(doc, "Each phase applies an AWS Well-Architected review at its design milestone, infrastructure-"
              "as-code (CloudFormation) for repeatable builds, and YAT's Change Management Procedure for "
              "all production changes.")
    h2("10.5 Alignment with the change management procedure")
    add_body_paragraph(doc, "Phases 2–4 raise change requests; the Phase 4 cutover is a high-risk change requiring "
              "Change Advisory Board endorsement and ICT-Manager approval, with the budget and option "
              "decision approved by the board via this Business Case.")

    h1("11. Next Steps and Decision Requested")
    add_uoc_evidence_tag(doc, "[ICTICT517 PC 3.3]")
    add_body_paragraph(doc, "We ask the board today to: (1) approve Option B — cloud migration to AWS — as the "
              "recommended approach; (2) approve the phased action plan in §10 as the implementation "
              "roadmap; and (3) authorise the Year-1 budget envelope (≈$135,750). The high-risk cutover "
              "change request is deferred to the appropriate phase gate, per the change management "
              "procedure.")

    h1("Sign-off")
    add_uoc_evidence_tag(doc, "[ICTICT517 PC 3.3]")
    add_data_table(doc, ["Role", "Name", "Date", "Signature"],
           [["Prepared by", "MTS Consultant", "", ""],
            ["Reviewed by", "Pat Lin (MTS Senior Consultant)", "", ""],
            ["Approved by", "Sam Walker (YAT ICT Manager)", "", ""]],
           widths=[5.0, 5.5, 2.5, 3.0])

    # ---- APPENDICES ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1("Appendix 1 — Cost-Benefit Analysis: Detailed Line Items")
    add_uoc_evidence_tag(doc, "[ICTICT517 FS Numeracy] · [ICTICT517 KE 3] · [ICTCLD401 KE 4] · [ICTCLD502 KE 3]")
    h2("A1.1 Option A — In-house renewal")
    add_body_paragraph(doc, "One-off capital (Year 1):")
    add_data_table(doc, ["Item", "Cost"],
           [["Replacement LMS server", "$25,000"],
            ["Backup tape library refresh + drive", "$8,000"],
            ["UPS upgrade", "$3,000"],
            ["One-off migration labour (in-house)", "$15,000"],
            ["Total Year-1 capital", "$51,000"]],
           widths=[11.5, 4.0])
    add_body_paragraph(doc, "Recurring operational (per year):")
    add_data_table(doc, ["Category", "Item", "Annual cost"],
           [["Software licensing", "Windows Server Standard", "$1,500"],
            ["", "Antivirus / EDR", "$300"],
            ["", "Monitoring / management tooling", "$2,000"],
            ["Power & facilities", "Electricity + cooling", "$1,200"],
            ["", "Server-room rent allocation", "$5,000"],
            ["", "UPS battery (amortised)", "$500"],
            ["Backup", "Tape media", "$1,500"],
            ["", "Offsite tape storage", "$2,400"],
            ["", "Backup software maintenance", "$1,500"],
            ["Staff (YAT ICT)", "LMS admin (0.20 FTE)", "$23,000"],
            ["", "Incident response (0.05 FTE)", "$5,750"],
            ["External support", "MTS application support", "$30,000"],
            ["Total recurring per year", "", "$74,650"]],
           widths=[4.5, 7.0, 4.0])
    h2("A1.2 Option B — Cloud migration to AWS")
    add_body_paragraph(doc, "One-off project costs (Year 1):")
    add_data_table(doc, ["Item", "Cost"],
           [["MTS migration labour (design + build)", "$60,000"],
            ["YAT ICT staff time during migration (oversight, deploy, data, cutover, change mgmt)", "$38,000"],
            ["Parallel running of on-prem + cloud during cutover", "$7,000"],
            ["Decommissioning + secure data destruction", "$3,500"],
            ["Total Year-1 project", "$108,500"]],
           widths=[11.5, 4.0])
    add_body_paragraph(doc, "AWS direct recurring (per year, indicative — ap-southeast-2):")
    add_data_table(doc, ["Category", "Service / item", "Annual cost"],
           [["Compute", "EC2 baseline (2× Windows m6i.xlarge, Multi-AZ, Savings Plan)", "$8,400"],
            ["", "EC2 auto-scaling (assessment-window peaks)", "$1,500"],
            ["Database", "RDS for MySQL — Multi-AZ (1-yr RI)", "$6,800"],
            ["", "RDS storage + automated backups", "$1,400"],
            ["Network", "Application Load Balancer", "$850"],
            ["", "NAT Gateway (×2, AZ-redundant)", "$1,300"],
            ["", "Data transfer out", "$900"],
            ["Storage", "EBS volumes (web tier)", "$700"],
            ["", "S3 (attachments + snapshots)", "$600"],
            ["", "AWS Backup (incl. cross-Region copy)", "$1,200"],
            ["Monitoring", "CloudWatch (metrics + logs + alarms)", "$700"],
            ["Support", "AWS Business Support (≤1h sev-1)", "$2,900"],
            ["AWS direct per year", "", "$27,250"]],
           widths=[3.5, 8.0, 4.0])
    add_body_paragraph(doc, "Staff + ongoing external support (per year, Years 2–5):")
    add_data_table(doc, ["Item", "Annual cost"],
           [["YAT ICT staff time (admin/monitoring, reduced — 0.10 FTE)", "$11,500"],
            ["YAT ICT staff time (incident response — 0.03 FTE)", "$3,450"],
            ["MTS ongoing support (application + cloud infrastructure)", "$30,000"],
            ["Staff + external per year", "$44,950"]],
           widths=[11.5, 4.0])

    h1("Appendix 2 — Knowledge Evidence")
    add_body_paragraph(doc, "Underlying knowledge as applied to the work presented above.")
    h2("A2.1 Selection-style questions")
    add_data_table(doc, ["Question", "Response"],
           [["Service model of each AWS component (IaaS/PaaS/SaaS) and why [ICTCLD401 KE 3]",
             "EC2 = IaaS (full OS control is required to host DOODLE on Windows Server 2016). "
             "RDS, ALB, S3, AWS Backup, CloudWatch = PaaS-style managed services (chosen to offload "
             "operations from a thin in-house team). Office 365 (email) = SaaS."],
            ["Deployment model and why [ICTCLD401 KE 11]",
             "Public cloud. YAT has no regulatory reason to run private infrastructure, already uses "
             "SaaS (O365), and gains managed resilience and elasticity it cannot achieve on-prem; data "
             "residency is met by an Australian region."],
            ["Industry standards that informed the approach (≥3) [ICTCLD401 KE 1] · [ICTCLD502 KE 1]",
             "NIST SP 800-145 (service-model definitions); AWS Well-Architected (design + the Reliability "
             "pillar behind the HA design); ISO/IEC 27017 and the ACSC Essential Eight (cloud security "
             "controls); the Privacy Act / APP 8 (data-residency region choice)."],
            ["Industry-standard products in the solution and their role [ICTCLD401 KE 2] · [ICTCLD502 KE 2]",
             "EC2 (compute hosting DOODLE), RDS for MySQL (managed database), Application Load Balancer "
             "(traffic distribution across AZs), S3 (object storage for attachments/snapshots), "
             "CloudWatch (monitoring/alarms)."],
            ["Cloud cost model(s) applied and why [ICTCLD401 KE 4] · [ICTCLD502 KE 3]",
             "Reserved Instances / Savings Plans for the always-on baseline (compute, RDS); on-demand "
             "auto-scaling for the assessment-window peaks; pay-as-you-go for storage and data transfer. "
             "The model follows the workload's demand shape, so cost tracks usage."]],
           widths=[6.5, 9.0])
    h2("A2.2 Demonstration-style questions")
    add_data_table(doc, ["Question", "Response"],
           [["How does §10 demonstrate the key sections of an action plan? [ICTICT517 KE 1]",
             "§10.1 prioritised changes; §10.2 schedule with durations/dependencies/owners; §10.3 "
             "standards, targets and success metrics; §10.4 methods; §10.5 change-management alignment."],
            ["Evaluation and planning methods applied [ICTICT517 KE 2]",
             "Five-year CBA (§7), avoided-downtime quantification (§7.4), sensitivity analysis (§7.6), "
             "intangibles comparison (§8.1), and a risk register (§8.2)."],
            ["How competing products were evaluated [ICTICT517 KE 3]",
             "Option A vs Option B compared on cost (§7), strategic impact and difficulty (§6.4), "
             "intangibles (§8.1) and risk (§8.2)."],
            ["Emerging trends that informed the recommendation [ICTICT517 KE 4]",
             "Sector shift to managed cloud and OPEX-over-CAPEX, and elastic capacity as a baseline "
             "expectation — material to choosing a cloud operating model over a hardware refresh."]],
           widths=[6.5, 9.0])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "scenario/exemplars/internal-document-exemplar-business-case.docx"
    build(out)
