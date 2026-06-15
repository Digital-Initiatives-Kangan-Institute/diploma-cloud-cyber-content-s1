#!/usr/bin/env python3
"""Build the AT2 Deployment Report EXEMPLAR (.docx) — assessor marking reference.

A worked model answer for the YAT LMS Cloud Migration *foundation build* (S1-CL1 AT2,
ICTCLD401), built on the generic Deployment Report superset template. Demonstrates the
"Not applicable — [reason]" convention by marking the HA-only sections deferred to the
AT3 hardening phase. Evidence is described, not captured: "[SCREENSHOT — should show ...]".

Assessor-facing: retains UoC `Evidences:` tags, §8 Knowledge Evidence responses, and
reflections (the assessment layers a real org template omits). Reuses the docx brand
helpers (build_bc_template) and the prose/table helpers (build_s1_cl1_at1_bc_exemplar).

Usage:  python scripts/build_at2_deployment_exemplar.py [output.docx]
Default: S1-CL1-Cloud-Design-Build/assessments/AT2/AT2-exemplar-deployment-report.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_bullet_list  # noqa: E402
import build_bc_template as bc   # noqa: E402
import build_s1_cl1_at1_bc_exemplar as ex   # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def na(doc, reason):
    p = doc.add_paragraph()
    r = p.add_run(f"Not applicable — {reason}")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    p.paragraph_format.space_after = Pt(6)
    return p


def ev(doc, kind, desc):
    """Described-evidence placeholder, e.g. [SCREENSHOT — should show ...]."""
    p = doc.add_paragraph()
    tag = p.add_run(f"[{kind}] ")
    tag.bold = True
    tag.font.size = Pt(9.5)
    tag.font.color.rgb = RGBColor.from_string(bc.TEAL)
    d = p.add_run("— " + desc)
    d.font.size = Pt(9.5)
    d.italic = True
    d.font.color.rgb = RGBColor.from_string(bc.GREY)
    p.paragraph_format.space_after = Pt(6)
    return p


def build(path):
    doc = Document()
    bc.configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    bc.build_header_footer(sec)

    # ---- COVER ----
    bc.wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(bc.ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(bc.GREY)
    bc.paragraph_bottom_rule(doc.add_paragraph(), bc.TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Deployment Report")
    sub = doc.add_paragraph().add_run("YAT LMS Cloud Migration — Foundation Build Phase")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    note = doc.add_paragraph().add_run("Assessor exemplar — internal marking reference (not for distribution to students)")
    note.italic = True; note.font.size = Pt(10); note.font.color.rgb = RGBColor.from_string(bc.GREY)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT LMS Cloud Migration — Foundation Build"),
        ("Engagement reference", "YAT-LMS-MIG-2026"),
        ("Report version", "v1.0"),
        ("Prepared by", "MTS Consultant"),
        ("Date submitted", "[DD/MM/YYYY]"),
        ("Submitted to", "Sam Walker (YAT ICT Manager) · Pat Lin (MTS Senior Consultant)"),
        ("Related documents", "YAT LMS Cloud Architecture — Baseline Design v1.0; YAT LMS Migration Business Case"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        bc.set_cell_borders(cells[0]); bc.set_cell_borders(cells[1]); bc.shade_cell(cells[0], bc.CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    bc.add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h2 = lambda t: doc.add_paragraph(t, style="Heading 2")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Executive Summary")
    ex.uoc(doc, "[ICTCLD401 PC 4.1] — partial (document and communicate the work)")
    ex.para(doc, "This report documents the foundation build phase of the YAT LMS cloud migration: the "
                 "baseline AWS infrastructure for the LMS, deployed in the Sydney region (ap-southeast-2), "
                 "availability zone ap-southeast-2a. The build delivered identity and access management, the "
                 "network (a VPC with public, private-app and private-data subnets), the compute tier (an EC2 "
                 "Auto Scaling Group behind an Application Load Balancer), a managed MySQL database on RDS, "
                 "S3 and EBS storage, a three-tier security-group model, and a baseline of CloudWatch "
                 "monitoring. All connectivity, autoscaling and smoke tests passed and the infrastructure is "
                 "ready for YAT IT to deploy the DOODLE application onto.")
    ex.para(doc, "This phase is deliberately single-AZ: high availability (Multi-AZ database, cross-AZ "
                 "compute, failover testing, DR runbook and cross-Region backup) is out of scope here and is "
                 "the subject of the next phase. The engagement is ready to proceed to the HA design phase.")

    h1("2. Engagement Context")
    ex.uoc(doc, "[ICTCLD502 AC 3, AC 5] (scenario data sources used)")
    ex.para(doc, "This build implements the board-approved direction from the YAT LMS Migration Business "
                 "Case and the supplied YAT LMS Cloud Architecture — Baseline Design (v1.0, approved by Pat "
                 "Lin and Sam Walker). As the MTS consultant I implemented that architecture in the AWS "
                 "Academy environment. Per the MTS scope, the LMS application install, the MySQL data "
                 "migration, and the cutover remain YAT IT's responsibility; HA hardening is the next MTS "
                 "phase, covered in a separate report.")

    h1("3. Scope of Deployment")
    ex.para(doc, "In scope of this report: the baseline cloud foundation — IAM, network, compute, database, "
                 "autoscaling, storage, and a monitoring baseline. Out of scope and deferred to the HA "
                 "hardening phase: Multi-AZ database, cross-AZ resilience, failover and resize simulation, "
                 "the DR runbook, cross-Region backup copies, and HA-tuned monitoring.")
    h3("3.1 Maintenance-window context")
    na(doc, "this was a greenfield build in an empty account, not a change to a running system, so no "
            "maintenance window applied.")

    h1("4. Build / Change Narrative")
    ex.uoc(doc, "[ICTCLD401 PC 1.1–1.7, 2.1–2.5, 3.1] · [ICTCLD502 PC 1.3]")
    h3("4.1 Identity and access management (IAM)")
    ex.para(doc, "I created four IAM groups reflecting the engagement's job functions — YAT-ICT-Admins "
                 "(full administration), MTS-Consultants (build/operate within a permission boundary), "
                 "Application-Service (the EC2 instance role), and Read-Only-Auditors. MFA is enforced on all "
                 "human users via an IAM policy that denies actions without MFA. The EC2 instances assume an "
                 "instance profile granting only the S3 and CloudWatch access the application needs.")
    ev(doc, "SCREENSHOT", "should show the IAM Groups list with the four groups and member counts, and an "
                          "IAM user with an MFA device enabled; region ap-southeast-2 visible.")
    h3("4.2 Network topology")
    ex.para(doc, "I deployed a VPC (10.0.0.0/16) with three subnets in ap-southeast-2a: public-web-a "
                 "(10.0.1.0/24) for the ALB and NAT Gateway, private-app-a (10.0.2.0/24) for the EC2 "
                 "instances, and private-data-a (10.0.3.0/24) for RDS. An Internet Gateway serves the public "
                 "subnet; a NAT Gateway gives the private app tier outbound internet (e.g. Windows Update) "
                 "without inbound exposure. Route tables route the public subnet to the IGW and the private "
                 "subnets to the NAT.")
    ev(doc, "SCREENSHOT", "should show the VPC subnets list with the three subnets, their CIDRs and AZ, and "
                          "the route tables showing 0.0.0.0/0 → IGW (public) and → NAT (private-app).")
    h3("4.3 Compute (EC2 + Auto Scaling)")
    ex.para(doc, "A launch template defines the LMS application instance (Windows Server 2016, m6i.large). "
                 "The Auto Scaling Group runs min 1 / desired 1 / max 3 in this single-AZ baseline, with a "
                 "target-tracking policy that scales out when average CPU exceeds 65%.")
    ev(doc, "SCREENSHOT", "should show the EC2 Auto Scaling Group with its launch template, min/desired/max, "
                          "and the scaling policy; and the running instance(s) with tags.")
    h3("4.4 Load balancing (ALB)")
    ex.para(doc, "An internet-facing Application Load Balancer in the public subnet forwards HTTPS to the app "
                 "tier. The target group health-checks the instances on an HTTP path; the HTTPS listener uses "
                 "an ACM certificate.")
    ev(doc, "SCREENSHOT", "should show the ALB target group with the EC2 instance(s) reporting Healthy.")
    h3("4.5 Database (RDS)")
    ex.para(doc, "The LMS database runs on RDS for MySQL 8.0.35, instance class db.m6i.large, 150 GB gp3 "
                 "storage, encryption at rest enabled, and 7-day automated backup retention. This phase is "
                 "Single-AZ; Multi-AZ is enabled in the HA phase.")
    ev(doc, "SCREENSHOT", "should show the RDS database in 'available' state, engine MySQL 8.0.35, "
                          "Single-AZ, storage encrypted.")
    h3("4.6 Storage (EBS + S3)")
    ex.para(doc, "The EC2 instance has an encrypted gp3 root volume (50 GB) and a gp3 data volume (100 GB). "
                 "Two S3 buckets hold LMS attachments and backups; both have Block Public Access enabled and "
                 "default encryption on.")
    ev(doc, "SCREENSHOT", "should show the S3 buckets list with Block Public Access enabled, and the EBS "
                          "volumes attached to the instance.")
    h3("4.7 Security (security groups + encryption)")
    ex.para(doc, "A three-tier security-group model: sg-alb allows 443 from the internet; sg-app allows the "
                 "application port only from sg-alb; sg-db allows 3306 only from sg-app. Traffic is encrypted "
                 "in transit (HTTPS at the ALB, TLS to RDS) and at rest (EBS, RDS, S3).")
    ev(doc, "SCREENSHOT", "should show the three security groups with inbound rules expanded, demonstrating "
                          "the tier-to-tier restriction.")
    h3("4.8 Monitoring (baseline)")
    ex.para(doc, "Baseline CloudWatch alarms cover EC2 CPU, the ASG, RDS CPU / connections / free storage, "
                 "and ALB 5xx and unhealthy-host counts, notifying an SNS topic. HA-tuned monitoring is added "
                 "in the next phase.")
    ev(doc, "SCREENSHOT", "should show the CloudWatch Alarms list with the baseline alarms configured.")
    h3("4.9 Cross-Region backup / replication")
    na(doc, "cross-Region resilience is deferred to the HA hardening phase (AT3).")

    h1("5. Configuration Decisions")
    ex.uoc(doc, "[ICTCLD401 PC 1.3] (justified service selection)")
    ex.etable(doc, ["#", "Decision", "Choice", "Rationale"],
              [["C1", "EC2 instance type", "m6i.large (2 vCPU, 8 GB)",
                "Comfortably serves 200–300 typical concurrent users; the ASG absorbs assessment-window peaks."],
               ["C2", "RDS instance class", "db.m6i.large",
                "Sufficient for the ~68 GB MySQL workload and connection load, with headroom for growth."],
               ["C3", "Storage sizing", "RDS 150 GB; EBS data 100 GB",
                "DB ~68 GB + ~25 GB/yr growth ≈ 3 years' headroom; bulk attachments go to S3, not the DB."],
               ["C4", "ASG scaling threshold", "Scale out at 65% CPU",
                "Below the point where response time degrades for this workload."],
               ["C5", "MTS-Consultants permission boundary", "Limited to ap-southeast-2 + LMS-tagged resources",
                "Least privilege; no IAM or billing actions."],
               ["C6", "Admin access to instances", "SSM Session Manager (no bastion)",
                "Removes a public RDP jump host to harden and audit; access is logged."],
               ["C7", "MySQL engine version", "MySQL 8.0.35",
                "DOODLE-supported and a current RDS-supported minor version."],
               ["C8", "DNS + certificate", "Route 53 alias → ALB; ACM public cert",
                "lms.yat.edu.au resolves to the ALB; HTTPS terminated with a managed certificate."]],
              widths=[1.0, 4.2, 4.0, 6.3])

    h1("6. Testing, Simulation and Validation")
    ex.uoc(doc, "[ICTCLD401 PC 2.6, 3.2]")
    h3("6.1 Connectivity tests")
    ex.etable(doc, ["Test", "Outcome", "Notes"],
              [["ALB → EC2 health check", "Pass", "Target healthy"],
               ["EC2 → RDS connection (private)", "Pass", "Connected over private-data subnet"],
               ["EC2 → internet via NAT", "Pass", "Windows Update reachable"],
               ["RDS not publicly reachable (negative)", "Pass", "Connection from outside the VPC timed out"]],
              widths=[7.0, 3.0, 5.5])
    ev(doc, "TEST EVIDENCE", "should show a terminal capture of the EC2 → RDS mysql connection succeeding, "
                            "and an external connection attempt to RDS timing out.")
    h3("6.2 Autoscaling test")
    ex.para(doc, "Generating load against the ALB pushed average CPU above 65%; the ASG launched a second "
                 "instance, which became healthy in the target group within ~4 minutes. When load dropped, "
                 "the ASG scaled back to one instance.")
    ev(doc, "SCREENSHOT", "should show the ASG activity history with the scale-out event, plus a CloudWatch "
                          "CPU graph showing the spike and the new instance entering service.")
    h3("6.3 Database connectivity and basic operations")
    ex.para(doc, "From the EC2 instance, a mysql client connected over the private network; SELECT VERSION() "
                 "confirmed 8.0.35; the connection used TLS.")
    ev(doc, "TEST EVIDENCE", "should show the SELECT VERSION() output and the TLS-enabled connection.")
    h3("6.4 Infrastructure end-to-end smoke test")
    ex.para(doc, "A placeholder page served from the instances returned HTTP 200 via the ALB's DNS name, "
                 "confirming the request path (internet → ALB → EC2) works and the instance can reach RDS over "
                 "the private network. The DOODLE application itself is installed by YAT IT.")
    ev(doc, "TEST EVIDENCE", "should show a browser (or curl -I) reaching the placeholder page via the ALB "
                            "DNS name with HTTP 200/302.")
    h3("6.5 Failure simulation")
    na(doc, "no Multi-AZ resilience exists yet to fail over to; failure simulation is performed in the HA "
            "hardening phase (AT3).")
    h3("6.6 Resize simulation")
    na(doc, "deferred to the HA hardening phase (AT3).")
    h3("6.7 Availability measurement")
    na(doc, "formal availability measurement against the 99.9% target is part of the HA phase; this baseline "
            "build was verified functionally.")
    h3("6.8 Simulation findings vs the design")
    na(doc, "deferred to the HA hardening phase (AT3).")
    h3("6.9 Adjustments made per simulation outcomes")
    na(doc, "deferred to the HA hardening phase (AT3).")

    h1("7. Operational Handover")
    ex.uoc(doc, "[ICTCLD401 PC 4.3] (save documentation per organisational policies)")
    h3("7.1 Access")
    ex.para(doc, "YAT-ICT-Admins retain full administration; MTS-Consultants remain within the permission "
                 "boundary; MFA stays enforced; instance access is via SSM Session Manager only.")
    h3("7.2 Runbook references")
    add_bullet_list(doc, [
        "The supplied Baseline Design document (operational reference).",
        "The naming and tagging conventions, so YAT ICT can identify resources.",
        "Backup arrangements: RDS automated backups (7-day retention) and the EBS snapshot policy.",
        "The CloudWatch baseline alarms list and the SNS notification topic.",
    ])
    h3("7.3 Known limitations and what's next")
    ex.para(doc, "This foundation is single-AZ and is not highly available: the RDS instance and each AZ are "
                 "single points of failure. The HA hardening phase will enable RDS Multi-AZ, extend the ASG "
                 "and ALB across two AZs, add HA-tuned monitoring and cross-Region backup, and establish a DR "
                 "runbook with failure-simulation testing.")
    h3("7.4 Documentation filing")
    ex.etable(doc, ["Item", "Filed in", "Reference"],
              [["This Deployment Report (v1.0)", "YAT ICT shared documentation", "[ref]"],
               ["Configuration exports (Appendix B)", "YAT ICT shared documentation", "[ref]"],
               ["Test evidence (Appendix C)", "YAT ICT shared documentation", "[ref]"]],
              widths=[6.5, 5.0, 4.0])
    h3("7.5 Feedback record")
    na(doc, "the formal feedback loop and end-of-engagement sign-off are captured at the close of the HA "
            "phase; this handover is an interim phase handover.")
    h3("7.6 Sign-off")
    ex.etable(doc, ["Role", "Name", "Date", "Acceptance"],
              [["Prepared by", "MTS Consultant", "[date]", "Submitted"],
               ["Accepted by", "Sam Walker (YAT ICT Manager)", "[date]",
                "Foundation build accepted; ready for HA design"]],
              widths=[4.5, 5.0, 2.5, 3.5])

    h1("8. Knowledge Evidence Responses")
    ex.uoc(doc, "[ICTCLD401 KE 5, 6, 7, 8, 9, 10]")
    h3("Q1 — Virtual machine, networking and scaling features")
    ex.para(doc, "EC2 provides the compute capacity that hosts the LMS application; I chose m6i.large to suit "
                 "the typical concurrent load while keeping cost down, leaning on the ASG for peaks. The Auto "
                 "Scaling Group provides automatic scaling — it adds instances when CPU passes 65% (the "
                 "assessment-window peak pattern) and removes them when load falls, so capacity tracks demand. "
                 "The Application Load Balancer distributes traffic across healthy instances and, via its "
                 "health check, stops routing to an unhealthy one — which is also what lets the ASG replace an "
                 "instance without downtime to users.")
    h3("Q2 — Managed services, storage options, and scaling model")
    ex.para(doc, "RDS is preferred over self-hosting MySQL on EC2 because AWS manages patching, backups and "
                 "(later) Multi-AZ failover — valuable given YAT ICT's thin cloud experience. EBS and S3 are "
                 "used together because they solve different problems: EBS is block storage attached to the "
                 "instance for the OS and working data, while S3 is durable, cheap object storage for the bulk "
                 "LMS attachments and backups that don't need to sit on a volume. The ASG scales horizontally; "
                 "vertical scaling (resizing the instance) was viable but would mean downtime on each resize "
                 "and a hard ceiling, whereas horizontal scaling matches the spiky assessment-window load "
                 "without interruption.")
    h3("Q3 — Shared security responsibility")
    ex.para(doc, "Two responsibilities that remain YAT's in this environment: configuring the security groups "
                 "and IAM permissions correctly (AWS provides the controls; using them well is the customer's "
                 "job), and managing the guest OS — patching Windows Server and the application stack on the "
                 "EC2 instances. One responsibility that shifted to AWS by moving to RDS: the database host "
                 "OS patching and the underlying hardware, which YAT previously owned on the on-prem server.")
    h3("Q4 — User access protocols and organisational hierarchy")
    ex.para(doc, "The MTS-Consultants group has build-and-operate permissions within a permission boundary "
                 "limiting it to ap-southeast-2 and LMS-tagged resources, with no IAM or billing rights. It "
                 "serves the external consultant function. Its permissions must differ from YAT-ICT-Admins "
                 "(who hold full administration as the system owner) so that an external party can do the "
                 "delivery work without being able to alter identity, billing, or resources outside the "
                 "engagement.")
    h3("Q5 — Security policies and network traffic limits")
    ex.para(doc, "sg-db allows inbound 3306 only from sg-app, and no outbound to the internet. Traffic is "
                 "restricted this way so the database is reachable only from the application tier over the "
                 "private network — never from the internet or even the web tier directly. If the restriction "
                 "were removed (e.g. 3306 open to the VPC or the internet), the database would be exposed to "
                 "any compromised host or external attacker, putting the student personal information and "
                 "records at risk.")
    h3("Q6 — Role of DNS in the deployment")
    ex.para(doc, "Two points where DNS resolution happens: end users browse to lms.yat.edu.au, which Route 53 "
                 "resolves to the ALB's address — if misconfigured, users couldn't reach the LMS at all. And "
                 "the EC2 instances resolve the RDS endpoint name to its current IP — if that failed (or, in "
                 "the HA phase, didn't update on failover) the application couldn't reach its database.")

    # ---- APPENDICES ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1("Appendix A — Build evidence (screenshots)")
    ex.para(doc, "Evidence captures are described below in lieu of live screenshots for this exemplar.")
    for kind, d in [
        ("SCREENSHOT A1", "IAM dashboard / groups — the four groups with member counts; an MFA-enabled user; region ap-southeast-2."),
        ("SCREENSHOT A2", "VPC subnets — public-web-a, private-app-a, private-data-a with CIDRs and AZ; route tables (IGW / NAT)."),
        ("SCREENSHOT A3", "EC2 instances + ASG — running instance(s) with tags; ASG min/desired/max and scaling policy."),
        ("SCREENSHOT A4", "ALB target group — instance(s) reporting Healthy."),
        ("SCREENSHOT A5", "RDS database — available, MySQL 8.0.35, Single-AZ, encryption enabled."),
        ("SCREENSHOT A6", "S3 buckets — attachments + backups with Block Public Access enabled; EBS volumes attached."),
        ("SCREENSHOT A7", "Security groups — sg-alb / sg-app / sg-db with inbound rules expanded."),
        ("SCREENSHOT A8", "CloudWatch Alarms — the baseline alarm set."),
    ]:
        ev(doc, kind, d)
    h1("Appendix B — Configuration exports")
    ex.para(doc, "Each export is described; in a live submission these are CLI/console JSON dumps attached as files or code blocks.")
    for kind, d in [
        ("CONFIG EXPORT B1", "IAM group permission policies (one per group)."),
        ("CONFIG EXPORT B2", "Security-group rules for sg-alb / sg-app / sg-db (describe-security-groups)."),
        ("CONFIG EXPORT B3", "VPC, subnet and route-table configuration."),
        ("CONFIG EXPORT B4", "Launch template + ASG configuration."),
        ("CONFIG EXPORT B5", "ALB, target group and listener configuration."),
        ("CONFIG EXPORT B6", "RDS instance configuration (describe-db-instances)."),
        ("CONFIG EXPORT B7", "S3 bucket policy, public-access-block and encryption configuration."),
    ]:
        ev(doc, kind, d)
    h1("Appendix C — Test and simulation evidence")
    for kind, d in [
        ("TEST EVIDENCE C1", "ALB target group healthy (connectivity)."),
        ("TEST EVIDENCE C2", "EC2 → RDS mysql connection succeeding; external attempt timing out (negative test)."),
        ("TEST EVIDENCE C3", "ASG scale-out activity history + CloudWatch CPU graph from the autoscaling test."),
        ("TEST EVIDENCE C4", "SELECT VERSION() = 8.0.35 over a TLS connection."),
        ("TEST EVIDENCE C5", "Placeholder page reached via the ALB DNS name (HTTP 200/302)."),
    ]:
        ev(doc, kind, d)
    ex.para(doc, "Note: HA failure/resize simulation evidence (Appendix C of the HA report) is not applicable "
                 "to this foundation-build phase — see §6.5–6.9.")

    h1("Appendix D — Reflections")
    ex.uoc(doc, "[ICTCLD401 FS Learning, FS Planning and organising, FS Self-management]")
    h3("R1 — Lessons applicable beyond this build")
    ex.para(doc, "Tagging discipline paid off: tagging every resource with the engagement reference from the "
                 "start made it trivial to scope the permission boundary, find resources for the evidence "
                 "screenshots, and reason about cost. I arrived at it after an early hunt for an untagged "
                 "security group; next time I'd enforce tagging via the launch template and an SCP from "
                 "day one.")
    h3("R2 — Decisions in hindsight")
    ex.para(doc, "Choosing SSM Session Manager over a bastion host was the right call — it removed a public "
                 "RDP entry point I would otherwise have had to harden and monitor, and it logged every "
                 "session. If I were to do it again I'd size the EBS data volume with a little more headroom; "
                 "100 GB is adequate now but a single resize during a quiet period would have been cheap "
                 "insurance against the LMS's steady attachment growth.")

    h1("Document control")
    ex.etable(doc, ["Field", "Value"],
              [["Document version", "v1.0 — Initial submission"],
               ["Author", "MTS Consultant"],
               ["Engagement", "YAT LMS Cloud Migration — Foundation Build Phase"],
               ["Distribution", "Sam Walker (YAT ICT Manager), Pat Lin (MTS Senior Consultant)"],
               ["Successor document", "YAT LMS HA Hardening — HA Design + HA Deployment Report (AT3)"]],
              widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/assessments/AT2/AT2-exemplar-deployment-report.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
