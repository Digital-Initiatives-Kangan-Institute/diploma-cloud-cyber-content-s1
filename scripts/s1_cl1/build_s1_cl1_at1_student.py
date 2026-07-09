#!/usr/bin/env python3
"""Build the S1-CL1 AT1 STUDENT instrument (.docx) by populating the Kangan template.

The student-facing version of AT1. It imports the assessor module (build_s1_cl1_at1_assessor)
to single-source the shared content — DETAILS, the CHECK box, and the 'Instructions to Student'
prose (a.PROSE) — and defines its own student-facing wording for the instruction table and the
marking criteria (the authored student copy is phrased in the second person and omits the UoC
traceability). It loads the official Kangan 'Project Assessment - Student.docx' template. The
assessor-only material — the Marking Benchmark, the observation process and the UoC reverse-map
— is NOT included.

Usage:  python scripts/s1_cl1/build_s1_cl1_at1_student.py [output.docx]
Default: S1-CL1-Cloud-Design-Build/assessments/AT1/AT1-BusinessCase-Student.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_s1_cl1_at1_assessor as a  # noqa: E402  (shared helpers + content — single source of truth)
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_tables import add_section_row, clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Student.docx")


# ---------- student-facing content (second person; no UoC traceability) ----------

OVERVIEW = [
    'You are being assessed on the production and presentation of a Business Case for the YAT LMS Cloud Migration. The assessment is split into two parts (Parts A and B — see Task(s) to be assessed below):',
    'Part A (written): the Business Case document, including all four appendices (CBA detail, Knowledge Evidence responses, supporting research, Feedback Record)',
    'Part B (observed): the Business Case presentation delivered to the role-played YAT board, with sign-off captured from the role-played YAT ICT Manager',
    'This is an open-book assessment. You may use the YAT intranet, the AWS Pricing Calculator, AWS Academy lab environments, course reference materials, and external research (which must be cited) throughout. Where the Business Case requires synthesis (e.g. the current-state summary in §4), you must produce content in your own words — verbatim reproduction of intranet material is not satisfactory.',
    'Reasonable adjustment for this assessment may include extending the time allowed, varying the location or format of the presentation event (e.g. in-person vs video conference), allowing you to present to a recorded audience for asynchronous assessment where appropriate, or providing one-on-one verbal explanation of board questions where needed.',
    'Teacher/assessor support level: the teacher/assessor may clarify task requirements and the scenario but must not guide you to specific recommendations or correct answers. During the presentation, the teacher/assessor (in the role of the YAT ICT Manager) may ask substantive questions but must not coach you through your responses.',
    'Submission: Part A (Business Case) is submitted via the LMS as a .docx with all four appendices populated. Part B (presentation deck) is submitted via the LMS as a .pptx. The completed Feedback Record and the signed Sign-off block are attached to Part A.',
    'The assessment will not proceed if for any reason it is not safe to do so. The assessor must advise you of the reason for suspending the assessment, and what safety action should be taken and of revised arrangements for the assessment when it is safe to do so.',
    'There is a zero tolerance for plagiarism, cheating and collusion. You will be expected to make a declaration that all work is your own prior to submission. Refer to the Training and Assessment Policy for further information.',
]

TASKS = [
    'In this project, you take on the role of a consultant from MP Tech Solutions (MTS) engaged by YAT College to lead a migration of their mission-critical Learning Management System (LMS) to the cloud. AT1 is the analysis and planning phase of that engagement — you will produce a Business Case for the migration and present it to the YAT board for approval of the action plan.',
    'There are two parts to this assessment (Parts A and B). Each part is individually assessed as Satisfactory / Not Yet Satisfactory.',
    'Part A — Business Case',
    "Produce a written Business Case for the YAT LMS Cloud Migration using YAT's standard Business Case template.",
    'Part B — Presentation to the YAT board',
    'Prepare a slide deck and present your Business Case to the YAT board to seek approval of the action plan in Business Case §10.',
]

RESOURCES = [
    'Teacher/assessor supplied resources',
    "Access to the YAT intranet — provided as the engagement's reference site",
    'AWS Academy lab access — Cloud Foundations [104469] + Cloud Architecting [172221]',
    'Student supplied resources',
    'Computer with web browser',
    'Word-processing and presentation software (e.g. Microsoft Word + PowerPoint, or equivalents)',
]

CRITERIA = [
    'To receive a Satisfactory outcome for this assessment you must:',
    'Achieve Satisfactory on every criterion in the Part A Marking Guide (Business Case — covers Business Case sections §1–§11, Sign-off, and Appendices 1–4)',
    'Achieve Satisfactory on every criterion in the Part B Marking Guide (Presentation — covers the observed presentation event including Q&A, feedback capture, and sign-off)',
    'Submit all required artefacts:',
    'Completed Business Case (.docx) with all four appendices populated',
    'Presentation deck (.pptx) with populated speaker notes',
    'Completed Feedback Record (attached as Business Case Appendix 4)',
    'Signed Sign-off block from the role-played YAT ICT Manager (within the Business Case)',
]

RESULTS = [
    'If you are deemed not satisfactory for any of the observations, you will be given one (1) more attempt at this assessment (or part thereof) or your teacher/assessor will negotiate a further assessment with you. The second attempt must be completed within 10 working days from the date your feedback is given.',
]

SUBMIT = [
    'Part A is submitted to the LMS as the populated Business Case (.docx) with all four appendices completed.',
    'Part B is submitted to the LMS as the presentation deck (.pptx or equivalent) with populated speaker notes. The completed Feedback Record and the signed Sign-off block are attached to the Business Case (as Appendix 4 and the Sign-off block respectively).',
]

MARKING_A = [
    ["A1 - Strategic Alignment Analysis — you analysed YAT's ICT Strategic Plan against the industry environment and organisational objectives, with citations "],
    ["A2 -Current State of YAT's ICT — you produced a synthesised summary of YAT's current ICT systems and practices in your own words (not verbatim from the intranet) "],
    ["A3 - Gap Analysis — you compared YAT's strategic plan objectives against the current state and identifies gaps, improvement opportunities, and proposed changes"],
    ['A4 - Options Considered and Evaluation — you defined the LMS workload, considers both options (in-house renewal vs cloud migration to AWS), names and justifies the evaluation method, and produces an initial impact + difficulty assessment for both options'],
    ['A5 - Appendix 1 Cost-Benefit Analysis — you completed the 5-year CBA covering both options with line-item detail in Appendix 1; assumptions stated; year-by-year projections; comparison summary; avoided-downtime benefit; Year-1 cash-flow comparison; sensitivity analysis. AWS Pricing Calculator export attached in Appendix 3. Numbers reconcile between Appendix 1 detail, summary, and Executive Summary '],
    ['A6 - Risk and Impact Assessment — you populated intangibles comparison for both options (honest about trade-offs of the recommended option) and produces a risk register for the recommended option '],
    ['A7 - Recommendation — you state a clear recommendation with rationale connecting back to the CBA findings, intangibles, and risk register'],
    ["A8 - Action Plan — you developed a phased action plan with prioritised changes, implementation schedule with dependencies, standards / targets / success metrics, implementation methods, alignment with YAT's change-management procedure, and an action-plan risk register"],
    ['A9 - Next Steps and Decision Requested — you explicitly named what the board is being asked to approve today and what is deferred to later sign-off gates'],
    ['A10 - Executive Summary — one-page summary covering the recommendation, 5-year cost position, top 2–3 risks of the recommended option, and the asked-for decision. Numbers reconcile with A5'],
    ['A11 - Appendix 2 Knowledge Evidence — you answered all 9 KE questions (5 selection-style + 4 demonstration-style) with reference to specific sections of your own Business Case (not generic textbook answers)'],
    ['A12 - Appendix 3 Supporting Research — AWS Pricing Calculator export attached; external research sources cited with URLs and access dates'],
    ['A13 - Appendix 4 Feedback Record — completed during/after the Part B presentation event using the YAT Feedback Record template; captures actual board feedback'],
    ['A14 - Sign-off block — signed by the role-played YAT ICT Manager at the end of the Part B presentation event'],
    ['A15 - Document quality — Business Case uses plain English, appropriate vocabulary, terminology, diagrams, numerical information, formatting and structure relevant to a professional consulting deliverable for an RTO client'],
]

MARKING_B = [
    ['B1 - You reported on proposed changes, gaps and improvement opportunities to the role-played superior during the meeting (walks the board through the relevant Business Case content)'],
    ['B2 - The documented evaluation process is provided to the superior at/before the meeting; you walked the superior through it for feedback during the meeting '],
    ["B3 - Action plan is provided to the superior during the meeting; you explicitly sought the superior's feedback and approval on it"],
    ['B4 - You document and communicate the work through the presentation event (deck + verbal delivery to the role-played stakeholder)'],
    ['B5 – You seek feedback during Q&A and respond substantively to feedback received'],
    ['B6 – You confirm, seek, and respond to feedback with the role-played superior'],
    ['B7 – You use plain English and translate technical terminology when necessary, communicating with the role-played superior to articulate ideas, requirements, and plans '],
    ['B8 - You elicit information using effective listening and questioning techniques during the Q&A'],
    ['B9 – You use listening and questioning techniques to articulate complex concepts and requirements using industry language for the intended audience'],
]


def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details (student name / ID / assessor name+email left blank to fill) ----
    t_details = doc.tables[0]
    set_cell_content(t_details.rows[2].cells[1], a.DETAILS["qualification"])
    set_cell_content(t_details.rows[3].cells[1], a.DETAILS["units"])
    set_cell_content(t_details.rows[4].cells[1], a.DETAILS["task_title"])
    set_cell_content(t_details.rows[5].cells[1], a.DETAILS["task_number"])

    # ---- Table 1: Student instructions ----
    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Time allowed"), "")
    set_cell_content(find_instruction_row(t_instr, "Location"), "")
    set_cell_content(find_instruction_row(t_instr, "Resources required"), RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA)
    set_cell_content(find_instruction_row(t_instr, "Results"), RESULTS)
    # 'Important information' is left as the template's standard text.
    # add a 'How to Submit' row at the end (matching the table style)
    submit_row = t_instr.add_row()
    set_cell_content(submit_row.cells[0], "How to Submit")
    for r in submit_row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell_content(submit_row.cells[1], SUBMIT)

    # ---- drop the template's marking-guide intro line (the authored student copy omits it) ----
    for p in doc.paragraphs:
        if p.text.strip() == "List the criteria the student will be assessed against for the project assessment.":
            p._element.getparent().remove(p._element)
            break

    # ---- Table 2: Assessment criteria (what the student is marked against — no UoC/benchmark) ----
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)  # keep the two header rows
    add_section_row(t_mark, "Part A – Business Case")
    for lines in MARKING_A:
        row = t_mark.add_row()
        set_cell_content(row.cells[0], lines)
        set_cell_content(row.cells[1], a.CHECK)
    add_section_row(t_mark, "Part B - Presentation")
    for lines in MARKING_B:
        row = t_mark.add_row()
        set_cell_content(row.cells[0], lines)
        set_cell_content(row.cells[1], a.CHECK)

    # ---- drop the template's trailing marking-guide boilerplate ----
    for text in ("Add or delete rows as required",
                 "If questioning or observation is incorporated into this assessment task, "
                 "you can incorporate a Practical Observation Checklist."):
        for p in doc.paragraphs:
            if p.text.strip() == text:
                p._element.getparent().remove(p._element)
                break

    # ---- Detailed task instructions (shared prose; NO marking benchmark / UoC) ----
    doc.add_paragraph("Instructions", style="Heading 1")
    for text, style in a.PROSE:
        doc.add_paragraph(text, style=style)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/assessments/AT1/AT1-BusinessCase-Student.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
