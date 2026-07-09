#!/usr/bin/env python3
"""Build the S1-CL1 AT3 STUDENT instrument (.docx) by populating the Kangan template.

The student-facing version of AT3. It imports the assessor module (build_s1_cl1_at3_assessor)
to single-source the shared content — DETAILS and the CHECK box — and defines its own
student-facing wording for the instruction table, the assessment-criteria table, and the
detailed instructions (phrased in the second person, with NO UoC traceability). It loads the
official Kangan 'Project Assessment - Student.docx' template. The assessor-only material — the
Marking Guide UoC lines, the HA Design + HA Deployment Report Benchmarks, and the UoC
reverse-map — is NOT included.

(Unlike CL1 AT1, the AT3 student copy carries a DIFFERENT, fuller instructions body from the
assessor copy, so the instruction prose is authored here rather than imported.)

EVERGREEN: emits NO literal scenario-site URL. Where the committed student instrument carried
the placeholder intranet URL parenthetical (after "the YAT intranet"), the generator drops only
the URL + its wrapper and keeps every other word verbatim.

Usage:  python scripts/s1_cl1/build_s1_cl1_at3_student.py [output.docx]
Default: S1-CL1-Cloud-Design-Build/assessments/AT3/AT3-High-Availability-Student.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_s1_cl1_at3_assessor as a  # noqa: E402  (shared DETAILS + CHECK — single source of truth)
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_tables import add_section_row, clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Student.docx")


# ---------- student-facing content (second person; no UoC traceability) ----------

# EVERGREEN: the committed source read "... the YAT intranet (<url>), AWS Academy ..."; the URL
# parenthetical is dropped, every other word verbatim.
OVERVIEW = [
    'You are being assessed on the design and implementation of HA improvements to the cloud infrastructure you (in-scenario, with MTS) deployed in AT2. AT3 is the third and final assessment task in the S1-CL1 Cloud Design and Build cluster — it closes out the YAT LMS Cloud Migration engagement.',
    'The assessment has two parts, both written deliverables:',
    'Part A (written design): the HA Design document — you produce the HA-equivalent architecture that supersedes the AT2 baseline. Covers HA requirements, review of the AT2 baseline, designed HA architecture, implementation sequencing, and simulation plan.',
    'Part B (written deployment report): the HA Deployment Report — documents the implementation of your HA design during a simulated maintenance window, the simulation outcomes (failure + resize), any post-simulation adjustments, the feedback you received, and the final sign-off.',
    'Both parts are submitted together as one bundle. There is no presentation or observation event.',
    'This is an open-book practical assessment. You may use the YAT intranet, AWS Academy lab environments, AWS documentation, course reference materials, and external research (which must be cited) throughout.',
    'AT3 is the third of three assessment tasks in the cluster. It builds on AT1 (Business Case) and AT2 (Foundation Build). You continue in the same MTS consultant role across all three.',
    'Submission: the completed HA Design (.docx) AND HA Deployment Report (.docx) bundle, submitted via the LMS.',
    'The assessment will not proceed if for any reason it is not safe to do so. Your assessor must advise you of the reason for suspending the assessment, and what safety action should be taken, and of revised arrangements when it is safe to resume.',
    'There is zero tolerance for plagiarism, cheating and collusion. You will be expected to make a declaration that all work is your own prior to submission. Refer to the Training and Assessment Policy for further information.',
]

TASKS = [
    "Following the AT2 foundation-build phase, you return to YAT for the final phase of the engagement — HA hardening. The cloud infrastructure you handed over at the end of AT2 is single-AZ and non-HA; YAT's strategic target of 99.9% availability has not yet been achieved. AT3's HA work closes that gap and formally closes out the engagement.",
    'The task has two parts that combine into a single submission:',
    'Part A — HA Design. Produce the HA Design document using the YAT-provided HA Design template. Covers HA requirements, a review of the AT2 baseline (architecture, single points of failure, recovery objectives, vertical-scaling impact), your HA-equivalent architecture, implementation sequencing, and a simulation plan.',
    "Part B — HA Deployment Report. Implement your HA design in the AWS Academy lab during a simulated Saturday late-night maintenance window of approximately 3.5 hours. Run the failure and resize simulations from your design's §6 simulation plan. Document everything in the HA Deployment Report (using the YAT-provided template), secure feedback and final sign-off from the role-played YAT ICT Manager, and file the documentation per YAT's records procedures.",
    'The HA Deployment Report uses the same shape as the AT2 Deployment Report (same section structure in the same places) but with HA-specific content within each section.',
    "This phase is the HA hardening only. The cloud infrastructure HA work is in scope. LMS application installation, MySQL data migration, cutover, and organisational change management are YAT in-house IT's responsibility — out of scope for this assessment.",
]

TIME = ["Design Phase - ", "Implementation Window – 3.5 hours"]

RESOURCES = [
    'Provided to you (download from the YAT intranet)',
    'HA Design template (intranet Templates section) — Part A template',
    'HA Deployment Report template (intranet Templates section) — Part B template',
    'AT2 supplied baseline design (internal-lms-cloud-architecture-design-S1-CL1-AT2.md, in Engagement Documents) — your starting point for the architecture review',
    'Example previous Deployment Report (intranet Document Archive) — reference for what good looks like (when authored)',
    'HA database requirements (internal-ha-database-requirements-S1-CL1-AT3.md) — HA expectations specific to the database',
    'AT2 baseline CloudFormation template — provided by your assessor at the start of the assessment day; deploy in your AWS Academy lab to get the AT2 baseline as a consistent starting state',
    'All other scenario materials (LMS application spec, LMS cloud migration requirements, organisational policies, your AT2 Deployment Report)',
    'Provided externally',
    'AWS Academy lab access — Cloud Foundations [104469] + Cloud Architecting [172221]',
    'You supply',
    'Computer with web browser',
    'Word-processing software (Microsoft Word or equivalent)',
    'A screenshot tool',
]

CRITERIA = [
    'To receive a Satisfactory result for AT3 you must:',
    'Achieve Satisfactory on every criterion in the Part A Assessment Criteria table (HA Design — A1 through A7)',
    'Achieve Satisfactory on every criterion in the Part B Assessment Criteria table (HA Deployment Report — B1 through B16)',
    'Submit both deliverables (.docx) with all sections and appendices populated, using the YAT-provided templates',
]

RESULTS = [
    'If you are deemed not satisfactory for any of the observations, you will be given one (1) more attempt at this assessment (or part thereof) or your teacher/assessor will negotiate a further assessment with you. The second attempt must be completed within 10 working days from the date your feedback is given.',
]

MARKING_A = [
    "A1 - §2 HA requirements — you documented the reliability, recoverability and service-level requirements your HA design must meet, derived from the LMS application specification and the LMS cloud migration requirements",
    'A2 - §3.1–§3.4 Review of AT2 baseline — you reviewed the AT2 baseline architecture, identified all single points of failure, estimated current-state recovery objectives per component, and identified components requiring vertical scaling with their availability impact',
    'A3 - §3.5 Review findings documented — you produced a clear summary of the review findings (the gap between AT2 baseline and the §2 HA requirements) suitable for stakeholder review',
    'A4 - §4.1–§4.12 HA architecture designed — you produced an HA-equivalent architecture covering all layers (IAM, network cross-AZ, compute cross-AZ, ALB cross-AZ, RDS Multi-AZ, storage, security, monitoring HA-tuned, naming, backup)',
    'A5 - §4.13–§4.15 HA design state — recovery objectives, vertical-scaling considerations, SPOFs removed — you documented the recovery objectives the HA design achieves, the components still requiring vertical scaling, and the SPOFs from the baseline that the HA design removes',
    "A6 - §4 (overall) Design documented per business needs — your design is complete, internally consistent, and addresses YAT's business needs (99.9% availability, ≤1h RPO, ≤4h RTO, data residency)",
    'A7 - HA Design document quality — uses plain English, appropriate technical vocabulary, structure, and depth relevant to a professional consulting design deliverable',
]

MARKING_B = [
    'B1 - §1 Executive Summary — you produced a concise (≤ 1 page) summary of the HA hardening delivered, including the simulation outcomes in headline form',
    'B2 - §4 Build narrative — HA implementation across all tiers (cross-AZ network, ASG cross-AZ expansion, Multi-AZ database, ALB cross-AZ targets, monitoring HA-tuning, storage adjustments) — you documented the implementation work that translated your HA Design into running infrastructure',
    'B3 - §6.1 Connectivity tests — you verified post-implementation connectivity between resources at all tiers',
    'B4 - §6.2 Failure simulation — you executed the failure simulations planned in your HA Design §6.1, observed outcomes, captured evidence',
    'B5 - §6.3 Resize simulation — you executed the resize simulations planned in your HA Design §6.2, measured availability impact, captured evidence',
    'B6 - §6.4 Availability measurement — you set up and used CloudWatch (or equivalent) to define, monitor, and record resource availability across the maintenance window',
    'B7 - §6.5 Simulation findings vs design — you compared the actual simulation outcomes against the expected outcomes specified in your HA Design §6 simulation plan, and documented the comparison',
    'B8 - §6.6 Adjustments per simulations — you documented any adjustments made to the architecture, configuration, or monitoring based on simulation outcomes (or explicitly justified that no adjustments were needed)',
    "B9 - §7.4 Documentation filed per organisational policies — you documented the filing of the HA Design + HA Deployment Report + appendices per YAT's records management procedures",
    'B10 - §7.5 Feedback record — you captured the feedback received from the role-played YAT ICT Manager / MTS Senior Consultant on the HA Design and HA implementation, including your response',
    'B11 - §7.6 Final sign-off obtained — the role-played YAT ICT Manager has signed off on the HA hardening; the sign-off block in §7.6 is populated and signed',
    'B12 - §8 Knowledge Evidence Responses — you answered all 6 KE questions with specific reference to your own HA design and implementation, not generic textbook answers',
    "B13 - Appendix A Build Evidence — all named screenshots present (cross-AZ subnets, Multi-AZ RDS, cross-AZ ASG instances, HA-tuned alarms, etc.), correct AWS region visible, named items visible per the template's requirements",
    'B14 - Appendix C Simulation Evidence — all simulation evidence items present and consistent with the §6 narrative outcomes',
    'B15 - Appendix D Reflections — two honest reflective responses present (decisions in hindsight + problem solving under time pressure), both specific to your own HA work',
    'B16 - HA Deployment Report document quality — uses plain English, appropriate technical vocabulary, structure, formatting, and depth relevant to a professional consulting deliverable',
]

MARKING_NOTE = ('Note: §2 (Engagement Context), §3 (Scope of Deployment), §5 (Configuration Decisions), '
                '§7.1–§7.3 (Access, Runbook, Limitations) are required template sections — you must complete '
                'them — but they are not separately graded above. Their quality manifests in criterion B16 '
                '(Document quality).')

# Detailed student instructions (second person; no benchmark / UoC). Emitted after the marking table.
INSTRUCTIONS = [
    'The engagement — picking up where AT2 left off',
    'YAT College is migrating their mission-critical Learning Management System (LMS) from on-premises to AWS. You are an MTS Consultant on this engagement, reporting to Pat Lin (MTS Senior Consultant). Sam Walker (YAT ICT Manager) is your primary YAT-side stakeholder.',
    "At the end of AT2 you handed YAT the cloud foundation — single-AZ, non-HA, ready to operate but not yet meeting YAT's strategic 99.9% availability target. YAT IT deployed the LMS application on the infrastructure you provided and ran the cutover; the LMS is now in production on AWS. After reviewing your AT2 Deployment Report, the YAT board has approved the next phase: HA hardening. You are returning to MTS to lead this final phase of the engagement.",
    'Scope of your work in this assessment',
    'The LMS application is already deployed (by YAT IT after the AT2 handover) and is running in production on the AWS infrastructure you built in AT2. Your AT3 work is HA hardening of the existing running infrastructure — Multi-AZ database, cross-AZ compute, HA-tuned monitoring, and so on — performed in-place during the simulated maintenance window. No application re-deployment, data migration, or cutover is required: those happened in earlier phases.',
    "What remains YAT IT's responsibility post-handover: re-validating the LMS application against your HA-hardened infrastructure (e.g. verifying DOODLE behaves correctly during a Multi-AZ database failover, tuning session-affinity at the application layer if needed). That re-validation is out of MTS scope for this assessment per the LMS Migration Role Brief on the YAT intranet.",
    'Your AT3 deliverable stops at HA-hardened infrastructure handed over to YAT IT. You design the HA improvements, implement them in-place, run the simulations, capture feedback, obtain sign-off, file the documentation — and the engagement closes.',
    'Part 1 — HA Design (Part A)',
    "Download the HA Design template from the YAT intranet's Templates section. The template mirrors the same section structure as the AT2 supplied design (internal-lms-cloud-architecture-design-S1-CL1-AT2.md) — so you have a worked example of what a good design looks like, and you know exactly what shape your design needs to take.",
    'Cover in your design:',
    '§2 The HA requirements you must meet (from the LMS application spec + cloud migration requirements — 99.9% availability, RPO ≤ 1h, RTO ≤ 4h)',
    '§3 A review of the AT2 baseline: ',
    '§3.1 Architecture review against HA requirements',
    '§3.2 Single points of failure identified (this is the hardest section — be thorough)',
    '§3.3 Recovery objectives the baseline currently achieves (quantified per component)',
    '§3.4 Components requiring vertical scaling, with availability impact',
    '§3.5 Summary of the gap between baseline and HA requirements',
    '§4 Your HA-equivalent architecture: ',
    '§4.1–§4.12 Each layer of the design (IAM, network cross-AZ, compute cross-AZ ASG, ALB cross-AZ, RDS Multi-AZ, storage, security, monitoring HA-tuned, naming, backup baseline)',
    '§4.13 Recovery objectives the HA design achieves',
    '§4.14 Components still requiring vertical scaling (post-HA)',
    '§4.15 Single points of failure removed (cross-referenced to §3.2)',
    "§5 Implementation sequencing — the order you'll apply changes during the maintenance window",
    "§6 Simulation plan — the failure and resize simulations you'll run to verify the HA design works",
    "Read the AT2 supplied design end-to-end first. Your HA Design follows the same shape but addresses the HA gap. Your AT2 design is your reference — you're building the HA-equivalent of it.",
    'Part 2 — HA Deployment Report (Part B)',
    'After your HA Design is complete and ready for implementation, you start the maintenance-window work. You implement the HA design in the AWS Academy lab and document everything in the HA Deployment Report.',
    'Maintenance window framing',
    'For the implementation work, you have a simulated Saturday late-night maintenance window of approximately 3.5 hours, framed as the YAT LMS being in its lowest-traffic period.',
    'Plan your changes to minimise service disruption — most HA changes on AWS are non-disruptive (adding subnets, expanding ASGs, adjusting monitoring), but the RDS Single-AZ → Multi-AZ conversion incurs a brief failover blip (~30–60 seconds)',
    'Brief service blips within the window are acceptable. A few seconds of LMS unreachability during the RDS failover, or during an ALB target re-registration, is normal HA-hardening behaviour',
    'At the end of the window, the LMS infrastructure must be back online — either with HA hardening complete or rolled back to the pre-window state. Leaving it half-done is not acceptable',
    "If you reach the 2.5-hour mark and aren't yet stable, plan your rollback rather than push to completion",
    "This framing is deliberate. At this qualification level, if you were doing this on truly mission-critical infrastructure, you would be supervised by a senior engineer; the simulated maintenance window is a fair proxy for the real-world risk envelope you'll experience early in your career.",
    'What you do during the window',
    'Execute your implementation sequence from HA Design §5 — apply the HA changes one at a time, verifying each before proceeding',
    'Take screenshots as you go (Appendix A of the Deployment Report has the list — capture them at the moment they exist, not retrospectively)',
    'Run your simulations (failure sim + resize sim) per HA Design §6',
    'Measure availability across the window using CloudWatch',
    'Adjust the architecture if simulations reveal gaps (or explicitly justify that no adjustments were needed)',
    'What you document in the Deployment Report',
    'Use the YAT-provided HA Deployment Report template. The structure mirrors the AT2 Deployment Report (same sections in the same places) with HA-specific content inside each section. Cover:',
    '§1 Executive Summary (write this last)',
    '§2 Engagement Context (references AT1, AT2, and your HA Design)',
    '§3 Scope of Deployment (what HA work is in this phase)',
    '§4 Build narrative (HA implementation across all tiers)',
    '§5 Configuration decisions made during implementation',
    '§6 Testing, Simulation and Validation — the heart of the AT3 evidence: ',
    '§6.1 Post-implementation connectivity tests',
    '§6.2 Failure simulation outcomes',
    '§6.3 Resize simulation outcomes',
    '§6.4 Availability measurement across the maintenance window',
    "§6.5 Simulation findings compared against the HA Design's expected outcomes",
    '§6.6 Adjustments made per simulation results (or justified statement that none were needed)',
    '§7 Operational Handover: ',
    '§7.1–§7.3 Access, runbook references, known limitations',
    "§7.4 Documentation filed per YAT's records procedures",
    '§7.5 Feedback record from the role-played stakeholders',
    '§7.6 Final sign-off block — signed by the role-played YAT ICT Manager',
    '§8 Knowledge Evidence responses (six contextual questions about HA concepts applied to your own work)',
    'Appendix A — HA build evidence (screenshots)',
    'Appendix B — Configuration exports',
    'Appendix C — Simulation evidence (failure sim, resize sim, availability data)',
    'Appendix D — Reflections (two short responses on your own learning during this phase)',
    'AT2 baseline starting state',
    'Your assessor will provide a CloudFormation template at the start of the assessment day. Deploy it in your AWS Academy lab during the morning prep slot — this instantiates the AT2 baseline (single-AZ, non-HA) so you start the afternoon HA work from a consistent known state.',
    'Steps:',
    'Download the CloudFormation template (.yaml) from where your assessor provides it',
    'AWS Academy → AWS Console → CloudFormation → Create stack → Upload template',
    'Provide any parameters (stack name, key pair, etc.) and launch',
    'Wait ~10–15 minutes for deployment to complete',
    'Verify the baseline is healthy (curl the ALB DNS name; check RDS is available)',
    'Take a break, then return for the afternoon HA work',
    'Submit',
    'Submit both deliverables together via the LMS at the end of the assessment:',
    'HA Design document (.docx) — Part A',
    'HA Deployment Report (.docx) — Part B, with all four appendices populated',
    'Tips for success',
    "Read the AT2 supplied design + your own AT2 Deployment Report first. They are the inputs to your HA Design — the baseline you're hardening",
    'The architecture review in HA Design §3 is the hardest section to do well. Take time to identify SPOFs carefully — every one you miss is one your simulation will fail to catch',
    'Recovery objectives must be quantified. "Better than before" is not an answer. Cite specific RPO/RTO targets and what your design delivers',
    "The implementation sequence matters. Plan additive changes first (new subnets, ASG capacity); leave the RDS Multi-AZ conversion until after you've added the new AZ infrastructure",
    "Simulate honestly. A simulation that confirms the design works is great; a simulation that exposes a gap is also great — both demonstrate competent engineering. Don't fake outcomes",
    "Watch the clock during the maintenance window. If you're 2.5 hours in and not yet stable, start the rollback rather than push to completion",
    "Knowledge Evidence responses (§8) are about your own work. Generic textbook answers don't pass",
    "File the report per YAT's records procedures (§7.4) — this is part of what's being assessed",
    'Get the sign-off block (§7.6) signed at the end of the engagement — without it Part B cannot be Satisfactory',
    'Take screenshots as you go, not at the end. This is the single most common failure mode',
]


def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details (student name / ID / assessor name+email left blank to fill) ----
    t_details = doc.tables[0]
    set_cell_content(t_details.rows[2].cells[1], a.DETAILS["qualification"])
    set_cell_content(t_details.rows[3].cells[1], a.DETAILS["units"])
    set_cell_content(t_details.rows[4].cells[1], a.DETAILS["task_title"])
    set_cell_content(t_details.rows[5].cells[1], a.DETAILS["task_number"])

    # ---- Table 1: Student instructions ----
    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Time allowed"), TIME)
    set_cell_content(find_instruction_row(t_instr, "Location"), "")
    set_cell_content(find_instruction_row(t_instr, "Resources required"), RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA)
    set_cell_content(find_instruction_row(t_instr, "Results"), RESULTS)
    # 'Important information' is left as the template's standard text.

    # ---- drop the template's marking-guide intro line (the authored student copy omits it) ----
    for p in doc.paragraphs:
        if p.text.strip() == "List the criteria the student will be assessed against for the project assessment.":
            p._element.getparent().remove(p._element)
            break

    # ---- Table 2: Assessment criteria (what the student is marked against — no UoC/benchmark) ----
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)  # keep the two header rows
    add_section_row(t_mark, "Part A")
    for line in MARKING_A:
        row = t_mark.add_row()
        set_cell_content(row.cells[0], line)
        set_cell_content(row.cells[1], a.CHECK)
    add_section_row(t_mark, "Part B")
    for line in MARKING_B:
        row = t_mark.add_row()
        set_cell_content(row.cells[0], line)
        set_cell_content(row.cells[1], a.CHECK)

    # ---- drop the template's trailing marking-guide boilerplate ----
    for text in ("Add or delete rows as required",
                 "If questioning or observation is incorporated into this assessment task, "
                 "you can incorporate a Practical Observation Checklist."):
        for p in doc.paragraphs:
            if p.text.strip() == text:
                p._element.getparent().remove(p._element)
                break

    # ---- Note on ungraded-but-required template sections (follows the marking table) ----
    doc.add_paragraph(MARKING_NOTE, style="Assessor text")

    # ---- Detailed task instructions (student prose; NO marking benchmark / UoC) ----
    doc.add_paragraph("Instructions", style="Heading 1")
    for text in INSTRUCTIONS:
        doc.add_paragraph(text, style="Assessor text")

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/assessments/AT3/AT3-High-Availability-Student.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
