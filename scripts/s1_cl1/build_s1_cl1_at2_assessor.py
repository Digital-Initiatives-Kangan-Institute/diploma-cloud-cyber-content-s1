#!/usr/bin/env python3
"""Build the S1-CL1 AT2 ASSESSOR instrument (.docx) by populating the Kangan template.

This is an institutional compliance document, NOT a YAT-branded artefact: it loads the
official Kangan 'Project Assessment - Assessor' template and fills it in, preserving the
Kangan structure and styles (Details, Teacher/Assessor instructions, Marking Guide,
Instructions to Student, Benchmark). It retro-fits a generator to the approved, hand-authored
CL1 AT2 instrument, reproducing it exactly at the content level. Mirrors the CL1 AT1 and
CL2 AT2 generators' mechanics.

EVERGREEN: unlike the committed source, this generator emits NO scenario-site URL. Where the
authored text named the intranet URL, only the URL (and its immediate " (…)"/" at …" wrapper)
is dropped; every other word is verbatim. There is no brand.WEBSITE_URL import.

AT2 = Cloud Foundation Build: YAT LMS Migration (ICTICT517 + ICTCLD401 + ICTCLD502) — a single
written Deployment Report (with four evidence appendices) documenting the implementation of a
supplied AWS foundation architecture in the AWS Academy lab. There is no presentation/observation
event. The assessor instrument carries the task instructions, the Marking Guide, the Deployment
Report Benchmark (per-section guidance + KE model answers), and the UoC coverage reverse-map. The
student instrument (build_s1_cl1_at2_student) carries only the student-facing content.

Usage:  python scripts/s1_cl1/build_s1_cl1_at2_assessor.py [output.docx]
Default: S1-CL1-Cloud-Design-Build/assessments/AT2/AT2-Deployment-Assessor.docx
"""
import sys
from pathlib import Path

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Assessor.docx")

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_tables import add_section_row, clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402


# ---------- content ----------

CHECK = "☐ Yes  ☐ No"  # marking-guide Satisfactory? cell (matches the Kangan template)

DETAILS = {
    "qualification": "ICT50220 Diploma of Information Technology",
    "units": [
        "ICTICT517 Match ICT needs with the strategic direction of the organisation",
        "ICTCLD401 Configure cloud services",
        "ICTCLD502 Design and implement highly-available cloud infrastructure",
    ],
    "task_title": "AT2 — Cloud Foundation Build: YAT LMS Migration",
    "task_number": "2 of 3",
}

# NOTE (evergreen): the committed doc named the intranet URL here — " (https://www.placeholder.com.au)"
# after "YAT intranet". This generator omits the URL and its wrapper; all other words verbatim.
OVERVIEW = [
    'Students are assessed on the implementation of a supplied AWS cloud architecture for the YAT LMS migration, and the production of a Deployment Report documenting the build.',
    'The assessment is a single-task project delivered as one written artefact (the Deployment Report) with four populated appendices. Implementation occurs in the AWS Academy lab environment authorised for this cluster. There is no presentation or observation event — all assessment evidence is captured in the Deployment Report and its appendices.',
    'This is an open-book practical assessment. Students may use the YAT intranet, the AWS Pricing Calculator, AWS Academy lab environments, AWS documentation, course reference materials, and external research (which must be cited) throughout.',
    'AT2 is the second of three assessment tasks in the S1-CL1 Cloud Design and Build cluster. It builds on AT1 (Business Case + presentation event) and feeds into AT3 (HA hardening + project closure). The student remains in the same MTS consultant role across all three.',
    'Reasonable adjustment for this assessment may include extending the time allowed, varying the lab access arrangement (e.g. extended lab hours), allowing alternative screenshot-evidence formats where assistive technology requires it, or providing one-on-one verbal explanation of the supplied design where needed.',
    'Teacher/assessor support level: the teacher/assessor may clarify task requirements, scenario context, or the supplied design but must not guide students to specific configuration decisions or correct knowledge-evidence answers. The eight Configuration Decisions left to the implementer (per the supplied design §14) must be made by the student.',
    'Submission: the completed Deployment Report (.docx) with all four appendices populated is submitted via the LMS.',
    'The assessment will not proceed if for any reason it is not safe to do so. You must advise the student of the reason for suspending the assessment, and what safety action should be taken. Advise the student of revised arrangements for the assessment when it is safe to do so.',
    'There is a zero tolerance for plagiarism, cheating and collusion. Students will be expected to make a declaration that all work is their own prior to submission. Refer to the Training and Assessment Policy for further information.',
]

# NOTE (evergreen): the committed doc named the intranet URL in the last line — " at https://www.placeholder.com.au".
TASKS = [
    "Following the YAT board's approval of the action plan in the student's AT1 Business Case engagement, the student took a period of planned annual leave (in-scenario narrative).",
    'During that time MTS Senior Architecture worked with YAT IT to translate the approved direction into a detailed technical design — the YAT LMS Cloud Architecture — Baseline Design — which is now the student\'s build specification.',
    'The student has returned to MTS to lead the foundation-build phase of the engagement. The task has two parts that combine into a single deliverable:',
    'Implement the supplied AWS architecture for the YAT LMS migration foundation build, in the AWS Academy lab environment authorised for this engagement.',
    "Produce a Deployment Report documenting the build, using the YAT-provided Deployment Report template. The report covers what was built, configuration decisions made where the supplied design left them open, testing and validation outcomes, an operational handover for YAT IT, written responses to six contextual Knowledge Evidence questions about the student's own build, and four appendices of evidence.",
    "The Deployment Report is the single submitted deliverable. All build evidence (screenshots, configuration exports, test results, reflections) is captured in the report's appendices — there is no separate portfolio submission and no presentation event.",
    'The architecture being implemented is intentionally non-HA at this stage. HA hardening is the next phase (AT3) and is explicitly out of scope for AT2.',
    "MTS scope: cloud infrastructure provisioning only. Per the LMS Migration Role Brief on the YAT intranet (§ Scope of the MTS consulting engagement), students must not perform LMS application installation, MySQL data migration, cutover activities, or organisational change management as part of AT2. Those are YAT in-house IT's responsibility in-scenario. The AT2 deliverable stops at infrastructure ready for application deployment. Assessors should not award credit for application-deployment work that falls outside this scope; equally, students who include such work in their report should be redirected to focus their evidence on the in-scope infrastructure deliverables.",
    'Note:',
    'All scenario materials, organisational policies, supplied design, report template, and previous-project examples for YAT College are available on the YAT intranet — students sign in to the intranet at the start of the engagement and refer to it throughout.',
]

# NOTE (evergreen): the committed doc named the intranet URL in the first item — " (https://www.placeholder.com.au)".
RESOURCES = [
    'Teacher/assessor supplied resources',
    'Access to the YAT intranet — supplying the supplied baseline design, the Deployment Report template, the example previous deployment report, and all other scenario materials',
    'AWS Academy lab access — Cloud Foundations [104469] + Cloud Architecting [172221]',
    'Deployment Report Benchmark (later in this document) for marking the submitted report',
    'Student supplied resources',
    'Computer with web browser',
    'Word-processing software (e.g. Microsoft Word or equivalent)',
    'A screenshot tool',
]

CRITERIA = [
    'To receive a Satisfactory outcome for this assessment the student must:',
    'Achieve Satisfactory on every criterion in the Marking Guide below',
    'Submit a completed Deployment Report (.docx) using the YAT-supplied template, with every section and every appendix populated',
]

# NOTE (evergreen): C2 named the intranet URL — " at https://www.placeholder.com.au".
CONDITIONS = [
    'These are conditions the assessor verifies as present before marking begins. They are not student-performance criteria — they are the conditions under which the assessment can validly be conducted.',
    'C1 Lab environment is accessible to the student throughout the assessment — AWS Academy Cloud Foundations [104469] + AWS Academy Cloud Architecting [172221] — providing cloud vendor service provider access, cloud managed database service (RDS), IDE / console / CLI / SSH-RDP tooling, and internet/web browser access',
    '[ICTCLD401 AC 1] · [ICTCLD401 AC 2] · [ICTCLD401 AC 3] · [ICTCLD502 AC 1] · [ICTCLD502 AC 2] · [ICTCLD502 AC 4] · [ICTCLD502 AC 6] · [ICTCLD502 AC 7]',
    'C2 The YAT scenario site / intranet is accessible to the student throughout the assessment — supplying the baseline design, Deployment Report template, example previous deployment report, LMS application specification, organisational policies, and reference documents',
    '[ICTCLD401 AC 4] · [ICTCLD502 AC 3] · [ICTCLD502 AC 5] · [ICTCLD502 AC 8]',
]

# Marking Guide — single-part (A1-A13); each criterion is [statement, UoC-traceability line].
MARKING = [
    ['A1 §1 Executive Summary — student produces a concise (≤ 1 page) summary of what was delivered, with explicit acknowledgement that HA hardening is deferred to AT3, and numbers/components reconcile with the body of the report.', '[ICTCLD401 PC 4.1] (document and communicate work — written communication evidence)'],
    ['A2 §4.1 + §4.2 + §4.7 Build narrative — Foundation tier (IAM, network, security) — student describes the IAM groups/users/MFA enforcement, the VPC + subnets + gateways + routing, and the security-group model; cross-references to Appendix A screenshots and Appendix B configuration exports', '[ICTCLD401 PC 1.4] · [ICTCLD401 PC 1.5] · [ICTCLD401 PC 1.6] · [ICTCLD401 PC 1.7] · [ICTCLD401 PC 2.1] · [ICTCLD401 PC 2.2] · [ICTCLD502 PC 1.3]'],
    ['A3 §4.3 + §4.4 + §4.5 + §4.6 Build narrative — Workload tier (compute, load balancing, database, storage) — student describes the EC2 + ALB + ASG, the RDS managed database, and the EBS + S3 storage; cross-references to Appendix A screenshots and Appendix B configuration exports', '[ICTCLD401 PC 2.3] · [ICTCLD401 PC 2.4] · [ICTCLD401 PC 2.5] · [ICTCLD502 PC 4.1]'],
    ['A4 §4.8 Build narrative — Operability (autoscaling configuration, baseline CloudWatch monitoring) — student describes the ASG scaling policy + the baseline CloudWatch alarms; identifies the shared security responsibility outcome of the build', '[ICTCLD401 PC 1.2] · [ICTCLD401 PC 3.1] · [ICTCLD502 PC 4.3] (monitoring baseline only — full availability monitoring is AT3)'],
    ['A5 §5 Configuration Decisions — student justifies all 8 decision points (C1–C8 from the supplied design §14) with rationale tied to the YAT LMS workload (referencing the LMS Application Specification), not generic justifications', '[ICTCLD401 PC 1.1] (discuss and compare cloud computing solutions) · [ICTCLD401 PC 1.3] (select best cloud computing solution and service)'],
    ['A6 §6 Testing and Validation — student ran the four required test categories (connectivity, autoscaling, database connectivity, end-to-end smoke); results documented with cross-reference to Appendix C evidence', '[ICTCLD401 PC 2.6] · [ICTCLD401 PC 3.2] · [ICTCLD502 PC 4.2]'],
    ["A7 §7 Operational Handover — student documents access information for YAT IT, names known limitations carried into AT3, and confirms filing of the report per YAT's documented records procedures", '[ICTCLD401 PC 4.3]'],
    ['A8 §8 Knowledge Evidence Responses — student answers all 6 KE questions with specific reference to their own build, not generic textbook answers', '[ICTCLD401 KE 5] · [ICTCLD401 KE 6] · [ICTCLD401 KE 7] · [ICTCLD401 KE 8] · [ICTCLD401 KE 9] · [ICTCLD401 KE 10]'],
    ["A9 Appendix A Build Evidence — all 17 named screenshots (A1–A17) present, correct AWS region visible, named items visible per the template's requirements", '[ICTCLD401 PE 1] · [ICTCLD401 PE 2] (the built artefacts are the evidence)'],
    ['A10 Appendix B Configuration Exports — all 7 named exports (B1–B7) present, exported from the actual deployed environment (consistent with the screenshots in Appendix A)', '[ICTCLD401 PE 2] (deeper configuration evidence)'],
    ['A11 Appendix C Test Evidence — test outcomes documented with supporting evidence; outcomes consistent with §6 narrative', '[ICTCLD401 PE 3] · [ICTCLD502 PC 4.2] (connectivity evidence) · [ICTCLD502 PC 4.3] (baseline monitoring evidence)'],
    ["A12 Appendix D Reflections — two honest reflective responses present, both specific to the student's own build experience", '[ICTCLD401 FS Learning] · [ICTCLD401 FS Planning and organising] · [ICTCLD401 FS Self-management skills]'],
    ['A13 Document quality — Deployment Report uses plain English, appropriate technical vocabulary, structure, formatting, and depth relevant to a professional consulting deliverable for an RTO client; student has correctly interpreted the supplied design and other complex technical documentation in the build', '[ICTCLD401 FS Writing] · [ICTCLD401 FS Reading] · [ICTCLD502 FS Reading]'],
]

# ---- Shared 'Instructions to Student' prose (single-sourced; the student builder imports these) ----
# The intro + Part 2/Tips blocks are IDENTICAL across the assessor and student copies; only the
# Part 1 block differs (student adds detail), so Part 1 is defined per-builder.
PROSE_INTRO = [
    ('The engagement — picking up where AT1 left off', 'Heading 2'),
    ('YAT College is migrating their mission-critical Learning Management System (LMS) from on-premises to AWS. You are an MTS Consultant on this engagement, reporting to Pat Lin (MTS Senior Consultant). Sam Walker (YAT IT Manager) is your primary YAT-side stakeholder.', 'Assessor text'),
    ('At the end of the AT1 engagement, the YAT board approved your action plan. You then took planned annual leave. During that period, MTS Senior Architecture worked with YAT IT to translate the approved direction into a detailed cloud architecture design — the YAT LMS Cloud Architecture — Baseline Design, available from the Engagement Documents section of the YAT intranet.', 'Assessor text'),
    ('You have returned to MTS, and your assignment is the foundation build phase of the engagement: stand up the supplied architecture in the AWS Academy lab environment, then hand it over to YAT IT with a Deployment Report.', 'Assessor text'),
    ('Scope of your work in this assessment', 'Heading 2'),
    ("Per the LMS Migration Role Brief on the YAT intranet (§ Scope of the MTS consulting engagement), your assessment scope is cloud infrastructure provisioning only. The following activities are YAT in-house IT's responsibility, not yours:", 'Assessor text'),
    ('LMS application installation on the AWS infrastructure you build', 'Assessor text'),
    ('Migration of the MySQL database from on-prem to the AWS RDS instance', 'Assessor text'),
    ('Cutover from legacy to new (DNS switch, parallel running, decommissioning)', 'Assessor text'),
    ('Organisational change management around the cutover (CAB process, end-user communications, training)', 'Assessor text'),
    ('Ongoing application support post-handover', 'Assessor text'),
    ('Your AT2 deliverable stops at infrastructure ready for application deployment. You hand the infrastructure over to YAT IT with the Deployment Report — they take it from there. Do not install the LMS application, migrate production data, or perform cutover activities as part of this assessment.', 'Assessor text'),
]

PART1_ASSESSOR = [
    ('Part 1 — Build the supplied design', 'Heading 2'),
    ("Using AWS Academy, implement the architecture exactly as specified in the supplied design. The design is opinionated where it matters (region, network topology, IAM model, service categories, security controls, tagging) and intentionally silent where you must demonstrate professional judgement. The supplied design's §14 Configuration decisions left to the implementer enumerates the eight decisions you must make and justify (C1 through C8).", 'Assessor text'),
    ('For each of those configuration decisions, make a deliberate, justified choice based on the YAT LMS workload as described in the LMS Application Specification on the YAT intranet. You will document the choice and rationale in §5 of your Deployment Report.', 'Assessor text'),
    ('Important constraints on the build:', 'Assessor text'),
    ('The deployment is single-AZ and non-HA by design. HA hardening is the next phase (AT3). Do not pre-empt that work.', 'Assessor text'),
    ('All security, encryption, tagging, and naming conventions in the supplied design are mandatory. These are non-negotiable.', 'Assessor text'),
    ('Take the screenshots listed in Appendix A of the Deployment Report template as you go, not at the end.', 'Assessor text'),
    ('Test outcomes (§6 of the template) require you to actually run the tests — do not fabricate.', 'Assessor text'),
]

PROSE_PART2_TIPS = [
    ('Part 2 — Produce the Deployment Report', 'Heading 2'),
    ("Download the Deployment Report template from the YAT intranet's Templates section. Populate every section and every appendix. The template provides explicit prompts in each section — follow them.", 'Assessor text'),
    ('Your report covers:', 'Assessor text'),
    ('§1 Executive Summary (write this last)', 'Assessor text'),
    ('§2 Engagement Context (referring back to your AT1 Business Case + the supplied design)', 'Assessor text'),
    ("§3 Scope of Deployment (what's in this phase, what's deferred to AT3)", 'Assessor text'),
    ('§4 Build Narrative (chronological account by layer — IAM, network, compute, ALB, DB, storage, security, monitoring)', 'Assessor text'),
    ('§5 Configuration Decisions (justifications for the eight points the design left to you — C1 through C8)', 'Assessor text'),
    ('§6 Testing and Validation (results from the tests in §6 of the template)', 'Assessor text'),
    ("§7 Operational Handover (information YAT IT needs to operate the environment, including filing the report per YAT's records procedures)", 'Assessor text'),
    ('§8 Knowledge Evidence Responses (six contextual questions about choices in your own build)', 'Assessor text'),
    ('Appendix A — Build Evidence (17 named AWS console screenshots — the template lists them explicitly)', 'Assessor text'),
    ('Appendix B — Configuration Exports (IAM policies, security group rules, VPC config, etc.)', 'Assessor text'),
    ('Appendix C — Test Evidence (logs, screenshots of test outcomes)', 'Assessor text'),
    ('Appendix D — Reflections (two short reflective responses on the build experience)', 'Assessor text'),
    ('Before starting your own report, review at least one Example previous Deployment Report from the YAT intranet\'s Document Archive section. The example shows what a completed report looks like in style, depth, and structure — model your work on it.', 'Assessor text'),
    ('Tips for success', 'Heading 2'),
    ('Review the example past Deployment Report first. It shows what good looks like and saves you guessing about depth and structure.', 'Assessor text'),
    ('Read the supplied design end-to-end before building — including the "out of scope" section (§13) and the "configuration decisions left to the implementer" section (§14). The build is faster when you know the constraints upfront.', 'Assessor text'),
    ('Screenshots as you go, not at the end. This is the single most common failure mode in this assessment.', 'Assessor text'),
    ('Justify every configuration decision against the YAT workload. Generic justifications ("m6i.large is a good choice for general-purpose workloads") are not credible. Specific justifications ("m6i.large handles the expected 200–300 concurrent users with headroom for assessment-window peaks per the LMS application spec, at a unit cost consistent with the AT1 CBA Year-1 budget") are.', 'Assessor text'),
    ("Don't paper over the eight Configuration Decisions. The §5 table is one of the most-scrutinised sections.", 'Assessor text'),
    ('The Knowledge Evidence responses (§8) are about your own build. Every answer should reference specific sections, components, or decisions in your report. Generic textbook answers about cloud concepts won\'t pass.', 'Assessor text'),
    ('The reflections in Appendix D are honest, not promotional. A genuine "here\'s a decision I\'d revise" earns more credit than a varnished "everything was perfect".', 'Assessor text'),
    ("File the report per YAT's records procedures (§7.4 of the template) — this is part of what's being assessed.", 'Assessor text'),
]

# Assessor-only body — Deployment Report Benchmark + UoC coverage reverse-map.
# Each block is ('h1'|'h2'|'p', text) or ('tbl', [[row cells], ...]).
ASSESSOR_BODY = [
    ('h1', 'Deployment Report Benchmark'),
    ('p', 'Per-section marking guidance for the assessor. What "Satisfactory" looks like for each section of the report. Model answers for the contextual KE questions. Common deductions.'),
    ('p', '1. Marking framework'),
    ('p', 'The Deployment Report is marked Satisfactory / Not Yet Satisfactory as a whole, with per-section judgements feeding the overall.'),
    ('p', 'Each report section evidences specific UoC items (see the section-level UoC mapping in the Marking Guide above).'),
    ('p', "A section is Satisfactory when the UoC items it evidences are demonstrated through the student's content."),
    ('p', "Students may arrive at different specific configuration choices (e.g. m6i.large vs m5.large, scaling thresholds 60% vs 70%) — that's acceptable as long as their justifications are internally consistent and tied to the YAT workload."),
    ('p', 'Critical sections (where NYS forces overall NYS regardless of strengths elsewhere): §4 Build Narrative (workload tier — criterion A3), §5 Configuration Decisions (A5), §6 Testing and Validation (A6), §8 Knowledge Evidence (A8), Appendix A Build Evidence (A9).'),
    ('p', 'Supporting sections (where NYS may be balanced against compensating evidence elsewhere): §1 Executive Summary (A1), Appendix D Reflections (A12).'),
    ('p', '§2 Engagement Context + §3 Scope of Deployment are required template sections but not separately graded — they support the rest of the report and any weakness shows up under A13 (Document Quality).'),
    ('p', '2. Marking §1 Executive Summary'),
    ('p', 'UoC evidenced: [ICTCLD401 PC 4.1] (document and communicate work — written executive-level communication evidence).'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'One page or less'),
    ('p', 'Names what was delivered (the AWS foundation for the LMS)'),
    ('p', 'Region + AZ stated'),
    ('p', 'Top 2–3 highlights of the build called out'),
    ('p', 'Explicit acknowledgement that HA is deferred to AT3'),
    ('p', 'Confirmation the engagement is ready for the next phase'),
    ('p', 'NYS:'),
    ('p', 'Missing or absent'),
    ('p', 'Reads as a full report rather than a summary'),
    ('p', 'Does not reconcile with the body (e.g. mentions multi-AZ when the body shows single-AZ)'),
    ('p', '3. Marking §2 Engagement Context'),
    ('p', 'UoC evidenced: scoping context (supports the rest of the report).'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'References the AT1 Business Case engagement and its board-approved action plan'),
    ('p', 'Names the supplied baseline design as the build specification (and notes its approval status)'),
    ('p', 'Names Pat Lin and Sam Walker as the in-scenario stakeholders'),
    ('p', 'Frames the AT2 work as the foundation-build phase'),
    ('p', 'NYS: vague about the prior work, or treats AT2 as a greenfield engagement disconnected from AT1.'),
    ('p', '4. Marking §3 Scope of Deployment'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'Clear list of what is in this build (IAM, network, compute, ALB, DB, storage, autoscaling, baseline monitoring)'),
    ('p', 'Clear list of what is deferred to AT3 (Multi-AZ DB, cross-AZ resilience, failure simulation, DR, cross-Region backup, HA-tuned monitoring)'),
    ('p', 'NYS: scope unclear; HA elements implicitly assumed to be in this phase.'),
    ('p', '5. Marking §4 Build Narrative'),
    ('p', 'UoC evidenced: [ICTCLD401 PC 1.2, 1.4, 1.5, 1.6, 1.7, 2.1, 2.2, 2.3, 2.4, 2.5, 3.1] · [ICTCLD502 PC 1.3, 4.1, 4.3 — partial].'),
    ('p', 'Satisfactory looks like, per sub-section:'),
    ('p', '5.1 §4.1 IAM (criterion A4)'),
    ('p', 'IAM groups named match the supplied design (or rationale provided for any addition)'),
    ('p', 'MFA enforcement evidenced'),
    ('p', 'EC2 instance role / Application-Service permissions described'),
    ('p', 'Cross-references Screenshot A1, A2, A3 + Configuration export B1, B2'),
    ('p', '5.2 §4.2 Network (criterion A4)'),
    ('p', 'VPC CIDR + subnet layout matches supplied design'),
    ('p', 'Internet Gateway + NAT Gateway documented'),
    ('p', 'Route table behaviour documented'),
    ('p', 'Cross-references A4, A5, A6, A7 + B3'),
    ('p', '5.3 §4.3 + §4.4 Compute + ALB (criterion A5)'),
    ('p', 'EC2 instance type stated (per C1 decision)'),
    ('p', 'ASG configuration stated (min/desired/max + scaling policy threshold per C4)'),
    ('p', 'ALB configuration + target group health check'),
    ('p', 'Cross-references A9, A13, A14 + B4, B5'),
    ('p', '5.4 §4.5 Database (criterion A5)'),
    ('p', 'RDS instance class stated (per C2)'),
    ('p', 'MySQL engine version (per C7)'),
    ('p', 'Single-AZ deployment confirmed'),
    ('p', 'Encryption at rest confirmed'),
    ('p', 'Cross-references A12 + B6'),
    ('p', '5.5 §4.6 Storage (criterion A5)'),
    ('p', 'EBS volumes (root + data) sizing documented'),
    ('p', 'S3 buckets (attachments + backups) with block-public-access confirmed'),
    ('p', 'Cross-references A10, A11 + B7'),
    ('p', '5.6 §4.7 Security (criterion A4)'),
    ('p', 'Three-tier security group model documented'),
    ('p', 'Encryption-in-transit between tiers documented'),
    ('p', 'Shared responsibility table from §9.4 of the supplied design referenced'),
    ('p', 'Cross-references A8 + B2'),
    ('p', '5.7 §4.8 Monitoring baseline (criterion A6)'),
    ('p', "Baseline CloudWatch alarms named (per the supplied design's §10.2 list)"),
    ('p', 'Acknowledged that HA-tuned monitoring is deferred to AT3'),
    ('p', 'Cross-references A17'),
    ('p', 'NYS for the build narrative:'),
    ('p', 'Any sub-section missing'),
    ('p', 'Any sub-section that contradicts the screenshot evidence in Appendix A'),
    ('p', '"Build narrative" reads like generic AWS documentation rather than a description of what was actually built'),
    ('p', '6. Marking §5 Configuration Decisions'),
    ('p', 'UoC evidenced: [ICTCLD401 PC 1.1] · [ICTCLD401 PC 1.3]. Critical section.'),
    ('p', 'Satisfactory looks like: All 8 decision points justified with rationale tied to the YAT LMS workload (not generic):'),
    ('tbl', [
        ['#', 'Decision', 'What a credible justification references'],
        ['C1', 'EC2 instance type', 'Concurrent-user load from the LMS application spec (200–300 typical, 500–700 peak); cost envelope from AT1 CBA; growth headroom'],
        ['C2', 'RDS instance class', 'Database workload characteristics (relatively read-heavy for an LMS); concurrent-connection load'],
        ['C3', 'EBS data volume + RDS storage sizing', 'The 178 GB current footprint + ~25 GB/year growth from the LMS app spec; 12-month projection + headroom'],
        ['C4', 'ASG scaling threshold', 'CPU profile expectations during assessment-week peaks; cooldown/cool-off implications'],
        ['C5', 'Permission boundary for MTS-Consultants', 'Specifically the AWS Academy lab access scope'],
        ['C6', 'Bastion / RDP design', 'Security trade-off: bastion host vs SSM Session Manager vs VPN-only'],
        ['C7', 'MySQL engine version', 'DOODLE compatibility (the migration brief says no application upgrade)'],
        ['C8', 'DNS strategy + ACM cert', 'LMS hostname strategy; ACM auto-renewal considerations'],
    ]),
    ('p', 'Acceptable variance: specific choices can differ between students (e.g. m5.large vs m6i.large), as long as the justification holds. The benchmark is the quality of reasoning, not a single correct answer.'),
    ('p', 'NYS:'),
    ('p', 'Any of C1–C8 left blank or with "TBD"'),
    ('p', 'Generic justifications ("m6i.large is a good general-purpose instance") not tied to YAT workload'),
    ('p', 'Justification contradicts the workload (e.g. picking a tiny instance after acknowledging 700 peak concurrent users)'),
    ('p', 'Sizing decisions without calculations shown'),
    ('p', '7. Marking §6 Testing and Validation'),
    ('p', 'UoC evidenced: [ICTCLD401 PC 2.6, 3.2] · [ICTCLD502 PC 4.2, 4.3 partial]. Critical section.'),
    ('p', 'Satisfactory looks like, per test category:'),
    ('tbl', [
        ['Test category', 'Satisfactory evidence'],
        ['6.1 Connectivity', 'All four tests recorded with Pass/Fail; Appendix C C1, C2 evidence attached'],
        ['6.2 Autoscaling', 'Load applied, scale-out occurred, instance entered service, scale-in followed; Appendix C C3 evidence attached'],
        ['6.3 DB connectivity', 'MySQL client connection established from EC2 over private network; version confirmation matches C7 decision; Appendix C C4 evidence attached'],
        ['6.4 End-to-end smoke', 'LMS reachable via ALB DNS name; HTTP success status; Appendix C C5 evidence attached'],
    ]),
    ('p', 'NYS:'),
    ('p', 'Any test category left blank or "not tested"'),
    ('p', 'Tests recorded as Pass without evidence to support'),
    ('p', 'Evidence contradicts the recorded outcome'),
    ('p', 'No negative test (RDS NOT publicly reachable) — this is a deliberate test of security posture'),
    ('p', '8. Marking §7 Operational Handover'),
    ('p', 'UoC evidenced: [ICTCLD401 PC 4.3] (save and store user documentation per organisational policies and procedures).'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'Access information for YAT IT documented (which IAM groups, who has access post-handover)'),
    ('p', 'Runbook references to baseline design + tagging + backup baseline + alarms'),
    ('p', '§7.4 Documentation filing — student names how/where the report was filed per YAT records procedures (refer to the Backup and Retention Policy or the Records Management Policy on the YAT intranet)'),
    ('p', 'NYS:'),
    ('p', '§7.4 filing not addressed (PC 4.3 not evidenced)'),
    ('p', 'Handover section reads as a checklist without actual handover content'),
    ('p', "Known limitations section claims AT2 is HA when it isn't"),
    ('p', '9. Marking §8 Knowledge Evidence Responses'),
    ('p', 'UoC evidenced: [ICTCLD401 KE 5, 6, 7, 8, 9, 10]. Critical section.'),
    ('p', 'Below are model answers for each question. Student answers may differ in detail but must demonstrate the underlying knowledge applied to their own Deployment Report.'),
    ('p', 'Q1 — VM/networking/scaling features [ICTCLD401 KE 5]'),
    ('p', 'Model coverage: Student names each of EC2, ASG, ALB with the feature they provide and explains how the choice supports YAT LMS specifically.'),
    ('p', 'EC2: compute capacity to run the Windows Server 2016 + DOODLE LMS stack; sized for the 200–300 typical / 500–700 peak concurrent user load (per LMS app spec).'),
    ('p', 'ASG: horizontal auto-scaling so that the LMS can scale-out during assessment-window peaks without manual intervention; CPU threshold chosen at X% based on Y reasoning.'),
    ('p', 'ALB: distributes traffic across the ASG members; HTTPS termination via ACM cert; health-check approach ensures unhealthy instances are removed from rotation quickly.'),
    ('p', "NYS: generic feature explanations without ties to YAT; doesn't address all three components; ignores the workload context."),
    ('p', 'Q2 — Vertical vs horizontal scaling; managed services; storage [ICTCLD401 KE 6]'),
    ('p', 'Model coverage:'),
    ('p', '(a) RDS over self-hosted: managed backups, patching, monitoring; reduces operational burden on YAT IT (which lacks cloud expertise per the scenario context); HA-readiness (Multi-AZ is one config change away — coming in AT3).'),
    ('p', '(b) EBS + S3 together: EBS for OS + application binaries + frequently-accessed application data (low latency, block-level); S3 for course attachments + student submissions + backups (cheaper at scale, virtually unlimited, versioning + lifecycle to Glacier supports records retention).'),
    ('p', '(c) Vertical vs horizontal: vertical would require downtime to resize (poor fit for an LMS that must stay available during assessment windows); horizontal via ASG scales without downtime and handles the assessment-week 3× spike pattern. Trade-off: horizontal needs the application to support stateless/sticky-session patterns; LMS sessions may need session affinity at the ALB.'),
    ('p', 'NYS: textbook definitions only; no reference to YAT specifics; missing one or more of the three sub-parts.'),
    ('p', 'Q3 — Shared security responsibility [ICTCLD401 KE 7]'),
    ('p', 'Model coverage:'),
    ('p', "(a) Two YAT responsibilities: e.g. (1) Windows Server 2016 OS patching post-handover; (2) IAM users + MFA management for YAT IT staff. Both fall on YAT's side of the shared responsibility line."),
    ('p', "(b) One responsibility shifted from YAT to AWS: hardware refresh + data centre cooling/power (used to be YAT's server-room overhead; now AWS's). Or: backup-tape rotation (now AWS-managed RDS backups). Or: physical data centre security."),
    ('p', "NYS: copy-paste from supplied design without identifying specific responsibilities; gets the responsibility line wrong (e.g. claims patching is AWS's)."),
    ('p', 'Q4 — User access protocols [ICTCLD401 KE 8]'),
    ('p', 'Model coverage: Student picks one IAM group, describes its permissions, names the job function, explains why permissions differ from another group.'),
    ('p', "Example: YAT-ICT-Admins has read-only on infrastructure + full access to CloudWatch and RDS console + no IAM modifications. Job function: day-to-day operations after handover. Different from MTS-Consultants (full admin during build) because YAT IT shouldn't be modifying IAM (that's a security boundary), and different from Read-Only-Auditors (no service mutations) because YAT IT need to acknowledge/clear alarms."),
    ('p', 'NYS: group described without job function context; permissions stated without rationale.'),
    ('p', 'Q5 — Security policies + network traffic limits [ICTCLD401 KE 9]'),
    ('p', 'Model coverage: Student picks one SG (e.g. sg-db), describes its rules, explains why traffic is restricted that way, names the risk if removed.'),
    ('p', 'Example: sg-db allows MySQL:3306 inbound only from sg-app (no public ingress at all), no outbound. Restricted because the database must only be reachable from the application tier — direct database access from the internet would expose student PII (Privacy Act 1988 risk). Removing the restriction would mean anyone on the internet could attempt to brute-force MySQL credentials.'),
    ('p', 'NYS: SG rules stated without rationale; risk statement absent or generic.'),
    ('p', 'Q6 — Role of DNS [ICTCLD401 KE 10]'),
    ('p', 'Model coverage: Student identifies two DNS resolution points and explains what fails if misconfigured.'),
    ('p', 'Examples:'),
    ('p', "ALB DNS name — end-users browse to the LMS hostname which resolves (via Route 53 or YAT's authoritative DNS) to the ALB DNS name, which AWS resolves to one of the ALB's IPs. If misconfigured, end-users get the wrong IP or no IP → can't reach the LMS."),
    ('p', 'RDS endpoint — the EC2 instance connects to RDS via its DNS endpoint (which AWS uses to abstract away which underlying instance is the primary). If misconfigured (e.g. wrong endpoint in app config), the application can\'t reach the database.'),
    ('p', 'NYS: generic "DNS resolves names to IPs" without YAT-specific resolution points named.'),
    ('p', 'Marking checks for Appendix 2 (§8) overall'),
    ('p', 'Award credit for:'),
    ('p', 'All 6 questions answered'),
    ('p', "Each answer references specific sections of the student's own report"),
    ('p', 'The cited content is actually present in their report'),
    ('p', 'Common deductions:'),
    ('p', 'Generic textbook answers ("IaaS provides infrastructure")'),
    ('p', "Citations to report sections that don't contain what the answer claims"),
    ('p', 'Repeating the supplied design verbatim instead of synthesising'),
    ('p', '10. Marking Appendix A — Build Evidence (Screenshots)'),
    ('p', 'UoC evidenced: [ICTCLD401 PE 1, 2]. Critical section.'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'All 17 screenshots (A1–A17) present'),
    ('p', 'AWS region indicator visible in each (top-right of console)'),
    ('p', 'Named items visible per the template\'s screenshot description'),
    ('p', 'Screenshots are from the actual deployed environment (not Photoshopped, not from AWS marketing material)'),
    ('p', 'Common deductions:'),
    ('p', "Region indicator not visible — can't verify ap-southeast-2"),
    ('p', "Screenshot doesn't show the named items (e.g. A14 doesn't show the launch template attached to the ASG)"),
    ('p', "Screenshot from a tutorial / shared resource rather than the student's own deployment"),
    ('p', '11. Marking Appendix B — Configuration Exports'),
    ('p', 'UoC evidenced: [ICTCLD401 PE 2].'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'All 7 exports (B1–B7) present'),
    ('p', 'Exports are consistent with the screenshots in Appendix A (e.g. the security group rules in B2 match the rules visible in A8)'),
    ('p', 'Exports are recognisably from the AWS environment (JSON/text dump format consistent with aws ec2 describe-* output)'),
    ('p', 'NYS:'),
    ('p', 'Exports missing'),
    ('p', 'Exports inconsistent with screenshots (suggests fabrication)'),
    ('p', "Generic example configurations rather than the student's own"),
    ('p', '12. Marking Appendix C — Test Evidence'),
    ('p', 'UoC evidenced: [ICTCLD401 PE 3] · [ICTCLD502 PC 4.2, 4.3].'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'All 6 evidence items (C1–C6) present'),
    ('p', 'Evidence is consistent with the outcomes recorded in §6'),
    ('p', 'Test C6 (negative test — RDS not publicly reachable) explicitly demonstrated'),
    ('p', 'NYS:'),
    ('p', 'Evidence missing for one or more tests'),
    ('p', 'Evidence shows the opposite of the recorded §6 outcome'),
    ('p', 'Negative test missing — security posture not verified'),
    ('p', '13. Marking Appendix D — Reflections'),
    ('p', 'UoC evidenced: [ICTCLD401 FS Learning, Planning and organising, Self-management skills].'),
    ('p', 'UoC evidenced (foundation skill — closest-fit): [ICTCLD401 FS Reading] (A13).'),
    ('p', 'Satisfactory looks like:'),
    ('p', 'Both reflections present'),
    ('p', "Each is specific to the student's own build (not generic)"),
    ('p', 'R1 names a transferable lesson and how it was arrived at'),
    ('p', 'R2 names a decision validated AND a decision the student would revise'),
    ('p', 'NYS:'),
    ('p', 'Reflections absent or one-sentence-each'),
    ('p', 'All reflections positive — no honest "would revise" content'),
    ('p', 'Generic reflections that don\'t engage with the actual build (could be cut-pasted into any deployment report)'),
    ('h1', 'UoC coverage verification (reverse map)'),
    ('p', 'This table closes the loop on bidirectional traceability: every UoC requirement AT2 claims to evidence is named below with the marking criterion (or criteria) that evidence it. No UoC requirement claimed by AT2 is left without a criterion.'),
    ('p', 'ICTCLD401 — Configure cloud services (AT2-evidenced items)'),
    ('tbl', [
        ['UoC item', 'Evidenced by criterion(ia)'],
        ['PC 1.1 — Discuss and compare different cloud computing solutions, models and services', 'A5 (justifications in §5 Configuration Decisions implicitly compare options)'],
        ['PC 1.2 — Identify impact of shared security responsibility models', 'A4'],
        ['PC 1.3 — Select best cloud computing solution and service', 'A5'],
        ['PC 1.4 — Access account on cloud platform', 'A2'],
        ['PC 1.5 — Identify user access protocols and policies', 'A2'],
        ['PC 1.6 — Configure access functions within cloud environment', 'A2'],
        ['PC 1.7 — Identify and assign security responsibilities', 'A2'],
        ['PC 2.1 — Create users and groups', 'A2'],
        ['PC 2.2 — Create virtual multi-tier network', 'A2'],
        ['PC 2.3 — Create virtual machine', 'A3'],
        ['PC 2.4 — Define, add and expand storage', 'A3'],
        ['PC 2.5 — Deploy a managed database', 'A3'],
        ['PC 2.6 — Test external network access', 'A6'],
        ['PC 3.1 — Configure and apply autoscaling', 'A4'],
        ['PC 3.2 — Test automatic scaling', 'A6'],
        ['PC 4.3 — Save and store user documentation per organisational policies', 'A7'],
        ['PE 1 — Build at least one simple virtual network capable of supporting a workload', 'A9'],
        ['PE 2 — Configure compute, storage, database and autoscaling resources within virtual network', 'A9 · A10'],
        ['PE 3 — Conduct simple tests to confirm access to resources', 'A11'],
        ['KE 5 — VM/networking/scaling features (VM sizing, load balancing, autoscaling, monitoring, storage backups, virtual networks/traffic routing)', 'A8 (Q1)'],
        ['KE 6 — Vertical vs horizontal scaling; VM vs physical; RDBMS/DW/NoSQL; self-hosted vs managed vs cloud-native DB; storage options (block/object/archive/network filesystems)', 'A8 (Q2)'],
        ['KE 7 — User, business and vendor responsibilities according to shared security responsibility models', 'A8 (Q3)'],
        ['KE 8 — User access protocols and policies according to organisation hierarchy and job function', 'A8 (Q4)'],
        ['KE 9 — Security policies, protocols and mechanisms as they relate to cloud (network traffic limits + security responsibilities per work function/user access)', 'A8 (Q5)'],
        ['KE 10 — Purpose of DNS for connecting remote servers when web browsing', 'A8 (Q6)'],
        ['FS Reading', 'A13'],
        ['FS Writing', 'A13'],
        ['FS Learning', 'A12'],
        ['FS Planning and organising', 'A12'],
        ['FS Self-management skills', 'A12'],
        ['AC 1 — Cloud vendor service provider', 'C1 (pre-condition)'],
        ['AC 2 — Cloud managed database service', 'C1 (pre-condition)'],
        ['AC 3 — Internet and web browser', 'C1 (pre-condition)'],
        ['AC 4 — Data to gather information from to determine output and user requirements', 'C2 (pre-condition)'],
    ]),
    ('p', '(Note: 401 PCs not listed above (PC 1.8, PC 4.1, PC 4.2) are evidenced in AT1. 401 KEs not listed (KE 1, 2, 3, 4, 11) are evidenced in AT1 Appendix 2.)'),
    ('p', 'ICTCLD502 — Design and implement highly-available cloud infrastructure (AT2-evidenced items only)'),
    ('tbl', [
        ['UoC item', 'Evidenced by criterion(ia)'],
        ['PC 1.3 — Identify level of shared security responsibility models according to business needs', 'A2 (co-evidenced with [ICTCLD401 PC 1.2] · [ICTCLD401 PC 1.7] via Group 2 of consolidated UoC)'],
        ['PC 4.1 — Implement architecture design in cloud environment', 'A3'],
        ['PC 4.2 — Demonstrate connectivity between resources at all tiers', 'A6 · A11'],
        ['PC 4.3 — Monitor and measure availability of resources', 'A4 · A11 (baseline monitoring only — HA-tuned monitoring is AT3)'],
        ['FS Reading', 'A13'],
        ['AC 1 — Cloud vendor service provider', 'C1 (pre-condition)'],
        ['AC 2 — Cloud managed database service', 'C1 (pre-condition)'],
        ['AC 3 — Information and data sources required to design and implement cloud infrastructure', 'C2 (pre-condition)'],
        ['AC 4 — Integrated development environment (IDE)', 'C1 (pre-condition — AWS Console + CLI counts as IDE)'],
        ['AC 5 — Specific requirements + industry standards + organisational procedures + legislative requirements + business and functionality requirements', 'C2 (pre-condition — supplied via YAT intranet)'],
        ['AC 6 — Internet and web browser', 'C1 (pre-condition)'],
        ['AC 7 — Secure shell (SSH) or remote desktop protocol (RDP) client', 'C1 (pre-condition — RDP for Windows EC2 instances)'],
        ['AC 8 — Data to gather information from to determine output and user requirements', 'C2 (pre-condition)'],
    ]),
    ('p', '(Note: 502 PCs not listed above (PC 1.1, 1.2, 2.1–2.5, 3.1–3.5, 4.4–4.6, 5.1, 5.2, 5.3) are evidenced in AT1 (PC 1.2, 5.2) or AT3 (the rest). 502 KEs not listed (KE 1, 2, 3) are evidenced in AT1 Appendix 2; KEs 4–9 are evidenced in AT3.)'),
]


# ---------- build helpers ----------

STYLE = {"h1": "Heading 1", "h2": "Heading 2", "p": "Normal"}


def add_marking_row(table, lines):
    """A marking-guide criterion row: criterion (+ UoC line) in col 0, the check box in col 1."""
    row = table.add_row()
    set_cell_content(row.cells[0], lines)
    set_cell_content(row.cells[1], CHECK)


def delete_body_paragraph(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            p._element.getparent().remove(p._element)
            return


def render_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    for r, cells in enumerate(rows):
        for c, val in enumerate(cells):
            set_cell_content(t.rows[r].cells[c], val)
    doc.add_paragraph()


def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details ----
    t_details = doc.tables[0]
    set_cell_content(t_details.rows[1].cells[1], DETAILS["qualification"])
    set_cell_content(t_details.rows[2].cells[1], DETAILS["units"])
    set_cell_content(t_details.rows[3].cells[1], DETAILS["task_title"])
    set_cell_content(t_details.rows[4].cells[1], DETAILS["task_number"])

    # ---- Table 1: Teacher/Assessor instructions ----
    # (Second attempt + Assessment retention rows keep the template's standard text.)
    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Time allowed"), "")
    set_cell_content(find_instruction_row(t_instr, "Location"), "")
    set_cell_content(find_instruction_row(t_instr, "Resources required"), RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA)
    # add a Conditions row at the end (matching the table style), as CL1/CL2 do
    cond_row = t_instr.add_row()
    set_cell_content(cond_row.cells[0], "Assessment Conditions & Setup Requirements")
    for r in cond_row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell_content(cond_row.cells[1], CONDITIONS)

    # ---- Table 2: Marking Guide (single-part A1-A13; the 'Project ' heading + intro line stay) ----
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)  # keep 'Assessment criteria' + 'Criteria | Satisfactory?' header rows
    for lines in MARKING:
        add_marking_row(t_mark, lines)

    # ---- drop the template's trailing marking-guide boilerplate ----
    delete_body_paragraph(doc, "Add or delete rows as required")
    delete_body_paragraph(doc, "If questioning or observation is incorporated into this assessment task, "
                                "you can incorporate a Practical Observation Checklist.")

    # ---- Instructions to Student (shared prose; assessor Part 1 variant) ----
    doc.add_paragraph("Instructions to Student", style="Heading 1")
    for text, style in (PROSE_INTRO + PART1_ASSESSOR + PROSE_PART2_TIPS):
        doc.add_paragraph(text, style=style)

    # ---- Assessor-only body (Benchmark + reverse-map) ----
    for kind, payload in ASSESSOR_BODY:
        if kind == "tbl":
            render_table(doc, payload)
        else:
            doc.add_paragraph(payload, style=STYLE[kind])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/assessments/AT2/AT2-Deployment-Assessor.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
