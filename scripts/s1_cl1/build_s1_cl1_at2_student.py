#!/usr/bin/env python3
"""Build the S1-CL1 AT2 STUDENT instrument (.docx) by populating the Kangan template.

The student-facing version of AT2. It imports the assessor module (build_s1_cl1_at2_assessor)
to single-source the shared content — DETAILS, the CHECK box, and the identical parts of the
'Instructions' prose (the engagement intro, Part 2 and Tips blocks) — and defines its own
student-facing wording for the instruction table, the marking criteria (second person, no UoC
traceability), and the Part 1 build block (which the student copy phrases with extra detail). It
loads the official Kangan 'Project Assessment - Student.docx' template. The assessor-only material
— the Deployment Report Benchmark and the UoC reverse-map — is NOT included.

EVERGREEN: emits NO scenario-site URL and imports no brand.WEBSITE_URL — where the authored copy
named the intranet URL, only the URL and its immediate wrapper are dropped.

Usage:  python scripts/s1_cl1/build_s1_cl1_at2_student.py [output.docx]
Default: S1-CL1-Cloud-Design-Build/assessments/AT2/AT2-Deployment-Student.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_s1_cl1_at2_assessor as a  # noqa: E402  (shared content — single source of truth)
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_tables import clear_table_rows, find_instruction_row, set_cell_content  # noqa: E402

from docx import Document  # noqa: E402

TEMPLATE = str(Path(__file__).resolve().parents[2] / "kangan-templates" / "Project Assessment - Student.docx")


# ---------- student-facing content (second person; no UoC traceability) ----------

# NOTE (evergreen): the committed doc named the intranet URL — " (https://www.placeholder.com.au)".
OVERVIEW = [
    'You are being assessed on the implementation of a supplied AWS cloud architecture for the YAT LMS migration, and the production of a Deployment Report documenting your build.',
    'This is an open-book practical assessment. You may use the YAT intranet, AWS Academy lab environments, AWS documentation, course reference materials, and external research (which must be cited) throughout.',
    'AT2 is the second of three assessment tasks in the S1-CL1 Cloud Design and Build cluster. It builds on AT1 (the Business Case engagement) and feeds into AT3 (HA hardening and project closure). You continue in the same MTS consultant role across all three.',
    'Submission: the completed Deployment Report (.docx) with all appendices populated, submitted via the LMS.',
    'The assessment will not proceed if for any reason it is not safe to do so. The assessor must advise you of the reason for suspending the assessment, what safety action should be taken, and of revised arrangements when it is safe to resume.',
    'There is zero tolerance for plagiarism, cheating and collusion. You will be expected to make a declaration that all work is your own prior to submission. Refer to the Training and Assessment Policy for further information.',
]

TASKS = [
    "Following the YAT board's approval of the action plan in your AT1 Business Case engagement, you took a period of planned annual leave. During that time MTS Senior Architecture worked with YAT IT to translate the approved direction into a detailed technical design — the YAT LMS Cloud Architecture — Baseline Design — which has been approved by Pat Lin and Sam Walker and is now your build specification.",
    'You have returned to MTS to lead the foundation-build phase of the engagement. Your task has two parts that combine into a single deliverable:',
    'Implement the supplied AWS architecture for the YAT LMS migration foundation build, in the AWS Academy lab environment authorised for this engagement.',
    'Produce a Deployment Report documenting your build, using the YAT-provided Deployment Report template. The report includes Configuration Decision justifications (where the supplied design left choices to you), testing outcomes, an operational handover for YAT IT, written responses to six Knowledge Evidence questions about your own build, and four appendices of evidence (build screenshots, configuration exports, test results, reflections).',
    "The Deployment Report is your single submitted deliverable. All build evidence (screenshots, configuration exports, test results) is captured in the report's appendices — there is no separate portfolio submission.",
    'This phase is the foundation build only. The architecture is intentionally non-HA at this stage — HA hardening is the next phase (AT3) and is out of scope here.',
]

# NOTE (evergreen): the committed doc named the intranet URL in the first header — " (download from the YAT intranet)"
# does NOT contain the URL; there is no URL in the student Resources list, so nothing is dropped here.
RESOURCES = [
    'Provided to you (download from the YAT intranet)',
    'YAT LMS Cloud Architecture — Baseline Design (intranet Engagement Documents section) — the design you implement',
    'Deployment Report template (intranet Templates section) — the template you fill in',
    'Example previous Deployment Report (intranet Document Archive section) — review at least one before starting your own',
    'All other scenario materials (LMS application spec, current ICT environment, organisational policies, your AT1 Business Case)',
    'Provided externally',
    'AWS Academy lab access — Cloud Foundations [104469] + Cloud Architecting [172221]',
    'You supply',
    'Computer with web browser',
    'Word-processing software (Microsoft Word or equivalent)',
    'A screenshot tool',
]

CRITERIA = [
    'To receive a Satisfactory result for AT2 you must:',
    'Achieve Satisfactory on every criterion in the Assessment Criteria table (below)',
    'Submit a completed Deployment Report (.docx) with every section and every appendix populated, using the YAT-provided Deployment Report template',
]

RESULTS = [
    'If you are deemed not satisfactory for this assessment, you will be given one (1) more attempt at this assessment (or part thereof), or your teacher/assessor will negotiate a further assessment with you. The second attempt must be completed within 10 working days from the date your feedback is given.',
]

SUBMIT = [
    'Submit the completed Deployment Report (.docx) to the LMS by the due date. The report includes all four appendices populated; no separate files are submitted.',
]

# Single-part marking criteria (A1-A13; second person; no UoC traceability line).
MARKING = [
    'A1 - §1 Executive Summary — you produced a concise (≤ 1 page) summary of what was delivered, with explicit acknowledgement that HA hardening is deferred to AT3, and numbers/components reconcile with the body of the report',
    'A2 - §4.1 + §4.2 + §4.7 Build narrative — Foundation tier (IAM, network, security) — you described the IAM model, network topology, and security-group model, with cross-references to Appendix A screenshots and Appendix B configuration exports',
    'A3 - §4.3 + §4.4 + §4.5 + §4.6 Build narrative — Workload tier (compute, load balancing, database, storage) — you described the EC2 + ALB + ASG, the RDS managed database, and the EBS + S3 storage',
    'A4 - §4.8 Build narrative — Operability (autoscaling configuration, baseline monitoring) — you described the ASG scaling policy and the baseline CloudWatch alarms; you identified the shared security responsibility outcome of your build',
    'A5 - §5 Configuration Decisions — you justified all 8 decision points (C1–C8 from the supplied design §14) with rationale tied to the YAT LMS workload',
    'A6 - §6 Testing and Validation — you ran the four required test categories (connectivity, autoscaling, database connectivity, end-to-end smoke); results are documented with cross-reference to Appendix C evidence',
    "A7 - §7 Operational Handover — you documented access information for YAT IT, named known limitations carried into AT3, and confirmed filing of the report per YAT's documented records procedures",
    'A8 - §8 Knowledge Evidence Responses — you answered all 6 KE questions with specific reference to your own build, not generic textbook answers',
    "A9 - Appendix A Build Evidence — all 17 named screenshots (A1–A17) are present, the AWS region indicator is visible, and the named items are visible per the template's requirements",
    'A10 - Appendix B Configuration Exports — all 7 named exports (B1–B7) are present, exported from the actual deployed environment',
    'A11 - Appendix C Test Evidence — test outcomes are documented with supporting evidence; outcomes are consistent with §6',
    'A12 - Appendix D Reflections — two honest reflective responses are present, both specific to your own build experience',
    'A13 - Document quality — the Deployment Report uses plain English, appropriate technical vocabulary, structure, formatting, and depth relevant to a professional consulting deliverable for an RTO client',
]

# The Part 1 build block — the student copy carries extra detail vs the assessor's. The intro and
# Part 2/Tips blocks are single-sourced from the assessor module (identical text).
PART1_STUDENT = [
    ('Part 1 — Build the supplied design', 'Heading 2'),
    ("Using AWS Academy, implement the architecture exactly as specified in the supplied design. The design is opinionated where it matters — region, network topology, IAM model, service categories, security controls, tagging — and intentionally silent where you must demonstrate professional judgement (specific instance types, sizing, scaling thresholds, etc.). The supplied design's §14 Configuration decisions left to the implementer enumerates the eight decisions you must make and justify (C1 through C8).", 'Assessor text'),
    ('For each of those configuration decisions, make a deliberate, justified choice based on the YAT LMS workload as described in the LMS Application Specification on the YAT intranet. You will document the choice and rationale in §5 of your Deployment Report.', 'Assessor text'),
    ('Important constraints on the build:', 'Assessor text'),
    ('The deployment is single-AZ and non-HA by design. HA hardening is the next phase (AT3). Do not pre-empt that work — Multi-AZ database, cross-AZ subnets, cross-Region backup copies, failure simulation, DR runbook are all out of scope for this phase.', 'Assessor text'),
    ('All security, encryption, tagging, and naming conventions in the supplied design are mandatory. These are non-negotiable.', 'Assessor text'),
    ("Take the screenshots listed in Appendix A of the Deployment Report template as you go, not at the end. Trying to recreate the state for a screenshot after you've moved on is painful and sometimes impossible.", 'Assessor text'),
    ('Test outcomes (§6 of the template) require you to actually run the tests — do not fabricate. Screenshot the test evidence as you run.', 'Assessor text'),
]


def build(path):
    doc = Document(TEMPLATE)

    # ---- Table 0: Details (student name / ID / assessor name+email left blank to fill) ----
    t_details = doc.tables[0]
    set_cell_content(t_details.rows[2].cells[1], a.DETAILS["qualification"])
    set_cell_content(t_details.rows[3].cells[1], a.DETAILS["units"])
    set_cell_content(t_details.rows[4].cells[1], a.DETAILS["task_title"])
    set_cell_content(t_details.rows[5].cells[1], a.DETAILS["task_number"])

    # ---- Table 1: Student instructions ----
    t_instr = doc.tables[1]
    set_cell_content(find_instruction_row(t_instr, "Assessment overview"), OVERVIEW)
    set_cell_content(find_instruction_row(t_instr, "Task"), TASKS)
    set_cell_content(find_instruction_row(t_instr, "Time allowed"), "")
    set_cell_content(find_instruction_row(t_instr, "Location"), "")
    set_cell_content(find_instruction_row(t_instr, "Resources required"), RESOURCES)
    set_cell_content(find_instruction_row(t_instr, "Assessment criteria"), CRITERIA)
    set_cell_content(find_instruction_row(t_instr, "Results"), RESULTS)
    # 'Important information' is left as the template's standard text.
    # add a 'How To Submit' row at the end (matching the table style)
    submit_row = t_instr.add_row()
    set_cell_content(submit_row.cells[0], "How To Submit")
    for r in submit_row.cells[0].paragraphs[0].runs:
        r.bold = True
    set_cell_content(submit_row.cells[1], SUBMIT)

    # ---- drop the template's marking-guide intro line (the authored student copy omits it) ----
    for p in doc.paragraphs:
        if p.text.strip() == "List the criteria the student will be assessed against for the project assessment.":
            p._element.getparent().remove(p._element)
            break

    # ---- Table 2: Assessment criteria (single-part A1-A13; no UoC/benchmark) ----
    t_mark = doc.tables[2]
    clear_table_rows(t_mark, 2)  # keep the two header rows
    for line in MARKING:
        row = t_mark.add_row()
        set_cell_content(row.cells[0], line)
        set_cell_content(row.cells[1], a.CHECK)

    # ---- drop the template's trailing marking-guide boilerplate ----
    for text in ("Add or delete rows as required",
                 "If questioning or observation is incorporated into this assessment task, "
                 "you can incorporate a Practical Observation Checklist."):
        for p in doc.paragraphs:
            if p.text.strip() == text:
                p._element.getparent().remove(p._element)
                break

    # ---- Detailed task instructions (shared prose; NO marking benchmark / UoC) ----
    doc.add_paragraph("Instructions", style="Heading 1")
    for text, style in (a.PROSE_INTRO + PART1_STUDENT + a.PROSE_PART2_TIPS):
        doc.add_paragraph(text, style=style)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL1-Cloud-Design-Build/assessments/AT2/AT2-Deployment-Student.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
