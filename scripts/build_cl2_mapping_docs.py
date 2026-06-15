#!/usr/bin/env python3
"""Build the S1-CL2 per-UoC Assessment Mapping docx files (ICTCLD501/503/505).

Mirrors the CL1 mapping docs: the institutional "Assessment Mapping Tool" template,
filled with each unit's items (Details, AT descriptions, and the AC/PC/PE/KE/FS rows in
the description columns) and the AT1/AT2 columns carrying the criterion code(s) that
evidence each item.

Two data sources, both already in the repo (no hand-typed UoC text):
  * item text + order  — parsed from units_of_competency/<UNIT>_Complete_R1.md
  * item -> AT/criterion — inverted from the AT1 + AT2 assessor benchmarks
    (build_s1_cl2_at1_assessor.BENCHMARK, build_s1_cl2_at2_assessor.BENCHMARK), with FS/AC
    handled with judgement (the benchmarks don't tag AC, and FS coverage is partial).

Staged: `--dump <unit>` prints the parsed items + the inverted mapping for review;
generation of the docx is layered on once the data checks out.

USAGE:
    python scripts/build_cl2_mapping_docs.py --dump 501
    python scripts/build_cl2_mapping_docs.py --dump all
"""
import copy
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_s1_cl2_at1_assessor as a1   # noqa: E402
import build_s1_cl2_at2_assessor as a2   # noqa: E402

from docx import Document  # noqa: E402

REPO = Path(__file__).resolve().parent.parent
UOC = REPO / "S1-CL2-Cloud-Disaster-Recovery" / "units_of_competency"
MAPPINGS = REPO / "S1-CL2-Cloud-Disaster-Recovery" / "mappings"
TEMPLATE = REPO / "templates" / "Assessment Mapping Tool.docx"

AT1_TITLE = "Cloud Expansion: Design & DR Plan"
AT2_TITLE = "Cloud Microservice & IaC Implementation"

# FS closest-fit map (CL1 precedent: map to where the skill is genuinely exercised;
# the assessor instruments are NOT updated to claim these — noted inconsistency).
# Reading/Writing -> document/review criteria; Oral comm -> presentation/sign-off;
# Self-management/Planning/Problem solving -> the design/analysis/troubleshoot criteria.
FS_MAP = {
    "ICTCLD501": {
        "Reading": {"AT1": "B1"},
        "Oral communication": {"AT1": "C6"},
        "Self-management": {"AT1": "B8"},
        "Planning and organising": {"AT1": "B12"},
        "Problem solving": {"AT1": "B5"},
    },
    "ICTCLD503": {
        "Reading": {"AT1": "A3", "AT2": "D2"},
        "Writing": {"AT1": "A11, A13", "AT2": "D13"},
        "Problem solving": {"AT1": "A8", "AT2": "D5"},
        "Self-management": {"AT1": "A7", "AT2": "D8"},
    },
    "ICTCLD505": {
        "Oral communication": {"AT2": "D11"},
        "Reading": {"AT2": "D2"},
        "Writing": {"AT2": "D10, D13"},
        "Problem solving": {"AT2": "D5"},
        "Self-management": {"AT2": "D8"},
    },
}

# AC -> assessment Conditions map (ACs map to the instrument Conditions / scenario
# access, not the marking criteria). One entry per AC in source order. The final
# assessor-requirement AC is the tick row.
AC_MAP = {
    "ICTCLD501": [
        {"AT1": "C1"},          # data to assess risk events (scenario)
        {"AT1": "C1"},          # legislation applicable (scenario / residency reqs)
        {"AT1": "C2"},          # reporting standards for the DR plan (templates)
        {"AT1": "✓"},      # assessor requirement
    ],
    "ICTCLD503": [
        {"AT1": "C3", "AT2": "C2"},   # cloud vendor service provider (lab)
        {"AT1": "C3", "AT2": "C2"},   # cloud managed database service (lab)
        {"AT2": "C2"},                # cloud serverless environment (lab; build = AT2)
        {"AT2": "C3"},                # pre-prepared code elements (provided artefacts)
        {"AT1": "C1", "AT2": "C1"},   # information and data sources (scenario)
        {"AT1": "C3", "AT2": "C2"},   # IDE (lab)
        {"AT1": "C1", "AT2": "C1"},   # specific requirements / standards (scenario)
        {"AT1": "C3", "AT2": "C2"},   # internet and web browser (lab)
        {"AT1": "C1", "AT2": "C1"},   # data re user requirements (scenario)
        {"AT1": "✓", "AT2": "✓"},  # assessor requirement
    ],
    "ICTCLD505": [
        {"AT2": "C2"},          # cloud vendor service provider (lab)
        {"AT2": "C2"},          # IaC service (lab)
        {"AT2": "C1"},          # specific requirements / standards (scenario)
        {"AT2": "C1"},          # information and data sources (scenario)
        {"AT2": "C2"},          # IDE (lab)
        {"AT2": "C2"},          # internet and web browser (lab)
        {"AT2": "C2"},          # SSH / RDP client (lab)
        {"AT2": "C2"},          # cloud console / SDK / CLI (lab)
        {"AT2": "✓"},      # assessor requirement
    ],
}

UNITS = {
    "501": ("ICTCLD501", "Develop cloud disaster recovery plans"),
    "503": ("ICTCLD503", "Implement web-scale cloud infrastructure"),
    "505": ("ICTCLD505", "Implement cloud infrastructure with code"),
}


# ----------------------------------------------------------------------------
# Source .md parser
# ----------------------------------------------------------------------------

def _sections(text):
    """Split the markdown into {heading: body} by top-level '# ' headings."""
    out, cur, buf = {}, None, []
    for line in text.splitlines():
        m = re.match(r"^#\s+(.*)", line)
        if m:
            if cur is not None:
                out.setdefault(cur, "\n".join(buf))
            cur, buf = m.group(1).strip(), []
        else:
            buf.append(line)
    if cur is not None:
        out.setdefault(cur, "\n".join(buf))
    return out


def _bullets(body):
    """Return top-level '- ...' items in order. Indented sub-bullets are folded
    into their parent item (parent + children), matching the consolidated-UoC
    discipline of quoting a nested KE as a single item under the parent's tag."""
    items = []
    for ln in body.splitlines():
        stripped = ln.strip()
        if not stripped.startswith("- "):
            continue
        text = re.sub(r"^-\s+", "", stripped).strip()
        indented = len(ln) - len(ln.lstrip()) > 0
        if indented and items:
            items[-1] = items[-1] + "\n  - " + text
        else:
            items.append(text)
    return items


def parse_source_unit(unit_code):
    text = (UOC / f"{unit_code}_Complete_R1.md").read_text(encoding="utf-8")
    secs = _sections(text)

    # --- PCs: Elements & Performance Criteria table ---
    pcs = []          # [(pc_num, pc_text, element_text)]
    epc = secs.get("Elements and Performance Criteria", "")
    for ln in epc.splitlines():
        if not ln.strip().startswith("|"):
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if len(cells) < 2:
            continue
        elem, pc_cell = cells[0], cells[1]
        if elem.upper().startswith("ELEMENTS") or elem.startswith("Elements describe"):
            continue
        if not re.match(r"^\d+\.", elem):
            continue
        for piece in pc_cell.split("<br>"):
            piece = piece.strip()
            m = re.match(r"^(\d+\.\d+)\s+(.*)", piece)
            if m:
                pcs.append((m.group(1), m.group(2).strip(), elem))

    # --- FS: Foundation Skills table ---
    fss = []          # [(skill, description)]
    fsec = secs.get("Foundation Skills", "")
    for ln in fsec.splitlines():
        if not ln.strip().startswith("|"):
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if len(cells) < 2:
            continue
        skill, desc = cells[0], cells[1]
        if skill.upper() == "SKILL" or not skill or skill.startswith("---"):
            continue
        fss.append((skill, desc.replace("<br>", " ")))

    # --- PE / KE: bullets ---
    pes = _bullets(secs.get("Performance Evidence", ""))
    kes = _bullets(secs.get("Knowledge Evidence", ""))

    # --- AC: access bullets + the assessor-requirements statement ---
    acsec = secs.get("Assessment Conditions", "")
    acs = _bullets(acsec)
    for ln in acsec.splitlines():
        if ln.strip().lower().startswith("assessors of this unit must satisfy"):
            acs.append(ln.strip())
    return {"pcs": pcs, "fss": fss, "pes": pes, "kes": kes, "acs": acs}


# ----------------------------------------------------------------------------
# Benchmark inverter: (unit, section, item) -> set of criterion codes
# ----------------------------------------------------------------------------

def _expand(section, numbering):
    """Expand a numbering fragment into individual item ids."""
    out = []
    for part in numbering.split(","):
        part = part.strip().replace("–", "-").replace("—", "-")
        if not part:
            continue
        if "-" in part and section != "FS":
            a, b = (x.strip() for x in part.split("-", 1))
            if section == "PC" and "." in a and "." in b:
                ea, ia = a.split("."); eb, ib = b.split(".")
                if ea == eb:
                    out += [f"{ea}.{i}" for i in range(int(ia), int(ib) + 1)]
                else:
                    out += [a, b]
            else:
                try:
                    out += [str(i) for i in range(int(a), int(b) + 1)]
                except ValueError:
                    out += [a, b]
        else:
            out.append(part)
    return out


def _parse_tags(s):
    """Yield (unit, section, item) from a benchmark UoC string (inherits unit prefix)."""
    cur_unit = None
    for m in re.finditer(r"\[([^\]]+)\]", s):
        inner = m.group(1).strip()
        um = re.match(r"(ICTCLD\d{3})\s+(.*)", inner)
        if um:
            cur_unit, rest = um.group(1), um.group(2)
        else:
            rest = inner
        sm = re.match(r"(PC|PE|KE|AC|FS)\s+(.*)", rest)
        if not sm or cur_unit is None:
            continue
        section, numbering = sm.group(1), sm.group(2).strip()
        if section == "FS":
            yield (cur_unit, "FS", numbering)
        else:
            for item in _expand(section, numbering):
                yield (cur_unit, section, item)


def invert_benchmarks():
    """Return {(unit, section, item): [criteria...]} across AT1 + AT2."""
    out = {}
    for bench in (a1.BENCHMARK, a2.BENCHMARK):
        for _part_title, rows in bench:
            for cid, uoc in rows:
                for tag in _parse_tags(uoc):
                    out.setdefault(tag, [])
                    if cid not in out[tag]:
                        out[tag].append(cid)
    return out


# ----------------------------------------------------------------------------
# Dump (verification)
# ----------------------------------------------------------------------------

def dump(unit_key):
    code, title = UNITS[unit_key]
    items = parse_source_unit(code)
    inv = invert_benchmarks()
    print(f"\n================  {code}  {title}  ================")
    print(f"counts: PC={len(items['pcs'])} FS={len(items['fss'])} "
          f"PE={len(items['pes'])} KE={len(items['kes'])} AC={len(items['acs'])}")

    def codes_for(section, item):
        return ", ".join(inv.get((code, section, item), [])) or "-- (unmapped)"

    print("\n-- PCs --")
    for num, txt, _elem in items["pcs"]:
        print(f"  {num:<5} [{codes_for('PC', num):<12}] {txt[:70]}")
    print("\n-- PEs --")
    for i, txt in enumerate(items["pes"], 1):
        print(f"  PE{i}  [{codes_for('PE', str(i)):<12}] {txt[:70]}")
    print("\n-- KEs --")
    for i, txt in enumerate(items["kes"], 1):
        print(f"  KE{i}  [{codes_for('KE', str(i)):<12}] {txt[:70]}")
    print("\n-- FSs --")
    for skill, _desc in items["fss"]:
        print(f"  [{codes_for('FS', skill):<12}] {skill}")
    print("\n-- ACs --")
    for i, txt in enumerate(items["acs"], 1):
        print(f"  AC{i}  {txt[:80]}")


# ----------------------------------------------------------------------------
# docx generation
# ----------------------------------------------------------------------------

# Table indices (confirmed against the blank template)
T_DETAILS, T_DESC, T_AC, T_PC, T_PE, T_KE, T_FS = 1, 2, 3, 4, 5, 6, 7


def _put(cell, text):
    """Replace a cell's content with `text` (newlines become separate paragraphs)."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r.text = ""
    lines = (text or "").split("\n")
    (p0.runs[0] if p0.runs else p0.add_run("")).text = lines[0]
    for ln in lines[1:]:
        cell.add_paragraph().add_run(ln)


def _reset(table, n_header=2):
    """Remove data rows after the headers; return a deep-copied template row XML."""
    rows = list(table.rows)
    template_tr = copy.deepcopy(rows[n_header]._tr)
    for row in rows[n_header:]:
        row._tr.getparent().remove(row._tr)
    return template_tr


def _add(table, template_tr):
    table._tbl.append(copy.deepcopy(template_tr))
    return table.rows[-1]


def _split_codes(crits):
    """Split a criterion-code list into (AT1, AT2) strings. A/B/C -> AT1, D -> AT2."""
    at1 = [c for c in crits if c[:1] in ("A", "B", "C")]
    at2 = [c for c in crits if c[:1] == "D"]
    return ", ".join(at1), ", ".join(at2)


def build_unit(unit_key):
    code, title = UNITS[unit_key]
    items = parse_source_unit(code)
    inv = invert_benchmarks()
    doc = Document(str(TEMPLATE))

    # --- Table 1: Details ---
    det = doc.tables[T_DETAILS]
    _put(det.rows[1].cells[1], code)
    _put(det.rows[1].cells[3], title)
    _put(det.rows[1].cells[5], "1")
    _put(det.rows[2].cells[1], "ICT50220")
    _put(det.rows[2].cells[3], "Diploma of Information Technology")

    # --- Table 2: Description of assessment tasks ---
    desc = doc.tables[T_DESC]
    _put(desc.rows[2].cells[1], AT1_TITLE); _put(desc.rows[2].cells[2], "1.0")
    _put(desc.rows[3].cells[1], AT2_TITLE); _put(desc.rows[3].cells[2], "1.0")

    # --- Table 4: PCs ---
    pc_tbl = doc.tables[T_PC]
    tmpl = _reset(pc_tbl)
    last_elem = None
    for num, txt, elem in items["pcs"]:
        row = _add(pc_tbl, tmpl)
        _put(row.cells[0], elem if elem != last_elem else "")
        last_elem = elem
        _put(row.cells[1], f"{num} {txt}")
        at1, at2 = _split_codes(inv.get((code, "PC", num), []))
        _put(row.cells[2], at1); _put(row.cells[3], at2)
        for c in (4, 5, 6):
            _put(row.cells[c], "")

    # --- Table 5: PEs / Table 6: KEs (col0 desc; col1 AT1; col2 AT2) ---
    for tnum, key, sect in ((T_PE, "pes", "PE"), (T_KE, "kes", "KE")):
        tbl = doc.tables[tnum]
        tmpl = _reset(tbl)
        for i, txt in enumerate(items[key], 1):
            row = _add(tbl, tmpl)
            _put(row.cells[0], txt)
            at1, at2 = _split_codes(inv.get((code, sect, str(i)), []))
            _put(row.cells[1], at1); _put(row.cells[2], at2)
            for c in (3, 4, 5):
                _put(row.cells[c], "")

    # --- Table 7: FSs (col0 skill; col2 description; col3 AT1; col4 AT2) ---
    fs_tbl = doc.tables[T_FS]
    tmpl = _reset(fs_tbl)
    for skill, fdesc in items["fss"]:
        row = _add(fs_tbl, tmpl)
        m = FS_MAP.get(code, {}).get(skill, {})
        _put(row.cells[0], skill)
        _put(row.cells[1], "")
        _put(row.cells[2], fdesc)
        _put(row.cells[3], m.get("AT1", ""))
        _put(row.cells[4], m.get("AT2", ""))
        for c in (5, 6, 7):
            _put(row.cells[c], "")

    # --- Table 3: ACs (col0 desc; col1 AT1; col2 AT2) ---
    ac_tbl = doc.tables[T_AC]
    tmpl = _reset(ac_tbl)
    ac_codes = AC_MAP.get(code, [])
    for i, txt in enumerate(items["acs"]):
        row = _add(ac_tbl, tmpl)
        _put(row.cells[0], txt)
        m = ac_codes[i] if i < len(ac_codes) else {}
        _put(row.cells[1], m.get("AT1", ""))
        _put(row.cells[2], m.get("AT2", ""))
        for c in (3, 4, 5):
            _put(row.cells[c], "")

    MAPPINGS.mkdir(parents=True, exist_ok=True)
    out = MAPPINGS / f"{code}_Assessment_Mapping.docx"
    doc.save(str(out))
    print(f"Wrote {out}  (PC {len(items['pcs'])}, PE {len(items['pes'])}, "
          f"KE {len(items['kes'])}, FS {len(items['fss'])}, AC {len(items['acs'])})")


def main():
    if len(sys.argv) >= 3 and sys.argv[1] == "--build":
        targets = ["501", "503", "505"] if sys.argv[2] == "all" else [sys.argv[2]]
        for t in targets:
            build_unit(t)
        return
    if len(sys.argv) >= 3 and sys.argv[1] == "--dump":
        targets = ["501", "503", "505"] if sys.argv[2] == "all" else [sys.argv[2]]
        for t in targets:
            dump(t)
    else:
        print("Usage: build_cl2_mapping_docs.py --dump {501|503|505|all}")
        sys.exit(2)


if __name__ == "__main__":
    main()
