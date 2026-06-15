"""Tests for helpers.docx_styling — OOXML shading, borders, fields."""
import pathlib
import sys

sys.path.insert(0, str(next(
    d for d in pathlib.Path(__file__).resolve().parents
    if (d / "helpers" / "__init__.py").exists()
)))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from helpers.docx_styling import (  # noqa: E402
    shade_paragraph, shade_cell, set_cell_borders, paragraph_bottom_rule, add_field,
)
from helpers.yat_brand import TEAL, STONE  # noqa: E402


def _cell():
    return Document().add_table(rows=1, cols=1).rows[0].cells[0]


def test_shade_cell_sets_fill():
    c = _cell()
    shade_cell(c, TEAL)
    shd = c._tc.tcPr.find(qn("w:shd"))
    assert shd is not None and shd.get(qn("w:fill")) == TEAL


def test_shade_paragraph_sets_fill():
    p = Document().add_paragraph()
    shade_paragraph(p, TEAL)
    shd = p._p.pPr.find(qn("w:shd"))
    assert shd is not None and shd.get(qn("w:fill")) == TEAL


def test_set_cell_borders_draws_all_four_edges():
    c = _cell()
    set_cell_borders(c)
    borders = c._tc.tcPr.find(qn("w:tcBorders"))
    assert borders is not None
    for edge in ("top", "left", "bottom", "right"):
        e = borders.find(qn(f"w:{edge}"))
        assert e is not None and e.get(qn("w:color")) == STONE


def test_paragraph_bottom_rule_adds_bottom_border():
    p = Document().add_paragraph()
    paragraph_bottom_rule(p, TEAL, sz=6)
    pbdr = p._p.pPr.find(qn("w:pBdr"))
    assert pbdr is not None
    bottom = pbdr.find(qn("w:bottom"))
    assert bottom is not None and bottom.get(qn("w:color")) == TEAL


def test_add_field_inserts_field_code():
    p = Document().add_paragraph()
    add_field(p, "PAGE", placeholder="x")
    instr = p._p.find(".//" + qn("w:instrText"))
    assert instr is not None and instr.text == "PAGE"
    fld_types = [e.get(qn("w:fldCharType")) for e in p._p.iter(qn("w:fldChar"))]
    assert fld_types == ["begin", "separate", "end"]
