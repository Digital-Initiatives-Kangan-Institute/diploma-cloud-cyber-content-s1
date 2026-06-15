"""Tests for helpers.docx_callouts."""
import pathlib
import sys

sys.path.insert(0, str(next(
    d for d in pathlib.Path(__file__).resolve().parents
    if (d / "helpers" / "__init__.py").exists()
)))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from helpers.docx_callouts import add_applicability_note, add_convention_box  # noqa: E402
from helpers.yat_brand import TERRACOTTA, CREAM, TEAL  # noqa: E402


def test_applicability_note_text_and_colour():
    p = add_applicability_note(Document(), "high-availability deployments")
    r = p.runs[0]
    assert r.text.startswith("Applicability: complete for high-availability deployments")
    assert r.italic is True
    assert str(r.font.color.rgb) == TERRACOTTA


def test_convention_box_shaded_cell_and_bold_head():
    t = add_convention_box(Document(), [("Convention", "use ISO dates"), ("Tip", "")])
    cell = t.rows[0].cells[0]
    assert cell._tc.tcPr.find(qn("w:shd")).get(qn("w:fill")) == CREAM
    head_run = cell.paragraphs[0].runs[0]
    assert head_run.text == "Convention" and head_run.bold is True
    assert str(head_run.font.color.rgb) == TEAL
