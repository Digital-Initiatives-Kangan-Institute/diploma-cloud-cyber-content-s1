"""Tests for helpers.docx_tables."""
import pathlib
import sys

sys.path.insert(0, str(next(
    d for d in pathlib.Path(__file__).resolve().parents
    if (d / "helpers" / "__init__.py").exists()
)))

import pytest  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402

from helpers.docx_tables import (  # noqa: E402
    add_template_table, add_data_table, set_cell_content, find_instruction_row, clear_table_rows,
)
from helpers.yat_brand import GREY  # noqa: E402


def test_template_table_shape_and_placeholder_styling():
    t = add_template_table(Document(), ["A", "B"], [["[x]", "plain"]])
    assert len(t.columns) == 2 and len(t.rows) == 2  # header + 1 data row
    placeholder_run = t.rows[1].cells[0].paragraphs[0].runs[0]
    assert placeholder_run.italic is True and str(placeholder_run.font.color.rgb) == GREY


def test_data_table_first_col_bold_and_money_right_aligned():
    t = add_data_table(Document(), ["Item", "Cost"], [["Compute", "$100"]])
    assert t.rows[1].cells[0].paragraphs[0].runs[0].bold is True
    assert t.rows[1].cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_set_cell_content_multiline_replaces():
    cell = Document().add_table(rows=1, cols=1).rows[0].cells[0]
    cell.paragraphs[0].add_run("old")
    set_cell_content(cell, ["line one", "line two"])
    assert cell.text == "line one\nline two"


def test_find_instruction_row_matches_and_raises():
    t = Document().add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Assessment overview"
    t.rows[0].cells[1].text = "marker"
    found = find_instruction_row(t, "Assessment overview")
    assert found._tc is t.rows[0].cells[1]._tc  # same underlying cell
    assert found.text == "marker"
    with pytest.raises(KeyError):
        find_instruction_row(t, "nonexistent")


def test_clear_table_rows_keeps_n():
    t = Document().add_table(rows=5, cols=1)
    clear_table_rows(t, 2)
    assert len(t.rows) == 2
