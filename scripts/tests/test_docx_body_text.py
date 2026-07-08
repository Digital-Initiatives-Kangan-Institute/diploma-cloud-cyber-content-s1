"""Tests for helpers.docx_body_text."""
import pathlib
import sys

# Put the scripts/ root (the dir containing the `helpers` package) on sys.path.
sys.path.insert(0, str(next(
    d for d in pathlib.Path(__file__).resolve().parents
    if (d / "helpers" / "__init__.py").exists()
)))

from docx import Document  # noqa: E402
from docx.shared import Pt  # noqa: E402

from helpers.docx_body_text import (  # noqa: E402
    add_body_paragraph,
    add_bullet_list,
    add_guidance_text,
    add_response_placeholder,
    BODY_FONT_SIZE_PT,
    GUIDANCE_FONT_SIZE_PT,
)
from brand import GREY  # noqa: E402


def test_guidance_text_is_italic_grey_at_guidance_size():
    p = add_guidance_text(Document(), "Write this last.")
    run = p.runs[0]
    assert run.italic is True
    assert run.font.size == Pt(GUIDANCE_FONT_SIZE_PT)
    assert str(run.font.color.rgb) == GREY


def test_response_placeholder_default_label_italic():
    p = add_response_placeholder(Document())
    assert p.runs[0].text == "[ Write your response here ]"
    assert p.runs[0].italic is True


def test_body_paragraph_text_and_size():
    p = add_body_paragraph(Document(), "Hello world")
    assert p.text == "Hello world"
    assert p.runs[0].font.size == Pt(BODY_FONT_SIZE_PT)


def test_body_paragraph_returns_one_paragraph():
    doc = Document()
    before = len(doc.paragraphs)
    add_body_paragraph(doc, "x")
    assert len(doc.paragraphs) == before + 1


def test_one_paragraph_per_item():
    paras = add_bullet_list(Document(), ["First", "Second", "Third"])
    assert len(paras) == 3


def test_text_and_bullet_style():
    paras = add_bullet_list(Document(), ["Alpha", "Beta"])
    assert [p.text for p in paras] == ["Alpha", "Beta"]
    assert all(p.style.name == "List Bullet" for p in paras)


def test_body_font_size():
    paras = add_bullet_list(Document(), ["x"])
    assert paras[0].runs[0].font.size == Pt(BODY_FONT_SIZE_PT)


def test_empty_is_noop():
    assert add_bullet_list(Document(), []) == []
