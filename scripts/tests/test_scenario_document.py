"""Tests for helpers.scenario_document."""
import pathlib
import sys

sys.path.insert(0, str(next(
    d for d in pathlib.Path(__file__).resolve().parents
    if (d / "helpers" / "__init__.py").exists()
)))

from docx import Document  # noqa: E402
from docx.shared import Pt  # noqa: E402

from helpers.scenario_document import configure_styles, build_header_footer, wordmark  # noqa: E402
from brand import FONT, TEAL, TERRACOTTA, DISCLOSURE  # noqa: E402


def test_configure_styles_sets_normal_and_title():
    doc = Document()
    configure_styles(doc)
    assert doc.styles["Normal"].font.name == FONT
    assert doc.styles["Normal"].font.size == Pt(10.5)
    assert str(doc.styles["Title"].font.color.rgb) == TEAL


def test_wordmark_two_runs_branded():
    p = Document().add_paragraph()
    wordmark(p)
    assert [r.text for r in p.runs] == ["YAT", "  COLLEGE"]
    assert str(p.runs[0].font.color.rgb) == TEAL
    assert str(p.runs[1].font.color.rgb) == TERRACOTTA


def test_header_footer_disclosure_and_page_field():
    doc = Document()
    section = doc.sections[0]
    build_header_footer(section)
    assert DISCLOSURE in section.header.paragraphs[0].text
    assert "YAT College" in section.footer.paragraphs[0].text
