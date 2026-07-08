#!/usr/bin/env python3
"""Build the AccentLoitte LMS Replacement Business Case (.docx) — student download example.

A complete, polished, *approved* worked Business Case for the 2022 AccentLoitte engagement
that replaced YAT's end-of-life GrayBoard LMS with DOODLE on-premises. Provided to students
as an example of what a finished business case looks and feels like.

Student-facing → in-world only: NO UoC tags, NO assessment/course language, sign-off filled
as an approved past deliverable. Reuses the YAT brand template helpers (build_bc_template)
and the filled-table/prose helpers (build_s1_cl1_at1_bc_exemplar) so it matches the house style exactly.

Usage:  python scripts/scenario/build_lms_replacement_bc.py [output.docx]
Default: ../diploma-cloud-cyber-website-s1/public/documents/YAT-LMS-Replacement-Business-Case.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_body_paragraph, add_bullet_list  # noqa: E402
from helpers.docx_tables import add_data_table  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from brand import ADDRESS, CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document            # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def build(path):
    doc = Document()
    configure_styles(doc)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Business Case")
    sub = doc.add_paragraph().add_run("YAT LMS Replacement — GrayBoard to DOODLE")
    sub.font.size = Pt(15); sub.font.color.rgb = RGBColor.from_string(TERRACOTTA); sub.bold = True
    doc.add_paragraph()

    cover = [
        ("Prepared for", "YAT Board, via Sam Walker (YAT ICT Manager)"),
        ("Prepared by", "AccentLoitte Consulting"),
        ("Reviewed by", "Jamie O'Donnell, AccentLoitte Senior Consultant"),
        ("Engagement", "YAT LMS Replacement Project — Phase 1"),
        ("Document version", "v1.0 — Approved"),
        ("Date", "May 2022"),
        ("Analysis period", "5 years"),
        ("Currency", "AUD, exclusive of GST (2022 price levels)"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h2 = lambda t: doc.add_paragraph(t, style="Heading 2")

    h1("1. Executive Summary")
    add_body_paragraph(doc, "YAT's Learning Management System, GrayBoard, is a bespoke in-house platform that has "
                 "reached end of supportable life. Its availability and performance have been degrading, "
                 "it has no vendor support or product roadmap, and it does not meet current accessibility "
                 "obligations — placing YAT's teaching delivery and compliance at growing risk. This "
                 "Business Case evaluates two options over a five-year horizon: Option A — remediate and "
                 "retain GrayBoard; Option B — replace it with a supported contemporary LMS (DOODLE) "
                 "deployed on-premises at the Cremorne campus.")
    add_body_paragraph(doc, "Option A is ~$63,600 cheaper over five years on direct cost, but it only defers the "
                 "problem: it leaves YAT on an unsupported, dead-end platform with a heavy in-house "
                 "maintenance burden and a fragile accessibility retrofit, and a full replacement would "
                 "still be required. Option B delivers a vendor-supported, WCAG 2.1 AA-compliant, "
                 "performant LMS with an active product roadmap, preserves all student records, and "
                 "actually reduces YAT's ongoing maintenance burden — partly offsetting the higher "
                 "upfront cost.")
    add_body_paragraph(doc, "We recommend Option B. We ask the board to approve the replacement, the phased "
                 "implementation plan in §10 (with cutover scheduled into the December–January teaching "
                 "break), and the Year-1 budget envelope.")

    h1("2. Engagement Context")
    add_body_paragraph(doc, "YAT College engaged AccentLoitte to deliver the end-to-end replacement of its legacy "
                 "GrayBoard LMS, structured in three phases: (1) selection and implementation planning; "
                 "(2) deployment and data migration; (3) stabilisation and closure. This Business Case is "
                 "the Phase 1 deliverable: it presents the options analysis and the replacement "
                 "recommendation for the YAT Board's approval, via the Engagement Sponsor, the YAT ICT "
                 "Manager (Sam Walker).")

    h1("3. Strategic Alignment Analysis")
    add_body_paragraph(doc, "YAT's ICT direction in 2022 prioritises reliable, well-supported core systems; the "
                 "capacity to support continued student growth; and meeting the College's statutory "
                 "obligations as a Registered Training Organisation. The LMS is the platform through "
                 "which teaching, assessment and student records are delivered and held, so its "
                 "reliability and compliance are foundational to that direction.")
    add_body_paragraph(doc, "Across the RTO sector, institutions have moved decisively off bespoke in-house LMS "
                 "platforms onto established, vendor-supported products with active communities and "
                 "accessibility conformance — driven by the maintenance cost and compliance risk of "
                 "home-grown systems, and by rising accessibility expectations under the Disability "
                 "Discrimination Act 1992. YAT's continued reliance on GrayBoard runs against this "
                 "direction. Cloud-hosted delivery is the emerging next step in the sector; YAT notes it "
                 "as a future consideration but, consistent with its then-current on-premises operating "
                 "model, this engagement is scoped to an on-premises replacement. Replacing GrayBoard now "
                 "with a supported product both resolves the immediate risk and positions YAT for a "
                 "future move to cloud hosting.")

    h1("4. Current State — the GrayBoard LMS")
    add_body_paragraph(doc, "GrayBoard is a bespoke LMS built and maintained in-house by YAT ICT, running on "
                 "on-premises infrastructure at the Cremorne campus and integrated with Active Directory "
                 "(authentication), Office 365 (email notifications), and AVETMISS/ASQA statutory "
                 "reporting. It has reached end of supportable life: effective availability has been "
                 "degrading, typical and peak-load performance is declining (the platform struggles with "
                 "the ~3× concurrent load during end-of-term assessment windows), and it depends on a "
                 "small number of YAT staff for continued maintenance with no external support or product "
                 "roadmap. It does not meet WCAG 2.1 Level AA, creating a compliance exposure under the "
                 "Disability Discrimination Act. It does, however, hold YAT's complete student-record "
                 "history, which carries long-tail retention obligations under the Standards for RTOs "
                 "2015 — so any replacement must preserve those records with full integrity. These "
                 "limitations — no support, declining availability and performance, accessibility "
                 "non-conformance, and concentration risk on in-house maintenance — are what drive the "
                 "gap analysis.")

    h1("5. Gap Analysis")
    add_data_table(doc,
              ["Objective", "Current state (GrayBoard)", "Desired future state", "Gap",
               "Improvement opportunity", "Proposed change"],
              [["Supported, current LMS platform", "Bespoke, end-of-life, no vendor support/roadmap",
                "Vendor-supported product with active roadmap", "No support or roadmap; concentration risk",
                "Adopt an established commercial LMS", "Replace with DOODLE (on-prem)"],
               ["Availability ≥ 99%", "Degrading, below target", "≥ 99%",
                "Reliability shortfall", "Modern supported stack", "Deploy DOODLE on supported infrastructure"],
               ["WCAG 2.1 AA accessibility", "Non-conformant (DDA exposure)", "Conformant",
                "Compliance gap", "Product with native AA conformance", "Replace with conformant LMS"],
               ["Peak-load performance", "Degrades at ~3× term-end load", "Meets/exceeds, holds at peak",
                "Performance shortfall", "Right-sized supported platform", "Provision DOODLE to peak profile"],
               ["Preserve all student records", "Held in GrayBoard", "Migrated intact + verified",
                "Migration risk if not managed", "Verified ETL migration", "Migrate with reconciliation + parallel run"]],
              widths=[2.7, 2.8, 2.6, 2.4, 2.5, 2.5])
    add_body_paragraph(doc, "The gaps cluster around supportability, reliability, compliance and performance — none "
                 "of which an in-house remediation truly closes, and all of which a supported replacement "
                 "addresses while also preserving the student-record history.")

    h1("6. Options Considered and Evaluation")
    h2("6.1 Solution requirements")
    add_body_paragraph(doc, "The replacement must: be a current, vendor-supported LMS product; run on YAT's Windows "
                 "Server 2016 on-premises hosting standard; preserve all GrayBoard student records with "
                 "full integrity; achieve ≥ 99% availability; meet WCAG 2.1 AA; match or exceed "
                 "GrayBoard's performance (including ~3× term-end peaks); integrate with Active Directory, "
                 "Office 365 and AVETMISS reporting; and carry vendor support with severity-1 response "
                 "within 4 hours during teaching hours.")
    h2("6.2 Options considered")
    add_bullet_list(doc, [
        "Option A — Remediate and retain GrayBoard: invest in an accessibility retrofit and a hardware "
        "refresh, and continue in-house maintenance of the legacy platform.",
        "Option B — Replace with a contemporary on-prem LMS (DOODLE): select an established product via "
        "market evaluation (see Appendix 2), deploy it on YAT infrastructure, migrate the data, and train staff.",
        "A SaaS / cloud-hosted LMS was considered but not carried forward in this engagement: it is "
        "inconsistent with YAT's current on-premises operating model and is noted as a future option.",
    ])
    h2("6.3 Evaluation method")
    add_body_paragraph(doc, "Each option is evaluated with a five-year cost comparison (§7), an intangibles "
                 "comparison covering support, compliance, reliability and maintenance burden (§8.1), and "
                 "a risk register (§8.2). Because the decision turns largely on supportability and "
                 "compliance rather than on cost alone, the intangibles carry significant weight.")
    h2("6.4 Initial impact and difficulty assessment")
    add_data_table(doc, ["", "Option A — Remediate GrayBoard", "Option B — Replace (DOODLE)"],
              [["Strategic impact", "Defers the problem; remains a dead-end platform",
                "Resolves it; supported, compliant, future-ready"],
               ["Implementation difficulty", "Low — incremental work on a known system",
                "Moderate — selection, deployment, data migration, training"],
               ["Headline pros", "Lower upfront cost; familiar platform",
                "Vendor support + roadmap; compliant; lower ongoing maintenance"],
               ["Headline cons", "Unsupported, non-compliant, replacement still looms",
                "Higher Year-1 cost; migration and change effort"]],
              widths=[4.0, 6.25, 6.25])

    h1("7. Cost Comparison")
    add_body_paragraph(doc, "A five-year cost comparison of the two options (AUD ex GST, 2022 price levels; "
                 "inflation excluded for comparison clarity). Detailed line items are in Appendix 1.")
    yr = ["", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5", "5-year"]
    h2("7.1 Option A — Remediate and retain GrayBoard")
    add_data_table(doc, yr,
              [["One-off remediation", "$55,000", "—", "—", "—", "—", "$55,000"],
               ["Recurring (maintenance)", "$67,000", "$67,000", "$67,000", "$67,000", "$67,000", "$335,000"],
               ["Annual total", "$122,000", "$67,000", "$67,000", "$67,000", "$67,000", "$390,000"]])
    h2("7.2 Option B — Replace with DOODLE (on-premises)")
    add_data_table(doc, yr,
              [["One-off project", "$160,000", "—", "—", "—", "—", "$160,000"],
               ["Recurring (run + support)", "(in project)", "$73,400", "$73,400", "$73,400", "$73,400", "$293,600"],
               ["Annual total", "$160,000", "$73,400", "$73,400", "$73,400", "$73,400", "$453,600"]])
    h2("7.3 Comparison summary")
    add_data_table(doc, ["", "Option A", "Option B", "Delta (B − A)"],
              [["One-off (Year 1)", "$55,000", "$160,000", "+$105,000"],
               ["Recurring (per year)", "$67,000", "$73,400", "+$6,400"],
               ["5-year direct cost", "$390,000", "$453,600", "+$63,600"]],
              widths=[5.0, 3.8, 3.8, 3.8])
    add_body_paragraph(doc, "Option B costs more on every measure — ~$105,000 more upfront and ~$6,400 more per "
                 "year, ~$63,600 more over five years. The case for it is therefore not cost but "
                 "supportability, compliance and performance: Option A's lower spend buys only a "
                 "temporary reprieve on an unsupported, non-compliant platform that must still be "
                 "replaced. And as the sensitivity below shows, Option A's cost advantage erodes — or "
                 "reverses — as the ageing platform's maintenance burden grows.")
    h2("7.4 Sensitivity")
    add_body_paragraph(doc, "Sensitivity 1 — GrayBoard maintenance effort rising from ~0.4 to ~0.6 FTE: Option A's "
                 "recurring cost rises by ~$22,000/year (~$110,000 over five years), making Option A more "
                 "expensive than Option B. Given the platform's trajectory this is the more likely case, "
                 "and it reinforces the recommendation.")
    add_body_paragraph(doc, "Sensitivity 2 — AccentLoitte fees rising from $95,000 to $115,000: Option B's five-year "
                 "cost rises to ~$473,600. The cost gap widens, but the supportability and compliance case "
                 "for Option B is unchanged.")

    h1("8. Risk and Impact Assessment")
    h2("8.1 Intangibles comparison")
    add_data_table(doc, ["Factor", "Option A — Remediate GrayBoard", "Option B — Replace (DOODLE)"],
              [["Vendor support & roadmap", "None — bespoke dead-end", "Active vendor support + roadmap"],
               ["Accessibility (WCAG 2.1 AA / DDA)", "Fragile retrofit", "Native conformance"],
               ["Availability (≥ 99%)", "Degrading", "Meets target on supported stack"],
               ["Peak-load performance", "Degrades at term-end", "Match/exceed; provisioned to peak"],
               ["YAT ICT maintenance burden", "High (in-house bespoke)", "Reduced (supported product)"],
               ["Future-proofing", "Replacement still looms", "Current platform; cloud-ready later"]],
              widths=[4.5, 5.75, 6.0])
    h2("8.2 Risk register (recommended option)")
    add_data_table(doc, ["Risk", "Likelihood", "Impact", "Mitigation"],
              [["Data loss/corruption in migration", "Low", "High",
                "Verified ETL with record reconciliation; parallel run; tested backups"],
               ["Cutover disrupts teaching", "Medium", "High",
                "Cutover in the Dec–Jan break; rollback plan; staged go-live"],
               ["Low user adoption / support spike", "Medium", "Medium",
                "Role-based training before cutover; stabilisation support period"],
               ["Integration gaps (AD / O365 / AVETMISS)", "Medium", "Medium",
                "Integration testing during parallel run before cutover"],
               ["Schedule slips into a teaching term", "Low", "Medium",
                "Phase gates; the fixed Dec–Jan cutover window"]],
              widths=[5.2, 2.3, 2.0, 6.7])

    h1("9. Recommendation")
    add_body_paragraph(doc, "Recommended option: Option B — replace GrayBoard with DOODLE, deployed on-premises. "
                 "Although ~$63,600 more over five years, it is the only option that delivers a supported, "
                 "accessible, performant LMS with a product roadmap, reduces YAT's ongoing maintenance "
                 "burden, and preserves the student-record history — while Option A merely defers an "
                 "unavoidable replacement on an unsupported, non-compliant platform. This recommendation "
                 "drives the phased action plan below.")

    h1("10. Action Plan")
    h2("10.1 Prioritised phases")
    add_data_table(doc, ["#", "Phase", "Priority rationale"],
              [["1", "Selection & implementation plan", "Confirm the product and the deployment/migration approach first"],
               ["2", "Deployment & data migration", "Stand up DOODLE, migrate and verify records, run in parallel"],
               ["3", "Stabilisation & closure", "Support, training, as-built docs, knowledge transfer, closure"]],
              widths=[1.2, 6.0, 8.5])
    h2("10.2 Implementation schedule")
    add_data_table(doc, ["Phase", "Activities", "Window", "Dependencies", "Owner"],
              [["1", "Requirements, market evaluation, implementation plan", "2022 H1", "—", "AccentLoitte"],
               ["2", "Infra provisioning, DOODLE install, ETL migration, parallel run, cutover",
                "Dec 2022 – Jan 2023 break", "Phase 1 approval", "AccentLoitte + YAT ICT"],
               ["3", "Stabilisation support, training, as-built docs, closure pack",
                "Jan 2023", "Phase 2 cutover", "AccentLoitte + YAT ICT"]],
              widths=[1.5, 5.8, 3.2, 2.7, 2.6])
    add_body_paragraph(doc, "Staff training is delivered before cutover to contain the post-cutover support burden, "
                 "and the cutover is scheduled into the December–January teaching break.")
    h2("10.3 Standards, targets, and success metrics")
    add_data_table(doc, ["Aspect", "Standard / target", "How measured"],
              [["Availability", "≥ 99%", "LMS uptime monitoring over the stabilisation period"],
               ["Accessibility", "WCAG 2.1 AA", "Accessibility audit of the deployed LMS"],
               ["Record preservation", "100% reconciliation", "Post-migration record counts vs pre-migration; sample fidelity checks"],
               ["Vendor support", "Sev-1 ≤ 4h (teaching hours)", "Support contract terms; incident logs"]],
              widths=[4.5, 5.5, 5.5])
    h2("10.4 Alignment with the change management procedure")
    add_body_paragraph(doc, "The cutover is a high-risk change requiring Change Advisory Board endorsement and YAT "
                 "senior-management sign-off; the Phase 3 Closure Pack is signed off under the YAT Change "
                 "Management Procedure.")

    h1("11. Next Steps and Decision Requested")
    add_body_paragraph(doc, "We ask the board to: (1) approve Option B — replacing GrayBoard with DOODLE on-premises; "
                 "(2) approve the phased implementation plan in §10; and (3) authorise the Year-1 budget "
                 "envelope (≈$160,000). The high-risk cutover change request is submitted at the Phase 2 "
                 "gate per the change management procedure.")

    h1("Sign-off")
    add_data_table(doc, ["Role", "Name", "Date", "Signature"],
              [["Prepared by", "AccentLoitte Consultant", "12 May 2022", "(signed)"],
               ["Reviewed by", "Jamie O'Donnell (AccentLoitte Senior Consultant)", "13 May 2022", "(signed)"],
               ["Approved by", "Sam Walker (YAT ICT Manager)", "20 May 2022", "(signed)"]],
              widths=[4.0, 6.0, 3.0, 2.5])

    # ---- APPENDICES ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1("Appendix 1 — Cost Detail")
    h2("A1.1 Option A — Remediate and retain GrayBoard")
    add_body_paragraph(doc, "One-off remediation (Year 1):")
    add_data_table(doc, ["Item", "Cost"],
              [["Accessibility retrofit to WCAG 2.1 AA", "$35,000"],
               ["Legacy server hardware refresh", "$20,000"],
               ["Total one-off", "$55,000"]],
              widths=[11.5, 4.0])
    add_body_paragraph(doc, "Recurring (per year):")
    add_data_table(doc, ["Item", "Annual cost"],
              [["In-house bespoke maintenance (~0.4 FTE)", "$44,000"],
               ["Incident response (~0.10 FTE, rising)", "$11,000"],
               ["Hosting, facilities, backup", "$12,000"],
               ["Total recurring per year", "$67,000"]],
              widths=[11.5, 4.0])
    h2("A1.2 Option B — Replace with DOODLE (on-premises)")
    add_body_paragraph(doc, "One-off project (Year 1):")
    add_data_table(doc, ["Item", "Cost"],
              [["AccentLoitte fees (Phases 1–3: selection, deployment, migration, training, closure)", "$95,000"],
               ["On-premises infrastructure (server, storage, UPS, backup)", "$35,000"],
               ["YAT ICT staff time (oversight, integration, acceptance, parallel run)", "$22,000"],
               ["Training delivery", "$8,000"],
               ["Total Year-1 project", "$160,000"]],
              widths=[11.5, 4.0])
    add_body_paragraph(doc, "Recurring (per year, Years 2–5):")
    add_data_table(doc, ["Item", "Annual cost"],
              [["DOODLE vendor support & maintenance contract", "$30,000"],
               ["Windows Server licensing + antivirus + monitoring", "$3,800"],
               ["Power, facilities, server-room allocation", "$6,700"],
               ["Backup (media, offsite storage, software maintenance)", "$5,400"],
               ["YAT ICT staff (admin 0.20 FTE + incident 0.05 FTE)", "$27,500"],
               ["Total recurring per year", "$73,400"]],
              widths=[11.5, 4.0])

    h1("Appendix 2 — LMS Market Evaluation (summary)")
    add_body_paragraph(doc, "Candidate replacement platforms were scored against the §6.1 requirements. DOODLE was "
                 "selected on the balance of functional fit, total cost of ownership, RTO-sector fit, "
                 "accessibility conformance, and vendor longevity.")
    add_data_table(doc, ["Criterion", "DOODLE (selected)", "Alternative 1", "Alternative 2"],
              [["Functional fit to YAT requirements", "Strong", "Strong", "Moderate"],
               ["On-prem deployment supported", "Yes", "Yes", "SaaS only"],
               ["WCAG 2.1 AA conformance", "Yes", "Partial", "Yes"],
               ["AD / O365 / AVETMISS integration", "Yes", "Yes", "Limited"],
               ["5-year total cost of ownership", "Competitive", "Higher", "Lower (but SaaS — out of scope)"],
               ["Vendor longevity & roadmap", "Strong", "Moderate", "Strong"]],
              widths=[5.0, 3.8, 3.4, 3.3])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website-s1/public/documents/YAT-LMS-Replacement-Business-Case.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
