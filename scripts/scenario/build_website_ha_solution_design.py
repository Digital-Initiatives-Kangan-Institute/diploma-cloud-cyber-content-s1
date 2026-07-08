#!/usr/bin/env python3
"""Build the Website HA-Hardened Solution Design (.docx) — in-world document.

The as-built cloud architecture of the YAT website AFTER high-availability hardening —
the single-AZ 2023 pilot taken to a Multi-AZ, fault-tolerant deployment (ALB + cross-AZ
Auto Scaling Group + Multi-AZ RDS + media offloaded from instance EBS to S3). This is the
output of the Website HA Hardening engagement.

It deliberately stops at IN-REGION high availability. Global serving / CDN, cross-region
disaster recovery, the audit/access-log microservice, and infrastructure-as-code are the
subject of the subsequent Website Global Expansion engagement (out of scope here, §7).

This is the provided STARTING STATE for the CL2 assessment (DR plan + microservice + IaC
against a hardened website) — students look this up; they do not build it. In CL3 the
website-hardening practice arrives at an equivalent state from the single-AZ baseline.
An in-world supplied artefact (no UoC tags) — the branded Solution Design document.

Output to the website documents folder for printing to PDF, then linked from the
website-ha-hardening project's HA design page.

Usage:  python scripts/scenario/build_website_ha_solution_design.py [output.docx]
Default: ../diploma-cloud-cyber-website-s1/public/documents/YAT-Website-HA-Solution-Design.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_body_paragraph, add_bullet_list  # noqa: E402
from helpers.docx_tables import add_data_table  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from brand import ADDRESS, CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def na(doc, reason):
    p = doc.add_paragraph()
    r = p.add_run(f"Not applicable — {reason}")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    p.paragraph_format.space_after = Pt(6)
    return p


def diagram_placeholder(doc, caption, source):
    """A bordered, shaded drop-zone for a network diagram — replace with the exported image."""
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    set_cell_borders(cell)
    shade_cell(cell, CREAM)
    cell.width = Cm(16.6)
    t.rows[0].height = Cm(9.5)
    t.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[ NETWORK TOPOLOGY DIAGRAM — PASTE HERE ]")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(source)
    r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = RGBColor.from_string(GREY)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = RGBColor.from_string(GREY)
    return t


def build(path):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Solution Design")
    sub = doc.add_paragraph().add_run("YAT Website Cloud Architecture — HA-Hardened Design")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT Website HA Hardening"),
        ("Document type", "Technical design (Solution Design)"),
        ("Version", "v1.0 — As-built (HA-hardened)"),
        ("Authored by", "MTS Consulting, in consultation with YAT ICT"),
        ("Approved by", "Pat Lin (MTS Senior Consultant) · Sam Walker (YAT ICT Manager)"),
        ("Predecessor", "YAT Website Cloud Architecture — Baseline Design (single-AZ)"),
        ("Classification", "Internal — YAT ICT, and MTS personnel on signed MSA"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Purpose and Scope")
    add_body_paragraph(doc, "This document records the cloud architecture of the YAT website after high-availability "
                 "hardening. It takes the single-Availability-Zone 2023 pilot (the Website Cloud Architecture — "
                 "Baseline Design) and hardens it into a Multi-AZ, fault-tolerant deployment: the single web "
                 "instance becomes a load-balanced, auto-scaling group spread across two Availability Zones; "
                 "the single-AZ database becomes a Multi-AZ deployment with an automatic-failover standby; and "
                 "uploaded media moves off the instance's local disk into shared object storage.")
    h3("In scope of this design")
    add_bullet_list(doc, [
        "Removal of the single-instance, single-Availability-Zone, and single-database points of failure on the website's serving path.",
        "An Application Load Balancer, a cross-AZ Auto Scaling Group, a Multi-AZ managed database, and media offloaded to object storage.",
        "Single-region, two-Availability-Zone deployment in ap-southeast-2 (Sydney).",
    ])
    h3("Out of scope — the subsequent Website Global Expansion engagement")
    add_bullet_list(doc, [
        "Global serving and edge content delivery (CDN) for an international audience.",
        "Cross-region disaster recovery (a second region, recovery objectives, DR runbook).",
        "The audit/access-log microservice and the in-India log store.",
        "Infrastructure-as-code templating of the footprint.",
    ])
    h3("Out of MTS scope entirely — YAT responsibility")
    add_bullet_list(doc, [
        "Website content authoring and CMS configuration (YAT Marketing).",
        "Ongoing operation and support of the website after handover (YAT ICT).",
    ])

    h1("2. Design Inputs and Requirements")
    h3("2.1 Inputs")
    add_bullet_list(doc, [
        "Website Cloud Architecture — Baseline Design — the single-AZ pilot being hardened.",
        "Website Specification — workload, traffic, and data profile of the public website.",
        "Website Infrastructure Specifications — the as-deployed single-AZ resources.",
    ])
    h3("2.2 Requirements the design must meet")
    add_data_table(doc, ["Requirement", "Target / note"],
              [["Region / data residency", "ap-southeast-2 (Sydney); enquiry / application PII within Australia (Privacy Act, APP 8)"],
               ["Application stack", "Preserved — the existing PHP / MySQL CMS on a LAMP stack"],
               ["Availability", "Survive the loss of a single instance or a single Availability Zone with no manual intervention"],
               ["Recovery", "Automatic database failover; serving continues from the surviving AZ"],
               ["Scaling", "Scale the web tier with demand (the Jan–Feb enrolment-enquiry peak)"],
               ["Disaster recovery", "Out of scope here — cross-region DR is the subsequent expansion engagement"]],
              widths=[5.0, 11.0])

    h1("3. Review of the Existing (Single-AZ) Architecture")
    add_body_paragraph(doc, "The starting point is the 2023 single-AZ pilot: a single EC2 instance (LAMP / CMS) with an "
                 "Elastic IP serving traffic directly, a single-AZ Amazon RDS for MySQL database, and uploaded "
                 "media stored on the instance's EBS volume, with nightly backups to S3. Its known single points "
                 "of failure — the single instance, the single Availability Zone, the single database, and media "
                 "held on one instance's disk — are precisely what this engagement removes.")

    h1("4. Architecture Design")
    h3("4.1 Assumptions and constraints")
    add_data_table(doc, ["#", "Assumption / constraint", "Source"],
              [["A1", "Region must be ap-southeast-2 (Sydney) for data residency", "Website Specification; Privacy Policy"],
               ["A2", "CMS application stack preserved: PHP / MySQL on a LAMP stack", "Baseline Design"],
               ["A3", "Two Availability Zones (ap-southeast-2a, ap-southeast-2b) used for fault tolerance", "HA scoping decision"],
               ["A4", "Light load with a Jan–Feb enrolment-enquiry peak; web tier scales with demand", "Website Specification"],
               ["A5", "Enquiry / application PII must remain within Australia", "Privacy Act 1988, APP 8"],
               ["A6", "Cross-region DR, global serving, and the microservice are out of scope (later engagement)", "Scoping decision"]],
              widths=[1.0, 9.0, 6.0])
    h3("4.2 AWS account and region")
    add_body_paragraph(doc, "Region ap-southeast-2 (Sydney); no deployment outside Australian regions. The deployment now "
                 "spans two Availability Zones (ap-southeast-2a and ap-southeast-2b) so that the loss of one AZ "
                 "does not take the website offline.")
    h3("4.3 Identity and Access Management (IAM)")
    add_data_table(doc, ["Group / role", "Purpose", "Indicative permissions"],
              [["YAT-ICT-Admins", "YAT ICT day-to-day ops", "Console access to EC2/RDS/S3/ELB/CloudWatch; no IAM changes"],
               ["MTS-Consultants", "MTS during the hardening + support", "Admin during build; reduced at handover"],
               ["Website-Instance role", "EC2 instance profile for the web tier", "S3 read/write (media + backups); CloudWatch logs/metrics"]],
              widths=[3.5, 5.0, 7.5])
    add_bullet_list(doc, [
        "MFA required for all human users (User Access Policy; ACSC Essential Eight).",
        "EC2 access to AWS via the instance profile; no long-lived human access keys.",
    ])
    h3("4.4 Network topology")
    add_body_paragraph(doc, "VPC 10.0.0.0/16, DNS resolution and VPC flow logs enabled. Subnets across two AZs:")
    add_data_table(doc, ["Subnet", "CIDR", "Tier", "Internet-facing?"],
              [["public-web-a / -b", "10.0.1.0/24, 10.0.2.0/24", "Load balancer + NAT (one per AZ)", "Yes (ALB)"],
               ["private-app-a / -b", "10.0.11.0/24, 10.0.12.0/24", "Website EC2 (ASG, one per AZ)", "No"],
               ["private-data-a / -b", "10.0.21.0/24, 10.0.22.0/24", "Database (RDS primary / standby)", "No"]],
              widths=[4.0, 4.0, 5.0, 3.0])
    add_bullet_list(doc, [
        "Internet Gateway for the public subnets; a NAT Gateway in each AZ for the private subnets' outbound traffic.",
        "The web instances move OUT of the public subnet (where the single pilot instance sat) into private app subnets behind the load balancer.",
    ])
    diagram_placeholder(doc,
                        "Figure 4.4 — Website HA-hardened network topology (Multi-AZ).",
                        "Source diagram: website-ha-hardened.drawio")
    h3("4.5 Compute (ALB + cross-AZ Auto Scaling)")
    add_bullet_list(doc, [
        "Application Load Balancer, internet-facing, across public-web-a/b; HTTPS:443 listener; ACM-issued certificate (TLS now terminates at the load balancer, not on the instance).",
        "Auto Scaling Group of website instances (LAMP / CMS) across private-app-a/b: min 2 (one per AZ), scaling on CPU / request count; target group health-checked by the ALB.",
        "This removes the single-instance SPOF: an instance or AZ failure leaves healthy capacity serving in the other AZ.",
    ])
    h3("4.6 Database (RDS — Multi-AZ)")
    add_bullet_list(doc, [
        "Amazon RDS for MySQL converted to a Multi-AZ deployment: a primary in private-data-a with a synchronously-replicated standby in private-data-b.",
        "Automatic failover to the standby on primary failure or AZ loss (typically under two minutes), with no application reconfiguration.",
        "Storage encryption (KMS); not publicly accessible; automated daily backups retained.",
    ])
    h3("4.7 Storage (media offloaded)")
    add_data_table(doc, ["Resource", "Type", "Purpose"],
              [["Media bucket", "S3", "Uploaded media (images, brochures, course PDFs) — moved off instance EBS so every instance serves the same media and no single disk is a SPOF"],
               ["Backups bucket", "S3", "Nightly database and media backups"]],
              widths=[4.0, 3.0, 9.0])
    add_bullet_list(doc, ["Both buckets: block all public access; SSE-S3 encryption; versioning enabled. Edge delivery (CDN) of media is deferred to the global-expansion engagement."])
    h3("4.8 Security")
    add_data_table(doc, ["Security group", "Inbound", "Outbound"],
              [["sg-alb", "HTTPS:443 from 0.0.0.0/0", "HTTP to sg-app"],
               ["sg-app", "from sg-alb; SSH from an admin source (implementer decision)", "MySQL:3306 to sg-db; HTTPS via NAT to AWS APIs / updates"],
               ["sg-db", "MySQL:3306 from sg-app only", "none"]],
              widths=[3.0, 7.5, 5.5])
    add_bullet_list(doc, [
        "Encryption in transit (HTTPS to the ALB; TLS app→RDS) and at rest (EBS, RDS, S3).",
        "Operates under the AWS Shared Responsibility Model.",
    ])
    h3("4.9 Monitoring (HA)")
    add_body_paragraph(doc, "CloudWatch metrics for the ALB, the Auto Scaling Group, and RDS, with HA-relevant alarms:")
    add_data_table(doc, ["Alarm", "Threshold"],
              [["ALB unhealthy host count", "≥ 1 unhealthy target"],
               ["ASG in-service instances", "< 2 (capacity below one-per-AZ)"],
               ["ALB 5XX", "> 10 / min"],
               ["RDS failover / replica lag", "Failover event, or lag above threshold"],
               ["RDS free storage low", "< 15%"]],
              widths=[8.0, 8.0])
    h3("4.10 Naming and tagging conventions")
    add_body_paragraph(doc, "Naming pattern yat-web-<resource-type>-<env>-<az> (e.g. yat-web-ec2-prod-a). Mandatory tags: "
                 "Project=YAT-Website-HA, Environment=Production, Owner=YAT-ICT, DataClassification=Internal; PII.")
    h3("4.11 Backup")
    add_data_table(doc, ["Resource", "Mechanism", "Retention"],
              [["RDS database", "Automated daily backups + transaction logs (Multi-AZ)", "7 days"],
               ["Media (S3)", "Versioning + nightly snapshot to the backups bucket", "Versioned"]],
              widths=[4.5, 7.5, 4.0])
    add_bullet_list(doc, ["Cross-Region backup copies remain out of scope here — they are part of the disaster-recovery work in the subsequent expansion engagement."])
    h3("4.12 Recovery objectives — HA state")
    add_body_paragraph(doc, "Within the region, the website now tolerates instance and AZ failure automatically: serving "
                 "continues from the surviving AZ and the database fails over to its standby. Recovery from the "
                 "loss of the whole region is NOT addressed by this design — that is the cross-region disaster "
                 "recovery work of the subsequent expansion engagement, and no cross-region recovery objective "
                 "is committed here.")
    h3("4.13 Single points of failure removed")
    add_body_paragraph(doc, "The hardening removes the baseline's single points of failure:")
    add_data_table(doc, ["Baseline SPOF", "Removed by"],
              [["Single EC2 instance", "Cross-AZ Auto Scaling Group (min 2) behind an ALB"],
               ["Single Availability Zone", "Two AZs; serving continues from the survivor"],
               ["Single RDS database", "Multi-AZ deployment with automatic failover"],
               ["Media on one instance's EBS volume", "Media offloaded to shared S3"]],
              widths=[7.0, 9.0])
    h3("4.14 Configuration decisions left to the implementer")
    add_data_table(doc, ["#", "Decision", "Why left open"],
              [["C1", "EC2 instance type and ASG min/max/scaling policy", "Size against the website's load + peak"],
               ["C2", "RDS instance class", "Size from the CMS data footprint + growth"],
               ["C3", "Administrative SSH access path to the private instances", "Adapt to YAT ICT's admin arrangements"],
               ["C4", "Media-migration cutover approach (EBS → S3)", "Minimise disruption to the live site"]],
              widths=[1.0, 7.0, 8.0])

    h1("5. Implementation Sequencing")
    add_body_paragraph(doc, "Because this hardens a live website, the work is sequenced to avoid an outage and to steer "
                 "clear of the enrolment-enquiry peak:")
    add_bullet_list(doc, [
        "Add the second Availability Zone's subnets, a NAT Gateway per AZ, and the security groups.",
        "Convert the RDS instance to a Multi-AZ deployment (a managed change, taken in a maintenance window).",
        "Offload uploaded media from the instance EBS volume to the S3 media bucket and repoint the CMS.",
        "Stand up the Auto Scaling Group across both AZs and the internet-facing ALB; move TLS to an ACM certificate on the ALB.",
        "Cut traffic over to the ALB (DNS), retire the single instance's Elastic IP, and verify.",
    ])

    h1("6. Verification Plan")
    add_body_paragraph(doc, "The HA properties are verified by simulation: terminate an instance and confirm the ASG "
                 "replaces it with no loss of service; fail an Availability Zone and confirm serving continues "
                 "from the survivor; trigger an RDS failover and confirm the standby takes over within the "
                 "expected window; confirm media is served identically from any instance via S3.")

    h1("7. Out of Scope")
    add_body_paragraph(doc, "Stated explicitly — these are deliberately left for the subsequent Website Global Expansion "
                 "engagement, which works from this HA-hardened state:")
    add_bullet_list(doc, [
        "Global serving and edge content delivery (CDN) for an international / India audience.",
        "Cross-region disaster recovery: a second region, recovery objectives (RTO/RPO), and a DR runbook.",
        "The audit/access-log microservice and the in-India log store.",
        "Infrastructure-as-code templating of the footprint.",
        "And, out of MTS scope entirely: website content / CMS and ongoing support (YAT).",
    ])

    h1("8. References")
    add_bullet_list(doc, [
        "Website Cloud Architecture — Baseline Design (single-AZ) — the predecessor state.",
        "Website Specification; Website Infrastructure Specifications.",
        "Privacy Policy; User Access Policy; Industry Standards Reference (AWS Well-Architected, ACSC Essential Eight).",
    ])

    h1("Document control")
    add_data_table(doc, ["Field", "Value"],
              [["Document version", "v1.0 — As-built (HA-hardened)"],
               ["Authored by", "MTS Consulting, in consultation with YAT ICT"],
               ["Approved by", "Pat Lin (MTS Senior Consultant) · Sam Walker (YAT ICT Manager)"],
               ["Predecessor", "YAT Website Cloud Architecture — Baseline Design (single-AZ)"],
               ["Status", "The website's HA-hardened state of record (Multi-AZ; in-region HA, no cross-region DR)"]],
              widths=[5.0, 11.0])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website-s1/public/documents/YAT-Website-HA-Solution-Design.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
