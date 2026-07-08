#!/usr/bin/env python3
"""Build the AccentLoitte LMS Replacement Solution Design (.docx) — student model.

The on-premises solution design for the 2022 GrayBoard -> DOODLE replacement, rendered on
the generic Solution Design template. Student-facing model (in-world, no UoC tags),
companion to the AccentLoitte Business Case + presentation. Demonstrates the template on a
NON-cloud deployment: cloud-only sections (load balancing, SPOF-removal, simulation) are
marked "Not applicable"; §5 Implementation sequencing IS applicable (cutover from a live
GrayBoard). Output to public/documents/ for printing to PDF.

Usage:  python scripts/scenario/build_lms_replacement_solution_design.py [output.docx]
Default: ../diploma-cloud-cyber-website-s1/public/documents/YAT-LMS-Replacement-Solution-Design.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))  # content-repo scripts/ (brand + registry)  # noqa: E402
sys.path.insert(0, str(next(d / "scripts" for d in Path(__file__).resolve().parents if (d / "scripts" / "helpers" / "__init__.py").exists())))  # umbrella scripts/ (engine)  # noqa: E402
from helpers.docx_body_text import add_body_paragraph, add_bullet_list  # noqa: E402
from helpers.docx_tables import add_data_table  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from brand import ADDRESS, CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def na(doc, reason):
    p = doc.add_paragraph()
    r = p.add_run(f"Not applicable — {reason}")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    p.paragraph_format.space_after = Pt(6)
    return p


def build(path):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Solution Design")
    sub = doc.add_paragraph().add_run("LMS Replacement — GrayBoard to DOODLE (on-premises)")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT LMS Replacement Project — Phase 1"),
        ("Document type", "Technical design (Solution Design)"),
        ("Version", "v1.0 — Approved for implementation"),
        ("Authored by", "AccentLoitte Consulting, in consultation with YAT ICT"),
        ("Approved by", "Jamie O'Donnell (AccentLoitte) · Sam Walker (YAT ICT Manager)"),
        ("Date", "June 2022"),
        ("Implemented by", "the LMS Replacement Deployment Report"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Purpose and Scope")
    add_body_paragraph(doc, "This document specifies the solution for replacing YAT's end-of-life GrayBoard LMS with "
                 "the DOODLE platform, deployed on-premises at the Cremorne campus. It translates the "
                 "board-approved recommendation from the LMS Replacement Business Case into an implementable "
                 "design covering the infrastructure, the DOODLE installation approach, the migration of "
                 "student records from GrayBoard, and the cutover.")
    h3("In scope of this design")
    add_bullet_list(doc, [
        "On-premises infrastructure for DOODLE at the Cremorne campus (server, storage, network placement, backup).",
        "DOODLE installation and base configuration; integration with Active Directory, Office 365 and AVETMISS reporting.",
        "Migration of all student records from the legacy GrayBoard LMS, with reconciliation.",
        "The cutover plan, scheduled into the December–January teaching break.",
    ])
    h3("Out of scope of this design")
    add_bullet_list(doc, [
        "Cloud hosting — noted as a future consideration, not part of this engagement.",
        "Application-layer high availability / clustering.",
        "Ongoing operational support after acceptance (separately contracted).",
    ])

    h1("2. Design Inputs and Requirements")
    h3("2.1 Inputs")
    add_bullet_list(doc, [
        "LMS Replacement Requirements — the functional and non-functional targets.",
        "Engagement Role Brief — scope, phases, on-premises deployment direction.",
        "ICT Environment Overview — the current campus environment the LMS sits within.",
    ])
    h3("2.2 Requirements the design must meet")
    add_data_table(doc, ["Requirement", "Target / note"],
              [["Replacement platform", "A current, vendor-supported LMS — DOODLE (selected via market evaluation)"],
               ["Deployment", "On-premises at the Cremorne campus; Windows Server 2016 hosting standard"],
               ["Record preservation", "All GrayBoard student records migrated with full integrity (counts reconcile)"],
               ["Availability", "≥ 99% (a meaningful step up from the degrading GrayBoard)"],
               ["Accessibility", "WCAG 2.1 Level AA (Disability Discrimination Act 1992)"],
               ["Integrations", "Active Directory SSO · Office 365 SMTP · AVETMISS/ASQA export · LDAP user groups"],
               ["Performance", "Match or exceed GrayBoard, including ~3× concurrent load at end-of-term peaks"],
               ["Vendor support", "Severity-1 response ≤ 4 h (teaching hours) / ≤ 8 h (outside)"],
               ["Cutover", "During the December–January teaching break; training delivered before cutover"]],
              widths=[4.5, 11.5])

    h1("3. Review of Existing Architecture")
    na(doc, "this is a replacement, not a hardening of an existing system — GrayBoard is decommissioned, "
            "not extended. The assessment of the legacy GrayBoard platform is in the LMS Replacement "
            "Business Case.")

    h1("4. Architecture Design")
    h3("4.1 Assumptions and constraints")
    add_data_table(doc, ["#", "Assumption / constraint", "Source"],
              [["A1", "Deployed on-premises at the Cremorne campus (current operating model)", "Role Brief"],
               ["A2", "Windows Server 2016 application-hosting standard", "Replacement Requirements"],
               ["A3", "All GrayBoard student records must be preserved with integrity", "Standards for RTOs 2015"],
               ["A4", "Authentication via Active Directory; preserve LDAP group structures", "Replacement Requirements"],
               ["A5", "Cutover confined to the December–January teaching break", "Replacement Requirements"]],
              widths=[1.0, 9.0, 6.0])
    h3("4.2 Deployment environment")
    add_body_paragraph(doc, "The DOODLE server is deployed in the YAT Cremorne campus computer room — physically "
                 "secured, air-conditioned, UPS-protected — consistent with the existing ICT environment. "
                 "It sits in the staff network zone, with student-zone access to the LMS application "
                 "governed by Active Directory permissions.")
    h3("4.3 Identity and access")
    add_bullet_list(doc, [
        "DOODLE authenticates against Active Directory via LDAP bind — single sign-on for staff and students.",
        "Existing LDAP user-group structures are preserved (per the User Access Policy in force).",
        "MFA for staff with grading or course-management roles, per the User Access Policy.",
    ])
    h3("4.4 Network")
    add_bullet_list(doc, [
        "The server uses the redundant campus network (dual NICs) with no single point of failure at the network layer.",
        "Students reach the LMS across the zone boundary under AD permissions; staff reach it from the staff zone.",
        "Office 365 SMTP relay for LMS notifications; outbound access for vendor updates via the campus firewall.",
    ])
    h3("4.5 Compute (server)")
    add_bullet_list(doc, [
        "A new mid-range enterprise server (dual PSU, RAID-protected disks), running Windows Server 2016, hosts DOODLE and its MySQL database.",
        "Sized to match or exceed GrayBoard's performance, including the ~3× concurrent load during end-of-term assessment windows.",
    ])
    h3("4.6 Load balancing")
    na(doc, "single-server on-premises deployment — no load balancer in scope (this is not a "
            "horizontally-scaled or highly-available design).")
    h3("4.7 Database")
    add_bullet_list(doc, [
        "MySQL on the application server (DOODLE's supported database); engine version confirmed against DOODLE's compatibility matrix.",
        "Populated by migrating the student records from the legacy GrayBoard LMS (see §5), with post-migration reconciliation.",
    ])
    h3("4.8 Storage")
    add_bullet_list(doc, [
        "RAID-protected local disk for the OS, the DOODLE application, and the MySQL database.",
        "The campus NAS (RAID-5) holds course attachments, student submissions, and backup copies.",
    ])
    h3("4.9 Security")
    add_bullet_list(doc, [
        "Active Directory permissions and the two-zone separation control access; antivirus/EDR on the server.",
        "HTTPS for LMS application traffic; TLS for the MySQL connection.",
        "WCAG 2.1 Level AA conformance at the application layer (a DOODLE selection criterion).",
    ])
    h3("4.10 Monitoring")
    add_body_paragraph(doc, "The existing YAT system-management and monitoring server tracks the DOODLE server's health "
                 "and runs the nightly backups; availability is tracked against the ≥ 99% target.")
    h3("4.11 Naming and asset conventions")
    add_body_paragraph(doc, "The server and related assets are recorded in the YAT ICT asset register under the "
                 "standard naming and ownership conventions.")
    h3("4.12 Backup")
    add_bullet_list(doc, [
        "Nightly backup of the DOODLE database and application data to the system-management server, with offsite tape rotation — aligned to the existing Backup and Retention Policy.",
        "GrayBoard is retained (read-only) until acceptance, as the migration rollback position (see §5).",
    ])
    h3("4.13 Recovery objectives")
    add_body_paragraph(doc, "Recovery is backup-based: nightly backups give an RPO of up to one day, with restore from "
                 "the system-management server or offsite tape. The design targets ≥ 99% availability — a "
                 "step up from the degrading GrayBoard.")
    h3("4.14 Components requiring vertical scaling")
    add_body_paragraph(doc, "The server is provisioned with headroom for growth; storage on the NAS is expandable. A "
                 "future capacity increase is a server resource upgrade in a maintenance window.")
    h3("4.15 Single points of failure removed")
    na(doc, "this is a supported single-server replacement, not a high-availability design; resilience "
            "beyond the existing campus posture (redundant network, RAID, UPS) is out of scope — cloud HA "
            "is noted as a future consideration.")
    h3("4.16 Configuration decisions left to the implementer")
    add_data_table(doc, ["#", "Decision", "Why left open"],
              [["C1", "Server specification (CPU / RAM / disk)", "Size against the GrayBoard performance baseline + peak load"],
               ["C2", "MySQL engine version", "Confirm against the DOODLE compatibility matrix"],
               ["C3", "Storage sizing (local + NAS)", "Compute from the migrated data footprint + growth"],
               ["C4", "Backup schedule specifics", "Align with the existing backup windows (per the Backup and Retention Policy)"]],
              widths=[1.0, 7.0, 8.0])

    h1("5. Implementation Sequencing")
    add_body_paragraph(doc, "The replacement is a change to a live service — GrayBoard is in use until cutover — so the "
                 "sequence and rollback matter. The work is confined to the December–January teaching break, "
                 "with GrayBoard retained read-only as the rollback position until acceptance.")
    add_data_table(doc, ["#", "Step", "Impact", "Verification", "Rollback"],
              [["1", "Provision and harden the server; install Windows Server 2016", "None (parallel)", "Server build checklist", "n/a"],
               ["2", "Install and configure DOODLE; wire AD / O365 / AVETMISS integrations", "None (parallel)", "Integration tests", "n/a"],
               ["3", "Migrate GrayBoard student records; reconcile counts + sample fidelity", "None (parallel)", "Record reconciliation", "Re-run migration"],
               ["4", "Parallel-run DOODLE alongside GrayBoard; deliver staff training", "None", "Acceptance testing; trained staff", "Continue on GrayBoard"],
               ["5", "Cutover to DOODLE (Dec–Jan break); redirect users", "Planned downtime in break", "Smoke + acceptance", "Revert to GrayBoard"],
               ["6", "Decommission GrayBoard after acceptance", "None", "Sign-off received", "n/a"]],
              widths=[0.8, 5.2, 3.0, 3.5, 3.5])

    h1("6. Simulation and Verification Plan")
    na(doc, "no high-availability failure/resize simulation applies to this single-server on-premises "
            "design; verification (record reconciliation, performance baseline, accessibility audit) is "
            "documented in the Deployment Report's testing section.")

    h1("7. Out of Scope")
    add_bullet_list(doc, [
        "Cloud hosting — a future consideration, not this engagement.",
        "Application-layer high availability / clustering.",
        "Ongoing operational support after acceptance of the Closure Pack (separately contracted).",
    ])

    h1("8. References")
    add_bullet_list(doc, [
        "LMS Replacement Requirements; Engagement Role Brief; ICT Environment Overview.",
        "Privacy / Data Handling Policy; User Access Policy; Backup and Retention Policy.",
        "Standards for RTOs 2015 (record retention); Disability Discrimination Act 1992 (WCAG 2.1 AA).",
    ])

    h1("Document control")
    add_data_table(doc, ["Field", "Value"],
              [["Document version", "v1.0 — Approved for implementation"],
               ["Authored by", "AccentLoitte Consulting, in consultation with YAT ICT"],
               ["Approved by", "Jamie O'Donnell (AccentLoitte) · Sam Walker (YAT ICT Manager)"],
               ["Implemented by", "the LMS Replacement Deployment Report"],
               ["Successor document", "LMS Replacement Deployment Report (Phase 2–3)"]],
              widths=[5.0, 11.0])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website-s1/public/documents/YAT-LMS-Replacement-Solution-Design.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
