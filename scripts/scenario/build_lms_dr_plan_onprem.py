#!/usr/bin/env python3
"""Build the YAT on-premises LMS Disaster Recovery Plan (.docx) — in-world document.

The existing YAT LMS DR plan for the pre-cloud, on-premises platform, expressed in the YAT
DR Plan template structure so the same format runs end-to-end (existing plan → the plan
students develop). In-world (no UoC tags, no exemplar marking). Content is derived from the
on-prem Backup and Recovery Process: nightly tape backups, a 7–11 h worst-case RTO and a
24 h RPO — deliberately short of the ICT Strategic Plan targets (RTO ≤ 4 h / RPO ≤ 1 h),
which is why it needs replacing once the LMS moves to the cloud.

Pass deprecated=True (or `--deprecated`) to stamp the superseded banner for the post-cutover
states (CL1 AT3, CL2 AT1), where the LMS has migrated and this plan no longer applies.

Usage:  python scripts/scenario/build_lms_dr_plan_onprem.py [output.docx] [--deprecated]
Default: ../diploma-cloud-cyber-website-s1/public/documents/YAT-LMS-DR-Plan-OnPrem.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_body_paragraph  # noqa: E402
from helpers.docx_tables import add_data_table  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from brand import ADDRESS, CHARCOAL, CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def deprecated_banner(doc):
    """Full-width terracotta DEPRECATED banner."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade_cell(cell, "F6E5DF")  # pale terracotta
    set_cell_borders(cell, TERRACOTTA, sz=12)
    p = cell.paragraphs[0]
    h = p.add_run("DEPRECATED — REPLACEMENT REQUIRED")
    h.bold = True; h.font.size = Pt(12); h.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    b = cell.add_paragraph().add_run(
        "This disaster recovery plan applies to the former on-premises LMS, which has since been "
        "migrated to the cloud. It no longer reflects the current platform, its recovery options, "
        "or the recovery objectives now achievable, and must not be relied on for recovery. A "
        "replacement DR plan for the cloud-hosted LMS is required.")
    b.font.size = Pt(10); b.font.color.rgb = RGBColor.from_string(CHARCOAL)
    doc.add_paragraph()


def build(path, deprecated=False):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(2):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Disaster Recovery Plan")
    sub = doc.add_paragraph().add_run("YAT Learning Management System (On-Premises)")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(TEAL)
    doc.add_paragraph()
    if deprecated:
        deprecated_banner(doc)
    status = ("Deprecated — superseded by the LMS cloud migration; replacement required"
              if deprecated else "Current")
    cover = [
        ("System", "YAT Learning Management System (on-premises)"),
        ("Document owner", "Sam Walker, ICT Manager"),
        ("Prepared by", "YAT ICT"),
        ("Plan version", "v2.3"),
        ("Last reviewed", "[ DD/MM/YYYY ]"),
        ("Status", status),
        ("Classification", "Internal — ICT, and engaged consultants on signed MSA"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    if deprecated:
        deprecated_banner(doc)
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Executive Summary")
    add_body_paragraph(doc, "This plan sets out the disaster recovery arrangements for the YAT on-premises Learning "
                 "Management System (LMS). The LMS is backed up in full every night to tape, with tapes "
                 "rotated offsite, and can be fully recovered from backup following a major incident such as "
                 "server hardware failure, ransomware, or irrecoverable database corruption.")
    add_body_paragraph(doc, "Current recovery performance — a worst-case Recovery Time Objective of 7–11 hours and a "
                 "Recovery Point Objective of up to 24 hours — sits materially below the ICT Strategic Plan "
                 "targets of RTO ≤ 4 hours and RPO ≤ 1 hour. The dominant gaps are the offsite tape-retrieval "
                 "step (RTO) and the nightly-only backup cadence (RPO). Closing these gaps is an explicit "
                 "driver of the planned LMS cloud migration.")

    h1("2. Engagement Context and Scope")
    add_body_paragraph(doc, "The YAT LMS (the DOODLE application on Windows Server 2016 with a MySQL database) runs "
                 "on-premises at the Cremorne campus and is a mission-critical system for staff and students. "
                 "This plan covers disaster recovery for that system — the backup arrangements, the recovery "
                 "steps, and the recovery objectives currently achieved.")
    add_body_paragraph(doc, "In scope: the LMS server, its MySQL database, the LMS application, and the LMS attachments "
                 "(course materials and student submissions). Out of scope: unrelated campus systems, which "
                 "are covered by their own arrangements, and any future cloud-hosted LMS, which is not yet in "
                 "place.")

    h1("3. Current Recovery Position")
    h3("3.1 Backup arrangements")
    add_body_paragraph(doc, "A complete, full backup of the LMS environment runs nightly (22:00 – 04:00 local time), "
                 "captured to tape by the campus System Management server. The backup covers the Windows "
                 "Server 2016 system state, the MySQL database (DOODLE schema and content), the LMS "
                 "application binaries and configuration, and the LMS attachments. Tapes are rotated offsite "
                 "each business day under the Backup and Retention Policy. The LMS remains available during "
                 "the backup window; the database is captured against a brief consistency snapshot.")
    h3("3.2 Suppliers and dependencies")
    add_body_paragraph(doc, "Recovery depends on the offsite tape-storage provider (tape retrieval to the Cremorne "
                 "campus takes 4–8 hours), on the availability of replacement server hardware, and on the "
                 "campus Active Directory for post-restore authentication. There is no cloud or vendor-provided "
                 "disaster-recovery capability for the on-premises LMS.")

    h1("4. Impact Analysis")
    h3("4.1 Recovery objectives")
    add_body_paragraph(doc, "The ICT Strategic Plan sets the targets at RTO ≤ 4 hours and RPO ≤ 1 hour. The current "
                 "on-premises process achieves a worst-case RTO of 7–11 hours and an RPO of up to 24 hours — "
                 "both materially short of target.")
    add_data_table(doc, ["Objective", "Current achievement", "Target"],
              [["RTO — worst case", "7–11 hours from incident declaration to LMS operational", "≤ 4 hours"],
               ["RPO — worst case", "Up to 24 hours (anything written since the last nightly backup)", "≤ 1 hour"]],
              widths=[4.0, 7.5, 4.0])
    h3("4.2 Data managed")
    add_body_paragraph(doc, "The protected data comprises the MySQL DOODLE database (schema and content) and the LMS "
                 "attachments — course materials and student submissions. This includes student personal "
                 "information and academic records, handled under the Privacy / Data Handling Policy and the "
                 "Privacy Act.")
    h3("4.3 Risk assessment")
    add_data_table(doc, ["#", "Risk event", "Likelihood", "Impact", "Severity", "What it drives"],
              [["RE1", "LMS server hardware failure", "Medium", "High", "High", "Restore to replacement hardware"],
               ["RE2", "Ransomware / malware compromise", "Medium", "Critical", "High", "Clean restore from offsite tape"],
               ["RE3", "Irrecoverable database corruption or deletion", "Medium", "High", "High", "Database restore from nightly backup"],
               ["RE4", "Campus-level event (fire / flood at Cremorne)", "Low", "Critical", "High", "Offsite tape copy; single-site exposure"]],
              widths=[0.9, 4.6, 2.0, 1.8, 1.8, 4.4])
    h3("4.4 Exclusions")
    add_body_paragraph(doc, "This plan excludes unrelated campus systems (covered by their own arrangements) and routine "
                 "single-file or partial restores (handled operationally). It assumes the offsite tapes are "
                 "intact and retrievable, and that replacement hardware can be sourced.")

    h1("5. Recovery Strategy")
    h3("5.1 Current strategy")
    add_body_paragraph(doc, "The current strategy is restore-from-tape onto replacement hardware at the Cremorne campus: "
                 "a single-site, tape-based recovery. It relies entirely on the nightly tape backup and the "
                 "offsite rotation for any protection against a campus-level event.")
    h3("5.2 Assessment against the objectives")
    add_body_paragraph(doc, "The strategy does not meet the targets. The offsite tape-retrieval step (4–8 hours) is the "
                 "dominant contributor to the RTO gap; the nightly-only backup cadence is the dominant "
                 "contributor to the RPO gap; and a single campus site offers no rapid recovery option for a "
                 "site-level disaster. The arrangement is adequate as a last-resort restore but cannot meet "
                 "the strategic RTO/RPO targets.")
    h3("5.3 Improvement direction")
    add_body_paragraph(doc, "Closing the RTO/RPO gap is an explicit driver of the planned LMS cloud migration. A "
                 "cloud-hosted platform is expected to provide faster recovery and a tighter recovery point "
                 "than tape allows; the forward design work for the migrated LMS is required to demonstrate "
                 "how the targets will be met. This plan will need to be replaced once the LMS is migrated.")

    h1("6. The Disaster Recovery Plan")
    h3("6.1 Detection and declaration")
    add_body_paragraph(doc, "A major LMS incident is identified through system monitoring and the ICT helpdesk. The ICT "
                 "Manager declares a disaster and authorises the recovery, notifying affected stakeholders.")
    h3("6.2 Recovery steps")
    add_body_paragraph(doc, "To completely recover the LMS from backup, the following steps are executed in sequence:")
    add_data_table(doc, ["Step", "Action", "Owner", "Duration"],
              [["1", "Request the relevant tape from offsite storage and have it delivered to the Cremorne campus", "YAT ICT", "4–8 hours"],
               ["2", "Restore the Windows Server 2016 system state and LMS application binaries from tape onto replacement hardware", "YAT ICT", "1 hour"],
               ["3", "Restore the MySQL database from the latest nightly backup and verify schema integrity", "YAT ICT", "1 hour"],
               ["4", "Verify LMS connectivity to MySQL, re-authenticate against Active Directory, and confirm end-user sign-in", "YAT ICT", "1 hour"]],
              widths=[1.0, 8.5, 2.5, 2.8])
    h3("6.3 Meeting the recovery objectives")
    add_body_paragraph(doc, "The plan does not currently meet the strategic objectives: the cumulative recovery time of "
                 "7–11 hours exceeds the RTO ≤ 4 hours, and the nightly backup cadence allows up to 24 hours "
                 "of data loss against the RPO ≤ 1 hour. The gap is documented here as the basis for the "
                 "improvement work.")

    h1("7. Plan Validation and Approval")
    h3("7.1 Review")
    add_body_paragraph(doc, "This plan is reviewed annually, or on a material change to the LMS environment or backup "
                 "tooling, and is walked through with YAT ICT as part of the review.")
    h3("7.2 Sign-off")
    add_data_table(doc, ["Role", "Name", "Date", "Acceptance"],
              [["Prepared by", "YAT ICT", "[date]", "Issued"],
               ["Approved by", "Sam Walker (YAT ICT Manager)", "[date]", "Current DR arrangements accepted; gap to target noted"]],
              widths=[4.0, 5.0, 2.5, 4.0])

    h1("Document control")
    add_data_table(doc, ["Field", "Value"],
              [["Document version", "v2.3"],
               ["Document owner", "Sam Walker, ICT Manager"],
               ["Review cycle", "Annual, or on material change to the LMS environment or backup tooling"],
               ["Status", status],
               ["Related documents", "Backup and Retention Policy; ICT Strategic Plan; LMS Server Specifications and Current Status"]],
              widths=[5.0, 10.5])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    args = [a for a in sys.argv[1:] if a != "--deprecated"]
    dep = "--deprecated" in sys.argv[1:]
    default = "../diploma-cloud-cyber-website-s1/public/documents/YAT-LMS-DR-Plan-OnPrem.docx"
    out = args[0] if args else default
    build(out, deprecated=dep)
