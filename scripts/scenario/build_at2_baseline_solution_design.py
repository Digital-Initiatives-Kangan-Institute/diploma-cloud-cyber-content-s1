#!/usr/bin/env python3
"""Build the supplied AT2 Baseline Solution Design (.docx) — in-world document.

The cloud architecture design a second MTS consultant produced while the student was on
leave, which the student implements in AT2. An in-world supplied artefact (NOT an
exemplar — no UoC tags): the branded Solution Design document, rendered from the existing
baseline-design content. Being a greenfield single-AZ design, the HA sections (§3 review,
§5 sequencing, §6 simulation, §4.15) are marked "Not applicable" — demonstrating the
convention from the design side, and modelling the Solution Design genre for AT3 (where
the student authors their own HA Solution Design).

Output to the website documents folder for printing to PDF, then wiring into the AT2+
intranet state in place of the current markdown page.

Usage:  python scripts/scenario/build_at2_baseline_solution_design.py [output.docx]
Default: ../diploma-cloud-cyber-website-s1/public/documents/YAT-LMS-Baseline-Solution-Design.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_body_paragraph, add_bullet_list  # noqa: E402
from helpers.docx_tables import add_data_table  # noqa: E402
from helpers.docx_styling import add_field, paragraph_bottom_rule, set_cell_borders, shade_cell  # noqa: E402
from brand import ADDRESS, CREAM, GREY, TEAL, TERRACOTTA  # noqa: E402
from helpers.scenario_document import build_header_footer, configure_styles, wordmark  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def na(doc, reason):
    p = doc.add_paragraph()
    r = p.add_run(f"Not applicable — {reason}")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    p.paragraph_format.space_after = Pt(6)
    return p


def diagram_placeholder(doc, caption, source):
    """A bordered, shaded drop-zone for a network diagram — replace with the exported image."""
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    set_cell_borders(cell)
    shade_cell(cell, CREAM)
    cell.width = Cm(16.6)
    t.rows[0].height = Cm(9.5)
    t.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[ NETWORK TOPOLOGY DIAGRAM — PASTE HERE ]")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(source)
    r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = RGBColor.from_string(GREY)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = RGBColor.from_string(GREY)
    return t


def build(path):
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(GREY)
    paragraph_bottom_rule(doc.add_paragraph(), TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Solution Design")
    sub = doc.add_paragraph().add_run("YAT LMS Cloud Architecture — Baseline Design")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(TERRACOTTA)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT LMS Cloud Migration — Foundation Build"),
        ("Document type", "Technical design (Solution Design)"),
        ("Version", "v1.0 — Approved for implementation"),
        ("Authored by", "MTS Senior Architecture Team, in consultation with YAT ICT"),
        ("Approved by", "Pat Lin (MTS Senior Consultant) · Sam Walker (YAT ICT Manager)"),
        ("Implemented by", "the AT2 Deployment Report (foundation build)"),
        ("Classification", "Internal — YAT ICT, and MTS personnel on signed MSA"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1]); shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Purpose and Scope")
    add_body_paragraph(doc, "This document specifies the baseline AWS architecture MTS will implement as the first "
                 "build phase of the YAT LMS cloud migration. It translates the action plan approved at the "
                 "close of the Business Case engagement into a concrete, implementable design. The design "
                 "stops at “infrastructure ready for application deployment”: the EC2 instance is provisioned "
                 "with the OS, the RDS instance with an empty MySQL schema, and the ALB with placeholder "
                 "health checks — no application binaries or production data are placed by the MTS build.")
    h3("In scope of this design")
    add_bullet_list(doc, [
        "The production cloud foundation for the DOODLE LMS application.",
        "All compute, networking, identity, storage, database, autoscaling and monitoring needed to run the LMS as a multi-tier web workload in AWS.",
        "Single-region, single-Availability-Zone deployment in ap-southeast-2 (Sydney).",
    ])
    h3("Out of scope — deferred to the follow-on HA design phase")
    add_bullet_list(doc, [
        "High-availability hardening (Multi-AZ database, cross-AZ compute resilience, failure-simulation testing).",
        "Disaster recovery to a second AWS region; DR runbook and tabletop testing.",
        "Application re-platforming (the LMS remains Windows Server 2016 + DOODLE + MySQL).",
    ])
    h3("Out of MTS scope entirely — YAT ICT responsibility")
    add_bullet_list(doc, [
        "LMS application installation onto the EC2 instance(s) after handover.",
        "Database migration (extract from on-prem MySQL, load into RDS).",
        "Cutover — DNS switch, parallel running, decommissioning, user redirection.",
        "Organisational change management — CAB approvals, communications, training, post-cutover support.",
    ])

    h1("2. Design Inputs and Requirements")
    h3("2.1 Inputs")
    add_bullet_list(doc, [
        "LMS Application Specification — workload, SLAs, data footprint, integration points.",
        "LMS Cloud Migration Requirements — SLA, RPO/RTO targets.",
        "Engagement Role Brief — engagement scope, OS and application preservation.",
        "ICT Environment Overview and the On-Premises Network Diagram — the current state being migrated from.",
    ])
    h3("2.2 Requirements the design must meet")
    add_data_table(doc, ["Requirement", "Target / note"],
              [["Region / data residency", "ap-southeast-2 (Sydney); all student PII within Australia (Privacy Act, APP 8)"],
               ["Application stack", "Preserved — Windows Server 2016 · MySQL · DOODLE"],
               ["Concurrent users", "200–300 typical; 500–700 peak during assessment windows"],
               ["Data footprint", "~178 GB, growing ~25 GB/year"],
               ["Authentication", "Preserve AD-LDAP integration over a private network to campus AD"],
               ["High availability", "Out of baseline scope — the 99.9% target is deferred to the follow-on HA design"]],
              widths=[5.0, 11.0])

    h1("3. Review of Existing Architecture")
    na(doc, "this is a greenfield cloud foundation; there is no existing cloud baseline to review. The "
            "on-premises current state is documented separately in the ICT Environment Overview.")

    h1("4. Architecture Design")
    h3("4.1 Assumptions and constraints")
    add_data_table(doc, ["#", "Assumption / constraint", "Source"],
              [["A1", "Region must be ap-southeast-2 (Sydney) for data residency", "LMS App Spec; Privacy Policy"],
               ["A2", "Application stack preserved: Windows Server 2016, MySQL, DOODLE", "Engagement Role Brief"],
               ["A3", "Data footprint ~178 GB, +~25 GB/year (attachments + submissions dominate)", "LMS App Spec"],
               ["A4", "Concurrent users 200–300 typical, 500–700 peak", "LMS App Spec"],
               ["A5", "All student PII must remain within Australia", "Privacy Act 1988, APP 8"],
               ["A6", "Preserve AD-LDAP integration over private network to campus AD", "LMS App Spec"],
               ["A7", "Baseline is NOT high-availability; 99.9% deferred to the HA design", "Scoping decision"]],
              widths=[1.0, 9.0, 6.0])
    h3("4.2 AWS account and region")
    add_body_paragraph(doc, "An AWS account scoped to the migration engagement, provisioned by YAT and shared with MTS "
                 "for the build. Region ap-southeast-2 (Sydney); no deployment outside Australian regions. "
                 "Availability Zone ap-southeast-2a for the baseline; the follow-on HA design introduces a "
                 "second AZ.")
    h3("4.3 Identity and Access Management (IAM)")
    add_data_table(doc, ["Group", "Purpose", "Indicative permissions"],
              [["YAT-ICT-Admins", "YAT ICT day-to-day ops post-handover", "Read-only on infra; full CloudWatch/RDS/EC2 console; no IAM changes"],
               ["MTS-Consultants", "MTS during build + support", "Full admin during build; reduced post-handover"],
               ["Application-Service", "EC2 instance role for the LMS", "RDS read/write; S3 read/write (attachments); CloudWatch logs"],
               ["Read-Only-Auditors", "Compliance / external auditors", "Read-only on logs, metrics, configs"]],
              widths=[3.5, 5.0, 7.5])
    add_bullet_list(doc, [
        "MFA required for all YAT-ICT-Admins and MTS-Consultants users (Essential Eight; User Access Policy).",
        "No long-lived access keys for humans; programmatic access via IAM roles only; EC2 access via instance profile.",
        "Configuration decision left to the implementer: the MTS-Consultants permission boundary during build vs after handover.",
    ])
    h3("4.4 Network topology")
    add_body_paragraph(doc, "VPC 10.0.0.0/16 (room to expand), with DNS hostnames/resolution and VPC flow logs enabled. "
                 "Single-AZ subnets in ap-southeast-2a:")
    add_data_table(doc, ["Subnet", "CIDR", "Tier", "Internet-facing?"],
              [["public-web-a", "10.0.1.0/24", "Web / public load balancer", "Yes"],
               ["private-app-a", "10.0.11.0/24", "Application / LMS EC2", "No"],
               ["private-data-a", "10.0.21.0/24", "Database (RDS)", "No"]],
              widths=[4.0, 3.5, 5.5, 3.0])
    add_bullet_list(doc, [
        "Internet Gateway for public-subnet traffic; NAT Gateway in public-web-a for private-app outbound (Windows Update etc.).",
        "Route tables: public → IGW; private-app → NAT; private-data → no internet route.",
        "Connectivity to campus AD: Site-to-Site VPN (baseline choice); Direct Connect deferred unless latency requires it.",
        "The follow-on HA design adds the corresponding -b subnets in ap-southeast-2b.",
    ])
    diagram_placeholder(doc,
                        "Figure 4.4 — LMS baseline network topology (single-AZ).",
                        "Source diagram: network-at3-start-non-hardened.drawio")
    h3("4.5 Compute (EC2 + Auto Scaling)")
    add_bullet_list(doc, [
        "EC2: general-purpose x86 (e.g. m6i.large — final type a C1 implementer decision); Windows Server 2016 AMI; placed in private-app-a (no public IP).",
        "EBS: gp3 root 100 GB + a gp3 data volume sized by the implementer (footprint + 12-month growth + 50% headroom).",
        "Auto Scaling Group: min 1 / desired 1 / max 2 (baseline); target-tracking on CPU at 70%; ELB+EC2 health checks; 300 s cooldown.",
        "The follow-on HA design expands the ASG (min 2, multi-AZ) and tunes scaling for assessment-window peaks.",
    ])
    h3("4.6 Load balancing (ALB)")
    add_bullet_list(doc, [
        "Internet-facing ALB in public-web-a; HTTPS:443 listener forwarding to the LMS target group.",
        "Target group = the ASG instances; HTTP health check on the LMS health endpoint (30 s; 2 unhealthy → out of service).",
        "TLS via an ACM-issued certificate for the LMS DNS name (DNS strategy a C8 implementer decision).",
    ])
    h3("4.7 Database (RDS)")
    add_bullet_list(doc, [
        "Amazon RDS for MySQL (preserves the existing data/schema); engine version confirmed against DOODLE at build time.",
        "General-purpose instance class (e.g. db.m6i.large — a C2 implementer decision); gp3 storage sized to footprint + growth.",
        "Multi-AZ DISABLED for the baseline (enabled in the HA design); storage encryption enabled (KMS).",
        "Placed in private-data-a; not publicly accessible; 7-day automated backups (window 22:00–04:00 AEST); maintenance Sun 02:00–06:00 AEST.",
        "Schema and data migration are YAT ICT's responsibility, not MTS's — MTS provisions an empty instance.",
    ])
    h3("4.8 Storage")
    add_data_table(doc, ["Resource", "Type", "Purpose"],
              [["EC2 root volume", "EBS gp3 100 GB", "OS and LMS application install"],
               ["EC2 data volume", "EBS gp3 (sized by implementer)", "LMS application data, attachments staging"],
               ["Attachments bucket", "S3", "Course attachments + submissions; lifecycle to Glacier Deep Archive after 24 months"],
               ["Backups bucket", "S3", "Off-instance mysqldump exports, file snapshots"]],
              widths=[4.0, 4.5, 7.5])
    add_bullet_list(doc, ["Both buckets: block all public access; SSE-S3 encryption; versioning enabled; access logging to a log bucket."])
    h3("4.9 Security")
    add_data_table(doc, ["Security group", "Inbound", "Outbound"],
              [["sg-alb", "HTTPS:443 from 0.0.0.0/0", "HTTP/HTTPS to sg-app"],
               ["sg-app", "from sg-alb; RDP:3389 from MTS bastion (design left to implementer)", "MySQL:3306 to sg-db; HTTPS via NAT; LDAP/LDAPS to campus AD"],
               ["sg-db", "MySQL:3306 from sg-app only", "none"]],
              widths=[3.0, 7.5, 5.5])
    add_bullet_list(doc, [
        "Encryption in transit: HTTPS ALB→EC2; TLS EC2→RDS; AWS calls over TLS via VPC endpoints where available.",
        "Encryption at rest: EBS, RDS, and S3 (SSE-S3) all enabled.",
        "Operates under the AWS Shared Responsibility Model — AWS secures the cloud; YAT/MTS secure the OS, application, IAM, data and access in the cloud.",
    ])
    h3("4.10 Monitoring (baseline)")
    add_body_paragraph(doc, "Standard CloudWatch metrics for EC2, RDS, ALB and Auto Scaling. Baseline alarms (HA-tuned "
                 "alarms come in the follow-on HA design):")
    add_data_table(doc, ["Alarm", "Threshold"],
              [["EC2 CPU high", "≥ 80% over 10 min"],
               ["RDS CPU high", "≥ 80% over 10 min"],
               ["RDS free storage low", "< 15%"],
               ["ALB 5XX", "> 10 / min"],
               ["RDS connections high", "> 80% of max_connections"]],
              widths=[8.0, 8.0])
    add_bullet_list(doc, ["Logging: VPC flow logs and RDS logs → CloudWatch Logs (90-day retention); ALB access logs → S3; EC2 OS logs via the CloudWatch Agent."])
    h3("4.11 Naming and tagging conventions")
    add_body_paragraph(doc, "Naming pattern yat-lms-<resource-type>-<env>-<az-or-purpose> (e.g. yat-lms-alb-prod). Mandatory tags:")
    add_data_table(doc, ["Tag", "Value"],
              [["Project", "YAT-LMS-Migration"], ["Environment", "Production"], ["Owner", "YAT-ICT"],
               ["ManagedBy", "MTS-Migration during build → YAT-ICT post-handover"],
               ["CostCentre", "YAT-LMS"], ["DataClassification", "Confidential (PII) or Internal"]],
              widths=[5.0, 11.0])
    h3("4.12 Backup")
    add_data_table(doc, ["Resource", "Mechanism", "Retention"],
              [["RDS database", "Automated daily backups + transaction logs", "7 days"],
               ["EC2 EBS volumes", "Daily AMI snapshot (Data Lifecycle Manager)", "14 days"],
               ["LMS attachments (S3)", "Versioning + lifecycle to Glacier Deep Archive", "Versioned; archive after 24 months"]],
              widths=[4.5, 7.5, 4.0])
    add_bullet_list(doc, ["Cross-Region backup copies are out of scope for the baseline — addressed in the follow-on HA design."])
    h3("4.13 Recovery objectives — baseline state")
    add_body_paragraph(doc, "The baseline's recovery posture is backup-based: RPO is up to the last automated backup "
                 "(daily) and RTO is restore-based (single-AZ restore). The HA targets (99.9% availability, "
                 "RPO ≤ 1 h, RTO ≤ 4 h) are deliberately not met by this baseline and are the objective of "
                 "the follow-on HA design.")
    h3("4.14 Components requiring vertical scaling")
    add_body_paragraph(doc, "RDS instance-class changes require a modify-and-apply (a brief interruption, taken in the "
                 "maintenance window); EBS volumes support online resize with no downtime.")
    h3("4.15 Single points of failure removed")
    na(doc, "this baseline is single-AZ by design; the single RDS instance and the single AZ are known "
            "single points of failure, deliberately deferred to the follow-on HA design.")
    h3("4.16 Configuration decisions left to the implementer")
    add_body_paragraph(doc, "The design is opinionated where it matters and silent where the implementer must show "
                 "judgement. Each decision below is to be made and evidenced in the Deployment Report.")
    add_data_table(doc, ["#", "Decision", "Why left open"],
              [["C1", "EC2 instance type (general-purpose family)", "Size against the LMS workload"],
               ["C2", "RDS instance class (general-purpose family)", "Size against the workload"],
               ["C3", "EBS data volume + RDS storage size", "Compute from data footprint + growth"],
               ["C4", "ASG scaling threshold", "Rationalise against the expected CPU profile"],
               ["C5", "MTS-Consultants permission boundary", "Adapt to the lab access scope"],
               ["C6", "Bastion / jump-host design for RDP", "Left to the implementer"],
               ["C7", "MySQL engine version", "Confirm against the DOODLE compatibility matrix"],
               ["C8", "DNS strategy + ACM certificate domain", "Confirm the LMS hostname with YAT ICT"]],
              widths=[1.0, 7.0, 8.0])

    h1("5. Implementation Sequencing")
    na(doc, "this is a greenfield build in a new account, not a change to a running system; build order is "
            "at the implementer's discretion and there is no live service to sequence around or roll back.")

    h1("6. Simulation and Verification Plan")
    na(doc, "failure/resize simulation and availability verification are the subject of the follow-on HA "
            "design and its deployment report; the baseline build is verified functionally per the "
            "Deployment Report's testing section.")

    h1("7. Out of Scope")
    add_body_paragraph(doc, "Stated explicitly so the implementer knows what not to build (these are the deliberate "
                 "inputs to the follow-on HA design):")
    add_bullet_list(doc, [
        "Multi-AZ database; cross-AZ subnets (private-app-b, private-data-b); ASG capacity ≥ 2 across AZs.",
        "HA-tuned monitoring (cross-AZ latency, RDS replica lag); cross-Region backup copies and DR runbook.",
        "Failure-simulation testing; automated availability reporting against the 99.9% target.",
        "And, out of MTS scope entirely: LMS application install, data migration, cutover, and change management (YAT ICT).",
    ])

    h1("8. References")
    add_bullet_list(doc, [
        "LMS Application Specification; LMS Cloud Migration Requirements; Engagement Role Brief.",
        "ICT Environment Overview; On-Premises Network Diagram.",
        "Privacy Policy; User Access Policy; Security and Incident Response; Industry Standards Reference (AWS Well-Architected, ACSC Essential Eight).",
    ])

    h1("Document control")
    add_data_table(doc, ["Field", "Value"],
              [["Document version", "v1.0 — Approved for implementation"],
               ["Authored by", "MTS Senior Architecture Team, in consultation with YAT ICT"],
               ["Approved by", "Pat Lin (MTS Senior Consultant) · Sam Walker (YAT ICT Manager)"],
               ["Implemented by", "the AT2 Deployment Report (foundation build)"],
               ["Successor document", "YAT LMS Cloud Architecture — HA Design (follow-on phase)"]],
              widths=[5.0, 11.0])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website-s1/public/documents/YAT-LMS-Baseline-Solution-Design.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
