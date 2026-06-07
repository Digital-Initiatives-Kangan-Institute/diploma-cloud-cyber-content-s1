#!/usr/bin/env python3
"""Build the S1-CL2 AT2 STUDENT instrument (.docx) by populating the Kangan template.

Student-facing version of AT2, derived from the assessor instrument (build_at2_assessor) so
the shared content is single-sourced. Loads the official Kangan 'Project Assessment - Student'
template; includes Details, Student instructions, the assessment criteria (what they're marked
against), and the detailed task instructions. The assessor-only material — the marking
benchmark and its UoC traceability — is NOT included.

Usage:  python scripts/build_at2_student.py [output.docx]
Default: S1-CL2-Cloud-Disaster-Recovery/assessments/AT2/AT2-DR-Implementation-Student.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import build_at1_assessor as a1   # noqa: E402  shared cell/table helpers
import build_at2_assessor as a2   # noqa: E402  shared AT2 content (single source of truth)

from docx import Document  # noqa: E402

TEMPLATE = "templates/Project Assessment - Student.docx"


def build(path):
    doc = Document(TEMPLATE)

    # ---- Details (student name / ID / assessor name+email left blank) ----
    t = doc.tables[0]
    a1.set_cell(t.rows[2].cells[1], a2.DETAILS["qualification"])
    a1.set_cell(t.rows[3].cells[1], a2.DETAILS["units"])
    a1.set_cell(t.rows[4].cells[1], a2.DETAILS["task_title"])
    a1.set_cell(t.rows[5].cells[1], a2.DETAILS["task_number"])

    # ---- Student instructions ----
    ti = doc.tables[1]
    a1.set_cell(a1.instr_row(ti, "Assessment overview"), a2.OVERVIEW)
    a1.set_cell(a1.instr_row(ti, "Task"), a2.TASKS)
    a1.set_cell(a1.instr_row(ti, "Resources required"), a2.RESOURCES)
    a1.set_cell(a1.instr_row(ti, "Assessment criteria"), a2.CRITERIA_STATEMENT)

    # ---- Assessment criteria (no UoC / benchmark) ----
    tm = doc.tables[2]
    a1.clear_table_rows(tm, 2)
    a1.add_section_row(tm, "Deployment Report")
    for c in a2.CRITERIA:
        a1.add_criterion_row(tm, c)

    # ---- Detailed task instructions (no marking benchmark) ----
    doc.add_paragraph("Instructions to Student", style="Heading 1")
    for text, style in (a2.STUDENT_INTRO + a2.STUDENT_TASK + a2.TIPS):
        doc.add_paragraph(text, style=style)

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "S1-CL2-Cloud-Disaster-Recovery/assessments/AT2/AT2-DR-Implementation-Student.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
