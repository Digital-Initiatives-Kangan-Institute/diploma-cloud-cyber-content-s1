#!/usr/bin/env python3
"""Build the YAT / MTS Business Case template (.docx) from the brand pack.

Generates a professionally branded, system-agnostic Business Case document template:
cover sheet, table-of-contents field, all narrative sections, and the standard tables
(gap analysis, CBA summaries, risk register, action plan, sign-off). Branding follows
scenario/branding/brand-pack.md (§4 visual identity, §5.2 document templates,
§5.3 disclosure). Source of structure: the BC section outline used across S1-CL1.

This is the in-world template (no assessment/UoC scaffolding) — the canonical shape of a
YAT/MTS Business Case, suitable for handing students per-section snippets for exercises
and as the basis for the AT1 deliverable variant.

Usage:  python scripts/build_bc_template.py [output.docx]
Default output: scenario/templates/YAT-Business-Case-Template.docx
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

# --- Brand palette (brand-pack.md §4.1) ---------------------------------------
TEAL = "1F5A5C"        # Primary — headings, lockup, table header rows
TERRACOTTA = "C5613B"  # Accent — emphasis / links
OCHRE = "C99932"       # Highlight — disclosure banner background
CREAM = "F8F4ED"       # Page background (not applied to body; cover band only)
CHARCOAL = "1F2329"    # Body text
GREY = "6B6660"        # Muted / guidance text
STONE = "E4DED3"       # Borders / dividers
WHITE = "FFFFFF"
FONT = "Source Sans 3"  # brand-pack.md §4.2 (Word falls back to Calibri if absent)

DISCLOSURE = "This is a fictional organisation used as a case study in an educational context."
ADDRESS = "YAT College  ·  175 Cremorne St, Cremorne VIC 3121"


# --- low-level OOXML helpers --------------------------------------------------
def _shade(element_pr, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    element_pr.append(shd)


def shade_paragraph(paragraph, hex_fill):
    _shade(paragraph._p.get_or_add_pPr(), hex_fill)


def shade_cell(cell, hex_fill):
    _shade(cell._tc.get_or_add_tcPr(), hex_fill)


def set_cell_borders(cell, hex_color=STONE, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), hex_color)
        borders.append(e)
    tcPr.append(borders)


def paragraph_bottom_rule(paragraph, hex_color=TEAL, sz=12):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_field(paragraph, instr, placeholder=""):
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(it); run._r.append(sep)
    if placeholder:
        ph = paragraph.add_run(placeholder)
        ph.font.color.rgb = RGBColor.from_string(GREY)
    run2 = paragraph.add_run()
    run2._r.append(end)


# --- content helpers ----------------------------------------------------------
def guidance(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = RGBColor.from_string(GREY)
    r.font.size = Pt(9.5)
    p.space_after = Pt(6)
    return p


def response(doc, label="[ Write your response here ]"):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.italic = True
    r.font.color.rgb = RGBColor.from_string(GREY)
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], TEAL)
        set_cell_borders(hdr[i])
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_borders(cells[i])
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
            if val.startswith("[") or val.startswith("$"):
                run.italic = True
                run.font.color.rgb = RGBColor.from_string(GREY)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


# --- styles -------------------------------------------------------------------
def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(CHARCOAL)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, bold in (
        ("Title", 30, TEAL, True),
        ("Heading 1", 17, TEAL, True),
        ("Heading 2", 13, TEAL, True),
        ("Heading 3", 11.5, CHARCOAL, True),
    ):
        st = doc.styles[name]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = RGBColor.from_string(color)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def build_header_footer(section):
    # Header: ochre disclosure banner (brand-pack §5.3) on every page.
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(hp, OCHRE)
    r = hp.add_run(DISCLOSURE)
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(CHARCOAL)
    r.bold = True

    # Footer: YAT lockup (left) + page x of y (right), thin teal rule above.
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    paragraph_bottom_rule(fp, TEAL, sz=6)  # rule sits as a top divider via empty line below
    fp.text = ""
    # tab stops: left lockup, right page number
    tabs = fp.paragraph_format.tab_stops
    from docx.enum.text import WD_TAB_ALIGNMENT
    tabs.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
    lock = fp.add_run("YAT College")
    lock.bold = True
    lock.font.size = Pt(8.5)
    lock.font.color.rgb = RGBColor.from_string(TEAL)
    fp.add_run("\t")
    pg = fp.add_run("Page ")
    pg.font.size = Pt(8.5)
    pg.font.color.rgb = RGBColor.from_string(GREY)
    add_field(fp, "PAGE")
    of = fp.add_run(" of ")
    of.font.size = Pt(8.5)
    of.font.color.rgb = RGBColor.from_string(GREY)
    add_field(fp, "NUMPAGES")


def wordmark(paragraph):
    yat = paragraph.add_run("YAT")
    yat.bold = True
    yat.font.size = Pt(34)
    yat.font.name = FONT
    yat.font.color.rgb = RGBColor.from_string(TEAL)
    coll = paragraph.add_run("  COLLEGE")
    coll.bold = True
    coll.font.size = Pt(16)
    coll.font.color.rgb = RGBColor.from_string(TERRACOTTA)


# --- document body ------------------------------------------------------------
def build(path):
    doc = Document()
    configure_styles(doc)

    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(1.0)
    build_header_footer(sec)

    # ---- COVER ----
    lock_p = doc.add_paragraph()
    wordmark(lock_p)
    addr = doc.add_paragraph()
    ar = addr.add_run(ADDRESS)
    ar.font.size = Pt(9)
    ar.font.color.rgb = RGBColor.from_string(GREY)
    rule = doc.add_paragraph()
    paragraph_bottom_rule(rule, TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph(style="Title")
    title.add_run("Business Case")
    sub = doc.add_paragraph()
    sr = sub.add_run("[ Project / initiative name ]")
    sr.font.size = Pt(15)
    sr.italic = True
    sr.font.color.rgb = RGBColor.from_string(GREY)
    for _ in range(2):
        doc.add_paragraph()

    cover_rows = [
        ("Prepared for", "[ Client / board — e.g. YAT board, via the engagement sponsor ]"),
        ("Prepared by", "[ Consultant name, MP Tech Solutions (MTS) ]"),
        ("Engagement", "[ Engagement name ]"),
        ("Document version", "[ e.g. v0.1 ]"),
        ("Date", "[ Date ]"),
        ("Analysis period", "[ e.g. 5 years ]"),
        ("Currency", "AUD ex GST"),
        ("Classification", "Commercial-in-confidence"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in cover_rows:
        cells = ct.add_row().cells
        set_cell_borders(cells[0]); set_cell_borders(cells[1])
        shade_cell(cells[0], CREAM)
        kr = cells[0].paragraphs[0].add_run(k)
        kr.bold = True
        kr.font.size = Pt(10)
        vr = cells[1].paragraphs[0].add_run(v)
        vr.font.size = Pt(10)
        vr.italic = True
        vr.font.color.rgb = RGBColor.from_string(GREY)
        cells[0].width = Cm(4.5)
        cells[1].width = Cm(12.0)

    # ---- TABLE OF CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE)
    build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    toc_p = doc.add_paragraph()
    add_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u',
              placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE)
    build_header_footer(doc.sections[-1])

    def h1(text):
        doc.add_paragraph(text, style="Heading 1")

    def h2(text):
        doc.add_paragraph(text, style="Heading 2")

    h1("1. Executive Summary")
    guidance(doc, "Write this section last. One page the board reads first: the problem in a "
                  "sentence; the options considered; your recommendation; the high-level 5-year "
                  "cost position for each option; the top 2–3 risks of the recommended option and "
                  "how you'll manage them; and the decision you are asking the board to take today.")
    response(doc)

    h1("2. Engagement Context")
    guidance(doc, "Establish who you are, the engagement, who you report to, and what this "
                  "Business Case asks the board to decide. Approx. one short paragraph.")
    response(doc)

    h1("3. Strategic Alignment Analysis")
    guidance(doc, "Analyse the organisation's ICT Strategic Plan against the broader industry "
                  "environment and the organisation's objectives. Identify the objectives material "
                  "to this initiative (cite them); characterise the relevant industry direction; "
                  "state where the plan aligns with industry and where it diverges; draw out the "
                  "implications for this initiative. Approx. 250–400 words.")
    response(doc)

    h1("4. Current State")
    guidance(doc, "Synthesise the current-state material into a board-level summary — do not "
                  "reproduce source documents. Decide what is material to the decision, describe it "
                  "in your own words, and focus the reader on what drives the gap analysis. Cover "
                  "topology and security posture; the system's current status (what it runs, "
                  "availability, recovery objectives, condition); supporting infrastructure and "
                  "integrations; and staff capability and external dependencies. Approx. 150–250 words.")
    response(doc)

    h1("5. Gap Analysis")
    guidance(doc, "Compare each in-scope strategic objective against the current state. Identify the "
                  "gap, the improvement opportunity, and the proposed change. Add rows as needed.")
    table(doc,
          ["Strategic objective", "Current state", "Desired future state", "Gap",
           "Improvement opportunity", "Proposed change"],
          [["[ Objective ]", "[ … ]", "[ … ]", "[ … ]", "[ … ]", "[ … ]"]] * 3,
          widths=[3.0, 2.6, 2.6, 2.6, 2.6, 2.6])
    guidance(doc, "Narrative summary tying the table together (100–200 words).")
    response(doc)

    h1("6. Options Considered and Evaluation")
    h2("6.1 Workload definition")
    guidance(doc, "Define the workload the chosen option must support: user load, peak patterns, "
                  "data volumes, integration requirements, availability target.")
    response(doc)
    h2("6.2 Options considered")
    guidance(doc, "State the options you assessed (e.g. Option A — renew on-premises; "
                  "Option B — migrate to AWS). Note any other options and why they were / weren't assessed.")
    response(doc)
    h2("6.3 Evaluation method")
    guidance(doc, "State the evaluation method applied (CBA + intangibles + risk register, and any "
                  "weighted decision matrix). Justify why it suits this decision.")
    response(doc)
    h2("6.4 Initial impact and difficulty assessment")
    table(doc, ["", "Option A", "Option B"],
          [["Strategic impact", "[ … ]", "[ … ]"],
           ["Implementation difficulty", "[ … ]", "[ … ]"],
           ["Headline pros", "[ … ]", "[ … ]"],
           ["Headline cons", "[ … ]", "[ … ]"]],
          widths=[4.5, 6.0, 6.0])

    h1("7. Cost-Benefit Analysis")
    guidance(doc, "A full 5-year CBA comparing the options. Detailed line items live in Appendix 1; "
                  "the summary tables here are the headline view the board reads. All figures AUD ex GST.")
    h2("7.1 Assumptions")
    table(doc, ["Assumption", "Value", "Used as-is?"],
          [["[ Assumption ]", "[ Value ]", "[ ] yes  /  [ ] adjusted — see note"]] * 4,
          widths=[7.0, 5.0, 4.5])
    h2("7.2 Option A — 5-year summary")
    yr = ["", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5", "5-year"]
    table(doc, yr,
          [["One-off capital", "$_____", "—", "—", "—", "—", "$_____"],
           ["Recurring", "$_____", "$_____", "$_____", "$_____", "$_____", "$_____"],
           ["Annual total", "$_____", "$_____", "$_____", "$_____", "$_____", "$_____"]])
    h2("7.3 Option B — 5-year summary")
    table(doc, yr,
          [["One-off project", "$_____", "—", "—", "—", "—", "$_____"],
           ["Cloud / direct", "$_____", "$_____", "$_____", "$_____", "$_____", "$_____"],
           ["Staff + external", "$_____", "$_____", "$_____", "$_____", "$_____", "$_____"],
           ["Annual total", "$_____", "$_____", "$_____", "$_____", "$_____", "$_____"]])
    h2("7.4 Quantified benefit (e.g. avoided downtime)")
    guidance(doc, "Quantify the material benefit of the recommended option (for an availability "
                  "uplift, the avoided-downtime cost; otherwise the relevant quantified benefit).")
    response(doc)
    h2("7.5 Comparison summary")
    table(doc, ["", "Option A", "Option B", "Delta (B − A)"],
          [["5-year direct cost", "$_____", "$_____", "$_____"],
           ["Quantified benefit", "—", "$_____", "$_____"],
           ["Net 5-year position", "$_____", "$_____", "$_____"]],
          widths=[5.0, 3.8, 3.8, 3.8])
    h2("7.6 Sensitivity analysis")
    guidance(doc, "Pick at least two assumptions where a reasonable change would materially affect "
                  "the outcome. Show the recalculation and the effect on the recommendation.")
    response(doc)

    h1("8. Risk and Impact Assessment")
    h2("8.1 Intangibles comparison")
    guidance(doc, "One to two sentences per factor. Be honest about trade-offs.")
    table(doc, ["Factor", "Option A", "Option B"],
          [["[ Availability target ]", "[ … ]", "[ … ]"],
           ["[ Recovery objectives ]", "[ … ]", "[ … ]"],
           ["[ Capacity for growth / peaks ]", "[ … ]", "[ … ]"],
           ["[ Staff capability impact ]", "[ … ]", "[ … ]"],
           ["[ Vendor lock-in ]", "[ … ]", "[ … ]"],
           ["[ Security posture ]", "[ … ]", "[ … ]"]],
          widths=[5.0, 5.75, 5.75])
    h2("8.2 Risk register (recommended option)")
    table(doc, ["Risk", "Likelihood", "Impact", "Mitigation"],
          [["[ Risk ]", "[ H/M/L ]", "[ H/M/L ]", "[ … ]"]] * 3,
          widths=[6.5, 2.5, 2.5, 5.0])

    h1("9. Recommendation")
    guidance(doc, "State the recommended option and the headline rationale (3–5 sentences) — connect "
                  "it to the CBA findings, the intangibles, and the risk register. This drives the "
                  "prioritisation and action plan in §10.")
    response(doc)

    h1("10. Action Plan")
    h2("10.1 Prioritised changes")
    table(doc, ["#", "Change", "Priority rationale"],
          [["1", "[ … ]", "[ … ]"], ["2", "[ … ]", "[ … ]"], ["3", "[ … ]", "[ … ]"]],
          widths=[1.2, 7.0, 8.0])
    h2("10.2 Implementation schedule")
    guidance(doc, "Indicative schedule — phase the work; durations and dependencies. Do not propose "
                  "a single all-at-once cutover unless the risk register supports it.")
    table(doc, ["Phase", "Activities", "Start", "Duration", "Dependencies", "Owner"],
          [["1", "[ … ]", "[ … ]", "[ … ]", "—", "[ … ]"],
           ["2", "[ … ]", "[ … ]", "[ … ]", "[ … ]", "[ … ]"],
           ["3", "[ … ]", "[ … ]", "[ … ]", "[ … ]", "[ … ]"]],
          widths=[1.5, 5.0, 2.2, 2.2, 3.0, 2.4])
    h2("10.3 Standards, targets, and success metrics")
    table(doc, ["Aspect", "Standard / target", "How measured"],
          [["[ Availability target ]", "[ … ]", "[ … ]"],
           ["[ Recovery (RPO/RTO) ]", "[ … ]", "[ … ]"],
           ["[ Cost envelope ]", "[ … ]", "[ … ]"],
           ["[ Data residency ]", "[ … ]", "[ … ]"]],
          widths=[5.0, 5.0, 5.5])
    h2("10.4 Implementation methods")
    guidance(doc, "For each phase, identify the methods, tools, and standards applied (e.g. a "
                  "Well-Architected review at design milestones; the change-management procedure "
                  "for production changes).")
    response(doc)
    h2("10.5 Alignment with change-management procedure")
    guidance(doc, "State how the action plan aligns with the organisation's change-management "
                  "procedure: which phases require change requests, approval, and sign-off.")
    response(doc)

    h1("11. Next Steps and Decision Requested")
    guidance(doc, "State precisely what you are asking the board to decide today, and list any items "
                  "deferred to later sign-off gates.")
    response(doc)

    h1("Sign-off")
    table(doc, ["Role", "Name", "Date", "Signature"],
          [["Prepared by", "", "", ""],
           ["Reviewed by", "", "", ""],
           ["Approved by", "", "", ""]],
          widths=[5.0, 4.5, 3.0, 3.5])

    # ---- APPENDICES ----
    doc.add_section(WD_SECTION.NEW_PAGE)
    build_header_footer(doc.sections[-1])
    h1("Appendix 1 — Cost-Benefit Analysis: Detailed Line Items")
    guidance(doc, "The detailed line items underpinning the §7 summary tables — one-off and "
                  "recurring costs per option, and the year-by-year projection whose totals feed §7. "
                  "Build out the line-item tables here per option.")
    response(doc, "[ Detailed CBA line-item tables ]")

    h1("Appendix 2 — Supporting Research")
    guidance(doc, "Attach or link the supporting research: pricing-calculator exports for the "
                  "Appendix 1 figures, industry-environment sources cited in §3 (with access dates), "
                  "and any other supporting material.")
    response(doc, "[ Supporting research ]")

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website/public/templates/YAT-Business-Case-Template.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
