#!/usr/bin/env python3
"""Build the supplied Accounting System Baseline Solution Design (.docx) — in-world document.

The PRACTICE analogue of the YAT-LMS Baseline Solution Design: the supplied cloud architecture
students implement in the AT2 *practice* build (the Accounting System / Ledgerline engagement),
paralleling the real AT2's LMS baseline. Parallel-but-different: Microsoft SQL Server (not MySQL),
business-hours 99.5% (not 24/7 99.9%), internal staff-only over the campus VPN (not internet-facing),
commercial licensing + 7-year financial-records retention. Single-AZ baseline (non-HA); HA sections
marked "Not applicable" per the convention. In-world artefact (no UoC tags).

Output to the website documents folder for printing to PDF, then wiring into the AT2/AT3 intranet states.

Usage:  python scripts/build_accounting_solution_design.py [output.docx]
Default: ../diploma-cloud-cyber-website/public/documents/YAT-Accounting-Baseline-Solution-Design.docx
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
sys.path.insert(0, str(next(d for d in Path(__file__).resolve().parents if (d / "helpers" / "__init__.py").exists())))  # noqa: E402
from helpers.docx_body_text import add_bullet_list  # noqa: E402
import build_bc_template as bc   # noqa: E402
import build_s1_cl1_at1_bc_exemplar as ex   # noqa: E402  (etable, para, bullets)

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402


def na(doc, reason):
    p = doc.add_paragraph()
    r = p.add_run(f"Not applicable — {reason}")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    p.paragraph_format.space_after = Pt(6)
    return p


def diagram_placeholder(doc, caption, source):
    """A bordered, shaded drop-zone for a network diagram — replace with the exported image."""
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    bc.set_cell_borders(cell)
    bc.shade_cell(cell, bc.CREAM)
    cell.width = Cm(16.6)
    t.rows[0].height = Cm(9.5)
    t.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[ NETWORK TOPOLOGY DIAGRAM — PASTE HERE ]")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(source)
    r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = RGBColor.from_string(bc.GREY)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = RGBColor.from_string(bc.GREY)
    return t


def build(path):
    doc = Document()
    bc.configure_styles(doc)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)
    bc.build_header_footer(sec)

    # ---- COVER ----
    bc.wordmark(doc.add_paragraph())
    ar = doc.add_paragraph().add_run(bc.ADDRESS)
    ar.font.size = Pt(9); ar.font.color.rgb = RGBColor.from_string(bc.GREY)
    bc.paragraph_bottom_rule(doc.add_paragraph(), bc.TEAL, sz=12)
    for _ in range(3):
        doc.add_paragraph()
    doc.add_paragraph(style="Title").add_run("Solution Design")
    sub = doc.add_paragraph().add_run("YAT Accounting System Cloud Architecture — Baseline Design")
    sub.font.size = Pt(15); sub.bold = True; sub.font.color.rgb = RGBColor.from_string(bc.TERRACOTTA)
    doc.add_paragraph()
    cover = [
        ("Engagement", "YAT Accounting System Cloud Migration — Foundation Build"),
        ("Document type", "Technical design (Solution Design)"),
        ("Version", "v1.0 — Approved for implementation"),
        ("Authored by", "MTS Senior Architecture Team, in consultation with YAT ICT"),
        ("Approved by", "Pat Lin (MTS Senior Consultant) · Sam Walker (YAT ICT Manager)"),
        ("Implemented by", "the Accounting System foundation-build Deployment Report"),
        ("Classification", "Internal — YAT ICT, Finance, and MTS personnel on signed MSA"),
    ]
    ct = doc.add_table(rows=0, cols=2)
    for k, v in cover:
        cells = ct.add_row().cells
        bc.set_cell_borders(cells[0]); bc.set_cell_borders(cells[1]); bc.shade_cell(cells[0], bc.CREAM)
        kr = cells[0].paragraphs[0].add_run(k); kr.bold = True; kr.font.size = Pt(10)
        cells[1].paragraphs[0].add_run(v).font.size = Pt(10)
        cells[0].width = Cm(4.5); cells[1].width = Cm(12.0)

    # ---- CONTENTS ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    doc.add_paragraph("Contents", style="Heading 1")
    bc.add_field(doc.add_paragraph(), 'TOC \\o "1-3" \\h \\z \\u',
                 placeholder="Right-click and choose “Update Field” to build the table of contents.")

    # ---- BODY ----
    doc.add_section(WD_SECTION.NEW_PAGE); bc.build_header_footer(doc.sections[-1])
    h1 = lambda t: doc.add_paragraph(t, style="Heading 1")
    h3 = lambda t: doc.add_paragraph(t, style="Heading 3")

    h1("1. Purpose and Scope")
    ex.para(doc, "This document specifies the baseline AWS architecture MTS will implement as the foundation "
                 "build phase of the YAT Accounting System cloud migration. It translates the approved "
                 "direction into a concrete, implementable design for the Ledgerline finance and office-"
                 "administration system. The design stops at “infrastructure ready for application "
                 "deployment”: the EC2 instance is provisioned with the OS, the RDS instance with an empty "
                 "Microsoft SQL Server engine, and the load balancer with placeholder health checks — no "
                 "application binaries or financial data are placed by the MTS build.")
    h3("In scope of this design")
    add_bullet_list(doc, [
        "The production cloud foundation for the Ledgerline accounting application.",
        "All compute, networking, identity, storage, database, autoscaling and baseline monitoring needed to run Ledgerline as a staff-facing, business-hours workload in AWS.",
        "Single-region, single-Availability-Zone deployment in ap-southeast-2 (Sydney).",
    ])
    h3("Out of scope — deferred to the follow-on HA design phase")
    add_bullet_list(doc, [
        "High-availability hardening (Multi-AZ SQL Server, cross-AZ compute resilience, failure-simulation testing).",
        "Disaster recovery to a second AWS region; DR runbook and tabletop testing.",
        "Application re-platforming (Ledgerline remains Windows Server 2016 + Microsoft SQL Server).",
    ])
    h3("Out of MTS scope entirely — YAT ICT responsibility")
    add_bullet_list(doc, [
        "Ledgerline application installation onto the EC2 instance(s) after handover.",
        "Database migration (extract from on-prem SQL Server, load into RDS for SQL Server).",
        "Cutover — DNS switch, parallel running, decommissioning, user redirection (avoiding month-end and EOFY).",
        "Organisational change management — CAB approvals, communications, training, post-cutover support.",
    ])

    h1("2. Design Inputs and Requirements")
    h3("2.1 Inputs")
    add_bullet_list(doc, [
        "Accounting System Application Specification — functions, user load, data, integrations, SLAs, data residency.",
        "Accounting System Migration Requirements — platform preservation, availability, recovery, licensing, retention.",
        "Accounting System Server Specifications — current server, measured availability, utilisation, growth.",
        "Engagement Role Brief and Consultation Notes — engagement scope; OS, application and database preservation.",
    ])
    h3("2.2 Requirements the design must meet")
    ex.etable(doc, ["Requirement", "Target / note"],
              [["Region / data residency", "ap-southeast-2 (Sydney); financial records + staff/debtor PII within Australia (Privacy Act APP 8; financial-records retention)"],
               ["Application stack", "Preserved — Windows Server 2016 · Microsoft SQL Server (Standard) · Ledgerline"],
               ["Concurrent users", "15–25 typical; 45–55 peak at month-end close and EOFY; idle out of hours"],
               ["Availability", "≥ 99.5% business-hours (Mon–Fri ~07:30–18:00); no 24/7 requirement"],
               ["Recovery", "RPO ≤ 1 hour (no loss of financial transactions); RTO ≤ 1 business day (≤ 8 business hours)"],
               ["Data footprint", "~22 GB SQL Server data, growing ~5 GB/year, plus scanned-document attachments"],
               ["Retention", "Financial records and audit logs retained ≥ 7 years"],
               ["Licensing", "Commercial — Ledgerline per-user + SQL Server; cloud licence treatment (licence-included vs BYOL) is material"],
               ["Integrations", "AD authentication; O365 SMTP; LMS fee-status; payroll-bureau SFTP; banking/payment-gateway file exchange"],
               ["High availability", "Out of baseline scope — Multi-AZ resilience deferred to the follow-on HA design"]],
              widths=[4.5, 11.5])

    h1("3. Review of Existing Architecture")
    na(doc, "this is a greenfield cloud foundation; there is no existing cloud baseline to review. The "
            "on-premises current state is documented separately in the Accounting System Server "
            "Specifications and the ICT Environment Overview.")

    h1("4. Architecture Design")
    h3("4.1 Assumptions and constraints")
    ex.etable(doc, ["#", "Assumption / constraint", "Source"],
              [["A1", "Region must be ap-southeast-2 (Sydney) for financial-records data residency", "Migration Requirements; Privacy Policy"],
               ["A2", "Application stack preserved: Windows Server 2016, Microsoft SQL Server, Ledgerline", "Migration Requirements; Role Brief"],
               ["A3", "Data footprint ~22 GB SQL Server data, +~5 GB/year, plus scanned attachments", "Application Spec; Server Specs"],
               ["A4", "Concurrent users 15–25 typical, 45–55 at month-end / EOFY; idle out of hours", "Application Spec"],
               ["A5", "Financial records + staff/debtor PII must remain within Australia; retained ≥ 7 years", "Privacy Act 1988 APP 8; financial-records obligations"],
               ["A6", "Preserve AD authentication over a private network to campus AD; staff-only access", "Application Spec; Migration Requirements"],
               ["A7", "Baseline is NOT high-availability; Multi-AZ resilience deferred to the HA design", "Scoping decision"]],
              widths=[1.0, 9.5, 5.5])
    h3("4.2 AWS account and region")
    ex.para(doc, "An AWS account scoped to the migration engagement, provisioned by YAT and shared with MTS "
                 "for the build. Region ap-southeast-2 (Sydney); no deployment outside Australian regions, to "
                 "satisfy financial-records residency. Availability Zone ap-southeast-2a for the baseline; the "
                 "follow-on HA design introduces a second AZ.")
    h3("4.3 Identity and Access Management (IAM)")
    ex.etable(doc, ["Group", "Purpose", "Indicative permissions"],
              [["YAT-ICT-Admins", "YAT ICT day-to-day ops post-handover", "Read-only on infra; full CloudWatch/RDS/EC2 console; no IAM changes"],
               ["MTS-Consultants", "MTS during build + support", "Full admin during build; reduced post-handover"],
               ["Application-Service", "EC2 instance role for Ledgerline", "RDS read/write; S3 read/write (attachments); CloudWatch logs"],
               ["Finance-Auditors", "Finance / external auditors", "Read-only on logs, metrics, configs (financial-audit support)"]],
              widths=[3.5, 5.0, 7.5])
    add_bullet_list(doc, [
        "MFA required for all YAT-ICT-Admins and MTS-Consultants users (Essential Eight; User Access Policy).",
        "No long-lived access keys for humans; programmatic access via IAM roles only; EC2 access via instance profile.",
        "Configuration decision left to the implementer: the MTS-Consultants permission boundary during build vs after handover.",
    ])
    h3("4.4 Network topology")
    ex.para(doc, "VPC 10.0.0.0/16 (room to expand), with DNS hostnames/resolution and VPC flow logs enabled. "
                 "This is an internal, staff-only service — there is no public internet ingress to the "
                 "application. Single-AZ subnets in ap-southeast-2a:")
    ex.etable(doc, ["Subnet", "CIDR", "Tier", "Internet-facing?"],
              [["public-egress-a", "10.0.1.0/24", "NAT gateway only (outbound patching)", "Egress only"],
               ["private-app-a", "10.0.11.0/24", "Application / Ledgerline EC2 + internal ALB", "No"],
               ["private-data-a", "10.0.21.0/24", "Database (RDS for SQL Server)", "No"]],
              widths=[4.0, 3.5, 6.0, 2.5])
    add_bullet_list(doc, [
        "Internet Gateway used only for NAT egress (Windows Update, vendor patches); no inbound internet path to the app.",
        "Route tables: public-egress → IGW; private-app → NAT; private-data → no internet route.",
        "Staff reach the service over the campus Site-to-Site VPN; AD authentication runs over the same private link.",
        "The follow-on HA design adds the corresponding -b subnets in ap-southeast-2b.",
    ])
    diagram_placeholder(doc,
                        "Figure 4.4 — Ledgerline baseline network topology (single-AZ, internal / staff-only over the campus VPN).",
                        "Source diagram: network-accounting-baseline-singleaz.drawio")
    h3("4.5 Compute (EC2 + Auto Scaling)")
    add_bullet_list(doc, [
        "EC2: general-purpose x86 (e.g. m6i.large — final type a C1 implementer decision); Windows Server 2016 AMI; placed in private-app-a (no public IP).",
        "EBS: gp3 root 80 GB + a gp3 data volume sized by the implementer (footprint + 12-month growth + headroom).",
        "Auto Scaling Group: min 1 / desired 1 / max 2 (baseline); target-tracking on CPU at 70%; ELB+EC2 health checks; 300 s cooldown.",
        "The workload is business-hours and idle overnight; the follow-on HA design adds cross-AZ capacity for resilience (not for load).",
    ])
    h3("4.6 Load balancing (ALB)")
    add_bullet_list(doc, [
        "Internal (private) ALB in private-app-a; HTTPS:443 listener forwarding to the Ledgerline target group — reachable by staff over the campus VPN, not from the internet.",
        "Target group = the ASG instances; HTTP health check on the application health endpoint (30 s; 2 unhealthy → out of service).",
        "TLS via an ACM-issued certificate for the internal Ledgerline DNS name (DNS strategy a C8 implementer decision).",
    ])
    h3("4.7 Database (RDS for SQL Server)")
    add_bullet_list(doc, [
        "Amazon RDS for SQL Server, Standard edition (preserves the existing SQL Server engine and data); version confirmed against Ledgerline at build time.",
        "General-purpose instance class (e.g. db.m6i.large — a C2 implementer decision); gp3 storage sized to ~22 GB + ~5 GB/year growth.",
        "SQL Server licensing model — licence-included vs bring-your-own-licence — is a C3 implementer decision with a material cost impact.",
        "Multi-AZ DISABLED for the baseline (enabled in the HA design); storage encryption enabled (KMS).",
        "Placed in private-data-a; not publicly accessible. Automated backups + transaction-log backups sized to meet RPO ≤ 1 hour; backup window 22:00–04:00 AEST; maintenance Sun 02:00–06:00 AEST.",
        "Schema and data migration are YAT ICT's responsibility, not MTS's — MTS provisions an empty instance.",
    ])
    h3("4.8 Storage")
    ex.etable(doc, ["Resource", "Type", "Purpose"],
              [["EC2 root volume", "EBS gp3 80 GB", "OS and Ledgerline application install"],
               ["EC2 data volume", "EBS gp3 (sized by implementer)", "Ledgerline application data, report staging"],
               ["Documents bucket", "S3", "Scanned invoices / purchase orders / supporting documents; lifecycle to Glacier for 7-year financial retention"],
               ["Backups bucket", "S3", "Off-instance SQL backup exports, file snapshots"]],
              widths=[4.0, 4.5, 7.5])
    add_bullet_list(doc, ["Both buckets: block all public access; SSE-S3 (or SSE-KMS) encryption; versioning enabled; access logging to a log bucket; Object Lock considered for the 7-year financial-records hold."])
    h3("4.9 Security")
    ex.etable(doc, ["Security group", "Inbound", "Outbound"],
              [["sg-alb", "HTTPS:443 from the campus VPN / staff CIDR only", "HTTPS to sg-app"],
               ["sg-app", "from sg-alb; RDP:3389 from MTS bastion (design left to implementer)", "SQL:1433 to sg-db; HTTPS via NAT; LDAPS to campus AD; SMTP to O365; SFTP to payroll bureau; banking endpoints"],
               ["sg-db", "SQL Server:1433 from sg-app only", "none"]],
              widths=[3.0, 7.0, 6.0])
    add_bullet_list(doc, [
        "Encryption in transit: HTTPS staff→ALB→EC2; TLS EC2→RDS; AWS calls over TLS via VPC endpoints where available; LDAPS and SFTP for external links.",
        "Encryption at rest: EBS, RDS, and S3 all enabled (KMS where customer-managed keys are warranted for financial data).",
        "Operates under the AWS Shared Responsibility Model — AWS secures the cloud; YAT/MTS secure the OS, application, IAM, data and access in the cloud.",
    ])
    h3("4.10 Monitoring (baseline)")
    ex.para(doc, "Standard CloudWatch metrics for EC2, RDS, ALB and Auto Scaling. Baseline alarms (HA-tuned "
                 "alarms come in the follow-on HA design):")
    ex.etable(doc, ["Alarm", "Threshold"],
              [["EC2 CPU high", "≥ 80% over 10 min"],
               ["RDS CPU high", "≥ 80% over 10 min"],
               ["RDS free storage low", "< 15%"],
               ["ALB target 5XX / unhealthy host", "> 5 / min or any unhealthy host"],
               ["RDS connections high", "> 80% of max connections"]],
              widths=[8.0, 8.0])
    add_bullet_list(doc, ["Logging: VPC flow logs and RDS logs → CloudWatch Logs; ALB access logs → S3; EC2 OS logs via the CloudWatch Agent. Financial-audit-relevant logs retained to meet the 7-year obligation."])
    h3("4.11 Naming and tagging conventions")
    ex.para(doc, "Naming pattern yat-acct-<resource-type>-<env>-<az-or-purpose> (e.g. yat-acct-alb-prod). Mandatory tags:")
    ex.etable(doc, ["Tag", "Value"],
              [["Project", "YAT-Accounting-Migration"], ["Environment", "Production"], ["Owner", "YAT-ICT"],
               ["ManagedBy", "MTS-Migration during build → YAT-ICT post-handover"],
               ["CostCentre", "YAT-Accounting"], ["DataClassification", "Financial / Confidential (PII)"]],
              widths=[5.0, 11.0])
    h3("4.12 Backup")
    ex.etable(doc, ["Resource", "Mechanism", "Retention"],
              [["RDS database", "Automated daily backups + transaction-log backups (point-in-time recovery)", "Sized to RPO ≤ 1 h; long-term financial copies retained ≥ 7 years"],
               ["EC2 EBS volumes", "Daily AMI snapshot (Data Lifecycle Manager)", "14 days"],
               ["Documents (S3)", "Versioning + lifecycle to Glacier", "≥ 7 years (financial-records obligation)"]],
              widths=[4.0, 7.0, 5.0])
    add_bullet_list(doc, ["Cross-Region backup copies are out of scope for the baseline — addressed in the follow-on HA design."])
    h3("4.13 Recovery objectives — baseline state")
    ex.para(doc, "The baseline meets RPO ≤ 1 hour through RDS automated + transaction-log backups (point-in-"
                 "time recovery), and supports RTO ≤ 1 business day via a single-AZ restore. The single AZ and "
                 "the single RDS instance remain known single points of failure: tolerable for a business-"
                 "hours service in the short term, but resilience against an AZ failure is the objective of the "
                 "follow-on HA design.")
    h3("4.14 Components requiring vertical scaling")
    ex.para(doc, "RDS instance-class changes require a modify-and-apply (a brief interruption, taken in the "
                 "maintenance window, outside business hours); EBS volumes support online resize with no downtime.")
    h3("4.15 Single points of failure removed")
    na(doc, "this baseline is single-AZ by design; the single RDS instance and the single AZ are known "
            "single points of failure, deliberately deferred to the follow-on HA design.")
    h3("4.16 Configuration decisions left to the implementer")
    ex.para(doc, "The design is opinionated where it matters and silent where the implementer must show "
                 "judgement. Each decision below is to be made and evidenced in the Deployment Report, "
                 "justified against the Ledgerline workload.")
    ex.etable(doc, ["#", "Decision", "Why left open"],
              [["C1", "EC2 instance type (general-purpose family)", "Size against the Ledgerline workload (15–25 typical / 45–55 month-end peak)"],
               ["C2", "RDS for SQL Server instance class", "Size against the SQL Server workload (read-heavy at month-end)"],
               ["C3", "SQL Server licensing model — licence-included vs BYOL", "Material cost decision under the cloud operating model"],
               ["C4", "EBS data volume + RDS storage size", "Compute from the ~22 GB footprint + ~5 GB/year growth"],
               ["C5", "Backup retention + RPO mechanism", "Meet RPO ≤ 1 h and the 7-year financial-records retention"],
               ["C6", "Bastion / jump-host design for RDP admin access", "Left to the implementer"],
               ["C7", "Security-group ingress specifics", "Confirm AD/LDAPS, SMTP, payroll SFTP and banking endpoints"],
               ["C8", "DNS strategy + ACM certificate domain", "Confirm the internal Ledgerline hostname with YAT ICT"]],
              widths=[1.0, 7.0, 8.0])

    h1("5. Implementation Sequencing")
    na(doc, "this is a greenfield build in a new account, not a change to a running system; build order is "
            "at the implementer's discretion and there is no live service to sequence around or roll back.")

    h1("6. Simulation and Verification Plan")
    na(doc, "failure/resize simulation and availability verification are the subject of the follow-on HA "
            "design and its deployment report; the baseline build is verified functionally per the "
            "Deployment Report's testing section.")

    h1("7. Out of Scope")
    ex.para(doc, "Stated explicitly so the implementer knows what not to build (these are the deliberate "
                 "inputs to the follow-on HA design):")
    add_bullet_list(doc, [
        "Multi-AZ SQL Server; cross-AZ subnets (private-app-b, private-data-b); ASG capacity ≥ 2 across AZs.",
        "HA-tuned monitoring (cross-AZ latency, replica lag); cross-Region backup copies and DR runbook.",
        "Failure-simulation testing; automated availability reporting against the recovery objectives.",
        "And, out of MTS scope entirely: Ledgerline install, SQL Server data migration, cutover, and change management (YAT ICT).",
    ])

    h1("8. References")
    add_bullet_list(doc, [
        "Accounting System Application Specification; Migration Requirements; Server Specifications; Operational Costing.",
        "Engagement Role Brief; Consultation Notes; ICT Environment Overview.",
        "Privacy Policy; User Access Policy; Security and Incident Response; Industry Standards Reference (AWS Well-Architected, ACSC Essential Eight).",
    ])

    h1("Document control")
    ex.etable(doc, ["Field", "Value"],
              [["Document version", "v1.0 — Approved for implementation"],
               ["Authored by", "MTS Senior Architecture Team, in consultation with YAT ICT"],
               ["Approved by", "Pat Lin (MTS Senior Consultant) · Sam Walker (YAT ICT Manager)"],
               ["Implemented by", "the Accounting System foundation-build Deployment Report"],
               ["Successor document", "YAT Accounting System Cloud Architecture — HA Design (follow-on phase)"]],
              widths=[5.0, 11.0])

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"Wrote {path}")


if __name__ == "__main__":
    default = "../diploma-cloud-cyber-website/public/documents/YAT-Accounting-Baseline-Solution-Design.docx"
    out = sys.argv[1] if len(sys.argv) > 1 else default
    build(out)
